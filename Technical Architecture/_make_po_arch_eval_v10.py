# -*- coding: utf-8 -*-
"""Generate PO Architecture Evaluation Guide (programme-wide) as .docx."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\Prashanth\Official\Kaveri 3.0\Kaveri3Plan\Technical Architecture")
DST = BASE / "PO_Architecture_Evaluation_Guide_Kaveri3_v1.0.docx"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_runs(paragraph, text: str, base_size: Pt | None = None) -> None:
    """Support simple **bold** markers."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1
        if base_size is not None:
            run.font.size = base_size


def add_para(doc: Document, text: str = "", *, bold: bool = False, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_runs(p, text)
        if bold and p.runs:
            for r in p.runs:
                r.bold = True


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    if level:
        p.paragraph_format.left_indent = Pt(18 * (level + 1))
    add_runs(p, text)
    p.paragraph_format.space_after = Pt(2)


def add_table(doc: Document, rows: list[list[str]], col_widths: list[int] | None = None) -> None:
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    font_size = Pt(8.5) if width >= 4 else Pt(9.5)
    for r, row in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            add_runs(paragraph, row[c] if c < len(row) else "", font_size)
            if r == 0:
                shade_cell(cell, "D9E2F3")
                for run in paragraph.runs:
                    run.bold = True
            if col_widths and c < len(col_widths):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(col_widths[c]))
                tcW.set(qn("w:type"), "dxa")
                for child in list(tcPr):
                    if child.tag == qn("w:tcW"):
                        tcPr.remove(child)
                tcPr.append(tcW)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size in (
        ("Heading 1", 16),
        ("Heading 2", 13),
        ("Heading 3", 11.5),
    ):
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        style.font.name = "Calibri"
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(4)


def set_narrow_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)


def build() -> Path:
    BASE.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    set_narrow_margins(doc)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Product Owner Architecture Evaluation Guide")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("Kaveri 3.0 — Full Programme (All Phases)")
    sr.bold = True
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    add_para(
        doc,
        "This guide helps the Product Owner evaluate and challenge Technical Architecture "
        "produced by the Architecture team for the **entire Kaveri 3.0 programme**, not only "
        "the first go-live. It is informed by industry practice (Clean Architecture inside "
        "deployables; domain-aligned services / modular design between them) and by Kaveri "
        "programme constraints (phased delivery, shared platform, migration, e-Gov compliance).",
    )

    # Document control
    doc.add_heading("1. Document control", level=1)
    add_table(
        doc,
        [
            ["Field", "Value"],
            ["Document ID", "PO-K3-ARCH-EVAL-001"],
            ["Version", "1.0"],
            ["Status", "Draft for PO / Steering use"],
            ["Programme", "Kaveri 3.0 — Registration Department, Government of Karnataka"],
            ["Audience", "Product Owner, Steering, PM, Domain Expert (challenge); Architects (respond with evidence)"],
            ["Scope", "Full programme architecture across Phases 1–4 and shared platform"],
            ["Related inputs", "Programme Plan; Module BRDs; Marriage HLD; Clean Architecture / Microservices references"],
            ["Author", "Product Owner (Prashanth) — evaluation guide"],
            ["Last updated", "2026-08-26"],
        ],
    )

    doc.add_heading("1.1 Version history", level=2)
    add_table(
        doc,
        [
            ["Version", "Date", "Author", "Summary"],
            ["1.0", "2026-08-26", "PO", "Initial programme-wide PO architecture evaluation checklist, questions, red flags, artefacts and scorecard"],
        ],
    )

    # Purpose
    doc.add_heading("2. Purpose and PO role", level=1)
    add_para(
        doc,
        "Architecture is designed by a separate Architecture team. The Product Owner does **not** "
        "approve class diagrams or framework choices in detail. The PO judges whether the "
        "architecture **delivers the programme**, **protects statutory outcomes**, **reuses a "
        "shared platform**, **supports migration**, and **avoids irreversible mistakes**.",
    )

    doc.add_heading("2.1 PO vs Architect responsibilities", level=2)
    add_table(
        doc,
        [
            ["Product Owner decides / challenges", "Architects decide"],
            ["Priority, scope, go-live cuts across phases", "Technology stack and internal layering"],
            ["Acceptable risk and descope options", "API / event / database design details"],
            ["Business rules and non-bypassable hard gates", "CI/CD, container, broker implementation"],
            ["Business fallbacks for integrations", "Exact infra SKUs / cluster design"],
            ["Whether design supports all BRDs and phases", "Code structure and project layout"],
        ],
    )

    doc.add_heading("2.2 What “good” means for the full programme", level=2)
    add_para(doc, "Architecture is acceptable when the team can clearly show:")
    for item in [
        "All committed modules across Phases 1–4 (and explicit exclusions).",
        "One shared platform reused by Marriage, Stamp/Valuation, Document Registration, Firm, CC, Refund, Audit, MDM and related capabilities.",
        "Stable business boundaries that will not force a rewrite every phase.",
        "Statutory correctness for every Act / Rules set in scope.",
        "Migration from Kaveri 1.0 / 2.0 without stranding later modules.",
        "Operability and compliance at statewide scale for the full estate.",
        "An evolution path for new rules, fees, offices, integrations and modules after programme go-live.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "Phase dates matter for **delivery sequencing**. Architecture quality is judged on the "
        "**end-state programme**, with phased delivery of that same end state.",
    )

    # Checklist
    doc.add_heading("3. Programme evaluation checklist", level=1)

    doc.add_heading("3.1 End-state capability map", level=2)
    add_para(doc, "Check:")
    for item in [
        "Full catalogue of modules / sub-modules mapped to architecture building blocks.",
        "Each capability has exactly one system of record.",
        "No orphan BRD areas (Refund, EC, PoA, CVC/GIS, Firm, dashboards, scanning, etc.).",
        "Clear citizen vs officer vs admin / MIS surfaces for the whole estate.",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Ask:", bold=True)
    for q in [
        "Walk the full module list: which box owns each capability in the target architecture?",
        "Where do Phase 2–4 modules plug in without redesigning Phase 1 core?",
        "What is explicitly out of Kaveri 3.0 forever versus deferred?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.2 Platform vs domain (reuse across the programme)", level=2)
    add_para(doc, "Platform should be designed once for all phases. Verify ownership of:")
    add_table(
        doc,
        [
            ["Platform capability", "PO check"],
            ["Identity / User Management", "All officer modules consume same login, role, office, transfer, in-charge, DSC binding"],
            ["Citizen identity / eKYC", "Consistent across Marriage, Document, Firm (and other citizen journeys)"],
            ["Payment / Khajane / recon", "Same pattern for fees, stamps, registration payments, refunds"],
            ["Document / scan / store", "Reused by Marriage offline, Document Registration, CC, Firm"],
            ["eSign / DSC", "Citizen eSign and officer DSC as shared adapters"],
            ["Master data / MDM", "Offices, jurisdictions, fees, holidays, reason codes, templates"],
            ["Notification", "EN / KN templates programme-wide"],
            ["Audit / compliance", "Immutable trail across all modules"],
            ["Workflow", "Shared approach if multiple Acts need process engines"],
            ["Integration gateway", "External systems behind stable adapters"],
        ],
    )
    add_para(doc, "Ask:", bold=True)
    for q in [
        "What is built once for the programme, and what is rebuilt per module?",
        "If Document Registration starts, which platform contracts are already frozen?",
        "How does Refund use the same payment and audit model as Marriage and Stamp?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.3 Bounded contexts for the whole domain", level=2)
    add_para(doc, "Architecture should reflect department business, not delivery phases. Expected contexts include:")
    for item in [
        "User / office / post / privilege",
        "Marriage (and other personal-law Acts if in scope)",
        "Document registration and related (EC, CC, PoA, etc.)",
        "Stamp duty / valuation / e-stamp / guideline",
        "Firm / societies (as scoped)",
        "Payment, treasury, recon, refund",
        "Registers, certificates, certified copies",
        "MIS / dashboards per domain plus cross-cutting IGR views",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Ask:", bold=True)
    for q in [
        "Are boundaries drawn by business capability or by delivery phase?",
        "Will Stamp and Document Registration share data wrongly, or through published APIs / events?",
        "What prevents Marriage register logic leaking into Document Registration code?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.4 Cross-phase dependency and sequencing", level=2)
    add_para(doc, "Architecture must support phased go-lives without locking bad contracts.")
    add_para(doc, "Ask:", bold=True)
    for q in [
        "Which platform contracts are frozen early because later modules depend on them?",
        "What can change after an early go-live without breaking later phases?",
        "What is the dependency graph across UM, Payment, Document, Marriage, Stamp, Registration, Firm?",
        "If a later phase is the long critical path, what architecture work must finish before that coding peak?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.5 Statutory and process integrity (all Acts in programme)", level=2)
    add_para(doc, "For every major domain, hard gates must be enforceable in architecture — not only in UI:")
    for item in [
        "Registration / marriage workflow gates",
        "Payment / stamp / refund rules",
        "Role separation (SR, DEO, DR, IGR, etc.)",
        "Jurisdiction / office enforcement as hard filters",
        "Permanent registers versus operational data",
        "Form / template legal control (EN / KN)",
        "Maker-checker / dual control where BRDs require it",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Ask:", bold=True)
    for q in [
        "For each major Act in scope, where are non-bypassable rules enforced?",
        "How are statutory forms versioned and legally locked across modules?",
        "Can privilege or office changes break in-flight work in any module? What is the rule?",
        "Where is “cannot pay before approval” (or equivalent) enforced — UI only, or service / workflow?",
        "Can a restricted role’s token ever call approve / sign / register APIs?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.6 Data architecture for the full estate", level=2)
    add_para(doc, "Check:")
    for item in [
        "System-of-record map (office, fee, register serial, payment receipt, party identity, property identifiers, etc.).",
        "No shared mutable database across domains unless consciously accepted and controlled.",
        "Historical versus operational data; retention; AG / audit needs.",
        "Cross-module references via stable keys (application IDs, receipt IDs, office codes).",
        "MIS / reporting strategy (events, warehouse, replicas) for all dashboards.",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Ask:", bold=True)
    for q in [
        "Produce the programme data-ownership matrix.",
        "How will statewide MIS work without each module exposing ad-hoc database access?",
        "How are cross-module identifiers kept consistent after migration?",
        "Which data is permanent legal record versus purgeable operational data?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.7 Integration architecture (full external landscape)", level=2)
    add_para(
        doc,
        "Cover the whole programme, not a single module: Payment / Treasury / Khajane, eSign, DSC, "
        "Aadhaar / eKYC, SMS / Email, DigiLocker, Bhoomi or other land systems (if any), GIS / CVC, "
        "scanning vendors, State SSO, IGR / state reporting feeds.",
    )
    add_para(doc, "Ask:", bold=True)
    for q in [
        "Give the full integration register for Kaveri 3.0, with owner, pattern, fallback, and which phases use it.",
        "Which integrations are single points of failure for multiple modules?",
        "What is the standard pattern for sync versus async versus reconciliation across all externals?",
        "If a critical external is down for two hours on a go-live day, what still works and what is blocked?",
        "How do we prevent duplicate receipt, duplicate register serial, or duplicate certificate issue?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.8 Migration as a first-class programme concern", level=2)
    add_para(
        doc,
        "Architecture must support Kaveri 1.0 / 2.0 → 3.0 for all in-scope domains, not a one-time first-module cutover.",
    )
    add_para(doc, "Ask:", bold=True)
    for q in [
        "What is the migration architecture: coexistence, dual-run, cutover, rollback — per domain?",
        "How do legacy user / office / in-charge / history map into User Management for all later modules?",
        "How are historical registers and documents reachable after each phase go-live?",
        "What happens to in-flight legacy transactions during each cutover?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.9 Scale, resilience and DR for the finished system", level=2)
    add_para(doc, "Evaluate the statewide end state, then confirm phases can ramp to it.")
    add_para(doc, "Ask:", bold=True)
    for q in [
        "What concurrent users / transaction volumes is the end-state designed for?",
        "Which components scale horizontally when later modules create higher load than early ones?",
        "What are programme RPO / RTO, and which data is tier-1 (registers, payments) versus tier-2?",
        "Can one module’s outage be isolated so others keep running?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.10 Security, privacy and compliance (programme bar)", level=2)
    add_para(doc, "Check:")
    for item in [
        "Consistent AuthN / AuthZ model across all modules.",
        "Least privilege and jurisdiction as hard filters everywhere.",
        "Consistent PII / Aadhaar policy.",
        "Audit completeness across admin and registration actions.",
        "GIGW / WCAG / STQC / CERT-In path for the whole estate.",
        "Secrets, malware scanning, certificate binding as platform standards — not reinvented per module.",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Ask:", bold=True)
    for q in [
        "Is security a platform standard or reinvented in each module?",
        "How is access revoked programme-wide on transfer / relieving within the agreed window?",
        "Show how an officer of Office A is blocked from Office B’s work — at API level.",
        "What personal data is stored, for how long, and who can see it in logs?",
        "What is the evidence pack for STQC / security for the full system?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.11 Operability and support for the full estate", level=2)
    add_para(doc, "Ask:", bold=True)
    for q in [
        "How does operations support multiple phase go-lives without four different operating models?",
        "What is the common way to trace one citizen / officer journey across modules?",
        "What dashboards and alerts exist for stuck workflows, failed payments, DSC / eSign failures and SMS failures?",
        "Who is on-call and what is the rollback plan for each major go-live?",
        "How will ServiceDesk diagnose login, mapping, in-charge and digital-sign issues across modules?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.12 Evolution and total cost of ownership", level=2)
    add_para(doc, "Ask:", bold=True)
    for q in [
        "What changes are configuration / master-data versus code versus a new service?",
        "How do we add another personal-law Act or a new registration subtype?",
        "If the eSign or payment provider changes, which components change and what stays untouched?",
        "What is the expected cost of change after all modules are live?",
        "Where are we accepting intentional technical debt, and what is the payback plan across phases?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("3.13 Organisation and ownership fit", level=2)
    add_para(
        doc,
        "Architecture fails if ownership does not match boundaries. Prefer fewer deployables early; "
        "fine-grained split only with clear team and release reasons.",
    )
    add_para(doc, "Ask:", bold=True)
    for q in [
        "Who owns each platform service and each domain module for the life of the programme?",
        "Are we creating more independently deployable units than teams can run?",
        "How do BA / domain experts map to bounded contexts?",
        "What problem does a fine-grained split solve that a coarser modular design would not?",
    ]:
        add_bullet(doc, q)

    # Core questions
    doc.add_heading("4. Core questions for every programme architecture review", level=1)

    doc.add_heading("4.1 End state", level=2)
    for q in [
        "What is the target architecture when all modules are live?",
        "What will look the same in the years after go-live, and what is expected to evolve?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("4.2 Coverage", level=2)
    for q in [
        "Does every item in the modules catalogue have an architectural home?",
        "Where are the gaps and overlaps between modules?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("4.3 Platform", level=2)
    for q in [
        "What is the shared platform, and what is the contract freeze schedule by phase?",
        "What must never be duplicated module-by-module?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("4.4 Delivery without trapping the future", level=2)
    for q in [
        "How do early go-lives avoid freezing the wrong data model or API?",
        "What are the irreversible decisions, and why are they safe for later phases?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("4.5 Risk across the programme", level=2)
    for q in [
        "What are the top programme-level architecture risks (not only near-term)?",
        "Which failure takes down multiple modules at once?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("4.6 Acceptance", level=2)
    for q in [
        "What demos prove the architecture works for cross-module scenarios (for example, a User Management change affecting Marriage and later Document queues)?",
        "How will architecture readiness be gated before each phase go-live, not only the first?",
    ]:
        add_bullet(doc, q)

    doc.add_heading("4.7 Money and compliance", level=2)
    for q in [
        "How are payment, refund, reconciliation and audit consistent programme-wide?",
        "How do permanent legal records stay intact across modules and migrations?",
    ]:
        add_bullet(doc, q)

    # Red flags
    doc.add_heading("5. Red flags — push back hard", level=1)
    for item in [
        "Architecture story that only explains the first module / first phase.",
        "“We will redesign the platform when a later phase starts.”",
        "Separate login / role models per module.",
        "Separate payment / audit / document approaches per phase.",
        "Shared database growing into a programme-wide accidental monolith.",
        "No end-state integration or data-ownership map.",
        "Migration plan only for the first go-live.",
        "Module boundaries equal project phases, not business domains.",
        "More services than sustainable owners for the full estate.",
        "No answer on how Document Registration / Stamp / Firm reuse early platform.",
        "MIS that requires direct database access into every module.",
        "Security / compliance treated as a final-phase activity.",
        "Business rules enforced only in UI or reports.",
        "“We will handle audit / DR / reconciliation later.”",
        "Every external integration assumed always available.",
        "Architecture documents that do not reference BRD IDs / FRs / NFRs.",
    ]:
        add_bullet(doc, item)

    # Artefacts
    doc.add_heading("6. Artefacts the Architecture team must provide", level=1)
    add_table(
        doc,
        [
            ["Artefact", "Why the PO needs it"],
            ["End-state system context (all users + externals)", "Full scope clarity"],
            ["Target container / module view with phase tags (P1–P4)", "Sequencing without losing end state"],
            ["Capability → module ownership matrix (full catalogue)", "No orphans or overlaps"],
            ["Platform contract catalogue + freeze dates", "Reuse control across phases"],
            ["Programme data-ownership and retention matrix", "Legal / MIS safety"],
            ["Full integration register", "External risk management"],
            ["Cross-phase dependency and irreversible-decision log", "Protect later phases"],
            ["Migration architecture per domain / phase", "Cutover confidence"],
            ["NFR targets for end state + per-phase ramp", "Scale / DR honesty"],
            ["Hard-gate register by domain / Act", "Statutory enforcement"],
            ["Ops model for multi-module production", "Supportability"],
            ["Architecture Decision Records (ADRs)", "Why choices were made"],
        ],
    )

    # How to run review
    doc.add_heading("7. How to run an architecture review (PO script)", level=1)
    for i, step in enumerate(
        [
            "**Start with end state:** “Show Kaveri 3.0 when everything is live.”",
            "**Overlay phases:** “Show what exists at each go-live without contradicting the end state.”",
            "**Prove reuse:** “Pick Payment, User Management, Document, Audit — show consumption by every domain.”",
            "**Prove isolation:** “Show that one domain outage does not corrupt another domain’s legal registers.”",
            "**Prove change:** “Change office / post / fee / eSign provider — impact across the programme.”",
            "**Prove migration:** “Legacy coexistence for early and late domains.”",
            "**Close with decisions:** Irreversible choices, freeze dates, owners, and open questions across all phases.",
        ],
        start=1,
    ):
        add_para(doc, f"{i}. {step}")

    # Scorecard
    doc.add_heading("8. Pass / Conditional / Reject scorecard", level=1)
    add_table(
        doc,
        [
            ["Result", "When to use"],
            [
                "Pass",
                "End-state coverage complete; platform reuse clear; ownership / data / integration maps exist; "
                "phased delivery does not poison later phases; NFRs / DR / security are programme-wide.",
            ],
            [
                "Conditional",
                "End state is clear, but some later-phase details remain open — with dated decisions and no irreversible conflict.",
            ],
            [
                "Reject",
                "Only near-term design; platform duplication likely; no programme data / integration / migration story; "
                "statutory gates not enforceable; ownership unclear for major domains.",
            ],
        ],
    )

    doc.add_heading("8.1 Review record (copy per review meeting)", level=2)
    add_table(
        doc,
        [
            ["Item", "Entry"],
            ["Review date", ""],
            ["Architecture package / version reviewed", ""],
            ["Attendees", ""],
            ["Result (Pass / Conditional / Reject)", ""],
            ["Conditions / open decisions (with owners and due dates)", ""],
            ["Irreversible decisions accepted", ""],
            ["PO sign-off", ""],
            ["Architecture lead acknowledgement", ""],
        ],
    )

    # Industry anchors
    doc.add_heading("9. Industry anchors (for challenge, not for PO design work)", level=1)
    add_para(
        doc,
        "Use these only to challenge whether the Architecture team is applying industry-standard intent:",
    )
    add_table(
        doc,
        [
            ["Reference", "What the PO should listen for"],
            [
                "Clean Architecture (inside each deployable)",
                "Business rules independent of UI, database and external vendors; use cases / domain clear; adapters for Khajane, eSign, DSC.",
            ],
            [
                "Building Microservices — Sam Newman",
                "Independently releasable units modelled on business domains; information hiding; no casual shared mutable DB; "
                "sync for user actions and events / sagas for cross-module processes; BFFs for channels; "
                "do not default to excessive fine-grained services without team ownership.",
            ],
            [
                "Kaveri programme fit",
                "Shared platform once; domain modules for Marriage, Stamp, Document Registration, Firm; "
                "phased go-lives of one target architecture; migration and e-Gov compliance designed in.",
            ],
        ],
    )

    # Bottom line
    doc.add_heading("10. Bottom line for the Product Owner", level=1)
    add_para(
        doc,
        "Approve architecture when it is a **programme blueprint**: one department platform, multiple "
        "domain modules, phased go-lives of the same target design, safe migration, and long-term "
        "change without rewrite. Do not approve a design that only explains the first go-live.",
    )
    add_para(
        doc,
        "You approve when the design is **traceable to BRDs**, **deliverable across all phases**, "
        "**safe for registration / payment / identity**, and **operable as a statewide estate**.",
    )

    doc.save(DST)
    return DST


if __name__ == "__main__":
    out = build()
    print(f"Wrote {out}")
