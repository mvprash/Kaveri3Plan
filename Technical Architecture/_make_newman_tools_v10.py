# -*- coding: utf-8 -*-
"""Generate Newman microservices tooling-by-layer reference DOCX."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\Prashanth\Official\Kaveri 3.0\Kaveri3Plan\Technical Architecture")
DST = BASE / "Newman_Microservices_Tools_by_Layer_v1.0.docx"

# Source: Building Microservices, 2nd Edition, Sam Newman (O'Reilly, 2021)


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_runs(paragraph, text: str, base_size: Pt | None = None) -> None:
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1
        if base_size is not None:
            run.font.size = base_size


def add_para(doc: Document, text: str = "", *, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_runs(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    add_runs(p, text)
    p.paragraph_format.space_after = Pt(2)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    font_size = Pt(8) if width >= 4 else Pt(9)
    for r, row in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            add_runs(paragraph, row[c] if c < len(row) else "", font_size)
            if r == 0:
                shade_cell(cell, "D9E2F3")
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for name, size in (("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)):
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        style.font.name = "Calibri"
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(4)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)


def build() -> Path:
    BASE.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    set_margins(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Microservices Tooling by Architecture Layer")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("From Sam Newman — Building Microservices (2nd Edition)")
    sr.bold = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    add_para(
        doc,
        "This note summarises Newman’s **technology options and guidance by layer**: "
        "what the tool/pattern is for, how it is used, and **which problem it solves**. "
        "Newman repeatedly warns that tools are not the architecture — pick them when the "
        "problem appears, not because they are fashionable.",
    )

    add_table(
        doc,
        [
            ["Field", "Value"],
            ["Document ID", "REF-K3-NEWMAN-TOOLS-001"],
            ["Version", "1.0"],
            ["Source", "Building Microservices: Designing Fine-Grained Systems, 2nd ed., Sam Newman (O’Reilly, 2021)"],
            ["Purpose", "Architecture discussion aid for Kaveri 3.0 — tooling options by layer"],
            ["Last updated", "2026-08-26"],
        ],
    )

    doc.add_heading("0. Newman’s overarching caveats (read first)", level=1)
    for item in [
        "**Technology-agnostic services:** a microservice hides its implementation; consumers only see networked interfaces.",
        "**Avoid shared mutable databases** across services — that breaks information hiding and independent release.",
        "**Don’t rush to Kubernetes/containers** if you have only a few services; adopt when deployment overhead becomes painful, and prefer a **managed** cluster.",
        "**Don’t pick Kafka for request–response**; match the tool to the interaction style.",
        "**Schemas + contract tests** catch breakages early; without schemas, testing must do more work.",
        "Security tooling without a **threat model** is cargo-cult — start from threats, not from JWT/mTLS slides.",
    ]:
        add_bullet(doc, item)

    # ---------- Layer tables ----------
    doc.add_heading("1. Edge / user interface aggregation layer", level=1)
    add_para(
        doc,
        "Problem: UIs (especially mobile) would otherwise make many chatty calls to many microservices, "
        "pull too much data, and couple the frontend to backend topology.",
    )
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "Backend for Frontend (BFF)",
                "Dedicated server-side API shaped for one UI channel (web, mobile, officer workbench). Aggregates downstream calls.",
                "Stops browser/mobile from becoming an orchestration engine; allows channel-specific APIs without polluting core services.",
            ],
            [
                "GraphQL (perimeter)",
                "Single query endpoint at the edge; client asks for exactly the fields needed; server aggregates over microservices.",
                "Reduces round-trips and payload size for constrained clients; avoids bespoke aggregation endpoints for every screen.",
            ],
            [
                "REST for writes (hybrid)",
                "Often use GraphQL for reads and REST for writes when GraphQL write model fits poorly.",
                "Keeps mutations clear and cacheable/HTTP-verb semantics where GraphQL is awkward.",
            ],
            [
                "CDN / HTTP caching proxies (e.g. Varnish) with REST",
                "Cache GET responses using HTTP cache headers at the edge.",
                "Scales read-heavy public APIs; GraphQL is weaker here — caching is harder.",
            ],
            [
                "Apollo (example implementation)",
                "Cited as an implementation-specific GraphQL caching approach.",
                "Helps with GraphQL client/server caching gaps (still not as natural as REST+CDN).",
            ],
        ],
    )
    add_para(
        doc,
        "**Newman guidance:** GraphQL’s sweet spot is the **system perimeter** (GUIs / external APIs), "
        "not general microservice-to-microservice traffic. Don’t treat GraphQL as “SQL over microservices” "
        "or couple the schema to underlying databases.",
    )

    doc.add_heading("2. Synchronous service communication layer", level=1)
    add_para(
        doc,
        "Problem: services must call each other (or expose APIs to many client types) with clear contracts, "
        "without lockstep releases.",
    )
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "REST over HTTP (sensible default)",
                "Resource-oriented APIs; JSON (common) or other representations; wide client interoperability.",
                "Best default for **broad client variety**, external APIs, and caching-friendly interactions.",
            ],
            [
                "gRPC (+ Protocol Buffers)",
                "Strongly typed RPC with HTTP/2; codegen for clients/servers; Protolock for schema compatibility.",
                "High performance and strong contracts when **you control both ends**; great for internal sync calls.",
            ],
            [
                "OpenAPI / Swagger",
                "Document REST APIs; optionally generate clients; openapi-diff for breaking-change checks in CI.",
                "Makes REST contracts explicit; catches accidental API breaks before deploy.",
            ],
            [
                "Avoid SOAP / Java RMI (as primary choice)",
                "SOAP seen as heavyweight; RMI brittle / tech-locked.",
                "Avoids coupling and developer friction of older RPC stacks.",
            ],
            [
                "HTTP load balancers / proxies (e.g. mod_proxy, NGINX, HAProxy)",
                "Sit in front of service instances for routing and scale-out of HTTP APIs.",
                "Distributes load; enables rolling changes without clients knowing instance addresses.",
            ],
        ],
    )
    add_para(
        doc,
        "**Newman guidance:** Prefer **gRPC** when both client and server are under your control; "
        "prefer **REST over HTTP** when many unknown/external clients must integrate. "
        "Never hide the network completely in client libraries.",
    )

    doc.add_heading("3. Asynchronous messaging & streaming layer", level=1)
    add_para(
        doc,
        "Problem: not every interaction should be a blocking call; you need fan-out, temporal decoupling, "
        "and often near-real-time data movement without shared DBs.",
    )
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "Message brokers (RabbitMQ, ActiveMQ, cloud SQS/SNS/Kinesis, etc.)",
                "Producer gives message to broker with routing; consumers read queues/topics (often consumer groups).",
                "Temporal decoupling (receiver can be down); load buffering; pub/sub fan-out of events/commands.",
            ],
            [
                "Queues vs Topics",
                "Queue ≈ point-to-point work; Topic ≈ many subscribers each get a copy.",
                "Models competing consumers (scale workers) vs broadcast events (notify many services).",
            ],
            [
                "Apache Kafka (highlighted)",
                "High-throughput log; message permanence/retention; partitions; optional KSQL; often with Flink for processing.",
                "Moves large volumes of data; re-read history; shift from batch reports to near-real-time pipelines.",
            ],
            [
                "Debezium (+ Kafka)",
                "Change Data Capture from existing databases into Kafka topics.",
                "Brings legacy/shared DBs into a stream architecture without big-bang rewrite.",
            ],
            [
                "Apache Flink / stream processors",
                "Process Kafka (or similar) streams for transforms, joins, windows.",
                "Real-time analytics and derived views without each service polling databases.",
            ],
            [
                "AsyncAPI / CloudEvents",
                "Explicit schemas for evented endpoints (CloudEvents gaining CNCF / vendor support).",
                "Makes event contracts as clear as REST/RPC schemas; improves interoperability.",
            ],
        ],
    )
    add_para(
        doc,
        "**Newman guidance:** Kafka is a strong default for **streaming / high volume / permanence**, "
        "not a universal bus for every request. Prefer **fat events** (enough data in the message) over "
        "chatty follow-up GETs — with size limits in mind (Kafka default max message size historically ~1 MB).",
    )

    doc.add_heading("4. Serialization & contract layer", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "JSON (common for REST)",
                "Text payloads; easy for browsers and many clients.",
                "Interoperability and debuggability for human/tooling inspection.",
            ],
            [
                "Avro",
                "Schema-based format popular for message payloads; schema can travel with payload.",
                "Easier multi-version messaging formats; strong fit with Kafka ecosystems.",
            ],
            [
                "Protocol Buffers",
                "Binary format used by gRPC; also usable alone.",
                "Compact, efficient serialization when payload size/latency matters.",
            ],
            [
                "Schema compatibility tooling (Protolock, json-schema-diff-validator, openapi-diff, Confluent Schema Registry)",
                "Compare schemas in CI; fail build on incompatible changes.",
                "Catches **structural** breaking changes early so independent deployability survives.",
            ],
            [
                "Semantic Versioning (MAJOR.MINOR.PATCH)",
                "Communicate compatibility expectations in version numbers.",
                "Gives consumers a quick signal of breaking vs additive change (less proven in distributed systems, but useful).",
            ],
        ],
    )

    doc.add_heading("5. Service discovery, configuration & secrets layer", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "DNS (often enough)",
                "Resolve logical service names to instances / VIPs.",
                "Basic discovery without heavyweight platforms — start simple.",
            ],
            [
                "Consul (and similar)",
                "Service registry / discovery; health; often paired with config.",
                "Dynamic instance membership in changing microservice fleets.",
            ],
            [
                "etcd / ZooKeeper",
                "Distributed coordination / config stores (also under platforms like Kubernetes).",
                "Shared coordination state for clustered systems.",
            ],
            [
                "HashiCorp Vault (and cloud Key Vault equivalents)",
                "Central secrets store; dynamic DB credentials; rotation; consul-template style injection.",
                "Stops secrets in images/git; enables rotation and time-limited credentials after compromise risk.",
            ],
            [
                "Kubernetes Secrets (with caveats)",
                "Platform-native secret delivery into pods; still combine with real secret management practices.",
                "Operational delivery of credentials into running services.",
            ],
        ],
    )

    doc.add_heading("6. Edge gateway & service-mesh layer", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "API Gateway (e.g. Ambassador and similar edge proxies)",
                "Single entry for external traffic; routing, auth hand-off, rate limits, TLS termination.",
                "Hides internal topology; central place for cross-cutting edge concerns.",
            ],
            [
                "Envoy proxy",
                "High-performance proxy often used as data plane for gateways/meshes.",
                "Consistent L7 routing, retries, observability hooks at the network edge of services.",
            ],
            [
                "Service mesh (Istio, Linkerd — discussed cautiously)",
                "Sidecar proxies for mTLS, traffic policy, telemetry between services.",
                "Cross-cutting networking/security/observability without putting all logic in app code — **at operational cost**.",
            ],
            [
                "Sidecar pattern",
                "Companion process next to service instance (mesh proxy, log shipper, etc.).",
                "Adds platform capabilities without rewriting every microservice.",
            ],
        ],
    )
    add_para(
        doc,
        "**Newman guidance:** Gateways and meshes solve real cross-cutting problems, but they add "
        "complexity. Don’t assume a mesh is required on day one; justify with threat model and ops capacity.",
    )

    doc.add_heading("7. Workflow / multi-service process layer", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "Sagas (orchestration or choreography)",
                "Model long business processes as a sequence of local transactions + compensating actions — not 2PC distributed transactions.",
                "Coordinates multi-service business flows (orders, payments, registrations) without locking the whole system in one ACID transaction.",
            ],
            [
                "Process / workflow engines (e.g. Camunda mentioned in ecosystem)",
                "Explicit process definitions, timers, human tasks, compensation hooks.",
                "Makes complex stateful workflows visible, versionable, and operable.",
            ],
            [
                "Correlation IDs across saga steps",
                "One ID flows with every related call/event.",
                "Lets ops reconstruct an entire business transaction across services.",
            ],
        ],
    )

    doc.add_heading("8. Build & CI layer", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "CI servers (Jenkins, CircleCI, GitHub Actions, etc.)",
                "Build each microservice independently; run tests; publish artifacts.",
                "Enables independent releasability — the defining microservice property.",
            ],
            [
                "Per-service pipelines mapped to ownership",
                "One microservice → clear repo/build ownership boundaries.",
                "Avoids monolith build coupling that blocks independent teams.",
            ],
            [
                "Artifact repositories + Pact Broker",
                "Store build outputs and consumer-driven contract files.",
                "Shares verified contracts between consumer and provider pipelines.",
            ],
        ],
    )

    doc.add_heading("9. Deployment & runtime isolation layer", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "Containers (Docker)",
                "Package service + deps into a lightweight isolated unit.",
                "Faster spin-up and better isolation than full VMs for many small services; reproducible deploys.",
            ],
            [
                "Kubernetes (orchestration)",
                "Schedule containers across machines; desired state; scaling; self-healing.",
                "Manages many instances robustly — **when** you have enough services to justify the ops cost.",
            ],
            [
                "Managed Kubernetes / public cloud",
                "Let a provider run the control plane.",
                "Avoids “running your own Kubernetes cluster can be a significant amount of work.”",
            ],
            [
                "Helm / platform packaging",
                "Package k8s manifests for repeatable installs (ecosystem mention).",
                "Standardises how services are deployed onto the cluster.",
            ],
            [
                "FaaS / Serverless (AWS Lambda and equivalents)",
                "Deploy functions without managing machines; also serverless DBs/brokers/storage.",
                "Offloads ops for bursty or simple workloads; higher abstraction than VMs/containers.",
            ],
            [
                "Nomad / OpenShift / Knative (options in landscape)",
                "Alternative orchestrators or serverless-on-k8s layers.",
                "Different trade-offs for scheduling and scale-to-zero without forcing one vendor story.",
            ],
            [
                "Infrastructure as Code (e.g. Terraform)",
                "Declare cloud/network/data infra in code.",
                "Repeatable environments; rebuild after attack or drift.",
            ],
            [
                "Separate deployment from release (feature flags / progressive delivery; Spinnaker cited in landscape)",
                "Install new bits without exposing behaviour to all users immediately.",
                "Safer go-lives; quick disable without full rollback of every service.",
            ],
        ],
    )

    doc.add_heading("10. Testing layer", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "Test pyramid (unit → service → fewer E2E)",
                "Prefer fast, focused tests; limit brittle end-to-end suites.",
                "Keeps feedback fast as service count grows; E2E alone doesn’t scale.",
            ],
            [
                "Consumer-Driven Contracts — Pact",
                "Consumer defines expected interactions; provider verifies against Pact file; Pact Broker shares specs.",
                "Prevents breaking API changes without slow full-system E2E; clarifies producer/consumer ownership.",
            ],
            [
                "Schema diff in CI (see §4)",
                "Fail builds on incompatible OpenAPI/Protobuf/JSON Schema changes.",
                "Structural contract safety complementary to Pact’s behavioural contracts.",
            ],
            [
                "In-production testing / progressive exposure",
                "Test safely in prod with careful blast-radius controls (after separating deploy vs release).",
                "Finds issues staging cannot reproduce in complex distributed systems.",
            ],
        ],
    )

    doc.add_heading("11. Observability layer (monitoring → observability)", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "Log aggregation (Fluentd → Elasticsearch/Kibana; or cloud logging; Splunk/Datadog landscape)",
                "Ship logs from all instances to one searchable place.",
                "You can no longer SSH to one box — failures span many services.",
            ],
            [
                "Correlation IDs (implement early)",
                "Generate at gateway; pass on every call; put in every log line.",
                "Reconstruct one user/business journey across microservices; hard to retrofit later.",
            ],
            [
                "Distributed tracing (Jaeger; OpenTelemetry; commercial Honeycomb / Lightstep)",
                "Trace spans across services; find latency bottlenecks and causality better than logs alone.",
                "Answers “where is time spent?” and true call order despite clock skew.",
            ],
            [
                "Metrics (Prometheus cited in landscape; Grafana typically paired in industry)",
                "Time-series signals for RED/USE-style monitoring and alerting.",
                "Detect saturation, errors, and SLO burn across many instances.",
            ],
            [
                "NTP (with humility)",
                "Reduce clock skew across hosts.",
                "Improves log chronology somewhat — but tracing is still needed for accuracy.",
            ],
        ],
    )
    add_para(
        doc,
        "**Newman guidance:** Start with log aggregation + correlation IDs; add dedicated tracing when "
        "complexity warrants it. Prefer tools built for **exploration** of unknown questions (observability), "
        "not only static dashboards (classic monitoring).",
    )

    doc.add_heading("12. Security layer", level=1)
    add_table(
        doc,
        [
            ["Tool / pattern", "Usage (Newman)", "Problem solved"],
            [
                "Threat modelling first",
                "Classify data sensitivity; decide implicit trust vs zero trust by risk.",
                "Stops buying security tools without knowing what you’re defending against.",
            ],
            [
                "TLS in transit; harden data at rest",
                "Encrypt sensitive traffic and stored secrets/data per threat model.",
                "Confidentiality if network or disk is compromised.",
            ],
            [
                "Gateway-issued JWT for request context",
                "Edge authenticates user; gateway creates JWT for downstream call chain.",
                "Propagates identity/claims to microservices without each re-authenticating the human.",
            ],
            [
                "OAuth / OIDC / SAML (identity federation landscape)",
                "Standard protocols for user/app authentication and federation.",
                "Reusable identity with external IdPs instead of custom login per service.",
            ],
            [
                "mTLS (often via mesh or platform)",
                "Service-to-service mutual certificate authentication.",
                "Ensures caller really is the expected microservice in zero-trust / sensitive zones.",
            ],
            [
                "Vault / cloud KMS / Key Vault",
                "Issue, rotate, and inject secrets and short-lived DB credentials.",
                "Limits blast radius of credential theft; supports rebuild-and-rotate recovery.",
            ],
            [
                "Immutable rebuild of instances (container redeploy)",
                "Blow away and recreate compromised compute from known-good images + backups.",
                "Recovery after intrusion; reduces persistence of attackers on long-lived servers.",
            ],
        ],
    )

    doc.add_heading("13. Resiliency & scaling patterns (often implemented with libraries/platform)", level=1)
    add_table(
        doc,
        [
            ["Pattern / tool idea", "Usage (Newman)", "Problem solved"],
            [
                "Timeouts, retries (careful), idempotency",
                "Bound waiting; retry only safe operations; design APIs to tolerate duplicates.",
                "Prevents cascading hangs and double-charging / double-booking under failure.",
            ],
            [
                "Circuit breaker",
                "Stop calling a failing dependency for a cool-down period.",
                "Fails fast; gives unhealthy services time to recover; protects callers.",
            ],
            [
                "Bulkheads",
                "Isolate pools of resources (threads, connections, even service instances) so one failure doesn’t sink all.",
                "Contains blast radius — microservices themselves act as natural bulkheads vs a monolith.",
            ],
            [
                "Chaos engineering",
                "Deliberately inject failures in controlled ways.",
                "Proves resiliency claims before production incidents do.",
            ],
            [
                "Caching (e.g. Redis in landscape) / CQRS-style read models",
                "Serve hot reads without hammering source services.",
                "Scale read paths; reduce fan-out latency.",
            ],
            [
                "Horizontal scaling of instances + partitioning (Kafka partitions, DB shards as needed)",
                "Add workers; partition work/data.",
                "Throughput growth along Newman’s scaling axes without one giant process.",
            ],
        ],
    )

    doc.add_heading("14. How to use this in Kaveri 3.0 architecture discussions", level=1)
    add_para(doc, "Map Newman’s layers to Kaveri without over-buying tools early:")
    add_table(
        doc,
        [
            ["Kaveri concern", "Newman-aligned starting point"],
            ["Citizen / Officer / Admin UIs", "BFF per channel; REST; GraphQL only if aggregation pain is proven"],
            ["Module APIs (Marriage, UM, Payment…)", "REST (+ OpenAPI) as default; gRPC later for hot internal paths if needed"],
            ["Notify / Audit / MIS projections", "Broker/events (Kafka or managed equivalent); correlation IDs everywhere"],
            ["Pay-after-approve / multi-step registration", "Sagas (+ optional workflow engine) — not distributed DB transactions"],
            ["Khajane / eSign / DSC / eKYC", "Adapters behind service boundaries; timeouts, retries, idempotency, circuit breakers"],
            ["Deploy estate growth across phases", "Containers when deploy pain appears; managed Kubernetes when service count justifies it"],
            ["Security (gov / PII / Aadhaar)", "Threat model → Vault/secrets, TLS, gateway JWT, stronger zero-trust for sensitive domains"],
            ["Ops across many modules", "Log aggregation + correlation IDs first; tracing (Jaeger/OTel) as complexity grows"],
        ],
    )

    doc.add_heading("15. Bottom line from Newman", level=1)
    add_para(
        doc,
        "Pick tools **per problem and per layer**: REST/gRPC for sync calls, brokers/Kafka for async and "
        "streams, BFF/GraphQL at the UI edge, containers/Kubernetes when isolation and fleet management "
        "hurt, Pact/schemas to protect independent deployability, logs+correlation(+tracing) to operate, "
        "and Vault/JWT/mTLS according to a real threat model. "
        "**Do not** start from a product checklist and invent an architecture to justify it.",
    )

    doc.save(DST)
    return DST


if __name__ == "__main__":
    print("Wrote", build())
