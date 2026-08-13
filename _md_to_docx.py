"""Convert the Hindu Marriage BRD markdown to Word .docx (headings, tables, lists, images)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from PIL import Image

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan")
SRC = BASE / "BRD_Template_Marriage_Hindu_Marriage_Act_1955.md"
DST = BASE / "BRD_Template_Marriage_Hindu_Marriage_Act_1955.docx"

CONTENT_WIDTH_IN = 7.0
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
IMAGE_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_runs(paragraph, text: str, base_size: Pt | None = None) -> None:
    """Render **bold**, *italic* and `code` spans."""
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(token[1:-1]).italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    if base_size is not None:
        for run in paragraph.runs:
            run.font.size = base_size


def is_separator_row(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = [c.strip() for c in stripped.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c or "") for c in cells)


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    font_size = Pt(8) if width >= 6 else Pt(9)
    for r, row in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            add_runs(paragraph, row[c] if c < len(row) else "", font_size)
            if r == 0:
                shade_cell(cell, "D9E2F3")
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph()


def add_image(doc: Document, alt: str, rel_path: str) -> None:
    img_path = (SRC.parent / rel_path).resolve()
    if not img_path.exists():
        doc.add_paragraph(f"[Missing image: {rel_path}]")
        return

    with Image.open(img_path) as im:
        px_w, px_h = im.size
    dpi = 96
    width_in = min(CONTENT_WIDTH_IN, px_w / dpi)

    doc.add_picture(str(img_path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"Figure: {alt}" if alt else f"Figure: {img_path.name}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)

    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(4)


def convert(md_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    configure_styles(doc)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line or line in {"---", "***"}:
            i += 1
            continue

        if line.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                # Preserve ASCII/flow diagrams as monospace paragraphs
                code_line = lines[i].rstrip("\n")
                p = doc.add_paragraph()
                run = p.add_run(code_line if code_line else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                i += 1
            i += 1  # closing fence
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            add_runs(doc.add_paragraph(style=f"Heading {level}"), heading.group(2).strip())
            i += 1
            continue

        image = IMAGE_RE.match(line)
        if image:
            add_image(doc, image.group("alt"), image.group("path"))
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and is_separator_row(lines[i + 1]):
            rows = [split_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_separator_row(lines[i]):
                    rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if line.startswith(">"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            add_runs(paragraph, re.sub(r"^>\s?", "", line))
            for run in paragraph.runs:
                run.italic = True
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                item = re.sub(r"^[-*]\s+", "", lines[i].strip())
                add_runs(doc.add_paragraph(style="List Bullet"), item)
                i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                add_runs(doc.add_paragraph(style="List Number"), item)
                i += 1
            continue

        add_runs(doc.add_paragraph(), line)
        i += 1

    return doc


def main() -> None:
    doc = convert(SRC.read_text(encoding="utf-8"))
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_v0.3" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
