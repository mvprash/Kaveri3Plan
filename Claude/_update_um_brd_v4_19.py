# -*- coding: utf-8 -*-
"""Create BRD_User_Management_v4.19.docx from v4.18 with stakeholder-driven changes."""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

SRC = Path(r"Claude/BRD_User_Management_v4.18.docx")
DST = Path(r"Claude/BRD_User_Management_v4.19.docx")


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph text, keeping the first run's formatting when possible."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell: _Cell, text: str) -> None:
    """Replace all text in a cell with a single paragraph of text."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(text)
        return
    set_paragraph_text(paragraphs[0], text)
    # Clear remaining paragraphs
    for para in paragraphs[1:]:
        for run in para.runs:
            run.text = ""
        if not para.runs and para.text:
            para.text = ""


def replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    full = paragraph.text
    if old not in full:
        return False
    set_paragraph_text(paragraph, full.replace(old, new))
    return True


def replace_in_cell(cell: _Cell, old: str, new: str) -> bool:
    full = cell.text
    if old not in full:
        return False
    # Prefer single-paragraph replace to avoid losing structure when possible
    if len(cell.paragraphs) == 1:
        return replace_in_paragraph(cell.paragraphs[0], old, new)
    set_cell_text(cell, full.replace(old, new))
    return True


def find_req_row(table: Table, req_id: str):
    for row in table.rows:
        if row.cells[0].text.strip() == req_id:
            return row
    return None


def add_table_row_clone(table: Table, values: list[str]) -> None:
    """Append a row by cloning the last row's formatting, then set cell texts."""
    tbl = table._tbl
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    tbl.append(new_tr)
    new_row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(new_row.cells):
            set_cell_text(new_row.cells[i], val)


def insert_table_row_after(table: Table, after_index: int, values: list[str]) -> None:
    tbl = table._tbl
    ref_tr = table.rows[after_index]._tr
    new_tr = copy.deepcopy(ref_tr)
    ref_tr.addnext(new_tr)
    # After insert, find the new row (after_index + 1) and set values
    new_row = table.rows[after_index + 1]
    for i, val in enumerate(values):
        if i < len(new_row.cells):
            set_cell_text(new_row.cells[i], val)


def delete_table_row(table: Table, row_index: int) -> None:
    row = table.rows[row_index]
    tbl = table._tbl
    tbl.remove(row._tr)


def main() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # -------------------------------------------------------------------------
    # Document control
    # -------------------------------------------------------------------------
    t0 = doc.tables[0]
    set_cell_text(t0.rows[2].cells[1], "4.19")
    set_cell_text(t0.rows[11].cells[1], "2026-09-05")

    t1 = doc.tables[1]
    add_table_row_clone(
        t1,
        [
            "4.19",
            "05-Sep-2026",
            "Nandha Kumar",
            (
                "Stakeholder updates: face/biometric auth for DSR only (remove biometrics for Other "
                "Department); mandatory Aadhaar e-KYC for Citizens; remove citizen security questions; "
                "no email ID on DSR registration; DSR self-service mobile change after login; OTP resend "
                "cooldown 60s; idle timeout 10 min; absolute session 4 hours; remove temporary deputation "
                "End Date; Other Department Username = Department Code + Employee ID/KGID; Relieving "
                "Reason codes (Deputation, Transfer, Suspension, Superannuation, Death)."
            ),
        ],
    )

    # Stakeholders
    t2 = doc.tables[2]
    set_cell_text(
        t2.rows[3].cells[2],
        "Access departmental modules via OTP (Other Department) or OTP + face/biometric authentication (DSR Officers)",
    )

    # -------------------------------------------------------------------------
    # FR tables — registration (T4)
    # -------------------------------------------------------------------------
    t4 = doc.tables[4]

    set_cell_text(
        find_req_row(t4, "FR-UM-001").cells[1],
        (
            "The system shall support instant self-registration for Public users (Citizens) with no "
            "approval workflow. Registration shall capture a preferred Username chosen by the "
            "registering user (FR-UM-062), email address, and mobile number. Both the email address "
            "and the mobile number shall be verified by separate OTPs before the account is created "
            "(FR-UM-063). Aadhaar e-KYC shall be mandatory and must complete successfully before the "
            "account is created (FR-UM-085). No security questions shall be asked during Citizen "
            "registration (FR-UM-055 retired)."
        ),
    )

    set_cell_text(
        find_req_row(t4, "FR-UM-002").cells[1],
        (
            "The system shall allow Department users (DSR Officers) to be created only by authorised "
            "administrative roles. Creation shall capture KGID — which becomes the Username for this "
            "category (FR-UM-062, FR-UM-064) and is the governed unique login identifier — and mobile "
            "number. Official email ID shall not be captured during DSR Officer registration. Security "
            "questions are not captured for this category."
        ),
    )

    set_cell_text(
        find_req_row(t4, "FR-UM-003").cells[1],
        (
            "The system shall allow Other Department users (officers/staff from other government "
            "departments) to be created only by authorised administrative roles. Creation shall capture "
            "Employee ID or KGID, parent Department Code, official email ID of the parent department, "
            "and mobile number. The Username shall be formed by concatenating the Department Code with "
            "the Employee ID or KGID (FR-UM-062, FR-UM-064). Security questions are not captured for "
            "this category because no self-service mobile reset is offered (FR-UM-065)."
        ),
    )

    set_cell_text(
        find_req_row(t4, "FR-UM-062").cells[1],
        (
            "The system shall derive the Username by user category: for Public users (Citizens) the "
            "Username is a preferred Username entered by the registering user, and the system shall "
            "check availability against the whole User Master and reject a Username already in use; "
            "for DSR Officers the KGID shall be treated as the Username and shall not be user-selectable; "
            "for Other Department users the Username shall be the parent Department Code concatenated "
            "with the Employee ID or KGID (format: <DepartmentCode>-<EmployeeID|KGID>), and shall not "
            "be user-selectable. The Username shall not be changeable by the user after creation. A "
            "correction to a Username (including a KGID or Employee ID correction) shall be performed "
            "only by an authorised administrator, with reason and audit trail."
        ),
    )

    set_cell_text(
        find_req_row(t4, "FR-UM-063").cells[1],
        (
            "During Citizen self-registration the system shall verify both contact channels before "
            "creating the account: an OTP shall be sent to the entered email address and a separate "
            "OTP shall be sent to the entered mobile number, and both shall be entered correctly. "
            "Aadhaar e-KYC (FR-UM-085) shall also complete successfully before account creation. If "
            "either contact verification or Aadhaar e-KYC is not completed, the account shall not be "
            "created. Each verification OTP shall be time-limited and single-use, and failed attempts "
            "shall be rate-limited and audit-logged. No security questions shall be presented."
        ),
    )

    set_cell_text(
        find_req_row(t4, "FR-UM-064").cells[1],
        (
            "For DSR Officers the system shall validate that the KGID is present and unique in the "
            "User Master; official email ID shall not be required or captured at DSR registration. "
            "For Other Department users the system shall validate that Employee ID or KGID is present, "
            "that parent Department Code is selected from a configured department list, and that the "
            "derived Username (<DepartmentCode>-<EmployeeID|KGID>) is unique in the User Master; the "
            "system shall require an official email ID of the parent department (rather than a personal "
            "email address). Where a permitted official email domain list is configured, the system "
            "shall validate the entered Other Department email against that list and reject non-official "
            "domains."
        ),
    )

    # -------------------------------------------------------------------------
    # Auth / OTP / session FRs (T6)
    # -------------------------------------------------------------------------
    t6 = doc.tables[6]

    set_cell_text(
        find_req_row(t6, "FR-UM-006").cells[1],
        (
            "Department users (DSR Officers) shall authenticate using Username (KGID) + Captcha + OTP "
            "+ Face authentication or Biometric authentication on every login. Either face authentication "
            "or biometric authentication (e.g. fingerprint) shall satisfy the biometric factor. OTP "
            "shall be sent only to the registered mobile number."
        ),
    )

    set_cell_text(
        find_req_row(t6, "FR-UM-007").cells[1],
        (
            "Other Department users shall authenticate using Username (Department Code concatenated "
            "with Employee ID or KGID) + Captcha + OTP on every login. Biometric / face authentication "
            "shall not be required or offered for Other Department users. OTP shall be sent only to "
            "the registered mobile number."
        ),
    )

    set_cell_text(
        find_req_row(t6, "FR-UM-072").cells[1],
        (
            "The system shall enforce a sixty (60) second cooldown between successive OTP or PIN "
            "resend requests to the same user and channel (mobile or email), and a maximum of three "
            "(3) resends per channel per fifteen (15) minutes."
        ),
    )

    set_cell_text(
        find_req_row(t6, "FR-UM-074").cells[1],
        (
            "The system shall terminate a session after ten (10) minutes of no user activity (idle "
            "timeout). Re-authentication is then required (OTP, and face/biometric authentication for "
            "DSR Officers). This applies especially to shared SRO counter machines where "
            "FR-UM-052/FR-UM-053 fix the session post for the whole session."
        ),
    )

    set_cell_text(
        find_req_row(t6, "FR-UM-075").cells[1],
        (
            "The system shall terminate a session after an absolute maximum of four (4) hours from "
            "login, even if the user is still active. Re-authentication is required to continue."
        ),
    )

    set_cell_text(
        find_req_row(t6, "FR-UM-055").cells[1],
        (
            "RETIRED (v4.19) — Previously required five security questions during Citizen "
            "self-registration for lost-mobile reset. Security questions are no longer captured or "
            "used. Citizen identity proofing at registration and for lost-mobile reset is via "
            "mandatory Aadhaar e-KYC (FR-UM-085, FR-UM-056)."
        ),
    )
    set_cell_text(find_req_row(t6, "FR-UM-055").cells[2], "—")

    set_cell_text(
        find_req_row(t6, "FR-UM-056").cells[1],
        (
            "A Public user (Citizen) who has lost access to the registered mobile number shall be "
            "able to reset that mobile number before login as follows: the user enters the Username "
            "and Captcha; the system performs Aadhaar e-KYC (FR-UM-085) as identity verification; on "
            "successful e-KYC the system sends a time-limited, single-use PIN to the registered email "
            "address; on correct entry of that PIN the user may enter a new mobile number, which "
            "shall itself be verified by an OTP sent to that new number before the change takes "
            "effect. Only after the new mobile is verified shall the account's registered mobile be "
            "updated and login continue with an OTP to the new mobile. Failed e-KYC and failed PIN "
            "entries shall be rate-limited and shall lock the reset flow after a configured number of "
            "attempts. The completed change shall be notified to the registered email address and "
            "audit-logged. This flow is available only to Public users (Citizens) — see FR-UM-065 for "
            "departmental users. No security questions shall be used."
        ),
    )

    set_cell_text(
        find_req_row(t6, "FR-UM-065").cells[1],
        (
            "The lost-mobile reset flow (FR-UM-056) shall not be offered to DSR Officers or Other "
            "Department users. After login, a DSR Officer shall be able to change their registered "
            "mobile number from their own profile / login session: the new number shall be verified "
            "by OTP before the change takes effect, and the change shall be audit-logged (FR-UM-013, "
            "FR-UM-086). For Other Department users the registered mobile number shall be changed "
            "only by an authorised administrator through the User Management admin console, with a "
            "reason recorded and a full audit trail. No pre-login self-service path to change the "
            "mobile number, email address, or Username shall exist for DSR Officers or Other "
            "Department users."
        ),
    )

    # Insert FR-UM-085 after FR-UM-065 (last row of T6) — add at end of T6
    add_table_row_clone(
        t6,
        [
            "FR-UM-085",
            (
                "Aadhaar e-KYC shall be mandatory for Public users (Citizens) during self-registration "
                "and must complete successfully before the account is created (FR-UM-001, FR-UM-063). "
                "Aadhaar e-KYC shall also be used as the identity verification step in the Citizen "
                "lost-mobile reset flow (FR-UM-056). e-KYC shall comply with the Aadhaar Act, 2016 and "
                "UIDAI guidelines; Aadhaar number / VID handling, consent, and audit shall follow "
                "applicable MeitY / UIDAI norms. Aadhaar e-KYC is not required for DSR Officers or "
                "Other Department users at registration."
            ),
            "High",
        ],
    )
    add_table_row_clone(
        t6,
        [
            "FR-UM-086",
            (
                "After login, a DSR Officer shall be able to update their registered mobile number "
                "from within their authenticated session (profile / account settings) without "
                "administrator intervention. The system shall verify the new mobile number by OTP "
                "before saving, notify the officer of the change, and write a full audit trail "
                "(actor, timestamp, old and new number masked). Official email is out of scope for "
                "DSR registration (FR-UM-002); Other Department mobile changes remain "
                "administrator-only (FR-UM-065)."
            ),
            "High",
        ],
    )

    # Profile FR (T8)
    t8 = doc.tables[8]
    set_cell_text(
        find_req_row(t8, "FR-UM-013").cells[1],
        (
            "The system shall allow users to view and update their profile information after login. "
            "The Username shall be displayed read-only and shall not be user-editable (FR-UM-062). "
            "Public users (Citizens) may update their mobile number after login, verified by an OTP "
            "to the new number, and may update their email address, verified by an OTP to the new "
            "address. DSR Officers may update their mobile number after login under FR-UM-086 "
            "(OTP-verified). For Other Department users, mobile number and official email ID are "
            "maintained by an authorised administrator only (FR-UM-065). All such changes shall be "
            "audit-logged and notified to the user. Security questions are not part of the profile "
            "(FR-UM-055 retired)."
        ),
    )

    # Large FR table (T9) — deputation + relieving
    t9 = doc.tables[9]
    set_cell_text(
        find_req_row(t9, "FR-UM-030").cells[1],
        (
            "At least one sanctioned post with available capacity must be assigned to each DSR "
            "Officer at account creation (FR-UM-066(a)). Temporary / time-bound deputation via "
            "optional occupancy End Date is removed (v4.19) — post occupancy End Date and Deputation "
            "Reason shall not be captured at user creation or post assignment. Formal Transfer out / "
            "relieving shall follow FR-UM-057 and FR-UM-058 (including Relieving Reason under "
            "FR-UM-087). Transfer In shall follow FR-UM-060, FR-UM-061, and FR-UM-067. Occupancy "
            "de-allocation, reserved Transfer In becoming active, and Sanctioned Posts occupied-count "
            "updates shall be applied by the occupancy refresh job (FR-UM-068)."
        ),
    )

    set_cell_text(
        find_req_row(t9, "FR-UM-057").cells[1],
        (
            "The system shall support a Transfer out / relieving process for DSR Officers. An officer "
            "may be relieved from one or more currently assigned post occupancies (Post + Office). "
            "When the superior opens Transfer Out / Relieving, the system shall display only the "
            "offices that fall under the actor's session Office as per the Office Hierarchy Master "
            "(FR-UM-059) — i.e. the actor's own office and descendant offices. Within that office "
            "scope, the system shall list and allow relieving only of post occupancies where the "
            "actor's session Post is the immediate parent of the target Post in the DSR Officer "
            "Hierarchy Master (FR-UM-043). Office span is a first filter only: seeing a descendant "
            "office in the tree does not grant relieving of posts that do not report immediately to "
            "the actor. Example: District Registrar at DRO Mysuru sees SRO offices under DRO Mysuru "
            "and may relieve Sub-Registrar at those SROs (POST-SR reports to POST-DRO); Sub-Registrar "
            "at SRO Yeshwanthapura sees only that SRO and may relieve FDA/DEO under Sub-Registrar "
            "there. IGR or DIGR at Head Office may see statewide DRO/SRO offices in the span but may "
            "relieve only posts that report immediately to their session Post (e.g. IGR relieves "
            "DIGR; DIGR (Enforcement) relieves DRO) — not Sub-Registrar at an SRO. While relieving, "
            "the system shall capture Relieving Date, Relieving Reason (mandatory — FR-UM-087), and "
            "Relieving Order (order number / reference; upload of order document where applicable). "
            "Relieving actions shall be audit-logged. Citizens and Other Department users are out of "
            "scope for this process."
        ),
    )

    set_cell_text(
        find_req_row(t9, "FR-UM-068").cells[1],
        (
            "The system shall run a scheduled occupancy refresh job shortly after midnight IST each "
            "calendar day (after 12:00 AM IST). On each run the job shall: (1) de-allocate user–post "
            "occupancies whose Relieving Date has ended (FR-UM-058); (2) activate reserved Transfer In "
            "occupancies whose Joining Date is the current calendar day (FR-UM-060, FR-UM-061, "
            "FR-UM-067); (3) recalculate and persist occupied count, remaining capacity, and "
            "wholly-unoccupied flag on the Sanctioned Posts Master for every Post + Office affected "
            "(FR-UM-048, FR-UM-066); (4) refresh the officer's effective post assignments so that "
            "login post selection (FR-UM-052) reflects only occupancies that are active as of that "
            "day. The job shall be idempotent, audit-logged (run timestamp, occupancies ended, "
            "occupancies activated, before/after occupied counts), and shall raise an operational "
            "alert on failure so that relieving and joining dates are not left unenforced. Optional "
            "occupancy End Date / temporary deputation processing is removed (v4.19; see FR-UM-030)."
        ),
    )

    add_table_row_clone(
        t9,
        [
            "FR-UM-087",
            (
                "During Transfer Out / Relieving (FR-UM-057) the system shall require a Relieving "
                "Reason selected from the following enumerated list only: Deputation; Transfer; "
                "Suspension; Superannuation; Death. Free-text reasons shall not replace the "
                "enumerated code. The selected Relieving Reason shall be stored with the relieving "
                "record, shown in Transfer history reports, and included in the audit trail."
            ),
            "High",
        ],
    )

    # -------------------------------------------------------------------------
    # Narrative paragraphs
    # -------------------------------------------------------------------------
    replacements = [
        (
            "Citizens self-register instantly with no approval workflow. The registering user chooses their own Username, which the system checks for availability across the entire User Master. Both the email address and the mobile number are verified by separate OTPs before the account is created, so neither channel can be claimed without proof of control. Five security questions are captured at this point; they are used only for the lost-mobile reset in FR-UM-056 and are entered by the citizen alone — never by an administrator.",
            "Citizens self-register instantly with no approval workflow. The registering user chooses their own Username, which the system checks for availability across the entire User Master. Both the email address and the mobile number are verified by separate OTPs before the account is created, so neither channel can be claimed without proof of control. Aadhaar e-KYC is mandatory and must succeed before the account is created (FR-UM-085). No security questions are asked during registration (FR-UM-055 retired).",
        ),
        (
            "6.1.1 Citizen Self-Registration Workflow (FR-UM-001, FR-UM-055, FR-UM-062, FR-UM-063)",
            "6.1.1 Citizen Self-Registration Workflow (FR-UM-001, FR-UM-062, FR-UM-063, FR-UM-085)",
        ),
        (
            "Approved process diagram — P-01 Citizen Self-Registration — FR-UM-001, FR-UM-055, FR-UM-062, FR-UM-063 — ProcessDiagrams/User_Management/P-01_Citizen_Self_Registration.drawio",
            "Approved process diagram — P-01 Citizen Self-Registration — FR-UM-001, FR-UM-062, FR-UM-063, FR-UM-085 — ProcessDiagrams/User_Management/P-01_Citizen_Self_Registration.drawio",
        ),
        (
            "All users are stored in the single User Master, differentiated by User Category. All user categories authenticate without passwords. Login uses the Username; the login OTP is sent only to the registered mobile. Captcha is required before OTP dispatch. The Username is a preferred Username chosen by the Citizen at registration, and the KGID for DSR Officers and Other Department users (FR-UM-062). The Username is the single unique login identifier across the whole User Master; email address and mobile number carry no uniqueness constraint (FR-UM-004). DSR Officers and Other Department users also verify Biometrics. Citizens may update their mobile after login (FR-UM-013) and may reset a lost mobile before login using three of five security questions plus a PIN to the registered email (FR-UM-056); departmental users have no self-service mobile reset and are served by an administrator instead (FR-UM-065). DSR Officers with multiple active post occupancies must select one post (with office details) after authentication before entering the application (FR-UM-052). After login they may take additional charge of a wholly unoccupied subordinate post at the same office without logout (FR-UM-053). The home/header shows the assigned Post with mapped Role(s) and any active additional charge (FR-UM-054). OTP and session rules are defined in FR-UM-069–FR-UM-076 (login OTP 5 minutes, idle timeout 15 minutes, one session per Username).",
            "All users are stored in the single User Master, differentiated by User Category. All user categories authenticate without passwords. Login uses the Username; the login OTP is sent only to the registered mobile. Captcha is required before OTP dispatch. The Username is a preferred Username chosen by the Citizen at registration, the KGID for DSR Officers, and Department Code concatenated with Employee ID or KGID for Other Department users (FR-UM-062). The Username is the single unique login identifier across the whole User Master; email address and mobile number carry no uniqueness constraint (FR-UM-004). DSR Officers additionally verify face authentication or biometrics; Other Department users do not use biometrics (FR-UM-006, FR-UM-007). Citizens complete mandatory Aadhaar e-KYC at registration (FR-UM-085), may update their mobile after login (FR-UM-013), and may reset a lost mobile before login using Aadhaar e-KYC plus a PIN to the registered email (FR-UM-056). DSR Officers may change their mobile after login (FR-UM-086); Other Department users are served by an administrator (FR-UM-065). DSR Officers with multiple active post occupancies must select one post (with office details) after authentication before entering the application (FR-UM-052). After login they may take additional charge of a wholly unoccupied subordinate post at the same office without logout (FR-UM-053). The home/header shows the assigned Post with mapped Role(s) and any active additional charge (FR-UM-054). OTP and session rules are defined in FR-UM-069–FR-UM-076 (login OTP 5 minutes, idle timeout 10 minutes, absolute session 4 hours, one session per Username).",
        ),
        (
            "Login uses the Username — the preferred Username for Citizens and the KGID for DSR Officers and Other Department users. Because the Username is unique across the whole User Master, the account is resolved unambiguously from the Username alone and no category selection is needed. The login OTP is never emailed; it is sent only by SMS to the registered mobile.",
            "Login uses the Username — the preferred Username for Citizens, the KGID for DSR Officers, and Department Code concatenated with Employee ID or KGID for Other Department users. Because the Username is unique across the whole User Master, the account is resolved unambiguously from the Username alone and no category selection is needed. The login OTP is never emailed; it is sent only by SMS to the registered mobile. DSR Officers complete face or biometric authentication after OTP; Other Department users do not.",
        ),
        (
            "6.5.2.2 Citizen Lost-Mobile Reset — Three of Five Questions Plus PIN to Registered Email (FR-UM-056)",
            "6.5.2.2 Citizen Lost-Mobile Reset — Aadhaar e-KYC Plus PIN to Registered Email (FR-UM-056, FR-UM-085)",
        ),
        (
            "This flow exists because the login OTP goes only to the registered mobile, so a citizen who has lost that number cannot otherwise log in. It requires two independent proofs before the mobile number can be changed: three security questions chosen at random from the five registered under FR-UM-055 (all three must be answered correctly, and the selection differs on each attempt), and a single-use PIN sent to the registered email address. The new mobile number is itself verified by OTP before the change takes effect. This flow is available only to Public users (Citizens).",
            "This flow exists because the login OTP goes only to the registered mobile, so a citizen who has lost that number cannot otherwise log in. It requires two independent proofs before the mobile number can be changed: successful Aadhaar e-KYC (FR-UM-085), and a single-use PIN sent to the registered email address. The new mobile number is itself verified by OTP before the change takes effect. No security questions are used. This flow is available only to Public users (Citizens).",
        ),
        (
            "6.5.2.3 Departmental Mobile Change — Administrator Only, No Self-Service (FR-UM-065)",
            "6.5.2.3 Mobile Change — DSR Self-Service After Login; Other Department Administrator Only (FR-UM-065, FR-UM-086)",
        ),
        (
            "DSR Officers and Other Department users are not offered the FR-UM-056 reset path, and no security questions are held for them. A departmental user who has lost the registered mobile is served by an authorised administrator, who records a reason and whose action is audit-logged. This keeps the officer's authentication factors under departmental control and removes any pre-login route to redirect an officer's OTP.",
            "DSR Officers and Other Department users are not offered the FR-UM-056 pre-login reset path, and no security questions are held for them. A DSR Officer who needs to change the registered mobile after login does so from their own session (FR-UM-086) with OTP verification of the new number. An Other Department user who has lost the registered mobile is served by an authorised administrator, who records a reason and whose action is audit-logged (FR-UM-065). There is no pre-login self-service route to redirect a departmental OTP.",
        ),
        (
            "Applies only to DSR Officers. After Username (KGID) + Captcha + OTP (to mobile) + Biometrics succeed, the system loads the officer's active sanctioned-post occupancies. If more than one is active, the user must choose which post to work under for this session. Each choice shows Post Name and Office details so dual-charge / multi-office officers can pick the correct context.",
            "Applies only to DSR Officers. After Username (KGID) + Captcha + OTP (to mobile) + Face authentication or Biometrics succeed, the system loads the officer's active sanctioned-post occupancies. If more than one is active, the user must choose which post to work under for this session. Each choice shows Post Name and Office details so dual-charge / multi-office officers can pick the correct context.",
        ),
        (
            "The system shall provide dedicated step-by-step workflows to assign access during user creation. For DSR Officers and Other Department users the KGID is entered first because it becomes the Username (FR-UM-062); the Citizen self-registration flow is set out in 6.1.1. For DSR Officers, assign one or more sanctioned posts with available capacity (roles via Post–Role mapping; no Primary/Secondary; at login the officer selects one post when multiple are active — FR-UM-052). Transfer out / relieving and Transfer In are scoped to offices under the actor, then to posts whose immediate parent is the actor's session Post (6.6.3–6.6.4, FR-UM-057–FR-UM-061, FR-UM-067). Occupancies take effect or end via the occupancy refresh job after midnight IST (FR-UM-068). DSR user creation uses KGID as Username with no security questions; maker-checker is not required (FR-UM-051). For Other Department users, select exactly one role.",
            "The system shall provide dedicated step-by-step workflows to assign access during user creation. For DSR Officers the KGID is entered first because it becomes the Username; for Other Department users Employee ID or KGID is entered with Department Code and the Username is the concatenation (FR-UM-062). The Citizen self-registration flow is set out in 6.1.1. For DSR Officers, assign one or more sanctioned posts with available capacity (roles via Post–Role mapping; no Primary/Secondary; at login the officer selects one post when multiple are active — FR-UM-052). No temporary deputation End Date is captured (FR-UM-030). Transfer out / relieving and Transfer In are scoped to offices under the actor, then to posts whose immediate parent is the actor's session Post (6.6.3–6.6.4, FR-UM-057–FR-UM-061, FR-UM-067); Relieving Reason is mandatory (FR-UM-087). Occupancies take effect or end via the occupancy refresh job after midnight IST (FR-UM-068). DSR user creation uses KGID as Username with no email ID and no security questions; maker-checker is not required (FR-UM-051). For Other Department users, select exactly one role; biometrics are not captured.",
        ),
        (
            "A DSR Officer may be relieved from currently assigned post occupancy(ies). When Transfer Out / Relieving is opened, the system shall list offices under the actor's session Office per the Office Hierarchy Master — IGR Head Office → District Registrar Offices → Sub-Registrar Offices (FR-UM-059). Within that office span, the system shall list and allow relieving only where the actor's session Post is the immediate parent of the target Post in the Officer Hierarchy (6.5.7, FR-UM-057). Seeing a descendant office (for example an SRO under Head Office) does not allow relieving a Sub-Registrar there. Capture Relieving Date and Relieving Order. Mapping is removed after 11:59 PM IST of the Relieving Date by the occupancy refresh job (FR-UM-058, FR-UM-068).",
            "A DSR Officer may be relieved from currently assigned post occupancy(ies). When Transfer Out / Relieving is opened, the system shall list offices under the actor's session Office per the Office Hierarchy Master — IGR Head Office → District Registrar Offices → Sub-Registrar Offices (FR-UM-059). Within that office span, the system shall list and allow relieving only where the actor's session Post is the immediate parent of the target Post in the Officer Hierarchy (6.5.7, FR-UM-057). Seeing a descendant office (for example an SRO under Head Office) does not allow relieving a Sub-Registrar there. Capture Relieving Date, Relieving Reason (Deputation / Transfer / Suspension / Superannuation / Death — FR-UM-087), and Relieving Order. Mapping is removed after 11:59 PM IST of the Relieving Date by the occupancy refresh job (FR-UM-058, FR-UM-068).",
        ),
        (
            "This section covers temporary unavailability of a DSR Officer who remains the occupant of their post (Leave, OOD, or Other). It is distinct from Transfer Out / relieving (FR-UM-057), Transfer In (FR-UM-060), deputation End Date (FR-UM-030), and post-login additional charge of a wholly unoccupied subordinate post (FR-UM-053). Full HRMS leave balances and payroll are out of scope — User Management records absence and temporary charge for access control and operational continuity only.",
            "This section covers temporary unavailability of a DSR Officer who remains the occupant of their post (Leave, OOD, or Other). It is distinct from Transfer Out / relieving (FR-UM-057), Transfer In (FR-UM-060), and post-login additional charge of a wholly unoccupied subordinate post (FR-UM-053). Temporary / time-bound deputation via occupancy End Date has been removed (FR-UM-030, v4.19). Full HRMS leave balances and payroll are out of scope — User Management records absence and temporary charge for access control and operational continuity only.",
        ),
        (
            "UAT — FR-UM-069–FR-UM-076: Login OTP expires after 5 minutes; 6 digits; 3 incorrect entries invalidate the code; 30-second resend cooldown; 5 failed logins lock Username 15 minutes; idle timeout 15 minutes; session ends at 8 hours; a second login ends the prior session.",
            "UAT — FR-UM-069–FR-UM-076: Login OTP expires after 5 minutes; 6 digits; 3 incorrect entries invalidate the code; 60-second resend cooldown; 5 failed logins lock Username 15 minutes; idle timeout 10 minutes; session ends at 4 hours; a second login ends the prior session.",
        ),
    ]

    for p in doc.paragraphs:
        for old, new in replacements:
            if old in p.text:
                replace_in_paragraph(p, old, new)

    # Reporting paragraph that mentions security-question recovery
    for p in doc.paragraphs:
        if "Citizen lost-mobile / security-question recovery" in p.text:
            replace_in_paragraph(
                p,
                "Citizen lost-mobile / security-question recovery attempts (FR-UM-056) — Username, questions presented, PIN dispatch outcome, new mobile OTP outcome, success/failure, actor, and timestamp; (b) administrator-initiated mobile changes for departmental users (FR-UM-065)",
                "Citizen lost-mobile / Aadhaar e-KYC recovery attempts (FR-UM-056) — Username, e-KYC outcome, PIN dispatch outcome, new mobile OTP outcome, success/failure, actor, and timestamp; (b) DSR self-service mobile changes (FR-UM-086) and administrator-initiated mobile changes for Other Department users (FR-UM-065)",
            )

    # -------------------------------------------------------------------------
    # Process / summary tables
    # -------------------------------------------------------------------------
    # T5 Citizen registration steps
    t5 = doc.tables[5]
    set_cell_text(t5.rows[7].cells[1], "Complete mandatory Aadhaar e-KYC")
    set_cell_text(t5.rows[7].cells[2], "Citizen / System")
    set_cell_text(
        t5.rows[7].cells[3],
        "FR-UM-085; account not created unless e-KYC succeeds; no security questions",
    )
    set_cell_text(
        t5.rows[8].cells[1],
        "Save — account created; Citizen role assigned",
    )
    set_cell_text(
        t5.rows[8].cells[3],
        "FR-UM-001, FR-UM-063, FR-UM-085; no approval workflow; no security questions",
    )

    # T13 category matrix
    t13 = doc.tables[13]
    set_cell_text(
        t13.rows[1].cells[2],
        "Username + Captcha + OTP to mobile; Aadhaar e-KYC at registration (FR-UM-085)",
    )
    set_cell_text(
        t13.rows[1].cells[3],
        "Self-service: Aadhaar e-KYC + PIN to registered email + OTP to new mobile (FR-UM-056)",
    )
    set_cell_text(
        t13.rows[2].cells[2],
        "Username (KGID) + Captcha + OTP to mobile + Face authentication or Biometrics; then FR-UM-052",
    )
    set_cell_text(
        t13.rows[2].cells[3],
        "Self-service after login (FR-UM-086); no pre-login reset — admin path not required for routine mobile change",
    )
    set_cell_text(
        t13.rows[3].cells[1],
        "Department Code + Employee ID or KGID (FR-UM-062, FR-UM-064)",
    )
    set_cell_text(
        t13.rows[3].cells[2],
        "Username + Captcha + OTP to mobile (no biometrics)",
    )
    set_cell_text(
        t13.rows[3].cells[3],
        "Not available — administrator changes mobile with reason and audit (FR-UM-065)",
    )

    # T14 login steps
    t14 = doc.tables[14]
    set_cell_text(
        t14.rows[1].cells[3],
        "Citizen: preferred Username; DSR: KGID; Other Dept: Dept Code + Emp ID/KGID (FR-UM-062)",
    )
    set_cell_text(
        t14.rows[4].cells[1],
        "Enter OTP (+ Face authentication or Biometrics for DSR only)",
    )
    set_cell_text(
        t14.rows[4].cells[3],
        "FR-UM-005 / FR-UM-006 / FR-UM-007; Other Department — OTP only",
    )
    set_cell_text(
        t14.rows[7].cells[1],
        "After login — Citizen or DSR Officer may update mobile from profile",
    )
    set_cell_text(
        t14.rows[7].cells[3],
        "FR-UM-013, FR-UM-086; Other Department via administrator (FR-UM-065)",
    )

    # T15 lost-mobile steps — replace question steps with e-KYC
    t15 = doc.tables[15]
    set_cell_text(t15.rows[3].cells[1], "System initiates Aadhaar e-KYC for identity verification")
    set_cell_text(t15.rows[3].cells[2], "System / Citizen")
    set_cell_text(t15.rows[3].cells[3], "FR-UM-085, FR-UM-056; replaces security questions")
    set_cell_text(t15.rows[4].cells[1], "Complete Aadhaar e-KYC successfully")
    set_cell_text(t15.rows[4].cells[2], "Citizen")
    set_cell_text(t15.rows[4].cells[3], "Must succeed before PIN dispatch; rate-limited")

    # T16 departmental mobile — retitle for Other Dept / note DSR self-service
    t16 = doc.tables[16]
    set_cell_text(
        t16.rows[1].cells[1],
        "Other Department user reports the lost / changed mobile to the authorised administrator (DSR Officers use self-service after login — FR-UM-086)",
    )
    set_cell_text(t16.rows[1].cells[2], "Other Dept user")
    set_cell_text(
        t16.rows[1].cells[3],
        "Identity verified off-system per departmental procedure; DSR path is FR-UM-086",
    )
    set_cell_text(
        t16.rows[2].cells[1],
        "Administrator opens User Management and locates the Other Department user by Username",
    )
    set_cell_text(
        t16.rows[6].cells[1],
        "User logs in with Username + Captcha + OTP to new mobile (Other Dept: no biometrics; DSR self-service path uses face/biometric at next login)",
    )
    set_cell_text(t16.rows[6].cells[3], "FR-UM-007 / FR-UM-006; FR-UM-086 for DSR")

    # T17 DSR post selection auth step
    t17 = doc.tables[17]
    set_cell_text(
        t17.rows[1].cells[1],
        "Authenticate (Username (KGID) + Captcha + OTP to mobile + Face authentication or Biometrics)",
    )

    # T35 runtime
    t35 = doc.tables[35]
    set_cell_text(
        t35.rows[1].cells[1],
        "User authenticates (Username + Captcha + OTP to mobile; + Face/Biometrics if DSR only)",
    )

    # T42 DSR creation — remove email, remove deputation End Date step, keep biometrics
    t42 = doc.tables[42]
    set_cell_text(
        t42.rows[3].cells[1],
        "Enter remaining particulars (name, mobile, photo, ID proof) — no email ID",
    )
    set_cell_text(
        t42.rows[3].cells[3],
        "Email ID not captured for DSR (FR-UM-002, FR-UM-064); mobile need not be unique (FR-UM-004)",
    )
    set_cell_text(
        t42.rows[4].cells[3],
        "No security questions; DSR may change mobile after login (FR-UM-086)",
    )
    set_cell_text(
        t42.rows[5].cells[3],
        "At least one required; FR-UM-017, FR-UM-030, FR-UM-045, FR-UM-048, FR-UM-066(a); no temporary deputation End Date",
    )
    # Row 6 is deputation End Date — retire the step
    set_cell_text(t42.rows[6].cells[1], "Temporary deputation End Date — REMOVED (v4.19)")
    set_cell_text(t42.rows[6].cells[2], "—")
    set_cell_text(
        t42.rows[6].cells[3],
        "FR-UM-030: optional occupancy End Date / Deputation Reason removed; use Transfer Out / In instead",
    )
    set_cell_text(
        t42.rows[9].cells[1],
        "Capture face authentication template or biometrics",
    )
    set_cell_text(
        t42.rows[9].cells[3],
        "Mandatory for DSR users — face or biometric (FR-UM-006)",
    )

    # T43 Other Dept creation — username, remove biometrics
    t43 = doc.tables[43]
    set_cell_text(
        t43.rows[2].cells[1],
        "Enter Employee ID or KGID and parent Department Code — Username = Dept Code + ID",
    )
    set_cell_text(
        t43.rows[2].cells[3],
        "Username format <DepartmentCode>-<EmployeeID|KGID>; validated unique (FR-UM-062, FR-UM-064)",
    )
    set_cell_text(
        t43.rows[3].cells[1],
        "Enter remaining particulars (name, official email ID of parent department, mobile, photo, ID proof)",
    )
    set_cell_text(
        t43.rows[9].cells[1],
        "Biometrics — NOT required (removed v4.19)",
    )
    set_cell_text(t43.rows[9].cells[2], "—")
    set_cell_text(
        t43.rows[9].cells[3],
        "FR-UM-007: Other Department authenticates with Username + Captcha + OTP only",
    )

    # T45 relieving — add Relieving Reason
    t45 = doc.tables[45]
    set_cell_text(
        t45.rows[5].cells[1],
        "Enter Relieving Date, Relieving Reason, and Relieving Order (upload if applicable)",
    )
    set_cell_text(
        t45.rows[5].cells[3],
        "Mandatory (FR-UM-057, FR-UM-087); Reason = Deputation / Transfer / Suspension / Superannuation / Death",
    )

    # -------------------------------------------------------------------------
    # NFR / Risks / Glossary
    # -------------------------------------------------------------------------
    t52 = doc.tables[52]
    set_cell_text(
        t52.rows[2].cells[1],
        (
            "OTP and session policy is defined in FR-UM-069–FR-UM-076: login OTP 5 minutes; "
            "registration/recovery OTP and reset PIN 10 minutes; 6 digits; 3 incorrect entries per "
            "code; 60-second resend cooldown (max 3 per 15 minutes); 5 failed login attempts lock "
            "Username 15 minutes; idle timeout 10 minutes; absolute session 4 hours; one active "
            "session per Username."
        ),
    )
    set_cell_text(
        t52.rows[4].cells[1],
        (
            "Face authentication / biometric data for DSR Officers shall comply with the Aadhaar "
            "Act, 2016 and UIDAI guidelines where Aadhaar-linked biometrics are used. Citizen "
            "Aadhaar e-KYC (FR-UM-085) shall likewise comply with UIDAI / MeitY norms. Other "
            "Department users shall not use biometric authentication."
        ),
    )
    set_cell_text(
        t52.rows[5].cells[1],
        (
            "RETIRED (v4.19) — Security-question storage rules no longer apply; security questions "
            "are not captured (FR-UM-055 retired). Citizen identity proofing uses Aadhaar e-KYC "
            "(FR-UM-085)."
        ),
    )
    set_cell_text(
        t52.rows[6].cells[1],
        (
            "The Citizen lost-mobile reset shall require two independent proofs — successful "
            "Aadhaar e-KYC and a single-use PIN delivered to the registered email — before a new "
            "mobile number may be entered, and the new number shall itself be OTP-verified before "
            "the change takes effect (FR-UM-056, FR-UM-085)."
        ),
    )
    set_cell_text(
        t52.rows[7].cells[1],
        (
            "The reset PIN shall be single-use and time-limited; e-KYC and PIN attempts shall be "
            "rate-limited and the reset flow locked after a configured number of failures."
        ),
    )
    set_cell_text(
        t52.rows[8].cells[1],
        (
            "No pre-login self-service path shall exist to change the mobile number, email address, "
            "or Username of a DSR Officer or Other Department user. After login, DSR Officers may "
            "change mobile under FR-UM-086; Other Department mobile remains administrator-only "
            "(FR-UM-065)."
        ),
    )

    t53 = doc.tables[53]
    set_cell_text(
        t53.rows[1].cells[2],
        (
            "Retry mechanism; allow mobile update after login for Citizens (FR-UM-013) and DSR "
            "Officers (FR-UM-086); Citizen lost-mobile reset via Aadhaar e-KYC (FR-UM-056); "
            "administrator-initiated change for Other Department users (FR-UM-065)."
        ),
    )
    set_cell_text(
        t53.rows[2].cells[0],
        "Face / biometric device unavailability (DSR)",
    )
    set_cell_text(
        t53.rows[2].cells[2],
        (
            "Face authentication or biometrics are mandatory on every login for DSR Officers "
            "(FR-UM-006) with no exception path — offices must maintain working devices; Other "
            "Department users do not require biometrics (FR-UM-007)."
        ),
    )
    set_cell_text(
        t53.rows[3].cells[2],
        (
            "Two independent proofs required — Aadhaar e-KYC plus single-use PIN to registered "
            "email — and OTP verification of the new mobile; attempts rate-limited and flow locked "
            "after configured failures; change notified to registered email and audit-logged."
        ),
    )
    set_cell_text(
        t53.rows[4].cells[0],
        "Aadhaar e-KYC failure or UIDAI downtime",
    )
    set_cell_text(
        t53.rows[4].cells[2],
        (
            "Registration and lost-mobile reset depend on Aadhaar e-KYC (FR-UM-085). Provide clear "
            "user messaging, retry, and operational monitoring of UIDAI connectivity; no security-"
            "question fallback (FR-UM-055 retired)."
        ),
    )
    set_cell_text(
        t53.rows[12].cells[2],
        (
            "Enforce FR-UM-069–FR-UM-076 as specified (including 10-minute idle timeout and 4-hour "
            "absolute session on shared SRO counters with FR-UM-052/FR-UM-053)."
        ),
    )
    set_cell_text(
        t53.rows[21].cells[2],
        (
            "FR-UM-074 (10-minute idle timeout) combined with FR-UM-005–FR-UM-006 (OTP and "
            "face/biometrics on every fresh DSR login) adds authentication overhead at busy "
            "Sub-Registrar counters. Beyond functional UAT (9), measure end-to-end re-authentication "
            "time during performance testing with Kaveri IT Cell."
        ),
    )

    t54 = doc.tables[54]
    # Username / KGID / Security question / Lost-mobile / Deputation Reason / Relieving
    for row in t54.rows:
        key = row.cells[0].text.strip()
        if key == "Username":
            set_cell_text(
                row.cells[1],
                (
                    "The unique login identifier for every account, unique across the whole User "
                    "Master: a preferred Username chosen by the citizen for Public users; the KGID "
                    "for DSR Officers; and Department Code concatenated with Employee ID or KGID "
                    "(<DepartmentCode>-<EmployeeID|KGID>) for Other Department users (FR-UM-062)"
                ),
            )
        elif key == "KGID":
            set_cell_text(
                row.cells[1],
                (
                    "Karnataka Government Insurance Department number — used as the Username for "
                    "DSR Officers, and as one permitted identifier (with Department Code) for Other "
                    "Department Username formation (FR-UM-062, FR-UM-064)"
                ),
            )
        elif key == "Official email ID":
            set_cell_text(
                row.cells[1],
                (
                    "Government or parent-department email address recorded for Other Department "
                    "users; not captured during DSR Officer registration (FR-UM-002, FR-UM-064); "
                    "validated against the permitted domain list where configured"
                ),
            )
        elif key == "Security question":
            set_cell_text(
                row.cells[1],
                (
                    "RETIRED (v4.19) — Previously one of five questions selected by a citizen at "
                    "registration; no longer captured or used. Replaced by Aadhaar e-KYC "
                    "(FR-UM-085, FR-UM-056)"
                ),
            )
        elif key == "Lost-mobile reset":
            set_cell_text(
                row.cells[1],
                (
                    "Citizen-only pre-login flow to replace the registered mobile number: Aadhaar "
                    "e-KYC, then a PIN to the registered email, then OTP verification of the new "
                    "number (FR-UM-056, FR-UM-085); not available to departmental users"
                ),
            )
        elif key == "Deputation Reason":
            set_cell_text(
                row.cells[1],
                (
                    "RETIRED (v4.19) — Previously mandatory when optional occupancy End Date was "
                    "set under FR-UM-030; temporary deputation End Date feature removed. Deputation "
                    "as a Relieving Reason remains under FR-UM-087"
                ),
            )
        elif key == "Relieving / Transfer out":
            set_cell_text(
                row.cells[1],
                (
                    "Process by which a superior removes a DSR Officer from a post occupancy within "
                    "the superior's office span and immediate-parent posts only, capturing Relieving "
                    "Date, Relieving Reason (FR-UM-087), and Relieving Order (FR-UM-057–FR-UM-059); "
                    "mapping ends after 11:59 PM IST of that date via the occupancy refresh job "
                    "(FR-UM-058, FR-UM-068)"
                ),
            )
        elif key == "OTP":
            set_cell_text(
                row.cells[1],
                (
                    "One-Time Password — the login OTP is sent only to the registered mobile (SMS); "
                    "separate one-time OTPs verify the email address and mobile number at Citizen "
                    "registration (FR-UM-063) and verify a new mobile / email on change; resend "
                    "cooldown is 60 seconds (FR-UM-072)"
                ),
            )

    # Add glossary entries
    add_table_row_clone(
        t54,
        [
            "Aadhaar e-KYC",
            (
                "UIDAI electronic Know Your Customer verification mandatory for Citizen "
                "self-registration and used as identity proof in Citizen lost-mobile reset "
                "(FR-UM-085, FR-UM-056)"
            ),
        ],
    )
    add_table_row_clone(
        t54,
        [
            "Face authentication",
            (
                "Facial biometric verification permitted as the biometric factor for DSR Officer "
                "login in place of or as an alternative to fingerprint / other biometrics "
                "(FR-UM-006)"
            ),
        ],
    )
    add_table_row_clone(
        t54,
        [
            "Employee ID",
            (
                "Parent-department employee identifier that may be used (instead of or alongside "
                "KGID) when forming Other Department Username with Department Code (FR-UM-062)"
            ),
        ],
    )
    add_table_row_clone(
        t54,
        [
            "Department Code",
            (
                "Configured code for the parent government department of an Other Department user; "
                "concatenated with Employee ID or KGID to form the Username (FR-UM-062, FR-UM-064)"
            ),
        ],
    )
    add_table_row_clone(
        t54,
        [
            "Relieving Reason",
            (
                "Mandatory enumerated reason captured during Transfer Out / Relieving: Deputation; "
                "Transfer; Suspension; Superannuation; Death (FR-UM-087)"
            ),
        ],
    )

    # Broad residual string replacements across all tables
    residual = [
        ("thirty (30) second cooldown", "sixty (60) second cooldown"),
        ("30-second resend cooldown", "60-second resend cooldown"),
        ("thirty (30) second", "sixty (60) second"),
        ("fifteen (15) minutes of no user activity", "ten (10) minutes of no user activity"),
        ("idle timeout 15 minutes", "idle timeout 10 minutes"),
        ("idle timeout of 15 minutes", "idle timeout of 10 minutes"),
        ("FR-UM-074 (15-minute idle timeout)", "FR-UM-074 (10-minute idle timeout)"),
        ("15-minute idle timeout", "10-minute idle timeout"),
        ("absolute maximum of eight (8) hours", "absolute maximum of four (4) hours"),
        ("absolute session 8 hours", "absolute session 4 hours"),
        ("session ends at 8 hours", "session ends at 4 hours"),
        ("idle timeout 15 minutes; absolute session 8 hours", "idle timeout 10 minutes; absolute session 4 hours"),
        ("+ Biometrics for DSR / Other Department", "+ Face/Biometrics for DSR only"),
        ("+ Biometrics if DSR/Other Dept", "+ Face/Biometrics if DSR only"),
        ("OTP + biometrics", "OTP + face/biometrics (DSR) or OTP only (Other Department)"),
    ]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                new_text = text
                for old, new in residual:
                    if old in new_text:
                        new_text = new_text.replace(old, new)
                if new_text != text:
                    set_cell_text(cell, new_text)

    for p in doc.paragraphs:
        text = p.text
        new_text = text
        for old, new in residual:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(p, new_text)

    doc.save(str(DST))
    print(f"Saved {DST}")


if __name__ == "__main__":
    main()
