"""Generate NFR Client Questionnaire .docx for Hindu Marriage (HMA 1955) module."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DST = Path(
    r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan"
    r"\NFR_Questionnaire_Hindu_Marriage_HMA_1955.docx"
)


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, size: Pt = Pt(10), bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.size = size
    run.bold = bold
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color


def add_heading_styled(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: Pt = Pt(10)) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_font(run, size=Pt(10))


def add_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = True
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, bold=True, size=Pt(10))
        r1 = c1.paragraphs[0].add_run(value)
        set_run_font(r1, size=Pt(10))
        shade_cell(c0, "D6E3F0")
    doc.add_paragraph()


def add_question_table(
    doc: Document,
    questions: list[tuple[str, str, str, str]],
) -> None:
    """questions: (Q#, NFR ID, Question, Guidance / options)."""
    headers = ["Q#", "NFR ID", "Question for client", "Guidance / options", "Client response", "Owner", "Due"]
    table = doc.add_table(rows=1 + len(questions), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=Pt(8), color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "1F4E79")

    for r, (qid, nfr, question, guidance) in enumerate(questions, start=1):
        values = [qid, nfr, question, guidance, "", "", ""]
        for c, val in enumerate(values):
            cell = table.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=Pt(8), bold=(c <= 1))
            if r % 2 == 0:
                shade_cell(cell, "F2F2F2")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Non-Functional Requirements (NFR) Questionnaire\n"
        "Marriage Registration — Hindu Marriage (HMA 1955)\n"
        "Kaveri 3.0"
    )
    set_run_font(run, size=Pt(16), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    add_para(
        doc,
        "Purpose: Capture client / department decisions for BRD §13 so TBD targets can be "
        "confirmed before architecture and sizing. Map each answer back to the NFR ID shown.",
        italic=True,
    )

    add_heading_styled(doc, "1. Document control", 1)
    add_meta_table(
        doc,
        [
            ("Document ID", "QNR-K3-MRG-HMA-NFR-001"),
            ("Related BRD", "BRD-K3-MRG-HMA-001 (§13.1–13.11)"),
            ("Module", "Marriage Registration — Hindu Marriage (Online & Offline channels)"),
            ("Version", "0.1 (Draft for client workshop)"),
            ("Workshop date", "[Date]"),
            ("Facilitator (BA)", "[Name]"),
            ("Client / department respondents", "[Names / roles — PO, IGSR, SDC, Security, Ops, Legal]"),
            ("Status", "Draft / In progress / Completed"),
        ],
    )

    add_heading_styled(doc, "2. How to use this questionnaire", 1)
    add_bullet(doc, "Answer every question with a concrete target, “N/A”, or “Deferred” (with reason).")
    add_bullet(doc, "Prefer measurable values (%, seconds, hours, concurrency, TB) over qualitative wording.")
    add_bullet(doc, "Where Kaveri platform standards already exist, cite the platform NFR / SDC policy and confirm Marriage module inherits them.")
    add_bullet(doc, "After the workshop, update BRD §13 Target / measure columns and close NFR-OP-01…12.")
    add_bullet(doc, "Leave Owner and Due blank until a named respondent and date are agreed.")

    add_heading_styled(doc, "3. Scope reminder (for respondents)", 1)
    add_para(
        doc,
        "NFRs apply to citizen portal + SRO desk for Hindu Marriage Online (citizen eSign) and "
        "Hindu Marriage Offline (printout, appointment, DEO upload, two-stage SR verification), "
        "including payment, eSign/DSC, SMS, and related integrations.",
    )

    # --- Sections ---
    sections: list[tuple[str, str, list[tuple[str, str, str, str]]]] = [
        (
            "4. Availability",
            "BRD §13.1 — service uptime, maintenance, outage communication, HA.",
            [
                (
                    "A-01",
                    "NFR-HMA-AVA-001",
                    "What monthly / annual availability % is required for citizen portal and SRO desk (excluding approved maintenance)?",
                    "Options: 99.0% / 99.5% / 99.9% / Other: ___%. Confirm if Marriage follows platform SLA.",
                ),
                (
                    "A-02",
                    "NFR-HMA-AVA-002",
                    "What are the allowed planned maintenance windows (day, IST time band, max duration) and minimum citizen/SRO notice period?",
                    "e.g. Sunday 02:00–06:00 IST; notice ≥72 hours.",
                ),
                (
                    "A-03",
                    "NFR-HMA-AVA-003",
                    "How must unplanned outages be communicated (status page, SMS to SROs, email, WhatsApp group, helpdesk)? Who owns the RACI?",
                    "Name channels + primary owner.",
                ),
                (
                    "A-04",
                    "NFR-HMA-AVA-004",
                    "Is high availability required as active-active or active-passive for app, API and database? Any Karnataka SDC mandate?",
                    "Cite SDC / hosting standard if known.",
                ),
                (
                    "A-05",
                    "NFR-OP-01",
                    "Is there a separate availability target for peak seasons (e.g. auspicious marriage months) vs normal months?",
                    "Yes (specify) / No — same target year-round.",
                ),
            ],
        ),
        (
            "5. Performance",
            "BRD §13.2 — latency, concurrency, volumes, eSign/print/DEO/appointment.",
            [
                (
                    "P-01",
                    "NFR-HMA-PERF-001",
                    "What p95 response-time targets (seconds) apply to: page browse, application submit, SRO scrutiny save, certificate PDF generation?",
                    "Suggested defaults to confirm/reject: browse ≤2s; submit ≤3s; PDF ≤5s.",
                ),
                (
                    "P-02",
                    "NFR-HMA-PERF-002",
                    "What peak concurrent users must the Marriage module support (citizen + SRO combined)? What share of overall Kaveri 3.0 concurrency is Marriage expected to consume?",
                    "Programme reference often ≥10,000 concurrent platform-wide — confirm module share.",
                ),
                (
                    "P-03",
                    "NFR-HMA-PERF-003",
                    "What are expected peak volumes for applications submitted, payments completed, and certificates issued (per hour and per day)? Provide Kaveri 2.0 baseline if available.",
                    "Attach MIS extract or approximate numbers.",
                ),
                (
                    "P-04",
                    "NFR-HMA-PERF-004",
                    "By which calendar day/time each month must Form III (monthly duplicate bundle to Registrar-General) complete?",
                    "e.g. by 05th of month 06:00 IST + buffer.",
                ),
                (
                    "P-05",
                    "NFR-HMA-PERF-005",
                    "Is a formal performance / load-test gate mandatory before go-live? Who signs off pass/fail criteria?",
                    "Yes / No; name Perf Lead / PO.",
                ),
                (
                    "P-06",
                    "NFR-HMA-PERF-006",
                    "What maximum eSign round-trip time and timeout/retry behaviour is acceptable when the eSign provider is slow or unavailable?",
                    "e.g. ≤30s target; graceful retry; offline channel fallback message.",
                ),
                (
                    "P-07",
                    "NFR-HMA-PERF-007",
                    "What maximum time is acceptable to generate bilingual (Kannada + English) printouts of Form I, II and 1A?",
                    "e.g. ≤10s.",
                ),
                (
                    "P-08",
                    "NFR-HMA-PERF-008",
                    "At peak counter hours, what DEO scan-upload throughput is required per office (files/hour) and what is the max file size / page count per upload?",
                    "e.g. files/hr; max MB; PDF/JPG allowed.",
                ),
                (
                    "P-09",
                    "NFR-HMA-PERF-009",
                    "What response time is required for appointment slot search and booking under contention, and is double-booking strictly prohibited?",
                    "e.g. ≤2s; no double-booking.",
                ),
                (
                    "P-10",
                    "NFR-OP-12",
                    "What service-level turnaround is expected for Online SR verification, Offline Stage-1 verification, and Offline Stage-2 verification (working hours / working days)?",
                    "Define SLA per stage.",
                ),
            ],
        ),
        (
            "6. Scalability",
            "BRD §13.3 — growth of users, registrations, documents, integrations.",
            [
                (
                    "S-01",
                    "NFR-HMA-SCALE-001",
                    "What is the expected growth in registered portal users / citizens for Y1, Y2, Y3 after go-live?",
                    "Absolute numbers or % YoY.",
                ),
                (
                    "S-02",
                    "NFR-HMA-SCALE-002",
                    "What is the expected volume of Hindu marriage registrations per year for Y1–Y3 (state-wide)? Any district-wise hotspots?",
                    "Use IGSR / historical volumes where possible.",
                ),
                (
                    "S-03",
                    "NFR-HMA-SCALE-003",
                    "What document/attachment storage growth is expected (photos, proofs, DEO scans) and for how many years must operational attachments be kept online vs archived?",
                    "GB/TB + retention years.",
                ),
                (
                    "S-04",
                    "NFR-HMA-SCALE-004",
                    "What peak TPS / daily caps are expected for payment, Aadhaar/eKYC, DigiLocker, SMS and eSign integrations for this module?",
                    "Per integration if known; else “follow platform”.",
                ),
                (
                    "S-05",
                    "NFR-HMA-SCALE-005",
                    "Must the architecture support horizontal scale-out of app/API and DB read replicas / partitioning as load grows, or is vertical scale sufficient for Phase 1?",
                    "Confirm with Architecture / SDC.",
                ),
            ],
        ),
        (
            "7. Security",
            "BRD §13.4 — AuthN/AuthZ, encryption, hardening, eSign/DSC, DEO separation.",
            [
                (
                    "SEC-01",
                    "NFR-HMA-SEC-001 / NFR-OP-04",
                    "How must citizens authenticate (portal login, OTP, eKYC/Aadhaar as approved)? How must SRO / DEO / admin authenticate (department IdP / SSO)? Is MFA mandatory for privileged roles?",
                    "Describe mechanisms + MFA yes/no by role.",
                ),
                (
                    "SEC-02",
                    "NFR-HMA-SEC-002",
                    "Confirm RBAC roles and jurisdiction scoping: Citizen, SRO, DEO, Admin, IGSR. Any additional roles?",
                    "Approve role matrix in workshop.",
                ),
                (
                    "SEC-03",
                    "NFR-HMA-SEC-003",
                    "Confirm encryption: TLS version minimum for transit; encryption-at-rest standard for PII, documents and certificates.",
                    "e.g. TLS 1.2+; SDC disk/DB encryption.",
                ),
                (
                    "SEC-04",
                    "NFR-HMA-SEC-004",
                    "Where must secrets (DB creds, API keys, certificates) be stored? Is a vault / SDC secret store mandatory?",
                    "Name standard tool.",
                ),
                (
                    "SEC-05",
                    "NFR-HMA-SEC-005",
                    "Which hardening baseline applies (MeitY / CERT-In / SDC): OS, containers, WAF, SSL, privileged access?",
                    "Cite policy documents.",
                ),
                (
                    "SEC-06",
                    "NFR-HMA-SEC-006",
                    "What vulnerability scan cadence and patch SLAs (Critical / High / Medium) are required? Are third-party / CERT-In / STQC audits mandatory before go-live?",
                    "Cadence + SLA days + audit path.",
                ),
                (
                    "SEC-07",
                    "NFR-HMA-SEC-007 / NFR-OP-06",
                    "What integrity features are required on Form II-A certificates (QR code, digital seal, hash, public verification URL)?",
                    "Select approach and verification UX.",
                ),
                (
                    "SEC-08",
                    "NFR-HMA-SEC-008",
                    "Confirm Aadhaar / eKYC usage scope for Marriage registration and UIDAI-compliant handling constraints.",
                    "In / out of scope; masking rules.",
                ),
                (
                    "SEC-09",
                    "NFR-HMA-SEC-009 / NFR-OP-09",
                    "Which eSign provider will be used? What integrity / verifiability / audit requirements apply? What is the fallback if eSign is down?",
                    "Provider + SLA + contingency.",
                ),
                (
                    "SEC-10",
                    "NFR-HMA-SEC-010 / NFR-OP-10",
                    "What is the DSC custody process for Sub-Registrars (issuance, storage, expiry monitoring, revocation on transfer/retirement)?",
                    "Describe Ops process owner.",
                ),
                (
                    "SEC-11",
                    "NFR-HMA-SEC-011",
                    "Confirm DEO must not be able to approve, register or digitally sign — upload and signature-check only. Any exception?",
                    "Yes enforce / exception: ___.",
                ),
                (
                    "SEC-12",
                    "NFR-HMA-SEC-012",
                    "Is antivirus / malware scanning mandatory for all DEO-uploaded scans before storage and SR viewing?",
                    "Yes / No; tool preference.",
                ),
            ],
        ),
        (
            "8. Privacy",
            "BRD §13.5 — PII inventory, masking, retention, non-prod data.",
            [
                (
                    "PR-01",
                    "NFR-HMA-PRIV-001",
                    "Please confirm / complete the PII inventory for Marriage: parties, witnesses, Aadhaar refs, photos, addresses, contact numbers, bank/payment refs. Any additional sensitive fields?",
                    "List extras or “complete as listed”.",
                ),
                (
                    "PR-02",
                    "NFR-HMA-PRIV-002",
                    "Which fields must be masked by default in UI, logs, support tools and non-prod (e.g. Aadhaar, mobile, email)?",
                    "Masking pattern (e.g. XXXX-XXXX-1234).",
                ),
                (
                    "PR-03",
                    "NFR-HMA-PRIV-003 / NFR-OP-05",
                    "Confirm: statutory Hindu Marriage registers are permanent (Rule 10(2)). What retention applies to operational logs, attachments, drafts, rejected applications and payment artefacts?",
                    "Years / purge rules / archive policy.",
                ),
                (
                    "PR-04",
                    "NFR-HMA-PRIV-004",
                    "Who may export bulk PII / certificates? What approval path is required for MIS extracts containing personal data?",
                    "Role + approval authority.",
                ),
                (
                    "PR-05",
                    "NFR-HMA-PRIV-005",
                    "May production PII be copied to UAT/training environments, or must non-prod use anonymized / synthetic data only?",
                    "Policy: anonymized only / approved exceptions.",
                ),
            ],
        ),
        (
            "9. Audit",
            "BRD §13.6 — audit events, evidence, retention, channel/eSign/DEO specifics.",
            [
                (
                    "AU-01",
                    "NFR-HMA-AUD-001",
                    "Confirm immutable audit is required for status changes, SRO scrutiny, approvals/rejections and fee events. Any additional business events?",
                    "Append-only / WORM expectation.",
                ),
                (
                    "AU-02",
                    "NFR-HMA-AUD-002",
                    "Must certificate issuance, reprint/duplicate and corrections each create a separate auditable event linked to Form II-A?",
                    "Yes / No.",
                ),
                (
                    "AU-03",
                    "NFR-HMA-AUD-003",
                    "What retention period applies to login, privilege-change and configuration-change audits for SRO and admin accounts?",
                    "Years / align to security policy.",
                ),
                (
                    "AU-04",
                    "NFR-HMA-AUD-004",
                    "In what format must audit evidence be available for departmental / AG / security audits (export, report, SIEM)?",
                    "Name preferred format / consumers.",
                ),
                (
                    "AU-05",
                    "NFR-HMA-AUD-005",
                    "What audit MIS cadence is required (daily/weekly/monthly) and who owns reporting?",
                    "Cadence + owner.",
                ),
                (
                    "AU-06",
                    "NFR-HMA-AUD-006–010",
                    "Confirm audit is mandatory for: channel selection/change; eSign & SR DSC; DEO check/upload versions; both Offline SR verification stages; appointment book/reschedule/cancel/no-show.",
                    "Confirm all / list exclusions.",
                ),
            ],
        ),
        (
            "10. Disaster recovery (DR)",
            "BRD §13.7 — RPO, RTO, topology, failover, backup.",
            [
                (
                    "DR-01",
                    "NFR-HMA-DR-001 / NFR-OP-03",
                    "What Recovery Point Objective (RPO) is required for Marriage data (applications, payments, registers, certificates, scans)?",
                    "e.g. ≤15 min / ≤1 hour / Other.",
                ),
                (
                    "DR-02",
                    "NFR-HMA-DR-002",
                    "What Recovery Time Objective (RTO) is required to restore citizen and SRO services after a major outage?",
                    "e.g. ≤4 h / ≤8 h / Other.",
                ),
                (
                    "DR-03",
                    "NFR-HMA-DR-003",
                    "What DR topology applies (primary + DR site, replication mode)? Does Marriage inherit Karnataka SDC DR design unchanged?",
                    "Cite SDC design if applicable.",
                ),
                (
                    "DR-04",
                    "NFR-HMA-DR-004",
                    "Should failover be automatic or manual? Who has authority to declare DR and initiate failover?",
                    "Auto / Manual + decision authority.",
                ),
                (
                    "DR-05",
                    "NFR-HMA-DR-005",
                    "What failback and data-reconciliation expectations apply after DR (especially in-flight payments, eSign and DEO uploads)?",
                    "Describe critical reconciliation points.",
                ),
                (
                    "DR-06",
                    "NFR-HMA-DR-006",
                    "What backup schedule and restore-test frequency are required? Must evidence of last successful restore be retained?",
                    "e.g. daily full + continuous WAL; quarterly restore drill.",
                ),
            ],
        ),
        (
            "11. Operations",
            "BRD §13.8 — support model, monitoring, incidents, ownership.",
            [
                (
                    "OP-01",
                    "NFR-HMA-OPS-001 / NFR-OP-07",
                    "What L1 / L2 / L3 support model, hours of cover (IST) and escalation matrix are required for citizen and SRO issues?",
                    "Hours, channels (phone/email/portal), SLAs.",
                ),
                (
                    "OP-02",
                    "NFR-HMA-OPS-002",
                    "Which components must be monitored in dashboards (app health, API latency/errors, DB, payment, eSign, SMS, certificate jobs, appointment)?",
                    "Confirm list + add extras.",
                ),
                (
                    "OP-03",
                    "NFR-HMA-OPS-003",
                    "What alert severity matrix and acknowledgment SLAs apply (Critical / High / Medium / Low)?",
                    "Ack times per severity.",
                ),
                (
                    "OP-04",
                    "NFR-HMA-OPS-004",
                    "Confirm incident response expectations: classify, contain, communicate, RCA, post-incident review. Any department-specific ITIL deviations?",
                    "Yes follow ITIL / deviations: ___.",
                ),
                (
                    "OP-05",
                    "NFR-HMA-OPS-005",
                    "Name the service owner, application owner, infrastructure owner and data owner for Marriage module.",
                    "Four named roles/orgs.",
                ),
                (
                    "OP-06",
                    "NFR-HMA-OPS-006",
                    "Are runbooks mandatory before go-live for submit, pay, register, certificate issue, DEO upload, appointment and restore?",
                    "Yes (pack complete) / prioritized subset.",
                ),
            ],
        ),
        (
            "12. Capacity",
            "BRD §13.9 — compute, memory, storage, DB, network sizing.",
            [
                (
                    "CA-01",
                    "NFR-HMA-CAP-001 / NFR-OP-08",
                    "What compute sizing assumptions (vCPU / nodes / workers) should Architecture use for peak + headroom at go-live?",
                    "Provide numbers or “SDC to size from Perf targets”.",
                ),
                (
                    "CA-02",
                    "NFR-HMA-CAP-002",
                    "What memory sizing assumptions (GB per tier) apply?",
                    "App / API / DB.",
                ),
                (
                    "CA-03",
                    "NFR-HMA-CAP-003 / NFR-OP-11",
                    "What storage capacity (DB + object store for documents/scans) is required for Y1–Y3, including Offline DEO scan growth?",
                    "TB + growth %.",
                ),
                (
                    "CA-04",
                    "NFR-HMA-CAP-004",
                    "Any database constraints to plan for (max connections, IOPS, HA/replica footprint)?",
                    "Numbers or SDC defaults.",
                ),
                (
                    "CA-05",
                    "NFR-HMA-CAP-005",
                    "Any network bandwidth / latency requirements between SRO offices, SDC, DR and integration endpoints?",
                    "Especially for scan upload peak hours.",
                ),
                (
                    "CA-06",
                    "NFR-HMA-CAP-006",
                    "How often must capacity be reviewed (quarterly / before peak seasons)? Who chairs the review?",
                    "Cadence + chair.",
                ),
            ],
        ),
        (
            "13. Compliance",
            "BRD §13.10 — GIGW, WCAG, MeitY/CERT-In, STQC, localization, statutory forms.",
            [
                (
                    "CO-01",
                    "NFR-HMA-COMP-001",
                    "Is GIGW compliance mandatory for all citizen-facing Marriage screens? Who signs the checklist?",
                    "Yes / No; sign-off role.",
                ),
                (
                    "CO-02",
                    "NFR-HMA-COMP-002",
                    "Which WCAG level is required (A / AA / AAA) and for which channels (citizen portal, SRO desk, PDFs)?",
                    "Preferred: AA for citizen portal.",
                ),
                (
                    "CO-03",
                    "NFR-HMA-COMP-003",
                    "Confirm MeitY / CERT-In security guidelines and advisories apply; any additional Karnataka e-Gov norms?",
                    "List extra norms.",
                ),
                (
                    "CO-04",
                    "NFR-HMA-COMP-004",
                    "Is STQC / hosting / security clearance required before go-live? What is the clearance path and timeline?",
                    "Path + responsible party.",
                ),
                (
                    "CO-05",
                    "NFR-HMA-COMP-005",
                    "Confirm Aadhaar/UIDAI and Karnataka hosting/security norms that must be evidenced at go-live.",
                    "Evidence artefacts expected.",
                ),
                (
                    "CO-06",
                    "NFR-HMA-COMP-006",
                    "Confirm bilingual requirement: Kannada + English UI and Kannada fonts on screen and on PDF certificates/printouts. Any other languages?",
                    "Languages list.",
                ),
                (
                    "CO-07",
                    "NFR-HMA-COMP-007",
                    "Confirm: statutory form wording (Form I, IA, II, II-A, III, VI) must not change without Legal / Domain Expert approval.",
                    "Yes / process for exceptions.",
                ),
                (
                    "CO-08",
                    "NFR-HMA-COMP-008",
                    "Which government records / archival policies must registers and audit evidence align to?",
                    "Cite policy names.",
                ),
            ],
        ),
    ]

    for heading, intro, questions in sections:
        add_heading_styled(doc, heading, 1)
        add_para(doc, intro, italic=True, size=Pt(9))
        add_question_table(doc, questions)

    add_heading_styled(doc, "14. Priority open decisions (from BRD §13.11)", 1)
    add_para(
        doc,
        "Use this shortlist if workshop time is limited — these map to NFR-OP-01…12.",
        italic=True,
    )
    priority = [
        ("OD-01", "NFR-OP-01", "Confirm availability % and maintenance window with SDC / Ops.", "Must decide before SLA wording in BRD."),
        ("OD-02", "NFR-OP-02", "Confirm p95 latency, concurrency and peak TPS from Kaveri 2.0 + growth.", "Drives Perf test & capacity."),
        ("OD-03", "NFR-OP-03", "Confirm RPO / RTO and whether SDC DR standard is inherited.", "Drives architecture."),
        ("OD-04", "NFR-OP-04", "Confirm citizen + SRO AuthN and MFA for privileged roles.", "Security design gate."),
        ("OD-05", "NFR-OP-05", "Confirm PII retention beyond permanent registers.", "Legal + DBA."),
        ("OD-06", "NFR-OP-06", "Confirm Form II-A QR / digital seal approach.", "Product + Security."),
        ("OD-07", "NFR-OP-07", "Confirm L1/L2/L3 hours and incident SLAs.", "Support model."),
        ("OD-08", "NFR-OP-08", "Confirm go-live capacity numbers.", "Infra worksheet."),
        ("OD-09", "NFR-OP-09", "Confirm eSign provider, SLA and fallback.", "Integration contingency."),
        ("OD-10", "NFR-OP-10", "Confirm DSC provisioning & expiry monitoring for all SRs.", "Ops process."),
        ("OD-11", "NFR-OP-11", "Confirm storage growth from Offline DEO scans.", "Capacity model."),
        ("OD-12", "NFR-OP-12", "Confirm SLA per verification stage (Online / Offline 1 / Offline 2).", "Service standard."),
    ]
    add_question_table(doc, priority)

    add_heading_styled(doc, "15. Workshop attendees & sign-off", 1)
    sign = doc.add_table(rows=6, cols=4)
    sign.style = "Table Grid"
    headers = ["Role", "Name", "Organisation", "Signature / Date"]
    for c, h in enumerate(headers):
        cell = sign.rows[0].cells[c]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=Pt(9), color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "1F4E79")
    roles = [
        "Product Owner",
        "Domain Expert / IGSR nominee",
        "Architecture / SDC",
        "Security",
        "Operations / Support",
    ]
    for i, role in enumerate(roles, start=1):
        sign.rows[i].cells[0].text = role
        for c in range(4):
            for p in sign.rows[i].cells[c].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=Pt(9))

    doc.add_paragraph()
    add_para(
        doc,
        "End of questionnaire — return completed responses to BA for update of BRD-K3-MRG-HMA-001 §13.",
        italic=True,
        size=Pt(9),
    )

    doc.save(DST)
    return DST


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
