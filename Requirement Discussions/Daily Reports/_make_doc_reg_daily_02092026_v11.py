# -*- coding: utf-8 -*-
"""Create Document_Registration_requirement_02092026_v1.1.docx from v1.0.

Section A (Rule 17 filing) currently merges Parts IV and V into one row and has
no Part VI row. This splits Part IV and Part V into their own rows and adds a
Part VI row, with a note on its status in the available source.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(
    r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Requirement Discussions\Daily Reports"
)
SRC = BASE / "Document_Registration_requirement_02092026.docx"
DST = BASE / "Document_Registration_requirement_02092026_v1.1.docx"

PART_ROWS = [
    [
        "Rule 17(i) Part IV",
        "Institutional / revenue filing",
        "Copies of instruments and collateral securities executed under the Karnataka "
        "Land Improvement Loans Act, 1963 (Karnataka Act 16 of 1963) and the Karnataka "
        "Agriculturists Loans Act, 1963 (Karnataka Act 17 of 1963), received from "
        "Revenue officers",
    ],
    [
        "Rule 17(i) Part V",
        "Institutional / bank filing",
        "Copies of instruments received from Land Development Banks under Sec. 85-A of "
        "the Karnataka Co-operative Societies Act, 1959",
    ],
    [
        "Rule 17(i) Part VI",
        "Institutional filing — to be confirmed",
        "Referenced in departmental practice (ServiceDesk 16176 “Entry vide Rule 17 "
        "Part-VI”). Not present in the copy of the 1965 Rules held in "
        "Acts_Rules/Document, which lists Parts I–V only. Governing amendment / "
        "notification and the category of instruments to be confirmed with AIGR "
        "Computers before build.",
    ],
]

NOTE_PREFIX = "Note:"
NOTE_TEXT = (
    " Rule 17(i) in the available copy of the Karnataka Registration Rules, 1965 "
    "(Acts_Rules/Document) lists Parts I to V, with Parts IV and V substituted by the "
    "Karnataka Registration (Amendment) Rules, 1971. Part VI is referenced in "
    "departmental usage but is not in that copy, so its enabling amendment and scope "
    "need confirmation before the filing module is designed."
)


def set_cell_text(cell, text: str) -> None:
    """Replace cell text, keeping the formatting of the first run."""
    paras = cell.paragraphs
    first = paras[0]
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        first.add_run(text)


def row_index_containing(table: Table, needle: str) -> int:
    for i, row in enumerate(table.rows):
        if needle in row.cells[0].text:
            return i
    raise KeyError(f"Row not found: {needle!r}")


def replace_row_with(table: Table, index: int, contents: list[list[str]]) -> None:
    """Replace one row with several rows, reusing its formatting."""
    original = table.rows[index]._tr
    new_trs = []
    for _ in contents:
        new_tr = deepcopy(original)
        original.addnext(new_tr)
        new_trs.insert(0, new_tr)
    for new_tr, values in zip(new_trs, contents):
        row = next(r for r in table.rows if r._tr is new_tr)
        for ci, value in enumerate(values):
            set_cell_text(row.cells[ci], value)
    original.getparent().remove(original)


def clone_note_after(template: Paragraph, anchor_element, prefix: str, text: str) -> None:
    """Insert a note paragraph after `anchor_element`, styled like `template`."""
    new_p = deepcopy(template._p)
    anchor_element.addnext(new_p)
    para = Paragraph(new_p, template._parent)
    runs = para.runs
    for run in runs[2:]:
        run._element.getparent().remove(run._element)
    runs[0].text = prefix
    runs[1].text = text


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)

    doc = Document(str(DST))

    meta = doc.tables[0]
    meta._tbl.append(deepcopy(meta.rows[-1]._tr))
    set_cell_text(meta.rows[-1].cells[0], "Version")
    set_cell_text(
        meta.rows[-1].cells[1],
        "1.1 (04-09-2026) — Rule 17(i) Part IV and Part V listed separately; "
        "Part VI added in section A",
    )

    rule17 = doc.tables[2]
    idx = row_index_containing(rule17, "Parts IV–V")
    replace_row_with(rule17, idx, PART_ROWS)

    note_template = next(
        p
        for p in doc.paragraphs
        if p.text.strip().startswith("Note:") and "Explicit" in p.text
    )
    clone_note_after(note_template, rule17._tbl, NOTE_PREFIX, NOTE_TEXT)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    check = Document(str(DST))
    for row in check.tables[2].rows:
        print(" |", row.cells[0].text.strip(), "|", row.cells[1].text.strip()[:44])


if __name__ == "__main__":
    main()
