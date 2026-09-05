# -*- coding: utf-8 -*-
"""Append Acts/Rules/Notifications and ServiceDesk pain points to:

  Document_Registration_requirement_27082026 (1).docx
  Document_Registration_requirement_28082026.docx

Discussion topic (both): Registration core — Registration, Appointment,
Status tracking. Body notes on 28th also touch stamp duty / guidance value
and undervaluation & impound.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(
    r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Requirement Discussions\Daily Reports"
)

HEADING_FONT = "Segoe UI"
HEADING_SIZE_H3 = Pt(14.5)
HEADING_SIZE_H4 = Pt(13.5)
PAIN_TITLE_SIZE = Pt(24)


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    first = paras[0]
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        first.add_run(text)


def clone_table_after(paragraph, doc: Document, template: Table,
                      headers: list[str], rows: list[list[str]]) -> Table:
    """Insert a copy of `template` immediately after `paragraph`."""
    new_tbl = deepcopy(template._tbl)
    paragraph._p.addnext(new_tbl)
    table = Table(new_tbl, paragraph._parent)

    data_tr = table.rows[1]._tr if len(table.rows) > 1 else table.rows[0]._tr
    for row in list(table.rows)[1:]:
        new_tbl.remove(row._tr)

    for i, head in enumerate(headers):
        if i < len(table.rows[0].cells):
            set_cell_text(table.rows[0].cells[i], head)

    template_tr = deepcopy(data_tr)
    for _ in range(len(rows)):
        new_tbl.append(deepcopy(template_tr))

    for ri, values in enumerate(rows, start=1):
        for ci, value in enumerate(values):
            if ci < len(table.rows[ri].cells):
                set_cell_text(table.rows[ri].cells[ci], value)
    return table


def add_heading(doc: Document, text: str, size=HEADING_SIZE_H3):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = HEADING_FONT
    run.font.size = size
    return para


def add_pain_title(doc: Document, text: str = "Pain Points"):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = PAIN_TITLE_SIZE
    return para


def add_note(doc: Document, bold_prefix: str, rest: str):
    para = doc.add_paragraph()
    r1 = para.add_run(bold_prefix)
    r1.bold = True
    r1.font.name = HEADING_FONT
    r1.font.size = HEADING_SIZE_H4
    r2 = para.add_run(rest)
    r2.font.name = HEADING_FONT
    r2.font.size = HEADING_SIZE_H4
    return para


def make_3col_template(doc: Document) -> Table:
    """Build a detached 3-col styled template table (not left in the document body)."""
    style_name = doc.tables[0].style.name if doc.tables and doc.tables[0].style else None
    tbl = doc.add_table(rows=2, cols=3)
    if style_name:
        try:
            tbl.style = style_name
        except KeyError:
            pass
    for i, h in enumerate(["A", "B", "C"]):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.makeelement(
            qn("w:shd"),
            {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "D9E2F3",
            },
        )
        tc_pr.append(shd)
    for i, v in enumerate(["x", "y", "z"]):
        cell = tbl.rows[1].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(v)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
    # Detach from body; clone_table will re-append deep copies.
    tbl._tbl.getparent().remove(tbl._tbl)
    return tbl


def add_version_row(doc: Document, version_text: str) -> None:
    meta = doc.tables[0]
    meta._tbl.append(deepcopy(meta.rows[-1]._tr))
    set_cell_text(meta.rows[-1].cells[0], "Version")
    set_cell_text(meta.rows[-1].cells[1], version_text)


# ---------------------------------------------------------------------------
# Shared legal content for Registration / Appointment / Status tracking
# ---------------------------------------------------------------------------

ACTS_ROWS = [
    [
        "The Registration Act, 1908 (Central Act 16 of 1908)",
        "Secs. 17–22, 23, 25, 28–35, 51–52, 58–61, 71–77, 78",
        "Compulsory / optional registration; presentation time & place; examination; "
        "endorsements & Sec. 60 certificate; refusal / appeal; registration fees — "
        "core Registration and Status tracking",
    ],
    [
        "The Registration Act, 1908",
        "Secs. 5–7, 10–12, 28–31",
        "Districts / sub-districts; Registrars & Sub-Registrars; place of registration "
        "— Appointment office selection and jurisdiction routing",
    ],
    [
        "The Registration (Karnataka Amendment) Act, 2023 (Karnataka Act 47 of 2024)",
        "Secs. 22-B, 22-C, 22-D, 81-A, 81-B",
        "Refusal / cancellation of forged or prohibited documents; appeal; penalties "
        "— registration intake validation",
    ],
    [
        "The Karnataka Stamp Act, 1957",
        "Secs. 3, 10, 33–39, 45-A, Schedule",
        "Stamp duty chargeable; payment modes; impounding; undervaluation reference — "
        "touched in 28-08 discussion (stamp duty / guidance value / impound)",
    ],
    [
        "The Transfer of Property Act, 1882",
        "Sec. 54",
        "Sale of immovable property of value > ₹100 requires registered instrument — "
        "drives Article / transaction-type selection for conveyance",
    ],
]

RULES_ROWS = [
    [
        "Karnataka Registration Rules, 1965 — Ch. II",
        "Rules 3, 5",
        "Office hours and holidays — Appointment / slot calendar",
    ],
    [
        "Karnataka Registration Rules, 1965 — Ch. VI",
        "Rules 13–15",
        "Territorial divisions; survey / city survey description — Village Index and "
        "property details (KSRSAC / GIS)",
    ],
    [
        "Karnataka Registration Rules, 1965 — Ch. IX",
        "Rules 37, 40–46, 51–55",
        "Office where document may be registered; presentation; examination; "
        "endorsement; suspension (fine / stamp); condonation — Registration & Status",
    ],
    [
        "Karnataka Registration Rules, 1965 — Ch. XII / XVI",
        "Rules 71–73, 78–79, 94, 104",
        "Executant examination; thumb impression / photograph; endorsements; "
        "Sec. 60 certificate — Registration completion",
    ],
    [
        "Karnataka Registration Rules, 1965 — Ch. XVII / XXV",
        "Rules 110–118, 175–188",
        "Receipts; return of document; appeal against refusal — Status tracking termini",
    ],
    [
        "Karnataka Stamp Rules, 1958 / e-Stamping Rules, 2009",
        "Stamp payment procedures",
        "Impressed / franking / e-Stamp at presentation — Payment visibility (28-08)",
    ],
]

NOTIF_ROWS = [
    [
        "RGN 2/2002-03 (1 Apr 2002; w.e.f. 4 Apr 2002)",
        "Rule 19-A document sheets; Rules 22-A–22-C; Rule 40 photograph / digital photo",
        "Document format and photo at presentation",
    ],
    [
        "RD 403 ESR 85 (27 May 1986) as amended by RD/46/MNMU/2025 (29 Aug 2025)",
        "Table of Registration Fees under Sec. 78 — fee revision w.e.f. 31 Aug 2025",
        "Registration fee calculator / payment amount",
    ],
    [
        "RD 380 MUNOMU 2008 (8 Apr 2009)",
        "Karnataka Stamp (Payment of Duty by Means of e-Stamping) Rules, 2009",
        "e-Stamp payment and verification at SRO",
    ],
    [
        "Stamp Act Sec. 45-B (Act 8 of 2003) / Undervaluation Rules, 1977 (GSR 81)",
        "CVC market value guidelines; Sec. 45-A reference / Form I",
        "Guidance value & undervaluation / impound (noted on 28-08)",
    ],
]

PAIN_27 = [
    [
        "Article / sub-article selection & instrument mapping",
        "Release Deed Article Issue (17714); Stamp Duty wrong for partition sub-article "
        "39-b (31257); Gift Deed stamp duty double (93306)",
    ],
    [
        "Village / Index / road master gaps",
        "Village index absent / split villages (13949); road names wrong (95135); "
        "road option not showing for agriculture land (31819); Sy. No. not fetching "
        "from Bhoomi (20170, 30045); wrong village shown (29600)",
    ],
    [
        "Property schedule / boundary capture fails",
        "Unable to save the property schedule (24450); Unable to Save and Continue "
        "after 11E Sketch (24247); property details not in dept summary (95367)",
    ],
    [
        "Party information save / display errors",
        "Unable to save party info (30276); claimant details missing in summary "
        "(25784, 28927, 29588); executant name showing None (22789); claimant photo "
        "at executant place (27709)",
    ],
    [
        "Property valuation / fee after valuation",
        "Market valuation blank / not showing (29357, 29532, 27923); valuate → fee "
        "zero & Save disabled (28596); fee zero in summary (26292)",
    ],
    [
        "Pre-submission summary / review broken",
        "Unable to generate / view document summary (27360, 27805, 30609, 12009); "
        "consideration / market value wrong in summary (6187, 10531); sub-article "
        "name not in dept summary (92782); challan not in summary (31168, 29793)",
    ],
    [
        "Appointment / schedule after intake",
        "Unable to Schedule for Appointment (27283, 30478, 30908); schedule not "
        "reflecting (28240); schedule option missing after payment (Categorized: "
        "27283, 30478, 28812, 30908)",
    ],
    [
        "Status / step stuck after registration flow",
        "Check and Register (Step 7) disabled (94299, 94264, 92559); applications "
        "stuck in Step 5 (23243, 25183); pending shows in Step 1 instead of Minute "
        "Book (22599, 31762)",
    ],
]

PAIN_28 = [
    [
        "SR send to payment / send back for correction",
        "After verification unable to send to citizen for payment / rectify "
        "(94523, 92990, 92566, 88608, 83818, 31346 — pending count 6); after "
        "evaluation send does not proceed (30960, 31484); sent-back apps still ask "
        "for payment with amount 0 (88309, 27587)",
    ],
    [
        "Payment visibility & payment failures",
        "Unable to make payment for document registration (29349, 29407); payment "
        "error “Something went wrong” (94275); payment completed but still Make "
        "Payment / pending for payment (92384, 28335, 26151); payment details not "
        "in 10A (27824)",
    ],
    [
        "Slot booking / tatkal / reschedule",
        "Unable to schedule appointment (27283 P1, 30478, 30908); time slot issue "
        "(30841); unable to re-schedule (91929); after payment still pending for "
        "payment while citizen books slot (28335)",
    ],
    [
        "Withdrawal workflow",
        "Withdraw option not available for rejected application (93043); citizen "
        "unable to withdraw (24685, 23807); withdrawn app still pending in DEO "
        "(9494, 23080); withdraw auto-generated registration number (8691); "
        "payment withdraw issue (28630)",
    ],
    [
        "Stamp duty & registration fee calculation",
        "Unable to Calculate Stamp duty and Registration fee (23369); fees "
        "calculation category pending (74 tickets); partition / gift article "
        "wrong duty (31257, 93306, 31130); fee zero after valuation (28596, 26292)",
    ],
    [
        "Guidance value / undervaluation & impound",
        "Pending for 45-A undervaluation (93336… — count 6); unable to release "
        "45-A pending (13204, 14582, 31209–31212); Impound related issue (30024); "
        "CVC rates not displaying (31481, 31521); rural building valued at "
        "TMC/CMC rates (29393)",
    ],
    [
        "Summary / document number display",
        "Summary not generating (27360, 27805, 30609); document number generated "
        "but still in DEO Step-1 Generate Summary (28388); document number / "
        "challan / consideration missing or wrong in summary (31168, 6187, 10531)",
    ],
    [
        "Status tracking — wrong step / DSC / acknowledgement",
        "Digital Sign option / SR name missing at Step 9 (95322, 27050, 95349); "
        "Unable to Print acknowledgement at Step 10 (88602); slot allocation — "
        "registration complete but still in DEO Step 3 (28586); app number / "
        "status inconsistent across logins (Categorized Payment & Step issues)",
    ],
]


def append_common_legal(doc: Document, template: Table) -> None:
    doc.add_paragraph()
    h1 = add_heading(doc, "1. Primary Acts")
    clone_table_after(h1, doc, template, ["Act", "Sections", "Relevance"], ACTS_ROWS)

    doc.add_paragraph()
    h2 = add_heading(doc, "2. Primary Rules — Karnataka Registration Rules, 1965")
    clone_table_after(
        h2, doc, template, ["Rule", "Key provisions", "Relevance to topic"], RULES_ROWS
    )

    doc.add_paragraph()
    h3 = add_heading(doc, "3. Notifications / amendments")
    clone_table_after(
        h3, doc, template, ["Instrument", "Effect", "Relevance"], NOTIF_ROWS
    )


def append_pain_points(doc: Document, template: Table, title: str,
                       rows: list[list[str]], note: str | None = None) -> None:
    doc.add_paragraph()
    add_pain_title(doc)
    doc.add_paragraph()
    add_heading(doc, title)
    # 2-col pain table — build from template by using first 2 cols of a clone
    # Simpler: create dedicated 2-col table
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    style_name = doc.tables[0].style.name if doc.tables and doc.tables[0].style else None
    if style_name:
        try:
            tbl.style = style_name
        except KeyError:
            pass
    headers = ["Pain point", "Evidence (tickets / categorized)"]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.makeelement(
            qn("w:shd"),
            {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "D9E2F3"},
        )
        tc_pr.append(shd)
    for ri, values in enumerate(rows, start=1):
        for ci, value in enumerate(values):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
    if note:
        doc.add_paragraph()
        add_note(doc, "Note:", note)


def build_27() -> Path:
    src = BASE / "Document_Registration_requirement_27082026 (1).docx"
    dst = BASE / "Document_Registration_requirement_27082026_v1.1.docx"
    shutil.copy2(src, dst)
    doc = Document(str(dst))
    add_version_row(
        doc,
        "1.1 (04-09-2026) — Acts, Rules, notifications and ServiceDesk pain points "
        "added for Registration, Appointment, Status tracking",
    )
    template = make_3col_template(doc)
    # Remove seed template from body after we have a handle — actually keep it
    # out of final doc by removing after clones. First append content using it.
    append_common_legal(doc, template)
    append_pain_points(
        doc,
        template,
        "Registration core — Registration, Appointment, Status tracking "
        "(aligned to 27-08 discussion notes)",
        PAIN_27,
        " Source: Requirement Discussions/ServiceDesk Issues/ServiceDeskIssuesList.xlsx "
        "(OverallList + Categorized). Topics map to Schedule Sr.12 sub-modules "
        "#1 Registration, #2 Appointment, #19 Status tracking; village / valuation "
        "notes from the 27-08 walkthrough are included where tickets exist.",
    )
    doc.save(str(dst))
    return dst


def build_28() -> Path:
    src = BASE / "Document_Registration_requirement_28082026.docx"
    dst = BASE / "Document_Registration_requirement_28082026_v1.1.docx"
    shutil.copy2(src, dst)
    doc = Document(str(dst))
    add_version_row(
        doc,
        "1.1 (04-09-2026) — Acts, Rules, notifications and ServiceDesk pain points "
        "added for Registration, Appointment, Status tracking (incl. 28-08 notes on "
        "SR workflow, slot, payment, stamp / guidance value, impound)",
    )
    template = make_3col_template(doc)

    # Extra acts row emphasis for 28th stamp/guidance
    append_common_legal(doc, template)

    doc.add_paragraph()
    add_heading(doc, "4. Cross-reference (noted in 28-08 discussion)", HEADING_SIZE_H4)
    add_note(
        doc,
        "Stamp duty / Guidance value / Impound: ",
        "Detailed study of the Karnataka Stamp Act Schedule and CVC / Undervaluation "
        "Rules was flagged in the meeting (items 2, 9). Full legal tables for those "
        "topics are also in Document_Registration_requirement_01092026_v1.1.docx "
        "(Sr.13–14). Key hooks retained above: Stamp Act Secs. 3, 10, 33–39, 45-A; "
        "Sec. 45-B / Undervaluation Rules, 1977; Registration Act Sec. 78 + "
        "RD/46/MNMU/2025 fees.",
    )

    append_pain_points(
        doc,
        template,
        "Registration core — Registration, Appointment, Status tracking "
        "(aligned to 28-08 discussion notes)",
        PAIN_28,
        " Source: Requirement Discussions/ServiceDesk Issues/ServiceDeskIssuesList.xlsx "
        "(OverallList + Categorized). Includes SR approval → payment / correction, "
        "slot booking, withdrawal, stamp & fee calculation, 45-A / impound, and "
        "summary / status-step issues discussed on 28-08-2026.",
    )
    doc.save(str(dst))
    return dst


def main() -> None:
    d27 = build_27()
    d28 = build_28()
    print(f"Wrote {d27}")
    print(f"Wrote {d28}")
    for path in (d27, d28):
        doc = Document(str(path))
        print(path.name, "tables=", len(doc.tables), "paras=", len(doc.paragraphs))


if __name__ == "__main__":
    main()
