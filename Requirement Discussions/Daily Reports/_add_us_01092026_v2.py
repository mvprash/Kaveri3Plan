# -*- coding: utf-8 -*-
"""Create Document_Registration_requirement_01092026_v2.docx with user stories."""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Requirement Discussions\Daily Reports")
SRC = BASE / "Document_Registration_requirement_01092026.docx"
DST = BASE / "Document_Registration_requirement_01092026_v2.docx"


def add_para(doc: Document, text: str, style: str = "Normal", bold: bool = False) -> None:
    p = doc.add_paragraph()
    if style in ("Heading 3", "Normal", "List Paragraph", "ui-markdown__paragraph"):
        p.style = style
    else:
        p.style = "Normal"
    run = p.add_run(text)
    run.bold = bold


def add_heading3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def add_label(doc: Document, text: str) -> None:
    add_para(doc, text, "Normal", bold=True)


def add_story(doc: Document, sid: str, actor: str, want: str, so_that: str) -> None:
    text = f"{sid}\nAs a {actor},\nI want to {want},\nso that {so_that}."
    add_para(doc, text, "Normal")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Header note
    add_para(doc, "")
    add_heading3(doc, "User Stories")
    add_para(
        doc,
        "User stories mapped to the Acts, Rules, notifications and GIS notes listed "
        "in this discussion note (01-09-2026).",
    )

    # 1. Guideline value calculation
    add_heading3(doc, "1. Guideline value calculation")

    add_label(doc, "Karnataka Stamp Act, 1957 — Sec. 45-A")
    add_story(
        doc,
        "US-GV-01",
        "Sub-Registrar",
        "compare the consideration in the instrument with the market value guidelines published under Sec. 45-B",
        "if the property is undervalued I can estimate the value, collect the extra duty, or refer the case to the Deputy Commissioner",
    )

    add_label(doc, "Karnataka Stamp Act, 1957 — Sec. 45-B")
    add_story(
        doc,
        "US-GV-02",
        "citizen / Sub-Registrar",
        "use the official market value guidelines published and revised under Sec. 45-B in guideline value calculation",
        "stamp duty is computed on the current notified guideline rates and not on an unofficial figure",
    )

    add_label(doc, "Karnataka Stamp Act, 1957 — Sec. 3 + Schedule")
    add_story(
        doc,
        "US-GV-03",
        "citizen / Sub-Registrar",
        "calculate stamp duty ad valorem on market value for Schedule instruments (conveyance, gift, exchange and other listed articles)",
        "the correct Schedule article and rate are applied after guideline value is arrived at",
    )

    add_label(doc, "Karnataka Stamp (Prevention of Undervaluation of Instruments) Rules, 1977")
    add_story(
        doc,
        "US-GV-04",
        "Sub-Registrar",
        "capture property particulars and market value in Form I when starting a Sec. 45-A reference",
        "the undervaluation case follows the notified procedure and has a complete property statement",
    )

    add_label(doc, "Registration Rules 13–15 (supporting)")
    add_story(
        doc,
        "US-GV-05",
        "citizen / Sub-Registrar",
        "enter survey number, territorial division and property description as required by Registration Rules 13–15",
        "the system can pick the correct guideline rate for that property",
    )

    # 2. Valuation Data Entry Module (CVC)
    add_heading3(doc, "2. Valuation Data Entry Module (CVC)")

    add_label(doc, "Karnataka Stamp Act, 1957 — Sec. 45-B (primary)")
    add_story(
        doc,
        "US-CVC-01",
        "an IGR / Central Valuation Committee",
        "estimate, publish and revise market value guidelines and constitute district / sub-district market valuation sub-committees",
        "CVC remains the final authority for policy, methodology and administration of guideline rates in the State",
    )
    add_story(
        doc,
        "US-CVC-02",
        "CVC / valuation data-entry officer",
        "enter, update and publish guideline rates in the Valuation Module for use by SR and citizen logins",
        "rates are available for guideline value calculation and do not go missing in office or citizen screens",
    )

    add_label(doc, "Karnataka Stamp Act, 1957 — Sec. 2(ac)")
    add_story(
        doc,
        "US-CVC-03",
        "system administrator / department user",
        "maintain Central Valuation Committee as the statutory body defined in Sec. 2(ac)",
        "valuation masters and approvals are owned by the correct legal entity",
    )

    add_label(doc, "Karnataka Stamp Act, 1957 — Sec. 45-A")
    add_story(
        doc,
        "US-CVC-04",
        "Sub-Registrar",
        "apply the CVC-published guidelines at the time of registration for stamp duty",
        "duty is charged on the higher of consideration and guideline value where Sec. 45-A applies",
    )
    add_story(
        doc,
        "US-CVC-05",
        "citizen / Sub-Registrar",
        "see and use the District Registrar’s adjudicated Sec. 45-A value instead of an automatically higher guideline figure",
        "the summary and payment match the legally decided market value",
    )

    add_label(doc, "Prevention of Undervaluation Rules, 1977 (GSR 81 / RD 73 EST 74)")
    add_story(
        doc,
        "US-CVC-06",
        "Sub-Registrar / Deputy Commissioner",
        "run Form I capture, DC determination of market value, and the appeal path under the 1977 Rules",
        "a Sec. 45-A undervaluation case can be completed, paid and released without a data or payment stuck state",
    )

    add_label(doc, "Notification / amendment — Act 8 of 2003 (w.e.f. 1-4-2003)")
    add_story(
        doc,
        "US-CVC-07",
        "CVC administrator",
        "operate the Valuation Module on the Sec. 45-B framework as strengthened by Act 8 of 2003",
        "guideline publication, revision cycles and sub-committee working match the current CVC law",
    )

    add_label(doc, "Amendment — RD 264 MUNOMU 99 (18-8-1999)")
    add_story(
        doc,
        "US-CVC-08",
        "Deputy Commissioner / District Registrar",
        "follow the undervaluation workflow as amended by RD 264 MUNOMU 99 (no obsolete provisional / Rule 6 path)",
        "orders and screens match the Rules now in force",
    )

    # 3. GIS valuation
    add_heading3(doc, "3. GIS valuation")
    add_para(
        doc,
        "GIS valuation is an implementation layer on top of CVC guideline rates, "
        "to be integrated with KSRSAC (Karnataka State Remote Sensing Department).",
    )

    add_label(doc, "GIS + CVC guideline rates (KSRSAC integration)")
    add_story(
        doc,
        "US-GIS-01",
        "citizen / Sub-Registrar",
        "locate the property on GIS (KSRSAC) and apply the matching CVC guideline rate for that area",
        "valuation uses the correct spatial rate and not a neighbouring village, road or urban slab",
    )

    add_label(doc, "Registration Rules 13–15 / survey description")
    add_story(
        doc,
        "US-GIS-02",
        "citizen / Sub-Registrar",
        "map survey number, Pot Hissa / city survey and territorial division to the GIS layer",
        "guideline value can be calculated even for small extents (for example 1 gunta) with the correct rural or urban building rate",
    )

    add_story(
        doc,
        "US-GIS-03",
        "valuation / GIS administrator",
        "keep village, road and survey masters aligned between Kaveri, CVC rates and KSRSAC",
        "wrong road names, missing village splits and failed survey fetch do not produce a blank or wrong market value",
    )

    doc.save(str(DST))
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
