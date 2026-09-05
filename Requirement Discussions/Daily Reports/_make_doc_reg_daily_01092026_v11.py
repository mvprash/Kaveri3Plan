# -*- coding: utf-8 -*-
"""Create Document_Registration_requirement_01092026_v1.1.docx from v1.0.

Appends a User Stories section derived from the Acts / Rules / notifications and
pain points already listed in the 01-09-2026 daily report:
  1. Guideline value calculation
  2. Valuation Data Entry Module (CVC)
  3. GIS valuation
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(
    r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Requirement Discussions\Daily Reports"
)
SRC = BASE / "Document_Registration_requirement_01092026.docx"
DST = BASE / "Document_Registration_requirement_01092026_v1.1.docx"

HEADING_FONT = "Segoe UI"
HEADING_SIZE = Pt(14.5)

GV_ROWS = [
    [
        "GV-US-01",
        "As a citizen, I want the system to compute the guideline (market) value of my "
        "property automatically from the CVC published rates, so that I know the "
        "minimum value applicable before I pay stamp duty.",
        "Stamp Act Sec. 45-A; Sec. 45-B",
    ],
    [
        "GV-US-02",
        "As a Sub-Registrar, I want stamp duty and registration fee to be computed on "
        "the higher of consideration and guideline value for the applicable Schedule "
        "article, so that duty is correctly charged ad valorem.",
        "Stamp Act Sec. 3 + Schedule; Sec. 45-A",
    ],
    [
        "GV-US-03",
        "As a Sub-Registrar, I want the correct rate category to be selected based on "
        "property classification and local body (rural / TMC / CMC / corporation), so "
        "that rural properties are not valued at TMC/CMC building rates.",
        "CVC rates under Sec. 45-B; Registration Rules 13–15",
    ],
    [
        "GV-US-04",
        "As a Sub-Registrar, I want valuation to support all area units and fractions "
        "(acre, gunta, cent, sq.ft, sq.m), so that small and fractional extents such "
        "as 1 gunta are valued correctly.",
        "Registration Rules 13–15; CVC rate units",
    ],
    [
        "GV-US-05",
        "As a Sub-Registrar, I want the system to never display blank or zero "
        "valuation silently and instead show a clear reason when a rate is missing, "
        "so that I can act instead of being blocked.",
        "Sec. 45-A (duty on market value)",
    ],
    [
        "GV-US-06",
        "As a Sub-Registrar, I want to keep registration pending and refer an "
        "undervalued instrument to the Deputy Commissioner with the property "
        "particulars, so that market value is determined as per law.",
        "Sec. 45-A; Undervaluation Rules, 1977 — Rule 3 / Form I, Rule 4",
    ],
    [
        "GV-US-07",
        "As a Sub-Registrar and as a citizen, I want the value adjudicated under "
        "Sec. 45(A) to override the system estimate after determination, so that the "
        "summary and payment reflect the adjudicated value.",
        "Sec. 45-A; Undervaluation Rules — Rule 4 (DC order)",
    ],
    [
        "GV-US-08",
        "As a citizen, I want guideline value estimation to be available for all "
        "registrable instrument types (including trust deed, settlement, release), so "
        "that valuation is not blocked by instrument type.",
        "Sec. 45-A instrument list; Schedule articles",
    ],
    [
        "GV-US-09",
        "As a citizen, I want to proceed to payment immediately after valuation "
        "without the fee showing zero or the Save option being disabled, so that my "
        "application is not stuck.",
        "Sec. 3 + Schedule; Registration Act Sec. 78 (fees)",
    ],
]

CVC_ROWS = [
    [
        "CVC-US-01",
        "As the CVC secretariat (IGR & Commissioner of Stamps), I want to enter, "
        "publish and revise market value guidelines for an area with effective dates, "
        "so that registration offices apply the notified rates.",
        "Stamp Act Sec. 45-B(1)",
    ],
    [
        "CVC-US-02",
        "As a district / sub-district market valuation sub-committee member, I want to "
        "capture and submit proposed revised rates for my jurisdiction, so that "
        "revision follows the prescribed procedure.",
        "Sec. 45-B(2) and 45-B(3)",
    ],
    [
        "CVC-US-03",
        "As the Central Valuation Committee, I want to approve, modify or reject "
        "sub-committee rate proposals with methodology remarks, so that CVC remains "
        "the final authority on policy and methodology.",
        "Sec. 45-B(2)",
    ],
    [
        "CVC-US-04",
        "As DIGR-Valuation, I want versioned rate history with effective-from and "
        "effective-to dates, so that a document is valued using the rate valid on the "
        "relevant date.",
        "Sec. 45-B(1); Act 8 of 2003",
    ],
    [
        "CVC-US-05",
        "As a Sub-Registrar and as a citizen, I want all published CVC rates to be "
        "visible in both SR login and citizen login, so that valuation does not fail "
        "due to rates not displaying.",
        "Sec. 45-B (published guidelines); Sec. 45-A (application)",
    ],
    [
        "CVC-US-06",
        "As DIGR-Valuation, I want to correct an erroneous CVC rate through a "
        "controlled departmental workflow with audit trail, instead of corrections "
        "being attempted in citizen login.",
        "Sec. 45-B(2); Act 8 of 2003",
    ],
    [
        "CVC-US-07",
        "As a Sub-Registrar, I want to send an application back to the citizen for "
        "revaluation and have it move to the correct next step, so that revaluation "
        "does not leave the application stuck or force withdrawal.",
        "Sec. 45-A (SRO estimation and communication to parties)",
    ],
    [
        "CVC-US-08",
        "As a District Registrar / Deputy Commissioner user, I want to record a "
        "Sec. 45(A) undervaluation case with its order and appeal status, so that the "
        "statutory determination and appeal are tracked.",
        "Sec. 45-A; Undervaluation Rules — Rule 4, Rule 9; RD 264 MUNOMU 99",
    ],
    [
        "CVC-US-09",
        "As a citizen, I want to pay the differential duty after the undervaluation "
        "determination, so that my pending document can be released and registered.",
        "Sec. 45-A; Undervaluation Rules, 1977",
    ],
    [
        "CVC-US-10",
        "As IGR, I want every CVC master change to be auditable (who changed, when, "
        "under whose authority), so that guideline value decisions are defensible.",
        "Sec. 45-B; Act 8 of 2003 (w.e.f. 1-4-2003)",
    ],
]

GIS_ROWS = [
    [
        "GIS-US-01",
        "As a citizen, I want to identify my property on a GIS map or by survey "
        "number / PID, so that the correct guideline value zone and rate are applied "
        "automatically.",
        "Sec. 45-B rates applied via Sec. 45-A; Registration Rules 13–15",
    ],
    [
        "GIS-US-02",
        "As the system, I want to fetch parcel geometry and attributes from KSRSAC, so "
        "that valuation is based on authoritative spatial data of the State.",
        "KSRSAC integration (implementation layer over CVC rates)",
    ],
    [
        "GIS-US-03",
        "As a Sub-Registrar, I want the GIS-derived village / hobli / taluk and survey "
        "or city survey number to match the statutory property description, so that "
        "the document description is valid.",
        "Registration Act Secs. 21–22; Registration Rules 13–15",
    ],
    [
        "GIS-US-04",
        "As DIGR-Valuation, I want CVC rate zones mapped onto GIS layers, so that "
        "rural versus TMC / CMC boundary questions are resolved by geometry rather "
        "than manual choice.",
        "Sec. 45-B (area-wise guidelines); Registration Rules 13–15",
    ],
    [
        "GIS-US-05",
        "As a Sub-Registrar, I want a controlled manual fallback with a recorded "
        "reason when GIS data is unavailable, so that registration is not blocked.",
        "Sec. 45-A (SRO duty to value); Rule 41 (examination)",
    ],
    [
        "GIS-US-06",
        "As IGR, I want alerts when GIS extent differs from land-record extent, so "
        "that extent mismatches are detected before valuation and registration.",
        "Registration Rules 13–15; Sec. 45-A",
    ],
]


def set_cell_text(cell, text: str) -> None:
    """Replace cell text, preserving the formatting of the first run."""
    paras = cell.paragraphs
    first = paras[0]
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        first.add_run(text)


def add_heading_para(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = HEADING_FONT
    run.font.size = HEADING_SIZE


def clone_table(doc: Document, template: Table, headers: list[str],
                rows: list[list[str]]) -> None:
    """Append a copy of `template` (same borders / fonts) filled with new content."""
    new_tbl = deepcopy(template._tbl)
    doc.element.body.append(new_tbl)
    table = Table(new_tbl, doc._body)

    header_tr = table.rows[0]._tr
    data_tr = table.rows[1]._tr
    for row in list(table.rows)[2:]:
        new_tbl.remove(row._tr)

    for i, head in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], head)

    template_tr = deepcopy(data_tr)
    for _ in range(len(rows) - 1):
        new_tbl.append(deepcopy(template_tr))

    for ri, values in enumerate(rows, start=1):
        for ci, value in enumerate(values):
            set_cell_text(table.rows[ri].cells[ci], value)

    assert header_tr is not None


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)

    doc = Document(str(DST))
    template = doc.tables[1]

    meta = doc.tables[0]
    meta._tbl.append(deepcopy(meta.rows[-1]._tr))
    set_cell_text(meta.rows[-1].cells[0], "Version")
    set_cell_text(
        meta.rows[-1].cells[1],
        "1.1 (03-09-2026) — user stories added for the listed Acts, Rules and pain points",
    )

    doc.add_paragraph()
    add_heading_para(doc, "User Stories")
    doc.add_paragraph()

    headers = ["ID", "User story", "Act / Rule reference"]

    add_heading_para(doc, "1. Guideline value calculation")
    clone_table(doc, template, headers, GV_ROWS)
    doc.add_paragraph()

    add_heading_para(doc, "2. Valuation Data Entry Module (CVC)")
    clone_table(doc, template, headers, CVC_ROWS)
    doc.add_paragraph()

    add_heading_para(doc, "3. GIS valuation")
    clone_table(doc, template, headers, GIS_ROWS)
    doc.add_paragraph()

    doc.save(str(DST))
    print(f"Wrote {DST}")

    check = Document(str(DST))
    print("Tables:", len(check.tables))
    for ti, tbl in enumerate(check.tables):
        print(f"  table {ti}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")


if __name__ == "__main__":
    main()
