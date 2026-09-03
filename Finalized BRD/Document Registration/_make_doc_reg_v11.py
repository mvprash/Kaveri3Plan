# -*- coding: utf-8 -*-
"""Create BRD_Document_Registration_v1.1.docx — Legal and regulatory reference (§3).

Scope: Schedule Sr.12–15 discussion topics —
  Sr.12 Registration, Appointment, Status tracking (#1, #2, #19)
  Sr.13 Stamp duty / Registration fee; Guideline value (#4, #5)
  Sr.14 Valuation Module (CVC) and GIS valuation (#21, #22)
  Sr.15 Rule 17(2), Rule 17(3), Old pending release (#6, #7, #8)

Source: Acts_Rules/Document/; Schedule:
Requirement Discussions/Schedule/Kaveri_Requirements_Updated_Schedule_v3_DocumentSubModules.xlsx
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Document Registration")
DST = BASE / "BRD_Document_Registration_v1.1.docx"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i], "D9E2F3")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri].cells[ci], val)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_heading(doc, "Business Requirements Document (BRD)", 1)
    add_heading(doc, "Document Registration Module", 2)

    # Document control
    add_heading(doc, "Document control", 2)
    dc_rows = [
        ["Document ID", "BRD-K3-DOC-001"],
        ["Version", "1.1"],
        ["Status", "Draft / In review"],
        ["Module", "Document Registration"],
        [
            "Discussion topics (Schedule Sr.12–15)",
            "Sr.12 Registration, Appointment, Status tracking (25–27 Aug 2026); "
            "Sr.13 Stamp duty / Registration fee / Guideline value (28–29 Aug 2026); "
            "Sr.14 Valuation Module (CVC) and GIS valuation (31 Aug–01 Sep 2026); "
            "Sr.15 Rule 17(2)/(3) filing and Old pending release (02–03 Sep 2026)",
        ],
        [
            "Sub-modules in scope (this version)",
            "#1 Registration, #2 Appointment, #19 Status tracking; "
            "#4 Stamp duty and Registration fee calculation, #5 Guideline value calculation; "
            "#21 Valuation Module (CVC), #22 GIS valuation; "
            "#6 Rule 17(2) filing, #7 Rule 17(3) filing, #8 Old pending release",
        ],
        [
            "Legal basis (primary)",
            "The Registration Act, 1908; The Karnataka Registration Rules, 1965; "
            "The Karnataka Stamp Act, 1957; Karnataka Stamp Rules, 1958; "
            "Karnataka Stamp (Prevention of Undervaluation of Instruments) Rules, 1977; "
            "Karnataka Stamp (Payment of Duty by Means of e-Stamping) Rules, 2009",
        ],
        ["Author (BA)", "Nandha Kumar"],
        ["Product Owner", "Prashanth"],
        ["Domain expert / SRO reviewer", "Prabhakar Naik"],
        [
            "Target audience",
            "Kaveri IT Cell, Department of Stamps and Registration, Government of Karnataka",
        ],
        ["Last updated", "2026-09-03"],
        [
            "Schedule source",
            "Requirement Discussions/Schedule/Kaveri_Requirements_Updated_Schedule_v3_DocumentSubModules.xlsx",
        ],
    ]
    add_table(doc, ["Field", "Value"], dc_rows)

    add_para(doc, "Version history:")
    add_table(
        doc,
        ["Version", "Date", "Author", "Summary of change", "Approver"],
        [
            [
                "1.0",
                "2026-09-02",
                "Nandha Kumar",
                "Initial BRD — §3 Legal reference for Sr.12 only",
                "Prashanth",
            ],
            [
                "1.1",
                "2026-09-03",
                "Nandha Kumar",
                "Expand §2–§3 to Schedule Sr.12–15 (sub-modules #1–2, #4–8, #19, #21–22) "
                "— stamp duty/fees, guideline value, CVC/GIS, Rule 17(2)/(3), old pending release",
                "Prashanth",
            ],
        ],
    )

    add_para(doc, "Related documents:")
    add_table(
        doc,
        ["ID", "Title", "Link"],
        [
            ["BRD-K3-DOC-001", "This document", ""],
            [
                "SCH-K3-REQ-v3-DOC",
                "Kaveri_Requirements_Updated_Schedule_v3_DocumentSubModules.xlsx — Sr.12–15",
                "Requirement Discussions/Schedule/",
            ],
            [
                "MOD-K3-DOC-001",
                "Kaveri_2.0_Moduleslist.xlsx — Sub-modules #1, #2, #4–8, #19, #21, #22",
                "Requirement Discussions/Modules/DocumentRegistration/",
            ],
        ],
    )

    # Contents
    add_heading(doc, "Contents", 2)
    for item in [
        "1. Executive summary",
        "2. Scope",
        "2.1 In scope (Schedule Sr.12–15)",
        "2.2 Out of scope (this BRD version)",
        "2.3 Assumptions",
        "2.4 Constraints",
        "3. Legal and regulatory reference",
        "3.1 Applicable Acts",
        "3.2 Relevant sections followed by the Department for Document Registration",
        "3.3 Relevant rules followed by the Department for Document Registration",
        "3.4 Relevant notifications issued by the Department for Document Registration",
    ]:
        add_para(doc, item)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # 1. Executive summary
    add_heading(doc, "1. Executive summary", 1)
    add_para(
        doc,
        "This Business Requirements Document (BRD) captures the legal and regulatory "
        "foundation for Kaveri 3.0 Document Registration covering Schedule discussion "
        "topics Sr.12 through Sr.15 — Registration core (registration, appointment, "
        "status tracking); stamp duty, registration fee and guideline value calculation; "
        "Valuation Module (CVC) and GIS valuation; and Rule 17(2)/(3) filing with Old "
        "pending release. It is aligned to "
        "Kaveri_Requirements_Updated_Schedule_v3_DocumentSubModules.xlsx and sourced "
        "from Acts_Rules/Document/."
    )
    add_para(
        doc,
        "This version (v1.1) documents §3 Legal and regulatory reference for these "
        "four discussion topics. Subsequent BRD versions will expand functional "
        "requirements, process flows and remaining Document sub-modules (Sr.16 onwards)."
    )

    # 2. Scope
    add_heading(doc, "2. Scope", 1)
    add_heading(doc, "2.1 In scope (Schedule Sr.12–15)", 2)
    add_para(
        doc,
        "Per Kaveri_Requirements_Updated_Schedule_v3_DocumentSubModules.xlsx:"
    )
    add_table(
        doc,
        ["Sr.No", "Discussion topic", "Sub-modules", "Planned dates", "Legal hooks (summary)"],
        [
            [
                "12",
                "Registration core: Registration, Appointment, Status tracking",
                "#1 Registration; #2 Appointment; #19 Status tracking",
                "25–27 Aug 2026",
                "Registration Act 1908 Parts II–XI; Karnataka Registration Rules "
                "Ch. II–III, IX, XII, XVI–XVIII, XXV; Karnataka Amendment Act 2023",
            ],
            [
                "13",
                "Stamp duty and Registration fee calculation; Guideline value calculation",
                "#4 Stamp duty and Registration fee; #5 Guideline value calculation",
                "28–29 Aug 2026",
                "Karnataka Stamp Act 1957 (Secs. 3, 10, 28, 33–39, Schedule); "
                "Registration Act Sec. 78 (fees); Stamp Rules 1958; e-Stamping Rules 2009; "
                "fee notification RD/46/MNMU/2025",
            ],
            [
                "14",
                "Valuation Module (CVC) and GIS valuation",
                "#21 Valuation Module (CVC); #22 GIS valuation",
                "31 Aug–01 Sep 2026",
                "Stamp Act Secs. 45-A, 45-B (market value / CVC); Prevention of "
                "Undervaluation Rules 1977; Registration Rules 13–15 (territorial / "
                "survey description)",
            ],
            [
                "15",
                "Rule 17(2) filing, Rule 17(3) filing, and Old pending release module",
                "#6 Rule 17(2); #7 Rule 17(3); #8 Old pending release",
                "02–03 Sep 2026",
                "Karnataka Registration Rules Rule 17(ii)/(iii) and Rule 17(i) Parts I–V; "
                "Secs. 19, 21, 62, 64–67 of Registration Act; Rules 23–24, 110–118 "
                "(pending / return / unclaimed)",
            ],
        ],
    )

    add_heading(doc, "2.2 Out of scope (this BRD version)", 2)
    for item in [
        "FRUITS filing, scanning, memo transmission (Sr.16; #11–13)",
        "Re-registration; will after death of testator (Sr.17; #9/#23, #10)",
        "Sec. 68(2) correction; cross-reference Rule 123 (Sr.18; #14–15)",
        "Integration, exemption, court entry, liability (Sr.19; #3, #16–18)",
        "Investigation / search; verify document; PoA authentication (Sr.20–21)",
        "DRO undervaluation adjudication case-work, deposit of will, IGRO appeal "
        "(Sr.22–23) — undervaluation legal basis is cited for Sr.13–14 valuation "
        "intake only",
        "MIS, dashboards, Digital E-Stamp module, EC, CC (Sr.24–27)",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "2.3 Assumptions", 2)
    for item in [
        "SRO / DR / IGR operate under the Registration Act, 1908 (as applicable in "
        "Karnataka) and the Karnataka Registration Rules, 1965.",
        "Stamp duty is computed from the Karnataka Stamp Act Schedule using the "
        "higher of consideration and guideline (market) value where Sec. 45-A applies.",
        "Central Valuation Committee (CVC) market-value guidelines under Sec. 45-B "
        "are the authoritative rates for guideline-value calculation (Sr.13 #5) and "
        "the Valuation Module (Sr.14 #21).",
        "GIS valuation (Sr.14 #22) overlays survey / GIS spatial data on guideline "
        "rates; property description must still satisfy Registration Rules 13–15.",
        "Rule 17(2) and Rule 17(3) in the modules list map to Karnataka Registration "
        "Rules Rule 17(ii) (copies and translations file) and Rule 17(iii) "
        "(cancellation / modification / rectification communications file); supplements "
        "under Rule 17(i) Parts I–V remain part of filing design.",
        "Old pending release covers documents held pending fine, stamp adjudication, "
        "party non-appearance, or unclaimed return — linked to Rules 46, 51–55 and 110–118.",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "2.4 Constraints", 2)
    for item in [
        "System must not register documents falling under Sec. 22-B (Karnataka "
        "Amendment Act, 2023) — forged / prohibited / attached-property instruments.",
        "Presentation time limits (Secs. 23, 25) and fine schedules (Rules 51–55) "
        "must be enforced or surfaced before acceptance.",
        "Stamp duty and registration fee masters must reflect current Gazette "
        "notifications (including RD/46/MNMU/2025 fee revision).",
        "Where market value exceeds consideration, Sec. 45-A reference / undervaluation "
        "path must be available; CVC rates cannot be silently overridden.",
        "Rule 17 filing must preserve cross-references to register entries "
        "(Rule 17(ii)) and retain departmental cancellation / modification "
        "communications (Rule 17(iii)).",
        "Pending documents must remain accountable in Daily Register / Minute Book "
        "until released or otherwise disposed under the Rules.",
    ]:
        add_para(doc, f"• {item}")

    # 3. Legal and regulatory reference
    add_heading(doc, "3. Legal and regulatory reference", 1)
    add_para(
        doc,
        "Citation convention: Act = primary Central / State legislation (with Karnataka "
        "amendments where applicable); Rules = Karnataka subordinate legislation; "
        "Notification = Gazette / Government Order. Source folder: Acts_Rules/Document/. "
        "This section covers Schedule Sr.12–15 only."
    )

    # ---------- 3.1 Applicable Acts ----------
    add_heading(doc, "3.1 Applicable Acts", 2)
    add_para(
        doc,
        "The Department of Stamps and Registration (IGR, District Registrars and "
        "Sub-Registrars) administers the following Acts for Document Registration "
        "topics Sr.12–15:"
    )
    add_table(
        doc,
        ["Act", "Act No.", "Scope in Kaveri 3.0 (Sr.12–15)", "Schedule topics", "Source file"],
        [
            [
                "The Registration Act, 1908",
                "Central Act 16 of 1908",
                "Compulsory / optional registration; presentation; place and time; "
                "duties of registering officers; refusal and appeal; fees (Sec. 78); "
                "memoranda (Secs. 64–67); copies / translations (Secs. 19, 62)",
                "Sr.12, Sr.13 (fees), Sr.15",
                "the_registration_act,_1908.pdf",
            ],
            [
                "The Registration (Karnataka Amendment) Act, 2023",
                "Karnataka Act 47 of 2024",
                "Forged document definition; refusal (Sec. 22-B); cancellation "
                "(Sec. 22-C); appeal (Sec. 22-D); penalties (Secs. 81-A, 81-B)",
                "Sr.12",
                "TheRegistration(KarnatakaAmendment)Act2023(47of2024).pdf",
            ],
            [
                "The Karnataka Registration Rules, 1965",
                "Made under Registration Act Sec. 69",
                "Office hours; presentation; examination; register books; Rule 17 "
                "supplements and filing; endorsements; receipts; appeals; pending / "
                "unclaimed document handling",
                "Sr.12, Sr.15",
                "The Karnataka Registration Rules 1965.pdf",
            ],
            [
                "The Karnataka Stamp Act, 1957",
                "Karnataka Act 34 of 1957",
                "Stamp duty chargeable on instruments; payment of duty; adjudication; "
                "impounding; market value / undervaluation (Sec. 45-A); Central "
                "Valuation Committee (Sec. 45-B); Schedule of stamp duties",
                "Sr.13, Sr.14",
                "THE KARNATAKA STAMP ACT 1957.pdf",
            ],
            [
                "The Karnataka Stamp Rules, 1958",
                "Made under Karnataka Stamp Act",
                "Impressed / adhesive / franking stamps; Superintendent of Stamps; "
                "stamp payment procedures at SRO",
                "Sr.13",
                "Karnataka Stamps Rules 1958.pdf",
            ],
            [
                "The Karnataka Stamp (Prevention of Undervaluation of Instruments) Rules, 1977",
                "Under Stamp Act Secs. 45-A and 68",
                "Reference by registering officer; market-value determination; Form I "
                "property particulars; appeal under Sec. 45-A(5)",
                "Sr.13 (#5), Sr.14 (#21)",
                "Karnataka Stamp (Constitution of Central Valuation Committee…) Rules, 2003.docx "
                "(file contains Prevention of Undervaluation Rules content)",
            ],
            [
                "The Karnataka Stamp (Payment of Duty by Means of e-Stamping) Rules, 2009",
                "Under Stamp Act Secs. 10 and 68",
                "e-Stamp certificate; Central Record Keeping Agency; Authorised "
                "Collection Centres — stamp payment at presentation",
                "Sr.13",
                "THE KARNATAKA STAMP- Payment of Stamp Duty by means of e-Stamping.pdf",
            ],
            [
                "The Karnataka Stamp (Franking Impression Of Stamps) Rules, 2000",
                "Under Karnataka Stamp Act",
                "Franking impression as stamp-payment mode at registration offices",
                "Sr.13",
                "Karnataka Stamp (Franking Impression Of Stamps) Rules, 2000.docx",
            ],
            [
                "The Transfer of Property Act, 1882",
                "Central Act 4 of 1882",
                "Sec. 54 — sale of immovable property of value > ₹100 requires "
                "registered instrument (drives Sec. 17 compulsory registration)",
                "Sr.12",
                "Referenced in Registration Act Sec. 17; TPA not in folder",
            ],
            [
                "The Indian Stamp Act, 1899",
                "Central Act 2 of 1899",
                "Referenced for certain instruments and impounding at presentation "
                "(Rule 46)",
                "Sr.12, Sr.13",
                "INSTRUMENTS GOVERNED BY THE STAMP ACT 1899.docx",
            ],
            [
                "The Karnataka Court Fees and Suits Valuation Act, 1958",
                "Karnataka Act",
                "Court-fee stamps on certain documents — Rule 48 return for deficiency",
                "Sr.12, Sr.13",
                "Referenced in Karnataka Registration Rules Rule 48",
            ],
            [
                "Karnataka Registration (Amendment) Acts (1975–2002, 2023)",
                "Various Karnataka Acts",
                "State amendments — jurisdiction, computerised records (Sec. 16A), fees",
                "Sr.12–15",
                "Registration(KarnatakaAmendment)Act*.pdf",
            ],
            [
                "Karnataka Stamp (Amendment) Acts (2011, 2012, 2014) and Schedule 2022",
                "Various / Schedule notifications",
                "Updated stamp duty rates and Schedule articles used in fee engines",
                "Sr.13",
                "THE KARNATAKA STAMP (AMENDMENT) ACT, 2011/2012; "
                "The_Karnataka_Stamp_(Amendment)_Act,_2014.pdf; "
                "Karnataka Stamp Act 1957 Schedule 2022.pdf",
            ],
        ],
    )

    # ---------- 3.2 Sections ----------
    add_heading(
        doc,
        "3.2 Relevant sections followed by the Department for Document Registration",
        2,
    )
    add_para(
        doc,
        "Selected sections that drive business rules for Sr.12–15. Full text in the "
        "cited source files."
    )

    add_heading(doc, "3.2.1 Registration Act, 1908 — Registration, Appointment, Status (Sr.12)", 3)
    add_table(
        doc,
        ["Section", "Topic", "BRD relevance"],
        [
            ["Sec. 3, 5–7, 10–12", "IGR; districts / sub-districts; Registrars / Sub-Registrars; absence / vacancy", "Office hierarchy; appointment routing; charge arrangement"],
            ["Sec. 15, 16 / 16A", "Seal; register-books; electronic records", "Digital seal; electronic register foundation"],
            ["Sec. 17–22", "Compulsory / optional registration; language; property description", "Instrument-type and property validation"],
            ["Sec. 23, 25", "Time for presentation; delay and fine", "Appointment window; status 'Suspended (fine)'"],
            ["Sec. 28–31", "Place of registration; private residence", "Jurisdiction-based appointment office selection"],
            ["Sec. 32–35", "Presentant; photograph (32A); PoA; enquiry; admission / denial", "Intake and examination status transitions"],
            ["Sec. 51–52, 58–61", "Register-books; duties on presentation; endorsements; Sec. 60 certificate; return", "Core registration completion and status 'Registered'"],
            ["Sec. 64–66", "Land in several sub-districts / districts", "Cross-office memoranda (feeds Rule 17 Part I filing)"],
            ["Sec. 68, 71–77", "Registrar control; refusal; appeal", "Escalation; status 'Refused' / on appeal"],
            ["Sec. 78", "Fees for registration, etc.", "Registration fee calculation (Sr.13)"],
            ["Sec. 19, 62", "Documents in unknown language; translation / copy", "Rule 17(ii) filing inputs (Sr.15)"],
            ["Sec. 89", "Documents forwarded for registration", "Institutional presentation origin"],
        ],
    )

    add_heading(doc, "3.2.2 Registration Act — Karnataka Amendment 2023 (Sr.12)", 3)
    add_para(doc, "Source: TheRegistration(KarnatakaAmendment)Act2023(47of2024).pdf")
    add_table(
        doc,
        ["Section", "Topic", "BRD relevance"],
        [
            ["Sec. 2 (amend.)", "Definition — 'Forged document'", "Validation block at presentation / examination"],
            ["Sec. 22-B", "Refusal to register forged / prohibited / attached-property documents", "Mandatory refusal — status 'Refused (Sec. 22-B)'"],
            ["Sec. 22-C / 22-D", "DR cancellation; appeal", "Post-registration status reversal / appeal"],
            ["Sec. 81-A / 81-B", "Penalties for registration of forged documents", "Compliance / audit flags"],
        ],
    )

    add_heading(doc, "3.2.3 Karnataka Stamp Act, 1957 — Stamp duty, fee & guideline value (Sr.13)", 3)
    add_para(doc, "Source: THE KARNATAKA STAMP ACT 1957.pdf; Schedule 2022.pdf")
    add_table(
        doc,
        ["Section", "Topic", "BRD relevance (Sr.13)"],
        [
            ["Sec. 2", "Definitions — chargeable, conveyance, duly stamped, instrument, market-value related terms, CVC", "Glossary for duty engine and valuation"],
            ["Sec. 3", "Instruments chargeable with duty", "Instrument-type → Schedule article mapping"],
            ["Sec. 10 / 10-A", "Duties how to be paid; payment modes", "e-Stamp / franking / impressed stamp at presentation"],
            ["Sec. 28", "Direction as to duty in case of certain instruments", "Duty computation for complex instruments"],
            ["Sec. 30", "Duties by whom payable", "Citizen / party messaging on who bears stamp"],
            ["Sec. 31", "Adjudication as to proper stamp", "Optional pre-registration adjudication path"],
            ["Sec. 33–39", "Examination and impounding of instruments; procedure by Deputy Commissioner", "SRO impound → DR stamp case; status 'Impounded'"],
            ["Sec. 45-A", "Instrument of conveyance etc. undervalued — market value", "Guideline value vs consideration; reference for undervaluation"],
            ["Schedule", "Stamp duty on instruments (Articles)", "Master rate table for stamp-duty calculator (#4)"],
        ],
    )

    add_heading(doc, "3.2.4 Karnataka Stamp Act — CVC and market value (Sr.14)", 3)
    add_table(
        doc,
        ["Section", "Topic", "BRD relevance (Sr.14)"],
        [
            [
                "Sec. 2(ac)",
                "Definition — Central Valuation Committee",
                "CVC entity in Valuation Module master data",
            ],
            [
                "Sec. 45-A",
                "Market value for stamp duty; reference by registering officer to Deputy Commissioner",
                "Guideline value calculation (#5) and undervaluation trigger from SRO",
            ],
            [
                "Sec. 45-B",
                "Constitution of Central Valuation Committee under IGR & Commissioner of Stamps; "
                "estimation, publication and revision of market value guidelines; sub-committees "
                "at district / sub-district; CVC is final authority for policy and methodology",
                "Valuation Module (CVC) (#21) — rate publication, revision cycles, sub-committee hierarchy",
            ],
            [
                "Sec. 46",
                "Recovery of duties and penalties; charge on property; note in registration indices",
                "Post-valuation recovery linkage to registration indices",
            ],
        ],
    )

    add_heading(doc, "3.2.5 Sections feeding Rule 17 filing (Sr.15)", 3)
    add_table(
        doc,
        ["Section", "Topic", "BRD relevance (Sr.15)"],
        [
            ["Sec. 19 / 62", "Language unknown; translation and copy", "Rule 17(ii) file of copies and translations"],
            ["Sec. 21", "Maps or plans with documents", "Rule 17(i) Part II — copies of maps / plans"],
            ["Sec. 64–67", "Memoranda when property spans offices / districts", "Rule 17(i) Part I supplements"],
            [
                "Related filing inputs (Rule 17(i) Parts III–V)",
                "Court / Revenue sale certificates; Land Acquisition statements; "
                "Karnataka Land Improvement / Agriculturists Loans instruments; "
                "Land Development Bank instruments under Co-operative Societies Act Sec. 85-A",
                "Rule 17 supplement filing workflows (#6/#7 related institutional filing)",
            ],
        ],
    )

    # ---------- 3.3 Rules ----------
    add_heading(
        doc,
        "3.3 Relevant rules followed by the Department for Document Registration",
        2,
    )
    add_para(
        doc,
        "Selected Karnataka rules mapped to Sr.12–15 discussion topics."
    )

    add_heading(doc, "3.3.1 Karnataka Registration Rules, 1965 — Registration core (Sr.12)", 3)
    add_para(doc, "Source: The Karnataka Registration Rules 1965.pdf")
    add_table(
        doc,
        ["Rule", "Requirement", "System feature"],
        [
            ["Rules 3, 5", "Office hours; holidays", "Appointment slot windows; holiday calendar"],
            ["Rules 6–8", "Officer address; Joint Sub-Registrars", "Officer / charge master"],
            ["Rules 13–15", "Territorial divisions; property description; survey / city survey numbers", "Property capture; GIS / survey linkage (also Sr.14)"],
            ["Rules 16–17, 19 / 19-A, 22–24", "Register books; document sheets; daily register; cash book; minute book", "Register entry; end-of-day reconciliation"],
            ["Rules 37, 40–46, 51–55", "Presentation office; photograph; examination; endorsement; suspension; fines", "Intake, defect, suspend, fine workflows"],
            ["Rules 71–73, 78–79, 94, 104, 107", "Executant examination; thumb impression; endorsements; Sec. 60 certificate", "Examination and registration completion"],
            ["Rules 110–113, 175–188", "Receipts; return of document; appeal against refusal", "Status tracking termini"],
        ],
    )

    add_heading(doc, "3.3.2 Stamp Rules — duty calculation and payment (Sr.13)", 3)
    add_table(
        doc,
        ["Instrument", "Key provisions", "System feature (Sr.13)"],
        [
            [
                "Karnataka Stamp Rules, 1958",
                "Impressed and adhesive stamps; proper officer; issue of stamped paper; "
                "procedures under Stamp Act Secs. 10, 18, 36, 47, 68",
                "Stamp payment mode validation; stamp inventory / franking hooks",
            ],
            [
                "Karnataka Stamp (Franking Impression Of Stamps) Rules, 2000",
                "Franking machines at authorised offices",
                "Franking as stamp-payment option at SRO",
            ],
            [
                "Karnataka Stamp (Payment of Duty by Means of e-Stamping) Rules, 2009",
                "e-Stamp certificate; CRA; Authorised Collection Centres; verification "
                "at Sub-Registrar offices",
                "e-Stamp certificate capture / verify in duty calculator; CRA integration",
            ],
            [
                "Karnataka Stamp Act Schedule (as amended / Schedule 2022)",
                "Article-wise stamp duty rates",
                "Stamp-duty calculation engine master (#4)",
            ],
            [
                "Registration fee Table (Sec. 78) as amended",
                "Article I / III fees — see RD/46/MNMU/2025",
                "Registration-fee calculation engine (#4)",
            ],
        ],
    )

    add_heading(doc, "3.3.3 Prevention of Undervaluation Rules / CVC practice (Sr.13 #5, Sr.14)", 3)
    add_para(
        doc,
        "Source file in folder is titled as CVC Rules 2003 but the embedded text is "
        "the Karnataka Stamp (Prevention of Undervaluation of Instruments) Rules, 1977 "
        "(GSR 81; Notification No. RD 73 EST 74, dated 2-3-1977; amended by "
        "RD 264 MUNOMU 99). CVC constitution and market-value guidelines are governed "
        "by Stamp Act Sec. 45-B."
    )
    add_table(
        doc,
        ["Rule / provision", "Requirement", "System feature"],
        [
            [
                "Stamp Act Sec. 45-B",
                "CVC under IGR & Commissioner of Stamps; publish and revise market "
                "value guidelines; district / sub-district market valuation sub-committees",
                "Valuation Module (CVC) (#21) — guideline rate masters; revision workflow",
            ],
            [
                "Undervaluation Rules — Rule 3 / Form I",
                "Statement of particulars of property and its market value",
                "Property valuation capture form at SRO / citizen intake",
            ],
            [
                "Undervaluation Rules — Rule 4",
                "On Sec. 45-A reference, Deputy Commissioner communicates / determines value",
                "Reference case from registration to DR stamp / valuation",
            ],
            [
                "Undervaluation Rules — principles for determination of market value",
                "Factors for land / building valuation",
                "Guideline value calculation (#5); CVC methodology alignment",
            ],
            [
                "Undervaluation Rules — appeal (Rule 9 onwards)",
                "Appeal under Sec. 45-A(5) to Divisional Commissioner",
                "Appeal status (detailed IGRO / DRO undervaluation — Sr.22–23)",
            ],
            [
                "Registration Rules 13–15",
                "Territorial divisions; survey / Pot Hissa / city survey numbers",
                "GIS valuation (#22) — spatial join of property to guideline polygon / survey",
            ],
        ],
    )

    add_heading(doc, "3.3.4 Rule 17 filing and Old pending release (Sr.15)", 3)
    add_para(
        doc,
        "Source: The Karnataka Registration Rules 1965.pdf — Rule 17. Module labels "
        "Rule 17(2) and Rule 17(3) correspond to Rule 17(ii) and Rule 17(iii)."
    )
    add_table(
        doc,
        ["Rule", "Requirement", "System feature (Sr.15)"],
        [
            [
                "Rule 17(i) Part I",
                "Supplements to Book 1 for Secs. 64–67 memoranda",
                "Cross-office memo filing (related to Sr.16 memo transmission)",
            ],
            [
                "Rule 17(i) Part II",
                "Copies of maps or plans under Sec. 21",
                "Map / plan filing with document",
            ],
            [
                "Rule 17(i) Part III",
                "(a) Court / Revenue sale certificates; (b) Land Acquisition statements "
                "from Deputy Commissioner",
                "Institutional Rule 17 supplement filing",
            ],
            [
                "Rule 17(i) Parts IV–V",
                "Copies of instruments under Land Improvement / Agriculturists Loans Acts; "
                "Land Development Bank instruments (Co-operative Societies Act Sec. 85-A)",
                "Departmental instrument filing under Rule 17",
            ],
            [
                "Rule 17(ii) — modules list #6 Rule 17(2)",
                "Separate file for copies and translations under Secs. 19 and 62 / Rule 12(1); "
                "cross-reference with register entry",
                "Rule 17(2) filing module — store copy/translation and link to register number",
            ],
            [
                "Rule 17(iii) — modules list #7 Rule 17(3)",
                "Separate file for communications from other departments intimating "
                "cancellation, modification or rectification of previously filed / registered papers",
                "Rule 17(3) filing module — ingest and index departmental communications",
            ],
            [
                "Rule 23",
                "Minute Book — notes of suspension / deviation / refusal / summons / withdrawal",
                "Audit trail for pending cases",
            ],
            [
                "Rule 24 + Forms 9–11",
                "Daily Register; registers of impounded, unclaimed, and deficient fee / stamp documents",
                "Old pending inventory — impounded / unclaimed / deficient",
            ],
            [
                "Rules 46, 51–55",
                "Suspension for delay fine or stamp impounding; fine rates; condonation",
                "Pending reasons and release conditions",
            ],
            [
                "Rules 110–118",
                "Receipts; return of registered documents; loss of receipt; objection to return; "
                "registration after stamp adjudication",
                "Old pending release module (#8) — release / return / collection workflow",
            ],
        ],
    )

    # ---------- 3.4 Notifications ----------
    add_heading(
        doc,
        "3.4 Relevant notifications issued by the Department for Document Registration",
        2,
    )
    add_para(
        doc,
        "Gazette notifications, Government Orders and amendments relevant to Sr.12–15. "
        "Source folder: Acts_Rules/Document/."
    )
    add_table(
        doc,
        ["Instrument", "Date / No.", "Effect", "BRD relevance", "Topics", "Source file"],
        [
            [
                "The Karnataka Registration Rules, 1965",
                "Under Registration Act Sec. 69",
                "Procedural rules for all registration offices including Rule 17",
                "Presentation, register, Rule 17 filing, pending release",
                "Sr.12, Sr.15",
                "The Karnataka Registration Rules 1965.pdf",
            ],
            [
                "Notification No. RGN 2/2002-03",
                "1 Apr 2002; w.e.f. 4 Apr 2002",
                "Rule 19-A document sheets; Rules 22-A–22-C; Rule 40 photograph / digital photo",
                "Document format; photo at presentation",
                "Sr.12",
                "Cited in Karnataka Registration Rules 1965.pdf",
            ],
            [
                "Notification No. RGN/287/02-03",
                "29 Mar 2003",
                "Inserts / amends Rule 19-A document sheet specifications",
                "Presentation document standards",
                "Sr.12",
                "Cited in Karnataka Registration Rules 1965.pdf",
            ],
            [
                "RD 403 ESR 85",
                "27 May 1986",
                "Table of Registration Fees under Sec. 78",
                "Base registration fee schedule",
                "Sr.13",
                "Referenced in RegistrationActNotification.pdf",
            ],
            [
                "RD/46/MNMU/2025",
                "29 Aug 2025; w.e.f. 31 Aug 2025",
                "Substitutes registration fee: Art. I(4)(a) ₹10→₹20 per ₹1,000; "
                "Art. III(a)(i)/(ii) ₹1→₹2",
                "Current registration-fee master for calculator (#4)",
                "Sr.13",
                "RegistrationActNotification.pdf",
            ],
            [
                "RD 380 MUNOMU 2008",
                "8 Apr 2009",
                "Karnataka Stamp (Payment of Duty by Means of e-Stamping) Rules, 2009",
                "e-Stamp payment and verification at SRO",
                "Sr.13",
                "THE KARNATAKA STAMP- Payment of Stamp Duty by means of e-Stamping.pdf",
            ],
            [
                "RD 73 EST 74 (GSR 81)",
                "2 Mar 1977; Gazette 10 Mar 1977",
                "Karnataka Stamp (Prevention of Undervaluation of Instruments) Rules, 1977",
                "Form I; Sec. 45-A reference procedure",
                "Sr.13, Sr.14",
                "CVC-named docx in Acts_Rules/Document/ (Undervaluation Rules text)",
            ],
            [
                "RD 264 MUNOMU 99",
                "18 Aug 1999; Gazette 21 Aug 1999",
                "Amendments to Undervaluation Rules (omit provisional order path; Rule 6 omitted)",
                "Current undervaluation workflow shape",
                "Sr.14",
                "Cited in Undervaluation Rules docx",
            ],
            [
                "Stamp Act Sec. 45-B (Act 8 of 2003)",
                "w.e.f. 1 Apr 2003 (substituted)",
                "CVC constitution; market value guidelines; sub-committees",
                "Legal foundation for Valuation Module (CVC) (#21)",
                "Sr.14",
                "THE KARNATAKA STAMP ACT 1957.pdf",
            ],
            [
                "Karnataka Stamp Act Schedule 2022 / Stamp Amendment Acts 2011–2014",
                "Various",
                "Updated Schedule articles and rates",
                "Stamp-duty calculation master",
                "Sr.13",
                "Karnataka Stamp Act 1957 Schedule 2022.pdf; Amendment Act PDFs",
            ],
            [
                "The Registration (Karnataka Amendment) Act, 2023",
                "Gazette Extra-ordinary No. 480; 19 Oct 2024",
                "Secs. 22-B–22-D, 81-A–81-B",
                "Forged / prohibited document controls at registration",
                "Sr.12",
                "TheRegistration(KarnatakaAmendment)Act2023(47of2024).pdf",
            ],
            [
                "Karnataka Stamp (Franking Impression Of Stamps) Rules, 2000",
                "Under Stamp Act",
                "Franking at registration offices",
                "Stamp payment mode",
                "Sr.13",
                "Karnataka Stamp (Franking Impression Of Stamps) Rules, 2000.docx",
            ],
        ],
    )

    doc.save(str(DST))
    print(f"Wrote {DST}")


if __name__ == "__main__":
    build()
