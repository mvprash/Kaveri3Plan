# -*- coding: utf-8 -*-
"""Create BRD_Document_Registration_v1.0.docx — Legal and regulatory reference (§3).

Scope: Schedule Sr.12 discussion topic — Registration core (Registration,
Appointment, Status tracking). Source: Acts_Rules/Document/.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Document Registration")
DST = BASE / "BRD_Document_Registration_v1.0.docx"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i], "D9E2F3")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri].cells[ci], val)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_heading(doc, "Business Requirements Document (BRD)", 1)
    add_heading(doc, "Document Registration Module", 2)

    # Document control
    add_heading(doc, "Document control", 2)
    dc_headers = ["Field", "Value"]
    dc_rows = [
        ["Document ID", "BRD-K3-DOC-001"],
        ["Version", "1.0"],
        ["Status", "Draft / In review"],
        ["Module", "Document Registration"],
        [
            "Discussion topic (Schedule Sr.12)",
            "Registration core: Registration, Appointment, Status tracking "
            "(Current issues & process walkthrough) — 25–27 Aug 2026",
        ],
        [
            "Sub-modules in scope (this version)",
            "#1 Registration, #2 Appointment, #19 Status tracking",
        ],
        [
            "Legal basis (primary)",
            "The Registration Act, 1908; The Karnataka Registration Rules, 1965; "
            "The Karnataka Stamp Act, 1957",
        ],
        ["Author (BA)", "Nandha Kumar"],
        ["Product Owner", "Prashanth"],
        ["Domain expert / SRO reviewer", "Prabhakar Naik"],
        [
            "Target audience",
            "Kaveri IT Cell, Department of Stamps and Registration, Government of Karnataka",
        ],
        ["Last updated", "2026-09-02"],
    ]
    add_table(doc, dc_headers, dc_rows)

    add_para(doc, "Version history:")
    vh_headers = ["Version", "Date", "Author", "Summary of change", "Approver"]
    vh_rows = [
        [
            "1.0",
            "2026-09-02",
            "Nandha Kumar",
            "Initial BRD for Schedule Sr.12 — §3 Legal and regulatory reference only "
            "(Applicable Acts, sections, rules, notifications from Acts_Rules/Document/)",
            "Prashanth",
        ],
    ]
    add_table(doc, vh_headers, vh_rows)

    add_para(doc, "Related documents:")
    rd_headers = ["ID", "Title", "Link"]
    rd_rows = [
        ["BRD-K3-DOC-001", "This document", ""],
        [
            "SCH-K3-REQ-v3",
            "Kaveri_Requirements_Updated_Schedule_v3.xlsx — Sr.12",
            "Requirement Discussions/Schedule/Final/",
        ],
        [
            "MOD-K3-DOC-001",
            "Kaveri_2.0_Moduleslist.xlsx — Sub-modules #1, #2, #19",
            "Requirement Discussions/Modules/DocumentRegistration/",
        ],
    ]
    add_table(doc, rd_headers, rd_rows)

    # Contents
    add_heading(doc, "Contents", 2)
    toc = [
        "1. Executive summary",
        "2. Scope",
        "2.1 In scope (Registration, Appointment, Status tracking — Sr.12)",
        "2.2 Out of scope (this BRD version)",
        "2.3 Assumptions",
        "2.4 Constraints",
        "3. Legal and regulatory reference",
        "3.1 Applicable Acts",
        "3.2 Relevant sections followed by the Department for Document Registration",
        "3.3 Relevant rules followed by the Department for Document Registration",
        "3.4 Relevant notifications issued by the Department for Document Registration",
    ]
    for item in toc:
        add_para(doc, item)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # 1. Executive summary
    add_heading(doc, "1. Executive summary", 1)
    add_para(
        doc,
        "This Business Requirements Document (BRD) captures the legal and regulatory "
        "foundation for Kaveri 3.0 Document Registration — specifically the Schedule "
        "Sr.12 discussion topic covering Registration core workflows (document "
        "registration, citizen appointment booking, and application status tracking). "
        "It is derived from the requirements discussion held from 25 August 2026 and "
        "sources in Acts_Rules/Document/."
    )
    add_para(
        doc,
        "This version (v1.0) documents §3 Legal and regulatory reference only — "
        "applicable Central and State Acts, selected statutory sections, Karnataka "
        "Registration Rules, and Gazette notifications that govern presentation, "
        "examination, registration, office operations, fees, and status-related "
        "outputs. Subsequent BRD versions will expand functional requirements, "
        "process flows, and integration points as further schedule topics are covered."
    )

    # 2. Scope
    add_heading(doc, "2. Scope", 1)

    add_heading(doc, "2.1 In scope (Registration, Appointment, Status tracking — Sr.12)", 2)
    add_para(
        doc,
        "Per Kaveri_Requirements_Updated_Schedule_v3.xlsx (Sr.12; planned 25–27 Aug 2026):"
    )
    scope_rows = [
        ["Sub-module", "Modules list #", "Discussion focus", "Legal hooks (summary)"],
        [
            "Registration",
            "#1",
            "Core document registration — presentation, examination, admission/denial "
            "of execution, endorsement, register entry, certificate (Sec. 60), return of "
            "document; refusal and appeal triggers",
            "Registration Act 1908 Parts V–XI; Karnataka Registration Rules Ch. IX, XII, "
            "XVI–XVIII; Karnataka Amendment Act 2023 (forged documents)",
        ],
        [
            "Appointment",
            "#2",
            "Citizen appointment for registration office visit — slot booking, office hours, "
            "holiday calendar, jurisdiction-based office selection, queue management at SRO",
            "Registration Act Secs. 28–31 (place of registration); Karnataka Registration "
            "Rules Ch. II (office hours / holidays), Rules 37 (office where document may be "
            "registered), Rule 40 (presentation)",
        ],
        [
            "Status tracking",
            "#19",
            "End-to-end visibility of registration application — presented, under scrutiny, "
            "defect / re-presentation, suspended (fine / stamp), registered, refused, on appeal, "
            "document returned",
            "Registration Act Secs. 25, 34–35, 46, 71–77; Karnataka Registration Rules "
            "Rules 45–46, 110–118, 175–188",
        ],
    ]
    add_table(doc, scope_rows[0], scope_rows[1:])

    add_heading(doc, "2.2 Out of scope (this BRD version)", 2)
    for item in [
        "Stamp duty / registration fee calculation (Schedule Sr.13; sub-modules #4, #5)",
        "Guideline value / CVC / GIS valuation (Sr.14)",
        "Rule 17(2)/(3) filing, old pending release, e-filing (Sr.15)",
        "FRUITS filing, scanning, memo transmission (Sr.16)",
        "Re-registration, will after death of testator (Sr.17)",
        "Sec. 68(2) correction, cross-reference Rule 123 (Sr.18)",
        "Integration, exemption, court entry, liability (Sr.19)",
        "Investigation, search, verify document, PoA authentication (Sr.20–21)",
        "DRO / IGRO processes, MIS, dashboards (Sr.22–24)",
        "Digital E-Stamp module (Sr.25) — referenced only where stamp presentation "
        "intersects registration core",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "2.3 Assumptions", 2)
    for item in [
        "Sub-Registrars (SRO), District Registrars (DR) and Inspector General of "
        "Registration (IGR) operate under the Registration Act, 1908 as applicable in "
        "Karnataka together with the Karnataka Registration Rules, 1965.",
        "Stamp duty adequacy is verified at presentation; detailed stamp calculation "
        "engines are out of scope for this version but impounding / suspension rules "
        "apply at registration intake.",
        "Kaveri 3.0 will support both online pre-registration / appointment and "
        "in-person presentation at the SRO, consistent with Rule 40 and amendments "
        "permitting electronic processes where notified.",
        "Status tracking mirrors statutory milestones — presentation endorsement, "
        "examination, registration certificate, refusal order, and appeal — rather "
        "than internal IT statuses alone.",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "2.4 Constraints", 2)
    for item in [
        "System design must not permit registration of documents falling under "
        "Sec. 22-B (Karnataka Amendment Act, 2023) — forged documents and other "
        "prohibited transactions.",
        "Presentation time limits (Secs. 23, 25) and fine schedules (Rules 51–55) "
        "must be enforced or surfaced before acceptance.",
        "Office hours and gazetted holidays (Rules 3–5) constrain appointment slot "
        "availability.",
        "Register books, daily register, and endorsement forms (Rules 16–24, 94–104) "
        "define mandatory outputs that cannot be omitted in To-Be design.",
    ]:
        add_para(doc, f"• {item}")

    # 3. Legal and regulatory reference
    add_heading(doc, "3. Legal and regulatory reference", 1)
    add_para(
        doc,
        "Citation convention: Act = primary Central legislation (with Karnataka "
        "amendments where applicable); Rules = Karnataka subordinate legislation "
        "made under the Act; Notification = Gazette / Government Order cited in the "
        "Rules or issued separately. Source folder: Acts_Rules/Document/. "
        "This section is scoped to Schedule Sr.12 (Registration, Appointment, "
        "Status tracking)."
    )

    # 3.1 Applicable Acts
    add_heading(doc, "3.1 Applicable Acts", 2)
    add_para(
        doc,
        "The Department of Stamps and Registration (through IGR, District Registrars "
        "and Sub-Registrars) administers the following Acts in Kaveri 3.0 Document "
        "Registration — Registration core scope:"
    )
    acts_rows = [
        ["Act", "Act No.", "Scope in Kaveri 3.0 (Sr.12)", "Source file"],
        [
            "The Registration Act, 1908",
            "Central Act 16 of 1908",
            "Compulsory and optional registration; presentation; place and time of "
            "registration; duties of registering officers; refusal, appeal; inspection "
            "and certified copies; effects of registration / non-registration",
            "the_registration_act,_1908.pdf",
        ],
        [
            "The Registration (Karnataka Amendment) Act, 2023",
            "Karnataka Act 47 of 2024",
            "Forged document definition; refusal (Sec. 22-B); cancellation (Sec. 22-C); "
            "appeal (Sec. 22-D); penalties (Secs. 81-A, 81-B)",
            "TheRegistration(KarnatakaAmendment)Act2023(47of2024).pdf",
        ],
        [
            "The Karnataka Registration Rules, 1965",
            "Made under Registration Act Sec. 69",
            "Office hours, presentation, examination, endorsements, register books, "
            "indexes, receipts, appeals — operational procedures for SRO/DR",
            "The Karnataka Registration Rules 1965.pdf",
        ],
        [
            "The Transfer of Property Act, 1882",
            "Central Act 4 of 1882",
            "Sec. 54 — sale of immovable property of value > ₹100 requires registered "
            "instrument; drives compulsory registration under Registration Act Sec. 17",
            "Referenced in Registration Act Sec. 17; TPA not in folder",
        ],
        [
            "The Indian Stamp Act, 1899",
            "Central Act 2 of 1899",
            "Referenced for impounding inadequately stamped documents at presentation "
            "(Rule 46); applicable to certain instruments",
            "INSTRUMENTS GOVERNED BY THE STAMP ACT 1899.docx",
        ],
        [
            "The Karnataka Stamp Act, 1957",
            "Karnataka Act 34 of 1957",
            "Stamp duty chargeable on instruments presented for registration; "
            "adjudication and impounding at SRO",
            "THE KARNATAKA STAMP ACT 1957.pdf",
        ],
        [
            "The Karnataka Stamp Rules, 1958",
            "Made under Karnataka Stamp Act",
            "Stamp forms, impressed / adhesive stamps, Superintendent of Stamps procedures",
            "Karnataka Stamps Rules 1958.pdf",
        ],
        [
            "The Karnataka Court Fees and Suits Valuation Act, 1958",
            "Karnataka Act",
            "Court-fee stamps on certain documents — Rule 48 return for deficiency",
            "Referenced in Karnataka Registration Rules Rule 48",
        ],
        [
            "The Indian Evidence Act, 1872",
            "Central Act 1 of 1872",
            "Sec. 91 — public records / registered documents as evidence",
            "Referenced in Registration Act",
        ],
        [
            "Karnataka Registration (Amendment) Acts (1975–2002, 2023)",
            "Various Karnataka Acts",
            "State amendments to Registration Act — appointment, jurisdiction, "
            "computerised records (Sec. 16A), etc.",
            "Registration(KarnatakaAmendment)Act*.pdf in Acts_Rules/Document/",
        ],
    ]
    add_table(doc, acts_rows[0], acts_rows[1:])

    # 3.2 Sections
    add_heading(doc, "3.2 Relevant sections followed by the Department for Document Registration", 2)
    add_para(
        doc,
        "Selected sections of the Registration Act, 1908 (as applicable in Karnataka) "
        "that drive business rules, validations, appointment routing, and status "
        "milestones for Registration, Appointment and Status tracking. Full text in "
        "the cited source files."
    )

    add_heading(doc, "3.2.1 Registration establishment and offices (Part II)", 3)
    sec_part2 = [
        ["Section", "Topic", "BRD relevance (Sr.12)"],
        [
            "Sec. 3",
            "Inspector-General of Registration",
            "IGR oversight; rule-making; superintendence of registration offices",
        ],
        [
            "Sec. 5",
            "Districts and sub-districts",
            "Jurisdiction master — DR / SRO office hierarchy",
        ],
        [
            "Sec. 6",
            "Registrars and Sub-Registrars",
            "Appointment of registering officers; office assignment",
        ],
        [
            "Sec. 7",
            "Offices of Registrar and Sub-Registrar",
            "Office location for appointment booking and presentation",
        ],
        [
            "Sec. 10–12",
            "Absence / vacancy of Registrar or Sub-Registrar",
            "Charge arrangement; routing during vacancy — status and office master",
        ],
        [
            "Sec. 15",
            "Seal of registering officers",
            "Digital / physical seal on endorsements and certificates",
        ],
        [
            "Sec. 16 / 16A",
            "Register-books; computer floppies / electronic records",
            "Electronic register books; scanning / digital storage foundation",
        ],
    ]
    add_table(doc, sec_part2[0], sec_part2[1:])

    add_heading(doc, "3.2.2 Registerable documents and property description (Part III)", 3)
    sec_part3 = [
        ["Section", "Topic", "BRD relevance (Sr.12)"],
        [
            "Sec. 17",
            "Documents of which registration is compulsory",
            "Instrument-type validation — sale, mortgage, lease >1 year, etc.",
        ],
        [
            "Sec. 18",
            "Documents of which registration is optional",
            "Optional registration path",
        ],
        [
            "Sec. 19–20",
            "Language; interlineations, blanks, erasures",
            "Defect workflow — translation, attestation of alterations",
        ],
        [
            "Sec. 21–22",
            "Description of property; maps / Government surveys",
            "Property description validation; GIS / survey number linkage",
        ],
        [
            "Sec. 22-A",
            "Re-registration of certain documents",
            "Related sub-module (out of scope v1.0) — cross-reference only",
        ],
    ]
    add_table(doc, sec_part3[0], sec_part3[1:])

    add_heading(doc, "3.2.3 Time and place of registration (Parts IV–V)", 3)
    sec_part45 = [
        ["Section", "Topic", "BRD relevance (Sr.12)"],
        [
            "Sec. 23",
            "Time for presenting documents",
            "4-month presentation window — appointment scheduling constraint",
        ],
        [
            "Sec. 25",
            "Delay in presentation unavoidable",
            "Fine / condonation workflow — status 'Suspended pending fine'",
        ],
        [
            "Sec. 28",
            "Place for registering documents relating to land",
            "SRO jurisdiction by property location — appointment office selection",
        ],
        [
            "Sec. 29",
            "Place for registering other documents",
            "Jurisdiction for non-land documents",
        ],
        [
            "Sec. 30–31",
            "Registration by Registrars; private residence",
            "Special presentation routes — appointment types",
        ],
    ]
    add_table(doc, sec_part45[0], sec_part45[1:])

    add_heading(doc, "3.2.4 Presentation, examination and registration procedure (Parts VI–VII)", 3)
    sec_part67 = [
        ["Section", "Topic", "BRD relevance (Sr.12)"],
        [
            "Sec. 32",
            "Persons to present documents",
            "Presentant eligibility; agent / PoA rules",
        ],
        [
            "Sec. 32A",
            "Compulsory affixing of photograph",
            "Photograph on document and thumb-impression register — Rule 40",
        ],
        [
            "Sec. 33",
            "Power-of-attorney recognisable for presentation",
            "PoA validation at intake (detailed PoA module — Sr.21)",
        ],
        [
            "Sec. 34",
            "Enquiry before registration",
            "SRO examination workflow",
        ],
        [
            "Sec. 35",
            "Admission and denial of execution",
            "Status transitions — admitted / denied execution",
        ],
        [
            "Sec. 36–39",
            "Summons for executants and witnesses",
            "Enforcement of appearance — appointment / summons status",
        ],
    ]
    add_table(doc, sec_part67[0], sec_part67[1:])

    add_heading(doc, "3.2.5 Duties of registering officers and registration outputs (Part XI)", 3)
    sec_part11 = [
        ["Section", "Topic", "BRD relevance (Sr.12)"],
        [
            "Sec. 51",
            "Register-books to be kept",
            "Books 1–5; daily register — register entry on completion",
        ],
        [
            "Sec. 52",
            "Duties when document presented",
            "Checklist at presentation — drives intake validation",
        ],
        [
            "Sec. 58–61",
            "Endorsements; certificate of registration; copy in register; return of document",
            "Presentation endorsement (Rule 45); Sec. 60 certificate; status 'Registered'",
        ],
        [
            "Sec. 64–66",
            "Documents relating to land in several sub-districts / districts",
            "Cross-office memoranda — status propagation",
        ],
        [
            "Sec. 68",
            "Registrar superintendence and control of Sub-Registrars",
            "DR oversight; escalation from SRO",
        ],
        [
            "Sec. 71–77",
            "Refusal to register; appeal to Registrar / District Court",
            "Status 'Refused'; appeal workflow terminus",
        ],
    ]
    add_table(doc, sec_part11[0], sec_part11[1:])

    add_heading(doc, "3.2.6 Effects, inspection and copies (Parts X, XIV)", 3)
    sec_part10_14 = [
        ["Section", "Topic", "BRD relevance (Sr.12)"],
        [
            "Sec. 47–50",
            "Effect of registration / non-registration",
            "Legal effect messaging to citizens post-registration",
        ],
        [
            "Sec. 57",
            "Inspection of books; certified copies of entries",
            "Post-registration citizen services (CC module — Sr.27)",
        ],
        [
            "Sec. 78",
            "Fees for registration, etc.",
            "Registration fee at presentation — fee notification RD/46/MNMU/2025",
        ],
        [
            "Sec. 89",
            "Documents forwarded to Sub-Registrar for registration",
            "Institutional / departmental forwarding — status origin",
        ],
    ]
    add_table(doc, sec_part10_14[0], sec_part10_14[1:])

    add_heading(doc, "3.2.7 Karnataka Amendment Act, 2023 — selected new provisions", 3)
    add_para(doc, "Source: TheRegistration(KarnatakaAmendment)Act2023(47of2024).pdf")
    sec_kar2023 = [
        ["Section", "Topic", "BRD relevance (Sr.12)"],
        [
            "Sec. 2 (amend.)",
            "Definition — 'Forged document'",
            "Validation block at presentation / examination",
        ],
        [
            "Sec. 22-B",
            "Refusal to register forged / prohibited / attached-property documents",
            "Mandatory refusal grounds — status 'Refused (Sec. 22-B)'",
        ],
        [
            "Sec. 22-C",
            "Cancellation of registration by District Registrar",
            "Post-registration status reversal (DR workflow)",
        ],
        [
            "Sec. 22-D",
            "Appeal against District Registrar cancellation order",
            "Appeal status tracking",
        ],
        [
            "Sec. 81-A / 81-B",
            "Penalties for registration of forged documents",
            "Compliance / audit flags",
        ],
    ]
    add_table(doc, sec_kar2023[0], sec_kar2023[1:])

    # 3.3 Rules
    add_heading(doc, "3.3 Relevant rules followed by the Department for Document Registration", 2)
    add_para(
        doc,
        "Selected rules under the Karnataka Registration Rules, 1965 governing office "
        "operations, presentation, examination, registration procedure, receipts and "
        "appeals — mapped to Registration, Appointment and Status tracking."
    )

    add_heading(doc, "3.3.1 Office hours, holidays and registering officers (Chapters II–III)", 3)
    add_para(doc, "Source: The Karnataka Registration Rules 1965.pdf — Chapters II–III.")
    rules_ch2 = [
        ["Rule", "Requirement", "System feature (Sr.12)"],
        [
            "Rule 3",
            "Office hours of registration offices",
            "Appointment slot windows; citizen-facing office timing",
        ],
        [
            "Rule 5",
            "Holidays",
            "Holiday calendar integration — block appointment on non-working days",
        ],
        [
            "Rule 6",
            "Residential address of registering officers",
            "Office / officer master data",
        ],
        [
            "Rule 7–8",
            "Joint Sub-Registrars; registration by Joint Sub-Registrar",
            "Officer charge and routing during joint posting",
        ],
    ]
    add_table(doc, rules_ch2[0], rules_ch2[1:])

    add_heading(doc, "3.3.2 Books, forms and daily registers (Chapter VII)", 3)
    rules_ch7 = [
        ["Rule", "Requirement", "System feature (Sr.12)"],
        [
            "Rule 16–17",
            "Register Books 1–5; supplements and file copies",
            "Electronic register structure; filing copies",
        ],
        [
            "Rule 19 / 19-A",
            "Supply of books; document sheets",
            "Document sheet format; A4 / GSM specifications for presentation pack",
        ],
        [
            "Rule 22–22-C",
            "Additional register books; preparation and authentication of copy",
            "Copy comparison workflow before register entry",
        ],
        [
            "Rule 23",
            "Minute Book (Form 6) — suspension / deviation notes",
            "Audit trail for non-routine registration actions",
        ],
        [
            "Rule 24",
            "Daily Register (Form 7); Cash Book (Form 8)",
            "End-of-day registration count; fee reconciliation — status closure",
        ],
    ]
    add_table(doc, rules_ch7[0], rules_ch7[1:])

    add_heading(doc, "3.3.3 Presentation and examination (Chapter IX)", 3)
    rules_ch9 = [
        ["Rule", "Requirement", "System feature (Sr.12)"],
        [
            "Rule 37",
            "Office where document may be registered",
            "Jurisdiction routing — appointment at correct SRO",
        ],
        [
            "Rule 40",
            "Presentation of document — in person; photograph; no post except as per law",
            "Online pre-fill + in-person presentation; photo capture; presentation endorsement",
        ],
        [
            "Rule 41",
            "Examination before acceptance — defects remedied and re-presented",
            "Defect / re-presentation status; SRO advice workflow",
        ],
        [
            "Rule 42–44",
            "Interlineations; maps; duplicate presentation",
            "Document integrity checks at intake",
        ],
        [
            "Rule 45",
            "Endorsement — date, hour, place of presentation; presentant signature",
            "Presentation timestamp — status 'Presented'; audit trail",
        ],
        [
            "Rule 46",
            "Suspension — delay fine; impounding for insufficient stamp",
            "Status 'Suspended (fine)' / 'Impounded (stamp)'",
        ],
        [
            "Rule 51–55",
            "Registration on payment of fines; rate of fines; condonation of delay",
            "Fine calculation display; condonation application status",
        ],
    ]
    add_table(doc, rules_ch9[0], rules_ch9[1:])

    add_heading(doc, "3.3.4 Examination of parties, endorsements and certificates (Chapters XII, XVI)", 3)
    rules_ch12_16 = [
        ["Rule", "Requirement", "System feature (Sr.12)"],
        [
            "Rule 71–72",
            "Persons executing documents; place of registration",
            "Executant identification; jurisdiction confirmation",
        ],
        [
            "Rule 73",
            "Duties of Registering Officer — read document to executant",
            "SRO examination checklist; vernacular reading log",
        ],
        [
            "Rule 78–79",
            "Thumb impression of every person; manner of taking",
            "Biometric capture at presentation",
        ],
        [
            "Rule 94",
            "Manner of endorsing and certifying",
            "Standard endorsement templates on document",
        ],
        [
            "Rule 104",
            "Certificate under Section 60",
            "Registration certificate generation — status 'Registered'",
        ],
        [
            "Rule 107",
            "Endorsement when execution denied",
            "Status 'Execution denied'",
        ],
    ]
    add_table(doc, rules_ch12_16[0], rules_ch12_16[1:])

    add_heading(doc, "3.3.5 Receipts, return and appeals (Chapters XVII, XXV)", 3)
    rules_ch17_25 = [
        ["Rule", "Requirement", "System feature (Sr.12)"],
        [
            "Rule 110–112",
            "Receipts for documents and fees; procedure for obtaining registered document",
            "Citizen receipt; document collection status",
        ],
        [
            "Rule 113",
            "Procedure on loss of receipt",
            "Recovery workflow for document return",
        ],
        [
            "Rule 175–176",
            "Appeal against refusal; by whom preferred",
            "Appeal initiation from refused status",
        ],
        [
            "Rule 179–188",
            "Procedure of disposing appeal; orders directing registration",
            "Appeal outcome status transitions",
        ],
    ]
    add_table(doc, rules_ch17_25[0], rules_ch17_25[1:])

    # 3.4 Notifications
    add_heading(doc, "3.4 Relevant notifications issued by the Department for Document Registration", 2)
    add_para(
        doc,
        "Gazette notifications, Government Orders and amendments cited in the "
        "Karnataka Registration Rules or issued under the Registration Act, 1908 and "
        "Karnataka Stamp Act, 1957. Source folder: Acts_Rules/Document/."
    )
    notif_rows = [
        ["Instrument", "Date / No.", "Effect", "BRD relevance (Sr.12)", "Source file"],
        [
            "The Karnataka Registration Rules, 1965",
            "Original notification under Sec. 69, Registration Act",
            "Comprehensive procedural rules for all registration offices",
            "Foundation for presentation, register, endorsement, appeal procedures",
            "The Karnataka Registration Rules 1965.pdf",
        ],
        [
            "Karnataka Registration (Amendment) Rules, 1971",
            "Various GSR notifications cited in Rules header",
            "Amendments to property description, presentation, power-of-attorney rules",
            "Rule 13–15 property description; Rule 44 duplicate presentation",
            "Embedded in Karnataka Registration Rules 1965.pdf",
        ],
        [
            "Notification No. RGN 2/2002-03",
            "1 Apr 2002; w.e.f. 4 Apr 2002",
            "Inserts Rule 19-A (document sheets); Rules 22-A–22-C; Rule 40 photograph "
            "sub-rules; digital photograph option",
            "Document format; photo at presentation; electronic copy workflow",
            "Cited in Karnataka Registration Rules 1965.pdf",
        ],
        [
            "Notification No. RGN/287/02-03",
            "29 Mar 2003; w.e.f. 29 Mar 2003",
            "Inserts Rule 19-A — document sheet specifications",
            "Presentation document paper / format standards",
            "Cited in Karnataka Registration Rules 1965.pdf",
        ],
        [
            "RD 403 ESR 85",
            "27 May 1986",
            "Table of Registration Fees under Sec. 78 — Articles I, III",
            "Base registration fee schedule (superseded in part by 2025 notification)",
            "Referenced in RegistrationActNotification.pdf",
        ],
        [
            "RD/46/MNMU/2025",
            "29 Aug 2025; w.e.f. 31 Aug 2025",
            "Substitutes registration fee: Art. I(4)(a) ₹10→₹20 per ₹1,000; "
            "Art. III(a)(i)/(ii) ₹1→₹2",
            "Registration fee at presentation; appointment fee display",
            "RegistrationActNotification.pdf",
        ],
        [
            "RD 380 MUNOMU 2008",
            "8 Apr 2009",
            "Karnataka Stamp (Payment of Duty by Means of e-Stamping) Rules, 2009",
            "e-Stamp certificate at presentation; CRA / ACC integration touchpoint",
            "THE KARNATAKA STAMP- Payment of Stamp Duty by means of e-Stamping.pdf",
        ],
        [
            "Karnataka Stamp Rules, 1958 — commencement",
            "G.S.R. notifications cited in Rules header (1960–2003)",
            "Stamp rule framework for instruments presented at SRO",
            "Stamp impounding reference at Rule 46",
            "Karnataka Stamps Rules 1958.pdf",
        ],
        [
            "Karnataka Stamp (Franking Impression Of Stamps) Rules, 2000",
            "Under Karnataka Stamp Act",
            "Franking at registration offices",
            "Stamp payment method at intake",
            "Karnataka Stamp (Franking Impression Of Stamps) Rules, 2000.docx",
        ],
        [
            "The Registration (Karnataka Amendment) Act, 2023",
            "Karnataka Gazette Extra-ordinary No. 480; 19 Oct 2024; "
            "Presidential assent 8 Oct 2024",
            "Secs. 22-B–22-D, 81-A–81-B — forged / prohibited document controls",
            "Mandatory refusal and cancellation workflows in registration core",
            "TheRegistration(KarnatakaAmendment)Act2023(47of2024).pdf",
        ],
        [
            "Karnataka Registration (Amendment) Acts — 1975 to 2002",
            "Various Karnataka Acts in folder",
            "State amendments to Registration Act — jurisdiction, fees, computerisation",
            "Historical / enabling amendments referenced in registration core",
            "Registration(KarnatakaAmendment)Act*.pdf",
        ],
    ]
    add_table(doc, notif_rows[0], notif_rows[1:])

    doc.save(str(DST))
    print(f"Wrote {DST}")


if __name__ == "__main__":
    build()
