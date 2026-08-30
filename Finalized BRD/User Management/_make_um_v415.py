# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v4.15.docx from v4.14.

Align the cover / front matter with BRD_Marriage_v1.10.docx: Heading 1 title,
Heading 2 module name, Field/Value document-control table, Related documents.
"""
from __future__ import annotations

import copy
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(__file__).resolve().parent
SRC = BASE / "BRD_User_Management_v4.14.docx"
MARRIAGE = (
    BASE.parent / "Marriage" / "RFP" / "BRD_Marriage_v1.10.docx"
)
DST = BASE / "BRD_User_Management_v4.15.docx"

FIELD_ROWS = [
    ("Field", "Value"),
    ("Document ID", "BRD-K3-UM-001"),
    ("Version", "4.15"),
    ("Status", "In review — pending Domain Expert sign-off"),
    ("Module", "User Management"),
    (
        "Legal basis (primary)",
        "Information Technology Act, 2000; Indian Registration Act, 1908 "
        "(appointment of Sub-Registrars); Aadhaar Act, 2016 (where biometric / "
        "Aadhaar is used)",
    ),
    (
        "State rules (primary)",
        "Karnataka e-Governance hosting and security norms; MeitY / CERT-In / "
        "STQC / GIGW; Government Orders for office and post creation",
    ),
    ("Author (BA)", "Nandha Kumar"),
    ("Product Owner", "Prashanth"),
    ("Domain expert / reviewer", "Prabhakar Naik"),
    (
        "Target audience",
        "Kaveri IT Cell, Department of Stamps and Registration, Government of Karnataka",
    ),
    ("Last updated", "2026-08-30"),
]

RELATED_ROWS = [
    ("ID", "Title", "Link"),
    ("BRD-K3-UM-001", "This document", ""),
    (
        "PROC-K3-UM-TOBE-001",
        "Process flows",
        "ProcessDiagrams/User_Management/ (P-01–P-13, S-01–S-06)",
    ),
    (
        "ERD-K3-UM-001",
        "Entity-relationship diagrams",
        "ERD_User_Management_v1.0.docx",
    ),
    (
        "RTM-K3-UM-001",
        "Requirements traceability",
        "Functional requirements §6 of this document",
    ),
]


def shade_cell(cell: _Cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shading)


def set_cell_text(cell: _Cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def add_table_row(table: Table, values: list[str]) -> None:
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    last_tr.addnext(new_tr)
    row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell_text(row.cells[i], val)


def clear_center(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return
    jc = p_pr.find(qn("w:jc"))
    if jc is not None:
        p_pr.remove(jc)


def set_heading(paragraph: Paragraph, text: str, style: str) -> None:
    clear_center(paragraph)
    try:
        paragraph.style = style
    except KeyError:
        pass
    replace_paragraph_text(paragraph, text)
    for run in paragraph.runs:
        run.bold = None
        run.font.size = None


def fill_cloned_table(src_tbl_el, parent, rows: list[tuple[str, ...]]) -> None:
    """Clone extra rows if needed, then write cell text preserving formatting."""
    trs = src_tbl_el.findall(qn("w:tr"))
    while len(trs) < len(rows):
        src_tbl_el.append(copy.deepcopy(trs[-1]))
        trs = src_tbl_el.findall(qn("w:tr"))
    while len(trs) > len(rows):
        src_tbl_el.remove(trs[-1])
        trs = src_tbl_el.findall(qn("w:tr"))

    tbl = Table(src_tbl_el, parent)
    for ri, values in enumerate(rows):
        header = ri == 0
        for ci, val in enumerate(values):
            if ci >= len(tbl.rows[ri].cells):
                continue
            cell = tbl.rows[ri].cells[ci]
            paras = cell.paragraphs
            if not paras:
                cell.add_paragraph(val)
                continue
            p = paras[0]
            if p.runs:
                p.runs[0].text = val
                p.runs[0].bold = True if header else bool(p.runs[0].bold) and ci == 0
                if header:
                    p.runs[0].bold = True
                else:
                    p.runs[0].bold = False
                for extra in p.runs[1:]:
                    extra.text = ""
            else:
                run = p.add_run(val)
                run.bold = header
            for extra_p in paras[1:]:
                for run in extra_p.runs:
                    run.text = ""
        if header:
            for cell in tbl.rows[ri].cells:
                shade_cell(cell, "D9E2F3")


def drop_tbl_style(tbl_el) -> None:
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        return
    style = tbl_pr.find(qn("w:tblStyle"))
    if style is not None:
        style.set(qn("w:val"), "TableNormal")


def insert_after(ref_el, new_el) -> None:
    ref_el.addnext(new_el)


def build() -> Document:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if not MARRIAGE.exists():
        raise FileNotFoundError(MARRIAGE)

    shutil.copy2(SRC, DST)
    doc = Document(str(DST))
    marriage = Document(str(MARRIAGE))

    # --- Cover headings (Marriage layout) ---
    paras = doc.paragraphs
    set_heading(paras[0], "", "Heading 1")
    set_heading(paras[1], "Business Requirements Document (BRD)", "Heading 1")
    set_heading(paras[2], "User Management Module", "Heading 2")

    # Remove Version / Date / Prepared-by lines from the old centred cover
    to_remove = []
    for p in paras[3:7]:
        t = p.text.strip()
        if t.startswith("Version ") or t.startswith("Date:") or t.startswith("Prepared by:"):
            to_remove.append(p._element)
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # --- Field / Value table: clone Marriage table 0 ---
    field_tbl_el = copy.deepcopy(marriage.tables[0]._tbl)
    drop_tbl_style(field_tbl_el)

    old_tbl = doc.tables[0]._tbl
    old_tbl.addnext(field_tbl_el)
    old_tbl.getparent().remove(old_tbl)
    fill_cloned_table(field_tbl_el, doc, FIELD_ROWS)

    # --- Related documents (Marriage cover) ---
    related_tbl_el = copy.deepcopy(marriage.tables[2]._tbl)
    drop_tbl_style(related_tbl_el)

    rel_label = OxmlElement("w:p")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_pr.append(OxmlElement("w:b"))
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = "Related documents:"
    r.append(t)
    rel_label.append(r)

    insert_after(field_tbl_el, rel_label)
    insert_after(rel_label, related_tbl_el)
    fill_cloned_table(related_tbl_el, doc, RELATED_ROWS)

    # --- Revision history (after cover tables) ---
    rev = None
    for tbl in doc.tables:
        if (
            len(tbl.columns) >= 4
            and tbl.rows
            and tbl.rows[0].cells[0].text.strip() == "Version"
            and tbl.rows[0].cells[3].text.strip() in ("Description", "Summary of change")
        ):
            rev = tbl
            break
    if rev is None:
        raise RuntimeError("revision history table not found")
    add_table_row(
        rev,
        [
            "4.15",
            "30-Aug-2026",
            "Nandha Kumar",
            "Cover / front matter aligned to Marriage BRD v1.10: Heading 1 "
            "Business Requirements Document (BRD), Heading 2 User Management "
            "Module, Field/Value document-control table, Related documents",
        ],
    )

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v4.15"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    claude_dir = BASE.parent.parent / "Claude"
    if claude_dir.is_dir():
        claude_dst = claude_dir / target.name
        try:
            shutil.copy2(target, claude_dst)
            print(f"Mirrored: {claude_dst}")
        except Exception as exc:
            print(f"Claude mirror skipped: {exc}")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
