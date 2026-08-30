# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v4.5.docx — consistency fixes (review items 7–13);
placeholder OTP/session FRs (item 14)."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(__file__).resolve().parent
TEMPLATE = BASE / "Template" / "User_Management_Module_BRD_Template.docx"
DST = BASE / "BRD_User_Management_v4.5.docx"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_table_rows(table, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "D9E2F3")
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(cells[i], val)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph_after(ref_paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    para = Paragraph(new_p, ref_paragraph._parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_table_after(ref_paragraph, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_paragraph._element.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def insert_after_element(ref_el, parent, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_el.addnext(new_p)
    para = Paragraph(new_p, parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_heading_after(ref, text: str, style: str = "Heading 3", parent=None) -> Paragraph:
    if isinstance(ref, Paragraph):
        return insert_paragraph_after(ref, text, style)
    if parent is None:
        raise ValueError("parent required when ref is an XML element")
    return insert_after_element(ref, parent, text, style)


def insert_table_after_ref(ref_el, parent, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_el.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def build() -> Document:
    doc = Document(str(TEMPLATE))

    cover_map = {
        1: "BUSINESS REQUIREMENTS DOCUMENT",
        2: "User Management Module — KAVERI 3.0",
        3: "Prepared for: Department of Stamps & Registration, Government of Karnataka",
        4: "Version 4.5",
        5: "Date: 30 August 2026",
        6: "Prepared by: Nandha Kumar",
    }
    for idx, text in cover_map.items():
        replace_paragraph_text(doc.paragraphs[idx], text)

    replace_paragraph_text(
        doc.paragraphs[12],
        "This Business Requirements Document (BRD) defines the business requirements for the "
        "User Management Module of KAVERI 3.0 — the integrated platform of the Department of "
        "Stamps and Registration (DSR), Government of Karnataka. It describes a single unified "
        "Role Master and User Master differentiated by user category (Public/Citizen, DSR Officer, "
        "Other Department), OTP-based authentication (Username-based login — preferred Username for "
        "Citizens, KGID for departmental users; login OTP to mobile only; no password management), "
        "a separate Posts Master "
        "with Post–Role mapping (one post may map to multiple roles), Sanctioned Posts Master "
        "(Post + Office + strength), DSR Officer assignment to one or more sanctioned posts with "
        "available capacity (roles via Post–Role mapping; no Primary/Secondary assignment; multi-post "
        "login selection with office details), two distinct vacancy tests — available capacity for "
        "assignment and Transfer In, and a wholly unoccupied post for post-login additional charge (FR-53) — "
        "single-role assignment "
        "for Other Department users with optional account "
        "end date, a Module Master (Registration of Documents, Marriage Registration, Encumbrance Search, "
        "Certified Copy, and others), Role-to-Module-Function privilege mapping, Module Function and "
        "Resource (API/URL) masters, runtime access enforcement, DSR officer organisational hierarchy "
        "maintenance, and dedicated user-creation workflows. It serves as the agreed basis for design, "
        "development, testing, and sign-off.",
    )
    replace_paragraph_text(
        doc.paragraphs[14],
        "KAVERI 3.0 requires a centralized mechanism to manage user identities, roles, sanctioned "
        "posts, and access permissions. There shall be one Role Master and one User Master for all "
        "categories — Public users (Citizens), Department users (DSR Officers), and Other Department "
        "users — differentiated by User Category (and Role Category on roles). Authentication shall "
        "be passwordless for all three categories: Username + Captcha + OTP sent "
        "only to the registered mobile; DSR Officers and Other Department users additionally verify "
        "Biometrics. Citizens choose their own preferred Username at registration, where both the email "
        "address and the mobile number are verified by separate OTPs before the account is created; "
        "KGID is the Username for DSR Officers and Other Department users, who record official email "
        "IDs. The Username is unique across the entire User Master and is the only unique login "
        "identifier — email and mobile number need not be unique. A Citizen who has lost the registered "
        "mobile may reset it before login by correctly answering three security questions selected at "
        "random from the five captured at registration and entering a PIN sent to the registered email; "
        "this reset path is not available to DSR Officers or Other Department users, whose mobile number "
        "is changed only by an authorised administrator. DSR officers are assigned to sanctioned "
        "posts from the Posts Master that still have capacity (strength per office in Sanctioned Posts "
        "Master). Posts are mapped "
        "to one or more roles in the Role Master via Post–Role mapping; a user may hold multiple "
        "sanctioned posts with no Primary/Secondary distinction and shall select one post (with office "
        "details) at login when multiple are active. Post assignment and Transfer In require available "
        "capacity at the office (Occupied < Sanctioned Strength), except that a future-dated Transfer "
        "In may reserve capacity when relieving is already recorded for that Post + Office (FR-67). "
        "Additional charge under a "
        "subordinate post after login requires that post to be wholly unoccupied at the office "
        "(Occupied = 0). After login, an officer may take additional charge of such a subordinate post "
        "without logout or order (FR-53). A scheduled job after midnight refreshes post assignments "
        "and Sanctioned Posts occupied counts from Transfer Out, Transfer In, and occupancy End Date "
        "(FR-68). Other Department users are assigned "
        "exactly one role and may have an optional account end "
        "date that deactivates the user when reached. Application modules (e.g. Registration of Documents, "
        "Marriage Registration, Encumbrance Search, Certified Copy) are maintained in a Module Master "
        "and mapped to roles via Module Functions and Resources (APIs/URLs) — DSR organisational roles "
        "are not named after modules. The system shall maintain the DSR officer reporting hierarchy "
        "(organisational chart). Application Admin maintains these masters. Password management "
        "is explicitly out of scope.",
    )

    scope_in = {
        17: "User registration for three categories in a single User Master: Public users (Citizens), Department users (DSR Officers), and Other Department users",
        18: "Username-based login for all categories (preferred Username for Citizens, KGID for departmental users); OTP to mobile only; Captcha; biometrics for departmental users; OTP and session policy (FR-69–FR-76); Citizen-only lost-mobile reset via three random security questions plus PIN to registered email; FR-52 assigned-post login; FR-53 post-login additional charge; FR-54 display",
        19: "Single unified Role Master with Role Category differentiating Citizen, DSR, and Other Department roles (RBAC)",
        20: "Module Master, Module Function Master, Resource Master (API/URL), and Role–Module–Function mapping maintained by Application Admin",
        21: "DSR Officer Hierarchy Master (posts) and Office Hierarchy Master (IGR Head Office → District Registrar Offices → Sub-Registrar Offices)",
        22: "Transfer Out / Relieving and Transfer In (office span then immediate-parent post only; future-dated Transfer In may reserve capacity when relieving is already recorded; Transfer/Reporting Order + Joining Date; login from 12:00 AM on Joining Date); midnight occupancy refresh job; office span via Office Hierarchy",
        23: "Dedicated step-by-step workflows for role assignment during user creation (DSR and Other Department)",
    }
    for idx, text in scope_in.items():
        replace_paragraph_text(doc.paragraphs[idx], text)
    admin_scope = insert_paragraph_after(
        doc.paragraphs[23],
        "Administrative user management (create, edit, suspend, deactivate users)",
        "List Paragraph",
    )
    insert_paragraph_after(
        admin_scope,
        "Audit logging of user-related activities",
        "List Paragraph",
    )

    for p in doc.paragraphs:
        if "password resets" in p.text.lower():
            replace_paragraph_text(
                p,
                "Reduce support overhead related to account access issues through reliable OTP delivery.",
            )

    for p in doc.paragraphs:
        if "Eliminate password-related" in p.text:
            remove_paragraph(p)
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.3 Password Management":
            replace_paragraph_text(p, "6.3 OTP Login and Account Recovery")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.6 Administrative User Management":
            replace_paragraph_text(p, "6.7 Administrative User Management")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "Enforce role-based access control to protect sensitive data and functionality.":
            insert_paragraph_after(
                p,
                "Support Module Function and Resource (API/URL) masters with Role–Module–Function mapping and runtime access enforcement maintained by Application Admin.",
                "List Paragraph",
            )
            insert_paragraph_after(
                p,
                "Maintain the DSR Officer Hierarchy Master aligned to the Department organisational chart.",
                "List Paragraph",
            )
            break

    for p in doc.paragraphs:
        if p.text.strip() == "The system shall provide a report of role and permission assignments across all users.":
            sp_note = insert_paragraph_after(
                p,
                "The system shall provide a sanctioned post occupancy report showing, for each office, "
                "the number of sanctioned posts per Post (from Posts Master), occupied count, remaining "
                "capacity, and whether the post is wholly unoccupied (FR-66).",
                "List Paragraph",
            )
            rm_note = insert_paragraph_after(
                sp_note,
                "The system shall provide a Role-to-Module mapping report showing which modules and "
                "functions are assigned to each role.",
                "List Paragraph",
            )
            ac_note = insert_paragraph_after(
                rm_note,
                "The system shall provide a contact-change and recovery report over a selected date range "
                "covering Citizen lost-mobile resets (FR-56), administrator-initiated mobile changes for "
                "departmental users (FR-65), and email changes — showing actor, user, timestamp, outcome, "
                "and failed attempt counts.",
                "List Paragraph",
            )
            job_note = insert_paragraph_after(
                ac_note,
                "The system shall provide an additional charge report (FR-53) showing, for each officer "
                "and date range, when additional charge was taken or cleared, the assigned post, the "
                "subordinate post, office, and audit timestamp.",
                "List Paragraph",
            )
            insert_paragraph_after(
                job_note,
                "The system shall provide an occupancy-refresh report (FR-68) showing, for each midnight "
                "job run, occupancies de-allocated (relieving / End Date), reserved Transfer In "
                "occupancies activated, and Sanctioned Posts occupied-count changes (Post + Office, "
                "before/after counts).",
                "List Paragraph",
            )
            break

    replace_table_rows(
        doc.tables[0],
        ["Document Control", ""],
        [
            ("Status", "In review — pending Domain Expert sign-off"),
            ("Prior versions", "Versions 1.0–1.6 predate template alignment (Aug 2026); revision history begins at v1.7"),
            ("Owner", "Nandha Kumar (BA) / Prashanth (PO)"),
            ("Reviewers", "Prabhakar Naik (Domain Expert); Kaveri IT Cell"),
            ("Approved By", "Prashanth (Product Owner)"),
        ],
    )

    replace_table_rows(
        doc.tables[1],
        ["Version", "Date", "Author", "Description"],
        [
            (
                "1.7",
                "28-Aug-2026",
                "Nandha Kumar",
                "Aligned to User Management BRD template; three user categories with OTP-only "
                "authentication (no password management); DSR division and role catalogue",
            ),
            (
                "1.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added sanctioned posts master: capture number of sanctioned posts (approved strength) "
                "for each role at each office; DSR user mapping to sanctioned posts only; "
                "over-capacity blocking and occupancy reporting",
            ),
            (
                "1.9",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Other Department role catalogue; Primary and Secondary role assignment; "
                "dedicated step-by-step user-creation workflows for DSR and Other Department users",
            ),
            (
                "2.0",
                "29-Aug-2026",
                "Nandha Kumar",
                "Primary/Secondary roles apply to DSR Officers only; Other Department users get "
                "exactly one role; optional account end date on Other Department user creation "
                "deactivates the user when reached (§6.6.2)",
            ),
            (
                "2.1",
                "29-Aug-2026",
                "Nandha Kumar",
                "Unified single Role Master and User Master for Citizens, DSR Officers, and Other "
                "Department users — differentiated by User Category / Role Category; no separate "
                "masters per category",
            ),
            (
                "2.2",
                "29-Aug-2026",
                "Nandha Kumar",
                "Listed seed roles in Role Master for Role Category = Citizen and Role Category = "
                "Other Department (§6.5.4)",
            ),
            (
                "2.3",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added DSR Department roles listing to §6.5.4 alongside Citizen and Other Department roles",
            ),
            (
                "2.4",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Module Master and Role-to-Module mapping (§6.5.6): Document Registration, "
                "Marriage Registration, Encumbrance Search, Certified Copy and related modules "
                "mapped to organisational / Citizen / Other Department roles",
            ),
            (
                "2.5",
                "29-Aug-2026",
                "Nandha Kumar",
                "Removed online/offline classification for document registration in §6.5.6 — "
                "module described as Registration of Documents only",
            ),
            (
                "2.6",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Module Function Master, Resource Master (API/URL), Role–Module–Function "
                "mapping with example rows, Application Admin CRUD, and runtime enforcement (§6.5.6)",
            ),
            (
                "2.7",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added DSR Officer Hierarchy Master (§6.5.7) to maintain Department officer reporting "
                "structure aligned to DSR organisational chart (Admin/Law/Computers, Vigilance, "
                "Computers, Enforcement, Intelligence & Audit, DIGR CVC)",
            ),
            (
                "2.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "DSR Officers may be assigned to multiple sanctioned posts when each post is vacant "
                "(Primary plus optional additional posts); over-capacity still blocked",
            ),
            (
                "2.9",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added separate Posts Master; Post–Role mapping (one post to multiple roles); "
                "Sanctioned Posts Master references Posts Master (Post + Office + strength)",
            ),
            (
                "3.0",
                "29-Aug-2026",
                "Nandha Kumar",
                "Removed Primary and Secondary role/post assignment; DSR Officers are assigned one or "
                "more vacant sanctioned posts equally; roles derived only via Post–Role mapping",
            ),
            (
                "3.1",
                "29-Aug-2026",
                "Nandha Kumar",
                "Resolved structural conflicts: division-specific Posts (no generic DIGR/AIGR over-provisioning); "
                "Post–Role mapping typically 1:1 to matching unique roles; Hierarchy Master references Posts "
                "not Roles; ACS/Secretary root role added; unique FDA/SDA/Typist role names per division",
            ),
            (
                "3.2",
                "29-Aug-2026",
                "Nandha Kumar",
                "Role–Module–Function mapping aligned to unique Role Master names (FDA (Enforcement) not bare FDA); "
                "FR-50 referential integrity for Role–Module–Function ↔ Role Master; Application Admin treated as "
                "system-level actor outside Role Master (Option B — not a DSR seed role); "
                "FR-31 retired — was tied to Primary/Secondary (additional) post assignment with mandatory end date, "
                "removed in v3.0 (gap left intentional for audit traceability)",
            ),
            (
                "3.3",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added FR-52: after authentication, if a DSR Officer has multiple active post occupancies, "
                "the user must select one post for the session; each choice is displayed with office details; "
                "session roles/privileges derive from the selected post via Post–Role mapping",
            ),
            (
                "3.4",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added FR-53: after login and post selection, DSR Officer may switch session to vacant "
                "subordinate posts under the login-selected post using Hierarchy Master (cascade through "
                "vacant nodes, same office); example SR → vacant FDA → vacant SDA; Enforcement seed "
                "hierarchy aligned so FDA/SDA report under Sub-Registrar",
            ),
            (
                "3.5",
                "29-Aug-2026",
                "Nandha Kumar",
                "FR-53 revised: vacant subordinate post/role selection occurs only during login (after "
                "authentication and FR-52 post selection, before home); after login completes the user "
                "cannot switch post for the remainder of the session",
            ),
            (
                "3.6",
                "29-Aug-2026",
                "Nandha Kumar",
                "FR-53 UI: during login post/role switch, options shown in a dropdown as Role with Post "
                "(and Office); FR-54: after login, home/header displays the logged-in Post with mapped Role(s)",
            ),
            (
                "3.7",
                "29-Aug-2026",
                "Nandha Kumar",
                "Login for all categories uses Email as Username; OTP sent only to registered mobile "
                "(not email); mobile updatable after login; lost-mobile recovery before login via "
                "security question selected at registration (FR-55, FR-56); FR-04/05/06/07/10/12/13 updated",
            ),
            (
                "3.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Transfer out / relieving process (FR-57, FR-58): only hierarchy superior may relieve "
                "an officer from assigned post(s); capture Relieving Date and Relieving Order; system "
                "de-allocates user–post mapping after 11:59 PM of the Relieving Date; §6.6.3 workflow "
                "with hierarchy examples (e.g. DIGR Enforcement relieves Sub-Registrar; SR relieves FDA)",
            ),
            (
                "3.9",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Office Hierarchy Master (FR-59): IGR Head Office → District Registrar Offices → "
                "Sub-Registrar Offices; Transfer Out / Relieving (FR-57) displays only offices and post "
                "occupancies under the actor's session office; Enforcement post chain aligned DRO→SR→FDA",
            ),
            (
                "4.0",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Transfer In (FR-60, FR-61): capture Transfer/Reporting Order and Joining Date; "
                "assign only to vacant post within office span; user may login only from 12:00 AM on "
                "the Joining Date; §6.6.4 workflow",
            ),
            (
                "4.1",
                "30-Aug-2026",
                "Nandha Kumar",
                "Email is no longer the Username. Login for all categories uses Username (FR-62): "
                "Citizens choose a preferred Username at registration; KGID is the Username for DSR "
                "Officers and Other Department users. Username is unique across the whole User Master; "
                "email and mobile number uniqueness removed. Citizen registration verifies both email "
                "and mobile by separate OTPs (FR-63) and captures five security questions (FR-55). "
                "Citizen lost-mobile reset asks three of the five questions at random and additionally "
                "requires a PIN sent to the registered email (FR-56); mobile reset is not offered to "
                "DSR Officers or Other Department users — their mobile is changed only by an authorised "
                "administrator (FR-65). Departmental users record official email IDs (FR-64). "
                "Vacancy is now defined in two distinct ways (FR-66): available capacity "
                "(Occupied < Sanctioned Strength) governs assignment and Transfer In, while FR-53 "
                "login-time subordinate selection requires an unoccupied post (Occupied = 0). "
                "FR-01–FR-07, FR-10, FR-12, FR-13, FR-27, FR-38, FR-45, FR-52–FR-56, FR-60 updated",
            ),
            (
                "4.2",
                "30-Aug-2026",
                "Nandha Kumar",
                "FR-53 revised: login completes at assigned post only (FR-52) — no subordinate selection "
                "during login. After login, a DSR Officer may take additional charge of a wholly unoccupied "
                "subordinate post at the same office without logout and without an order (in-office "
                "additional charge for the session). Session privileges are the union of the assigned "
                "post and the additional charge post (FR-38). Officer may revert to assigned-post-only "
                "without logout. FR-54 updated to show assigned post and optional additional charge. "
                "§6.5.2 workflows updated",
            ),
            (
                "4.3",
                "30-Aug-2026",
                "Nandha Kumar",
                "FR-38 / FR-53 clarified: when additional charge is active, session Module Function "
                "claims derive from the additional charge post only — the officer does not retain "
                "assigned-post privileges (e.g. an SR acting as FDA loses SR signing until switching "
                "back to assigned-post context without logout). When additional charge is cleared, "
                "claims revert to the assigned post.",
            ),
            (
                "4.4",
                "30-Aug-2026",
                "Nandha Kumar",
                "FR-57 office-span examples narrowed so seeing descendant offices does not imply "
                "authority over posts that do not report immediately to the actor (e.g. IGR/DIGR "
                "cannot relieve or Transfer In a Sub-Registrar). FR-60 / FR-67: future-dated Transfer "
                "In may be recorded when relieving is already captured, reserving capacity until "
                "Joining Date. FR-68: scheduled job after midnight refreshes post assignments from "
                "Transfer Out / Transfer In / occupancy End Date and recalculates Sanctioned Posts "
                "occupied counts. FR-45, FR-58, FR-61, FR-66, §6.5.8, §6.6.3–6.6.4 updated",
            ),
            (
                "4.5",
                "30-Aug-2026",
                "Nandha Kumar",
                "Consistency fixes (review items 7–13): FR-50 extended to Post Name integrity across "
                "Posts Master, Post–Role mapping, and Hierarchy; SRO role names renamed to Sub Registrar "
                "(Admin/Computers); Sanctioned Posts reference Office Code; FR-47 blocks unmapped posts; "
                "FR-30 End Date restricted to deputations; FR-41 deny-by-default with Resource Master "
                "Is Public and path-match precedence; FR-15 retired into FR-20; unified UI labels for "
                "FR-52/53/54; DEO under Division 4 — Enforcement. Placeholder OTP/session policy "
                "FR-69–FR-76 added.",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[2],
        ["Name / Role", "Department", "Responsibility"],
        [
            ("Prashanth", "DSR / Product", "Prioritizes requirements; approves scope"),
            ("Nandha Kumar", "Business Analysis", "Documents and validates requirements"),
            ("Prabhakar Naik", "Domain Expert", "Validates DSR roles, Other Department roles, sanctioned posts, and organizational structure"),
            ("Kaveri IT Cell", "Engineering", "Reviews technical feasibility"),
            ("Citizens (Public users)", "External", "Self-register and access citizen portal services"),
            (
                "DSR Officers & Other Department users",
                "Government",
                "Access departmental modules via OTP + biometrics",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[3],
        ["ID", "Requirement", "Priority"],
        [
            (
                "FR-01",
                "The system shall support instant self-registration for Public users (Citizens) with no "
                "approval workflow. Registration shall capture a preferred Username chosen by the "
                "registering user (FR-62), email address, mobile number, and five security questions "
                "with answers (FR-55). Both the email address and the mobile number shall be verified "
                "by separate OTPs before the account is created (FR-63).",
                "High",
            ),
            (
                "FR-02",
                "The system shall allow Department users (DSR Officers) to be created only by authorised "
                "administrative roles. Creation shall capture KGID — which shall be the Username for this "
                "category (FR-62, FR-64) — official email ID, and mobile number. Security questions are "
                "not captured for this category because no self-service mobile reset is offered (FR-65).",
                "High",
            ),
            (
                "FR-03",
                "The system shall allow Other Department users (officers/staff from other government "
                "departments) to be created only by authorised administrative roles. Creation shall "
                "capture KGID — which shall be the Username for this category (FR-62, FR-64) — official "
                "email ID of the parent department, and mobile number. Security questions are not "
                "captured for this category because no self-service mobile reset is offered (FR-65).",
                "High",
            ),
            (
                "FR-04",
                "The Username shall be the login identifier for all user categories. The Username shall "
                "be unique across the entire User Master (a single namespace covering Citizens, DSR "
                "Officers, and Other Department users) and the system shall reject creation of a "
                "duplicate Username. Email address and mobile number shall not be subject to a "
                "uniqueness constraint — the same email or mobile may legitimately appear on more than "
                "one account (for example family members sharing a mobile, or an office landline). "
                "See FR-62.",
                "High",
            ),
            (
                "FR-62",
                "The system shall derive the Username by user category: for Public users (Citizens) the "
                "Username is a preferred Username entered by the registering user, and the system shall "
                "check availability against the whole User Master and reject a Username already in use; "
                "for DSR Officers and Other Department users the KGID shall be treated as the Username "
                "and shall not be user-selectable. The Username shall not be changeable by the user "
                "after creation. A correction to a Username (including a KGID correction) shall be "
                "performed only by an authorised administrator, with reason and audit trail.",
                "High",
            ),
            (
                "FR-63",
                "During Citizen self-registration the system shall verify both contact channels before "
                "creating the account: an OTP shall be sent to the entered email address and a separate "
                "OTP shall be sent to the entered mobile number, and both shall be entered correctly. "
                "If either verification is not completed, the account shall not be created. Each "
                "verification OTP shall be time-limited and single-use, and failed attempts shall be "
                "rate-limited and audit-logged.",
                "High",
            ),
            (
                "FR-64",
                "For DSR Officers and Other Department users the system shall validate that the KGID is "
                "present and unique in the User Master, and shall require an official email ID "
                "(government / parent-department email address) rather than a personal email address. "
                "Where a permitted official email domain list is configured, the system shall validate "
                "the entered email against that list and reject non-official domains.",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[4],
        ["ID", "Requirement", "Priority"],
        [
            (
                "FR-05",
                "Public users (Citizens) shall authenticate using Username + Captcha + OTP "
                "on every login. OTP shall be sent only to the registered mobile number.",
                "High",
            ),
            (
                "FR-06",
                "Department users (DSR Officers) shall authenticate using Username (KGID) + Captcha + "
                "OTP + Biometrics on every login. OTP shall be sent only to the registered mobile number.",
                "High",
            ),
            (
                "FR-07",
                "Other Department users shall authenticate using Username (KGID) + Captcha + OTP + "
                "Biometrics on every login. OTP shall be sent only to the registered mobile number.",
                "High",
            ),
            ("FR-08", "The system shall allow users to log out and terminate their active session.", "High"),
            ("FR-09", "The system shall not provide password-based login, password reset, or password change for any user category.", "High"),
            (
                "FR-69",
                "PLACEHOLDER — The system shall define and enforce a maximum validity period for login OTPs "
                "and registration/recovery OTPs (including Citizen email and mobile verification OTPs under "
                "FR-63 and reset PIN validity under FR-56). Exact duration(s) shall be configured by Kaveri "
                "IT Cell / security policy and documented before UAT; this BRD reserves the requirement "
                "only.",
                "High",
            ),
            (
                "FR-70",
                "PLACEHOLDER — The system shall define the OTP length (numeric digits) for login and "
                "registration/recovery flows. Exact length shall be configured before UAT.",
                "High",
            ),
            (
                "FR-71",
                "PLACEHOLDER — The system shall enforce a maximum number of incorrect OTP entry attempts "
                "per OTP issuance before that OTP is invalidated and the user must request a new one. "
                "Exact limit shall be configured before UAT.",
                "High",
            ),
            (
                "FR-72",
                "PLACEHOLDER — The system shall enforce a resend throttle / cooldown between successive OTP "
                "requests to the same user and channel (mobile or email). Exact interval shall be "
                "configured before UAT.",
                "High",
            ),
            (
                "FR-73",
                "PLACEHOLDER — The system shall define an account or flow lockout policy after repeated "
                "failed login or OTP verification attempts (separate from FR-56 lost-mobile reset lockout). "
                "Lockout duration, scope (per Username / per IP), and unlock procedure shall be "
                "configured before UAT.",
                "High",
            ),
            (
                "FR-74",
                "PLACEHOLDER — The system shall enforce a session idle timeout: after a configured period "
                "of no user activity the session shall terminate and re-authentication (including OTP, "
                "and biometrics for departmental users) shall be required. Exact idle period shall be "
                "configured before UAT; this is especially important for shared SRO counter machines "
                "where FR-52/FR-53 fix the session post for the whole session.",
                "High",
            ),
            (
                "FR-75",
                "PLACEHOLDER — The system shall enforce an absolute maximum session lifetime regardless "
                "of activity. After this period the session shall terminate even if the user is active. "
                "Exact lifetime shall be configured before UAT.",
                "High",
            ),
            (
                "FR-76",
                "PLACEHOLDER — The system shall define a concurrent-session rule per Username (e.g. allow "
                "only one active session, or allow multiple with last-login-wins / notify prior session). "
                "Exact rule shall be configured before UAT.",
                "High",
            ),
            (
                "FR-52",
                "After successful authentication, if a DSR Officer has more than one active sanctioned-post "
                "occupancy, the system shall present a mandatory post-selection step before entering the "
                "application. Each option shall be labelled consistently as "
                "\"Role — Post Name — Office Name (Office Code)\" (see FR-53, FR-54). The user shall "
                "select exactly one assigned post for the session. If the officer has only one active "
                "occupancy, the system shall auto-select that post and skip the selection screen. Login "
                "shall complete under that assigned post only — additional charge of a subordinate post "
                "(FR-53) shall not be offered during the login flow. Initial Module Function claims on "
                "entering the home page shall derive from the selected assigned post via Post–Role "
                "mapping only — not from other post occupancies the officer may hold concurrently. When "
                "a post maps to multiple roles, claims are the union of Module Functions for those "
                "roles of that single selected post (not a union across multiple assigned posts). "
                "Other Department and Citizen users are not subject to post selection.",
                "High",
            ),
            (
                "FR-53",
                "After login — while the DSR Officer remains in the same session and without requiring "
                "logout — the system shall allow the officer to take additional charge of a wholly "
                "unoccupied subordinate post under the officer's assigned post (the post selected or "
                "auto-selected at login under FR-52) at the same Office, using the DSR Officer Hierarchy "
                "Master (FR-43). This is temporary in-office additional charge for the session; no "
                "Transfer Order, Reporting Order, or other order document is required (distinct from "
                "Transfer In under FR-60). Eligibility: for the assigned Post at the same Office, examine "
                "each immediate child post in the hierarchy; that child post shall be offered only if it "
                "is wholly unoccupied at that Office — the occupied count is zero (FR-66(b)). A post "
                "that has any occupant shall not be offered, even where sanctioned strength has not been "
                "fully used. Where a child post is unoccupied, the system shall continue one level further "
                "under it (cascade while posts remain unoccupied); a post with any occupant shall block "
                "cascading beneath it. Example: an SR logged in at SRO Yeshwanthapura; if no FDA is posted "
                "there, additional charge as FDA is allowed; if no SDA is posted under that unoccupied "
                "FDA, additional charge as FDA or SDA is allowed; if even one FDA is posted, FDA is not "
                "offered and SDA beneath it is not reachable. The officer may hold at most one additional "
                "charge subordinate post at a time. While additional charge is active, effective Module "
                "Function claims shall be derived from the additional charge post via Post–Role mapping "
                "only — assigned-post privileges shall not apply (e.g. a Sub-Registrar acting as FDA under "
                "additional charge shall not retain Sub-Registrar signing or approval functions until "
                "switching back). The officer may switch back to assigned-post context at any time during "
                "the session without logout; on switch-back, claims revert to the assigned post. When no "
                "additional charge is active, claims derive from the assigned post only. Selection, "
                "switch-back, and reversion shall be audit-logged (actor, assigned post, additional "
                "charge post or cleared, office, timestamp). The additional-charge picker shall label "
                "each option as \"Role — Post Name — Office Name (Office Code)\" — the same format as "
                "FR-52 post selection and FR-54 header display.",
                "High",
            ),
            (
                "FR-54",
                "After login, the home page and/or application header shall display the assigned Post "
                "(from FR-52) with its mapped Role(s) and Office details using the format "
                "\"Role — Post Name — Office Name (Office Code)\" for the assigned post. When additional "
                "charge is active under FR-53, the header shall show which context is currently driving "
                "session privileges — either assigned post only, or additional charge post (with assigned "
                "post still shown for reference but not active for access). Example with additional charge "
                "active: \"Assigned: Sub-Registrar (SR) — Sub-Registrar — SRO Yeshwanthapura (OFF-SRO-YESH) "
                "| Active context: Additional charge — FDA (Enforcement) — FDA (Enforcement) — "
                "SRO Yeshwanthapura (OFF-SRO-YESH)\". The officer may switch between assigned-post "
                "context and additional charge, or clear additional charge, from this header or an "
                "equivalent in-application control without logging out.",
                "High",
            ),
            (
                "FR-55",
                "During Citizen self-registration the system shall require the registering user to select "
                "five distinct security questions from a predefined list and to provide an answer to each. "
                "All five questions and answers are mandatory. Answers shall be stored hashed/encrypted, "
                "shall never be displayed or retrievable in clear text, and shall be used only for the "
                "Citizen lost-mobile reset flow (FR-56). Security questions shall not be captured for DSR "
                "Officers or Other Department users, and shall never be entered on a user's behalf by an "
                "administrator.",
                "High",
            ),
            (
                "FR-56",
                "A Public user (Citizen) who has lost access to the registered mobile number shall be "
                "able to reset that mobile number before login as follows: the user enters the Username "
                "and Captcha; the system selects three of the five security questions registered under "
                "FR-55 at random and presents them; all three answers must be correct; the system then "
                "sends a time-limited, single-use PIN to the registered email address; on correct entry "
                "of that PIN the user may enter a new mobile number, which shall itself be verified by "
                "an OTP sent to that new number before the change takes effect. Only after the new mobile "
                "is verified shall the account's registered mobile be updated and login continue with an "
                "OTP to the new mobile. Failed answers and failed PIN entries shall be rate-limited and "
                "shall lock the reset flow after a configured number of attempts. The completed change "
                "shall be notified to the registered email address and audit-logged. This flow is "
                "available only to Public users (Citizens) — see FR-65 for departmental users.",
                "High",
            ),
            (
                "FR-65",
                "The lost-mobile reset flow (FR-56) shall not be offered to DSR Officers or Other "
                "Department users. For these two categories the registered mobile number shall be "
                "changed only by an authorised administrator through the User Management admin console, "
                "with a reason recorded and a full audit trail. No pre-login self-service path to change "
                "the mobile number, email address, or Username shall exist for these categories.",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[5],
        ["ID", "Requirement", "Priority"],
        [
            (
                "FR-10",
                "The system shall dispatch the login OTP only to the user's registered mobile number "
                "(SMS) within 5 seconds of request. The login OTP shall never be sent to email. Email "
                "shall be used for exactly two one-time verification purposes: the email verification "
                "OTP during Citizen registration (FR-63), and the reset PIN in the Citizen lost-mobile "
                "flow (FR-56).",
                "High",
            ),
            ("FR-11", "The system shall validate Captcha before OTP dispatch for all user categories.", "High"),
            (
                "FR-12",
                "Lost-mobile recovery is available only to Public users (Citizens) and shall follow "
                "FR-56 (three random security questions of five, plus a PIN to the registered email, "
                "plus OTP verification of the new mobile). DSR Officers and Other Department users have "
                "no self-service recovery path; their mobile number is changed by an authorised "
                "administrator only (FR-65). The system shall not provide password reset. The login OTP "
                "remains mobile-only in all cases.",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[6],
        ["ID", "Requirement", "Priority"],
        [
            (
                "FR-13",
                "The system shall allow users to view and update their profile information after login. "
                "The Username shall be displayed read-only and shall not be user-editable (FR-62). "
                "Public users (Citizens) may update their mobile number after login, verified by an OTP "
                "to the new number, and may update their email address, verified by an OTP to the new "
                "address; they may also revise their five security questions and answers, which shall "
                "require an OTP to the registered mobile. For DSR Officers and Other Department users, "
                "mobile number and official email ID are maintained by an authorised administrator only "
                "(FR-65). All such changes shall be audit-logged and notified to the user.",
                "High",
            ),
            ("FR-14", "The system shall allow users to upload and update a profile photo where applicable.", "Medium"),
            ("FR-15", "RETIRED (v4.5) — Previously required administrators to deactivate user accounts with reason and audit trail. Superseded by FR-20, which shall require a deactivation reason and audit trail. Number retained for traceability; do not implement separately.", "—"),
        ],
    )

    replace_table_rows(
        doc.tables[7],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-16", "The system shall support RBAC using a single unified Role Master. Roles shall be differentiated by Role Category (Citizen, DSR, Other Department) as defined in Section 6.5.1. Users shall be stored in a single User Master differentiated by User Category.", "High"),
            (
                "FR-17",
                "The system shall allow administrators to assign DSR Officers to one or more sanctioned "
                "posts with available capacity during user creation (at least one post required; "
                "FR-66(a)). There is no "
                "Primary/Secondary role or post distinction. At login, the officer selects one active "
                "post for the session when multiple posts are assigned (FR-52); session roles shall be "
                "Those mapped to the selected assigned post via Post–Role mapping. Additional charge "
                "under FR-53 is only after login.",
                "High",
            ),
            (
                "FR-28",
                "There shall be no separate Role Masters or User Masters for Citizens, DSR Officers, "
                "or Other Department users. All roles and users shall reside in one Role Master and one "
                "User Master respectively, differentiated by Role Category / User Category.",
                "High",
            ),
            (
                "FR-29",
                "During Other Department user creation, the system shall allow administrators to assign "
                "exactly one role from the Role Master filtered by Role Category = Other Department. "
                "Posts Master and sanctioned post occupancy do not apply to Other Department users.",
                "High",
            ),
            (
                "FR-30",
                "At least one sanctioned post with available capacity must be assigned to each DSR Officer "
                "at account creation (FR-66(a)). Optional End Date may be set on a post occupancy only for "
                "genuinely time-bound deputations — not as a substitute for Transfer Out / relieving under "
                "FR-57. When End Date is set, the system shall require a Deputation Reason code. The "
                "occupancy shall remain active until the end of the End Date and shall be freed after "
                "11:59 PM IST of that date by the occupancy refresh job (FR-68) — the same day boundary "
                "as FR-58 relieving. Formal Transfer out / relieving shall follow FR-57 and FR-58. "
                "Transfer In shall follow FR-60, FR-61, and FR-67. Occupancy de-allocation, reserved "
                "Transfer In becoming active, and Sanctioned Posts occupied-count updates shall be applied "
                "by the occupancy refresh job (FR-68).",
                "High",
            ),
            (
                "FR-31",
                "RETIRED (v3.0) — Previously required mandatory end date on Secondary / additional "
                "post assignment for DSR Officers. Superseded by removal of Primary/Secondary assignment; "
                "optional End Date on any post occupancy is covered by FR-30. Number retained for audit "
                "traceability; do not implement.",
                "—",
            ),
            (
                "FR-45",
                "A DSR Officer may be assigned to multiple sanctioned posts concurrently, provided each "
                "selected post has available capacity at assignment (Occupied < Sanctioned Strength — "
                "FR-66). Over-capacity assignment shall be blocked. Occupied count shall increase per "
                "assigned post (including a reserved Transfer In under FR-67) and decrease when occupancy "
                "ends (including after relieving under FR-58). Transfer In assignments shall also require "
                "available capacity except where FR-67 reservation against a recorded relieving applies. "
                "Occupied counts on the Sanctioned Posts Master shall be recalculated by the occupancy "
                "refresh job (FR-68).",
                "High",
            ),
            (
                "FR-66",
                "The system shall apply two distinct vacancy tests and shall not treat them as "
                "interchangeable. (a) Available capacity — used for post assignment during user creation, "
                "for Transfer In, and for the post pick lists in those workflows — is satisfied "
                "when the occupied count for a Post at an Office is less than the sanctioned strength "
                "for that Post at that Office (Occupied < Sanctioned Strength). Occupied for this test "
                "shall include active occupancies and reserved (pending-join) Transfer In occupancies "
                "(FR-67). An occupancy with a recorded Relieving Date still counts as occupied until "
                "the occupancy refresh job de-allocates it (FR-58, FR-68), except that Transfer In may "
                "reserve the slot being freed as defined in FR-67. (b) Unoccupied post — "
                "used only for the FR-53 post-login additional charge selection — is satisfied only when "
                "the occupied count for that Post at that Office is zero (Occupied = 0), irrespective of "
                "sanctioned strength. Reserved Transfer In occupancies also prevent Occupied = 0. A Post "
                "at an Office with sanctioned strength 2 and one occupant "
                "therefore has available capacity for assignment but is not unoccupied, and shall not be "
                "offered under FR-53. Both counts shall be shown on sanctioned post occupancy screens and "
                "reports.",
                "High",
            ),
            (
                "FR-57",
                "The system shall support a Transfer out / relieving process for DSR Officers. An officer "
                "may be relieved from one or more currently assigned post occupancies (Post + Office). "
                "When the superior opens Transfer Out / Relieving, the system shall display only the "
                "offices that fall under the actor's session Office as per the Office Hierarchy Master "
                "(FR-59) — i.e. the actor's own office and descendant offices. Within that office "
                "scope, the system shall list and allow relieving only of post occupancies where the "
                "actor's session Post is the immediate parent of the target Post in the DSR Officer "
                "Hierarchy Master (FR-43). Office span is a first filter only: seeing a descendant "
                "office in the tree does not grant relieving of posts that do not report immediately "
                "to the actor. Example: District Registrar at DRO Mysuru sees SRO offices under DRO "
                "Mysuru and may relieve Sub-Registrar at those SROs (POST-SR reports to POST-DRO); "
                "Sub-Registrar at SRO Yeshwanthapura sees only that SRO and may relieve FDA/DEO under "
                "Sub-Registrar there. IGR or DIGR at Head Office may see statewide DRO/SRO offices in "
                "the span but may relieve only posts that report immediately to their session Post "
                "(e.g. IGR relieves DIGR; DIGR (Enforcement) relieves DRO) — not Sub-Registrar at an "
                "SRO. While relieving, the system "
                "shall capture Relieving Date and Relieving Order (order number / reference; upload of "
                "order document where applicable). Relieving actions shall be audit-logged. Citizens and "
                "Other Department users are out of scope for this process.",
                "High",
            ),
            (
                "FR-58",
                "After a relieving is recorded (FR-57), the system shall retain the user–post occupancy "
                "until the end of the Relieving Date. The occupancy refresh job (FR-68) shall "
                "de-allocate / remove the mapping of the user to the relieved post after 11:59 PM IST of "
                "the Relieving Date (i.e. occupancy ends at the close of that calendar day; the job "
                "runs shortly after midnight). Occupied count for that Post + Office shall then "
                "decrease unless a reserved Transfer In for the same Post + Office becomes active on "
                "the same job run (FR-67). Where the count reaches zero the post "
                "also becomes wholly unoccupied and therefore eligible under FR-53 (FR-66). If the user has no remaining "
                "active post occupancies after de-allocation, login shall be blocked or limited per "
                "policy until a new post is assigned. De-allocation shall be audit-logged.",
                "High",
            ),
            (
                "FR-59",
                "The system shall maintain an Office Hierarchy Master for DSR offices. Structure: "
                "IGR Office is the Head Office (root); under the Head Office are District Registrar "
                "Offices (DRO); under each District Registrar Office are Sub-Registrar Offices (SRO). "
                "Each Office shall have Office Code, Office Name, Office Type (Head Office / District "
                "Registrar Office / Sub-Registrar Office), and optional Parent Office. Application Admin "
                "shall add, edit, enable/disable, and update offices with audit trail. Sanctioned Posts "
                "and Transfer Out / Transfer In office scoping shall use this hierarchy (FR-24, FR-57, FR-60).",
                "High",
            ),
            (
                "FR-60",
                "The system shall support a Transfer In process for DSR Officers. A superior (within "
                "office span per FR-59 and immediate-parent post parentage per FR-43, same rules as "
                "Transfer Out — office span does not expand parentage) may "
                "assign an officer to a sanctioned post (Post + Office). The system shall allow "
                "Transfer In to a post that has available capacity at that office (Occupied < "
                "Sanctioned Strength — FR-66(a)), or, when the post is at full strength, only as a "
                "reserved future-dated occupancy under FR-67 where relieving is already recorded. "
                "While processing Transfer In, the system shall capture Transfer Order / "
                "Reporting Order (order number / reference; upload of order document where applicable) "
                "and Joining Date. A reserved occupancy shall count toward occupied (FR-66(a)) from "
                "the moment it is recorded; login remains blocked until Joining Date (FR-61). The "
                "occupancy becomes active and Sanctioned Posts occupied counts are refreshed by the "
                "midnight job (FR-68). Transfer In actions shall be audit-logged. Citizens and Other "
                "Department users are out of scope for this process.",
                "High",
            ),
            (
                "FR-61",
                "For a Transfer In occupancy (FR-60), including a reserved occupancy under FR-67, the "
                "user shall be allowed to log in and use that post only from 12:00 AM IST (start) of the "
                "Joining Date onwards. Before 12:00 AM IST of the Joining Date, login for that post "
                "shall be blocked (or the occupancy shall not appear as selectable under FR-52). The "
                "occupancy refresh job (FR-68) shall activate the reserved occupancy as of Joining Date "
                "so that from 12:00 AM on the Joining Date the post occupancy is available for login "
                "and session selection.",
                "High",
            ),
            (
                "FR-67",
                "Where a Transfer In is required against a Post + Office that is currently at full "
                "sanctioned strength, the system shall allow a future-dated Transfer In only if a "
                "relieving (FR-57) has already been recorded for an occupancy of that same Post + "
                "Office, and that relieving slot is not already reserved by another Transfer In. The "
                "Joining Date shall be on or after the calendar day following the Relieving Date of "
                "the occupancy being replaced (i.e. after 11:59 PM of Relieving Date — example: "
                "outgoing officer relieved 31-Aug, incoming Joining Date 01-Sep or later). Recording "
                "the Transfer In shall immediately reserve one unit of capacity: the reserved occupancy "
                "counts toward Occupied for FR-66(a) so that a second incoming officer cannot be "
                "assigned to the same slot, but the incoming officer cannot log in until Joining Date "
                "(FR-61). If no relieving is recorded for that Post + Office, Transfer In to a full "
                "post shall remain blocked. Transfer In to a post that already has available capacity "
                "does not require a prior relieving.",
                "High",
            ),
            (
                "FR-68",
                "The system shall run a scheduled occupancy refresh job shortly after midnight IST each "
                "calendar day (after 12:00 AM IST). On each run the job shall: (1) de-allocate user–post "
                "occupancies whose Relieving Date has ended (FR-58) or whose optional occupancy End "
                "Date has been reached (FR-30); (2) activate reserved Transfer In occupancies whose "
                "Joining Date is the current calendar day (FR-60, FR-61, FR-67); (3) recalculate and "
                "persist occupied count, remaining capacity, and wholly-unoccupied flag on the "
                "Sanctioned Posts Master for every Post + Office affected (FR-48, FR-66); (4) refresh "
                "the officer's effective post assignments so that login post selection (FR-52) reflects "
                "only occupancies that are active as of that day. The job shall be idempotent, "
                "audit-logged (run timestamp, occupancies ended, occupancies activated, before/after "
                "occupied counts), and shall raise an operational alert on failure so that relieving "
                "and joining dates are not left unenforced.",
                "High",
            ),
            (
                "FR-46",
                "The system shall maintain a separate Posts Master for DSR establishment posts "
                "(e.g. Sub-Registrar (SR), FDA (Enforcement), DEO, DRO, HQA (Enforcement)). "
                "Application Admin shall add, edit, enable/disable, and update posts. Posts Master "
                "is distinct from Role Master.",
                "High",
            ),
            (
                "FR-47",
                "The system shall maintain a Post–Role mapping table. Each Post shall map to the "
                "corresponding unique Role(s) from the Role Master. Division-specific posts "
                "(e.g. POST-DIGR-ADMIN, POST-DIGR-ENF) shall not be mapped to roles of other divisions, "
                "so that assigning one post does not grant Admin + Vigilance + Enforcement access "
                "together. Application Admin shall add, edit, and remove mappings with audit trail. "
                "One post may map to multiple roles only when those roles intentionally belong to the "
                "same post's functional scope (not across unrelated divisions). The system shall block "
                "sanctioning or assignment of any Post that has no active Post–Role mapping row — an "
                "officer assigned to an unmapped post would receive zero roles (FR-17). Example rows in "
                "this BRD are incomplete; Domain Expert (Prabhakar Naik) shall complete the full seed "
                "mapping list for all active Posts before go-live.",
                "High",
            ),
            (
                "FR-48",
                "The Sanctioned Posts Master shall reference a Post from the Posts Master (not Role "
                "directly), an Office from the Office Hierarchy Master (by Office Code), and sanctioned "
                "strength. Strength, occupied count, remaining "
                "capacity, and whether the post is wholly unoccupied are maintained per Post per Office "
                "(FR-66). Occupied count shall include active occupancies and reserved Transfer In "
                "occupancies (FR-67) and shall be recalculated by the occupancy refresh job (FR-68).",
                "High",
            ),
            (
                "FR-49",
                "Posts Master entries for DSR leadership and staff shall be division-specific where "
                "the organisation chart distinguishes them (e.g. POST-DIGR-ADMIN, POST-DIGR-VIG, "
                "POST-DIGR-ENF, POST-FDA-ADMIN, POST-FDA-ENF). Generic posts that would map to all "
                "division variants of DIGR/AIGR/FDA/SDA are not permitted in seed data.",
                "High",
            ),
            (
                "FR-33",
                "During Other Department user creation, the system shall allow an optional End Date. "
                "If an End Date is entered, the system shall automatically deactivate the user account "
                "on that date and block login thereafter.",
                "High",
            ),
            (
                "FR-34",
                "When assigning a role to a user, the system shall present only roles whose Role Category "
                "matches the user's User Category (Citizen roles for Citizens, DSR roles for DSR Officers, "
                "Other Department roles for Other Department users).",
                "High",
            ),
            (
                "FR-35",
                "The Role Master shall include the seed roles listed in Section 6.5.4 for Role Category = "
                "Citizen, Role Category = Other Department, and Role Category = DSR (Department). "
                "Application Admin may add, edit, or deactivate roles within each Role Category.",
                "High",
            ),
            (
                "FR-36",
                "The system shall maintain a Module Master of application modules (including Registration "
                "of Documents, Marriage Registration, Encumbrance Search, Certified Copy, and others "
                "listed in Section 6.5.6). Modules are not DSR organisational roles; they define "
                "functional areas to which roles are mapped. Registration of Documents is a single "
                "module with no online/offline classification in this master.",
                "High",
            ),
            (
                "FR-37",
                "The system shall allow Application Admin to map roles to Module Functions "
                "(Role–Module–Function mapping). Access to a module capability is granted only through "
                "this mapping (except Application Admin system privileges — see FR-51). Each mapping "
                "row shall reference an exact Role Master name (unique role name) and a valid Module "
                "Function code. Application Admin shall add, edit, and update mappings with audit trail.",
                "High",
            ),
            (
                "FR-39",
                "The system shall maintain a Module Function Master under each module (e.g. VIEW, ADD, "
                "EDIT, APPROVE, SIGN, PRINT, DOWNLOAD). Application Admin shall add, edit, and update "
                "module functions.",
                "High",
            ),
            (
                "FR-40",
                "The system shall maintain a Resource Master of APIs and URLs linked to Module Functions "
                "(resource type API or URL, HTTP method, path pattern, Is Public flag). Application Admin "
                "shall add, edit, and update resources. Roles shall not store raw API/URL lists; access "
                "is resolved via Role → Module Function → Resource. Is Public = Yes marks resources "
                "accessible without authentication (deny-by-default otherwise — FR-41).",
                "High",
            ),
            (
                "FR-41",
                "At runtime the application shall enforce access as follows: resolve the user's session "
                "roles (for DSR Officers: assigned post from FR-52 when no additional charge is active; "
                "additional charge post only when FR-53 is active); load Role–Module–Function claims per "
                "FR-38; for each API/URL request look up the Resource Master using Type + Method + Path "
                "pattern matching with the most specific matching pattern taking precedence when multiple "
                "patterns match (longer / more specific path wins over shorter prefix). If no resource "
                "matches, access shall be denied (HTTP 403) unless the resource is explicitly marked "
                "Is Public = Yes in the Resource Master — deny-by-default. Allow only if the user holds "
                "the required Module Function for a matched non-public resource; otherwise deny (HTTP "
                "403) and audit the attempt. Application Admin system principals are authorised for "
                "FN-UM-ADMIN outside Role–Module–Function mapping (FR-51).",
                "High",
            ),
            (
                "FR-50",
                "Role–Module–Function mapping entries shall be validated against exact Role Master names "
                "(referential integrity). The system shall reject save of a mapping whose Role name does "
                "not match an active Role Master entry character-for-character (e.g. bare FDA is invalid "
                "when only FDA (Admin), FDA (Enforcement), etc. exist). Post–Role mapping and Hierarchy "
                "Master entries shall reference Posts Master Post Name character-for-character (same "
                "discipline as Role names). Same integrity pattern as Hierarchy ↔ Posts (FR-44).",
                "High",
            ),
            (
                "FR-51",
                "Application Admin is a system-level / deployment-seeded actor, not a Role Master seed "
                "role and not created via the normal User Master creation workflow. Privileges for "
                "maintaining masters (including FN-UM-ADMIN resources) are granted to this principal "
                "outside Role–Module–Function mapping. Application Admin shall not appear as a DSR "
                "(or other) row in the Role–Module–Function example catalogue.",
                "High",
            ),
            (
                "FR-42",
                "Application Admin shall create, edit, enable/disable, and update Module Master, "
                "Module Function Master, Resource Master, and Role–Module–Function mapping tables "
                "from the User Management admin console. All changes shall be audit-logged.",
                "High",
            ),
            (
                "FR-38",
                "A user's effective module access for a session shall be determined as follows: for DSR "
                "Officers, when no additional charge is active under FR-53, Module Functions mapped to "
                "the assigned post (FR-52) apply; when additional charge is active, only Module Functions "
                "mapped to the additional charge subordinate post apply — assigned-post privileges are "
                "suspended until the officer switches back to assigned-post context without logout; for "
                "Other Department users, the single assigned role; for Citizens, assigned Citizen roles.",
                "High",
            ),
            ("FR-18", "The system shall restrict access to features, APIs, and URLs based on the user's assigned role(s) and Role–Module–Function / Resource mappings.", "High"),
            ("FR-19", "The system shall maintain DSR organizational divisions and DSR roles within the unified Role Master (Role Category = DSR) as defined in Section 6.5.1 and the DSR Officer Hierarchy Master in Section 6.5.7.", "High"),
            (
                "FR-43",
                "The system shall maintain a DSR Officer Hierarchy Master representing the reporting "
                "structure of Department officers (organisational chart). Each hierarchy node shall "
                "reference a Post from the Posts Master (not a Role) and an optional parent hierarchy "
                "node (also a Post). Application Admin shall add, edit, reorder, enable/disable, and "
                "update hierarchy nodes with audit trail. Reporting lines are position-based so that "
                "a specific post holder has one parent post in the tree.",
                "High",
            ),
            (
                "FR-44",
                "The DSR Officer Hierarchy Master shall seed and support the structure: Additional Chief "
                "Secretary / Principal Secretary / Secretary → Inspector General of Registration & "
                "Commissioner of Stamps → Divisions (Admin Law & Computers; Vigilance; Computers; "
                "Enforcement; Intelligence & Audit; DIGR CVC) with subordinate posts as listed in "
                "Section 6.5.7. Every hierarchy node Post must exist in Posts Master; every mapped Role "
                "in Post–Role mapping must exist in Role Master with a unique role name.",
                "High",
            ),
            (
                "FR-24",
                "The system shall maintain a Sanctioned Posts Master that references Posts Master entries "
                "and records sanctioned strength (approved headcount) for each Post at each office. "
                "Sanctioned strength applies only to DSR Posts Master entries.",
                "High",
            ),
            (
                "FR-25",
                "The system shall allow Application Admin to configure and update sanctioned strength "
                "per Post per office in the Sanctioned Posts Master, with audit trail of changes.",
                "High",
            ),
            (
                "FR-26",
                "DSR department users shall be assigned only to posts defined in the Posts Master and "
                "sanctioned in the Sanctioned Posts Master. A user may hold multiple post occupancies. "
                "Assignment to an unlisted post, or to a sanctioned post already at full strength "
                "(over-capacity), shall be blocked (FR-66(a)).",
                "High",
            ),
            (
                "FR-27",
                "The system shall display sanctioned strength, occupied count, and remaining vacancies "
                "per Post per office, and prevent over-capacity assignment. When assigning posts to a DSR "
                "Officer, only sanctioned posts with available capacity shall be selectable (FR-66(a)). "
                "Screens shall make clear whether a post is wholly unoccupied or only partly filled, "
                "since FR-53 requires a wholly unoccupied post (FR-66(b)). Occupied count shall include "
                "reserved Transfer In occupancies (FR-67) and shall be kept current by the occupancy "
                "refresh job (FR-68).",
                "High",
            ),
            (
                "FR-32",
                "The system shall provide dedicated step-by-step workflows for role assignment during "
                "user creation for DSR Officers (Section 6.6.1) and Other Department users (Section 6.6.2).",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[8],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-20", "The system shall allow administrators to create, edit, suspend, and deactivate user accounts by category. Deactivation shall require a reason and shall be audit-logged (supersedes retired FR-15).", "High"),
            ("FR-21", "The system shall allow administrators to search and filter users by Username (including KGID for departmental users), category, role, office, division, and status.", "Medium"),
            ("FR-22", "The system shall log all administrative actions performed on user accounts.", "High"),
            ("FR-23", "The system shall notify a user via SMS/email when their account is created, suspended, or deactivated.", "Medium"),
        ],
    )

    replace_table_rows(
        doc.tables[9],
        ["Category", "Requirement"],
        [
            ("Security", "No password storage — authentication is OTP-based only; the login OTP is sent only to the registered mobile."),
            ("Security", "OTP validity, length, attempt limits, resend throttle, lockout, session idle timeout, absolute session lifetime, and concurrent-session rules shall be defined in FR-69–FR-76 before UAT."),
            ("Security", "All data in transit shall be encrypted using TLS 1.2 or higher."),
            ("Security", "Biometric data for departmental users shall comply with Aadhaar Act, 2016 and UIDAI guidelines."),
            ("Security", "Security-question answers shall be stored hashed/encrypted, shall never be displayed or retrievable in clear text, and shall never be entered or viewed by an administrator."),
            ("Security", "The Citizen lost-mobile reset shall require two independent proofs — three of five security questions answered correctly and a single-use PIN delivered to the registered email — before a new mobile number may be entered, and the new number shall itself be OTP-verified before the change takes effect."),
            ("Security", "The reset PIN shall be single-use and time-limited; question and PIN attempts shall be rate-limited and the reset flow locked after a configured number of failures."),
            ("Security", "No pre-login self-service path shall exist to change the mobile number, email address, or Username of a DSR Officer or Other Department user."),
            ("Performance", "OTP dispatch (SMS to mobile only) shall complete within 5 seconds of user request."),
            ("Performance", "Registration verification OTPs (email and mobile) and the reset PIN to email shall be dispatched within 30 seconds of request."),
            ("Performance", "Login and authentication requests shall complete within 2 seconds after OTP/biometric verification."),
            ("Availability", "The occupancy refresh job (FR-68) shall complete shortly after midnight so that relieving ended the previous day and Transfer In joining the current day are reflected before the first login of the day; job failure shall raise an operational alert."),
            ("Usability", "Registration and OTP login workflows shall be completable on desktop and mobile browsers."),
            ("Usability", "Username availability shall be checked and reported to the citizen during registration before submission, with suggestions where the chosen Username is taken."),
            ("Auditability", "All create, update, delete, login, mobile-update, email-update, Citizen lost-mobile reset, administrator-initiated mobile change, access-control, Transfer Out / Transfer In, and occupancy-refresh job actions shall be logged with timestamp and actor."),
            ("Compliance", "The module shall comply with Karnataka e-Governance, MeitY/CERT-In, and applicable data protection norms."),
        ],
    )

    replace_table_rows(
        doc.tables[10],
        ["Risk", "Impact", "Mitigation"],
        [
            ("OTP delivery failure (SMS)", "High", "Retry mechanism; allow mobile update after login (FR-13); Citizen lost-mobile reset (FR-56); administrator-initiated change for departmental users (FR-65)."),
            ("Biometric device unavailability", "Medium", "Define fallback procedure for DSR and Other Department users per security policy."),
            ("Account takeover via lost-mobile reset", "High", "Two independent proofs required — three of five questions plus single-use PIN to registered email — and OTP verification of the new mobile; attempts rate-limited and flow locked after configured failures; change notified to registered email and audit-logged (FR-56)."),
            ("Security-question answers known to a third party", "Medium", "Answers are self-entered by the citizen only, stored hashed, never displayed, and never captured by an administrator; departmental users hold no security questions at all (FR-55, FR-65)."),
            ("Citizen has lost both mobile and access to registered email", "Medium", "Assisted recovery at a designated counter with documentary identity verification and supervisory approval, fully audit-logged; no purely remote path."),
            ("Security-question answers forgotten", "Medium", "Assisted recovery with strong documentary identity verification and audit; no override of the two-proof rule."),
            ("Username squatting or impersonating Usernames", "Medium", "Reject reserved and misleading Username patterns (e.g. names resembling official designations); Username uniqueness enforced across the whole User Master (FR-62)."),
            ("Over-capacity post assignment", "High", "Enforce sanctioned strength validation; block assignment when the post at that office is at full strength (FR-66(a))."),
            ("Vacancy tests confused in build", "High", "FR-66 defines available capacity and wholly unoccupied separately; FR-53 additional charge uses Occupied = 0; active context is either assigned post or additional charge post, not both (FR-38)"),
            ("Post occupancy end date / relieving not enforced", "High", "Occupancy refresh job shortly after midnight (FR-68) de-allocates relieving (FR-58) and occupancy End Date (FR-30); reserved Transfer In activated on Joining Date; audit alert on job failure."),
            ("Unmapped post assigned or sanctioned", "High", "FR-47 blocks sanction/assignment without Post–Role mapping; Domain Expert to complete seed mappings before go-live."),
            ("OTP/session policy undefined at build", "High", "Configure FR-69–FR-76 before UAT; especially session idle timeout on shared SRO counters with FR-52/FR-53 fixed session post."),
            ("Unauthorised relieving / transfer in", "High", "Scope to offices under actor (FR-59) then immediate-parent posts only (FR-43, FR-57); seeing a descendant office does not grant SR transfer to IGR/DIGR; available capacity or FR-67 reservation required for Transfer In (FR-60, FR-66(a))."),
            ("Joining Date login gate not enforced", "High", "Block login for Transfer In occupancy until 12:00 AM of Joining Date (FR-61); FR-68 activates reserved occupancy that day."),
            ("Double-booking a full post during handover", "High", "FR-67 reserves capacity when relieving is recorded; reserved occupancy counts toward Occupied so a second Transfer In cannot take the same slot."),
            ("Other Department end date not enforced", "Medium", "Scheduled job to deactivate Other Department users on optional End Date; audit alert."),
            ("Scope creep beyond OTP-only auth", "Medium", "Maintain strict change-control; no password management in KAVERI 3.0."),
            ("Delayed stakeholder sign-off", "Low", "Set clear review deadlines and escalation path."),
        ],
    )

    replace_table_rows(
        doc.tables[11],
        ["Term", "Definition"],
        [
            ("BRD", "Business Requirements Document"),
            ("RBAC", "Role-Based Access Control"),
            ("OTP", "One-Time Password — the login OTP is sent only to the registered mobile (SMS); separate one-time OTPs verify the email address and mobile number at Citizen registration (FR-63) and verify a new mobile number; no static password exists"),
            ("Username", "The unique login identifier for every account, unique across the whole User Master: a preferred Username chosen by the citizen for Public users, and the KGID for DSR Officers and Other Department users (FR-04, FR-62)"),
            ("Preferred Username", "The Username a citizen chooses for themselves during self-registration, subject to an availability check against the entire User Master (FR-62)"),
            ("KGID", "Karnataka Government Insurance Department number — the unique government employee identifier, used as the Username for DSR Officers and Other Department users (FR-62, FR-64)"),
            ("Official email ID", "Government or parent-department email address recorded for DSR Officers and Other Department users in place of a personal address; validated against the permitted domain list where configured (FR-64)"),
            ("Security question", "One of five questions selected by a citizen at registration with a secret answer; three are asked at random during lost-mobile reset; not held for departmental users (FR-55, FR-56)"),
            ("Reset PIN", "Single-use, time-limited code sent to a citizen's registered email address as the second independent proof in the lost-mobile reset flow (FR-56)"),
            ("Lost-mobile reset", "Citizen-only pre-login flow to replace the registered mobile number: three of five security questions, then a PIN to the registered email, then OTP verification of the new number (FR-56); not available to departmental users (FR-65)"),
            ("DSR", "Department of Stamps and Registration, Government of Karnataka"),
            ("UAT", "User Acceptance Testing"),
            ("IGR", "Inspector General of Registration"),
            ("DIGR", "Deputy Inspector General of Registration"),
            ("AIGR", "Assistant Inspector General of Registration"),
            ("Sanctioned post", "A Posts Master entry sanctioned at a specific office with approved strength; Post + Office + Strength"),
            ("Sanctioned strength", "The approved number of occupants for a given Post at a given office"),
            ("Available capacity", "Vacancy test used for post assignment and Transfer In: occupied count (active occupancies plus reserved Transfer In occupancies) for a Post at an Office is less than its sanctioned strength (FR-66(a), FR-67)"),
            ("Wholly unoccupied post", "Vacancy test used only for FR-53 post-login additional charge: the occupied count for a Post at an Office is zero, irrespective of sanctioned strength (FR-66(b)); reserved Transfer In also counts as occupied"),
            ("Posts Master", "Separate catalogue of DSR establishment posts (distinct from Role Master)"),
            ("Post–Role mapping", "Table linking a Post to one or more Roles; one post may map to multiple roles"),
            ("Office", "A concrete DSR office instance (e.g. IGR Head Office, DRO Mysuru, SRO Yeshwanthapura)"),
            ("Office Hierarchy Master", "Tree of DSR offices: IGR Head Office (root) → District Registrar Offices → Sub-Registrar Offices (FR-59)"),
            ("Office span", "The actor's session Office plus all descendant offices in the Office Hierarchy Master; first filter for Transfer Out / Transfer In lists — actions remain limited to posts whose immediate parent is the actor's session Post (FR-57, FR-60)"),
            ("Role Master", "Single master of all roles for Citizens, DSR Officers, and Other Department users; differentiated by Role Category"),
            ("User Master", "Single master of all users across categories; differentiated by User Category"),
            ("Role Category", "Attribute on a role in the Role Master: Citizen, DSR, or Other Department"),
            ("User Category", "Attribute on a user in the User Master: Public (Citizen), DSR Officer, or Other Department"),
            ("Module Master", "Catalogue of application modules (e.g. Registration of Documents, Marriage Registration, Encumbrance Search, Certified Copy)"),
            ("Module Function Master", "Catalogue of privilege verbs under a module — e.g. VIEW, ADD, EDIT, APPROVE, SIGN, PRINT, DOWNLOAD"),
            ("Resource Master", "Catalogue of APIs and URLs linked to a Module Function (type, HTTP method, path pattern)"),
            ("Role–Module–Function mapping", "Association of a Role Master role (exact unique name) to one or more Module Functions; Application Admin system privileges are not represented here (FR-51)"),
            ("Module function", "Privilege verb within a module — e.g. VIEW, ADD, EDIT, APPROVE, SIGN, PRINT, DOWNLOAD"),
            ("Resource", "An API endpoint or UI URL/route protected by a Module Function"),
            ("Access enforcement", "Runtime check: request path/method → Resource → required Module Function → user's role claims"),
            ("Application Admin", "System-level / deployment-seeded actor who maintains masters; not a Role Master seed role; privileges outside Role–Module–Function mapping (FR-51)"),
            ("DSR Officer Hierarchy Master", "Master of reporting relationships among DSR Posts (Posts Master) as per the Department organisational chart"),
            ("Hierarchy node", "A node in the DSR Officer Hierarchy Master linking a Post to an optional parent Post node"),
            ("Unique role name", "Distinct Role Master name including division context where needed (e.g. FDA (Admin), FDA (Enforcement)) — no duplicate bare FDA/SDA/Typist entries"),
            ("Other Department role", "A role in the Role Master with Role Category = Other Department; exactly one such role is assigned per Other Department user"),
            ("Post occupancy", "Assignment of a DSR user to a sanctioned post at an office; roles follow from Post–Role mapping for the assigned post at login (FR-52); no Primary/Secondary distinction"),
            ("Relieving / Transfer out", "Process by which a superior removes a DSR Officer from a post occupancy within the superior's office span and immediate-parent posts only, capturing Relieving Date and Relieving Order (FR-57–FR-59); mapping ends after 11:59 PM of that date via the occupancy refresh job (FR-58, FR-68)"),
            ("Relieving Date", "Calendar date captured during relieving; user–post mapping is removed after 11:59 PM of this date by the occupancy refresh job (FR-68)"),
            ("Relieving Order", "Order number / reference (and optional uploaded order document) recorded when relieving an officer"),
            ("Transfer In", "Process by which a superior assigns a DSR Officer to a post within office span and immediate-parent posts only, capturing Transfer/Reporting Order and Joining Date (FR-60, FR-66(a)); when the post is full, allowed only if relieving is already recorded and capacity is reserved (FR-67); login from 12:00 AM on Joining Date (FR-61)"),
            ("Reserved occupancy", "A future-dated Transfer In recorded against a Post + Office after relieving is captured; counts toward Occupied immediately so the slot cannot be double-booked; incoming officer cannot log in until Joining Date (FR-67, FR-61)"),
            ("Joining Date", "Calendar date captured during Transfer In; login for that post occupancy is allowed only from 12:00 AM of this date onwards; the occupancy refresh job activates reserved occupancies on this date (FR-61, FR-68)"),
            ("Sub-Registrar Office (SRO)", "Office Type in Office Hierarchy Master — not a role name; do not use SRO as a Role Master name (use Sub Registrar (Admin) or Sub Registrar (Computers) instead)"),
            ("Deputation Reason", "Mandatory code when optional occupancy End Date is set under FR-30; End Date is for time-bound deputations only, not a substitute for relieving"),
            ("Is Public", "Resource Master flag: Yes allows unauthenticated access; default No — deny-by-default at runtime (FR-40, FR-41)"),
            ("Occupancy refresh job", "Scheduled job shortly after midnight IST that de-allocates relieved and ended occupancies, activates reserved Transfer In occupancies whose Joining Date is today, recalculates Sanctioned Posts occupied counts, and refreshes officers' effective post assignments for login (FR-68)"),
            ("Transfer / Reporting Order", "Order number / reference (and optional uploaded document) recorded when transferring an officer into a post"),
            ("Login post selection", "After authentication, when a DSR Officer has multiple active post occupancies, the mandatory step to choose one Post (shown with Office details) for the session (FR-52)"),
            ("Assigned post", "The sanctioned post (Post + Office) selected or auto-selected at login under FR-52; fixed for the session unless the officer logs out and logs in again with a different choice"),
            ("Additional charge", "Temporary in-office cover of a wholly unoccupied subordinate post under the assigned post, taken after login under FR-53 without logout or order; session privileges derive from the additional charge post only while active; switch back to assigned post to restore assigned-post privileges (FR-38)"),
            ("Session post", "Synonym for assigned post (FR-52) in audit and reporting contexts; distinct from additional charge post under FR-53"),
            ("Additional charge switch", "Post-login action allowing a DSR Officer to take or clear additional charge of a wholly unoccupied subordinate post at the same office (FR-53, FR-66(b)); no logout required"),
            ("Session Post–Role display", "After login, persistent display of assigned Post with Role(s); when FR-53 is active, shows which context drives session privileges — assigned post or additional charge post (FR-54)"),
            ("Parent department", "The government department to which an Other Department user belongs (e.g. Revenue, Treasury, Police)"),
            ("Account End Date", "Optional date on an Other Department user account; if set, the system deactivates the user on that date"),
        ],
    )

    replace_table_rows(
        doc.tables[12],
        ["Name", "Role", "Signature", "Date"],
        [
            ("Prashanth", "Product Owner", "", ""),
            ("Prabhakar Naik", "Domain Expert", "", ""),
            ("Kaveri IT Cell Lead", "IT Security / Engineering", "", ""),
        ],
    )

    rbac_fr_tbl_el = doc.tables[7]._tbl
    rbac_heading = next(p for p in doc.paragraphs if p.text.strip() == "6.5 Role-Based Access Control (RBAC)")
    rbac_parent = rbac_heading._parent
    rbac_fr_tbl_el.getparent().remove(rbac_fr_tbl_el)
    rbac_heading._element.addnext(rbac_fr_tbl_el)

    sub1 = insert_heading_after(
        rbac_fr_tbl_el, "6.5.1 Unified Role Master and User Master", "Heading 3", parent=rbac_parent
    )
    note1 = insert_paragraph_after(
        sub1,
        "The system shall maintain a single Role Master and a single User Master for all user types. "
        "There shall be no separate masters for Citizens, DSR Officers, or Other Department users. "
        "Differentiation is by Role Category on roles and User Category on users. When assigning a role, "
        "the system shall filter the Role Master to roles matching the user's User Category (FR-34).",
        "Normal",
    )
    cat_diff_tbl = insert_table_after(
        note1,
        ["Master", "Differentiator", "Values", "Purpose"],
        [
            (
                "User Master",
                "User Category",
                "Public (Citizen); DSR Officer; Other Department",
                "One user store; category drives auth, creation path, and eligible roles",
            ),
            (
                "Role Master",
                "Role Category",
                "Citizen; DSR; Other Department",
                "One role store; category filters which roles can be assigned to which users",
            ),
        ],
    )
    div_intro = insert_heading_after(
        cat_diff_tbl,
        "DSR roles in the Role Master (Role Category = DSR) are organised by division. Role names "
        "shall be unique (e.g. FDA (Admin), FDA (Enforcement) — not a single shared FDA). Summary:",
        "Normal",
        parent=rbac_parent,
    )
    div_tbl = insert_table_after(
        div_intro,
        ["Division", "Roles (Role Category = DSR) — unique names"],
        [
            ("Secretariat", "ACS / Principal Secretary / Secretary"),
            ("Top Management", "IGR"),
            (
                "Division 1 — Admin, Law & Computers",
                "DIGR (Admin, Law & Computers), AIGR (Admin), HQA (Admin), Sub Registrar (Admin), FDA (Admin), SDA (Admin), Typist (Admin), HQA (RTI), FDA (RTI), SDA (RTI), Statistical Inspector, Accountant Superintendent (Admin)",
            ),
            ("Division 2 — Vigilance", "DIGR (Vigilance), Law Officer"),
            (
                "Division 3 — Computers",
                "AIGR (Computers), System Integrator, PMU, Application Developer, HQA / Project Manager (Comp), Sub Registrar (Computers), FDA (Computers), SDA (Computers)",
            ),
            ("Division 4 — Enforcement", "DIGR (Enforcement), DRO, HQA (Enforcement), Sub-Registrar (SR), FDA (Enforcement), SDA (Enforcement), DEO"),
            (
                "Division 5 — Intelligence & Audit",
                "DIGR (Intelligence), AIGR (Audit), HQA (Audit), Superintendent (Audit), FDA (Audit), SDA (Audit), Typist (Audit)",
            ),
            ("Division 6 — CVC", "DIGR CVC, JD Town Planning"),
        ],
    )

    sub2 = insert_heading_after(
        div_tbl, "6.5.2 User Categories and Authentication", "Heading 3", parent=rbac_parent
    )
    note2 = insert_paragraph_after(
        sub2,
        "All users are stored in the single User Master, differentiated by User Category. "
        "All user categories authenticate without passwords. Login uses the Username; the login OTP is "
        "sent only to the registered mobile. Captcha is required before OTP dispatch. "
        "The Username is a preferred Username chosen by the Citizen at registration, and the KGID for "
        "DSR Officers and Other Department users (FR-62). The Username is the single unique login "
        "identifier across the whole User Master; email address and mobile number carry no uniqueness "
        "constraint (FR-04). "
        "DSR Officers and Other Department users also verify Biometrics. Citizens may update their mobile "
        "after login (FR-13) and may reset a lost mobile before login using three of five security "
        "questions plus a PIN to the registered email (FR-56); departmental users have no self-service "
        "mobile reset and are served by an administrator instead (FR-65). "
        "DSR Officers with multiple active post occupancies must select one post (with office details) "
        "after authentication before entering the application (FR-52). After login they may take "
        "additional charge of a wholly unoccupied subordinate post at the same office without logout "
        "(FR-53). The home/header shows the assigned Post with mapped Role(s) and any active additional "
        "charge (FR-54). OTP validity, resend limits, lockout, and session timeout rules are reserved "
        "in FR-69–FR-76 (to be configured before UAT).",
        "Normal",
    )
    sro_note = insert_paragraph_after(
        note2,
        "Terminology: Sub-Registrar Office (abbreviated SRO in Office Hierarchy) is an office type only. "
        "Role Master names use Sub Registrar (Admin) or Sub Registrar (Computers) — not SRO (Admin) or "
        "SRO (Comp) — to avoid confusion with the office abbreviation.",
        "Normal",
    )
    cat_tbl = insert_table_after(
        sro_note,
        ["User Category", "Username", "Authentication", "Lost-mobile reset"],
        [
            (
                "Public users (Citizens) — User Category = Public (Citizen)",
                "Preferred Username chosen at registration (FR-62)",
                "Username + Captcha + OTP to mobile",
                "Self-service: 3 of 5 security questions + PIN to registered email + OTP to new mobile (FR-56)",
            ),
            (
                "Department users (DSR Officers) — User Category = DSR Officer",
                "KGID (FR-62, FR-64)",
                "Username (KGID) + Captcha + OTP to mobile + Biometrics; then FR-52",
                "Not available — administrator changes mobile with reason and audit (FR-65)",
            ),
            (
                "Other Department users — User Category = Other Department",
                "KGID (FR-62, FR-64)",
                "Username (KGID) + Captcha + OTP to mobile + Biometrics",
                "Not available — administrator changes mobile with reason and audit (FR-65)",
            ),
        ],
    )

    reg_intro = insert_heading_after(
        cat_tbl,
        "Citizen registration — preferred Username, dual OTP verification, security questions "
        "(FR-01, FR-55, FR-62, FR-63):",
        "Normal",
        parent=rbac_parent,
    )
    reg_note = insert_paragraph_after(
        reg_intro,
        "Citizens self-register instantly with no approval workflow. The registering user chooses their "
        "own Username, which the system checks for availability across the entire User Master. Both the "
        "email address and the mobile number are verified by separate OTPs before the account is created, "
        "so neither channel can be claimed without proof of control. Five security questions are captured "
        "at this point; they are used only for the lost-mobile reset in FR-56 and are entered by the "
        "citizen alone — never by an administrator.",
        "Normal",
    )
    reg_tbl = insert_table_after(
        reg_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "Open citizen registration; enter name and personal particulars", "Citizen", "Instant self-registration (FR-01)"),
            ("2", "Enter preferred Username; system checks availability", "Citizen / System", "Rejected if already in use anywhere in User Master (FR-62)"),
            ("3", "Enter email address and mobile number", "Citizen", "No uniqueness constraint on either (FR-04)"),
            ("4", "System dispatches an OTP to the entered email address", "System", "FR-63; time-limited, single-use"),
            ("5", "System dispatches a separate OTP to the entered mobile number", "System", "FR-63; time-limited, single-use"),
            ("6", "Enter both OTPs — email and mobile", "Citizen", "Account not created unless both verify (FR-63)"),
            ("7", "Select five distinct security questions and answer each", "Citizen", "All five mandatory; self-entered only (FR-55)"),
            ("8", "Save — account created; Citizen role assigned; answers stored hashed", "System", "FR-01, FR-55; no approval workflow"),
        ],
    )

    auth_flow_intro = insert_heading_after(
        reg_tbl,
        "Login (all categories — FR-04–FR-07, FR-10, FR-11, FR-62):",
        "Normal",
        parent=rbac_parent,
    )
    auth_flow_note = insert_paragraph_after(
        auth_flow_intro,
        "Login uses the Username — the preferred Username for Citizens and the KGID for DSR Officers and "
        "Other Department users. Because the Username is unique across the whole User Master, the account "
        "is resolved unambiguously from the Username alone and no category selection is needed. The login "
        "OTP is never emailed; it is sent only by SMS to the registered mobile.",
        "Normal",
    )
    login_flow_tbl = insert_table_after(
        auth_flow_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "Enter Username and Captcha", "User", "Citizen: preferred Username; DSR / Other Dept: KGID (FR-62)"),
            ("2", "Validate Captcha; look up the account by Username", "System", "FR-11; Username unique across User Master (FR-04)"),
            ("3", "Dispatch login OTP to the registered mobile only", "System", "FR-10; never to email"),
            ("4", "Enter OTP (+ Biometrics for DSR / Other Department)", "User", "FR-05 / FR-06 / FR-07"),
            ("5", "If a Citizen cannot receive the OTP — leave login for lost-mobile reset", "Citizen / System", "FR-56; Citizens only"),
            ("6", "On success — enter home under assigned post (DSR: FR-52 only at login)", "System", "Session established; FR-53 available after login"),
            ("7", "After login — Citizen may update mobile or email from profile", "Citizen", "FR-13; departmental users via administrator (FR-65)"),
        ],
    )

    reset_intro = insert_heading_after(
        login_flow_tbl,
        "Citizen lost-mobile reset — three of five questions plus PIN to registered email (FR-56):",
        "Normal",
        parent=rbac_parent,
    )
    reset_note = insert_paragraph_after(
        reset_intro,
        "This flow exists because the login OTP goes only to the registered mobile, so a citizen who has "
        "lost that number cannot otherwise log in. It requires two independent proofs before the mobile "
        "number can be changed: three security questions chosen at random from the five registered under "
        "FR-55 (all three must be answered correctly, and the selection differs on each attempt), and a "
        "single-use PIN sent to the registered email address. The new mobile number is itself verified by "
        "OTP before the change takes effect. This flow is available only to Public users (Citizens).",
        "Normal",
    )
    reset_tbl = insert_table_after(
        reset_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "On the login screen choose \"Lost / changed mobile number\"", "Citizen", "Offered to Citizens only (FR-56, FR-65)"),
            ("2", "Enter Username and Captcha", "Citizen", "Username identifies the account (FR-62)"),
            ("3", "System selects three of the five registered security questions at random", "System", "Selection varies on each attempt (FR-56)"),
            ("4", "Answer all three questions", "Citizen", "All three must be correct; rate-limited"),
            ("5", "On success — system sends a single-use, time-limited PIN to the registered email", "System", "Second independent factor (FR-56)"),
            ("6", "Enter the PIN", "Citizen", "Flow locked after configured failed attempts"),
            ("7", "Enter the new mobile number", "Citizen", "Replaces the lost number"),
            ("8", "System sends an OTP to the new mobile; citizen enters it", "System / Citizen", "Change takes effect only after this verification"),
            ("9", "Update registered mobile; notify registered email; write audit log", "System", "FR-56; actor, timestamp, old and new number (masked)"),
            ("10", "Continue login with OTP to the new mobile", "System", "FR-05"),
        ],
    )

    dept_reset_intro = insert_heading_after(
        reset_tbl,
        "Departmental mobile change — administrator only, no self-service (FR-65):",
        "Normal",
        parent=rbac_parent,
    )
    dept_reset_note = insert_paragraph_after(
        dept_reset_intro,
        "DSR Officers and Other Department users are not offered the FR-56 reset path, and no security "
        "questions are held for them. A departmental user who has lost the registered mobile is served by "
        "an authorised administrator, who records a reason and whose action is audit-logged. This keeps "
        "the officer's authentication factors under departmental control and removes any pre-login route "
        "to redirect an officer's OTP.",
        "Normal",
    )
    auth_flow_tbl = insert_table_after(
        dept_reset_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "Officer reports the lost / changed mobile to the authorised administrator", "DSR / Other Dept user", "Identity verified off-system per departmental procedure"),
            ("2", "Administrator opens User Management and locates the user by KGID", "Admin", "Authorised admin role only (FR-65)"),
            ("3", "Administrator enters the new mobile number and a reason", "Admin", "Reason mandatory (FR-65)"),
            ("4", "System verifies the new number by OTP before saving", "System / Officer", "OTP to the new mobile"),
            ("5", "Save — mobile updated; audit-logged; user notified on official email", "System", "FR-65, FR-13"),
            ("6", "Officer logs in with Username (KGID) + Captcha + OTP to new mobile + Biometrics", "Officer", "FR-06 / FR-07"),
        ],
    )

    login_post_intro = insert_heading_after(
        auth_flow_tbl,
        "Login post selection (DSR Officers — FR-52):",
        "Normal",
        parent=rbac_parent,
    )
    login_post_note = insert_paragraph_after(
        login_post_intro,
        "Applies only to DSR Officers. After Username (KGID) + Captcha + OTP (to mobile) + Biometrics "
        "succeed, the system "
        "loads the officer's active sanctioned-post occupancies. If more than one is active, the user "
        "must choose which post to work under for this session. Each choice shows Post Name and Office "
        "details so dual-charge / multi-office officers can pick the correct context.",
        "Normal",
    )
    login_post_tbl = insert_table_after(
        login_post_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "Authenticate (Username (KGID) + Captcha + OTP to mobile + Biometrics)", "DSR Officer / Auth", "FR-06"),
            ("2", "Load active post occupancies (Post + Office)", "System", "Relieved / ended occupancies excluded; reserved Transfer In excluded until Joining Date (FR-58, FR-61, FR-68)"),
            ("3a", "If exactly one active post — auto-select; continue to home", "System", "No selection UI; exclude posts before Joining Date 12:00 AM (FR-61)"),
            ("3b", "If two or more active posts — show post selection list labelled Role — Post Name — Office Name (Office Code)", "System / UI", "Mandatory; cannot skip (FR-52); only posts effective for login"),
            ("4", "Display each option as Post Name + Office Name (+ Office Code)", "UI", "Office details mandatory on each row"),
            ("5", "User selects exactly one assigned post and confirms", "DSR Officer", "Assigned post stored on session (FR-52)"),
            ("6", "Resolve Module Function claims from assigned post; enter home", "UM / Auth", "FR-38; FR-53 not offered at login"),
            ("7", "Display assigned Post with mapped Role(s) on home/header", "UI", "FR-54"),
        ],
    )

    switch_intro = insert_heading_after(
        login_post_tbl,
        "Additional charge of unoccupied subordinate post after login (DSR Officers — FR-53) and display (FR-54):",
        "Normal",
        parent=rbac_parent,
    )
    switch_note = insert_paragraph_after(
        switch_intro,
        "Additional charge is temporary in-office cover for a wholly unoccupied subordinate post — like "
        "holding an extra desk for the session — and does not require an order. It happens only after "
        "login under the assigned post (FR-52), without logout. A subordinate post qualifies only when "
        "nobody is posted against it at that Office (Occupied = 0 — FR-66(b)). While additional charge "
        "is active, session privileges derive from the additional charge post only — assigned-post "
        "privileges do not apply (FR-38). At most one additional charge post may be active at a time; "
        "the officer may switch back to assigned-post context without logging out to restore "
        "assigned-post privileges.",
        "Normal",
    )
    switch_tbl = insert_table_after(
        switch_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "Officer is logged in under assigned Post + Office (FR-52)", "DSR Officer / System", "FR-06, FR-52"),
            ("2", "From header or control — open Additional charge", "DSR Officer / UI", "Post-login only; no logout"),
            ("3", "Read Hierarchy Master children of assigned Post at same Office", "System", "FR-43"),
            ("4", "List wholly unoccupied subordinate posts (cascade rules; FR-66(b))", "System", "Partial vacancy does not qualify"),
            ("5", "Show options: Role + Post (+ Office); option to clear additional charge", "UI", "At most one additional charge active"),
            ("6", "Officer selects subordinate post, switches back to assigned post, or clears additional charge", "DSR Officer", "No order required; one active context at a time"),
            ("7", "Recompute Module Function claims for active context only — additional charge OR assigned post (FR-38)", "UM / Auth", "SR acting as FDA loses SR signing until switch-back"),
            ("8", "Update header — assigned post + active context label (FR-54)", "UI", "Shows which context drives access"),
            ("9", "Audit log selection or reversion", "System", "Mandatory (FR-53)"),
        ],
    )
    switch_ex_intro = insert_heading_after(
        switch_tbl,
        "Illustrative example (SR at SRO Yeshwanthapura — additional charge after login) — wholly unoccupied only:",
        "Normal",
        parent=rbac_parent,
    )
    switch_ex_tbl = insert_table_after(
        switch_ex_intro,
        [
            "Office",
            "Post under the login post (Sub-Registrar)",
            "Sanctioned",
            "Occupied",
            "Offered for additional charge (FR-53)?",
            "Drop-down entry (Role — Post Name — Office Name (Office Code))",
        ],
        [
            (
                "SRO Yeshwanthapura",
                "Sub-Registrar (assigned post)",
                "1",
                "1 (self)",
                "Not applicable — assigned post is not additional charge",
                "Assigned: Sub-Registrar (SR) — Sub-Registrar — SRO Yeshwanthapura (OFF-SRO-YESH)",
            ),
            (
                "SRO Yeshwanthapura",
                "→ FDA (Enforcement)",
                "2",
                "1",
                "No — an FDA is posted, so the post is not unoccupied",
                "Not listed",
            ),
            (
                "SRO Yeshwanthapura",
                "→ → SDA (Enforcement) under FDA",
                "1",
                "0",
                "No — unreachable; the occupied FDA blocks the cascade",
                "Not listed",
            ),
            (
                "SRO Yeshwanthapura",
                "→ Data Entry Operator",
                "2",
                "0",
                "Yes — nobody is posted against it",
                "DEO — Data Entry Operator / SRO Yeshwanthapura",
            ),
            (
                "SRO Jayanagar",
                "→ FDA (Enforcement)",
                "2",
                "0",
                "Yes — nobody is posted, even though strength is 2",
                "FDA (Enforcement) — FDA (Enforcement) / SRO Jayanagar",
            ),
            (
                "SRO Jayanagar",
                "→ → SDA (Enforcement) under FDA",
                "1",
                "0",
                "Yes — cascade continues under the unoccupied FDA",
                "SDA (Enforcement) — SDA (Enforcement) / SRO Jayanagar",
            ),
            (
                "After login — additional charge",
                "—",
                "—",
                "—",
                "SR may switch to additional charge as DEO — SR signing not available until switch-back (FR-38)",
                "Active context: Additional charge — DEO — Data Entry Operator — SRO Yeshwanthapura (OFF-SRO-YESH); switch back for SR privileges (FR-54)",
            ),
        ],
    )

    display_intro = insert_heading_after(
        switch_ex_tbl,
        "Assigned post and additional charge display (FR-54):",
        "Normal",
        parent=rbac_parent,
    )
    display_note = insert_paragraph_after(
        display_intro,
        "On the home page and/or application header, the system shall always show the assigned Post Name, "
        "Office, and Role(s) from login (FR-52). When additional charge is active, the header shall "
        "clearly indicate which context is driving session privileges — assigned post or additional "
        "charge post — while still showing the assigned post for reference. Example without additional "
        "charge: \"Sub-Registrar (SR) — Sub-Registrar — SRO Yeshwanthapura (OFF-SRO-YESH)\". Example with "
        "additional charge active: \"Assigned: Sub-Registrar (SR) — Sub-Registrar — SRO Yeshwanthapura "
        "(OFF-SRO-YESH) | Active context: Additional charge — FDA (Enforcement) — FDA (Enforcement) — "
        "SRO Yeshwanthapura (OFF-SRO-YESH)\". The officer opens the additional-charge control from here "
        "without logging out.",
        "Normal",
    )

    sub3 = insert_heading_after(
        display_note, "6.5.3 Posts Master, Post–Role Mapping, and Sanctioned Posts", "Heading 3", parent=rbac_parent
    )
    note3 = insert_paragraph_after(
        sub3,
        "Posts Master is separate from Role Master. Posts shall be division-specific where the "
        "organisation chart distinguishes them (FR-49), so that Post–Role mapping does not over-provision "
        "access across Admin, Vigilance, and Enforcement. Sanctioned Posts Master references Posts Master "
        "(Post + Office + strength). Post–Role mapping links each Post to its corresponding unique Role(s). "
        "DSR Officers are assigned sanctioned posts that have available capacity; at login the officer "
        "selects one active post when multiple are assigned (FR-52). Effective access follows the active "
        "context only: assigned post when no additional charge is active, or additional charge post "
        "when FR-53 is active — not both simultaneously (FR-38). "
        "Over-capacity assignment is blocked (FR-26, FR-27, FR-45). FR-66: assignment and Transfer In "
        "need available capacity (Occupied < Sanctioned Strength), except FR-67 reservation against a "
        "recorded relieving; FR-53 additional charge needs a wholly "
        "unoccupied post (Occupied = 0). Occupied counts are refreshed after midnight (FR-68).",
        "Normal",
    )
    posts_intro = insert_paragraph_after(note3, "Posts Master — seed rows (division-specific; Application Admin maintains):", "Normal")
    posts_tbl = insert_table_after(
        posts_intro,
        ["Post Code", "Post Name", "Division / Branch", "Status"],
        [
            ("POST-ACS-SEC", "Additional Chief Secretary / Principal Secretary / Secretary", "Secretariat", "Active"),
            ("POST-IGR", "Inspector General of Registration & Commissioner of Stamps", "Top Management", "Active"),
            ("POST-DIGR-ADMIN", "DIGR (Admin, Law & Computers)", "Admin, Law & Computers", "Active"),
            ("POST-AIGR-ADMIN", "AIGR (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-HQA-ADMIN", "HQA (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-SR-ADMIN", "Sub Registrar (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-ACCT-SUP", "Accountant Superintendent (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-FDA-ADMIN", "FDA (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-SDA-ADMIN", "SDA (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-TYPIST-ADMIN", "Typist (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-HQA-RTI", "HQA (RTI)", "Admin, Law & Computers — RTI & Statistics", "Active"),
            ("POST-FDA-RTI", "FDA (RTI)", "Admin, Law & Computers — RTI & Statistics", "Active"),
            ("POST-SDA-RTI", "SDA (RTI)", "Admin, Law & Computers — RTI & Statistics", "Active"),
            ("POST-SI", "Statistical Inspector", "Admin, Law & Computers — RTI & Statistics", "Active"),
            ("POST-DIGR-VIG", "DIGR (Vigilance)", "Vigilance", "Active"),
            ("POST-LAW-OFF", "Law Officer", "Vigilance", "Active"),
            ("POST-AIGR-COMP", "AIGR (Computers)", "Computers", "Active"),
            ("POST-SI-INT", "System Integrator", "Computers — Development", "Active"),
            ("POST-APP-DEV", "Application Developer", "Computers — Development", "Active"),
            ("POST-PMU", "PMU", "Computers — Development", "Active"),
            ("POST-HQA-COMP", "HQA / Project Manager (Comp)", "Computers — Operations", "Active"),
            ("POST-SR-COMP", "Sub Registrar (Comp)", "Computers — Operations", "Active"),
            ("POST-FDA-COMP", "FDA (Computers)", "Computers — Operations", "Active"),
            ("POST-SDA-COMP", "SDA (Computers)", "Computers — Operations", "Active"),
            ("POST-DIGR-ENF", "DIGR (Enforcement)", "Enforcement", "Active"),
            ("POST-DRO", "District Registrar", "Enforcement", "Active"),
            ("POST-HQA-ENF", "HQA (Enforcement)", "Enforcement", "Active"),
            ("POST-SR", "Sub-Registrar", "Enforcement / Field", "Active"),
            ("POST-FDA-ENF", "FDA (Enforcement)", "Enforcement", "Active"),
            ("POST-SDA-ENF", "SDA (Enforcement)", "Enforcement", "Active"),
            ("POST-DEO", "Data Entry Operator", "Field / SRO", "Active"),
            ("POST-DIGR-INT", "DIGR (Intelligence)", "Intelligence & Audit", "Active"),
            ("POST-AIGR-AUDIT", "AIGR (Audit)", "Intelligence & Audit", "Active"),
            ("POST-HQA-AUDIT", "HQA (Audit)", "Intelligence & Audit", "Active"),
            ("POST-SUP-AUDIT", "Superintendent (Audit)", "Intelligence & Audit", "Active"),
            ("POST-FDA-AUDIT", "FDA (Audit)", "Intelligence & Audit", "Active"),
            ("POST-SDA-AUDIT", "SDA (Audit)", "Intelligence & Audit", "Active"),
            ("POST-TYPIST-AUDIT", "Typist (Audit)", "Intelligence & Audit", "Active"),
            ("POST-DIGR-CVC", "DIGR CVC", "DIGR CVC", "Active"),
            ("POST-JD-TP", "JD Town Planning", "DIGR CVC", "Active"),
        ],
    )
    map_pr_intro = insert_heading_after(
        posts_tbl,
        "Post–Role mapping — example rows (division-aligned; avoid cross-division over-provisioning). "
        "Post Name must match Posts Master character-for-character (FR-50). Incomplete — Domain Expert "
        "to complete all active Posts before go-live (FR-47):",
        "Normal",
        parent=rbac_parent,
    )
    post_role_tbl = insert_table_after(
        map_pr_intro,
        ["Post Code", "Post Name", "Mapped Role (unique name)", "Notes"],
        [
            ("POST-ACS-SEC", "Additional Chief Secretary / Principal Secretary / Secretary", "ACS / Principal Secretary / Secretary", "1:1 — hierarchy root"),
            ("POST-IGR", "Inspector General of Registration & Commissioner of Stamps", "IGR", "1:1"),
            ("POST-DIGR-ADMIN", "DIGR (Admin, Law & Computers)", "DIGR (Admin, Law & Computers)", "Not mapped to Vigilance/Enforcement DIGR roles"),
            ("POST-DIGR-VIG", "DIGR (Vigilance)", "DIGR (Vigilance)", "1:1 — division-specific"),
            ("POST-DIGR-ENF", "DIGR (Enforcement)", "DIGR (Enforcement)", "1:1 — division-specific"),
            ("POST-DIGR-INT", "DIGR (Intelligence)", "DIGR (Intelligence)", "1:1"),
            ("POST-AIGR-ADMIN", "AIGR (Admin)", "AIGR (Admin)", "1:1"),
            ("POST-AIGR-COMP", "AIGR (Computers)", "AIGR (Computers)", "1:1"),
            ("POST-AIGR-AUDIT", "AIGR (Audit)", "AIGR (Audit)", "1:1"),
            ("POST-FDA-ADMIN", "FDA (Admin)", "FDA (Admin)", "Unique role — not shared FDA"),
            ("POST-FDA-ENF", "FDA (Enforcement)", "FDA (Enforcement)", "Unique role"),
            ("POST-FDA-AUDIT", "FDA (Audit)", "FDA (Audit)", "Unique role"),
            ("POST-SDA-ADMIN", "SDA (Admin)", "SDA (Admin)", "Unique role"),
            ("POST-SDA-ENF", "SDA (Enforcement)", "SDA (Enforcement)", "Unique role"),
            ("POST-SR", "Sub-Registrar", "Sub-Registrar (SR)", "Field signing post"),
            ("POST-DEO", "Data Entry Operator", "DEO", "1:1"),
            ("POST-DRO", "District Registrar", "DRO", "1:1"),
        ],
    )
    sanc_intro = insert_heading_after(
        post_role_tbl,
        "Sanctioned Posts Master — references Posts Master (example rows; the last two columns show the "
        "two distinct vacancy tests of FR-66 — note that FDA (Enforcement) at SRO Yeshwanthapura has "
        "capacity for another appointment yet is not available for FR-53 acting):",
        "Normal",
        parent=rbac_parent,
    )
    sanc_tbl = insert_table_after(
        sanc_intro,
        [
            "Office Code",
            "Office Name",
            "Post Code",
            "Post Name",
            "Sanctioned Strength",
            "Occupied",
            "Vacancies remaining",
            "Available capacity — FR-66(a)",
            "Wholly unoccupied — FR-66(b)",
        ],
        [
            ("OFF-SRO-YESH", "SRO Yeshwanthapura", "POST-SR", "Sub-Registrar", "1", "1", "0", "No — full", "No"),
            ("OFF-SRO-YESH", "SRO Yeshwanthapura", "POST-FDA-ENF", "FDA (Enforcement)", "2", "1", "1", "Yes — may assign one more", "No — an FDA is posted"),
            ("OFF-SRO-YESH", "SRO Yeshwanthapura", "POST-SDA-ENF", "SDA (Enforcement)", "1", "0", "1", "Yes", "Yes"),
            ("OFF-SRO-YESH", "SRO Yeshwanthapura", "POST-DEO", "Data Entry Operator", "2", "0", "2", "Yes", "Yes"),
            ("OFF-SRO-JAY", "SRO Jayanagar", "POST-FDA-ENF", "FDA (Enforcement)", "2", "0", "2", "Yes", "Yes"),
            ("OFF-SRO-JAY", "SRO Jayanagar", "POST-SDA-ENF", "SDA (Enforcement)", "1", "0", "1", "Yes", "Yes"),
            ("OFF-DRO-MYS", "DRO Mysuru", "POST-DRO", "District Registrar", "1", "1", "0", "No — full", "No"),
            ("OFF-DRO-MYS", "DRO Mysuru", "POST-HQA-ENF", "HQA (Enforcement)", "1", "0", "1", "Yes", "Yes"),
            ("OFF-IGR", "IGR Office (Head Office)", "POST-DIGR-ADMIN", "DIGR (Admin, Law & Computers)", "1", "1", "0", "No — full", "No"),
            ("OFF-IGR", "IGR Office (Head Office)", "POST-DIGR-ENF", "DIGR (Enforcement)", "1", "0", "1", "Yes", "Yes"),
        ],
    )
    admin_posts_intro = insert_heading_after(
        sanc_tbl,
        "Application Admin maintenance:",
        "Normal",
        parent=rbac_parent,
    )
    admin_posts_tbl = insert_table_after(
        admin_posts_intro,
        ["Master / Table", "Application Admin may", "FR"],
        [
            ("Posts Master", "Add, edit, enable/disable posts", "FR-46"),
            ("Post–Role mapping", "Map one Post to one or more Roles; block sanction/assignment if unmapped (FR-47)", "FR-47"),
            ("Sanctioned Posts Master", "Set strength per Post per Office; view occupied count (including reserved Transfer In), remaining capacity, and whether the post is wholly unoccupied; occupied counts updated by occupancy refresh job", "FR-24, FR-25, FR-48, FR-67, FR-68"),
        ],
    )

    sub4 = insert_heading_after(
        admin_posts_tbl, "6.5.4 Role Master — Seed Roles by Role Category", "Heading 3", parent=rbac_parent
    )
    note4 = insert_paragraph_after(
        sub4,
        "All roles below reside in the single Role Master, differentiated by Role Category. Seed roles "
        "are listed for Citizen, Other Department, and DSR (Department). Application Admin may add further "
        "roles under the same Role Categories. Other Department users are assigned exactly one role with "
        "Role Category = Other Department. DSR Officers receive roles only through Post–Role mapping "
        "after sanctioned post assignment. Sanctioned posts reference Posts Master and apply only to "
        "DSR establishment posts.",
        "Normal",
    )
    citizen_intro = insert_paragraph_after(
        note4,
        "Roles with Role Category = Citizen:",
        "Normal",
    )
    citizen_tbl = insert_table_after(
        citizen_intro,
        ["Role", "Description"],
        [
            ("Citizen", "Default role assigned on instant self-registration; access to citizen portal services"),
            ("Marriage Applicant", "Apply for Hindu Marriage / Special Marriage registration and related citizen actions"),
            ("Document Registration Applicant", "Initiate and track document registration applications"),
            ("Certified Copy / EC Applicant", "Apply for Encumbrance Certificate (EC) and certified copy of registered documents"),
            ("Firm / Society Applicant", "Apply for firm / society related registration services where offered on the portal"),
            ("Stamp Duty / Challan Payer", "Pay stamp duty / registration fees and view payment history for citizen transactions"),
        ],
    )
    od_intro = insert_heading_after(
        citizen_tbl,
        "Roles with Role Category = Other Department:",
        "Normal",
        parent=rbac_parent,
    )
    other_tbl = insert_table_after(
        od_intro,
        ["Role", "Typical Parent Department", "Description"],
        [
            ("Revenue Verification Officer", "Revenue", "Verify land / revenue particulars linked to registration applications"),
            ("Bhoomi Cross-check User", "Revenue", "Cross-check Bhoomi / RTC data against registration records"),
            ("Treasury Payment Verifier", "Treasury / Finance", "Verify treasury / payment status for registration fees"),
            ("Khajane Reconciliation User", "Treasury / Finance", "Reconcile Khajane-II receipts with Kaveri transactions"),
            ("Police Enquiry Officer", "Police", "View / respond to police enquiry requests on registered documents"),
            ("FIR Verification User", "Police", "Verify FIR / crime particulars where required for registration workflow"),
            ("ULB Document Verifier", "Urban Local Body (ULB)", "Verify municipal / ULB documents submitted with applications"),
            ("Read-only MIS Viewer", "Any (cross-department)", "Read-only access to assigned MIS / dashboards; no transactional actions"),
            ("Document Upload User", "Any (cross-department)", "Upload supporting documents for assigned inter-department workflows"),
            ("Inter-Department Enquiry User", "Any (cross-department)", "Raise / respond to inter-department enquiries on applications"),
        ],
    )
    dsr_intro = insert_heading_after(
        other_tbl,
        "Roles with Role Category = DSR (Department) — unique role names (no duplicate FDA/SDA/Typist):",
        "Normal",
        parent=rbac_parent,
    )
    dsr_roles_tbl = insert_table_after(
        dsr_intro,
        ["Division", "Role (unique name)", "Description"],
        [
            ("Secretariat", "ACS / Principal Secretary / Secretary", "Additional Chief Secretary / Principal Secretary / Secretary — hierarchy root"),
            ("Top Management", "IGR", "Inspector General of Registration & Commissioner of Stamps"),
            ("Division 1 — Admin, Law & Computers", "DIGR (Admin, Law & Computers)", "Deputy IGR for Admin, Law & Computers"),
            ("Division 1 — Admin, Law & Computers", "AIGR (Admin)", "Assistant IGR (Admin)"),
            ("Division 1 — Admin, Law & Computers", "HQA (Admin)", "Head Quarter Assistant (Admin)"),
            ("Division 1 — Admin, Law & Computers", "Sub Registrar (Admin)", "Sub Registrar (Admin) — not to be confused with Sub-Registrar Office (SRO) office type"),
            ("Division 1 — Admin, Law & Computers", "Accountant Superintendent (Admin)", "Accountant Superintendent (Admin)"),
            ("Division 1 — Admin, Law & Computers", "FDA (Admin)", "First Division Assistant — Admin"),
            ("Division 1 — Admin, Law & Computers", "SDA (Admin)", "Second Division Assistant — Admin"),
            ("Division 1 — Admin, Law & Computers", "Typist (Admin)", "Typist — Admin"),
            ("Division 1 — Admin, Law & Computers", "HQA (RTI)", "Head Quarter Assistant (RTI)"),
            ("Division 1 — Admin, Law & Computers", "FDA (RTI)", "First Division Assistant — RTI & Statistics"),
            ("Division 1 — Admin, Law & Computers", "SDA (RTI)", "Second Division Assistant — RTI & Statistics"),
            ("Division 1 — Admin, Law & Computers", "Statistical Inspector", "Statistical Inspector"),
            ("Division 2 — Vigilance", "DIGR (Vigilance)", "Deputy IGR (Vigilance)"),
            ("Division 2 — Vigilance", "Law Officer", "Departmental Law Officer"),
            ("Division 3 — Computers", "AIGR (Computers)", "Assistant IGR (Computers)"),
            ("Division 3 — Computers", "System Integrator", "System Integrator / SI support"),
            ("Division 3 — Computers", "PMU", "Project Management Unit"),
            ("Division 3 — Computers", "Application Developer", "Application development support"),
            ("Division 3 — Computers", "HQA / Project Manager (Comp)", "Head Quarter Assistant / Project Manager (Computers)"),
            ("Division 3 — Computers", "Sub Registrar (Computers)", "Sub Registrar (Computers)"),
            ("Division 3 — Computers", "FDA (Computers)", "First Division Assistant — Computers"),
            ("Division 3 — Computers", "SDA (Computers)", "Second Division Assistant — Computers"),
            ("Division 4 — Enforcement", "DIGR (Enforcement)", "Deputy IGR (Enforcement)"),
            ("Division 4 — Enforcement", "DRO", "District Registrar / DRO"),
            ("Division 4 — Enforcement", "HQA (Enforcement)", "Head Quarter Assistant (Enforcement)"),
            ("Division 4 — Enforcement", "Sub-Registrar (SR)", "Sub-Registrar (office head / signing)"),
            ("Division 4 — Enforcement", "FDA (Enforcement)", "First Division Assistant — Enforcement"),
            ("Division 4 — Enforcement", "SDA (Enforcement)", "Second Division Assistant — Enforcement"),
            ("Division 4 — Enforcement", "DEO", "Data Entry Operator — SRO operational role"),
            ("Division 5 — Intelligence & Audit", "DIGR (Intelligence)", "Deputy IGR (Intelligence)"),
            ("Division 5 — Intelligence & Audit", "AIGR (Audit)", "Assistant IGR (Audit)"),
            ("Division 5 — Intelligence & Audit", "HQA (Audit)", "Head Quarter Assistant (Audit)"),
            ("Division 5 — Intelligence & Audit", "Superintendent (Audit)", "Superintendent (Audit)"),
            ("Division 5 — Intelligence & Audit", "FDA (Audit)", "First Division Assistant — Audit"),
            ("Division 5 — Intelligence & Audit", "SDA (Audit)", "Second Division Assistant — Audit"),
            ("Division 5 — Intelligence & Audit", "Typist (Audit)", "Typist — Audit"),
            ("Division 6 — CVC", "DIGR CVC", "Deputy IGR (CVC)"),
            ("Division 6 — CVC", "JD Town Planning", "Joint Director, Town Planning"),
        ],
    )

    sub5 = insert_heading_after(
        dsr_roles_tbl, "6.5.5 DSR Officer Post Assignment", "Heading 3", parent=rbac_parent
    )
    primary_note = insert_paragraph_after(
        sub5,
        "DSR Officers are assigned one or more sanctioned posts that have available capacity (Posts Master "
        "via Sanctioned Posts Master). There is no Primary or Secondary role/post assignment. At login, if multiple "
        "active posts exist, the officer selects one post shown with office details (FR-52); session "
        "roles = roles mapped to the assigned post at login; may take additional charge of a wholly "
        "unoccupied subordinate after login without logout (FR-53, FR-38). Division-specific posts "
        "prevent over-provisioning across Admin/Vigilance/Enforcement (FR-47, FR-49). Other Department "
        "users and Citizens use Role Master directly (see Sections 6.5.1, 6.5.4 and 6.6.2).",
        "Normal",
    )
    primary_tbl = insert_table_after(
        primary_note,
        ["User Category", "What is assigned", "Roles in session", "End Date", "Notes"],
        [
            (
                "DSR Officer",
                "One or more sanctioned posts with available capacity (Post + Office) — FR-66(a)",
                "Active context: assigned post (FR-52) or additional charge (FR-53) — one at a time per FR-38",
                "Optional per post occupancy",
                "Multi-post → choose assigned post at login; switch context after login without logout",
            ),
            (
                "Other Department",
                "Exactly one role from Role Master",
                "That single role",
                "Optional account End Date",
                "No Posts Master / sanctioned posts",
            ),
            (
                "Citizen",
                "Citizen role(s) from Role Master (e.g. on self-registration)",
                "Assigned Citizen role(s)",
                "N/A",
                "No sanctioned posts",
            ),
        ],
    )

    sub6 = insert_heading_after(
        primary_tbl, "6.5.6 Module, Function, Resource Masters and Access Enforcement", "Heading 3", parent=rbac_parent
    )
    note6 = insert_paragraph_after(
        sub6,
        "DSR organisational roles (SR, FDA (Enforcement), DEO, etc.) are not named after application "
        "services. Access is modelled as: User → Role(s) → Module Function(s) → Resource(s) (API/URL). "
        "Registration of Documents is a single module (no online/offline classification). "
        "Role names in Role–Module–Function mapping must match Role Master exactly (FR-50). "
        "Application Admin is a system-level actor outside Role Master / Role–Module–Function mapping "
        "(FR-51) and maintains all masters in this section (FR-42).",
        "Normal",
    )
    hier_intro = insert_paragraph_after(
        note6,
        "Privilege hierarchy (masters):",
        "Normal",
    )
    hier_tbl = insert_table_after(
        hier_intro,
        ["Level", "Master", "Maintained by", "Purpose"],
        [
            ("1", "Module Master", "Application Admin", "Business modules (Registration of Documents, Marriage Registration, …)"),
            ("2", "Module Function Master", "Application Admin", "Functions under each module (VIEW, ADD, APPROVE, SIGN, …)"),
            ("3", "Resource Master", "Application Admin", "APIs and URLs linked to each Module Function"),
            ("4", "Role–Module–Function mapping", "Application Admin", "Which roles may perform which Module Functions"),
        ],
    )

    mod_intro = insert_heading_after(
        hier_tbl, "Module Master — seed rows:", "Normal", parent=rbac_parent
    )
    mod_tbl = insert_table_after(
        mod_intro,
        ["Module Code", "Module Name", "Description", "Status"],
        [
            ("MOD-DOC-REG", "Registration of Documents", "Registration of documents", "Active"),
            ("MOD-MARRIAGE", "Marriage Registration", "Marriage registration (Hindu Marriage and Special Marriage)", "Active"),
            ("MOD-EC", "Encumbrance Search", "Search and issue of Encumbrance Certificate (EC)", "Active"),
            ("MOD-CC", "Certified Copy", "Application and issue of certified copies of registered documents", "Active"),
            ("MOD-STAMP", "Stamp Duty / Payments", "Stamp duty assessment, challan, and payment reconciliation", "Active"),
            ("MOD-FIRM", "Firm / Society Registration", "Firm and society related registration services", "Active"),
            ("MOD-UM", "User Management", "Identity, roles, posts, modules, functions, resources, and mappings", "Active"),
            ("MOD-MIS", "MIS / Dashboards", "Operational and management reporting", "Active"),
        ],
    )

    fn_intro = insert_heading_after(
        mod_tbl, "Module Function Master — example rows:", "Normal", parent=rbac_parent
    )
    fn_tbl = insert_table_after(
        fn_intro,
        ["Function Code", "Module", "Function Name", "Description"],
        [
            ("FN-DOC-VIEW", "Registration of Documents", "VIEW", "View document registration applications / records"),
            ("FN-DOC-ADD", "Registration of Documents", "ADD", "Create / initiate document registration"),
            ("FN-DOC-EDIT", "Registration of Documents", "EDIT", "Edit draft / in-progress applications"),
            ("FN-DOC-APPROVE", "Registration of Documents", "APPROVE", "Approve application for registration"),
            ("FN-DOC-SIGN", "Registration of Documents", "SIGN", "Digitally sign registered document"),
            ("FN-DOC-PRINT", "Registration of Documents", "PRINT", "Print registration extracts / slips"),
            ("FN-MAR-VIEW", "Marriage Registration", "VIEW", "View marriage applications"),
            ("FN-MAR-ADD", "Marriage Registration", "ADD", "Create marriage application"),
            ("FN-MAR-APPROVE", "Marriage Registration", "APPROVE", "Approve / register marriage"),
            ("FN-EC-APPLY", "Encumbrance Search", "APPLY", "Apply for EC"),
            ("FN-EC-ISSUE", "Encumbrance Search", "ISSUE", "Issue / download EC"),
            ("FN-CC-APPLY", "Certified Copy", "APPLY", "Apply for certified copy"),
            ("FN-CC-ISSUE", "Certified Copy", "ISSUE", "Issue / download certified copy"),
            ("FN-UM-ADMIN", "User Management", "ADMIN", "Maintain users, roles, modules, functions, resources, mappings"),
        ],
    )

    res_intro = insert_heading_after(
        fn_tbl, "Resource Master (API / URL) — example rows:", "Normal", parent=rbac_parent
    )
    res_tbl = insert_table_after(
        res_intro,
        ["Resource Code", "Type", "Method", "Path / URL Pattern", "Module Function", "Is Public", "Status"],
        [
            ("RES-DOC-GET", "API", "GET", "/api/v1/documents/{id}", "FN-DOC-VIEW", "No", "Active"),
            ("RES-DOC-LIST", "API", "GET", "/api/v1/documents", "FN-DOC-VIEW", "No", "Active"),
            ("RES-DOC-CREATE", "API", "POST", "/api/v1/documents", "FN-DOC-ADD", "No", "Active"),
            ("RES-DOC-UPDATE", "API", "PUT", "/api/v1/documents/{id}", "FN-DOC-EDIT", "No", "Active"),
            ("RES-DOC-APPROVE", "API", "POST", "/api/v1/documents/{id}/approve", "FN-DOC-APPROVE", "No", "Active"),
            ("RES-DOC-SIGN", "API", "POST", "/api/v1/documents/{id}/sign", "FN-DOC-SIGN", "No", "Active"),
            ("RES-UI-DOC-VIEW", "URL", "GET", "/app/documents/view", "FN-DOC-VIEW", "No", "Active"),
            ("RES-UI-DOC-CREATE", "URL", "GET", "/app/documents/create", "FN-DOC-ADD", "No", "Active"),
            ("RES-MAR-CREATE", "API", "POST", "/api/v1/marriages", "FN-MAR-ADD", "No", "Active"),
            ("RES-MAR-APPROVE", "API", "POST", "/api/v1/marriages/{id}/approve", "FN-MAR-APPROVE", "No", "Active"),
            ("RES-EC-APPLY", "API", "POST", "/api/v1/encumbrance/applications", "FN-EC-APPLY", "No", "Active"),
            ("RES-CC-ISSUE", "API", "POST", "/api/v1/certified-copies/{id}/issue", "FN-CC-ISSUE", "No", "Active"),
            ("RES-UM-ROLES", "API", "PUT", "/api/v1/admin/roles/{id}/module-functions", "FN-UM-ADMIN", "No", "Active"),
            ("RES-PUBLIC-HEALTH", "API", "GET", "/api/v1/public/health", "—", "Yes", "Active"),
        ],
    )
    res_attr_intro = insert_heading_after(
        res_tbl,
        "Resource Master attributes (Application Admin maintains):",
        "Normal",
        parent=rbac_parent,
    )
    insert_table_after(
        res_attr_intro,
        ["Attribute", "Description", "Mandatory"],
        [
            ("Resource Code", "System or admin-defined unique identifier", "Yes"),
            ("Type", "API or URL", "Yes"),
            ("HTTP Method", "GET, POST, PUT, DELETE, … (for APIs)", "Yes for API"),
            ("Path / URL Pattern", "Ant-style or regex path pattern matched at runtime", "Yes"),
            ("Module Function", "Required function code for non-public resources (FR-41)", "Yes unless Is Public"),
            ("Is Public", "Yes = accessible without authentication; default No (deny-by-default)", "Yes"),
            ("Match Precedence", "When multiple patterns match, most specific (longest) pattern wins", "System rule (FR-41)"),
            ("Is Active", "Enable / disable resource without deleting history", "Yes"),
        ],
    )

    map_intro = insert_heading_after(
        res_attr_intro,
        "Role–Module–Function mapping — example rows (exact Role Master names only; FR-50):",
        "Normal",
        parent=rbac_parent,
    )
    map_note = insert_paragraph_after(
        map_intro,
        "Every Role column value must match an active Role Master name character-for-character. "
        "Document-registration functions for First Division Assistants are seeded for FDA (Enforcement) "
        "only (Enforcement / SRO document-registration path). FDA (Admin), FDA (RTI), FDA (Computers), "
        "and FDA (Audit) are not granted these functions unless Domain Expert adds explicit separate rows. "
        "Sub-Registrar (SR) and DEO match §6.5.4 unchanged. Application Admin is not listed here (FR-51).",
        "Normal",
    )
    map_tbl = insert_table_after(
        map_note,
        ["Role", "Role Category", "Module Function", "Allowed"],
        [
            ("Citizen", "Citizen", "FN-DOC-ADD", "Yes"),
            ("Citizen", "Citizen", "FN-DOC-VIEW", "Yes"),
            ("Citizen", "Citizen", "FN-EC-APPLY", "Yes"),
            ("Citizen", "Citizen", "FN-CC-APPLY", "Yes"),
            ("Document Registration Applicant", "Citizen", "FN-DOC-ADD", "Yes"),
            ("Document Registration Applicant", "Citizen", "FN-DOC-VIEW", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-VIEW", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-APPROVE", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-SIGN", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-PRINT", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-MAR-APPROVE", "Yes"),
            ("DEO", "DSR", "FN-DOC-VIEW", "Yes"),
            ("DEO", "DSR", "FN-DOC-ADD", "Yes"),
            ("DEO", "DSR", "FN-DOC-EDIT", "Yes"),
            ("DEO", "DSR", "FN-MAR-ADD", "Yes"),
            ("FDA (Enforcement)", "DSR", "FN-DOC-VIEW", "Yes"),
            ("FDA (Enforcement)", "DSR", "FN-DOC-ADD", "Yes"),
            ("FDA (Enforcement)", "DSR", "FN-DOC-EDIT", "Yes"),
            ("Revenue Verification Officer", "Other Department", "FN-DOC-VIEW", "Yes"),
        ],
    )

    app_admin_note = insert_heading_after(
        map_tbl,
        "Application Admin (system actor — Option B): Application Admin is not a Role Master seed role "
        "and is not assigned via sanctioned posts or the normal user-creation flow. It is seeded at "
        "deployment as a system principal with authority to maintain User Management masters "
        "(Module, Module Function, Resource, Role–Module–Function, Posts, Hierarchy, Sanctioned Posts, "
        "etc.). FN-UM-ADMIN and related Resource Master rows (e.g. RES-UM-ROLES) describe the capability "
        "surface for that principal; they are not granted through a Role–Module–Function mapping row "
        "labelled Application Admin | DSR (FR-51).",
        "Normal",
        parent=rbac_parent,
    )

    enf_intro = insert_heading_after(
        app_admin_note, "How the application enforces access (runtime):", "Normal", parent=rbac_parent
    )
    enf_tbl = insert_table_after(
        enf_intro,
        ["Step", "Action", "Actor / Component", "Notes"],
        [
            ("1", "User authenticates (Username + Captcha + OTP to mobile; + Biometrics if DSR/Other Dept)", "User / Auth service", "Username = preferred Username (Citizen) or KGID (departmental) — FR-62"),
            ("1a", "If DSR Officer with multiple active posts — select assigned Post + Office for session", "User / UI", "Mandatory at login only (FR-52)"),
            ("1b", "Show assigned Post with Role(s); offer Additional charge control after login", "UI", "FR-54, FR-53; no logout to switch"),
            ("2", "Load session role(s) for active context — assigned post OR additional charge post", "Auth / UM service", "One context at a time per FR-38"),
            ("3", "Resolve Role–Module–Function mappings into session claims (function codes)", "UM / Auth service", "Functions for session roles (FR-38)"),
            ("4", "User calls an API or opens a URL", "Client / Browser", "Method + path"),
            ("5", "Look up Resource Master for matching Type + Method + Path pattern (most specific wins)", "API Gateway / Middleware", "If no match → deny (403) unless Is Public = Yes (FR-41)"),
            ("6", "Obtain required Module Function from the matched Resource", "API Gateway / Middleware", "e.g. POST /api/v1/documents/{id}/approve → FN-DOC-APPROVE"),
            ("7", "Allow if session claims include that Module Function; else HTTP 403", "API Gateway / Middleware", "FR-41; Application Admin system principal may hold FN-UM-ADMIN outside mapping (FR-51); failed attempts audit-logged"),
            ("8", "UI menus / buttons shown only for Module Functions present in session claims", "Front-end", "Same mapping; hide unauthorised screens"),
        ],
    )

    admin_intro = insert_heading_after(
        enf_tbl, "Application Admin maintenance (CRUD):", "Normal", parent=rbac_parent
    )
    admin_crud_tbl = insert_table_after(
        admin_intro,
        ["Master / Table", "Application Admin may", "Audit"],
        [
            ("Module Master", "Add, edit, enable/disable modules", "Mandatory (FR-42)"),
            ("Module Function Master", "Add, edit, enable/disable functions under a module", "Mandatory (FR-42)"),
            ("Resource Master", "Add, edit, enable/disable API and URL resources; link to Module Function", "Mandatory (FR-42)"),
            ("Role–Module–Function mapping", "Grant or revoke Module Functions for Role Master roles only (exact names; FR-50); not used for Application Admin (FR-51)", "Mandatory (FR-42)"),
            ("DSR Officer Hierarchy Master", "Add, edit, reorder, enable/disable hierarchy nodes (§6.5.7)", "Mandatory (FR-43)"),
            ("Office Hierarchy Master", "Add, edit, enable/disable offices; set parent (IGR → DRO → SRO) (§6.5.8)", "Mandatory (FR-59)"),
        ],
    )

    sub7 = insert_heading_after(
        admin_crud_tbl,
        "6.5.7 DSR Officer Hierarchy Master",
        "Heading 3",
        parent=rbac_parent,
    )
    note7 = insert_paragraph_after(
        sub7,
        "The system shall maintain the reporting hierarchy for Department officers (DSR) as a "
        "Hierarchy Master aligned to the Department organisational chart. Hierarchy nodes reference "
        "Posts from the Posts Master (not Roles) — reporting is position-based (FR-43). Each child "
        "Post has exactly one parent Post in the tree. The same parent–child links drive unoccupied "
        "subordinate post additional charge after login (FR-53, FR-66(b)) and Transfer out / "
        "relieving by hierarchy superior (FR-57, FR-58). Application Admin shall add, edit, reorder, "
        "enable/disable, and update hierarchy nodes. The hierarchy does not replace Sanctioned Posts "
        "or Role–Module–Function access control.",
        "Normal",
    )
    hier_seed_intro = insert_paragraph_after(
        note7,
        "Seed hierarchy (Post → Parent Post) — uses unique Posts Master codes. Enforcement SRO chain "
        "is Sub-Registrar → FDA (Enforcement) → SDA (Enforcement) under DRO so FR-53 / FR-57 examples "
        "apply. Vacancy for the FR-53 cascade is judged per Post per Office as Occupied = 0 (FR-66(b)):",
        "Normal",
    )
    hier_seed_tbl = insert_table_after(
        hier_seed_intro,
        ["Level", "Division / Branch", "Post Code", "Post Name", "Reports To (Parent Post Code)"],
        [
            ("0", "Secretariat", "POST-ACS-SEC", "Additional Chief Secretary / Principal Secretary / Secretary", "— (root)"),
            ("1", "Top Management", "POST-IGR", "Inspector General of Registration & Commissioner of Stamps", "POST-ACS-SEC"),
            ("2", "Admin, Law & Computers", "POST-DIGR-ADMIN", "DIGR (Admin, Law & Computers)", "POST-IGR"),
            ("3", "Admin — Administration", "POST-AIGR-ADMIN", "AIGR (Admin)", "POST-DIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-HQA-ADMIN", "HQA (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-SR-ADMIN", "Sub Registrar (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-ACCT-SUP", "Accountant Superintendent (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-FDA-ADMIN", "FDA (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-SDA-ADMIN", "SDA (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-TYPIST-ADMIN", "Typist (Admin)", "POST-AIGR-ADMIN"),
            ("3", "Admin — RTI & Statistics", "POST-HQA-RTI", "HQA (RTI)", "POST-DIGR-ADMIN"),
            ("4", "Admin — RTI & Statistics", "POST-FDA-RTI", "FDA (RTI)", "POST-HQA-RTI"),
            ("4", "Admin — RTI & Statistics", "POST-SDA-RTI", "SDA (RTI)", "POST-HQA-RTI"),
            ("4", "Admin — RTI & Statistics", "POST-SI", "Statistical Inspector", "POST-HQA-RTI"),
            ("2", "Vigilance", "POST-DIGR-VIG", "DIGR (Vigilance)", "POST-IGR"),
            ("3", "Vigilance", "POST-LAW-OFF", "Law Officer", "POST-DIGR-VIG"),
            ("2", "Computers", "POST-AIGR-COMP", "AIGR (Computers)", "POST-IGR"),
            ("3", "Computers — Development", "POST-SI-INT", "System Integrator", "POST-AIGR-COMP"),
            ("3", "Computers — Development", "POST-APP-DEV", "Application Developer", "POST-AIGR-COMP"),
            ("3", "Computers — Development", "POST-PMU", "PMU", "POST-AIGR-COMP"),
            ("3", "Computers — Operations", "POST-HQA-COMP", "HQA / Project Manager (Comp)", "POST-AIGR-COMP"),
            ("4", "Computers — Operations", "POST-SR-COMP", "Sub Registrar (Comp)", "POST-HQA-COMP"),
            ("4", "Computers — Operations", "POST-FDA-COMP", "FDA (Computers)", "POST-HQA-COMP"),
            ("4", "Computers — Operations", "POST-SDA-COMP", "SDA (Computers)", "POST-HQA-COMP"),
            ("2", "Enforcement", "POST-DIGR-ENF", "DIGR (Enforcement)", "POST-IGR"),
            ("3", "Enforcement", "POST-DRO", "District Registrar", "POST-DIGR-ENF"),
            ("3", "Enforcement", "POST-HQA-ENF", "HQA (Enforcement)", "POST-DIGR-ENF"),
            ("4", "Enforcement — under DRO (field)", "POST-SR", "Sub-Registrar", "POST-DRO"),
            ("5", "Enforcement — SRO", "POST-FDA-ENF", "FDA (Enforcement)", "POST-SR"),
            ("6", "Enforcement — SRO", "POST-SDA-ENF", "SDA (Enforcement)", "POST-FDA-ENF"),
            ("5", "Enforcement — SRO", "POST-DEO", "Data Entry Operator", "POST-SR"),
            ("2", "Intelligence & Audit", "POST-DIGR-INT", "DIGR (Intelligence)", "POST-IGR"),
            ("3", "Intelligence & Audit", "POST-AIGR-AUDIT", "AIGR (Audit)", "POST-DIGR-INT"),
            ("4", "Intelligence & Audit", "POST-HQA-AUDIT", "HQA (Audit)", "POST-AIGR-AUDIT"),
            ("4", "Intelligence & Audit", "POST-SUP-AUDIT", "Superintendent (Audit)", "POST-AIGR-AUDIT"),
            ("4", "Intelligence & Audit", "POST-FDA-AUDIT", "FDA (Audit)", "POST-AIGR-AUDIT"),
            ("4", "Intelligence & Audit", "POST-SDA-AUDIT", "SDA (Audit)", "POST-AIGR-AUDIT"),
            ("4", "Intelligence & Audit", "POST-TYPIST-AUDIT", "Typist (Audit)", "POST-AIGR-AUDIT"),
            ("2", "DIGR CVC", "POST-DIGR-CVC", "DIGR CVC", "POST-IGR"),
            ("3", "DIGR CVC", "POST-JD-TP", "JD Town Planning", "POST-DIGR-CVC"),
        ],
    )
    hier_attr_intro = insert_heading_after(
        hier_seed_tbl,
        "Hierarchy Master attributes (Application Admin maintains):",
        "Normal",
        parent=rbac_parent,
    )
    hier_attr_tbl = insert_table_after(
        hier_attr_intro,
        ["Attribute", "Description", "Mandatory"],
        [
            ("Hierarchy Node ID", "System-generated unique identifier", "Yes (system)"),
            ("Post", "Post from Posts Master (unique Post Code)", "Yes"),
            ("Parent Node", "Immediate reporting parent node (null for root); parent must also be a Post", "No for root"),
            ("Division / Branch", "Division label (e.g. Enforcement, Computers)", "Yes"),
            ("Display Order", "Sort order among siblings under the same parent", "Yes"),
            ("Level", "Depth in the tree (0 = root)", "Yes (system/derived)"),
            ("Is Active", "Enable / disable node without deleting history", "Yes"),
            ("Effective From / To", "Optional validity of the hierarchy link", "No"),
        ],
    )
    hier_admin_intro = insert_heading_after(
        hier_attr_tbl,
        "Maintain hierarchy — Application Admin workflow:",
        "Normal",
        parent=rbac_parent,
    )
    hier_wf_tbl = insert_table_after(
        hier_admin_intro,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → DSR Officer Hierarchy", "Application Admin", "FR-43"),
            ("2", "View tree / list of hierarchy nodes by division", "System", "Filter by Division, Active status"),
            ("3", "Add node — select Post from Posts Master, parent Post node, division, display order", "Application Admin", "Post must exist; FR-43"),
            ("4", "Edit node — change parent, order, division, or active flag", "Application Admin", "Cannot create circular parent links"),
            ("5", "Disable node (soft) if post/structure superseded", "Application Admin", "History retained; audit logged"),
            ("6", "Save — hierarchy available for organisational views and admin span rules", "System", "FR-44 seed structure; each Post has one parent"),
        ],
    )

    sub8 = insert_heading_after(
        hier_wf_tbl,
        "6.5.8 Office Hierarchy Master",
        "Heading 3",
        parent=rbac_parent,
    )
    office_note = insert_paragraph_after(
        sub8,
        "The system shall maintain the hierarchy of DSR offices separately from the Officer (Post) "
        "Hierarchy. IGR Office is the Head Office. Under the Head Office are District Registrar "
        "Offices. Under each District Registrar Office are Sub-Registrar Offices (FR-59). This "
        "hierarchy scopes Transfer Out / Relieving and Transfer In so a superior sees only offices "
        "under them; within that list, only posts whose immediate parent is the actor's session Post "
        "are actionable (FR-57, FR-60).",
        "Normal",
    )
    office_seed_intro = insert_paragraph_after(
        office_note,
        "Seed Office Hierarchy (illustrative):",
        "Normal",
    )
    office_seed_tbl = insert_table_after(
        office_seed_intro,
        ["Level", "Office Type", "Office Code (example)", "Office Name (example)", "Parent Office"],
        [
            ("0", "Head Office", "OFF-IGR", "IGR Office (Head Office)", "— (root)"),
            ("1", "District Registrar Office", "OFF-DRO-MYS", "DRO Mysuru", "OFF-IGR"),
            ("1", "District Registrar Office", "OFF-DRO-BLR", "DRO Bengaluru", "OFF-IGR"),
            ("2", "Sub-Registrar Office", "OFF-SRO-YESH", "SRO Yeshwanthapura", "OFF-DRO-BLR"),
            ("2", "Sub-Registrar Office", "OFF-SRO-JAY", "SRO Jayanagar", "OFF-DRO-BLR"),
            ("2", "Sub-Registrar Office", "OFF-SRO-MYS-E", "SRO Mysuru East", "OFF-DRO-MYS"),
        ],
    )
    office_span_intro = insert_heading_after(
        office_seed_tbl,
        "Office span for Transfer Out / Transfer In (examples):",
        "Normal",
        parent=rbac_parent,
    )
    insert_table_after(
        office_span_intro,
        ["Actor session Office", "Offices shown in the tree", "May act on (immediate parent only)"],
        [
            (
                "IGR Head Office (OFF-IGR)",
                "DRO offices and SRO offices under Head Office (statewide tree)",
                "Only posts that report immediately to IGR (e.g. DIGR) — not Sub-Registrar at an SRO",
            ),
            (
                "DRO Bengaluru (OFF-DRO-BLR)",
                "OFF-DRO-BLR + SRO Yeshwanthapura, SRO Jayanagar, …",
                "Sub-Registrar at those SROs (POST-SR reports to POST-DRO) — not FDA/DEO (those report to SR)",
            ),
            (
                "SRO Yeshwanthapura (OFF-SRO-YESH)",
                "OFF-SRO-YESH only",
                "FDA / DEO under Sub-Registrar at that SRO only",
            ),
        ],
    )

    admin_heading = next(
        p for p in doc.paragraphs if p.text.strip() == "6.7 Administrative User Management"
    )
    wf_heading_el = OxmlElement("w:p")
    admin_heading._element.addprevious(wf_heading_el)
    wf_heading = Paragraph(wf_heading_el, admin_heading._parent)
    wf_heading.style = "Heading 2"
    wf_heading.add_run("6.6 User Creation and Role Assignment Workflow")

    wf_intro = insert_paragraph_after(
        wf_heading,
        "The system shall provide dedicated step-by-step workflows to assign access during user creation. "
        "For DSR Officers and Other Department users the KGID is entered first because it becomes the "
        "Username (FR-62); the Citizen self-registration flow is set out in §6.5.2. "
        "For DSR Officers, assign one or more sanctioned posts with available capacity (roles via Post–Role "
        "mapping; no Primary/Secondary; at login the officer selects one post when multiple are active — FR-52). "
        "Transfer out / relieving and Transfer In are scoped to offices under the actor, then to "
        "posts whose immediate parent is the actor's session Post "
        "(§6.6.3–6.6.4, FR-57–FR-61, FR-67). Occupancies take effect or end via the occupancy "
        "refresh job after midnight (FR-68). "
        "For Other Department users, select exactly one role.",
        "Normal",
    )

    dsr_sub = insert_heading_after(wf_intro, "6.6.1 DSR Officer User Creation with Post Assignment", "Heading 3")
    dsr_tbl = insert_table_after(
        dsr_sub,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → Add DSR Department User", "Admin", "Authorised admin role only (FR-02)"),
            ("2", "Enter KGID — this becomes the Username", "Admin", "Validated for presence and uniqueness across User Master (FR-62, FR-64)"),
            ("2a", "Enter remaining particulars (name, official email ID, mobile, photo, ID proof)", "Admin", "Official email domain validated where configured (FR-64); email and mobile need not be unique (FR-04)"),
            ("2b", "No security questions captured for this category", "System", "Departmental users have no self-service mobile reset (FR-55, FR-65)"),
            ("3", "Assign one or more sanctioned posts with available capacity (Post + Office from Posts / Sanctioned Posts Masters)", "Admin", "At least one required; FR-17, FR-30, FR-45, FR-48, FR-66(a)"),
            ("4", "Optionally set End Date per post occupancy for time-bound deputations only", "Admin", "Requires Deputation Reason; ends after 11:59 PM IST on End Date via FR-68 — not a substitute for FR-57 relieving"),
            ("5", "System shows roles available via Post–Role mapping for each assigned post", "System", "FR-47; login will use selected post only (FR-52)"),
            ("6", "Upload approval letter where applicable", "Admin", "Should"),
            ("7", "Capture biometrics", "Admin / Officer", "Mandatory for DSR users (FR-06)"),
            ("8", "Review post occupancies and mapped roles; confirm", "Admin", "Multi-post users choose post at each login"),
            ("9", "Save — account active; occupied count updated per assigned post", "System", "Blocked if no post assigned or any selected post is at full sanctioned strength (FR-66(a))"),
        ],
    )

    other_sub = insert_heading_after(
        dsr_tbl, "6.6.2 Other Department User Creation with Role Assignment", "Heading 3", parent=rbac_parent
    )
    other_note = insert_paragraph_after(
        other_sub,
        "Other Department users are stored in the same User Master (User Category = Other Department). "
        "Exactly one role is selected from the Role Master filtered by Role Category = Other Department. "
        "There is no Primary/Secondary model for Other Department users. An optional End Date may be "
        "entered; if entered, the system shall deactivate the user on that date.",
        "Normal",
    )
    other_tbl = insert_table_after(
        other_note,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → Add Other Department User", "Admin", "Authorised admin role only (FR-03)"),
            ("2", "Enter KGID — this becomes the Username", "Admin", "Validated for presence and uniqueness across User Master (FR-62, FR-64)"),
            ("2a", "Enter remaining particulars (name, official email ID of parent department, mobile, photo, ID proof)", "Admin", "User Category = Other Department; official email validated (FR-64); no uniqueness on email or mobile (FR-04)"),
            ("2b", "No security questions captured for this category", "System", "No self-service mobile reset (FR-55, FR-65)"),
            ("3", "Enter parent department and designation", "Admin", "e.g. Revenue, Treasury, Police"),
            ("4", "Assign exactly one role from Role Master (Role Category = Other Department)", "Admin", "Mandatory; FR-29, FR-34"),
            ("5", "Optionally enter Account End Date", "Admin", "Not mandatory; FR-33"),
            ("6", "Upload authorisation letter / NOC from parent department", "Admin", "Should"),
            ("7", "Capture biometrics", "Admin / User", "Mandatory (FR-07)"),
            ("8", "Review role and End Date (if any) and confirm", "Admin", ""),
            ("9", "Save — account active with module access for the assigned role", "System", "Blocked if no role selected"),
            ("10", "If End Date was entered and is reached — deactivate user; block login", "System", "FR-33; audit logged"),
        ],
    )

    relieve_sub = insert_heading_after(
        other_tbl,
        "6.6.3 Transfer Out / Relieving Process (DSR Officers)",
        "Heading 3",
        parent=rbac_parent,
    )
    relieve_note = insert_paragraph_after(
        relieve_sub,
        "A DSR Officer may be relieved from currently assigned post occupancy(ies). When Transfer Out / "
        "Relieving is opened, the system shall list offices under the actor's session Office per the "
        "Office Hierarchy Master — IGR Head Office → District Registrar Offices → Sub-Registrar "
        "Offices (FR-59). Within that office span, the system shall list and allow relieving only "
        "where the actor's session Post is the immediate parent of the target Post in the Officer "
        "Hierarchy (§6.5.7, FR-57). Seeing a descendant office (for example an SRO under Head Office) "
        "does not allow relieving a Sub-Registrar there. Capture Relieving Date and Relieving Order. "
        "Mapping is removed after 11:59 PM of the Relieving Date by the occupancy refresh job "
        "(FR-58, FR-68).",
        "Normal",
    )
    relieve_ex_intro = insert_paragraph_after(
        relieve_note,
        "Examples (Office Hierarchy + Officer Hierarchy):",
        "Normal",
    )
    relieve_ex_tbl = insert_table_after(
        relieve_ex_intro,
        ["Actor (session)", "Offices shown", "May relieve (posts under actor)", "Example"],
        [
            (
                "District Registrar @ DRO Bengaluru",
                "DRO Bengaluru + its SRO offices only",
                "Sub-Registrar at those SROs (POST-SR reports to POST-DRO)",
                "DRO Bengaluru relieves Sub-Registrar of SRO Yeshwanthapura — not SRO Mysuru East",
            ),
            (
                "Sub-Registrar @ SRO Yeshwanthapura",
                "SRO Yeshwanthapura only",
                "FDA / DEO under Sub-Registrar at that SRO",
                "SR Yeshwanthapura relieves FDA at Yeshwanthapura — not another SRO",
            ),
            (
                "IGR / DIGR @ IGR Head Office",
                "DRO and SRO offices under Head Office (tree only)",
                "Only posts whose immediate parent is the actor's session Post (e.g. IGR relieves DIGR; DIGR (Enforcement) relieves DRO / HQA (Enforcement))",
                "Seeing SRO Yeshwanthapura in the tree does not allow IGR or DIGR to relieve the Sub-Registrar there",
            ),
            (
                "FDA (Enforcement) @ SRO Yeshwanthapura",
                "SRO Yeshwanthapura only",
                "SDA under FDA at that SRO",
                "FDA relieves SDA at same office",
            ),
        ],
    )
    relieve_wf_intro = insert_heading_after(
        relieve_ex_tbl,
        "Relieving workflow:",
        "Normal",
        parent=rbac_parent,
    )
    relieve_wf_tbl = insert_table_after(
        relieve_wf_intro,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open Transfer Out / Relieving", "Superior", "FR-57"),
            ("2", "System lists only offices in actor's office span (Office Hierarchy)", "System", "FR-59; own office + descendants"),
            ("3", "Within those offices, list post occupancies where actor Post is immediate parent", "System", "Officer Hierarchy FR-43"),
            ("4", "Select officer / post occupancy to relieve", "Superior", "Cannot see outside span; cannot select posts that do not report immediately to actor"),
            ("5", "Enter Relieving Date and Relieving Order (upload if applicable)", "Superior", "Mandatory (FR-57)"),
            ("6", "Confirm relieving", "Superior", "Audit-logged; occupancy stays active until end of Relieving Date"),
            ("7", "Until end of Relieving Date — occupancy remains active", "System", "Officer may still use post until day closes"),
            ("8", "Shortly after midnight following Relieving Date — occupancy refresh job de-allocates mapping", "System", "FR-58, FR-68; occupied count −1 unless a reserved Transfer In activates the same run (FR-67)"),
            ("9", "If no remaining post occupancies — block / limit login until new assignment", "System", "Per policy; audit"),
        ],
    )

    transfer_in_sub = insert_heading_after(
        relieve_wf_tbl,
        "6.6.4 Transfer In Process (DSR Officers)",
        "Heading 3",
        parent=rbac_parent,
    )
    transfer_in_note = insert_paragraph_after(
        transfer_in_sub,
        "Transfer In assigns a DSR Officer to a post within the superior's office span and only where "
        "the actor's session Post is the immediate parent of the target Post (same rules as Transfer "
        "Out — FR-57, FR-59, FR-43). The system shall capture Transfer Order / Reporting Order and "
        "Joining Date (FR-60). Posts with available capacity (Occupied < Sanctioned Strength — "
        "FR-66(a)) are selectable without a prior relieving. A post at full strength may receive a "
        "future-dated Transfer In only when relieving is already recorded for that Post + Office; "
        "recording it reserves capacity immediately (FR-67). A post need not be wholly unoccupied "
        "to receive a Transfer In. The officer may log in for that post only from 12:00 AM on the "
        "Joining Date (FR-61); the occupancy refresh job activates reserved occupancies (FR-68).",
        "Normal",
    )
    transfer_in_ex = insert_paragraph_after(
        transfer_in_note,
        "Example (handover with reservation): District Registrar at DRO Bengaluru relieves the "
        "Sub-Registrar of SRO Yeshwanthapura with Relieving Date 31-Aug-2026. The same DRO then "
        "records Transfer In of the incoming officer to that Sub-Registrar post with Joining Date "
        "01-Sep-2026 — even though the outgoing occupancy has not yet ended. Capacity is reserved "
        "immediately (FR-67). The incoming officer cannot log in under that post before "
        "01-Sep-2026 12:00 AM. Shortly after midnight on 01-Sep the occupancy refresh job (FR-68) "
        "de-allocates the outgoing officer and activates the incoming occupancy; Sanctioned Posts "
        "occupied count stays 1. IGR at Head Office cannot perform this Transfer In because "
        "Sub-Registrar does not report immediately to IGR.",
        "Normal",
    )
    transfer_in_tbl = insert_table_after(
        transfer_in_ex,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open Transfer In", "Superior", "FR-60"),
            ("2", "System lists offices in actor's office span", "System", "FR-59; tree only"),
            ("3", "Select target Post + Office where actor Post is immediate parent", "Superior", "Same parentage as FR-57"),
            ("4", "If post has available capacity — proceed; if at full strength — allow only when relieving already recorded and slot not already reserved", "System", "FR-66(a), FR-67; else blocked"),
            ("5", "Select / identify the officer being transferred in", "Superior", "Existing DSR user (or create then Transfer In per policy)"),
            ("6", "Enter Transfer Order / Reporting Order (number + upload if applicable)", "Superior", "Mandatory (FR-60)"),
            ("7", "Enter Joining Date", "Superior", "If reserving against relieving: on or after the day after Relieving Date (FR-67)"),
            ("8", "Confirm Transfer In — occupancy recorded as reserved if Joining Date is in the future; occupied count includes reservation", "System", "Audit-logged (FR-67)"),
            ("9", "Before 12:00 AM of Joining Date — login with that post blocked", "System", "FR-61"),
            ("10", "Shortly after midnight on Joining Date — occupancy refresh job activates occupancy and recalculates Sanctioned Posts occupied count", "System", "FR-68, FR-52"),
        ],
    )

    job_sub = insert_heading_after(
        transfer_in_tbl,
        "6.6.5 Occupancy refresh job (after midnight)",
        "Heading 3",
        parent=rbac_parent,
    )
    job_note = insert_paragraph_after(
        job_sub,
        "Post assignments on the user record and occupied counts on the Sanctioned Posts Master shall "
        "not be left to wait for a user to log in. A scheduled job shall run shortly after midnight "
        "each day (FR-68) so that Transfer Out, Transfer In, and occupancy End Date take effect on "
        "the correct calendar day before the first login.",
        "Normal",
    )
    insert_table_after(
        job_note,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Start occupancy refresh job shortly after 12:00 AM", "System", "FR-68; idempotent"),
            ("2", "De-allocate occupancies whose Relieving Date ended the previous calendar day", "System", "FR-57, FR-58"),
            ("3", "De-allocate occupancies whose optional End Date has been reached", "System", "FR-30"),
            ("4", "Activate reserved Transfer In occupancies whose Joining Date is today", "System", "FR-60, FR-61, FR-67"),
            ("5", "Recalculate occupied count, remaining capacity, and wholly-unoccupied flag on Sanctioned Posts Master for affected Post + Office rows", "System", "FR-48, FR-66"),
            ("6", "Refresh each affected officer's effective post assignments for login selection", "System", "FR-52; relieved posts no longer selectable; joining posts now selectable"),
            ("7", "Write occupancy-refresh audit (run time, ended, activated, before/after occupied counts)", "System", "Mandatory"),
            ("8", "If the job fails — raise operational alert; do not leave relieving/joining unenforced", "System", "Risk mitigation"),
        ],
    )

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v4.5"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
