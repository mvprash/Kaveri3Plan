# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v2.1.docx — unified Role/User master with category differentiation."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\User Management")
TEMPLATE = BASE / "Template" / "User_Management_Module_BRD_Template.docx"
DST = BASE / "BRD_User_Management_v2.1.docx"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_table_rows(table, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "D9E2F3")
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(cells[i], val)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph_after(ref_paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    para = Paragraph(new_p, ref_paragraph._parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_table_after(ref_paragraph, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_paragraph._element.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def insert_after_element(ref_el, parent, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_el.addnext(new_p)
    para = Paragraph(new_p, parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_heading_after(ref, text: str, style: str = "Heading 3", parent=None) -> Paragraph:
    if isinstance(ref, Paragraph):
        return insert_paragraph_after(ref, text, style)
    if parent is None:
        raise ValueError("parent required when ref is an XML element")
    return insert_after_element(ref, parent, text, style)


def insert_table_after_ref(ref_el, parent, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_el.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def build() -> Document:
    doc = Document(str(TEMPLATE))

    cover_map = {
        1: "BUSINESS REQUIREMENTS DOCUMENT",
        2: "User Management Module — KAVERI 3.0",
        3: "Prepared for: Department of Stamps & Registration, Government of Karnataka",
        4: "Version 2.1",
        5: "Date: 29 August 2026",
        6: "Prepared by: Nandha Kumar",
    }
    for idx, text in cover_map.items():
        replace_paragraph_text(doc.paragraphs[idx], text)

    replace_paragraph_text(
        doc.paragraphs[12],
        "This Business Requirements Document (BRD) defines the business requirements for the "
        "User Management Module of KAVERI 3.0 — the integrated platform of the Department of "
        "Stamps and Registration (DSR), Government of Karnataka. It describes a single unified "
        "Role Master and User Master differentiated by user category (Public/Citizen, DSR Officer, "
        "Other Department), OTP-based authentication (no password management), sanctioned posts "
        "with approved strength per role for each office, Primary and Secondary role assignment "
        "for DSR Officers, single-role assignment for Other Department users with optional account "
        "end date, and dedicated user-creation workflows. It serves as the agreed basis for design, "
        "development, testing, and sign-off.",
    )
    replace_paragraph_text(
        doc.paragraphs[14],
        "KAVERI 3.0 requires a centralized mechanism to manage user identities, roles, sanctioned "
        "posts, and access permissions. There shall be one Role Master and one User Master for all "
        "categories — Public users (Citizens), Department users (DSR Officers), and Other Department "
        "users — differentiated by User Category (and Role Category on roles). Authentication shall "
        "be passwordless — Username with OTP and Captcha for citizens, and Username with OTP, "
        "Captcha, and Biometrics for departmental users. DSR officers are mapped to sanctioned posts "
        "with one Primary Role and optional Secondary Roles. Other Department users are assigned "
        "exactly one role (no Primary/Secondary distinction) and may have an optional account end "
        "date that deactivates the user when reached. Password management is explicitly out of scope.",
    )

    scope_in = {
        17: "User registration for three categories in a single User Master: Public users (Citizens), Department users (DSR Officers), and Other Department users",
        18: "OTP-based authentication (login, logout) — Username + OTP + Captcha; biometrics for departmental users",
        19: "Single unified Role Master with Role Category differentiating Citizen, DSR, and Other Department roles (RBAC)",
        20: "Sanctioned posts master capturing the number of sanctioned posts (approved strength) for each DSR role at each office",
        21: "Primary Role and Secondary Role assignment for DSR Officers only; single role for Other Department users with optional account end date",
        22: "Dedicated step-by-step workflows for role assignment during user creation (DSR and Other Department)",
        23: "Administrative user management (create, edit, suspend, deactivate users)",
    }
    for idx, text in scope_in.items():
        replace_paragraph_text(doc.paragraphs[idx], text)
    insert_paragraph_after(
        doc.paragraphs[23],
        "Audit logging of user-related activities",
        "List Paragraph",
    )

    for p in doc.paragraphs:
        if "password resets" in p.text.lower():
            replace_paragraph_text(
                p,
                "Reduce support overhead related to account access issues through reliable OTP delivery.",
            )

    for p in doc.paragraphs:
        if "Eliminate password-related" in p.text:
            remove_paragraph(p)
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.3 Password Management":
            replace_paragraph_text(p, "6.3 OTP Login and Account Recovery")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.6 Administrative User Management":
            replace_paragraph_text(p, "6.7 Administrative User Management")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "Enforce role-based access control to protect sensitive data and functionality.":
            insert_paragraph_after(
                p,
                "Support Primary and Secondary role assignment for DSR Officers; single role and optional end-date deactivation for Other Department users.",
                "List Paragraph",
            )
            break

    for p in doc.paragraphs:
        if p.text.strip() == "The system shall provide a report of role and permission assignments across all users.":
            insert_paragraph_after(
                p,
                "The system shall provide a sanctioned post occupancy report showing, for each office, "
                "the number of sanctioned posts per role, occupied count, and vacant slots.",
                "List Paragraph",
            )
            break

    replace_table_rows(
        doc.tables[0],
        ["Document Control", ""],
        [
            ("Status", "Draft / In review"),
            ("Owner", "Nandha Kumar (BA) / Prashanth (PO)"),
            ("Reviewers", "Prabhakar Naik (Domain Expert); Kaveri IT Cell"),
            ("Approved By", "Prashanth (Product Owner)"),
        ],
    )

    replace_table_rows(
        doc.tables[1],
        ["Version", "Date", "Author", "Description"],
        [
            (
                "1.7",
                "28-Aug-2026",
                "Nandha Kumar",
                "Aligned to User Management BRD template; three user categories with OTP-only "
                "authentication (no password management); DSR division and role catalogue",
            ),
            (
                "1.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added sanctioned posts master: capture number of sanctioned posts (approved strength) "
                "for each role at each office; DSR user mapping to sanctioned posts only; "
                "over-capacity blocking and occupancy reporting",
            ),
            (
                "1.9",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Other Department role catalogue; Primary and Secondary role assignment; "
                "dedicated step-by-step user-creation workflows for DSR and Other Department users",
            ),
            (
                "2.0",
                "29-Aug-2026",
                "Nandha Kumar",
                "Primary/Secondary roles apply to DSR Officers only; Other Department users get "
                "exactly one role; optional account end date on Other Department user creation "
                "deactivates the user when reached (§6.6.2)",
            ),
            (
                "2.1",
                "29-Aug-2026",
                "Nandha Kumar",
                "Unified single Role Master and User Master for Citizens, DSR Officers, and Other "
                "Department users — differentiated by User Category / Role Category; no separate "
                "masters per category",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[2],
        ["Name / Role", "Department", "Responsibility"],
        [
            ("Prashanth", "DSR / Product", "Prioritizes requirements; approves scope"),
            ("Nandha Kumar", "Business Analysis", "Documents and validates requirements"),
            ("Prabhakar Naik", "Domain Expert", "Validates DSR roles, Other Department roles, sanctioned posts, and organizational structure"),
            ("Kaveri IT Cell", "Engineering", "Reviews technical feasibility"),
            ("Citizens (Public users)", "External", "Self-register and access citizen portal services"),
            (
                "DSR Officers & Other Department users",
                "Government",
                "Access departmental modules via OTP + biometrics",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[3],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-01", "The system shall support instant self-registration for Public users (Citizens) with no approval workflow.", "High"),
            ("FR-02", "The system shall allow Department users (DSR Officers) to be created only by authorised administrative roles.", "High"),
            ("FR-03", "The system shall allow Other Department users (officers/staff from other government departments) to be created only by authorised administrative roles.", "High"),
            ("FR-04", "The system shall prevent duplicate registrations using the same Username within a user category.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[4],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-05", "Public users (Citizens) shall authenticate using Username + OTP + Captcha on every login.", "High"),
            ("FR-06", "Department users (DSR Officers) shall authenticate using Username + OTP + Captcha + Biometrics on every login.", "High"),
            ("FR-07", "Other Department users shall authenticate using Username + OTP + Captcha + Biometrics on every login.", "High"),
            ("FR-08", "The system shall allow users to log out and terminate their active session.", "High"),
            ("FR-09", "The system shall not provide password-based login, password reset, or password change for any user category.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[5],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-10", "The system shall dispatch OTP to the user's registered mobile and/or email within 5 seconds of request.", "High"),
            ("FR-11", "The system shall validate Captcha before OTP dispatch for all user categories.", "High"),
            ("FR-12", "The system shall allow account recovery via OTP verification to registered mobile/email (no password reset).", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[6],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-13", "The system shall allow users to view and update their profile information (name, mobile, email).", "High"),
            ("FR-14", "The system shall allow users to upload and update a profile photo where applicable.", "Medium"),
            ("FR-15", "The system shall allow administrators to deactivate user accounts with reason and audit trail.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[7],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-16", "The system shall support RBAC using a single unified Role Master. Roles shall be differentiated by Role Category (Citizen, DSR, Other Department) as defined in Section 6.5.1. Users shall be stored in a single User Master differentiated by User Category.", "High"),
            (
                "FR-17",
                "The system shall allow administrators to assign roles to Department users (DSR Officers) "
                "during user creation: one mandatory Primary Role (mapped to a vacant sanctioned post — "
                "role + office) and one or more optional Secondary Roles filtered from the Role Master "
                "where Role Category = DSR.",
                "High",
            ),
            (
                "FR-28",
                "There shall be no separate Role Masters or User Masters for Citizens, DSR Officers, "
                "or Other Department users. All roles and users shall reside in one Role Master and one "
                "User Master respectively, differentiated by Role Category / User Category.",
                "High",
            ),
            (
                "FR-29",
                "During Other Department user creation, the system shall allow administrators to assign "
                "exactly one role from the Role Master filtered by Role Category = Other Department. "
                "Primary and Secondary role distinction does not apply to Other Department users.",
                "High",
            ),
            (
                "FR-30",
                "A Primary Role must be selected for each DSR Officer at account creation. "
                "The Primary Role cannot have an end date and represents the officer's substantive assignment. "
                "Primary Role does not apply to Other Department users or Citizens.",
                "High",
            ),
            (
                "FR-31",
                "The system shall allow administrators to assign multiple Secondary Roles to DSR Officers. "
                "Each Secondary Role must have a mandatory end date; the system shall automatically "
                "revoke Secondary Role access on expiry. Secondary Roles do not apply to Other Department users or Citizens.",
                "High",
            ),
            (
                "FR-33",
                "During Other Department user creation, the system shall allow an optional End Date. "
                "If an End Date is entered, the system shall automatically deactivate the user account "
                "on that date and block login thereafter.",
                "High",
            ),
            (
                "FR-34",
                "When assigning a role to a user, the system shall present only roles whose Role Category "
                "matches the user's User Category (Citizen roles for Citizens, DSR roles for DSR Officers, "
                "Other Department roles for Other Department users).",
                "High",
            ),
            ("FR-18", "The system shall restrict access to features and data based on the user's assigned role(s).", "High"),
            ("FR-19", "The system shall maintain DSR organizational divisions and DSR roles within the unified Role Master (Role Category = DSR) as defined in Section 6.5.1.", "High"),
            (
                "FR-24",
                "The system shall maintain a sanctioned posts master listing all DSR roles/posts "
                "(e.g. IGR, DIGR, AIGR, Sub-Registrar, FDA, SDA, DRO, HQA) with the number of "
                "sanctioned posts (approved strength) for each role at each office. Sanctioned "
                "strength applies only to roles with Role Category = DSR.",
                "High",
            ),
            (
                "FR-25",
                "The system shall allow authorised administrators to configure and update sanctioned "
                "strength per role per office, with audit trail of changes.",
                "High",
            ),
            (
                "FR-26",
                "DSR department users shall be mapped only to posts defined in the sanctioned posts "
                "master. Assignment to an unlisted post or exceeding sanctioned strength for a "
                "role at an office shall be blocked.",
                "High",
            ),
            (
                "FR-27",
                "The system shall display vacant vs occupied sanctioned posts per role per office "
                "and prevent over-capacity assignment.",
                "High",
            ),
            (
                "FR-32",
                "The system shall provide dedicated step-by-step workflows for role assignment during "
                "user creation for DSR Officers (Section 6.6.1) and Other Department users (Section 6.6.2).",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[8],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-20", "The system shall allow administrators to create, edit, suspend, and deactivate user accounts by category.", "High"),
            ("FR-21", "The system shall allow administrators to search and filter users by category, role, office, division, and status.", "Medium"),
            ("FR-22", "The system shall log all administrative actions performed on user accounts.", "High"),
            ("FR-23", "The system shall notify a user via SMS/email when their account is created, suspended, or deactivated.", "Medium"),
        ],
    )

    replace_table_rows(
        doc.tables[9],
        ["Category", "Requirement"],
        [
            ("Security", "No password storage — authentication is OTP-based only for all user categories."),
            ("Security", "All data in transit shall be encrypted using TLS 1.2 or higher."),
            ("Security", "Biometric data for departmental users shall comply with Aadhaar Act, 2016 and UIDAI guidelines."),
            ("Performance", "OTP dispatch shall complete within 5 seconds of user request."),
            ("Performance", "Login and authentication requests shall complete within 2 seconds after OTP/biometric verification."),
            ("Availability", "The module shall maintain 99.9% uptime, excluding scheduled maintenance."),
            ("Usability", "Registration and OTP login workflows shall be completable on desktop and mobile browsers."),
            ("Auditability", "All create, update, delete, login, and access-control actions shall be logged with timestamp and actor."),
            ("Compliance", "The module shall comply with Karnataka e-Governance, MeitY/CERT-In, and applicable data protection norms."),
        ],
    )

    replace_table_rows(
        doc.tables[10],
        ["Risk", "Impact", "Mitigation"],
        [
            ("OTP delivery failure (SMS/email)", "High", "Multi-channel OTP delivery; retry mechanism; admin override procedure."),
            ("Biometric device unavailability", "Medium", "Define fallback procedure for DSR and Other Department users per security policy."),
            ("Over-capacity post assignment", "High", "Enforce sanctioned strength validation; block assignment when role at office is full."),
            ("Secondary role not expired on end date", "Medium", "Scheduled job to auto-revoke expired Secondary Roles for DSR Officers; audit alert."),
            ("Other Department end date not enforced", "Medium", "Scheduled job to deactivate Other Department users on optional End Date; audit alert."),
            ("Scope creep beyond OTP-only auth", "Medium", "Maintain strict change-control; no password management in KAVERI 3.0."),
            ("Delayed stakeholder sign-off", "Low", "Set clear review deadlines and escalation path."),
        ],
    )

    replace_table_rows(
        doc.tables[11],
        ["Term", "Definition"],
        [
            ("BRD", "Business Requirements Document"),
            ("RBAC", "Role-Based Access Control"),
            ("OTP", "One-Time Password — primary authentication credential (no static password)"),
            ("DSR", "Department of Stamps and Registration, Government of Karnataka"),
            ("UAT", "User Acceptance Testing"),
            ("IGR", "Inspector General of Registration"),
            ("DIGR", "Deputy Inspector General of Registration"),
            ("AIGR", "Assistant Inspector General of Registration"),
            ("Sanctioned post", "A DSR role/post defined in the posts master with approved strength (headcount) at a specific office"),
            ("Sanctioned strength", "The approved number of occupants for a given DSR role at a given office"),
            ("Office", "A concrete DSR office instance (e.g. SRO Yeshwanthapura, DRO Mysuru) in the organisational hierarchy"),
            ("Role Master", "Single master of all roles for Citizens, DSR Officers, and Other Department users; differentiated by Role Category"),
            ("User Master", "Single master of all users across categories; differentiated by User Category"),
            ("Role Category", "Attribute on a role in the Role Master: Citizen, DSR, or Other Department"),
            ("User Category", "Attribute on a user in the User Master: Public (Citizen), DSR Officer, or Other Department"),
            ("Other Department role", "A role in the Role Master with Role Category = Other Department; exactly one such role is assigned per Other Department user"),
            ("Primary Role", "The mandatory substantive role assigned to a DSR Officer at user creation; cannot have an end date; does not apply to Other Department users or Citizens"),
            ("Secondary Role", "An additional role assigned to a DSR Officer; must have a mandatory end date and auto-expires; does not apply to Other Department users or Citizens"),
            ("Parent department", "The government department to which an Other Department user belongs (e.g. Revenue, Treasury, Police)"),
            ("Account End Date", "Optional date on an Other Department user account; if set, the system deactivates the user on that date"),
        ],
    )

    replace_table_rows(
        doc.tables[12],
        ["Name", "Role", "Signature", "Date"],
        [
            ("Prashanth", "Product Owner", "", ""),
            ("Prabhakar Naik", "Domain Expert", "", ""),
            ("Kaveri IT Cell Lead", "IT Security / Engineering", "", ""),
        ],
    )

    rbac_fr_tbl_el = doc.tables[7]._tbl
    rbac_heading = next(p for p in doc.paragraphs if p.text.strip() == "6.5 Role-Based Access Control (RBAC)")
    rbac_parent = rbac_heading._parent
    rbac_fr_tbl_el.getparent().remove(rbac_fr_tbl_el)
    rbac_heading._element.addnext(rbac_fr_tbl_el)

    sub1 = insert_heading_after(
        rbac_fr_tbl_el, "6.5.1 Unified Role Master and User Master", "Heading 3", parent=rbac_parent
    )
    note1 = insert_paragraph_after(
        sub1,
        "The system shall maintain a single Role Master and a single User Master for all user types. "
        "There shall be no separate masters for Citizens, DSR Officers, or Other Department users. "
        "Differentiation is by Role Category on roles and User Category on users. When assigning a role, "
        "the system shall filter the Role Master to roles matching the user's User Category (FR-34).",
        "Normal",
    )
    cat_diff_tbl = insert_table_after(
        note1,
        ["Master", "Differentiator", "Values", "Purpose"],
        [
            (
                "User Master",
                "User Category",
                "Public (Citizen); DSR Officer; Other Department",
                "One user store; category drives auth, creation path, and eligible roles",
            ),
            (
                "Role Master",
                "Role Category",
                "Citizen; DSR; Other Department",
                "One role store; category filters which roles can be assigned to which users",
            ),
        ],
    )
    div_intro = insert_heading_after(
        cat_diff_tbl,
        "DSR roles in the Role Master (Role Category = DSR) are organised by division as follows:",
        "Normal",
        parent=rbac_parent,
    )
    div_tbl = insert_table_after(
        div_intro,
        ["Division", "Roles (Role Category = DSR)"],
        [
            ("Top Management", "IGR"),
            (
                "Division 1 — Admin, Law & Computers",
                "DIGR (Admin, Law & Computers), AIGR (Admin), HQA (Admin), SRO (Admin), FDA, SDA",
            ),
            ("Division 2 — Vigilance", "DIGR (Vigilance), Law Officer, HQA (RTI)"),
            (
                "Division 3 — Computers",
                "AIGR (Computers), System Integrator, PMU, Application Developer, SRO (Comp)",
            ),
            ("Division 4 — Enforcement", "DIGR (Enforcement), DRO, HQA, SRO, FDA, SDA"),
            (
                "Division 5 — Intelligence & Audit",
                "DIGR (Intelligence), AIGR (Audit), HQA (Audit), Superintendent (Audit)",
            ),
            ("Division 6 — CVC", "DIGR CVC, JD Town Planning"),
        ],
    )

    sub2 = insert_heading_after(
        div_tbl, "6.5.2 User Categories and Authentication", "Heading 3", parent=rbac_parent
    )
    note2 = insert_paragraph_after(
        sub2,
        "All users are stored in the single User Master, differentiated by User Category. "
        "All user categories authenticate without passwords. OTP is the sole login credential.",
        "Normal",
    )
    cat_tbl = insert_table_after(
        note2,
        ["User Category", "Description", "Authentication"],
        [
            ("Public users (Citizens)", "Citizens accessing Kaveri portal services — User Category = Public (Citizen)", "Username + OTP + Captcha"),
            ("Department users (DSR Officers)", "Officers and staff of DSR — User Category = DSR Officer", "Username + OTP + Captcha + Biometrics"),
            (
                "Other Department users",
                "Officers/staff from other government departments — User Category = Other Department",
                "Username + OTP + Captcha + Biometrics",
            ),
        ],
    )

    sub3 = insert_heading_after(
        cat_tbl, "6.5.3 Sanctioned Posts per Office", "Heading 3", parent=rbac_parent
    )
    note3 = insert_paragraph_after(
        sub3,
        "The system shall maintain a sanctioned posts master that records, for each office in the DSR "
        "hierarchy, the number of sanctioned posts (approved strength) for each role. DSR Officers may "
        "be assigned only to vacant sanctioned posts. The system shall block assignment when sanctioned "
        "strength for a role at an office is already fully occupied.",
        "Normal",
    )
    example_note = insert_paragraph_after(note3, "Example (illustrative):", "Normal")
    sanc_tbl = insert_table_after(
        example_note,
        ["Office", "Role", "Sanctioned Posts", "Occupied", "Vacant"],
        [
            ("SRO Yeshwanthapura", "Sub-Registrar (SR)", "1", "1", "0"),
            ("SRO Yeshwanthapura", "FDA", "2", "1", "1"),
            ("SRO Yeshwanthapura", "SDA", "1", "0", "1"),
            ("DRO Mysuru", "District Registrar (DR)", "1", "1", "0"),
            ("DRO Mysuru", "HQA", "1", "0", "1"),
        ],
    )

    sub4 = insert_heading_after(
        sanc_tbl, "6.5.4 Role Master — Other Department and Citizen Roles", "Heading 3", parent=rbac_parent
    )
    note4 = insert_paragraph_after(
        sub4,
        "Other Department and Citizen roles are maintained in the same Role Master (not separate masters), "
        "with Role Category = Other Department or Citizen respectively. Other Department users are assigned "
        "exactly one role with Role Category = Other Department. Primary and Secondary role distinction "
        "does not apply to Other Department users or Citizens. Sanctioned posts do not apply to these categories.",
        "Normal",
    )
    other_tbl = insert_table_after(
        note4,
        ["Role Category", "Example / Scope", "Example Roles"],
        [
            ("Citizen", "Public portal users", "Citizen (default self-registered access)"),
            ("Other Department", "Revenue Department", "Revenue Verification Officer, Bhoomi Cross-check User"),
            ("Other Department", "Treasury / Finance", "Treasury Payment Verifier, Khajane Reconciliation User"),
            ("Other Department", "Police", "Police Enquiry Officer, FIR Verification User"),
            ("Other Department", "Urban Local Body (ULB)", "ULB Document Verifier"),
            ("Other Department", "General (cross-department)", "Read-only MIS Viewer, Document Upload User"),
        ],
    )

    sub5 = insert_heading_after(
        other_tbl, "6.5.5 Primary and Secondary Role Assignment (DSR Officers Only)", "Heading 3", parent=rbac_parent
    )
    primary_note = insert_paragraph_after(
        sub5,
        "Primary and Secondary role assignment applies only to Department users (DSR Officers) selecting "
        "roles from the Role Master where Role Category = DSR. Other Department users and Citizens are "
        "assigned roles from the same Role Master filtered by their Role Category (see Sections 6.5.1, "
        "6.5.4 and 6.6.2).",
        "Normal",
    )
    primary_tbl = insert_table_after(
        primary_note,
        ["Role Type", "Mandatory?", "End Date", "Applies To", "Notes"],
        [
            (
                "Primary Role",
                "Yes — exactly one per DSR Officer",
                "Not permitted",
                "DSR Officers only",
                "Substantive assignment; mapped to a vacant sanctioned post (role + office)",
            ),
            (
                "Secondary Role",
                "No — zero or more",
                "Mandatory",
                "DSR Officers only",
                "Temporary/additional access; auto-revoked on end date",
            ),
            (
                "Other Department role",
                "Yes — exactly one per Other Department user",
                "N/A (account End Date is optional — see §6.6.2)",
                "Other Department users only",
                "Single role from Role Master (Role Category = Other Department); no Primary/Secondary split",
            ),
        ],
    )

    admin_heading = next(
        p for p in doc.paragraphs if p.text.strip() == "6.7 Administrative User Management"
    )
    wf_heading_el = OxmlElement("w:p")
    admin_heading._element.addprevious(wf_heading_el)
    wf_heading = Paragraph(wf_heading_el, admin_heading._parent)
    wf_heading.style = "Heading 2"
    wf_heading.add_run("6.6 User Creation and Role Assignment Workflow")

    wf_intro = insert_paragraph_after(
        wf_heading,
        "The system shall provide dedicated step-by-step workflows to assign roles during user creation. "
        "For DSR Officers, account creation cannot be completed without selecting a Primary Role. "
        "For Other Department users, account creation cannot be completed without selecting exactly one role.",
        "Normal",
    )

    dsr_sub = insert_heading_after(wf_intro, "6.6.1 DSR Officer User Creation with Role Assignment", "Heading 3")
    dsr_tbl = insert_table_after(
        dsr_sub,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → Add DSR Department User", "Admin", "Authorised admin role only (FR-02)"),
            ("2", "Enter user particulars (name, Username, mobile, email, KGID, photo, ID proof)", "Admin", ""),
            ("3", "Select Primary Role — vacant sanctioned post (role + office); Role Category = DSR", "Admin", "Mandatory; FR-30, FR-26, FR-34"),
            ("4", "Optionally add Secondary Role(s) from Role Master (Role Category = DSR)", "Admin", "Each requires end date; FR-31"),
            ("5", "Upload approval letter for Secondary Role(s) where applicable", "Admin", "Should"),
            ("6", "Capture biometrics", "Admin / Officer", "Mandatory for DSR users (FR-06)"),
            ("7", "Review role summary (Primary + Secondary list) and confirm", "Admin", ""),
            ("8", "Save — account active; sanctioned post occupancy updated", "System", "Blocked if Primary Role not selected"),
        ],
    )

    other_sub = insert_heading_after(
        dsr_tbl, "6.6.2 Other Department User Creation with Role Assignment", "Heading 3", parent=rbac_parent
    )
    other_note = insert_paragraph_after(
        other_sub,
        "Other Department users are stored in the same User Master (User Category = Other Department). "
        "Exactly one role is selected from the Role Master filtered by Role Category = Other Department. "
        "There is no Primary or Secondary role. An optional End Date may be entered; if entered, the "
        "system shall deactivate the user on that date.",
        "Normal",
    )
    insert_table_after(
        other_note,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → Add Other Department User", "Admin", "Authorised admin role only (FR-03)"),
            ("2", "Enter user particulars (name, Username, mobile, email, photo, ID proof)", "Admin", "User Category = Other Department in User Master"),
            ("3", "Enter parent department and designation", "Admin", "e.g. Revenue, Treasury, Police"),
            ("4", "Assign exactly one role from Role Master (Role Category = Other Department)", "Admin", "Mandatory; FR-29, FR-34 — no Primary/Secondary"),
            ("5", "Optionally enter Account End Date", "Admin", "Not mandatory; FR-33"),
            ("6", "Upload authorisation letter / NOC from parent department", "Admin", "Should"),
            ("7", "Capture biometrics", "Admin / User", "Mandatory (FR-07)"),
            ("8", "Review role and End Date (if any) and confirm", "Admin", ""),
            ("9", "Save — account active with module access for the assigned role", "System", "Blocked if no role selected"),
            ("10", "If End Date was entered and is reached — deactivate user; block login", "System", "FR-33; audit logged"),
        ],
    )

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v2.1"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
