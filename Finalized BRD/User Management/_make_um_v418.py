# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v4.18.docx from v4.17.

Refreshes the thirteen P-series process-diagram figures from the redrawn
.drawio sources, and relocates two risk rows and five glossary rows that a
previous build appended to the wrong tables.
"""
from __future__ import annotations

import copy
import shutil
import struct
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(__file__).resolve().parent
SRC = BASE / "BRD_User_Management_v4.17.docx"
DST = BASE / "BRD_User_Management_v4.18.docx"
DIAGRAMS = BASE / "ProcessDiagrams" / "User_Management"

NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
NS_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
NS_WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"

FIGURE_WIDTH_IN = 7.0  # full usable text width (8.5in page, 0.75in margins)

# Embedded media part -> redrawn diagram whose export replaces it.
# The six S-series figures are deliberately absent: their .drawio sources are
# unchanged, and the document's generated S-figures carry titles and footnotes
# that the plain drawio exports do not.
FIGURE_SOURCES = {
    "media/image1.png": "P-01_Citizen_Self_Registration",
    "media/image3.png": "P-02_Login_All_Categories",
    "media/image4.png": "P-03_Citizen_Lost_Mobile_Reset",
    "media/image5.png": "P-04_Departmental_Mobile_Change",
    "media/image6.png": "P-05_DSR_Login_Post_Selection",
    "media/image7.png": "P-06_Additional_Charge_After_Login",
    "media/image13.png": "P-07_DSR_Officer_User_Creation",
    "media/image14.png": "P-08_Other_Department_User_Creation",
    "media/image15.png": "P-09_Transfer_Out_Relieving",
    "media/image16.png": "P-12_Handover_Timeline",
    "media/image17.png": "P-10_Transfer_In",
    "media/image18.png": "P-11_Occupancy_Refresh_Job",
    "media/image19.png": "P-13_Temporary_Absence_Charge",
}

# Rows added to §6.6.6 by the v4.14 build that belong in §10 Risks.
MISPLACED_RISKS = [
    (
        "Leave / OOD treated as vacancy or FR-UM-053",
        "High",
        "FR-UM-081 keeps Occupied; FR-UM-082/FR-UM-083 temporary charge is "
        "superior-assigned and distinct from FR-UM-053; Transfer In must not "
        "open on absence alone",
    ),
    (
        "Absent officer still logs in on another post",
        "High",
        "FR-UM-080 blocks entire Username login while any effective absence exists",
    ),
]

# Rows added to §6.7 by the v4.14 build that belong in §11 Glossary.
MISPLACED_GLOSSARY = [
    (
        "Temporary Absence",
        "Superior-recorded Leave / OOD / Other period linked to an active post "
        "occupancy with reason and from–to dates (FR-UM-079); does not end the "
        "occupancy or free Transfer In capacity",
    ),
    (
        "Effective absence",
        "An Approved temporary absence whose from_date–to_date covers the "
        "current IST calendar day",
    ),
    (
        "OOD",
        "Out of Duty / Office Duty — an absence type under Temporary Absence "
        "(FR-UM-079)",
    ),
    (
        "Temporary Charge",
        "Superior-assigned, dated mapping of an absent Post + Office to another "
        "subordinate officer under the same superior (may be cross-office); "
        "cover officer selects it at FR-UM-052 (FR-UM-082, FR-UM-083); distinct "
        "from FR-UM-053 additional charge",
    ),
    (
        "Cover officer",
        "DSR Officer who holds Temporary Charge of another officer's Post + "
        "Office during Temporary Absence",
    ),
]

REVISION_NOTE = (
    "Refreshed the thirteen P-series process-diagram figures (P-01–P-13) from "
    "the redrawn .drawio sources and re-exported them at full text width; "
    "superseded sources archived under ProcessDiagrams/User_Management/"
    "_superseded_v4.17. Moved two Temporary Absence risk rows from the §6.6.6 "
    "workflow table into §10 Risks and Mitigations, and five Temporary Absence "
    "terms from the §6.7 requirements table into §11 Glossary. Process steps "
    "and functional requirements are otherwise unchanged."
)


def png_dimensions(data: bytes) -> tuple[int, int]:
    width, height = struct.unpack(">II", data[16:24])
    return width, height


def set_cell_text(cell: _Cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def add_table_row(table: Table, values: list[str]) -> None:
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    last_tr.addnext(new_tr)
    row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell_text(row.cells[i], val)


def remove_table_row(table: Table, index: int) -> None:
    table._tbl.remove(table.rows[index]._tr)


def find_table(doc: Document, predicate) -> Table:
    for tbl in doc.tables:
        if tbl.rows and predicate(tbl):
            return tbl
    raise RuntimeError("table not found")


def headers_of(tbl: Table) -> list[str]:
    return [c.text.strip() for c in tbl.rows[0].cells]


def set_field(table: Table, field: str, value: str) -> None:
    for row in table.rows:
        if row.cells[0].text.strip() == field:
            set_cell_text(row.cells[1], value)
            return
    raise KeyError(field)


def replace_figures(doc: Document) -> list[str]:
    """Swap image part bytes and rescale each drawing to the new aspect ratio."""
    replaced: list[str] = []
    for para in doc.paragraphs:
        for blip in para._element.iter(f"{NS_A}blip"):
            rid = blip.get(f"{NS_R}embed")
            if rid is None or rid not in doc.part.rels:
                continue
            target = doc.part.rels[rid].target_ref
            stem = FIGURE_SOURCES.get(target)
            if stem is None:
                continue

            png = DIAGRAMS / f"{stem}.drawio.png"
            if not png.exists():
                raise FileNotFoundError(png)
            data = png.read_bytes()
            width_px, height_px = png_dimensions(data)

            doc.part.related_parts[rid]._blob = data

            cx = int(Inches(FIGURE_WIDTH_IN))
            cy = int(round(cx * height_px / width_px))
            for tag in (f"{NS_WP}extent", f"{NS_A}ext"):
                for ext in para._element.iter(tag):
                    ext.set("cx", str(cx))
                    ext.set("cy", str(cy))
            replaced.append(f"{target} <- {png.name} ({width_px}x{height_px})")
    return replaced


def fix_misplaced_rows(doc: Document) -> None:
    workflow = find_table(
        doc,
        lambda t: headers_of(t) == ["Step", "Action", "Actor", "Notes"]
        and len(t.rows) > 1
        and t.rows[1].cells[1].text.strip().startswith("Open Temporary Absence"),
    )
    for index in range(len(workflow.rows) - 1, 0, -1):
        if workflow.rows[index].cells[1].text.strip() == "High":
            remove_table_row(workflow, index)

    admin = find_table(
        doc,
        lambda t: headers_of(t) == ["Req ID", "Requirement", "Priority"]
        and len(t.rows) > 1
        and t.rows[1].cells[0].text.strip() == "FR-UM-020",
    )
    glossary_terms = {term for term, _ in MISPLACED_GLOSSARY}
    for index in range(len(admin.rows) - 1, 0, -1):
        if admin.rows[index].cells[0].text.strip() in glossary_terms:
            remove_table_row(admin, index)

    risks = find_table(doc, lambda t: headers_of(t) == ["Risk", "Impact", "Mitigation"])
    existing_risks = {r.cells[0].text.strip() for r in risks.rows}
    for values in MISPLACED_RISKS:
        if values[0] not in existing_risks:
            add_table_row(risks, list(values))

    glossary = find_table(doc, lambda t: headers_of(t) == ["Term", "Definition"])
    existing_terms = {r.cells[0].text.strip() for r in glossary.rows}
    for values in MISPLACED_GLOSSARY:
        if values[0] not in existing_terms:
            add_table_row(glossary, list(values))


def build() -> Document:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    replaced = replace_figures(doc)
    missing = set(FIGURE_SOURCES) - {line.split(" <- ")[0] for line in replaced}
    if missing:
        raise RuntimeError(f"figures not replaced: {sorted(missing)}")
    for line in replaced:
        print(f"  figure {line}")

    fix_misplaced_rows(doc)

    fields = find_table(doc, lambda t: headers_of(t) == ["Field", "Value"])
    set_field(fields, "Version", "4.18")
    set_field(fields, "Last updated", "2026-09-04")

    revisions = find_table(
        doc, lambda t: headers_of(t) == ["Version", "Date", "Author", "Description"]
    )
    add_table_row(revisions, ["4.18", "04-Sep-2026", "Nandha Kumar", REVISION_NOTE])

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v4.18"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    claude_dir = BASE.parent.parent / "Claude"
    if claude_dir.is_dir():
        try:
            shutil.copy2(target, claude_dir / target.name)
            print(f"Mirrored: {claude_dir / target.name}")
        except Exception as exc:
            print(f"Claude mirror skipped: {exc}")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
