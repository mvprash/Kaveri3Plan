# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v4.14.docx from v4.13 — Temporary Absence / Leave / OOD
and Temporary Charge (§6.6.6, FR-79–FR-84, P-13)."""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(__file__).resolve().parent
SRC = BASE / "BRD_User_Management_v4.13.docx"
DST = BASE / "BRD_User_Management_v4.14.docx"
PNG_P13 = (
    BASE
    / "ProcessDiagrams"
    / "User_Management"
    / "P-13_Temporary_Absence_Charge.drawio.png"
)
DRAWIO_P13 = "ProcessDiagrams/User_Management/P-13_Temporary_Absence_Charge.drawio"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def insert_paragraph_after(ref_paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    para = Paragraph(new_p, ref_paragraph._parent)
    try:
        para.style = style
    except KeyError:
        pass
    if text:
        para.add_run(text)
    return para


def insert_table_after(ref_paragraph, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_paragraph._element.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    tbl_pr = OxmlElement("w:tblPr")
    tbl_el.append(tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return Paragraph(spacer, ref_paragraph._parent)


def insert_picture_after(ref_paragraph, image_path: Path, width_in: float = 6.5) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    para = Paragraph(new_p, ref_paragraph._parent)
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    return para


def find_para(doc: Document, prefix: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise KeyError(f"paragraph not found: {prefix!r}")


def find_para_exact(doc: Document, text: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise KeyError(f"paragraph not found exact: {text!r}")


def find_para_contains(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise KeyError(f"paragraph containing {needle!r} not found")


def add_table_row(table, values: list[str], shade_header_style: bool = False) -> None:
    """Append a row by cloning the last row (works even when tblGrid is absent)."""
    import copy

    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    last_tr.addnext(new_tr)
    row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell_text(row.cells[i], val, bold=shade_header_style)


FR_ROWS = [
    (
        "FR-79",
        "The system shall allow a hierarchy superior to record a Temporary Absence against an "
        "active DSR Officer post occupancy (Post + Office). Absence types shall include Leave, "
        "OOD (Out of Duty / Office Duty), and Other, with a mandatory reason code, inclusive "
        "from_date and to_date (IST calendar days), and optional order reference / uploaded "
        "order document. Absence shall be recorded by the superior only — the absent officer "
        "shall have no self-service create or approve path. The absence is a linked record on "
        "the occupancy; it shall not end the occupancy and shall not substitute for Transfer "
        "Out / relieving (FR-57) or deputation End Date (FR-30). Scope to create/amend/cancel "
        "shall follow office span (FR-59) and immediate-parent post parentage (FR-43) — the "
        "same filters as Transfer Out. Citizens and Other Department users are out of scope. "
        "All create, amend, and cancel actions shall be audit-logged.",
        "High",
    ),
    (
        "FR-80",
        "While an Approved temporary absence covers the current IST calendar day for any of "
        "an officer's active post occupancies (effective absence), the system shall deny login "
        "for that Username (KGID). The officer shall not select another assigned post under "
        "FR-52 for the same period. On cancel of the absence, or after 23:59 IST of to_date "
        "(FR-84 / FR-68), login shall be allowed again. The login denial message shall state "
        "that the officer is on Leave / OOD / Other until the to_date.",
        "High",
    ),
    (
        "FR-81",
        "Temporary absence shall not free sanctioned capacity. The absent occupancy shall "
        "continue to count toward Occupied for available-capacity tests (FR-66(a)) and for "
        "Transfer In / reserved occupancy (FR-67). Transfer In shall not treat the post as "
        "vacant solely because of absence. Wholly-unoccupied tests for FR-53 additional charge "
        "(FR-66(b)) remain Occupied = 0 and are unaffected — an occupied-but-absent post shall "
        "not be offered under FR-53.",
        "High",
    ),
    (
        "FR-82",
        "While an occupancy is under temporary absence (or as part of the same superior "
        "workflow), the superior may assign Temporary Charge of that absent Post + Office to "
        "another DSR Officer who holds an active occupancy under the same superior (office "
        "span FR-59 and posts under the superior in Officer Hierarchy FR-43). Cross-office "
        "assignment within the superior's span is allowed — for example a District Registrar "
        "may assign temporary charge of Sub-Registrar at SRO A to the Sub-Registrar of SRO B "
        "under the same district; an AIGR (Admin) may assign temporary charge of one District "
        "Registrar to another District Registrar under them. The covering officer shall not be "
        "the absentee and shall not themselves be under effective absence (FR-80). At most one "
        "active temporary charge shall exist per absent Post + Office at a time; the superior "
        "may cancel and reassign. Temporary charge is dated and linked to the absence; it is "
        "not a Transfer In (FR-60), not relieving (FR-57), and not FR-53 session additional "
        "charge. Assignment and cancellation shall be audit-logged (optional order reference).",
        "High",
    ),
    (
        "FR-83",
        "An officer holding an active temporary charge shall authenticate like any DSR Officer. "
        "At FR-52 post selection the system shall list their active assigned posts and each "
        "active temporary-charge Post + Office, labelled for example "
        "\"Temporary charge — Role — Post Name — Office Name (Office Code) — covering "
        "<Absent Officer Name>\". Selecting a temporary-charge row sets session Module Function "
        "claims from that covered post via Post–Role mapping only (one active context; same "
        "spirit as FR-38). The officer may switch back to an assigned-post context without "
        "logout where the product provides an in-session control (same pattern as FR-53 "
        "switch-back). FR-54 / header display shall show which context is active and the "
        "office of the covered post. Temporary charge is distinct from FR-53: it is "
        "superior-assigned, may be cross-office, and applies while the primary occupancy "
        "remains occupied.",
        "High",
    ),
    (
        "FR-84",
        "Temporary absence is effective through 23:59 IST of to_date. The occupancy refresh "
        "job (FR-68) — or equivalent day-boundary evaluation — shall clear effective absence "
        "and end linked temporary charge mappings after that boundary, and shall remove "
        "temporary-charge rows from FR-52 for subsequent logins. Superior cancel of absence "
        "or of temporary charge shall take effect immediately for login and post selection. "
        "The system shall provide an absence and temporary-charge report (officer, Post + "
        "Office, type, reason, dates, cover officer, actor, timestamps) and shall retain "
        "audit history for at least seven years consistent with other UM audit events.",
        "High",
    ),
]


def build() -> Document:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if not PNG_P13.exists():
        raise FileNotFoundError(PNG_P13)

    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Cover
    replace_paragraph_text(doc.paragraphs[3], "Version 4.14")
    replace_paragraph_text(doc.paragraphs[4], "Date: 30 August 2026")

    # Executive summary — append absence sentence if not present
    intro = doc.paragraphs[12].text
    if "Temporary Absence" not in intro and "temporary absence" not in intro.lower():
        replace_paragraph_text(
            doc.paragraphs[12],
            intro.rstrip(".")
            + ". It also defines Temporary Absence (Leave / OOD / Other) recorded by a "
            "hierarchy superior with temporary charge of the absent Post + Office to another "
            "subordinate under that superior (including cross-office peers such as SRO A → "
            "SRO B under a District Registrar), without freeing Transfer In capacity "
            "(§6.6.6, FR-79–FR-84).",
        )

    # Problem / context paragraph
    ctx = doc.paragraphs[14].text
    if "Temporary Absence" not in ctx and "FR-79" not in ctx:
        replace_paragraph_text(
            doc.paragraphs[14],
            ctx.rstrip(".")
            + " A hierarchy superior may record Temporary Absence (Leave / OOD / Other) on an "
            "active occupancy without freeing the slot; while absent the officer cannot log "
            "in; the superior may assign Temporary Charge of that Post + Office to another "
            "officer under them, including at a different office in the same span (FR-79–FR-84).",
        )

    # In scope — insert after Transfer Out line (para 22)
    scope_transfer = doc.paragraphs[22]
    if "Temporary Absence" not in scope_transfer.text:
        insert_paragraph_after(
            scope_transfer,
            "Temporary Absence (Leave / OOD / Other) recorded by hierarchy superior only; "
            "login blocked while any effective absence exists; Temporary Charge of the absent "
            "Post + Office to another subordinate under the superior (including cross-office "
            "peers, e.g. DRO assigns SRO A charge to SR of SRO B); capacity not freed for "
            "Transfer In; distinct from FR-53 additional charge (FR-79–FR-84, §6.6.6, P-13)",
        )

    # Out of scope — add HRMS note after SSO
    oos = find_para_exact(doc, "Out of Scope:")
    # next para is SSO
    sso = None
    found = False
    for p in doc.paragraphs:
        if found:
            sso = p
            break
        if p.text.strip() == "Out of Scope:":
            found = True
    if sso is not None and "HRMS" not in sso.text and "leave balance" not in sso.text.lower():
        insert_paragraph_after(
            sso,
            "Full HRMS leave-balance ledger, payroll, and attendance punching (UM records "
            "absence and temporary charge for access control only)",
        )

    # Revision history
    rev = doc.tables[1]
    add_table_row(
        rev,
        [
            "4.14",
            "30-Aug-2026",
            "Nandha Kumar",
            "Added §6.6.6 Temporary Absence / Leave / OOD and Temporary Charge (FR-79–FR-84): "
            "superior-only absence on occupancy with reason and date range; absence does not "
            "free Transfer In capacity; full login block while any effective absence exists; "
            "superior assigns temporary charge to another subordinate under them (cross-office "
            "peers allowed — e.g. DRO gives SRO A charge to SR of SRO B; AIGR gives one DRO "
            "charge to another DRO); cover officer selects temporary charge at FR-52; distinct "
            "from FR-53; P-13 process diagram; user stories; glossary; reports; UAT",
        ],
    )

    # Diagram index
    idx = doc.tables[3]
    add_table_row(
        idx,
        [
            "P-13",
            "Temporary Absence / Temporary Charge",
            "§6.6.6",
            DRAWIO_P13,
        ],
    )

    # Main FR table (table 9) — append FR-79–FR-84
    fr_tbl = doc.tables[9]
    for fr_id, req, pri in FR_ROWS:
        add_table_row(fr_tbl, [fr_id, req, pri])

    # Insert §6.6.6 after the occupancy-refresh job step table (end of §6.6.5).
    job_tbl_el = None
    seen_p11 = False
    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            texts = child.findall(".//" + qn("w:t"))
            text_p = "".join(t.text or "" for t in texts)
            if "P-11 Occupancy Refresh Job" in text_p:
                seen_p11 = True
        elif tag == "tbl" and seen_p11 and job_tbl_el is None:
            job_tbl_el = child
            break
    if job_tbl_el is None:
        raise RuntimeError("Could not find §6.6.5 occupancy refresh job table")

    def append_para_after(ref_el, text: str, style: str = "Normal"):
        new_p = OxmlElement("w:p")
        ref_el.addnext(new_p)
        para = Paragraph(new_p, doc.paragraphs[0]._parent)
        try:
            para.style = style
        except KeyError:
            pass
        if text:
            para.add_run(text)
        return new_p

    def append_table_after(ref_el, headers, rows):
        spacer = OxmlElement("w:p")
        ref_el.addnext(spacer)
        tbl_el = OxmlElement("w:tbl")
        spacer.addnext(tbl_el)
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.append(tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tbl_pr.append(borders)

        def add_row(values, header=False):
            tr = OxmlElement("w:tr")
            tbl_el.append(tr)
            for val in values:
                tc = OxmlElement("w:tc")
                tr.append(tc)
                p = OxmlElement("w:p")
                tc.append(p)
                r = OxmlElement("w:r")
                p.append(r)
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = val
                r.append(t)
                if header:
                    rpr = OxmlElement("w:rPr")
                    b = OxmlElement("w:b")
                    rpr.append(b)
                    r.insert(0, rpr)
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), "D9E2F3")
                    tc_pr = OxmlElement("w:tcPr")
                    tc_pr.append(shd)
                    tc.insert(0, tc_pr)

        add_row(list(headers), header=True)
        for row in rows:
            add_row(list(row))
        return tbl_el

    def append_png_after(ref_el, image_path: Path):
        new_p = OxmlElement("w:p")
        ref_el.addnext(new_p)
        para = Paragraph(new_p, doc.paragraphs[0]._parent)
        run = para.add_run()
        run.add_picture(str(image_path), width=Inches(6.5))
        return new_p

    content_forward: list[tuple[str, object]] = []

    content_forward.append(
        (
            "h3",
            "6.6.6 Temporary Absence / Leave / OOD and Temporary Charge (DSR Officers)",
        )
    )
    content_forward.append(
        (
            "p",
            "This section covers temporary unavailability of a DSR Officer who remains the "
            "occupant of their post (Leave, OOD, or Other). It is distinct from Transfer Out / "
            "relieving (FR-57), Transfer In (FR-60), deputation End Date (FR-30), and post-login "
            "additional charge of a wholly unoccupied subordinate post (FR-53). Full HRMS leave "
            "balances and payroll are out of scope — User Management records absence and "
            "temporary charge for access control and operational continuity only.",
        )
    )
    content_forward.append(("p", "Business rules (summary):"))
    content_forward.append(
        (
            "p",
            "1. Absence is entered by the hierarchy superior only (office span + immediate-parent "
            "parentage). The officer cannot self-record Leave / OOD.",
        )
    )
    content_forward.append(
        (
            "p",
            "2. Absence does not free a slot — Occupied is unchanged; Transfer In must not treat "
            "the post as vacant (FR-81).",
        )
    )
    content_forward.append(
        (
            "p",
            "3. While any effective absence exists on any of the officer's active occupancies, "
            "login is blocked entirely (FR-80).",
        )
    )
    content_forward.append(
        (
            "p",
            "4. The superior may assign Temporary Charge of the absent Post + Office to another "
            "subordinate under them, including cross-office peers (FR-82).",
        )
    )
    content_forward.append(
        (
            "p",
            "5. The covering officer logs in normally and selects the temporary-charge post at "
            "FR-52 (FR-83).",
        )
    )

    content_forward.append(("p", "User stories / examples:"))
    content_forward.append(
        (
            "table",
            (
                ["ID", "Actor", "Story", "Acceptance (system)"],
                [
                    (
                        "US-TA-01",
                        "District Registrar",
                        "As District Registrar of DRO Bengaluru, I need to record Leave for the "
                        "Sub-Registrar of SRO Yeshwanthapura (SRO A) from 01-Sep-2026 to "
                        "05-Sep-2026 so that the officer cannot access KAVERI while away, without "
                        "opening the post for Transfer In.",
                        "Absence saved (FR-79); Occupied unchanged (FR-81); SR of SRO A denied "
                        "login for those dates (FR-80); Transfer In to that SR post still blocked "
                        "unless relieving/capacity rules allow otherwise.",
                    ),
                    (
                        "US-TA-02",
                        "District Registrar",
                        "As the same District Registrar, I need to give temporary charge of "
                        "SRO Yeshwanthapura (SRO A) Sub-Registrar work to the Sub-Registrar of "
                        "SRO Jayanagar (SRO B) under my district for the leave period.",
                        "Temporary charge mapping created (FR-82); SR of SRO B sees a Temporary "
                        "charge row for SR @ SRO A at FR-52 (FR-83); claims are those of SR @ "
                        "SRO A while that context is selected; SRO B own post remains available "
                        "as a separate selectable context.",
                    ),
                    (
                        "US-TA-03",
                        "AIGR (Admin)",
                        "As AIGR (Admin), I need to record OOD / Leave for District Registrar of "
                        "DRO Mysuru and assign temporary charge of that DRO post to the District "
                        "Registrar of DRO Bengaluru (another district under me).",
                        "Same FR-79–FR-83 behaviour at DRO post level; cover officer is DRO "
                        "Bengaluru; absent DRO Mysuru cannot log in; capacity at DRO Mysuru not "
                        "freed for Transfer In.",
                    ),
                    (
                        "US-TA-04",
                        "Covering Sub-Registrar",
                        "As Sub-Registrar of SRO B holding temporary charge of SRO A, I need to "
                        "log in and choose whether to work as SR of B or under temporary charge "
                        "of A.",
                        "FR-52 lists both; one context per session; header shows active context "
                        "and office (FR-83 / FR-54 pattern).",
                    ),
                ],
            ),
        )
    )

    content_forward.append(("p", "Examples (who may assign temporary charge):"))
    content_forward.append(
        (
            "table",
            (
                [
                    "Superior (session)",
                    "Absent occupancy",
                    "May assign temporary charge to",
                    "Example",
                ],
                [
                    (
                        "District Registrar @ DRO Bengaluru",
                        "Sub-Registrar @ SRO A (under DRO)",
                        "Any other Sub-Registrar (or eligible subordinate) with an active "
                        "occupancy under the same DRO — including at SRO B",
                        "DRO gives charge of SR @ SRO Yeshwanthapura to SR @ SRO Jayanagar",
                    ),
                    (
                        "AIGR (Admin) @ Head Office / Admin span",
                        "District Registrar @ DRO Mysuru",
                        "Any other District Registrar under the same AIGR",
                        "AIGR gives charge of DRO Mysuru to DRO Bengaluru",
                    ),
                    (
                        "Sub-Registrar @ SRO Yeshwanthapura",
                        "FDA @ same SRO",
                        "Another subordinate under that SR at that SRO (office span)",
                        "SR records FDA leave and may assign temporary charge to another "
                        "eligible subordinate under the SR",
                    ),
                ],
            ),
        )
    )

    content_forward.append(
        (
            "p",
            "Approved process diagram — P-13 Temporary Absence / Temporary Charge — FR-79–FR-84 — "
            + DRAWIO_P13,
        )
    )
    content_forward.append(("png", PNG_P13))
    content_forward.append(
        (
            "p",
            "Figure — P-13 Temporary Absence / Leave / OOD and Temporary Charge — §6.6.6 — "
            "superior records absence; optional temporary charge to peer under superior "
            "(e.g. SR A → SR B); cover officer selects charge at login — " + DRAWIO_P13,
        )
    )

    content_forward.append(("p", "Temporary Absence / Temporary Charge workflow:"))
    content_forward.append(
        (
            "table",
            (
                ["Step", "Action", "Actor", "Notes"],
                [
                    (
                        "1",
                        "Open Temporary Absence / Leave / OOD",
                        "Superior",
                        "FR-79; not available to the officer for self-service",
                    ),
                    (
                        "2",
                        "System lists offices in actor's office span; within them, occupancies "
                        "where actor Post is immediate parent",
                        "System",
                        "FR-59, FR-43 — same filters as FR-57",
                    ),
                    (
                        "3",
                        "Select occupancy; enter type (Leave / OOD / Other), reason code, "
                        "from_date, to_date; optional order",
                        "Superior",
                        "Mandatory fields per FR-79",
                    ),
                    (
                        "4",
                        "Confirm absence — linked record on occupancy; Occupied unchanged",
                        "System",
                        "FR-79, FR-81; audit-logged",
                    ),
                    (
                        "5",
                        "From from_date — deny login for absent Username (any effective absence)",
                        "System",
                        "FR-80; applies even if other posts are assigned",
                    ),
                    (
                        "6",
                        "Optional: assign Temporary Charge of absent Post + Office to another "
                        "subordinate under the superior (may be different office)",
                        "Superior",
                        "FR-82; one active charge per Post + Office",
                    ),
                    (
                        "7",
                        "Cover officer authenticates; FR-52 lists own posts + Temporary charge row",
                        "Cover officer / System",
                        "FR-83",
                    ),
                    (
                        "8",
                        "Cover officer selects Temporary charge context — claims from covered "
                        "post only",
                        "Cover officer / System",
                        "FR-83; distinct from FR-53",
                    ),
                    (
                        "9",
                        "After to_date 23:59 IST (or superior cancel) — clear effective absence "
                        "and temporary charge; absentee may log in again",
                        "System",
                        "FR-84; FR-68 day-boundary / refresh",
                    ),
                ],
            ),
        )
    )

    content_forward.append(
        (
            "p",
            "FR-53 remains unchanged: session additional charge applies only to wholly "
            "unoccupied subordinate posts at the same office. Temporary charge under this "
            "section is the path used when the post is still occupied by an officer on Leave / "
            "OOD / Other.",
        )
    )

    # Insert forward after job table so heading appears first.
    cursor = job_tbl_el
    for kind, payload in content_forward:
        if kind == "h3":
            cursor = append_para_after(cursor, payload, "Heading 3")
        elif kind == "p":
            cursor = append_para_after(cursor, payload, "Normal")
        elif kind == "png":
            cursor = append_png_after(cursor, payload)
        elif kind == "table":
            headers, rows = payload
            cursor = append_table_after(cursor, headers, rows)

    # Extend FR-68 text lightly via paragraph search in FR table cell — already have FR-84 cross-ref

    # Reporting requirements
    try:
        ac_rep = find_para_contains(doc, "additional charge report (FR-53)")
        insert_paragraph_after(
            ac_rep,
            "The system shall provide a Temporary Absence and Temporary Charge report (FR-79–FR-84) "
            "over a selected date range showing each absence (officer, Post + Office, type, reason, "
            "from/to dates, actor) and each temporary charge assignment or clearance (cover officer, "
            "covered Post + Office, timestamps).",
        )
    except KeyError:
        pass

    # Acceptance criteria UAT bullets
    try:
        uat_rmf = find_para_contains(doc, "UAT — Role–Module–Function coverage")
        new_p = OxmlElement("w:p")
        uat_rmf._element.addprevious(new_p)
        para = Paragraph(new_p, uat_rmf._parent)
        para.add_run(
            "UAT — FR-79–FR-84: District Registrar records Leave for Sub-Registrar of SRO A; "
            "that SR cannot log in; Occupied unchanged; DRO assigns temporary charge to "
            "Sub-Registrar of SRO B; SRO B SR sees Temporary charge for SRO A at FR-52 and "
            "receives SR@A claims when selected; after to_date the charge disappears and SRO A "
            "SR may log in again. FR-53 still requires Occupied = 0 and is not used for this path."
        )
    except KeyError:
        pass

    # Risks table
    risks = doc.tables[50]
    add_table_row(
        risks,
        [
            "Leave / OOD treated as vacancy or FR-53",
            "High",
            "FR-81 keeps Occupied; FR-82/83 temporary charge is superior-assigned and distinct "
            "from FR-53; Transfer In must not open on absence alone",
        ],
    )
    add_table_row(
        risks,
        [
            "Absent officer still logs in on another post",
            "High",
            "FR-80 blocks entire Username login while any effective absence exists",
        ],
    )

    # Glossary
    gloss = doc.tables[51]
    for term, definition in [
        (
            "Temporary Absence",
            "Superior-recorded Leave / OOD / Other period linked to an active post occupancy "
            "with reason and from–to dates (FR-79); does not end the occupancy or free Transfer In capacity",
        ),
        (
            "Effective absence",
            "An Approved temporary absence whose from_date–to_date covers the current IST calendar day",
        ),
        (
            "OOD",
            "Out of Duty / Office Duty — an absence type under Temporary Absence (FR-79)",
        ),
        (
            "Temporary Charge",
            "Superior-assigned, dated mapping of an absent Post + Office to another subordinate "
            "officer under the same superior (may be cross-office); cover officer selects it at "
            "FR-52 (FR-82, FR-83); distinct from FR-53 additional charge",
        ),
        (
            "Cover officer",
            "DSR Officer who holds Temporary Charge of another officer's Post + Office during Temporary Absence",
        ),
    ]:
        add_table_row(gloss, [term, definition])

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v4.14"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    # Mirror to Claude folder if present
    claude_dir = BASE.parent.parent / "Claude"
    if claude_dir.is_dir():
        claude_dst = claude_dir / target.name
        try:
            shutil.copy2(target, claude_dst)
            print(f"Mirrored: {claude_dst}")
        except Exception as exc:
            print(f"Claude mirror skipped: {exc}")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
