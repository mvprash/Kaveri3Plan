# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v4.16.docx from v4.15.

Align requirement IDs with Hindu Marriage / Marriage BRD convention:
Req ID FR-UM-### (three digits, module prefix), matching FR-HMA-###.
"""
from __future__ import annotations

import copy
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(__file__).resolve().parent
SRC = BASE / "BRD_User_Management_v4.15.docx"
DST = BASE / "BRD_User_Management_v4.16.docx"

# FR-01 … FR-84 → FR-UM-001 … FR-UM-084 (do not match FR-UM- already)
FR_RE = re.compile(r"\bFR-(?!UM-)(\d{1,2})\b")


def rewrite_fr_ids(text: str) -> str:
    return FR_RE.sub(lambda m: f"FR-UM-{int(m.group(1)):03d}", text)


def set_cell_text(cell: _Cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def rewrite_paragraph(paragraph: Paragraph) -> bool:
    old = paragraph.text
    new = rewrite_fr_ids(old)
    if new == old:
        return False
    replace_paragraph_text(paragraph, new)
    return True


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                yield p
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def add_table_row(table: Table, values: list[str]) -> None:
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    last_tr.addnext(new_tr)
    row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell_text(row.cells[i], val)


def find_revision_table(doc: Document) -> Table:
    for tbl in doc.tables:
        if (
            len(tbl.columns) >= 4
            and tbl.rows
            and tbl.rows[0].cells[0].text.strip() == "Version"
            and tbl.rows[0].cells[3].text.strip() in ("Description", "Summary of change")
        ):
            return tbl
    raise RuntimeError("revision history table not found")


def find_field_table(doc: Document) -> Table:
    for tbl in doc.tables:
        if (
            tbl.rows
            and tbl.rows[0].cells[0].text.strip() == "Field"
            and len(tbl.columns) == 2
        ):
            return tbl
    raise RuntimeError("field/value table not found")


def set_field(table: Table, field: str, value: str) -> None:
    for row in table.rows:
        if row.cells[0].text.strip() == field:
            set_cell_text(row.cells[1], value)
            return
    raise KeyError(field)


def retitle_fr_tables(doc: Document) -> int:
    """First column 'ID' → 'Req ID' on functional-requirement tables."""
    n = 0
    for tbl in doc.tables:
        if not tbl.rows or len(tbl.columns) < 3:
            continue
        hdr = [c.text.strip() for c in tbl.rows[0].cells]
        if hdr[0] != "ID":
            continue
        if len(tbl.rows) < 2:
            continue
        first_data = tbl.rows[1].cells[0].text.strip()
        if not first_data.startswith("FR-"):
            continue
        set_cell_text(tbl.rows[0].cells[0], "Req ID", bold=True)
        n += 1
    return n


def build() -> Document:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    changed = 0
    for p in iter_all_paragraphs(doc):
        if rewrite_paragraph(p):
            changed += 1

    fr_tables = retitle_fr_tables(doc)

    fields = find_field_table(doc)
    set_field(fields, "Version", "4.16")
    set_field(fields, "Last updated", "2026-08-30")

    rev = find_revision_table(doc)
    add_table_row(
        rev,
        [
            "4.16",
            "30-Aug-2026",
            "Nandha Kumar",
            "Requirement IDs aligned to Hindu Marriage BRD convention: FR-nn → "
            "FR-UM-nnn (three-digit, module prefix, e.g. FR-UM-001); FR table "
            "column header ID → Req ID",
        ],
    )

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v4.16"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    print(f"paragraphs/cells rewritten: {changed}; FR tables retitled: {fr_tables}")
    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    claude_dir = BASE.parent.parent / "Claude"
    if claude_dir.is_dir():
        try:
            shutil.copy2(target, claude_dir / target.name)
            print(f"Mirrored: {claude_dir / target.name}")
        except Exception as exc:
            print(f"Claude mirror skipped: {exc}")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
