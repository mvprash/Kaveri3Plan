# -*- coding: utf-8 -*-
"""Create BRD_User_Management_v4.20.docx — remove OTP from FR-UM-006; Transfer In without Joining Date (available capacity only)."""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

SRC = Path(r"Finalized BRD/User Management/BRD_User_Management_v4.19.docx")
DST = Path(r"Finalized BRD/User Management/BRD_User_Management_v4.20.docx")


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell: _Cell, text: str) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(text)
        return
    set_paragraph_text(paragraphs[0], text)
    for para in paragraphs[1:]:
        for run in para.runs:
            run.text = ""


def replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    full = paragraph.text
    if old not in full:
        return False
    set_paragraph_text(paragraph, full.replace(old, new))
    return True


def find_req_row(table: Table, req_id: str):
    for row in table.rows:
        if row.cells[0].text.strip() == req_id:
            return row
    return None


def add_table_row_clone(table: Table, values: list[str]) -> None:
    tbl = table._tbl
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    tbl.append(new_tr)
    new_row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(new_row.cells):
            set_cell_text(new_row.cells[i], val)


def main() -> None:
    if not SRC.exists():
        # Fallback to Claude copy
        alt = Path(r"Claude/BRD_User_Management_v4.19.docx")
        shutil.copy2(alt, SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Document control
    t0 = doc.tables[0]
    set_cell_text(t0.rows[2].cells[1], "4.20")
    set_cell_text(t0.rows[11].cells[1], "2026-09-05")

    add_table_row_clone(
        doc.tables[1],
        [
            "4.20",
            "05-Sep-2026",
            "Nandha Kumar",
            (
                "FR-UM-006: remove OTP for DSR Officers (authenticate with Username + Captcha + "
                "Face/Biometric only). Transfer In: remove Joining Date; allow Transfer In only when "
                "the target post has available occupancy/capacity (FR-UM-060, FR-UM-066(a)); retire "
                "reserved future-dated Transfer In / Joining Date gate (FR-UM-061, FR-UM-067)."
            ),
        ],
    )

    # Stakeholders — DSR auth without OTP
    t2 = doc.tables[2]
    set_cell_text(
        t2.rows[3].cells[2],
        "Access departmental modules via OTP (Other Department) or Captcha + face/biometric authentication (DSR Officers — no OTP)",
    )

    # -------------------------------------------------------------------------
    # FR-UM-006 — remove OTP
    # -------------------------------------------------------------------------
    t6 = doc.tables[6]
    set_cell_text(
        find_req_row(t6, "FR-UM-006").cells[1],
        (
            "Department users (DSR Officers) shall authenticate using Username (KGID) + Captcha "
            "+ Face authentication or Biometric authentication on every login. Either face "
            "authentication or biometric authentication (e.g. fingerprint) shall satisfy the "
            "biometric factor. OTP shall not be required or dispatched for DSR Officer login."
        ),
    )

    # Idle timeout re-auth wording for DSR
    row074 = find_req_row(t6, "FR-UM-074")
    if row074:
        set_cell_text(
            row074.cells[1],
            (
                "The system shall terminate a session after ten (10) minutes of no user activity "
                "(idle timeout). Re-authentication is then required (OTP for Citizens and Other "
                "Department users; face/biometric authentication for DSR Officers — no OTP). This "
                "applies especially to shared SRO counter machines where FR-UM-052/FR-UM-053 fix "
                "the session post for the whole session."
            ),
        )

    # -------------------------------------------------------------------------
    # Transfer In FRs (T9)
    # -------------------------------------------------------------------------
    t9 = doc.tables[9]

    set_cell_text(
        find_req_row(t9, "FR-UM-060").cells[1],
        (
            "The system shall support a Transfer In process for DSR Officers. A superior "
            "(within office span per FR-UM-059 and immediate-parent post parentage per FR-UM-043, "
            "same rules as Transfer Out — office span does not expand parentage) may assign an "
            "officer to a sanctioned post (Post + Office) only when that post has available "
            "capacity at that office (Occupied < Sanctioned Strength — FR-UM-066(a)). Transfer In "
            "shall be blocked when the post is at full strength — there is no reserved / "
            "future-dated Transfer In path (FR-UM-067 retired). While processing Transfer In, the "
            "system shall capture Transfer Order / Reporting Order (order number / reference; "
            "upload of order document where applicable). Joining Date shall not be captured. On "
            "confirm, the occupancy becomes active immediately, occupied count is increased, and "
            "the officer may select the post at the next login under FR-UM-052. Transfer In "
            "actions shall be audit-logged. Citizens and Other Department users are out of scope "
            "for this process."
        ),
    )

    set_cell_text(
        find_req_row(t9, "FR-UM-061").cells[1],
        (
            "RETIRED (v4.20) — Previously gated Transfer In login until 12:00 AM IST of Joining "
            "Date and activated reserved occupancies via FR-UM-068. Joining Date is no longer "
            "captured; Transfer In is allowed only when capacity is available and takes effect "
            "immediately (FR-UM-060)."
        ),
    )
    set_cell_text(find_req_row(t9, "FR-UM-061").cells[2], "—")

    set_cell_text(
        find_req_row(t9, "FR-UM-067").cells[1],
        (
            "RETIRED (v4.20) — Previously allowed future-dated / reserved Transfer In against a "
            "full post after relieving was recorded. Transfer In is now permitted only when the "
            "target Post + Office already has available capacity (Occupied < Sanctioned Strength "
            "— FR-UM-066(a), FR-UM-060). If the post is full, Transfer In remains blocked until "
            "relieving has taken effect and capacity is free (FR-UM-058, FR-UM-068)."
        ),
    )
    set_cell_text(find_req_row(t9, "FR-UM-067").cells[2], "—")

    set_cell_text(
        find_req_row(t9, "FR-UM-066").cells[1],
        (
            "The system shall apply two distinct vacancy tests and shall not treat them as "
            "interchangeable. (a) Available capacity — used for post assignment during user "
            "creation, for Transfer In, and for the post pick lists in those workflows — is "
            "satisfied when the occupied count for a Post at an Office is less than the "
            "sanctioned strength for that Post at that Office (Occupied < Sanctioned Strength). "
            "Occupied for this test shall include active occupancies only. An occupancy with a "
            "recorded Relieving Date still counts as occupied until the occupancy refresh job "
            "de-allocates it (FR-UM-058, FR-UM-068). Transfer In requires available capacity at "
            "the time of recording (FR-UM-060); reserved / pending-join Transfer In occupancies "
            "are not used (FR-UM-067 retired). (b) Unoccupied post — used only for the FR-UM-053 "
            "post-login additional charge selection — is satisfied only when the occupied count "
            "for that Post at that Office is zero (Occupied = 0), irrespective of sanctioned "
            "strength. A Post at an Office with sanctioned strength 2 and one occupant therefore "
            "has available capacity for assignment but is not unoccupied, and shall not be "
            "offered under FR-UM-053. Both counts shall be shown on sanctioned post occupancy "
            "screens and reports."
        ),
    )

    set_cell_text(
        find_req_row(t9, "FR-UM-068").cells[1],
        (
            "The system shall run a scheduled occupancy refresh job shortly after midnight IST "
            "each calendar day (after 12:00 AM IST). On each run the job shall: (1) de-allocate "
            "user–post occupancies whose Relieving Date has ended (FR-UM-058); (2) recalculate "
            "and persist occupied count, remaining capacity, and wholly-unoccupied flag on the "
            "Sanctioned Posts Master for every Post + Office affected (FR-UM-048, FR-UM-066); "
            "(3) refresh the officer's effective post assignments so that login post selection "
            "(FR-UM-052) reflects only occupancies that are active as of that day. The job shall "
            "be idempotent, audit-logged (run timestamp, occupancies ended, before/after occupied "
            "counts), and shall raise an operational alert on failure so that relieving dates "
            "are not left unenforced. Reserved Transfer In / Joining Date activation is removed "
            "(v4.20; FR-UM-061 and FR-UM-067 retired). Optional occupancy End Date / temporary "
            "deputation processing remains removed (v4.19; see FR-UM-030)."
        ),
    )

    set_cell_text(
        find_req_row(t9, "FR-UM-030").cells[1],
        (
            "At least one sanctioned post with available capacity must be assigned to each DSR "
            "Officer at account creation (FR-UM-066(a)). Temporary / time-bound deputation via "
            "optional occupancy End Date is removed (v4.19) — post occupancy End Date and "
            "Deputation Reason shall not be captured at user creation or post assignment. Formal "
            "Transfer out / relieving shall follow FR-UM-057 and FR-UM-058 (including Relieving "
            "Reason under FR-UM-087). Transfer In shall follow FR-UM-060 and requires available "
            "capacity (FR-UM-066(a)); Joining Date / reserved Transfer In are retired "
            "(FR-UM-061, FR-UM-067). Occupancy de-allocation and Sanctioned Posts occupied-count "
            "updates shall be applied by the occupancy refresh job (FR-UM-068)."
        ),
    )

    row081 = find_req_row(t9, "FR-UM-081")
    if row081:
        text = row081.cells[1].text
        text = text.replace(
            "and for Transfer In / reserved occupancy (FR-UM-067). Transfer In shall not treat the post as vacant solely because of temporary absence.",
            "and for Transfer In available-capacity tests (FR-UM-066(a)). Transfer In shall not treat the post as vacant solely because of temporary absence.",
        )
        if "reserved occupancy (FR-UM-067)" in text:
            text = text.replace(
                "for Transfer In / reserved occupancy (FR-UM-067)",
                "for Transfer In available-capacity tests (FR-UM-066(a))",
            )
        set_cell_text(row081.cells[1], text)

    # -------------------------------------------------------------------------
    # Narrative paragraphs
    # -------------------------------------------------------------------------
    para_replacements = [
        (
            "DSR Officers additionally verify face authentication or biometrics; Other Department users do not use biometrics (FR-UM-006, FR-UM-007).",
            "DSR Officers authenticate with Captcha + face authentication or biometrics and without OTP (FR-UM-006); Other Department users use Captcha + OTP and do not use biometrics (FR-UM-007).",
        ),
        (
            "Login uses the Username — the preferred Username for Citizens, the KGID for DSR Officers, and Department Code concatenated with Employee ID or KGID for Other Department users. Because the Username is unique across the whole User Master, the account is resolved unambiguously from the Username alone and no category selection is needed. The login OTP is never emailed; it is sent only by SMS to the registered mobile. DSR Officers complete face or biometric authentication after OTP; Other Department users do not.",
            "Login uses the Username — the preferred Username for Citizens, the KGID for DSR Officers, and Department Code concatenated with Employee ID or KGID for Other Department users. Because the Username is unique across the whole User Master, the account is resolved unambiguously from the Username alone and no category selection is needed. Citizens and Other Department users receive a login OTP by SMS to the registered mobile (never by email). DSR Officers do not use OTP; they complete Captcha plus face or biometric authentication (FR-UM-006).",
        ),
        (
            "Applies only to DSR Officers. After Username (KGID) + Captcha + OTP (to mobile) + Face authentication or Biometrics succeed, the system loads the officer's active sanctioned-post occupancies. If more than one is active, the user must choose which post to work under for this session. Each choice shows Post Name and Office details so dual-charge / multi-office officers can pick the correct context.",
            "Applies only to DSR Officers. After Username (KGID) + Captcha + Face authentication or Biometrics succeed (no OTP), the system loads the officer's active sanctioned-post occupancies. If more than one is active, the user must choose which post to work under for this session. Each choice shows Post Name and Office details so dual-charge / multi-office officers can pick the correct context.",
        ),
        (
            "Transfer In assigns a DSR Officer to a post within the superior's office span and only where the actor's session Post is the immediate parent of the target Post (same rules as Transfer Out — FR-UM-057, FR-UM-059, FR-UM-043). The system shall capture Transfer Order / Reporting Order and Joining Date (FR-UM-060). Posts with available capacity (Occupied < Sanctioned Strength — FR-UM-066(a)) are selectable without a prior relieving. A post at full strength may receive a future-dated Transfer In only when relieving is already recorded for that Post + Office; recording it reserves capacity immediately (FR-UM-067). A post need not be wholly unoccupied to receive a Transfer In. The officer may log in for that post only from 12:00 AM IST on the Joining Date (FR-UM-061); the occupancy refresh job activates reserved occupancies (FR-UM-068).",
            "Transfer In assigns a DSR Officer to a post within the superior's office span and only where the actor's session Post is the immediate parent of the target Post (same rules as Transfer Out — FR-UM-057, FR-UM-059, FR-UM-043). The system shall capture Transfer Order / Reporting Order only — Joining Date is not captured (FR-UM-060). Transfer In is allowed only when the target post has available capacity (Occupied < Sanctioned Strength — FR-UM-066(a)); if the post is full, Transfer In is blocked until relieving has taken effect and capacity is free. Reserved / future-dated Transfer In is retired (FR-UM-061, FR-UM-067). On confirm, the occupancy is active immediately and the officer may use the post at the next login (FR-UM-052).",
        ),
        (
            "Example (handover with reservation): District Registrar at DRO Bengaluru relieves the Sub-Registrar of SRO Yeshwanthapura with Relieving Date 31-Aug-2026. The same DRO then records Transfer In of the incoming officer to that Sub-Registrar post with Joining Date 01-Sep-2026 — even though the outgoing occupancy has not yet ended. Capacity is reserved immediately (FR-UM-067). The incoming officer cannot log in under that post before 01-Sep-2026 12:00 AM IST. Shortly after midnight IST on 01-Sep the occupancy refresh job (FR-UM-068) de-allocates the outgoing officer and activates the incoming occupancy; Sanctioned Posts occupied count stays 1. IGR at Head Office cannot perform this Transfer In because Sub-Registrar does not report immediately to IGR.",
            "Example (handover after capacity frees): District Registrar at DRO Bengaluru relieves the Sub-Registrar of SRO Yeshwanthapura with Relieving Date 31-Aug-2026. Transfer In of the incoming officer to that Sub-Registrar post cannot be recorded while the post remains at full strength. Shortly after midnight IST following the Relieving Date, the occupancy refresh job (FR-UM-068) de-allocates the outgoing officer and frees capacity. The DRO then records Transfer In with Transfer Order only (no Joining Date); the occupancy is active immediately (FR-UM-060). IGR at Head Office cannot perform this Transfer In because Sub-Registrar does not report immediately to IGR.",
        ),
        (
            "Approved process diagram — P-10 Transfer In — FR-UM-060, FR-UM-061, FR-UM-066(a), FR-UM-067 — ProcessDiagrams/User_Management/P-10_Transfer_In.drawio",
            "Approved process diagram — P-10 Transfer In — FR-UM-060, FR-UM-066(a) — ProcessDiagrams/User_Management/P-10_Transfer_In.drawio",
        ),
        (
            "Post assignments on the user record and occupied counts on the Sanctioned Posts Master shall not be left to wait for a user to log in. A scheduled job shall run shortly after midnight each day (FR-UM-068) so that Transfer Out, Transfer In, and occupancy End Date take effect on the correct calendar day before the first login.",
            "Post assignments on the user record and occupied counts on the Sanctioned Posts Master shall not be left to wait for a user to log in. A scheduled job shall run shortly after midnight each day (FR-UM-068) so that Transfer Out / relieving takes effect on the correct calendar day before the first login. Transfer In itself takes effect immediately when recorded against available capacity (FR-UM-060).",
        ),
        (
            "The system shall provide an occupancy-refresh report (FR-UM-068) showing, for each midnight job run, occupancies de-allocated (relieving / End Date), reserved Transfer In occupancies activated, and Sanctioned Posts occupied-count changes (Post + Office).",
            "The system shall provide an occupancy-refresh report (FR-UM-068) showing, for each midnight job run, occupancies de-allocated on Relieving Date, and Sanctioned Posts occupied-count changes (Post + Office).",
        ),
        (
            "The system shall provide a Transfer Out / Transfer In history report (FR-UM-057–FR-UM-061, FR-UM-067) over a selected date range showing each relieving and Transfer In event — actor, officer, Post + Office, order reference, Relieving Date or Joining Date, and resulting occupancy state.",
            "The system shall provide a Transfer Out / Transfer In history report (FR-UM-057–FR-UM-060) over a selected date range showing each relieving and Transfer In event — actor, officer, Post + Office, order reference, Relieving Date / Relieving Reason or Transfer Order, and resulting occupancy state.",
        ),
        (
            "UAT — FR-UM-061 / FR-UM-067: Future-dated Transfer In against a full post is allowed only after relieving is recorded; reserved occupancy counts toward Occupied immediately; login for that post is blocked until 12:00 AM IST of Joining Date; the midnight job activates the occupancy on Joining Date.",
            "UAT — FR-UM-060 / FR-UM-066(a): Transfer In against a full post is blocked; after relieving takes effect and capacity is free, Transfer In with Transfer Order only succeeds immediately with no Joining Date; the officer can select the post at next login (FR-UM-052).",
        ),
    ]

    for p in doc.paragraphs:
        for old, new in para_replacements:
            if old in p.text:
                replace_in_paragraph(p, old, new)

    # Soft-fix remaining Joining / reserved phrases in paragraphs if exact match missed
    for p in doc.paragraphs:
        t = p.text
        if "Joining Date" in t or "reserved Transfer In" in t or "FR-UM-067" in t:
            # Skip if already handled; apply targeted small fixes
            nt = t
            nt = nt.replace(
                "Transfer out / relieving and Transfer In are scoped to offices under the actor, then to posts whose immediate parent is the actor's session Post (6.6.3–6.6.4, FR-UM-057–FR-UM-061, FR-UM-067); Relieving Reason is mandatory (FR-UM-087).",
                "Transfer out / relieving and Transfer In are scoped to offices under the actor, then to posts whose immediate parent is the actor's session Post (6.6.3–6.6.4, FR-UM-057–FR-UM-060); Relieving Reason is mandatory (FR-UM-087); Transfer In requires available capacity and has no Joining Date.",
            )
            if nt != t:
                set_paragraph_text(p, nt)

    # -------------------------------------------------------------------------
    # Process / matrix tables
    # -------------------------------------------------------------------------
    t13 = doc.tables[13]
    set_cell_text(
        t13.rows[2].cells[2],
        "Username (KGID) + Captcha + Face authentication or Biometrics (no OTP); then FR-UM-052",
    )

    t14 = doc.tables[14]
    set_cell_text(
        t14.rows[3].cells[1],
        "Dispatch login OTP to the registered mobile only (Citizens and Other Department)",
    )
    set_cell_text(
        t14.rows[3].cells[3],
        "FR-UM-010; never to email; not used for DSR Officer login (FR-UM-006)",
    )
    set_cell_text(
        t14.rows[4].cells[1],
        "Citizen / Other Dept: enter OTP; DSR: Face authentication or Biometrics (no OTP)",
    )
    set_cell_text(
        t14.rows[4].cells[3],
        "FR-UM-005 / FR-UM-006 / FR-UM-007",
    )

    t17 = doc.tables[17]
    set_cell_text(
        t17.rows[1].cells[1],
        "Authenticate (Username (KGID) + Captcha + Face authentication or Biometrics — no OTP)",
    )
    set_cell_text(
        t17.rows[2].cells[3],
        "Relieved occupancies excluded after Relieving Date ends (FR-UM-058, FR-UM-068)",
    )
    set_cell_text(
        t17.rows[3].cells[3],
        "No selection UI; only active occupancies (FR-UM-052)",
    )

    # T35 runtime
    t35 = doc.tables[35]
    set_cell_text(
        t35.rows[1].cells[1],
        "User authenticates (Citizen/Other Dept: Username + Captcha + OTP; DSR: Username + Captcha + Face/Biometrics — no OTP)",
    )

    # T16 last step may mention OTP + Biometrics for DSR
    t16 = doc.tables[16]
    for row in t16.rows:
        if "Biometrics" in row.cells[1].text or "face/biometric" in row.cells[1].text.lower():
            set_cell_text(
                row.cells[1],
                "User logs in: Other Dept with Username + Captcha + OTP; DSR with Username + Captcha + Face/Biometrics (no OTP)",
            )
            set_cell_text(row.cells[3], "FR-UM-007 / FR-UM-006; FR-UM-086 for DSR mobile change")

    # T45 relieving — remove reserved Transfer In note
    t45 = doc.tables[45]
    set_cell_text(
        t45.rows[8].cells[3],
        "FR-UM-058, FR-UM-068; occupied count −1; Transfer In may then proceed once capacity is free (FR-UM-060)",
    )

    # T46 Transfer In workflow — rewrite joining/reservation steps
    t46 = doc.tables[46]
    set_cell_text(
        t46.rows[4].cells[1],
        "If post has available capacity — proceed; if at full strength — block Transfer In",
    )
    set_cell_text(
        t46.rows[4].cells[3],
        "FR-UM-066(a), FR-UM-060; reserved / future-dated Transfer In retired (FR-UM-067)",
    )
    set_cell_text(
        t46.rows[7].cells[1],
        "Joining Date — NOT captured (removed v4.20)",
    )
    set_cell_text(t46.rows[7].cells[2], "—")
    set_cell_text(
        t46.rows[7].cells[3],
        "FR-UM-061 retired; occupancy takes effect immediately on confirm",
    )
    set_cell_text(
        t46.rows[8].cells[1],
        "Confirm Transfer In — occupancy active immediately; occupied count +1",
    )
    set_cell_text(
        t46.rows[8].cells[3],
        "Audit-logged (FR-UM-060); no reservation path",
    )
    set_cell_text(
        t46.rows[9].cells[1],
        "Officer may select the post at next login (FR-UM-052)",
    )
    set_cell_text(t46.rows[9].cells[2], "System")
    set_cell_text(
        t46.rows[9].cells[3],
        "No Joining Date gate (FR-UM-061 retired)",
    )
    set_cell_text(
        t46.rows[10].cells[1],
        "Occupancy refresh job — no Transfer In activation step (Joining Date removed)",
    )
    set_cell_text(
        t46.rows[10].cells[3],
        "FR-UM-068 handles relieving de-allocation only for day-boundary effects",
    )

    # T47 occupancy refresh steps
    t47 = doc.tables[47]
    for row in t47.rows:
        if "Activate reserved Transfer In" in row.cells[1].text or "Joining Date is today" in row.cells[1].text:
            set_cell_text(row.cells[1], "Reserved Transfer In activation — REMOVED (v4.20)")
            set_cell_text(row.cells[2], "—")
            set_cell_text(
                row.cells[3],
                "FR-UM-061 / FR-UM-067 retired; Transfer In is immediate when capacity available",
            )

    # T24 sanctioned posts description
    t24 = doc.tables[24]
    for row in t24.rows:
        if "reserved Transfer In" in row.cells[1].text:
            set_cell_text(
                row.cells[1],
                row.cells[1].text.replace(
                    "view occupied count (including reserved Transfer In), remaining capacity",
                    "view occupied count, remaining capacity",
                ),
            )

    # -------------------------------------------------------------------------
    # NFR / Risks / Glossary
    # -------------------------------------------------------------------------
    t52 = doc.tables[52]
    # Biometric NFR already mentions DSR — ensure OTP note if present
    for row in t52.rows:
        cell = row.cells[1].text if len(row.cells) > 1 else ""
        if "Login and authentication requests shall complete within 2 seconds after OTP/biometric verification" in cell:
            set_cell_text(
                row.cells[1],
                "Login and authentication requests shall complete within 2 seconds after OTP verification (Citizen / Other Department) or face/biometric verification (DSR).",
            )

    t53 = doc.tables[53]
    for row in t53.rows:
        risk = row.cells[0].text
        mit = row.cells[2].text if len(row.cells) > 2 else ""
        if risk.startswith("Face / biometric device") or "Biometric device" in risk:
            set_cell_text(
                row.cells[2],
                (
                    "Face authentication or biometrics are mandatory on every login for DSR Officers "
                    "(FR-UM-006) with no OTP and no exception path — offices must maintain working "
                    "devices; Other Department users do not require biometrics (FR-UM-007)."
                ),
            )
        if "Joining Date login gate" in risk:
            set_cell_text(row.cells[0], "Transfer In without available capacity")
            set_cell_text(
                row.cells[2],
                "Block Transfer In when Occupied >= Sanctioned Strength (FR-UM-060, FR-UM-066(a)); Joining Date / reservation path retired (FR-UM-061, FR-UM-067).",
            )
        if "Double-booking a full post" in risk:
            set_cell_text(
                row.cells[2],
                "Transfer In allowed only when available capacity exists (FR-UM-060); full posts remain blocked until relieving frees capacity via FR-UM-068.",
            )
        if "Post occupancy end date / relieving not enforced" in risk:
            set_cell_text(
                row.cells[2],
                "Occupancy refresh job shortly after midnight (FR-UM-068) de-allocates relieving (FR-UM-058); Transfer In is immediate when capacity is available (FR-UM-060); audit alert on job failure.",
            )
        if "Unauthorised relieving / transfer in" in risk:
            set_cell_text(
                row.cells[2],
                "Scope to offices under actor (FR-UM-059) then immediate-parent posts only (FR-UM-043, FR-UM-057); seeing a descendant office does not grant SR transfer to IGR/DIGR; available capacity required for Transfer In (FR-UM-060, FR-UM-066(a)).",
            )
        if "Repeated OTP + biometric" in risk or "face/biometrics on every fresh DSR login" in mit:
            set_cell_text(
                row.cells[2],
                (
                    "FR-UM-074 (10-minute idle timeout) combined with FR-UM-006 (Captcha + "
                    "face/biometrics on every fresh DSR login, no OTP) adds authentication overhead "
                    "at busy Sub-Registrar counters. Beyond functional UAT (9), measure end-to-end "
                    "re-authentication time during performance testing with Kaveri IT Cell."
                ),
            )

    t54 = doc.tables[54]
    for row in t54.rows:
        key = row.cells[0].text.strip()
        if key == "Available capacity":
            set_cell_text(
                row.cells[1],
                (
                    "Vacancy test used for post assignment and Transfer In: occupied count (active "
                    "occupancies) for a Post at an Office is less than its sanctioned strength "
                    "(FR-UM-066(a)); Transfer In requires this condition at the time of recording "
                    "(FR-UM-060)"
                ),
            )
        elif key == "Wholly unoccupied post":
            set_cell_text(
                row.cells[1],
                (
                    "Vacancy test used only for FR-UM-053 post-login additional charge: the occupied "
                    "count for a Post at an Office is zero, irrespective of sanctioned strength "
                    "(FR-UM-066(b))"
                ),
            )
        elif key == "Transfer In":
            set_cell_text(
                row.cells[1],
                (
                    "Process by which a superior assigns a DSR Officer to a post within office span "
                    "and immediate-parent posts only, capturing Transfer/Reporting Order (no Joining "
                    "Date), only when the post has available capacity; occupancy is active "
                    "immediately (FR-UM-060, FR-UM-066(a))"
                ),
            )
        elif key == "Reserved occupancy":
            set_cell_text(
                row.cells[1],
                (
                    "RETIRED (v4.20) — Previously a future-dated Transfer In that counted toward "
                    "Occupied until Joining Date; removed. Transfer In requires available capacity "
                    "and takes effect immediately (FR-UM-060)"
                ),
            )
        elif key == "Joining Date":
            set_cell_text(
                row.cells[1],
                (
                    "RETIRED (v4.20) — Previously captured during Transfer In to gate login until "
                    "12:00 AM IST of that date; no longer captured (FR-UM-061 retired)"
                ),
            )
        elif key == "Occupancy refresh job":
            set_cell_text(
                row.cells[1],
                (
                    "Scheduled job shortly after midnight IST that de-allocates relieved "
                    "occupancies, recalculates Sanctioned Posts occupied counts, and refreshes "
                    "officers' effective post assignments (FR-UM-068); no Joining Date / reserved "
                    "Transfer In activation"
                ),
            )
        elif key == "Face authentication":
            set_cell_text(
                row.cells[1],
                (
                    "Facial biometric verification used as the biometric factor for DSR Officer "
                    "login (with Captcha; no OTP) as an alternative to fingerprint / other "
                    "biometrics (FR-UM-006)"
                ),
            )

    # Broad residual replacements across tables/paragraphs for DSR OTP phrasing
    residual = [
        ("Username (KGID) + Captcha + OTP to mobile + Face authentication or Biometrics",
         "Username (KGID) + Captcha + Face authentication or Biometrics (no OTP)"),
        ("Username (KGID) + Captcha + OTP + Face authentication or Biometric authentication",
         "Username (KGID) + Captcha + Face authentication or Biometric authentication"),
        ("+ Face authentication or Biometrics for DSR only",
         "+ Face authentication or Biometrics for DSR only (no OTP)"),
        ("Enter OTP (+ Face authentication or Biometrics for DSR only)",
         "Citizen/Other Dept: enter OTP; DSR: Face/Biometrics only (no OTP)"),
        ("OTP, and face/biometric authentication for DSR Officers",
         "OTP for Citizens/Other Department; face/biometric for DSR Officers (no OTP)"),
        ("Authenticate (Username (KGID) + Captcha + OTP to mobile + Face authentication or Biometrics)",
         "Authenticate (Username (KGID) + Captcha + Face authentication or Biometrics — no OTP)"),
        ("+ Face/Biometrics if DSR only",
         "+ Face/Biometrics if DSR only (no OTP)"),
    ]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                new_text = text
                for old, new in residual:
                    if old in new_text:
                        new_text = new_text.replace(old, new)
                if new_text != text:
                    set_cell_text(cell, new_text)

    for p in doc.paragraphs:
        text = p.text
        new_text = text
        for old, new in residual:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(p, new_text)

    doc.save(str(DST))
    print(f"Saved {DST}")


if __name__ == "__main__":
    main()
