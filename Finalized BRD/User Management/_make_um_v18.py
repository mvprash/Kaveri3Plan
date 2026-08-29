# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v1.8.docx — v1.7 plus sanctioned posts per role per office."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\User Management")
TEMPLATE = BASE / "Template" / "User_Management_Module_BRD_Template.docx"
DST = BASE / "BRD_User_Management_v1.8.docx"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_table_rows(table, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "D9E2F3")
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(cells[i], val)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph_after(ref_paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    para = Paragraph(new_p, ref_paragraph._parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_table_after(ref_paragraph, headers: list[str], rows: list[tuple[str, ...]]):
    """Insert a styled table immediately after ref_paragraph. Returns the table element."""
    spacer = OxmlElement("w:p")
    ref_paragraph._element.addnext(spacer)

    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def build() -> Document:
    doc = Document(str(TEMPLATE))

    cover_map = {
        1: "BUSINESS REQUIREMENTS DOCUMENT",
        2: "User Management Module — KAVERI 3.0",
        3: "Prepared for: Department of Stamps & Registration, Government of Karnataka",
        4: "Version 1.8",
        5: "Date: 29 August 2026",
        6: "Prepared by: Nandha Kumar",
    }
    for idx, text in cover_map.items():
        replace_paragraph_text(doc.paragraphs[idx], text)

    replace_paragraph_text(
        doc.paragraphs[12],
        "This Business Requirements Document (BRD) defines the business requirements for the "
        "User Management Module of KAVERI 3.0 — the integrated platform of the Department of "
        "Stamps and Registration (DSR), Government of Karnataka. It describes user categories, "
        "OTP-based authentication (no password management), the DSR division and role catalogue, "
        "and sanctioned posts with approved strength per role for each office. It serves as the "
        "agreed basis for design, development, testing, and sign-off.",
    )
    replace_paragraph_text(
        doc.paragraphs[14],
        "KAVERI 3.0 requires a centralized mechanism to manage user identities, roles, sanctioned "
        "posts, and access permissions across citizen and departmental users. Authentication shall "
        "be passwordless — Username with OTP and Captcha for citizens, and Username with OTP, "
        "Captcha, and Biometrics for departmental users. Department users shall be mapped only to "
        "sanctioned posts defined in the posts master. Password management is explicitly out of scope.",
    )

    scope_in = {
        17: "User registration for three categories: Public users (Citizens), Department users (DSR Officers), and Other Department users",
        18: "OTP-based authentication (login, logout) — Username + OTP + Captcha; biometrics for departmental users",
        19: "Role-based access control (RBAC) aligned to DSR division and role catalogue",
        20: "Sanctioned posts master capturing the number of sanctioned posts (approved strength) for each role at each office",
        21: "Mapping of DSR department users exclusively to sanctioned posts; blocking of over-capacity assignment",
        22: "User profile management (view, update, deactivate)",
        23: "Administrative user management (create, edit, suspend, deactivate users)",
    }
    for idx, text in scope_in.items():
        replace_paragraph_text(doc.paragraphs[idx], text)
    insert_paragraph_after(
        doc.paragraphs[23],
        "Audit logging of user-related activities",
        "List Paragraph",
    )

    for p in doc.paragraphs:
        if "password resets" in p.text.lower():
            replace_paragraph_text(
                p,
                "Reduce support overhead related to account access issues through reliable OTP delivery.",
            )

    for p in doc.paragraphs:
        if "Eliminate password-related" in p.text:
            remove_paragraph(p)
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.3 Password Management":
            replace_paragraph_text(p, "6.3 OTP Login and Account Recovery")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "Enforce role-based access control to protect sensitive data and functionality.":
            insert_paragraph_after(
                p,
                "Maintain sanctioned post strength per role per office and prevent over-capacity user assignment.",
                "List Paragraph",
            )
            break

    for p in doc.paragraphs:
        if p.text.strip() == "The system shall provide a report of role and permission assignments across all users.":
            insert_paragraph_after(
                p,
                "The system shall provide a sanctioned post occupancy report showing, for each office, "
                "the number of sanctioned posts per role, occupied count, and vacant slots.",
                "List Paragraph",
            )
            break

    replace_table_rows(
        doc.tables[0],
        ["Document Control", ""],
        [
            ("Status", "Draft / In review"),
            ("Owner", "Nandha Kumar (BA) / Prashanth (PO)"),
            ("Reviewers", "Prabhakar Naik (Domain Expert); Kaveri IT Cell"),
            ("Approved By", "Prashanth (Product Owner)"),
        ],
    )

    replace_table_rows(
        doc.tables[1],
        ["Version", "Date", "Author", "Description"],
        [
            (
                "1.7",
                "28-Aug-2026",
                "Nandha Kumar",
                "Aligned to User Management BRD template; three user categories with OTP-only "
                "authentication (no password management); DSR division and role catalogue",
            ),
            (
                "1.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added sanctioned posts master: capture number of sanctioned posts (approved strength) "
                "for each role at each office; DSR user mapping to sanctioned posts only; "
                "over-capacity blocking and occupancy reporting",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[2],
        ["Name / Role", "Department", "Responsibility"],
        [
            ("Prashanth", "DSR / Product", "Prioritizes requirements; approves scope"),
            ("Nandha Kumar", "Business Analysis", "Documents and validates requirements"),
            ("Prabhakar Naik", "Domain Expert", "Validates DSR roles, sanctioned posts, and organizational structure"),
            ("Kaveri IT Cell", "Engineering", "Reviews technical feasibility"),
            ("Citizens (Public users)", "External", "Self-register and access citizen portal services"),
            (
                "DSR Officers & Other Department users",
                "Government",
                "Access departmental modules via OTP + biometrics",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[3],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-01", "The system shall support instant self-registration for Public users (Citizens) with no approval workflow.", "High"),
            ("FR-02", "The system shall allow Department users (DSR Officers) to be created only by authorised administrative roles.", "High"),
            ("FR-03", "The system shall allow Other Department users (officers/staff from other government departments) to be created only by authorised administrative roles.", "High"),
            ("FR-04", "The system shall prevent duplicate registrations using the same Username within a user category.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[4],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-05", "Public users (Citizens) shall authenticate using Username + OTP + Captcha on every login.", "High"),
            ("FR-06", "Department users (DSR Officers) shall authenticate using Username + OTP + Captcha + Biometrics on every login.", "High"),
            ("FR-07", "Other Department users shall authenticate using Username + OTP + Captcha + Biometrics on every login.", "High"),
            ("FR-08", "The system shall allow users to log out and terminate their active session.", "High"),
            ("FR-09", "The system shall not provide password-based login, password reset, or password change for any user category.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[5],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-10", "The system shall dispatch OTP to the user's registered mobile and/or email within 5 seconds of request.", "High"),
            ("FR-11", "The system shall validate Captcha before OTP dispatch for all user categories.", "High"),
            ("FR-12", "The system shall allow account recovery via OTP verification to registered mobile/email (no password reset).", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[6],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-13", "The system shall allow users to view and update their profile information (name, mobile, email).", "High"),
            ("FR-14", "The system shall allow users to upload and update a profile photo where applicable.", "Medium"),
            ("FR-15", "The system shall allow administrators to deactivate user accounts with reason and audit trail.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[7],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-16", "The system shall support RBAC with roles mapped to the DSR division and role catalogue (Section 6.5.1).", "High"),
            (
                "FR-17",
                "The system shall allow administrators to assign one or more roles to Department users, "
                "mapping each DSR officer to a sanctioned post (role + office) from the posts master.",
                "High",
            ),
            ("FR-18", "The system shall restrict access to features and data based on the user's assigned role(s).", "High"),
            ("FR-19", "The system shall maintain the DSR organizational divisions and roles as defined in Section 6.5.1.", "High"),
            (
                "FR-24",
                "The system shall maintain a sanctioned posts master listing all DSR roles/posts "
                "(e.g. IGR, DIGR, AIGR, Sub-Registrar, FDA, SDA, DRO, HQA) with the number of "
                "sanctioned posts (approved strength) for each role at each office.",
                "High",
            ),
            (
                "FR-25",
                "The system shall allow authorised administrators to configure and update sanctioned "
                "strength per role per office, with audit trail of changes.",
                "High",
            ),
            (
                "FR-26",
                "DSR department users shall be mapped only to posts defined in the sanctioned posts "
                "master. Assignment to an unlisted post or exceeding sanctioned strength for a "
                "role at an office shall be blocked.",
                "High",
            ),
            (
                "FR-27",
                "The system shall display vacant vs occupied sanctioned posts per role per office "
                "and prevent over-capacity assignment.",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[8],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-20", "The system shall allow administrators to create, edit, suspend, and deactivate user accounts by category.", "High"),
            ("FR-21", "The system shall allow administrators to search and filter users by category, role, office, division, and status.", "Medium"),
            ("FR-22", "The system shall log all administrative actions performed on user accounts.", "High"),
            ("FR-23", "The system shall notify a user via SMS/email when their account is created, suspended, or deactivated.", "Medium"),
        ],
    )

    replace_table_rows(
        doc.tables[9],
        ["Category", "Requirement"],
        [
            ("Security", "No password storage — authentication is OTP-based only for all user categories."),
            ("Security", "All data in transit shall be encrypted using TLS 1.2 or higher."),
            ("Security", "Biometric data for departmental users shall comply with Aadhaar Act, 2016 and UIDAI guidelines."),
            ("Performance", "OTP dispatch shall complete within 5 seconds of user request."),
            ("Performance", "Login and authentication requests shall complete within 2 seconds after OTP/biometric verification."),
            ("Availability", "The module shall maintain 99.9% uptime, excluding scheduled maintenance."),
            ("Usability", "Registration and OTP login workflows shall be completable on desktop and mobile browsers."),
            ("Auditability", "All create, update, delete, login, and access-control actions shall be logged with timestamp and actor."),
            ("Compliance", "The module shall comply with Karnataka e-Governance, MeitY/CERT-In, and applicable data protection norms."),
        ],
    )

    replace_table_rows(
        doc.tables[10],
        ["Risk", "Impact", "Mitigation"],
        [
            ("OTP delivery failure (SMS/email)", "High", "Multi-channel OTP delivery; retry mechanism; admin override procedure."),
            ("Biometric device unavailability", "Medium", "Define fallback procedure for DSR and Other Department users per security policy."),
            ("Over-capacity post assignment", "High", "Enforce sanctioned strength validation; block assignment when role at office is full."),
            ("Scope creep beyond OTP-only auth", "Medium", "Maintain strict change-control; no password management in KAVERI 3.0."),
            ("Delayed stakeholder sign-off", "Low", "Set clear review deadlines and escalation path."),
        ],
    )

    replace_table_rows(
        doc.tables[11],
        ["Term", "Definition"],
        [
            ("BRD", "Business Requirements Document"),
            ("RBAC", "Role-Based Access Control"),
            ("OTP", "One-Time Password — primary authentication credential (no static password)"),
            ("DSR", "Department of Stamps and Registration, Government of Karnataka"),
            ("UAT", "User Acceptance Testing"),
            ("IGR", "Inspector General of Registration"),
            ("DIGR", "Deputy Inspector General of Registration"),
            ("AIGR", "Assistant Inspector General of Registration"),
            ("Sanctioned post", "A role/post defined in the posts master with approved strength (headcount) at a specific office"),
            ("Sanctioned strength", "The approved number of occupants for a given role at a given office"),
            ("Office", "A concrete DSR office instance (e.g. SRO Yeshwanthapura, DRO Mysuru) in the organisational hierarchy"),
        ],
    )

    replace_table_rows(
        doc.tables[12],
        ["Name", "Role", "Signature", "Date"],
        [
            ("Prashanth", "Product Owner", "", ""),
            ("Prabhakar Naik", "Domain Expert", "", ""),
            ("Kaveri IT Cell Lead", "IT Security / Engineering", "", ""),
        ],
    )

    rbac_fr_tbl_el = doc.tables[7]._tbl
    rbac_heading = next(p for p in doc.paragraphs if p.text.strip() == "6.5 Role-Based Access Control (RBAC)")

    rbac_fr_tbl_el.getparent().remove(rbac_fr_tbl_el)
    rbac_heading._element.addnext(rbac_fr_tbl_el)

    sub1_el = OxmlElement("w:p")
    rbac_fr_tbl_el.addnext(sub1_el)
    sub1 = Paragraph(sub1_el, rbac_heading._parent)
    sub1.style = "Heading 3"
    sub1.add_run("6.5.1 DSR Division and Role Catalogue")
    note1 = insert_paragraph_after(
        sub1,
        "Department users (DSR Officers) shall be assigned roles from the following division structure:",
        "Normal",
    )
    div_tbl = insert_table_after(
        note1,
        ["Division", "Roles"],
        [
            ("Top Management", "IGR"),
            (
                "Division 1 — Admin, Law & Computers",
                "DIGR (Admin, Law & Computers), AIGR (Admin), HQA (Admin), SRO (Admin), FDA, SDA",
            ),
            ("Division 2 — Vigilance", "DIGR (Vigilance), Law Officer, HQA (RTI)"),
            (
                "Division 3 — Computers",
                "AIGR (Computers), System Integrator, PMU, Application Developer, SRO (Comp)",
            ),
            ("Division 4 — Enforcement", "DIGR (Enforcement), DRO, HQA, SRO, FDA, SDA"),
            (
                "Division 5 — Intelligence & Audit",
                "DIGR (Intelligence), AIGR (Audit), HQA (Audit), Superintendent (Audit)",
            ),
            ("Division 6 — CVC", "DIGR CVC, JD Town Planning"),
        ],
    )

    sub2_el = OxmlElement("w:p")
    div_tbl.addnext(sub2_el)
    sub2 = Paragraph(sub2_el, rbac_heading._parent)
    sub2.style = "Heading 3"
    sub2.add_run("6.5.2 User Categories and Authentication")

    note2 = insert_paragraph_after(
        sub2,
        "All user categories authenticate without passwords. OTP is the sole login credential.",
        "Normal",
    )
    cat_tbl = insert_table_after(
        note2,
        ["User Category", "Description", "Authentication"],
        [
            (
                "Public users (Citizens)",
                "Citizens accessing Kaveri portal services",
                "Username + OTP + Captcha",
            ),
            (
                "Department users (DSR Officers)",
                "Officers and staff of DSR",
                "Username + OTP + Captcha + Biometrics",
            ),
            (
                "Other Department users",
                "Officers/staff from other government departments",
                "Username + OTP + Captcha + Biometrics",
            ),
        ],
    )

    sub3_el = OxmlElement("w:p")
    cat_tbl.addnext(sub3_el)
    sub3 = Paragraph(sub3_el, rbac_heading._parent)
    sub3.style = "Heading 3"
    sub3.add_run("6.5.3 Sanctioned Posts per Office")

    note3 = insert_paragraph_after(
        sub3,
        "The system shall maintain a sanctioned posts master that records, for each office in the DSR "
        "hierarchy, the number of sanctioned posts (approved strength) for each role. Department "
        "users (DSR Officers) may be assigned only to vacant sanctioned posts. The system shall "
        "block assignment when sanctioned strength for a role at an office is already fully occupied.",
        "Normal",
    )
    example_note = insert_paragraph_after(
        note3,
        "Example (illustrative):",
        "Normal",
    )
    insert_table_after(
        example_note,
        ["Office", "Role", "Sanctioned Posts", "Occupied", "Vacant"],
        [
            ("SRO Yeshwanthapura", "Sub-Registrar (SR)", "1", "1", "0"),
            ("SRO Yeshwanthapura", "FDA", "2", "1", "1"),
            ("SRO Yeshwanthapura", "SDA", "1", "0", "1"),
            ("DRO Mysuru", "District Registrar (DR)", "1", "1", "0"),
            ("DRO Mysuru", "HQA", "1", "0", "1"),
        ],
    )

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v1.8"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
