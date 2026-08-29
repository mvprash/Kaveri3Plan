# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v2.8.docx — DSR users may hold multiple vacant sanctioned posts."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\User Management")
TEMPLATE = BASE / "Template" / "User_Management_Module_BRD_Template.docx"
DST = BASE / "BRD_User_Management_v2.8.docx"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_table_rows(table, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "D9E2F3")
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(cells[i], val)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph_after(ref_paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    para = Paragraph(new_p, ref_paragraph._parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_table_after(ref_paragraph, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_paragraph._element.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def insert_after_element(ref_el, parent, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_el.addnext(new_p)
    para = Paragraph(new_p, parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_heading_after(ref, text: str, style: str = "Heading 3", parent=None) -> Paragraph:
    if isinstance(ref, Paragraph):
        return insert_paragraph_after(ref, text, style)
    if parent is None:
        raise ValueError("parent required when ref is an XML element")
    return insert_after_element(ref, parent, text, style)


def insert_table_after_ref(ref_el, parent, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_el.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def build() -> Document:
    doc = Document(str(TEMPLATE))

    cover_map = {
        1: "BUSINESS REQUIREMENTS DOCUMENT",
        2: "User Management Module — KAVERI 3.0",
        3: "Prepared for: Department of Stamps & Registration, Government of Karnataka",
        4: "Version 2.8",
        5: "Date: 29 August 2026",
        6: "Prepared by: Nandha Kumar",
    }
    for idx, text in cover_map.items():
        replace_paragraph_text(doc.paragraphs[idx], text)

    replace_paragraph_text(
        doc.paragraphs[12],
        "This Business Requirements Document (BRD) defines the business requirements for the "
        "User Management Module of KAVERI 3.0 — the integrated platform of the Department of "
        "Stamps and Registration (DSR), Government of Karnataka. It describes a single unified "
        "Role Master and User Master differentiated by user category (Public/Citizen, DSR Officer, "
        "Other Department), OTP-based authentication (no password management), sanctioned posts "
        "with approved strength per role for each office, Primary and Secondary role assignment "
        "for DSR Officers, single-role assignment for Other Department users with optional account "
        "end date, a Module Master (Registration of Documents, Marriage Registration, Encumbrance Search, "
        "Certified Copy, and others), Role-to-Module-Function privilege mapping, Module Function and "
        "Resource (API/URL) masters, runtime access enforcement, DSR officer organisational hierarchy "
        "maintenance, and dedicated user-creation workflows. It serves as the agreed basis for design, "
        "development, testing, and sign-off.",
    )
    replace_paragraph_text(
        doc.paragraphs[14],
        "KAVERI 3.0 requires a centralized mechanism to manage user identities, roles, sanctioned "
        "posts, and access permissions. There shall be one Role Master and one User Master for all "
        "categories — Public users (Citizens), Department users (DSR Officers), and Other Department "
        "users — differentiated by User Category (and Role Category on roles). Authentication shall "
        "be passwordless — Username with OTP and Captcha for citizens, and Username with OTP, "
        "Captcha, and Biometrics for departmental users. DSR officers may be mapped to one or more "
        "vacant sanctioned posts — one mandatory Primary Role (substantive post) and optional "
        "additional sanctioned posts (Secondary Roles / additional charge), each only if vacant. "
        "Other Department users are assigned "
        "exactly one role (no Primary/Secondary distinction) and may have an optional account end "
        "date that deactivates the user when reached. Application modules (e.g. Registration of Documents, "
        "Marriage Registration, Encumbrance Search, Certified Copy) are maintained in a Module Master "
        "and mapped to roles via Module Functions and Resources (APIs/URLs) — DSR organisational roles "
        "are not named after modules. The system shall maintain the DSR officer reporting hierarchy "
        "(organisational chart). Application Admin maintains these masters. Password management "
        "is explicitly out of scope.",
    )

    scope_in = {
        17: "User registration for three categories in a single User Master: Public users (Citizens), Department users (DSR Officers), and Other Department users",
        18: "OTP-based authentication (login, logout) — Username + OTP + Captcha; biometrics for departmental users",
        19: "Single unified Role Master with Role Category differentiating Citizen, DSR, and Other Department roles (RBAC)",
        20: "Module Master, Module Function Master, Resource Master (API/URL), and Role–Module–Function mapping maintained by Application Admin",
        21: "DSR Officer Hierarchy Master — maintain reporting hierarchy for Department officers as per DSR organisational chart",
        22: "Sanctioned posts master; DSR Officers may be assigned multiple vacant sanctioned posts (one Primary + optional additional posts); single role for Other Department users with optional account end date",
        23: "Dedicated step-by-step workflows for role assignment during user creation (DSR and Other Department)",
    }
    for idx, text in scope_in.items():
        replace_paragraph_text(doc.paragraphs[idx], text)
    admin_scope = insert_paragraph_after(
        doc.paragraphs[23],
        "Administrative user management (create, edit, suspend, deactivate users)",
        "List Paragraph",
    )
    insert_paragraph_after(
        admin_scope,
        "Audit logging of user-related activities",
        "List Paragraph",
    )

    for p in doc.paragraphs:
        if "password resets" in p.text.lower():
            replace_paragraph_text(
                p,
                "Reduce support overhead related to account access issues through reliable OTP delivery.",
            )

    for p in doc.paragraphs:
        if "Eliminate password-related" in p.text:
            remove_paragraph(p)
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.3 Password Management":
            replace_paragraph_text(p, "6.3 OTP Login and Account Recovery")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.6 Administrative User Management":
            replace_paragraph_text(p, "6.7 Administrative User Management")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "Enforce role-based access control to protect sensitive data and functionality.":
            insert_paragraph_after(
                p,
                "Support Module Function and Resource (API/URL) masters with Role–Module–Function mapping and runtime access enforcement maintained by Application Admin.",
                "List Paragraph",
            )
            insert_paragraph_after(
                p,
                "Maintain the DSR Officer Hierarchy Master aligned to the Department organisational chart.",
                "List Paragraph",
            )
            break

    for p in doc.paragraphs:
        if p.text.strip() == "The system shall provide a report of role and permission assignments across all users.":
            sp_note = insert_paragraph_after(
                p,
                "The system shall provide a sanctioned post occupancy report showing, for each office, "
                "the number of sanctioned posts per role, occupied count, and vacant slots.",
                "List Paragraph",
            )
            insert_paragraph_after(
                sp_note,
                "The system shall provide a Role-to-Module mapping report showing which modules and "
                "functions are assigned to each role.",
                "List Paragraph",
            )
            break

    replace_table_rows(
        doc.tables[0],
        ["Document Control", ""],
        [
            ("Status", "Draft / In review"),
            ("Owner", "Nandha Kumar (BA) / Prashanth (PO)"),
            ("Reviewers", "Prabhakar Naik (Domain Expert); Kaveri IT Cell"),
            ("Approved By", "Prashanth (Product Owner)"),
        ],
    )

    replace_table_rows(
        doc.tables[1],
        ["Version", "Date", "Author", "Description"],
        [
            (
                "1.7",
                "28-Aug-2026",
                "Nandha Kumar",
                "Aligned to User Management BRD template; three user categories with OTP-only "
                "authentication (no password management); DSR division and role catalogue",
            ),
            (
                "1.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added sanctioned posts master: capture number of sanctioned posts (approved strength) "
                "for each role at each office; DSR user mapping to sanctioned posts only; "
                "over-capacity blocking and occupancy reporting",
            ),
            (
                "1.9",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Other Department role catalogue; Primary and Secondary role assignment; "
                "dedicated step-by-step user-creation workflows for DSR and Other Department users",
            ),
            (
                "2.0",
                "29-Aug-2026",
                "Nandha Kumar",
                "Primary/Secondary roles apply to DSR Officers only; Other Department users get "
                "exactly one role; optional account end date on Other Department user creation "
                "deactivates the user when reached (§6.6.2)",
            ),
            (
                "2.1",
                "29-Aug-2026",
                "Nandha Kumar",
                "Unified single Role Master and User Master for Citizens, DSR Officers, and Other "
                "Department users — differentiated by User Category / Role Category; no separate "
                "masters per category",
            ),
            (
                "2.2",
                "29-Aug-2026",
                "Nandha Kumar",
                "Listed seed roles in Role Master for Role Category = Citizen and Role Category = "
                "Other Department (§6.5.4)",
            ),
            (
                "2.3",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added DSR Department roles listing to §6.5.4 alongside Citizen and Other Department roles",
            ),
            (
                "2.4",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Module Master and Role-to-Module mapping (§6.5.6): Document Registration, "
                "Marriage Registration, Encumbrance Search, Certified Copy and related modules "
                "mapped to organisational / Citizen / Other Department roles",
            ),
            (
                "2.5",
                "29-Aug-2026",
                "Nandha Kumar",
                "Removed online/offline classification for document registration in §6.5.6 — "
                "module described as Registration of Documents only",
            ),
            (
                "2.6",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Module Function Master, Resource Master (API/URL), Role–Module–Function "
                "mapping with example rows, Application Admin CRUD, and runtime enforcement (§6.5.6)",
            ),
            (
                "2.7",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added DSR Officer Hierarchy Master (§6.5.7) to maintain Department officer reporting "
                "structure aligned to DSR organisational chart (Admin/Law/Computers, Vigilance, "
                "Computers, Enforcement, Intelligence & Audit, DIGR CVC)",
            ),
            (
                "2.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "DSR Officers may be assigned to multiple sanctioned posts when each post is vacant "
                "(Primary plus optional additional posts); over-capacity still blocked",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[2],
        ["Name / Role", "Department", "Responsibility"],
        [
            ("Prashanth", "DSR / Product", "Prioritizes requirements; approves scope"),
            ("Nandha Kumar", "Business Analysis", "Documents and validates requirements"),
            ("Prabhakar Naik", "Domain Expert", "Validates DSR roles, Other Department roles, sanctioned posts, and organizational structure"),
            ("Kaveri IT Cell", "Engineering", "Reviews technical feasibility"),
            ("Citizens (Public users)", "External", "Self-register and access citizen portal services"),
            (
                "DSR Officers & Other Department users",
                "Government",
                "Access departmental modules via OTP + biometrics",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[3],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-01", "The system shall support instant self-registration for Public users (Citizens) with no approval workflow.", "High"),
            ("FR-02", "The system shall allow Department users (DSR Officers) to be created only by authorised administrative roles.", "High"),
            ("FR-03", "The system shall allow Other Department users (officers/staff from other government departments) to be created only by authorised administrative roles.", "High"),
            ("FR-04", "The system shall prevent duplicate registrations using the same Username within a user category.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[4],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-05", "Public users (Citizens) shall authenticate using Username + OTP + Captcha on every login.", "High"),
            ("FR-06", "Department users (DSR Officers) shall authenticate using Username + OTP + Captcha + Biometrics on every login.", "High"),
            ("FR-07", "Other Department users shall authenticate using Username + OTP + Captcha + Biometrics on every login.", "High"),
            ("FR-08", "The system shall allow users to log out and terminate their active session.", "High"),
            ("FR-09", "The system shall not provide password-based login, password reset, or password change for any user category.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[5],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-10", "The system shall dispatch OTP to the user's registered mobile and/or email within 5 seconds of request.", "High"),
            ("FR-11", "The system shall validate Captcha before OTP dispatch for all user categories.", "High"),
            ("FR-12", "The system shall allow account recovery via OTP verification to registered mobile/email (no password reset).", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[6],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-13", "The system shall allow users to view and update their profile information (name, mobile, email).", "High"),
            ("FR-14", "The system shall allow users to upload and update a profile photo where applicable.", "Medium"),
            ("FR-15", "The system shall allow administrators to deactivate user accounts with reason and audit trail.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[7],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-16", "The system shall support RBAC using a single unified Role Master. Roles shall be differentiated by Role Category (Citizen, DSR, Other Department) as defined in Section 6.5.1. Users shall be stored in a single User Master differentiated by User Category.", "High"),
            (
                "FR-17",
                "The system shall allow administrators to assign roles to Department users (DSR Officers) "
                "during user creation: one mandatory Primary Role mapped to a vacant sanctioned post "
                "(role + office), and zero or more additional vacant sanctioned posts (Secondary Roles / "
                "additional charge) from the Role Master where Role Category = DSR.",
                "High",
            ),
            (
                "FR-28",
                "There shall be no separate Role Masters or User Masters for Citizens, DSR Officers, "
                "or Other Department users. All roles and users shall reside in one Role Master and one "
                "User Master respectively, differentiated by Role Category / User Category.",
                "High",
            ),
            (
                "FR-29",
                "During Other Department user creation, the system shall allow administrators to assign "
                "exactly one role from the Role Master filtered by Role Category = Other Department. "
                "Primary and Secondary role distinction does not apply to Other Department users.",
                "High",
            ),
            (
                "FR-30",
                "A Primary Role must be selected for each DSR Officer at account creation. "
                "The Primary Role is mapped to exactly one vacant sanctioned post, cannot have an end date, "
                "and represents the officer's substantive assignment. Primary Role does not apply to "
                "Other Department users or Citizens.",
                "High",
            ),
            (
                "FR-31",
                "The system shall allow administrators to assign multiple Secondary Roles to DSR Officers, "
                "each mapped to a vacant sanctioned post where applicable. Each Secondary Role must have "
                "a mandatory end date; on expiry the system shall revoke access and free the post occupancy. "
                "Secondary Roles do not apply to Other Department users or Citizens.",
                "High",
            ),
            (
                "FR-45",
                "A DSR Officer may be assigned to multiple sanctioned posts concurrently, provided each "
                "selected post is vacant at assignment. Non-vacant or over-capacity assignment shall be "
                "blocked. Occupied count shall increase per assigned post and decrease when occupancy ends "
                "(relieve, transfer, or Secondary Role expiry).",
                "High",
            ),
            (
                "FR-33",
                "During Other Department user creation, the system shall allow an optional End Date. "
                "If an End Date is entered, the system shall automatically deactivate the user account "
                "on that date and block login thereafter.",
                "High",
            ),
            (
                "FR-34",
                "When assigning a role to a user, the system shall present only roles whose Role Category "
                "matches the user's User Category (Citizen roles for Citizens, DSR roles for DSR Officers, "
                "Other Department roles for Other Department users).",
                "High",
            ),
            (
                "FR-35",
                "The Role Master shall include the seed roles listed in Section 6.5.4 for Role Category = "
                "Citizen, Role Category = Other Department, and Role Category = DSR (Department). "
                "Application Admin may add, edit, or deactivate roles within each Role Category.",
                "High",
            ),
            (
                "FR-36",
                "The system shall maintain a Module Master of application modules (including Registration "
                "of Documents, Marriage Registration, Encumbrance Search, Certified Copy, and others "
                "listed in Section 6.5.6). Modules are not DSR organisational roles; they define "
                "functional areas to which roles are mapped. Registration of Documents is a single "
                "module with no online/offline classification in this master.",
                "High",
            ),
            (
                "FR-37",
                "The system shall allow Application Admin to map roles to Module Functions "
                "(Role–Module–Function mapping). Access to a module capability is granted only through "
                "this mapping. Application Admin shall add, edit, and update mappings with audit trail.",
                "High",
            ),
            (
                "FR-39",
                "The system shall maintain a Module Function Master under each module (e.g. VIEW, ADD, "
                "EDIT, APPROVE, SIGN, PRINT, DOWNLOAD). Application Admin shall add, edit, and update "
                "module functions.",
                "High",
            ),
            (
                "FR-40",
                "The system shall maintain a Resource Master of APIs and URLs linked to Module Functions "
                "(resource type API or URL, HTTP method, path pattern). Application Admin shall add, "
                "edit, and update resources. Roles shall not store raw API/URL lists; access is resolved "
                "via Role → Module Function → Resource.",
                "High",
            ),
            (
                "FR-41",
                "At runtime the application shall enforce access as follows: resolve the user's assigned "
                "roles; load Role–Module–Function claims; for each API/URL request look up the Resource "
                "Master to obtain the required Module Function; allow only if the user holds that "
                "function; otherwise deny (HTTP 403) and audit the attempt.",
                "High",
            ),
            (
                "FR-42",
                "Application Admin shall create, edit, enable/disable, and update Module Master, "
                "Module Function Master, Resource Master, and Role–Module–Function mapping tables "
                "from the User Management admin console. All changes shall be audit-logged.",
                "High",
            ),
            (
                "FR-38",
                "A user's effective module access shall be the union of Module Functions mapped to all "
                "roles currently assigned to that user (Primary and Secondary for DSR Officers; the "
                "single assigned role for Other Department users; Citizen roles for Public users).",
                "High",
            ),
            ("FR-18", "The system shall restrict access to features, APIs, and URLs based on the user's assigned role(s) and Role–Module–Function / Resource mappings.", "High"),
            ("FR-19", "The system shall maintain DSR organizational divisions and DSR roles within the unified Role Master (Role Category = DSR) as defined in Section 6.5.1 and the DSR Officer Hierarchy Master in Section 6.5.7.", "High"),
            (
                "FR-43",
                "The system shall maintain a DSR Officer Hierarchy Master representing the reporting "
                "structure of Department officers (organisational chart). Each hierarchy node shall "
                "reference a DSR role (Role Category = DSR) and an optional parent node. Application "
                "Admin shall add, edit, reorder, enable/disable, and update hierarchy nodes with audit trail.",
                "High",
            ),
            (
                "FR-44",
                "The DSR Officer Hierarchy Master shall seed and support the structure: Additional Chief "
                "Secretary / Principal Secretary / Secretary → Inspector General of Registration & "
                "Commissioner of Stamps → Divisions (Admin Law & Computers; Vigilance; Computers; "
                "Enforcement; Intelligence & Audit; DIGR CVC) with subordinate posts as listed in "
                "Section 6.5.7.",
                "High",
            ),
            (
                "FR-24",
                "The system shall maintain a sanctioned posts master listing all DSR roles/posts "
                "(e.g. IGR, DIGR, AIGR, Sub-Registrar, FDA, SDA, DRO, HQA) with the number of "
                "sanctioned posts (approved strength) for each role at each office. Sanctioned "
                "strength applies only to roles with Role Category = DSR.",
                "High",
            ),
            (
                "FR-25",
                "The system shall allow authorised administrators to configure and update sanctioned "
                "strength per role per office, with audit trail of changes.",
                "High",
            ),
            (
                "FR-26",
                "DSR department users shall be mapped only to posts defined in the sanctioned posts "
                "master. A user may hold multiple post occupancies. Assignment to an unlisted post or "
                "to a post that is not vacant (would exceed sanctioned strength) shall be blocked.",
                "High",
            ),
            (
                "FR-27",
                "The system shall display vacant vs occupied sanctioned posts per role per office "
                "and prevent over-capacity assignment. When assigning an additional post to a DSR "
                "Officer, only vacant posts shall be selectable.",
                "High",
            ),
            (
                "FR-32",
                "The system shall provide dedicated step-by-step workflows for role assignment during "
                "user creation for DSR Officers (Section 6.6.1) and Other Department users (Section 6.6.2).",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[8],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-20", "The system shall allow administrators to create, edit, suspend, and deactivate user accounts by category.", "High"),
            ("FR-21", "The system shall allow administrators to search and filter users by category, role, office, division, and status.", "Medium"),
            ("FR-22", "The system shall log all administrative actions performed on user accounts.", "High"),
            ("FR-23", "The system shall notify a user via SMS/email when their account is created, suspended, or deactivated.", "Medium"),
        ],
    )

    replace_table_rows(
        doc.tables[9],
        ["Category", "Requirement"],
        [
            ("Security", "No password storage — authentication is OTP-based only for all user categories."),
            ("Security", "All data in transit shall be encrypted using TLS 1.2 or higher."),
            ("Security", "Biometric data for departmental users shall comply with Aadhaar Act, 2016 and UIDAI guidelines."),
            ("Performance", "OTP dispatch shall complete within 5 seconds of user request."),
            ("Performance", "Login and authentication requests shall complete within 2 seconds after OTP/biometric verification."),
            ("Availability", "The module shall maintain 99.9% uptime, excluding scheduled maintenance."),
            ("Usability", "Registration and OTP login workflows shall be completable on desktop and mobile browsers."),
            ("Auditability", "All create, update, delete, login, and access-control actions shall be logged with timestamp and actor."),
            ("Compliance", "The module shall comply with Karnataka e-Governance, MeitY/CERT-In, and applicable data protection norms."),
        ],
    )

    replace_table_rows(
        doc.tables[10],
        ["Risk", "Impact", "Mitigation"],
        [
            ("OTP delivery failure (SMS/email)", "High", "Multi-channel OTP delivery; retry mechanism; admin override procedure."),
            ("Biometric device unavailability", "Medium", "Define fallback procedure for DSR and Other Department users per security policy."),
            ("Over-capacity post assignment", "High", "Enforce sanctioned strength validation; block assignment when role at office is full."),
            ("Secondary role not expired on end date", "Medium", "Scheduled job to auto-revoke expired Secondary Roles for DSR Officers; audit alert."),
            ("Other Department end date not enforced", "Medium", "Scheduled job to deactivate Other Department users on optional End Date; audit alert."),
            ("Scope creep beyond OTP-only auth", "Medium", "Maintain strict change-control; no password management in KAVERI 3.0."),
            ("Delayed stakeholder sign-off", "Low", "Set clear review deadlines and escalation path."),
        ],
    )

    replace_table_rows(
        doc.tables[11],
        ["Term", "Definition"],
        [
            ("BRD", "Business Requirements Document"),
            ("RBAC", "Role-Based Access Control"),
            ("OTP", "One-Time Password — primary authentication credential (no static password)"),
            ("DSR", "Department of Stamps and Registration, Government of Karnataka"),
            ("UAT", "User Acceptance Testing"),
            ("IGR", "Inspector General of Registration"),
            ("DIGR", "Deputy Inspector General of Registration"),
            ("AIGR", "Assistant Inspector General of Registration"),
            ("Sanctioned post", "A DSR role/post defined in the posts master with approved strength (headcount) at a specific office"),
            ("Sanctioned strength", "The approved number of occupants for a given DSR role at a given office"),
            ("Office", "A concrete DSR office instance (e.g. SRO Yeshwanthapura, DRO Mysuru) in the organisational hierarchy"),
            ("Role Master", "Single master of all roles for Citizens, DSR Officers, and Other Department users; differentiated by Role Category"),
            ("User Master", "Single master of all users across categories; differentiated by User Category"),
            ("Role Category", "Attribute on a role in the Role Master: Citizen, DSR, or Other Department"),
            ("User Category", "Attribute on a user in the User Master: Public (Citizen), DSR Officer, or Other Department"),
            ("Module Master", "Catalogue of application modules (e.g. Registration of Documents, Marriage Registration, Encumbrance Search, Certified Copy)"),
            ("Module Function Master", "Catalogue of privilege verbs under a module — e.g. VIEW, ADD, EDIT, APPROVE, SIGN, PRINT, DOWNLOAD"),
            ("Resource Master", "Catalogue of APIs and URLs linked to a Module Function (type, HTTP method, path pattern)"),
            ("Role–Module–Function mapping", "Association of a role to one or more Module Functions; determines what a role-holder can do"),
            ("Module function", "Privilege verb within a module — e.g. VIEW, ADD, EDIT, APPROVE, SIGN, PRINT, DOWNLOAD"),
            ("Resource", "An API endpoint or UI URL/route protected by a Module Function"),
            ("Access enforcement", "Runtime check: request path/method → Resource → required Module Function → user's role claims"),
            ("DSR Officer Hierarchy Master", "Master of reporting relationships among DSR roles/posts as per the Department organisational chart"),
            ("Hierarchy node", "A node in the DSR Officer Hierarchy Master linking a DSR role to an optional parent role/node"),
            ("Other Department role", "A role in the Role Master with Role Category = Other Department; exactly one such role is assigned per Other Department user"),
            ("Primary Role", "The mandatory substantive DSR post assignment (exactly one vacant sanctioned post); cannot have an end date; does not apply to Other Department users or Citizens"),
            ("Secondary Role", "An additional DSR role/post assignment mapped to a vacant sanctioned post where applicable; mandatory end date; auto-expires and frees occupancy"),
            ("Post occupancy", "An active assignment of a DSR user to a specific sanctioned post at an office; a user may have multiple occupancies if each post is vacant"),
            ("Parent department", "The government department to which an Other Department user belongs (e.g. Revenue, Treasury, Police)"),
            ("Account End Date", "Optional date on an Other Department user account; if set, the system deactivates the user on that date"),
        ],
    )

    replace_table_rows(
        doc.tables[12],
        ["Name", "Role", "Signature", "Date"],
        [
            ("Prashanth", "Product Owner", "", ""),
            ("Prabhakar Naik", "Domain Expert", "", ""),
            ("Kaveri IT Cell Lead", "IT Security / Engineering", "", ""),
        ],
    )

    rbac_fr_tbl_el = doc.tables[7]._tbl
    rbac_heading = next(p for p in doc.paragraphs if p.text.strip() == "6.5 Role-Based Access Control (RBAC)")
    rbac_parent = rbac_heading._parent
    rbac_fr_tbl_el.getparent().remove(rbac_fr_tbl_el)
    rbac_heading._element.addnext(rbac_fr_tbl_el)

    sub1 = insert_heading_after(
        rbac_fr_tbl_el, "6.5.1 Unified Role Master and User Master", "Heading 3", parent=rbac_parent
    )
    note1 = insert_paragraph_after(
        sub1,
        "The system shall maintain a single Role Master and a single User Master for all user types. "
        "There shall be no separate masters for Citizens, DSR Officers, or Other Department users. "
        "Differentiation is by Role Category on roles and User Category on users. When assigning a role, "
        "the system shall filter the Role Master to roles matching the user's User Category (FR-34).",
        "Normal",
    )
    cat_diff_tbl = insert_table_after(
        note1,
        ["Master", "Differentiator", "Values", "Purpose"],
        [
            (
                "User Master",
                "User Category",
                "Public (Citizen); DSR Officer; Other Department",
                "One user store; category drives auth, creation path, and eligible roles",
            ),
            (
                "Role Master",
                "Role Category",
                "Citizen; DSR; Other Department",
                "One role store; category filters which roles can be assigned to which users",
            ),
        ],
    )
    div_intro = insert_heading_after(
        cat_diff_tbl,
        "DSR roles in the Role Master (Role Category = DSR) are organised by division as follows:",
        "Normal",
        parent=rbac_parent,
    )
    div_tbl = insert_table_after(
        div_intro,
        ["Division", "Roles (Role Category = DSR)"],
        [
            ("Top Management", "IGR"),
            (
                "Division 1 — Admin, Law & Computers",
                "DIGR (Admin, Law & Computers), AIGR (Admin), HQA (Admin), SRO (Admin), FDA, SDA",
            ),
            ("Division 2 — Vigilance", "DIGR (Vigilance), Law Officer, HQA (RTI)"),
            (
                "Division 3 — Computers",
                "AIGR (Computers), System Integrator, PMU, Application Developer, SRO (Comp)",
            ),
            ("Division 4 — Enforcement", "DIGR (Enforcement), DRO, HQA, SRO, FDA, SDA"),
            (
                "Division 5 — Intelligence & Audit",
                "DIGR (Intelligence), AIGR (Audit), HQA (Audit), Superintendent (Audit)",
            ),
            ("Division 6 — CVC", "DIGR CVC, JD Town Planning"),
        ],
    )

    sub2 = insert_heading_after(
        div_tbl, "6.5.2 User Categories and Authentication", "Heading 3", parent=rbac_parent
    )
    note2 = insert_paragraph_after(
        sub2,
        "All users are stored in the single User Master, differentiated by User Category. "
        "All user categories authenticate without passwords. OTP is the sole login credential.",
        "Normal",
    )
    cat_tbl = insert_table_after(
        note2,
        ["User Category", "Description", "Authentication"],
        [
            ("Public users (Citizens)", "Citizens accessing Kaveri portal services — User Category = Public (Citizen)", "Username + OTP + Captcha"),
            ("Department users (DSR Officers)", "Officers and staff of DSR — User Category = DSR Officer", "Username + OTP + Captcha + Biometrics"),
            (
                "Other Department users",
                "Officers/staff from other government departments — User Category = Other Department",
                "Username + OTP + Captcha + Biometrics",
            ),
        ],
    )

    sub3 = insert_heading_after(
        cat_tbl, "6.5.3 Sanctioned Posts per Office", "Heading 3", parent=rbac_parent
    )
    note3 = insert_paragraph_after(
        sub3,
        "The system shall maintain a sanctioned posts master that records, for each office in the DSR "
        "hierarchy, the number of sanctioned posts (approved strength) for each role. DSR Officers may "
        "be assigned to one or more vacant sanctioned posts (Primary plus optional additional posts). "
        "The system shall block assignment when a selected post is not vacant or when sanctioned "
        "strength for a role at an office would be exceeded (FR-45, FR-26).",
        "Normal",
    )
    example_note = insert_paragraph_after(note3, "Example (illustrative):", "Normal")
    sanc_tbl = insert_table_after(
        example_note,
        ["Office", "Role", "Sanctioned Posts", "Occupied", "Vacant"],
        [
            ("SRO Yeshwanthapura", "Sub-Registrar (SR)", "1", "1", "0"),
            ("SRO Yeshwanthapura", "FDA", "2", "1", "1"),
            ("SRO Yeshwanthapura", "SDA", "1", "0", "1"),
            ("DRO Mysuru", "District Registrar (DR)", "1", "1", "0"),
            ("DRO Mysuru", "HQA", "1", "0", "1"),
        ],
    )

    sub4 = insert_heading_after(
        sanc_tbl, "6.5.4 Role Master — Seed Roles by Role Category", "Heading 3", parent=rbac_parent
    )
    note4 = insert_paragraph_after(
        sub4,
        "All roles below reside in the single Role Master, differentiated by Role Category. Seed roles "
        "are listed for Citizen, Other Department, and DSR (Department). Application Admin may add further "
        "roles under the same Role Categories. Other Department users are assigned exactly one role with "
        "Role Category = Other Department. Primary and Secondary role distinction applies only to DSR "
        "Officers. Sanctioned posts apply only to Role Category = DSR.",
        "Normal",
    )
    citizen_intro = insert_paragraph_after(
        note4,
        "Roles with Role Category = Citizen:",
        "Normal",
    )
    citizen_tbl = insert_table_after(
        citizen_intro,
        ["Role", "Description"],
        [
            ("Citizen", "Default role assigned on instant self-registration; access to citizen portal services"),
            ("Marriage Applicant", "Apply for Hindu Marriage / Special Marriage registration and related citizen actions"),
            ("Document Registration Applicant", "Initiate and track document registration applications"),
            ("Certified Copy / EC Applicant", "Apply for Encumbrance Certificate (EC) and certified copy of registered documents"),
            ("Firm / Society Applicant", "Apply for firm / society related registration services where offered on the portal"),
            ("Stamp Duty / Challan Payer", "Pay stamp duty / registration fees and view payment history for citizen transactions"),
        ],
    )
    od_intro = insert_heading_after(
        citizen_tbl,
        "Roles with Role Category = Other Department:",
        "Normal",
        parent=rbac_parent,
    )
    other_tbl = insert_table_after(
        od_intro,
        ["Role", "Typical Parent Department", "Description"],
        [
            ("Revenue Verification Officer", "Revenue", "Verify land / revenue particulars linked to registration applications"),
            ("Bhoomi Cross-check User", "Revenue", "Cross-check Bhoomi / RTC data against registration records"),
            ("Treasury Payment Verifier", "Treasury / Finance", "Verify treasury / payment status for registration fees"),
            ("Khajane Reconciliation User", "Treasury / Finance", "Reconcile Khajane-II receipts with Kaveri transactions"),
            ("Police Enquiry Officer", "Police", "View / respond to police enquiry requests on registered documents"),
            ("FIR Verification User", "Police", "Verify FIR / crime particulars where required for registration workflow"),
            ("ULB Document Verifier", "Urban Local Body (ULB)", "Verify municipal / ULB documents submitted with applications"),
            ("Read-only MIS Viewer", "Any (cross-department)", "Read-only access to assigned MIS / dashboards; no transactional actions"),
            ("Document Upload User", "Any (cross-department)", "Upload supporting documents for assigned inter-department workflows"),
            ("Inter-Department Enquiry User", "Any (cross-department)", "Raise / respond to inter-department enquiries on applications"),
        ],
    )
    dsr_intro = insert_heading_after(
        other_tbl,
        "Roles with Role Category = DSR (Department):",
        "Normal",
        parent=rbac_parent,
    )
    dsr_roles_tbl = insert_table_after(
        dsr_intro,
        ["Division", "Role", "Description"],
        [
            ("Top Management", "IGR", "Inspector General of Registration — statewide oversight"),
            ("Division 1 — Admin, Law & Computers", "DIGR (Admin, Law & Computers)", "Deputy IGR for Admin, Law & Computers"),
            ("Division 1 — Admin, Law & Computers", "AIGR (Admin)", "Assistant IGR (Admin)"),
            ("Division 1 — Admin, Law & Computers", "HQA (Admin)", "Head Quarter Assistant (Admin)"),
            ("Division 1 — Admin, Law & Computers", "SRO (Admin)", "Sub-Registrar Office Admin functions"),
            ("Division 1 — Admin, Law & Computers", "FDA", "First Division Assistant"),
            ("Division 1 — Admin, Law & Computers", "SDA", "Second Division Assistant"),
            ("Division 1 — Admin, Law & Computers", "Accountant Superintendent (Admin)", "Accountant Superintendent (Admin)"),
            ("Division 1 — Admin, Law & Computers", "Typist", "Typist"),
            ("Division 1 — Admin, Law & Computers", "Statistical Inspector", "Statistical Inspector (RTI & Statistics)"),
            ("Division 2 — Vigilance", "DIGR (Vigilance)", "Deputy IGR (Vigilance)"),
            ("Division 2 — Vigilance", "Law Officer", "Departmental Law Officer"),
            ("Division 2 — Vigilance", "HQA (RTI)", "Head Quarter Assistant (RTI)"),
            ("Division 3 — Computers", "AIGR (Computers)", "Assistant IGR (Computers)"),
            ("Division 3 — Computers", "System Integrator", "System Integrator / SI support"),
            ("Division 3 — Computers", "PMU", "Project Management Unit"),
            ("Division 3 — Computers", "Application Developer", "Application development support"),
            ("Division 3 — Computers", "HQA / Project Manager (Comp)", "Head Quarter Assistant / Project Manager (Computers)"),
            ("Division 3 — Computers", "SRO (Comp)", "SRO Computers support"),
            ("Division 3 — Computers", "FDA", "First Division Assistant (Computers)"),
            ("Division 3 — Computers", "SDA", "Second Division Assistant (Computers)"),
            ("Division 4 — Enforcement", "DIGR (Enforcement)", "Deputy IGR (Enforcement)"),
            ("Division 4 — Enforcement", "DRO", "District Registrar / DRO"),
            ("Division 4 — Enforcement", "HQA", "Head Quarter Assistant (Enforcement)"),
            ("Division 4 — Enforcement", "SRO", "Sub-Registrar (office head)"),
            ("Division 4 — Enforcement", "FDA", "First Division Assistant (Enforcement offices)"),
            ("Division 4 — Enforcement", "SDA", "Second Division Assistant (Enforcement offices)"),
            ("Division 5 — Intelligence & Audit", "DIGR (Intelligence)", "Deputy IGR (Intelligence)"),
            ("Division 5 — Intelligence & Audit", "AIGR (Audit)", "Assistant IGR (Audit)"),
            ("Division 5 — Intelligence & Audit", "HQA (Audit)", "Head Quarter Assistant (Audit)"),
            ("Division 5 — Intelligence & Audit", "Superintendent (Audit)", "Superintendent (Audit)"),
            ("Division 5 — Intelligence & Audit", "FDA", "First Division Assistant (Audit)"),
            ("Division 5 — Intelligence & Audit", "SDA", "Second Division Assistant (Audit)"),
            ("Division 5 — Intelligence & Audit", "Typist", "Typist (Audit)"),
            ("Division 6 — CVC", "DIGR CVC", "Deputy IGR (CVC)"),
            ("Division 6 — CVC", "JD Town Planning", "Joint Director, Town Planning"),
            ("Field / SRO (common)", "DEO", "Data Entry Operator — SRO operational role"),
            ("Field / SRO (common)", "Sub-Registrar (SR)", "Sub-Registrar post holder for registration and signing"),
        ],
    )

    sub5 = insert_heading_after(
        dsr_roles_tbl, "6.5.5 Primary and Secondary Role Assignment (DSR Officers Only)", "Heading 3", parent=rbac_parent
    )
    primary_note = insert_paragraph_after(
        sub5,
        "Primary and Secondary role assignment applies only to Department users (DSR Officers). "
        "A DSR Officer shall have exactly one Primary Role (one vacant sanctioned post) and may be "
        "assigned additional vacant sanctioned posts as Secondary Roles / additional charge (FR-45). "
        "Other Department users and Citizens are assigned roles from the same Role Master filtered by "
        "their Role Category (see Sections 6.5.1, 6.5.4 and 6.6.2).",
        "Normal",
    )
    primary_tbl = insert_table_after(
        primary_note,
        ["Role Type", "Mandatory?", "Sanctioned post", "End Date", "Applies To", "Notes"],
        [
            (
                "Primary Role",
                "Yes — exactly one per DSR Officer",
                "Exactly one vacant post (role + office)",
                "Not permitted",
                "DSR Officers only",
                "Substantive assignment",
            ),
            (
                "Secondary Role / additional post",
                "No — zero or more",
                "Each must be a vacant sanctioned post",
                "Mandatory",
                "DSR Officers only",
                "Additional charge; auto-revoked on end date; occupancy freed",
            ),
            (
                "Other Department role",
                "Yes — exactly one per Other Department user",
                "N/A (not DSR sanctioned posts)",
                "N/A (account End Date optional — §6.6.2)",
                "Other Department users only",
                "Single role; no Primary/Secondary split",
            ),
        ],
    )

    sub6 = insert_heading_after(
        primary_tbl, "6.5.6 Module, Function, Resource Masters and Access Enforcement", "Heading 3", parent=rbac_parent
    )
    note6 = insert_paragraph_after(
        sub6,
        "DSR organisational roles (SR, FDA, DEO, etc.) are not named after application services. "
        "Access is modelled as: User → Role(s) → Module Function(s) → Resource(s) (API/URL). "
        "Registration of Documents is a single module (no online/offline classification). "
        "Application Admin shall create, edit, and update all masters in this section (FR-42).",
        "Normal",
    )
    hier_intro = insert_paragraph_after(
        note6,
        "Privilege hierarchy (masters):",
        "Normal",
    )
    hier_tbl = insert_table_after(
        hier_intro,
        ["Level", "Master", "Maintained by", "Purpose"],
        [
            ("1", "Module Master", "Application Admin", "Business modules (Registration of Documents, Marriage Registration, …)"),
            ("2", "Module Function Master", "Application Admin", "Functions under each module (VIEW, ADD, APPROVE, SIGN, …)"),
            ("3", "Resource Master", "Application Admin", "APIs and URLs linked to each Module Function"),
            ("4", "Role–Module–Function mapping", "Application Admin", "Which roles may perform which Module Functions"),
        ],
    )

    mod_intro = insert_heading_after(
        hier_tbl, "Module Master — seed rows:", "Normal", parent=rbac_parent
    )
    mod_tbl = insert_table_after(
        mod_intro,
        ["Module Code", "Module Name", "Description", "Status"],
        [
            ("MOD-DOC-REG", "Registration of Documents", "Registration of documents", "Active"),
            ("MOD-MARRIAGE", "Marriage Registration", "Marriage registration (Hindu Marriage and Special Marriage)", "Active"),
            ("MOD-EC", "Encumbrance Search", "Search and issue of Encumbrance Certificate (EC)", "Active"),
            ("MOD-CC", "Certified Copy", "Application and issue of certified copies of registered documents", "Active"),
            ("MOD-STAMP", "Stamp Duty / Payments", "Stamp duty assessment, challan, and payment reconciliation", "Active"),
            ("MOD-FIRM", "Firm / Society Registration", "Firm and society related registration services", "Active"),
            ("MOD-UM", "User Management", "Identity, roles, posts, modules, functions, resources, and mappings", "Active"),
            ("MOD-MIS", "MIS / Dashboards", "Operational and management reporting", "Active"),
        ],
    )

    fn_intro = insert_heading_after(
        mod_tbl, "Module Function Master — example rows:", "Normal", parent=rbac_parent
    )
    fn_tbl = insert_table_after(
        fn_intro,
        ["Function Code", "Module", "Function Name", "Description"],
        [
            ("FN-DOC-VIEW", "Registration of Documents", "VIEW", "View document registration applications / records"),
            ("FN-DOC-ADD", "Registration of Documents", "ADD", "Create / initiate document registration"),
            ("FN-DOC-EDIT", "Registration of Documents", "EDIT", "Edit draft / in-progress applications"),
            ("FN-DOC-APPROVE", "Registration of Documents", "APPROVE", "Approve application for registration"),
            ("FN-DOC-SIGN", "Registration of Documents", "SIGN", "Digitally sign registered document"),
            ("FN-DOC-PRINT", "Registration of Documents", "PRINT", "Print registration extracts / slips"),
            ("FN-MAR-VIEW", "Marriage Registration", "VIEW", "View marriage applications"),
            ("FN-MAR-ADD", "Marriage Registration", "ADD", "Create marriage application"),
            ("FN-MAR-APPROVE", "Marriage Registration", "APPROVE", "Approve / register marriage"),
            ("FN-EC-APPLY", "Encumbrance Search", "APPLY", "Apply for EC"),
            ("FN-EC-ISSUE", "Encumbrance Search", "ISSUE", "Issue / download EC"),
            ("FN-CC-APPLY", "Certified Copy", "APPLY", "Apply for certified copy"),
            ("FN-CC-ISSUE", "Certified Copy", "ISSUE", "Issue / download certified copy"),
            ("FN-UM-ADMIN", "User Management", "ADMIN", "Maintain users, roles, modules, functions, resources, mappings"),
        ],
    )

    res_intro = insert_heading_after(
        fn_tbl, "Resource Master (API / URL) — example rows:", "Normal", parent=rbac_parent
    )
    res_tbl = insert_table_after(
        res_intro,
        ["Resource Code", "Type", "Method", "Path / URL Pattern", "Module Function", "Status"],
        [
            ("RES-DOC-GET", "API", "GET", "/api/v1/documents/{id}", "FN-DOC-VIEW", "Active"),
            ("RES-DOC-LIST", "API", "GET", "/api/v1/documents", "FN-DOC-VIEW", "Active"),
            ("RES-DOC-CREATE", "API", "POST", "/api/v1/documents", "FN-DOC-ADD", "Active"),
            ("RES-DOC-UPDATE", "API", "PUT", "/api/v1/documents/{id}", "FN-DOC-EDIT", "Active"),
            ("RES-DOC-APPROVE", "API", "POST", "/api/v1/documents/{id}/approve", "FN-DOC-APPROVE", "Active"),
            ("RES-DOC-SIGN", "API", "POST", "/api/v1/documents/{id}/sign", "FN-DOC-SIGN", "Active"),
            ("RES-UI-DOC-VIEW", "URL", "GET", "/app/documents/view", "FN-DOC-VIEW", "Active"),
            ("RES-UI-DOC-CREATE", "URL", "GET", "/app/documents/create", "FN-DOC-ADD", "Active"),
            ("RES-MAR-CREATE", "API", "POST", "/api/v1/marriages", "FN-MAR-ADD", "Active"),
            ("RES-MAR-APPROVE", "API", "POST", "/api/v1/marriages/{id}/approve", "FN-MAR-APPROVE", "Active"),
            ("RES-EC-APPLY", "API", "POST", "/api/v1/encumbrance/applications", "FN-EC-APPLY", "Active"),
            ("RES-CC-ISSUE", "API", "POST", "/api/v1/certified-copies/{id}/issue", "FN-CC-ISSUE", "Active"),
            ("RES-UM-ROLES", "API", "PUT", "/api/v1/admin/roles/{id}/module-functions", "FN-UM-ADMIN", "Active"),
        ],
    )

    map_intro = insert_heading_after(
        res_tbl,
        "Role–Module–Function mapping — example rows (Application Admin maintains):",
        "Normal",
        parent=rbac_parent,
    )
    map_tbl = insert_table_after(
        map_intro,
        ["Role", "Role Category", "Module Function", "Allowed"],
        [
            ("Citizen", "Citizen", "FN-DOC-ADD", "Yes"),
            ("Citizen", "Citizen", "FN-DOC-VIEW", "Yes"),
            ("Citizen", "Citizen", "FN-EC-APPLY", "Yes"),
            ("Citizen", "Citizen", "FN-CC-APPLY", "Yes"),
            ("Document Registration Applicant", "Citizen", "FN-DOC-ADD", "Yes"),
            ("Document Registration Applicant", "Citizen", "FN-DOC-VIEW", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-VIEW", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-APPROVE", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-SIGN", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-PRINT", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-MAR-APPROVE", "Yes"),
            ("DEO", "DSR", "FN-DOC-VIEW", "Yes"),
            ("DEO", "DSR", "FN-DOC-ADD", "Yes"),
            ("DEO", "DSR", "FN-DOC-EDIT", "Yes"),
            ("DEO", "DSR", "FN-MAR-ADD", "Yes"),
            ("FDA", "DSR", "FN-DOC-VIEW", "Yes"),
            ("FDA", "DSR", "FN-DOC-ADD", "Yes"),
            ("FDA", "DSR", "FN-DOC-EDIT", "Yes"),
            ("Revenue Verification Officer", "Other Department", "FN-DOC-VIEW", "Yes"),
            ("Application Admin", "DSR", "FN-UM-ADMIN", "Yes"),
        ],
    )

    enf_intro = insert_heading_after(
        map_tbl, "How the application enforces access (runtime):", "Normal", parent=rbac_parent
    )
    enf_tbl = insert_table_after(
        enf_intro,
        ["Step", "Action", "Actor / Component", "Notes"],
        [
            ("1", "User authenticates (OTP / OTP+Biometrics as per category)", "User / Auth service", "Session established"),
            ("2", "Load assigned role(s) for the user from User Master", "Auth / UM service", "Primary + Secondary for DSR; one role for Other Department"),
            ("3", "Resolve Role–Module–Function mappings into session claims (function codes)", "UM / Auth service", "Union of all allowed functions (FR-38)"),
            ("4", "User calls an API or opens a URL", "Client / Browser", "Method + path"),
            ("5", "Look up Resource Master for matching Type + Method + Path pattern", "API Gateway / Middleware", "If no resource found → deny or treat as public only if explicitly marked"),
            ("6", "Obtain required Module Function from the matched Resource", "API Gateway / Middleware", "e.g. POST /api/v1/documents/{id}/approve → FN-DOC-APPROVE"),
            ("7", "Allow if session claims include that Module Function; else HTTP 403", "API Gateway / Middleware", "FR-41; failed attempts audit-logged"),
            ("8", "UI menus / buttons shown only for Module Functions present in session claims", "Front-end", "Same mapping; hide unauthorised screens"),
        ],
    )

    admin_intro = insert_heading_after(
        enf_tbl, "Application Admin maintenance (CRUD):", "Normal", parent=rbac_parent
    )
    admin_crud_tbl = insert_table_after(
        admin_intro,
        ["Master / Table", "Application Admin may", "Audit"],
        [
            ("Module Master", "Add, edit, enable/disable modules", "Mandatory (FR-42)"),
            ("Module Function Master", "Add, edit, enable/disable functions under a module", "Mandatory (FR-42)"),
            ("Resource Master", "Add, edit, enable/disable API and URL resources; link to Module Function", "Mandatory (FR-42)"),
            ("Role–Module–Function mapping", "Grant or revoke Module Functions for any role", "Mandatory (FR-42)"),
            ("DSR Officer Hierarchy Master", "Add, edit, reorder, enable/disable hierarchy nodes (§6.5.7)", "Mandatory (FR-43)"),
        ],
    )

    sub7 = insert_heading_after(
        admin_crud_tbl,
        "6.5.7 DSR Officer Hierarchy Master",
        "Heading 3",
        parent=rbac_parent,
    )
    note7 = insert_paragraph_after(
        sub7,
        "The system shall maintain the reporting hierarchy for Department officers (DSR) as a "
        "Hierarchy Master aligned to the Department of Stamps & Registration organisational chart. "
        "Hierarchy nodes reference roles in the Role Master (Role Category = DSR). Application Admin "
        "shall add, edit, reorder, enable/disable, and update hierarchy nodes. The hierarchy is used "
        "for organisational visibility, span of control, and administrative workflows — it does not "
        "replace sanctioned posts or Role–Module–Function access control.",
        "Normal",
    )
    hier_seed_intro = insert_paragraph_after(
        note7,
        "Seed hierarchy (from DSR organisational chart) — Parent → Child reporting:",
        "Normal",
    )
    hier_seed_tbl = insert_table_after(
        hier_seed_intro,
        ["Level", "Division / Branch", "Role / Post", "Reports To (Parent)"],
        [
            ("0", "Secretariat", "Additional Chief Secretary / Principal Secretary / Secretary", "— (root)"),
            ("1", "Top Management", "Inspector General of Registration & Commissioner of Stamps (IGR)", "Additional Chief Secretary / Principal Secretary / Secretary"),
            ("2", "Admin, Law & Computers", "DIGR (Admin, Law & Computers)", "IGR"),
            ("3", "Admin, Law & Computers — Administration", "AIGR (Admin)", "DIGR (Admin, Law & Computers)"),
            ("4", "Admin, Law & Computers — Administration", "HQA (Admin)", "AIGR (Admin)"),
            ("4", "Admin, Law & Computers — Administration", "Sub Registrar (Admin)", "AIGR (Admin)"),
            ("4", "Admin, Law & Computers — Administration", "Accountant Superintendent (Admin)", "AIGR (Admin)"),
            ("4", "Admin, Law & Computers — Administration", "First Division Assistant (FDA)", "AIGR (Admin)"),
            ("4", "Admin, Law & Computers — Administration", "Second Division Assistant (SDA)", "AIGR (Admin)"),
            ("4", "Admin, Law & Computers — Administration", "Typist", "AIGR (Admin)"),
            ("3", "Admin, Law & Computers — RTI & Statistics", "HQA (RTI)", "DIGR (Admin, Law & Computers)"),
            ("4", "Admin, Law & Computers — RTI & Statistics", "First Division Assistant (FDA)", "HQA (RTI)"),
            ("4", "Admin, Law & Computers — RTI & Statistics", "Second Division Assistant (SDA)", "HQA (RTI)"),
            ("4", "Admin, Law & Computers — RTI & Statistics", "Statistical Inspector", "HQA (RTI)"),
            ("2", "Vigilance", "DIGR (Vigilance)", "IGR"),
            ("3", "Vigilance", "Law Officer", "DIGR (Vigilance)"),
            ("3", "Vigilance", "HQA (RTI)", "DIGR (Vigilance)"),
            ("2", "Computers", "AIGR (Computers)", "IGR"),
            ("3", "Computers — Development", "System Integrator", "AIGR (Computers)"),
            ("3", "Computers — Development", "Application Developer", "AIGR (Computers)"),
            ("3", "Computers — Development", "PMU", "AIGR (Computers)"),
            ("3", "Computers — Operations", "HQA / Project Manager (Comp)", "AIGR (Computers)"),
            ("4", "Computers — Operations", "Sub Registrar (Comp)", "HQA / Project Manager (Comp)"),
            ("4", "Computers — Operations", "First Division Assistant (FDA)", "HQA / Project Manager (Comp)"),
            ("4", "Computers — Operations", "Second Division Assistant (SDA)", "HQA / Project Manager (Comp)"),
            ("2", "Enforcement", "DIGR (Enforcement)", "IGR"),
            ("3", "Enforcement", "District Registrar (DRO)", "DIGR (Enforcement)"),
            ("3", "Enforcement", "HQA", "DIGR (Enforcement)"),
            ("3", "Enforcement", "Sub Registrar (SR)", "DIGR (Enforcement)"),
            ("3", "Enforcement", "First Division Assistant (FDA)", "DIGR (Enforcement)"),
            ("3", "Enforcement", "Second Division Assistant (SDA)", "DIGR (Enforcement)"),
            ("2", "Intelligence & Audit", "DIGR (Intelligence)", "IGR"),
            ("3", "Intelligence & Audit", "AIGR (Audit)", "DIGR (Intelligence)"),
            ("4", "Intelligence & Audit", "HQA (Audit)", "AIGR (Audit)"),
            ("4", "Intelligence & Audit", "Superintendent (Audit)", "AIGR (Audit)"),
            ("4", "Intelligence & Audit", "First Division Assistant (FDA)", "AIGR (Audit)"),
            ("4", "Intelligence & Audit", "Second Division Assistant (SDA)", "AIGR (Audit)"),
            ("4", "Intelligence & Audit", "Typist", "AIGR (Audit)"),
            ("2", "DIGR CVC", "DIGR CVC", "IGR"),
            ("3", "DIGR CVC", "JD Town Planning", "DIGR CVC"),
        ],
    )
    hier_attr_intro = insert_heading_after(
        hier_seed_tbl,
        "Hierarchy Master attributes (Application Admin maintains):",
        "Normal",
        parent=rbac_parent,
    )
    hier_attr_tbl = insert_table_after(
        hier_attr_intro,
        ["Attribute", "Description", "Mandatory"],
        [
            ("Hierarchy Node ID", "System-generated unique identifier", "Yes (system)"),
            ("Role", "DSR role from Role Master (Role Category = DSR)", "Yes"),
            ("Parent Node", "Immediate reporting parent in the hierarchy (null for root)", "No for root"),
            ("Division / Branch", "Division label (e.g. Enforcement, Computers)", "Yes"),
            ("Display Order", "Sort order among siblings under the same parent", "Yes"),
            ("Level", "Depth in the tree (0 = root)", "Yes (system/derived)"),
            ("Is Active", "Enable / disable node without deleting history", "Yes"),
            ("Effective From / To", "Optional validity of the hierarchy link", "No"),
        ],
    )
    hier_admin_intro = insert_heading_after(
        hier_attr_tbl,
        "Maintain hierarchy — Application Admin workflow:",
        "Normal",
        parent=rbac_parent,
    )
    insert_table_after(
        hier_admin_intro,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → DSR Officer Hierarchy", "Application Admin", "FR-43"),
            ("2", "View tree / list of hierarchy nodes by division", "System", "Filter by Division, Active status"),
            ("3", "Add node — select DSR role, parent node, division, display order", "Application Admin", "Role must exist in Role Master"),
            ("4", "Edit node — change parent, order, division, or active flag", "Application Admin", "Cannot create circular parent links"),
            ("5", "Disable node (soft) if post/structure superseded", "Application Admin", "History retained; audit logged"),
            ("6", "Save — hierarchy available for organisational views and admin span rules", "System", "FR-44 seed structure must remain intact unless intentionally changed"),
        ],
    )

    admin_heading = next(
        p for p in doc.paragraphs if p.text.strip() == "6.7 Administrative User Management"
    )
    wf_heading_el = OxmlElement("w:p")
    admin_heading._element.addprevious(wf_heading_el)
    wf_heading = Paragraph(wf_heading_el, admin_heading._parent)
    wf_heading.style = "Heading 2"
    wf_heading.add_run("6.6 User Creation and Role Assignment Workflow")

    wf_intro = insert_paragraph_after(
        wf_heading,
        "The system shall provide dedicated step-by-step workflows to assign roles during user creation. "
        "For DSR Officers, account creation requires one Primary Role (vacant sanctioned post) and may "
        "include additional vacant sanctioned posts. For Other Department users, account creation cannot "
        "be completed without selecting exactly one role.",
        "Normal",
    )

    dsr_sub = insert_heading_after(wf_intro, "6.6.1 DSR Officer User Creation with Role Assignment", "Heading 3")
    dsr_tbl = insert_table_after(
        dsr_sub,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → Add DSR Department User", "Admin", "Authorised admin role only (FR-02)"),
            ("2", "Enter user particulars (name, Username, mobile, email, KGID, photo, ID proof)", "Admin", ""),
            ("3", "Select Primary Role — vacant sanctioned post (role + office); Role Category = DSR", "Admin", "Mandatory; FR-30, FR-26, FR-34"),
            ("4", "Optionally add one or more additional vacant sanctioned posts (Secondary Roles)", "Admin", "Each post must be vacant; end date mandatory; FR-31, FR-45"),
            ("5", "Upload approval letter for additional post(s) where applicable", "Admin", "Should"),
            ("6", "Capture biometrics", "Admin / Officer", "Mandatory for DSR users (FR-06)"),
            ("7", "Review all post occupancies (Primary + additional) and confirm", "Admin", "Only vacant posts listed"),
            ("8", "Save — account active; occupied count updated for each assigned post", "System", "Blocked if Primary missing or any post not vacant"),
        ],
    )

    other_sub = insert_heading_after(
        dsr_tbl, "6.6.2 Other Department User Creation with Role Assignment", "Heading 3", parent=rbac_parent
    )
    other_note = insert_paragraph_after(
        other_sub,
        "Other Department users are stored in the same User Master (User Category = Other Department). "
        "Exactly one role is selected from the Role Master filtered by Role Category = Other Department. "
        "There is no Primary or Secondary role. An optional End Date may be entered; if entered, the "
        "system shall deactivate the user on that date.",
        "Normal",
    )
    insert_table_after(
        other_note,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → Add Other Department User", "Admin", "Authorised admin role only (FR-03)"),
            ("2", "Enter user particulars (name, Username, mobile, email, photo, ID proof)", "Admin", "User Category = Other Department in User Master"),
            ("3", "Enter parent department and designation", "Admin", "e.g. Revenue, Treasury, Police"),
            ("4", "Assign exactly one role from Role Master (Role Category = Other Department)", "Admin", "Mandatory; FR-29, FR-34 — no Primary/Secondary"),
            ("5", "Optionally enter Account End Date", "Admin", "Not mandatory; FR-33"),
            ("6", "Upload authorisation letter / NOC from parent department", "Admin", "Should"),
            ("7", "Capture biometrics", "Admin / User", "Mandatory (FR-07)"),
            ("8", "Review role and End Date (if any) and confirm", "Admin", ""),
            ("9", "Save — account active with module access for the assigned role", "System", "Blocked if no role selected"),
            ("10", "If End Date was entered and is reached — deactivate user; block login", "System", "FR-33; audit logged"),
        ],
    )

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v2.8"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
