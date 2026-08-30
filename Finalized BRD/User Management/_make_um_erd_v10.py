# -*- coding: utf-8 -*-
"""Build ERD_User_Management_v1.0.docx — logical ERD for User Management BRD v4.8.

Standalone deliverable; does not modify BRD_User_Management_v4.8.docx.
"""
from __future__ import annotations

import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(__file__).resolve().parent
PNG_OVERVIEW = BASE / "ERD_User_Management_v1.0_overview.png"
PNG_ORG = BASE / "ERD_User_Management_v1.0_org_posts.png"
PNG_ID = BASE / "ERD_User_Management_v1.0_identity_occupancy.png"
PNG_RBAC = BASE / "ERD_User_Management_v1.0_rbac_runtime.png"
DST = BASE / "ERD_User_Management_v1.0.docx"
PNG = PNG_OVERVIEW  # retained for callers that expect a single figure

NAVY = (31, 78, 121)
NAVY_FILL = (31, 78, 121)
WHITE = (255, 255, 255)
INK = (33, 37, 41)
MUTED = (73, 80, 87)
LINE = (108, 117, 125)
GRID = (233, 236, 239)
BG = (248, 249, 250)

PALETTE = {
    "org": ((13, 110, 253), (232, 242, 255)),
    "est": ((25, 135, 84), (232, 245, 233)),
    "id": ((253, 126, 20), (255, 243, 224)),
    "occ": ((111, 66, 193), (243, 232, 255)),
    "rbac": ((13, 202, 240), (222, 247, 252)),
    "run": ((108, 117, 125), (241, 243, 244)),
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    names = ["segoeui.ttf", "segoeuib.ttf"] if bold else ["segoeui.ttf"]
    windir = Path(r"C:\Windows\Fonts")
    for name in names if bold else ["segoeui.ttf", "calibri.ttf", "arial.ttf"]:
        path = windir / (name if not bold else "segoeuib.ttf" if name.startswith("segoe") else name)
        if bold:
            for candidate in ("segoeuib.ttf", "calibrib.ttf", "arialbd.ttf"):
                p = windir / candidate
                if p.exists():
                    return ImageFont.truetype(str(p), size)
        else:
            for candidate in ("segoeui.ttf", "calibri.ttf", "arial.ttf"):
                p = windir / candidate
                if p.exists():
                    return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


class Box:
    def __init__(self, key: str, title: str, attrs: list[str], x: int, y: int, group: str, w: int = 250):
        self.key = key
        self.title = title
        self.attrs = attrs
        self.x = x
        self.y = y
        self.w = w
        self.group = group
        self.header_h = 28
        self.row_h = 18
        self.h = self.header_h + 8 + len(attrs) * self.row_h + 8

    def cx(self) -> int:
        return self.x + self.w // 2

    def cy(self) -> int:
        return self.y + self.h // 2

    def port(self, side: str) -> tuple[int, int]:
        if side == "n":
            return (self.cx(), self.y)
        if side == "s":
            return (self.cx(), self.y + self.h)
        if side == "w":
            return (self.x, self.cy())
        return (self.x + self.w, self.cy())


def draw_box(draw: ImageDraw.ImageDraw, box: Box, fonts: dict) -> None:
    accent, fill = PALETTE[box.group]
    draw.rounded_rectangle(
        [box.x, box.y, box.x + box.w, box.y + box.h],
        radius=6,
        fill=fill,
        outline=accent,
        width=2,
    )
    draw.rounded_rectangle(
        [box.x, box.y, box.x + box.w, box.y + box.header_h],
        radius=6,
        fill=accent,
        outline=accent,
    )
    draw.rectangle([box.x, box.y + box.header_h - 8, box.x + box.w, box.y + box.header_h], fill=accent)
    draw.text((box.x + 10, box.y + 5), box.title, font=fonts["h"], fill=WHITE)
    yy = box.y + box.header_h + 6
    for attr in box.attrs:
        colour = INK
        if attr.startswith("PK ") or attr.startswith("UK "):
            colour = NAVY
        elif attr.startswith("FK "):
            colour = MUTED
        draw.text((box.x + 10, yy), attr, font=fonts["a"], fill=colour)
        yy += box.row_h


def draw_rel(
    draw: ImageDraw.ImageDraw,
    a: Box,
    sa: str,
    b: Box,
    sb: str,
    left: str,
    right: str,
    label: str,
    fonts: dict,
    via: list[tuple[int, int]] | None = None,
) -> None:
    p1 = a.port(sa)
    p2 = b.port(sb)
    if via:
        pts = [p1] + via + [p2]
    elif sa in ("e", "w") and sb in ("e", "w"):
        mid = (p1[0] + p2[0]) // 2
        pts = [p1, (mid, p1[1]), (mid, p2[1]), p2]
    elif sa in ("n", "s") and sb in ("n", "s"):
        mid = (p1[1] + p2[1]) // 2
        pts = [p1, (p1[0], mid), (p2[0], mid), p2]
    elif sa in ("e", "w"):
        pts = [p1, (p2[0], p1[1]), p2]
    else:
        pts = [p1, (p1[0], p2[1]), p2]
    draw.line(pts, fill=LINE, width=2)
    draw.ellipse([p1[0] - 4, p1[1] - 4, p1[0] + 4, p1[1] + 4], fill=LINE)
    draw.ellipse([p2[0] - 4, p2[1] - 4, p2[0] + 4, p2[1] + 4], fill=LINE)
    lx = p1[0] + (8 if sa == "e" else -28 if sa == "w" else 8)
    ly = p1[1] - 18 if sa in ("e", "w", "n") else p1[1] + 6
    rx = p2[0] + (8 if sb == "e" else -28 if sb == "w" else 8)
    ry = p2[1] - 18 if sb in ("e", "w", "n") else p2[1] + 6
    draw.text((lx, ly), left, font=fonts["c"], fill=NAVY)
    draw.text((rx, ry), right, font=fonts["c"], fill=NAVY)
    if label:
        mx = (p1[0] + p2[0]) // 2
        my = (p1[1] + p2[1]) // 2 - 14
        draw.text((mx + 6, my), label, font=fonts["c"], fill=NAVY)


def draw_legend(draw: ImageDraw.ImageDraw, x: int, y: int, fonts: dict) -> None:
    items = [
        ("org", "Organisation"),
        ("est", "Establishment & mapping"),
        ("id", "Identity"),
        ("occ", "Occupancy & transfer"),
        ("rbac", "Privileges (RBAC)"),
        ("run", "Runtime / audit"),
    ]
    draw.text((x, y), "Subject areas", font=fonts["h"], fill=INK)
    yy = y + 28
    for key, label in items:
        accent, fill = PALETTE[key]
        draw.rounded_rectangle([x, yy, x + 18, yy + 18], radius=3, fill=accent)
        draw.text((x + 26, yy), label, font=fonts["a"], fill=INK)
        yy += 24
    yy += 8
    draw.text((x, yy), "Notation", font=fonts["h"], fill=INK)
    yy += 26
    for line in (
        "Crow’s-foot labels on connectors",
        "1 = exactly one   * = many",
        "0..1 = optional   0..* = optional many",
        "PK / UK / FK shown on the diagram",
        "Logical model — not physical DDL",
    ):
        draw.text((x, yy), line, font=fonts["a"], fill=MUTED)
        yy += 18


def fonts_pack() -> dict:
    return {
        "title": font(26, True),
        "sub": font(15, False),
        "h": font(13, True),
        "a": font(12, False),
        "c": font(12, True),
        "band": font(14, True),
    }


def canvas(title: str, subtitle: str, w: int, h: int) -> tuple:
    img = Image.new("RGB", (w, h), BG)
    draw = ImageDraw.Draw(img)
    fonts = fonts_pack()
    draw.rectangle([0, 0, w, 72], fill=NAVY_FILL)
    draw.text((24, 12), title, font=fonts["title"], fill=WHITE)
    draw.text((24, 44), subtitle, font=fonts["sub"], fill=(196, 216, 236))
    return img, draw, fonts


def finish(img: Image.Image, draw: ImageDraw.ImageDraw, fonts: dict, path: Path) -> None:
    draw.text(
        (24, img.height - 32),
        "KAVERI 3.0  ·  DSR  ·  Logical ERD v1.0  ·  Companion to BRD User Management v4.8  ·  Not part of the BRD",
        font=fonts["sub"],
        fill=MUTED,
    )
    img.save(path, "PNG", optimize=True)
    print(f"Wrote {path}")


def render_overview() -> None:
    img, draw, fonts = canvas(
        "Figure 1 — Overview  ·  User Management logical model",
        "Single User Master and Role Master  ·  DSR privileges flow occupancy → Post–Role → Module Function  ·  BRD v4.8",
        2400,
        1280,
    )
    boxes: dict[str, Box] = {}

    def add(b: Box) -> Box:
        boxes[b.key] = b
        return b

    def compact(title: str, group: str, x: int, y: int) -> Box:
        width = 250 if len(title) > 20 else 210
        return add(Box(title, title, [], x, y, group, width))

    compact("DIVISION_MASTER", "org", 80, 160)
    compact("OFFICE_TYPE", "org", 380, 160)
    compact("OFFICE_HIERARCHY", "org", 680, 160)
    compact("POSTS_MASTER", "est", 980, 160)
    compact("OFFICER_HIERARCHY", "est", 1280, 160)
    compact("POST_OFFICE_TYPE_ALLOWED", "est", 1580, 160)

    compact("SANCTIONED_POST", "est", 680, 430)
    compact("POST_ROLE_MAP", "est", 980, 430)
    compact("ROLE_MASTER", "rbac", 1280, 430)
    compact("ROLE_MODULE_FUNCTION", "rbac", 1580, 430)

    compact("USER_MASTER", "id", 80, 720)
    compact("USER_ROLE", "id", 380, 720)
    compact("USER_SECURITY_ANSWER", "id", 680, 720)
    compact("POST_OCCUPANCY", "occ", 980, 720)
    compact("USER_SESSION", "run", 1280, 720)
    compact("OTP_CHALLENGE", "run", 1580, 720)

    compact("MODULE_MASTER", "rbac", 980, 980)
    compact("MODULE_FUNCTION", "rbac", 1280, 980)
    compact("RESOURCE_MASTER", "rbac", 1580, 980)
    compact("AUDIT_LOG", "run", 80, 980)

    for b in boxes.values():
        b.h = 48
        b.header_h = 48
        draw_box(draw, b, fonts)

    r = fonts
    b = boxes
    draw_rel(draw, b["DIVISION_MASTER"], "e", b["POSTS_MASTER"], "w", "1", "*", "", r, via=[(320, 184), (960, 184)])
    draw_rel(draw, b["OFFICE_TYPE"], "e", b["OFFICE_HIERARCHY"], "w", "1", "*", "", r)
    draw_rel(draw, b["OFFICE_HIERARCHY"], "s", b["SANCTIONED_POST"], "n", "1", "*", "", r)
    draw_rel(draw, b["POSTS_MASTER"], "s", b["SANCTIONED_POST"], "e", "1", "*", "", r)
    draw_rel(draw, b["POSTS_MASTER"], "s", b["POST_ROLE_MAP"], "n", "1", "1..*", "", r)
    draw_rel(draw, b["POSTS_MASTER"], "e", b["OFFICER_HIERARCHY"], "w", "1", "1", "", r)
    draw_rel(draw, b["POSTS_MASTER"], "e", b["POST_OFFICE_TYPE_ALLOWED"], "w", "1", "*", "", r, via=[(1200, 140), (1570, 140)])
    draw_rel(draw, b["OFFICE_TYPE"], "n", b["POST_OFFICE_TYPE_ALLOWED"], "n", "1", "*", "FR-78", r, via=[(485, 120), (1685, 120)])
    draw_rel(draw, b["ROLE_MASTER"], "w", b["POST_ROLE_MAP"], "e", "1", "*", "", r)
    draw_rel(draw, b["ROLE_MASTER"], "e", b["ROLE_MODULE_FUNCTION"], "w", "1", "*", "", r)
    draw_rel(draw, b["USER_MASTER"], "e", b["USER_ROLE"], "w", "1", "0..*", "Citizen/OD", r)
    draw_rel(draw, b["ROLE_MASTER"], "s", b["USER_ROLE"], "n", "1", "*", "", r)
    draw_rel(draw, b["USER_MASTER"], "e", b["USER_SECURITY_ANSWER"], "w", "1", "0..5", "Citizen", r, via=[(350, 860), (670, 860)])
    draw_rel(draw, b["USER_MASTER"], "e", b["POST_OCCUPANCY"], "w", "1", "0..*", "DSR", r, via=[(350, 880), (970, 880)])
    draw_rel(draw, b["SANCTIONED_POST"], "s", b["POST_OCCUPANCY"], "n", "1", "*", "", r)
    draw_rel(draw, b["USER_MASTER"], "s", b["USER_SESSION"], "w", "1", "0..1", "FR-76", r, via=[(185, 900), (1270, 900)])
    draw_rel(draw, b["POST_OCCUPANCY"], "e", b["USER_SESSION"], "w", "1", "*", "assigned", r)
    draw_rel(draw, b["USER_MASTER"], "s", b["OTP_CHALLENGE"], "w", "1", "*", "", r, via=[(185, 920), (1570, 920)])
    draw_rel(draw, b["MODULE_MASTER"], "e", b["MODULE_FUNCTION"], "w", "1", "*", "", r)
    draw_rel(draw, b["MODULE_FUNCTION"], "e", b["RESOURCE_MASTER"], "w", "1", "*", "", r)
    draw_rel(draw, b["MODULE_FUNCTION"], "n", b["ROLE_MODULE_FUNCTION"], "s", "1", "*", "", r)
    draw_rel(draw, b["USER_MASTER"], "s", b["AUDIT_LOG"], "n", "1", "*", "actor", r)

    draw_legend(draw, 360, 980, fonts)
    finish(img, draw, fonts, PNG_OVERVIEW)


def render_org() -> None:
    img, draw, fonts = canvas(
        "Figure 2 — Organisation, posts and sanctioned strength",
        "Office tree is independent of officer (post) tree  ·  Occupied includes reserved Transfer In (FR-67)  ·  FR-78 validates Post + Office Type",
        2400,
        1180,
    )
    boxes: dict[str, Box] = {}

    def add(b: Box) -> Box:
        boxes[b.key] = b
        return b

    add(Box("DIV", "DIVISION_MASTER", ["PK division_code", "UK division_name", "display_order", "is_active"], 60, 120, "org", 250))
    add(Box("OT", "OFFICE_TYPE", ["PK office_type", "Secretariat", "Head Office", "District Registrar Office", "Sub-Registrar Office"], 380, 120, "org", 260))
    add(Box("OFF", "OFFICE_HIERARCHY", ["PK office_code", "office_name", "FK office_type", "FK parent_office_code", "is_active"], 720, 120, "org", 270))
    add(Box("POST", "POSTS_MASTER", ["PK post_code", "UK post_name", "FK division_code", "is_active"], 1080, 120, "est", 250))
    add(Box("HN", "OFFICER_HIERARCHY_NODE", ["PK node_id", "FK post_code", "FK parent_node_id", "display_order", "is_active"], 1420, 120, "est", 280))
    add(Box("POTA", "POST_OFFICE_TYPE_ALLOWED", ["PK (post_code, office_type)", "FK post_code", "FK office_type", "FR-78 seed in BRD §6.5.3"], 1780, 120, "est", 280))
    add(Box("SP", "SANCTIONED_POST", ["PK (post_code, office_code)", "FK post_code", "FK office_code", "sanctioned_strength", "occupied_count (FR-68 job)", "remaining_capacity (derived)"], 720, 520, "est", 300))
    add(Box("PR", "POST_ROLE_MAP", ["PK (post_code, role_id)", "FK post_code", "FK role_id", "Unmapped post blocked (FR-47)"], 1120, 540, "est", 300))
    add(Box("ROLE", "ROLE_MASTER", ["PK role_id", "UK role_name", "role_category C|DSR|OD", "FK division_code (optional)", "is_active"], 1520, 520, "rbac", 280))

    for b in boxes.values():
        draw_box(draw, b, fonts)
    r = fonts
    draw_rel(draw, boxes["DIV"], "e", boxes["POST"], "w", "1", "*", "FR-77", r, via=[(400, 400), (1070, 400)])
    draw_rel(draw, boxes["OT"], "e", boxes["OFF"], "w", "1", "*", "", r)
    draw_rel(draw, boxes["OFF"], "s", boxes["SP"], "n", "1", "*", "per office", r)
    draw_rel(draw, boxes["POST"], "s", boxes["SP"], "n", "1", "*", "per post", r)
    draw_rel(draw, boxes["POST"], "s", boxes["PR"], "n", "1", "1..*", "", r)
    draw_rel(draw, boxes["POST"], "e", boxes["HN"], "w", "1", "1", "", r)
    draw_rel(draw, boxes["POST"], "e", boxes["POTA"], "w", "1", "*", "", r, via=[(1360, 90), (1770, 90)])
    draw_rel(draw, boxes["OT"], "n", boxes["POTA"], "n", "1", "*", "", r, via=[(510, 80), (1920, 80)])
    draw_rel(draw, boxes["ROLE"], "w", boxes["PR"], "e", "1", "*", "FR-50", r)
    draw_rel(draw, boxes["DIV"], "s", boxes["ROLE"], "w", "1", "0..*", "DSR grouping", r, via=[(185, 900), (1510, 900)])
    finish(img, draw, fonts, PNG_ORG)


def render_identity() -> None:
    img, draw, fonts = canvas(
        "Figure 3 — Identity, occupancy, session and OTP",
        "Username unique (Citizen preferred / KGID)  ·  Email and mobile not unique  ·  Additional charge is session-only (FR-53)  ·  No password entity",
        2400,
        1280,
    )
    boxes: dict[str, Box] = {}

    def add(b: Box) -> Box:
        boxes[b.key] = b
        return b

    add(Box("USER", "USER_MASTER", [
        "PK user_id",
        "UK username (Citizen pref. / KGID)",
        "user_category",
        "email (not unique)",
        "mobile (not unique)",
        "parent_department (OD)",
        "account_end_date (OD FR-33)",
        "status / lockout_until",
        "biometric_ref (DSR / OD)",
    ], 80, 140, "id", 320))
    add(Box("UR", "USER_ROLE", [
        "PK (user_id, role_id)",
        "FK user_id",
        "FK role_id",
        "Citizen: 1..* Citizen roles",
        "Other Dept: exactly 1",
        "DSR: none (via Post–Role)",
    ], 500, 160, "id", 300))
    add(Box("SQC", "SECURITY_QUESTION", ["PK question_id", "question_text", "is_active", "Predefined catalogue"], 900, 160, "id", 260))
    add(Box("USA", "USER_SECURITY_ANSWER", [
        "PK (user_id, question_id)",
        "FK user_id (Citizen only)",
        "FK question_id",
        "answer_hash",
        "Exactly 5 per Citizen (FR-55)",
    ], 1260, 150, "id", 300))
    add(Box("DOM", "OFFICIAL_EMAIL_DOMAIN", ["PK domain", "is_active", "FR-64 allow-list"], 1660, 170, "id", 260))
    add(Box("OCC", "POST_OCCUPANCY", [
        "PK occupancy_id",
        "FK user_id (DSR only)",
        "FK post_code + office_code",
        "status Active | Reserved | Ended",
        "joining_date (login from 00:00 IST)",
        "relieving_date / relieving_order",
        "transfer_order + document",
        "end_date + deputation_reason",
        "reserved_flag (FR-67)",
    ], 80, 560, "occ", 340))
    add(Box("SES", "USER_SESSION", [
        "PK session_id",
        "UK user_id (one active FR-76)",
        "FK assigned_occupancy_id",
        "add_charge_post + office (opt.)",
        "login_at / last_activity_at",
        "idle 15 min · absolute 8 h",
    ], 520, 600, "run", 320))
    add(Box("OTP", "OTP_CHALLENGE", [
        "PK challenge_id",
        "FK user_id (or pending FR-63)",
        "purpose LOGIN|REG|PIN|NEW_MOB",
        "channel SMS | EMAIL",
        "code_hash · expires_at",
        "attempts · resend_count",
    ], 960, 600, "run", 300))
    add(Box("AUD", "AUDIT_LOG", [
        "PK audit_id (append-only)",
        "actor_id / actor_type",
        "action · entity · entity_id",
        "before_json / after_json",
        "reason · artefact_id",
        "occurred_at (retain 7 years)",
    ], 1380, 600, "run", 300))

    for b in boxes.values():
        draw_box(draw, b, fonts)
    r = fonts
    draw_rel(draw, boxes["USER"], "e", boxes["UR"], "w", "1", "0..*", "", r)
    draw_rel(draw, boxes["USER"], "e", boxes["USA"], "w", "1", "0 or 5", "Citizen", r, via=[(460, 500), (1250, 500)])
    draw_rel(draw, boxes["SQC"], "e", boxes["USA"], "w", "1", "*", "", r)
    draw_rel(draw, boxes["USER"], "s", boxes["OCC"], "n", "1", "0..*", "DSR", r)
    draw_rel(draw, boxes["USER"], "s", boxes["SES"], "n", "1", "0..1", "", r)
    draw_rel(draw, boxes["OCC"], "e", boxes["SES"], "w", "1", "*", "assigned post", r)
    draw_rel(draw, boxes["USER"], "s", boxes["OTP"], "n", "1", "*", "", r, via=[(240, 1100), (1110, 1100)])
    draw_rel(draw, boxes["USER"], "s", boxes["AUD"], "n", "1", "*", "actor", r, via=[(240, 1140), (1530, 1140)])
    finish(img, draw, fonts, PNG_ID)


def render_rbac() -> None:
    img, draw, fonts = canvas(
        "Figure 4 — Privileges (RBAC) and Application Admin",
        "Runtime: request → Resource (type + method + path) → Module Function → Role claims  ·  Deny-by-default unless Is Public  ·  FR-41 / FR-51",
        2400,
        980,
    )
    boxes: dict[str, Box] = {}

    def add(b: Box) -> Box:
        boxes[b.key] = b
        return b

    add(Box("ROLE", "ROLE_MASTER", ["PK role_id", "UK role_name", "role_category", "is_active"], 80, 160, "rbac", 260))
    add(Box("RMF", "ROLE_MODULE_FUNCTION", ["PK (role_id, function_id)", "FK role_id", "FK function_id", "Exact role name (FR-50)"], 460, 160, "rbac", 300))
    add(Box("MF", "MODULE_FUNCTION", ["PK function_id", "FK module_code", "function_code VIEW|ADD|…", "UK (module, function)"], 880, 150, "rbac", 300))
    add(Box("MOD", "MODULE_MASTER", ["PK module_code", "module_name", "is_active"], 1320, 170, "rbac", 260))
    add(Box("RES", "RESOURCE_MASTER", ["PK resource_id", "FK function_id", "type API | URL", "http_method · path_pattern", "is_public (deny-by-default)"], 880, 500, "rbac", 320))
    add(Box("AA", "APPLICATION_ADMIN", [
        "Seeded system principal",
        "Not a Role Master seed role",
        "Not User Master workflow",
        "FN-UM-ADMIN outside RMF",
        "FR-51",
    ], 1400, 470, "run", 300))

    for b in boxes.values():
        draw_box(draw, b, fonts)
    r = fonts
    draw_rel(draw, boxes["ROLE"], "e", boxes["RMF"], "w", "1", "*", "", r)
    draw_rel(draw, boxes["RMF"], "e", boxes["MF"], "w", "*", "1", "", r)
    draw_rel(draw, boxes["MOD"], "w", boxes["MF"], "e", "1", "*", "", r)
    draw_rel(draw, boxes["MF"], "s", boxes["RES"], "n", "1", "*", "", r)
    finish(img, draw, fonts, PNG_RBAC)


def render_png() -> None:
    render_overview()
    render_org()
    render_identity()
    render_rbac()


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text: str, bold: bool = False, size: Pt = Pt(9)) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = "Calibri"


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i], "1F4E79")
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            set_cell(table.rows[r].cells[c], val)
            if r % 2 == 0:
                shade_cell(table.rows[r].cells[c], "F2F2F2")
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(8)
    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(13.0)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    configure_styles(doc)

    title = doc.add_paragraph()
    run = title.add_run("Logical Entity-Relationship Diagram")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    r = sub.add_run("KAVERI 3.0 — User Management & Role-Based Access Control")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_heading("1. Document control", level=1)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ("Document ID", "ERD-K3-UM-001"),
            ("Version", "1.0"),
            ("Status", "Draft / companion to BRD v4.8"),
            ("Project", "KAVERI 3.0 — Department of Stamps and Registration"),
            ("Module", "User Management & RBAC"),
            ("Source BRD", "Finalized BRD/User Management/BRD_User_Management_v4.8.docx"),
            ("Relationship to BRD", "Separate companion document — not a section of the BRD"),
            ("Author (BA)", "Nandha Kumar"),
            ("Product Owner", "Prashanth"),
            ("Domain expert", "Prabhakar Naik"),
            ("Last updated", "2026-08-30"),
        ],
    )
    add_table(
        doc,
        ["Version", "Date", "Summary"],
        [
            ("1.0", "2026-08-30", "Initial logical ERD derived from User Management BRD v4.8 masters, FRs, and occupancy / RBAC rules."),
        ],
    )

    doc.add_heading("2. Purpose and scope", level=1)
    doc.add_paragraph(
        "This document is the logical data model for the User Management module. It is issued as a "
        "stand-alone engineering / BA companion so that the BRD remains a requirements document. "
        "Physical table names, indexes, and partitioning are out of scope; implementers may split or "
        "merge tables provided the business keys, cardinalities, and rules below are preserved."
    )
    doc.add_paragraph("In scope: masters, mappings, identity, occupancy / transfer, session / OTP, RBAC resources, and audit.")
    doc.add_paragraph(
        "Out of scope: citizen service case data (Marriage, Document Registration, etc.), Khajane-II "
        "DDO tables owned by Treasury, and Application Admin as a Role Master row (FR-51 — seeded system principal)."
    )

    doc.add_heading("3. Logical ERD", level=1)
    doc.add_paragraph(
        "Four figures follow. Colour bands group subject areas. Occupied count on SANCTIONED_POST is "
        "maintained by the occupancy refresh job (FR-68) and is shown as a stored derived attribute, "
        "not a separate entity. Full attributes and cardinalities are in sections 5 and 6."
    )

    def add_figure(path: Path, caption: str) -> None:
        if path.exists():
            doc.add_picture(str(path), width=Inches(11.5))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:
            doc.add_paragraph(f"[Missing figure: {path.name}]")

    add_figure(PNG_OVERVIEW, "Figure 1. Overview — all logical entities and cardinalities (BRD v4.8)")
    add_figure(PNG_ORG, "Figure 2. Organisation, posts, sanctioned strength and Post–Role mapping")
    add_figure(PNG_ID, "Figure 3. Identity, occupancy, session, OTP and audit")
    add_figure(PNG_RBAC, "Figure 4. RBAC path Role → Module Function → Resource; Application Admin (FR-51)")

    doc.add_heading("4. Subject-area summary", level=1)
    add_table(
        doc,
        ["Area", "Entities", "BRD anchors"],
        [
            (
                "Organisation",
                "DIVISION_MASTER, OFFICE_TYPE, OFFICE_HIERARCHY",
                "FR-59, FR-77; §6.5.1, §6.5.8",
            ),
            (
                "Establishment",
                "POSTS_MASTER, OFFICER_HIERARCHY_NODE, POST_OFFICE_TYPE_ALLOWED, SANCTIONED_POST, POST_ROLE_MAP",
                "FR-24, FR-43–FR-49, FR-66–FR-68, FR-78; §6.5.3, §6.5.7",
            ),
            (
                "Identity",
                "USER_MASTER, USER_ROLE, SECURITY_QUESTION, USER_SECURITY_ANSWER, OFFICIAL_EMAIL_DOMAIN",
                "FR-01–FR-04, FR-16, FR-33–FR-35, FR-55, FR-62–FR-65",
            ),
            (
                "Occupancy & transfer",
                "POST_OCCUPANCY (Transfer In / Relieving / deputation End Date as attributes)",
                "FR-26, FR-30, FR-52, FR-57–FR-61, FR-66–FR-68",
            ),
            (
                "RBAC",
                "ROLE_MASTER, MODULE_MASTER, MODULE_FUNCTION, RESOURCE_MASTER, ROLE_MODULE_FUNCTION",
                "FR-18, FR-36–FR-42, FR-50–FR-51",
            ),
            (
                "Runtime / audit",
                "USER_SESSION, OTP_CHALLENGE, AUDIT_LOG; APPLICATION_ADMIN (not a Role Master entity)",
                "FR-05–FR-13, FR-22, FR-51, FR-69–FR-76",
            ),
        ],
    )

    doc.add_heading("5. Entity catalogue", level=1)

    doc.add_heading("5.1 DIVISION_MASTER", level=2)
    doc.add_paragraph("Enumerated DSR organisational divisions. Free-text division labels are not permitted on new or edited records (FR-77).")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("division_code", "String, e.g. DIV-ENFORCEMENT", "PK"),
            ("division_name", "String", "UK; required"),
            ("display_order", "Integer", "Seed 1–8"),
            ("is_active", "Boolean", "Soft disable"),
        ],
    )

    doc.add_heading("5.2 OFFICE_TYPE", level=2)
    doc.add_paragraph("Lookup of Secretariat, Head Office, District Registrar Office, Sub-Registrar Office. Distinct from role names (SRO is an office type, not a Role Master name).")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("office_type", "Enum / code", "PK"),
            ("display_name", "String", "Required"),
        ],
    )

    doc.add_heading("5.3 OFFICE_HIERARCHY", level=2)
    doc.add_paragraph(
        "Concrete DSR offices. Tree: MS Building (Secretariat, root, OFF-MS-BLDG) → IGR Office (Head Office) "
        "→ District Registrar Offices → Sub-Registrar Offices (FR-59). Office span for Transfer Out / Transfer In "
        "is the session office plus descendants."
    )
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("office_code", "String, e.g. OFF-SRO-YESH", "PK"),
            ("office_name", "String", "Required"),
            ("office_type", "FK OFFICE_TYPE", "Required"),
            ("parent_office_code", "FK OFFICE_HIERARCHY", "Null only for MS Building"),
            ("is_active", "Boolean", "Soft disable with audit"),
        ],
    )

    doc.add_heading("5.4 POSTS_MASTER", level=2)
    doc.add_paragraph("DSR establishment posts, distinct from Role Master. Division-specific posts (FR-49). Post name is the referential-integrity key used by Post–Role and Hierarchy (FR-50).")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("post_code", "String, e.g. POST-SR", "PK"),
            ("post_name", "String, character-for-character match", "UK; required"),
            ("division_code", "FK DIVISION_MASTER", "Required (FR-77)"),
            ("is_active", "Boolean", "Soft disable"),
        ],
    )

    doc.add_heading("5.5 OFFICER_HIERARCHY_NODE", level=2)
    doc.add_paragraph("Reporting tree of Posts (not of people). Each node references a Post; optional parent is also a Post node (FR-43, FR-44). Used for FR-53 additional-charge cascade and Transfer Out / In immediate-parent filter.")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("node_id", "Surrogate", "PK"),
            ("post_code", "FK POSTS_MASTER", "UK recommended (one node per post)"),
            ("parent_node_id", "FK self", "Null at ACS / Secretary"),
            ("display_order", "Integer", "Optional"),
            ("is_active", "Boolean", "Soft disable; history retained"),
        ],
    )

    doc.add_heading("5.6 POST_OFFICE_TYPE_ALLOWED", level=2)
    doc.add_paragraph("Which Post codes may be sanctioned at which Office Types (FR-78). Example: POST-ACS-SEC only at Secretariat; POST-SR / POST-FDA-ENF / POST-DEO only at Sub-Registrar Office.")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("post_code", "FK POSTS_MASTER", "PK part"),
            ("office_type", "FK OFFICE_TYPE", "PK part"),
        ],
    )

    doc.add_heading("5.7 SANCTIONED_POST", level=2)
    doc.add_paragraph(
        "Approved headcount per Post at each Office (FR-24, FR-48). Occupied count includes active occupancies "
        "and reserved Transfer In (FR-67) and is recalculated by the midnight occupancy refresh job (FR-68). "
        "FR-66(a) available capacity = Occupied < Strength. FR-66(b) wholly unoccupied = Occupied = 0 (FR-53)."
    )
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("post_code", "FK POSTS_MASTER", "PK part; must have Post–Role mapping (FR-47)"),
            ("office_code", "FK OFFICE_HIERARCHY", "PK part"),
            ("sanctioned_strength", "Integer ≥ 0", "Required"),
            ("occupied_count", "Integer, job-maintained", "Includes reserved FR-67"),
            ("remaining_capacity", "Derived Strength − Occupied", "May be stored or computed"),
        ],
    )

    doc.add_heading("5.8 ROLE_MASTER", level=2)
    doc.add_paragraph("Single role catalogue for Citizen, DSR, and Other Department (FR-16, FR-35). Unique role names include division context (FDA (Enforcement), not bare FDA). Application Admin is not a seed row (FR-51).")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("role_id", "Surrogate", "PK"),
            ("role_name", "String", "UK; FR-50 exact match"),
            ("role_category", "Citizen | DSR | Other Department", "Required"),
            ("division_code", "FK DIVISION_MASTER", "Optional; DSR seed grouping"),
            ("description", "String", "Optional"),
            ("is_active", "Boolean", "Disable blocked if still in use"),
        ],
    )

    doc.add_heading("5.9 POST_ROLE_MAP", level=2)
    doc.add_paragraph("One Post may map to multiple Roles only within the same functional scope (FR-47). Sanction or assignment of an unmapped Post is blocked. DSR session privileges come from this mapping for the selected occupancy — not from USER_ROLE.")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("post_code", "FK POSTS_MASTER", "PK part"),
            ("role_id", "FK ROLE_MASTER", "PK part; name must match Role Master"),
        ],
    )

    doc.add_heading("5.10 USER_MASTER", level=2)
    doc.add_paragraph(
        "Single user store. Username is the only unique login identifier (FR-04, FR-62): preferred Username for "
        "Citizens; KGID for DSR Officers and Other Department users. Email and mobile are not unique. "
        "No password attribute exists (FR-09)."
    )
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("user_id", "Surrogate", "PK"),
            ("username", "Citizen preferred / departmental KGID", "UK across all categories"),
            ("user_category", "Public (Citizen) | DSR Officer | Other Department", "Required"),
            ("name_parts", "First / middle / last", "Required; name locked on standard edit"),
            ("email", "Citizen personal / departmental official", "Not unique; FR-64 domain check for DSR/OD"),
            ("mobile", "SMS login OTP target", "Not unique"),
            ("parent_department", "String", "Other Department only"),
            ("designation", "String", "Other Department only"),
            ("account_end_date", "Date", "Optional; OD auto-deactivate (FR-33)"),
            ("status", "Active | Suspended | Deactivated", "Deactivation requires reason"),
            ("lockout_until", "Timestamp", "FR-73 login lock 15 min"),
            ("reset_lockout_until", "Timestamp", "FR-56 / FR-73 reset lock 30 min"),
            ("biometric_ref", "UIDAI-compliant token", "DSR and Other Department; every login"),
            ("photo / id_proof_ref", "Document store keys", "Departmental creation"),
            ("authorisation_letter_ref", "Document store", "Other Department should"),
        ],
    )

    doc.add_heading("5.11 USER_ROLE", level=2)
    doc.add_paragraph(
        "Direct role assignment for Citizens (one or more Citizen-category roles) and Other Department users "
        "(exactly one Other Department role — FR-17 / FR-34). DSR Officers do not receive USER_ROLE rows; "
        "their roles are resolved Post occupancy → POST_ROLE_MAP at login (FR-52, FR-38)."
    )
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("user_id", "FK USER_MASTER", "PK part"),
            ("role_id", "FK ROLE_MASTER", "PK part; Role Category must match User Category"),
        ],
    )

    doc.add_heading("5.12 SECURITY_QUESTION and USER_SECURITY_ANSWER", level=2)
    doc.add_paragraph("Citizen self-registration captures five distinct questions from a predefined list; answers stored hashed only (FR-55). Not held for DSR or Other Department users.")
    add_table(
        doc,
        ["Entity", "Attribute", "Constraints"],
        [
            ("SECURITY_QUESTION", "question_id, question_text, is_active", "Predefined catalogue"),
            ("USER_SECURITY_ANSWER", "user_id, question_id, answer_hash", "Exactly 5 per Citizen; never displayed"),
        ],
    )

    doc.add_heading("5.13 OFFICIAL_EMAIL_DOMAIN", level=2)
    doc.add_paragraph("Optional allow-list for DSR and Other Department official email (FR-64).")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("domain", "e.g. karnataka.gov.in", "PK"),
            ("is_active", "Boolean", ""),
        ],
    )

    doc.add_heading("5.14 POST_OCCUPANCY", level=2)
    doc.add_paragraph(
        "Assignment of a DSR user to a sanctioned Post at an Office. Multiple occupancies per user are allowed. "
        "Transfer In, Relieving, reserved handover (FR-67), and deputation End Date (FR-30) are modelled as "
        "attributes and status on this entity rather than separate header tables. Login uses occupancies that "
        "are Active (Joining Date reached); Reserved occupancies count toward Occupied but cannot be selected "
        "until 12:00 AM IST on Joining Date (FR-61)."
    )
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("occupancy_id", "Surrogate", "PK"),
            ("user_id", "FK USER_MASTER", "User Category = DSR Officer"),
            ("post_code + office_code", "FK SANCTIONED_POST", "Must have available capacity unless FR-67"),
            ("status", "Active | Reserved | Ended", "Job transitions at midnight IST"),
            ("joining_date", "Date", "Transfer In; login from 00:00 IST that day"),
            ("transfer_order_no / document_id", "String / artefact", "FR-60"),
            ("relieving_date", "Date", "Mapping retained through 23:59 IST of that date"),
            ("relieving_order_no / document_id", "String / artefact", "FR-57"),
            ("end_date", "Date", "Deputation only (FR-30); not a substitute for relieving"),
            ("deputation_reason", "Code", "Mandatory if end_date set"),
            ("reserved_flag", "Boolean", "FR-67; counts as occupied immediately"),
            ("created_by / created_at", "Actor + timestamp", "Immediate-parent + office-span rules"),
        ],
    )

    doc.add_heading("5.15 USER_SESSION", level=2)
    doc.add_paragraph(
        "One active session per Username (FR-76). Idle 15 minutes (FR-74); absolute 8 hours (FR-75). "
        "For DSR: assigned_occupancy_id is the FR-52 post for the session. Additional charge (FR-53) is "
        "session-only — it does not create a POST_OCCUPANCY row and does not require an order. While "
        "additional charge is set, Module Function claims come only from that subordinate post (FR-38)."
    )
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("session_id", "Surrogate", "PK"),
            ("user_id", "FK USER_MASTER", "UK among active sessions"),
            ("assigned_occupancy_id", "FK POST_OCCUPANCY", "DSR only; null for Citizen / OD"),
            ("add_charge_post_code", "FK POSTS_MASTER", "Optional; same office as assigned"),
            ("add_charge_office_code", "FK OFFICE_HIERARCHY", "Must equal assigned occupancy office"),
            ("login_at / last_activity_at / expires_at", "Timestamps", "IST clock"),
        ],
    )

    doc.add_heading("5.16 OTP_CHALLENGE", level=2)
    doc.add_paragraph("OTP/PIN issuance for login (SMS only), Citizen registration (email + mobile), lost-mobile PIN (email), and new-mobile verification. Policy FR-69–FR-72: 6 digits; login 5 min; registration/PIN 10 min; 3 failures invalidate; 30 s resend; max 3 resends / 15 min.")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("challenge_id", "Surrogate", "PK"),
            ("user_id", "FK USER_MASTER (or pending registration token)", "Nullable during FR-63"),
            ("purpose", "LOGIN | REG_EMAIL | REG_MOBILE | RESET_PIN | NEW_MOBILE", "Required"),
            ("channel", "SMS | EMAIL", "Login OTP never emailed (FR-10)"),
            ("code_hash", "Hash", "Never stored clear"),
            ("expires_at / attempt_count / resend_count", "Policy attributes", "FR-69–FR-72"),
        ],
    )

    doc.add_heading("5.17 MODULE_MASTER, MODULE_FUNCTION, RESOURCE_MASTER, ROLE_MODULE_FUNCTION", level=2)
    doc.add_paragraph(
        "Access path: session roles → ROLE_MODULE_FUNCTION → MODULE_FUNCTION → RESOURCE. Runtime match is "
        "Type + Method + Path with most-specific pattern winning; unmatched non-public resources deny (FR-41). "
        "Is Public = Yes allows unauthenticated access. Application Admin uses FN-UM-ADMIN outside this mapping (FR-51)."
    )
    add_table(
        doc,
        ["Entity", "Key attributes", "Notes"],
        [
            ("MODULE_MASTER", "module_code, module_name, is_active", "e.g. Marriage Registration, Encumbrance Search"),
            ("MODULE_FUNCTION", "function_id, module_code, function_code", "VIEW, ADD, EDIT, APPROVE, SIGN, PRINT, DOWNLOAD, FN-UM-ADMIN"),
            ("RESOURCE_MASTER", "resource_id, function_id, type, method, path_pattern, is_public", "Roles do not store raw URL lists"),
            ("ROLE_MODULE_FUNCTION", "role_id, function_id", "Role name must match Role Master exactly (FR-50)"),
        ],
    )

    doc.add_heading("5.18 AUDIT_LOG", level=2)
    doc.add_paragraph("Append-only log of create/update/delete, login, mobile/email change, lost-mobile reset, Transfer Out / In, additional charge switch, occupancy-refresh job, and access denials. Retention at least seven years. Who, when, before/after, reason, artefact id.")
    add_table(
        doc,
        ["Attribute", "Type / notes", "Constraints"],
        [
            ("audit_id", "Surrogate", "PK; no updates/deletes"),
            ("actor_id / actor_type", "User or job / Application Admin", "Required"),
            ("action / entity / entity_id", "String", "Required"),
            ("before_json / after_json", "JSON", "Masked PII (mobile)"),
            ("reason / artefact_id", "String", "GO / order document"),
            ("occurred_at", "Timestamp IST", "Required"),
        ],
    )

    doc.add_heading("6. Relationship catalogue", level=1)
    add_table(
        doc,
        ["Parent", "Child", "Cardinality", "Rule"],
        [
            ("DIVISION_MASTER", "POSTS_MASTER", "1 : *", "Every post has a Division Code (FR-77)."),
            ("DIVISION_MASTER", "ROLE_MASTER", "1 : 0..*", "Optional grouping for DSR seed roles."),
            ("OFFICE_TYPE", "OFFICE_HIERARCHY", "1 : *", "Four office types only."),
            ("OFFICE_HIERARCHY", "OFFICE_HIERARCHY", "1 : 0..*", "Parent office; MS Building is root."),
            ("POSTS_MASTER", "OFFICER_HIERARCHY_NODE", "1 : 1", "Every hierarchy node Post exists in Posts Master (FR-44)."),
            ("OFFICER_HIERARCHY_NODE", "OFFICER_HIERARCHY_NODE", "1 : 0..*", "Immediate parent post (FR-43)."),
            ("POSTS_MASTER", "POST_OFFICE_TYPE_ALLOWED", "1 : *", "FR-78; required before sanction."),
            ("OFFICE_TYPE", "POST_OFFICE_TYPE_ALLOWED", "1 : *", "FR-78."),
            ("POSTS_MASTER", "SANCTIONED_POST", "1 : *", "Strength is per Post per Office."),
            ("OFFICE_HIERARCHY", "SANCTIONED_POST", "1 : *", "Office Type must be allowed for that Post."),
            ("POSTS_MASTER", "POST_ROLE_MAP", "1 : 1..*", "Unmapped post cannot be sanctioned or assigned (FR-47)."),
            ("ROLE_MASTER", "POST_ROLE_MAP", "1 : 0..*", "Exact unique role name (FR-50)."),
            ("USER_MASTER", "USER_ROLE", "1 : 0..*", "0 for DSR; 1 for Other Department; 1..* for Citizen."),
            ("ROLE_MASTER", "USER_ROLE", "1 : 0..*", "Role Category must equal User Category (FR-34)."),
            ("USER_MASTER", "USER_SECURITY_ANSWER", "1 : 0 or 5", "Citizens only (FR-55)."),
            ("SECURITY_QUESTION", "USER_SECURITY_ANSWER", "1 : *", "Five distinct questions per Citizen."),
            ("USER_MASTER", "POST_OCCUPANCY", "1 : 0..*", "DSR Officers only; multiple posts allowed (FR-26)."),
            ("SANCTIONED_POST", "POST_OCCUPANCY", "1 : 0..*", "Blocked when Occupied ≥ Strength unless FR-67 reservation."),
            ("USER_MASTER", "USER_SESSION", "1 : 0..1 active", "Last-login-wins (FR-76)."),
            ("POST_OCCUPANCY", "USER_SESSION", "1 : 0..*", "Assigned post for DSR session (FR-52)."),
            ("USER_MASTER", "OTP_CHALLENGE", "1 : 0..*", "Also used pre-account for FR-63."),
            ("MODULE_MASTER", "MODULE_FUNCTION", "1 : *", "FR-39."),
            ("MODULE_FUNCTION", "RESOURCE_MASTER", "1 : *", "FR-40."),
            ("ROLE_MASTER", "ROLE_MODULE_FUNCTION", "1 : *", "Except Application Admin (FR-51)."),
            ("MODULE_FUNCTION", "ROLE_MODULE_FUNCTION", "1 : *", "FR-37."),
            ("USER_MASTER", "AUDIT_LOG", "1 : *", "Actor; jobs may write with actor_type = SYSTEM."),
        ],
    )

    doc.add_heading("7. Category-specific population rules", level=1)
    add_table(
        doc,
        ["User Category", "Username", "USER_ROLE", "POST_OCCUPANCY", "Security answers", "Session extras"],
        [
            (
                "Public (Citizen)",
                "Preferred Username (FR-62)",
                "Citizen roles assigned at registration",
                "None",
                "Exactly 5 (FR-55)",
                "No post selection",
            ),
            (
                "DSR Officer",
                "KGID = Username (FR-64)",
                "None — roles via Post–Role",
                "1..* sanctioned posts",
                "None (FR-65)",
                "FR-52 assigned post; FR-53 additional charge on session only",
            ),
            (
                "Other Department",
                "KGID = Username",
                "Exactly one Other Department role",
                "None (Posts Master N/A)",
                "None",
                "No post selection; optional account_end_date (FR-33)",
            ),
        ],
    )

    doc.add_heading("8. Derived facts (not separate entities)", level=1)
    add_table(
        doc,
        ["Fact", "How obtained"],
        [
            ("Session Module Function claims", "Citizen/OD: USER_ROLE → RMF. DSR without additional charge: assigned occupancy → POST_ROLE_MAP → RMF. DSR with FR-53 active: additional-charge post only (FR-38)."),
            ("Office span", "Walk OFFICE_HIERARCHY descendants of session office (FR-59)."),
            ("Immediate-parent transfer filter", "OFFICER_HIERARCHY_NODE.parent of the target post = actor session post (FR-57, FR-60)."),
            ("Available capacity FR-66(a)", "SANCTIONED_POST.occupied_count < sanctioned_strength."),
            ("Wholly unoccupied FR-66(b)", "occupied_count = 0 (reserved Transfer In counts as occupied)."),
            ("Occupied count", "Count of POST_OCCUPANCY in Active or Reserved for that Post+Office; refreshed by FR-68 job."),
        ],
    )

    doc.add_heading("9. Explicit non-entities", level=1)
    doc.add_paragraph(
        "Password / credential secret — authentication is OTP (and biometrics for departmental users) only (FR-09)."
    )
    doc.add_paragraph(
        "Application Admin — deployment-seeded system principal with FN-UM-ADMIN; not created via User Master "
        "workflow and not a Role Master seed role (FR-51). Shown on the diagram for completeness only."
    )
    doc.add_paragraph(
        "Primary / Secondary role flags — removed; DSR context is assigned post vs additional charge (FR-52 / FR-53)."
    )
    doc.add_paragraph(
        "Separate Citizen vs Department user databases — forbidden; one USER_MASTER and one ROLE_MASTER (FR-16)."
    )

    doc.add_heading("10. Alignment to BRD v4.8", level=1)
    doc.add_paragraph(
        "This ERD is a companion to BRD_User_Management_v4.8.docx. If a functional requirement and this model "
        "ever diverge, the BRD requirement text prevails until this ERD is revised. Physical LLD / DDL should "
        "trace to the entity and relationship IDs in sections 5 and 6."
    )

    sig = doc.add_paragraph()
    run = sig.add_run("Related documents: BRD-K3-UM-001 v4.8 (requirements) · ERD-K3-UM-001 v1.0 (this document).")
    run.italic = True
    run.font.size = Pt(10)

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    render_png()
    build_docx()
