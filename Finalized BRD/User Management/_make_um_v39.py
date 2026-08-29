# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v3.9.docx — Office Hierarchy (IGR→DRO→SRO); relieving scoped to offices under actor."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\User Management")
TEMPLATE = BASE / "Template" / "User_Management_Module_BRD_Template.docx"
DST = BASE / "BRD_User_Management_v3.9.docx"


def shade_cell(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_table_rows(table, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "D9E2F3")
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_text(cells[i], val)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph_after(ref_paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    para = Paragraph(new_p, ref_paragraph._parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_table_after(ref_paragraph, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_paragraph._element.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def insert_after_element(ref_el, parent, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_el.addnext(new_p)
    para = Paragraph(new_p, parent)
    para.style = style
    if text:
        para.add_run(text)
    return para


def insert_heading_after(ref, text: str, style: str = "Heading 3", parent=None) -> Paragraph:
    if isinstance(ref, Paragraph):
        return insert_paragraph_after(ref, text, style)
    if parent is None:
        raise ValueError("parent required when ref is an XML element")
    return insert_after_element(ref, parent, text, style)


def insert_table_after_ref(ref_el, parent, headers: list[str], rows: list[tuple[str, ...]]):
    spacer = OxmlElement("w:p")
    ref_el.addnext(spacer)
    tbl_el = OxmlElement("w:tbl")
    spacer.addnext(tbl_el)

    def add_row(values: list[str], header: bool = False) -> None:
        tr = OxmlElement("w:tr")
        tbl_el.append(tr)
        for val in values:
            tc = OxmlElement("w:tc")
            tr.append(tc)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            if header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.insert(0, rpr)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr = OxmlElement("w:tcPr")
                tc_pr.append(shd)
                tc.insert(0, tc_pr)

    add_row(headers, header=True)
    for row in rows:
        add_row(list(row))
    return tbl_el


def build() -> Document:
    doc = Document(str(TEMPLATE))

    cover_map = {
        1: "BUSINESS REQUIREMENTS DOCUMENT",
        2: "User Management Module — KAVERI 3.0",
        3: "Prepared for: Department of Stamps & Registration, Government of Karnataka",
        4: "Version 3.9",
        5: "Date: 29 August 2026",
        6: "Prepared by: Nandha Kumar",
    }
    for idx, text in cover_map.items():
        replace_paragraph_text(doc.paragraphs[idx], text)

    replace_paragraph_text(
        doc.paragraphs[12],
        "This Business Requirements Document (BRD) defines the business requirements for the "
        "User Management Module of KAVERI 3.0 — the integrated platform of the Department of "
        "Stamps and Registration (DSR), Government of Karnataka. It describes a single unified "
        "Role Master and User Master differentiated by user category (Public/Citizen, DSR Officer, "
        "Other Department), OTP-based authentication (email as Username; OTP to mobile only; no password "
        "management), a separate Posts Master "
        "with Post–Role mapping (one post may map to multiple roles), Sanctioned Posts Master "
        "(Post + Office + strength), DSR Officer assignment to one or more vacant sanctioned posts "
        "(roles via Post–Role mapping; no Primary/Secondary assignment; multi-post login selection "
        "with office details), single-role assignment "
        "for Other Department users with optional account "
        "end date, a Module Master (Registration of Documents, Marriage Registration, Encumbrance Search, "
        "Certified Copy, and others), Role-to-Module-Function privilege mapping, Module Function and "
        "Resource (API/URL) masters, runtime access enforcement, DSR officer organisational hierarchy "
        "maintenance, and dedicated user-creation workflows. It serves as the agreed basis for design, "
        "development, testing, and sign-off.",
    )
    replace_paragraph_text(
        doc.paragraphs[14],
        "KAVERI 3.0 requires a centralized mechanism to manage user identities, roles, sanctioned "
        "posts, and access permissions. There shall be one Role Master and one User Master for all "
        "categories — Public users (Citizens), Department users (DSR Officers), and Other Department "
        "users — differentiated by User Category (and Role Category on roles). Authentication shall "
        "be passwordless for all three categories: Email (treated as Username) + Captcha + OTP sent "
        "only to the registered mobile; DSR Officers and Other Department users additionally verify "
        "Biometrics. Mobile may be updated after login; if the mobile is lost, recovery before login "
        "uses a security question chosen at registration. DSR officers are assigned to vacant sanctioned "
        "posts from the Posts Master (strength per office in Sanctioned Posts Master). Posts are mapped "
        "to one or more roles in the Role Master via Post–Role mapping; a user may hold multiple vacant "
        "sanctioned posts with no Primary/Secondary distinction and shall select one post (with office "
        "details) at login when multiple are active. Other Department users are assigned "
        "exactly one role and may have an optional account end "
        "date that deactivates the user when reached. Application modules (e.g. Registration of Documents, "
        "Marriage Registration, Encumbrance Search, Certified Copy) are maintained in a Module Master "
        "and mapped to roles via Module Functions and Resources (APIs/URLs) — DSR organisational roles "
        "are not named after modules. The system shall maintain the DSR officer reporting hierarchy "
        "(organisational chart). Application Admin maintains these masters. Password management "
        "is explicitly out of scope.",
    )

    scope_in = {
        17: "User registration for three categories in a single User Master: Public users (Citizens), Department users (DSR Officers), and Other Department users",
        18: "Email-as-Username login for all categories; OTP to mobile only; Captcha; biometrics for departmental users; security-question recovery for lost mobile; FR-52–FR-54 post selection/display",
        19: "Single unified Role Master with Role Category differentiating Citizen, DSR, and Other Department roles (RBAC)",
        20: "Module Master, Module Function Master, Resource Master (API/URL), and Role–Module–Function mapping maintained by Application Admin",
        21: "DSR Officer Hierarchy Master (posts) and Office Hierarchy Master (IGR Head Office → District Registrar Offices → Sub-Registrar Offices)",
        22: "Posts Master, Sanctioned Posts; Transfer out / relieving by superior — only offices (posts) under the actor per Office Hierarchy; Relieving Date & Order; de-allocate after 11:59 PM",
        23: "Dedicated step-by-step workflows for role assignment during user creation (DSR and Other Department)",
    }
    for idx, text in scope_in.items():
        replace_paragraph_text(doc.paragraphs[idx], text)
    admin_scope = insert_paragraph_after(
        doc.paragraphs[23],
        "Administrative user management (create, edit, suspend, deactivate users)",
        "List Paragraph",
    )
    insert_paragraph_after(
        admin_scope,
        "Audit logging of user-related activities",
        "List Paragraph",
    )

    for p in doc.paragraphs:
        if "password resets" in p.text.lower():
            replace_paragraph_text(
                p,
                "Reduce support overhead related to account access issues through reliable OTP delivery.",
            )

    for p in doc.paragraphs:
        if "Eliminate password-related" in p.text:
            remove_paragraph(p)
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.3 Password Management":
            replace_paragraph_text(p, "6.3 OTP Login and Account Recovery")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "6.6 Administrative User Management":
            replace_paragraph_text(p, "6.7 Administrative User Management")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "Enforce role-based access control to protect sensitive data and functionality.":
            insert_paragraph_after(
                p,
                "Support Module Function and Resource (API/URL) masters with Role–Module–Function mapping and runtime access enforcement maintained by Application Admin.",
                "List Paragraph",
            )
            insert_paragraph_after(
                p,
                "Maintain the DSR Officer Hierarchy Master aligned to the Department organisational chart.",
                "List Paragraph",
            )
            break

    for p in doc.paragraphs:
        if p.text.strip() == "The system shall provide a report of role and permission assignments across all users.":
            sp_note = insert_paragraph_after(
                p,
                "The system shall provide a sanctioned post occupancy report showing, for each office, "
                "the number of sanctioned posts per Post (from Posts Master), occupied count, and vacant slots.",
                "List Paragraph",
            )
            insert_paragraph_after(
                sp_note,
                "The system shall provide a Role-to-Module mapping report showing which modules and "
                "functions are assigned to each role.",
                "List Paragraph",
            )
            break

    replace_table_rows(
        doc.tables[0],
        ["Document Control", ""],
        [
            ("Status", "Draft / In review"),
            ("Owner", "Nandha Kumar (BA) / Prashanth (PO)"),
            ("Reviewers", "Prabhakar Naik (Domain Expert); Kaveri IT Cell"),
            ("Approved By", "Prashanth (Product Owner)"),
        ],
    )

    replace_table_rows(
        doc.tables[1],
        ["Version", "Date", "Author", "Description"],
        [
            (
                "1.7",
                "28-Aug-2026",
                "Nandha Kumar",
                "Aligned to User Management BRD template; three user categories with OTP-only "
                "authentication (no password management); DSR division and role catalogue",
            ),
            (
                "1.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added sanctioned posts master: capture number of sanctioned posts (approved strength) "
                "for each role at each office; DSR user mapping to sanctioned posts only; "
                "over-capacity blocking and occupancy reporting",
            ),
            (
                "1.9",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Other Department role catalogue; Primary and Secondary role assignment; "
                "dedicated step-by-step user-creation workflows for DSR and Other Department users",
            ),
            (
                "2.0",
                "29-Aug-2026",
                "Nandha Kumar",
                "Primary/Secondary roles apply to DSR Officers only; Other Department users get "
                "exactly one role; optional account end date on Other Department user creation "
                "deactivates the user when reached (§6.6.2)",
            ),
            (
                "2.1",
                "29-Aug-2026",
                "Nandha Kumar",
                "Unified single Role Master and User Master for Citizens, DSR Officers, and Other "
                "Department users — differentiated by User Category / Role Category; no separate "
                "masters per category",
            ),
            (
                "2.2",
                "29-Aug-2026",
                "Nandha Kumar",
                "Listed seed roles in Role Master for Role Category = Citizen and Role Category = "
                "Other Department (§6.5.4)",
            ),
            (
                "2.3",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added DSR Department roles listing to §6.5.4 alongside Citizen and Other Department roles",
            ),
            (
                "2.4",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Module Master and Role-to-Module mapping (§6.5.6): Document Registration, "
                "Marriage Registration, Encumbrance Search, Certified Copy and related modules "
                "mapped to organisational / Citizen / Other Department roles",
            ),
            (
                "2.5",
                "29-Aug-2026",
                "Nandha Kumar",
                "Removed online/offline classification for document registration in §6.5.6 — "
                "module described as Registration of Documents only",
            ),
            (
                "2.6",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Module Function Master, Resource Master (API/URL), Role–Module–Function "
                "mapping with example rows, Application Admin CRUD, and runtime enforcement (§6.5.6)",
            ),
            (
                "2.7",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added DSR Officer Hierarchy Master (§6.5.7) to maintain Department officer reporting "
                "structure aligned to DSR organisational chart (Admin/Law/Computers, Vigilance, "
                "Computers, Enforcement, Intelligence & Audit, DIGR CVC)",
            ),
            (
                "2.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "DSR Officers may be assigned to multiple sanctioned posts when each post is vacant "
                "(Primary plus optional additional posts); over-capacity still blocked",
            ),
            (
                "2.9",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added separate Posts Master; Post–Role mapping (one post to multiple roles); "
                "Sanctioned Posts Master references Posts Master (Post + Office + strength)",
            ),
            (
                "3.0",
                "29-Aug-2026",
                "Nandha Kumar",
                "Removed Primary and Secondary role/post assignment; DSR Officers are assigned one or "
                "more vacant sanctioned posts equally; roles derived only via Post–Role mapping",
            ),
            (
                "3.1",
                "29-Aug-2026",
                "Nandha Kumar",
                "Resolved structural conflicts: division-specific Posts (no generic DIGR/AIGR over-provisioning); "
                "Post–Role mapping typically 1:1 to matching unique roles; Hierarchy Master references Posts "
                "not Roles; ACS/Secretary root role added; unique FDA/SDA/Typist role names per division",
            ),
            (
                "3.2",
                "29-Aug-2026",
                "Nandha Kumar",
                "Role–Module–Function mapping aligned to unique Role Master names (FDA (Enforcement) not bare FDA); "
                "FR-50 referential integrity for Role–Module–Function ↔ Role Master; Application Admin treated as "
                "system-level actor outside Role Master (Option B — not a DSR seed role); "
                "FR-31 retired — was tied to Primary/Secondary (additional) post assignment with mandatory end date, "
                "removed in v3.0 (gap left intentional for audit traceability)",
            ),
            (
                "3.3",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added FR-52: after authentication, if a DSR Officer has multiple active post occupancies, "
                "the user must select one post for the session; each choice is displayed with office details; "
                "session roles/privileges derive from the selected post via Post–Role mapping",
            ),
            (
                "3.4",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added FR-53: after login and post selection, DSR Officer may switch session to vacant "
                "subordinate posts under the login-selected post using Hierarchy Master (cascade through "
                "vacant nodes, same office); example SR → vacant FDA → vacant SDA; Enforcement seed "
                "hierarchy aligned so FDA/SDA report under Sub-Registrar",
            ),
            (
                "3.5",
                "29-Aug-2026",
                "Nandha Kumar",
                "FR-53 revised: vacant subordinate post/role selection occurs only during login (after "
                "authentication and FR-52 post selection, before home); after login completes the user "
                "cannot switch post for the remainder of the session",
            ),
            (
                "3.6",
                "29-Aug-2026",
                "Nandha Kumar",
                "FR-53 UI: during login post/role switch, options shown in a dropdown as Role with Post "
                "(and Office); FR-54: after login, home/header displays the logged-in Post with mapped Role(s)",
            ),
            (
                "3.7",
                "29-Aug-2026",
                "Nandha Kumar",
                "Login for all categories uses Email as Username; OTP sent only to registered mobile "
                "(not email); mobile updatable after login; lost-mobile recovery before login via "
                "security question selected at registration (FR-55, FR-56); FR-04/05/06/07/10/12/13 updated",
            ),
            (
                "3.8",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Transfer out / relieving process (FR-57, FR-58): only hierarchy superior may relieve "
                "an officer from assigned post(s); capture Relieving Date and Relieving Order; system "
                "de-allocates user–post mapping after 11:59 PM of the Relieving Date; §6.6.3 workflow "
                "with hierarchy examples (e.g. DIGR Enforcement relieves Sub-Registrar; SR relieves FDA)",
            ),
            (
                "3.9",
                "29-Aug-2026",
                "Nandha Kumar",
                "Added Office Hierarchy Master (FR-59): IGR Head Office → District Registrar Offices → "
                "Sub-Registrar Offices; Transfer Out / Relieving (FR-57) displays only offices and post "
                "occupancies under the actor's session office; Enforcement post chain aligned DRO→SR→FDA",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[2],
        ["Name / Role", "Department", "Responsibility"],
        [
            ("Prashanth", "DSR / Product", "Prioritizes requirements; approves scope"),
            ("Nandha Kumar", "Business Analysis", "Documents and validates requirements"),
            ("Prabhakar Naik", "Domain Expert", "Validates DSR roles, Other Department roles, sanctioned posts, and organizational structure"),
            ("Kaveri IT Cell", "Engineering", "Reviews technical feasibility"),
            ("Citizens (Public users)", "External", "Self-register and access citizen portal services"),
            (
                "DSR Officers & Other Department users",
                "Government",
                "Access departmental modules via OTP + biometrics",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[3],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-01", "The system shall support instant self-registration for Public users (Citizens) with no approval workflow. Registration shall capture Email (Username), mobile, and a security question with answer (FR-55).", "High"),
            ("FR-02", "The system shall allow Department users (DSR Officers) to be created only by authorised administrative roles. Creation shall capture Email (Username), mobile, and a security question with answer (FR-55).", "High"),
            ("FR-03", "The system shall allow Other Department users (officers/staff from other government departments) to be created only by authorised administrative roles. Creation shall capture Email (Username), mobile, and a security question with answer (FR-55).", "High"),
            ("FR-04", "The system shall treat Email as the Username for all user categories. The system shall prevent duplicate registrations using the same Email (Username) within a user category.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[4],
        ["ID", "Requirement", "Priority"],
        [
            (
                "FR-05",
                "Public users (Citizens) shall authenticate using Email (Username) + Captcha + OTP "
                "on every login. OTP shall be sent only to the registered mobile number.",
                "High",
            ),
            (
                "FR-06",
                "Department users (DSR Officers) shall authenticate using Email (Username) + Captcha + "
                "OTP + Biometrics on every login. OTP shall be sent only to the registered mobile number.",
                "High",
            ),
            (
                "FR-07",
                "Other Department users shall authenticate using Email (Username) + Captcha + OTP + "
                "Biometrics on every login. OTP shall be sent only to the registered mobile number.",
                "High",
            ),
            ("FR-08", "The system shall allow users to log out and terminate their active session.", "High"),
            ("FR-09", "The system shall not provide password-based login, password reset, or password change for any user category.", "High"),
            (
                "FR-52",
                "After successful authentication, if a DSR Officer has more than one active sanctioned-post "
                "occupancy, the system shall present a mandatory post-selection step before entering the "
                "application. Each option shall display the Post Name together with Office details "
                "(Office Name and Office Code as available). The user shall select exactly one post for "
                "the session. If the officer has only one active occupancy, the system shall auto-select "
                "that post and skip the selection screen. Session roles and Module Function claims for "
                "that login shall be derived from the selected post via Post–Role mapping (not the union "
                "of all posts). Other Department and Citizen users are not subject to post selection.",
                "High",
            ),
            (
                "FR-53",
                "During login only — after authentication and after the DSR Officer's login post is "
                "determined (FR-52 multi-post selection or single-post auto-select) and before the "
                "home page is entered — the system shall allow the user to choose to act under a vacant "
                "subordinate post of that login-selected post, using the DSR Officer Hierarchy Master "
                "(FR-43). Eligibility: for the login-selected Post at the same Office, examine each "
                "immediate child post in the hierarchy; if that child post is vacant at that Office, it "
                "shall be offered as an option and the system shall continue one level further under that "
                "vacant child (cascade while posts remain vacant). Occupied subordinate posts shall not "
                "be offered and shall block cascading beneath them. Example: user logs in as Sub-Registrar "
                "at Office A; if FDA under Sub-Registrar is vacant, acting as FDA is allowed; if SDA under "
                "that vacant FDA is also vacant, acting as FDA or SDA is allowed. The user may also "
                "continue under the original login-selected post. The system shall present these options "
                "in a drop-down control where each entry displays the Role together with the Post "
                "(and Office details), e.g. \"FDA (Enforcement) — FDA (Enforcement) / SRO Yeshwanthapura\". "
                "Session roles and Module Function claims shall be derived from the post chosen at this "
                "login step via Post–Role mapping. After login is complete and the user has entered the "
                "application, the user shall not be allowed to switch post for the remainder of that "
                "session; a new login is required to choose a different post. Choice shall be "
                "audit-logged (actor, login post, chosen post, office).",
                "High",
            ),
            (
                "FR-54",
                "After login is complete, the system shall display the logged-in (session) Post together "
                "with the Role(s) mapped to that Post via Post–Role mapping (and Office details) on the "
                "home page and/or application header for the duration of the session. The display is "
                "read-only for post switching (FR-53 prohibits mid-session switch); it informs the user "
                "of the active Post and Role context.",
                "High",
            ),
            (
                "FR-55",
                "During registration / user creation for all three user categories, the system shall "
                "require the user (or Admin creating the user) to select one security question from a "
                "predefined list and provide an answer. The question and answer shall be stored securely "
                "(answer hashed/encrypted) and used only for lost-mobile recovery (FR-56).",
                "High",
            ),
            (
                "FR-56",
                "If a user has lost access to the registered mobile number, before login the system shall "
                "allow recovery as follows: user enters Email (Username); system presents the security "
                "question selected at registration; if the answer is correct, the system shall allow the "
                "user to update the mobile number before completing login; subsequent OTP for that login "
                "shall be sent only to the newly updated mobile. Incorrect answers shall be rate-limited "
                "and audit-logged. This flow applies to all three user categories.",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[5],
        ["ID", "Requirement", "Priority"],
        [
            (
                "FR-10",
                "The system shall dispatch OTP only to the user's registered mobile number (SMS) within "
                "5 seconds of request. OTP shall not be sent to email for login authentication.",
                "High",
            ),
            ("FR-11", "The system shall validate Captcha before OTP dispatch for all user categories.", "High"),
            (
                "FR-12",
                "Account recovery for lost mobile shall use the security-question flow (FR-56) to update "
                "mobile before login. The system shall not provide password reset. OTP remains mobile-only "
                "after the mobile is updated.",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[6],
        ["ID", "Requirement", "Priority"],
        [
            (
                "FR-13",
                "The system shall allow users to view and update their profile information after login. "
                "Mobile number may be updated after login when required (subject to verification policy, "
                "e.g. OTP to the new mobile). Email is the Username; any change to Email shall enforce "
                "uniqueness (FR-04) and be audit-logged.",
                "High",
            ),
            ("FR-14", "The system shall allow users to upload and update a profile photo where applicable.", "Medium"),
            ("FR-15", "The system shall allow administrators to deactivate user accounts with reason and audit trail.", "High"),
        ],
    )

    replace_table_rows(
        doc.tables[7],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-16", "The system shall support RBAC using a single unified Role Master. Roles shall be differentiated by Role Category (Citizen, DSR, Other Department) as defined in Section 6.5.1. Users shall be stored in a single User Master differentiated by User Category.", "High"),
            (
                "FR-17",
                "The system shall allow administrators to assign DSR Officers to one or more vacant "
                "sanctioned posts during user creation (at least one post required). There is no "
                "Primary/Secondary role or post distinction. At login, the officer selects one active "
                "post for the session when multiple posts are assigned (FR-52); session roles shall be "
                "those mapped to the selected post via Post–Role mapping.",
                "High",
            ),
            (
                "FR-28",
                "There shall be no separate Role Masters or User Masters for Citizens, DSR Officers, "
                "or Other Department users. All roles and users shall reside in one Role Master and one "
                "User Master respectively, differentiated by Role Category / User Category.",
                "High",
            ),
            (
                "FR-29",
                "During Other Department user creation, the system shall allow administrators to assign "
                "exactly one role from the Role Master filtered by Role Category = Other Department. "
                "Posts Master and sanctioned post occupancy do not apply to Other Department users.",
                "High",
            ),
            (
                "FR-30",
                "At least one vacant sanctioned post must be assigned to each DSR Officer at account "
                "creation. Optional End Date may be set per post occupancy; if set, the system shall "
                "free that occupancy on the end date and recalculate effective roles via Post–Role mapping. "
                "Formal Transfer out / relieving shall follow FR-57 and FR-58.",
                "High",
            ),
            (
                "FR-31",
                "RETIRED (v3.0) — Previously required mandatory end date on Secondary / additional "
                "post assignment for DSR Officers. Superseded by removal of Primary/Secondary assignment; "
                "optional End Date on any post occupancy is covered by FR-30. Number retained for audit "
                "traceability; do not implement.",
                "—",
            ),
            (
                "FR-45",
                "A DSR Officer may be assigned to multiple sanctioned posts concurrently, provided each "
                "selected post is vacant at assignment. Non-vacant or over-capacity assignment shall be "
                "blocked. Occupied count shall increase per assigned post and decrease when occupancy ends "
                "(including after relieving under FR-58).",
                "High",
            ),
            (
                "FR-57",
                "The system shall support a Transfer out / relieving process for DSR Officers. An officer "
                "may be relieved from one or more currently assigned post occupancies (Post + Office). "
                "When the superior opens Transfer Out / Relieving, the system shall display only the "
                "offices (and post occupancies at those offices) that fall under the actor's session "
                "Office as per the Office Hierarchy Master (FR-59) — i.e. the actor's own office and "
                "descendant offices (IGR Head Office sees District Registrar Offices and their "
                "Sub-Registrar Offices; a District Registrar Office sees only its subordinate "
                "Sub-Registrar Offices; a Sub-Registrar Office sees only that office). Within that "
                "office scope, relieving shall be allowed only where the actor's session Post is the "
                "immediate parent of the target Post in the DSR Officer Hierarchy Master (FR-43). "
                "Example: District Registrar at DRO Mysuru sees only SRO offices under DRO Mysuru and "
                "may relieve Sub-Registrar at those SROs; Sub-Registrar at SRO Yeshwanthapura sees only "
                "that SRO and may relieve FDA/DEO under Sub-Registrar there. While relieving, the system "
                "shall capture Relieving Date and Relieving Order (order number / reference; upload of "
                "order document where applicable). Relieving actions shall be audit-logged. Citizens and "
                "Other Department users are out of scope for this process.",
                "High",
            ),
            (
                "FR-58",
                "After a relieving is recorded (FR-57), the system shall retain the user–post occupancy "
                "until the end of the Relieving Date. The system shall de-allocate / remove the mapping "
                "of the user to the relieved post after 11:59 PM of the Relieving Date (i.e. occupancy "
                "ends at the close of that calendar day). Occupied count for that Post + Office shall "
                "decrease; the post becomes vacant for new assignment. If the user has no remaining "
                "active post occupancies after de-allocation, login shall be blocked or limited per "
                "policy until a new post is assigned. De-allocation shall be audit-logged.",
                "High",
            ),
            (
                "FR-59",
                "The system shall maintain an Office Hierarchy Master for DSR offices. Structure: "
                "IGR Office is the Head Office (root); under the Head Office are District Registrar "
                "Offices (DRO); under each District Registrar Office are Sub-Registrar Offices (SRO). "
                "Each Office shall have Office Code, Office Name, Office Type (Head Office / District "
                "Registrar Office / Sub-Registrar Office), and optional Parent Office. Application Admin "
                "shall add, edit, enable/disable, and update offices with audit trail. Sanctioned Posts "
                "and Transfer Out / Relieving office scoping shall use this hierarchy (FR-24, FR-57).",
                "High",
            ),
            (
                "FR-46",
                "The system shall maintain a separate Posts Master for DSR establishment posts "
                "(e.g. Sub-Registrar (SR), FDA (Enforcement), DEO, DRO, HQA (Enforcement)). "
                "Application Admin shall add, edit, enable/disable, and update posts. Posts Master "
                "is distinct from Role Master.",
                "High",
            ),
            (
                "FR-47",
                "The system shall maintain a Post–Role mapping table. Each Post shall map to the "
                "corresponding unique Role(s) from the Role Master. Division-specific posts "
                "(e.g. POST-DIGR-ADMIN, POST-DIGR-ENF) shall not be mapped to roles of other divisions, "
                "so that assigning one post does not grant Admin + Vigilance + Enforcement access "
                "together. Application Admin shall add, edit, and remove mappings with audit trail. "
                "One post may map to multiple roles only when those roles intentionally belong to the "
                "same post's functional scope (not across unrelated divisions).",
                "High",
            ),
            (
                "FR-48",
                "The Sanctioned Posts Master shall reference a Post from the Posts Master (not Role "
                "directly), plus Office and sanctioned strength. Strength and vacancy are maintained "
                "per Post per Office.",
                "High",
            ),
            (
                "FR-49",
                "Posts Master entries for DSR leadership and staff shall be division-specific where "
                "the organisation chart distinguishes them (e.g. POST-DIGR-ADMIN, POST-DIGR-VIG, "
                "POST-DIGR-ENF, POST-FDA-ADMIN, POST-FDA-ENF). Generic posts that would map to all "
                "division variants of DIGR/AIGR/FDA/SDA are not permitted in seed data.",
                "High",
            ),
            (
                "FR-33",
                "During Other Department user creation, the system shall allow an optional End Date. "
                "If an End Date is entered, the system shall automatically deactivate the user account "
                "on that date and block login thereafter.",
                "High",
            ),
            (
                "FR-34",
                "When assigning a role to a user, the system shall present only roles whose Role Category "
                "matches the user's User Category (Citizen roles for Citizens, DSR roles for DSR Officers, "
                "Other Department roles for Other Department users).",
                "High",
            ),
            (
                "FR-35",
                "The Role Master shall include the seed roles listed in Section 6.5.4 for Role Category = "
                "Citizen, Role Category = Other Department, and Role Category = DSR (Department). "
                "Application Admin may add, edit, or deactivate roles within each Role Category.",
                "High",
            ),
            (
                "FR-36",
                "The system shall maintain a Module Master of application modules (including Registration "
                "of Documents, Marriage Registration, Encumbrance Search, Certified Copy, and others "
                "listed in Section 6.5.6). Modules are not DSR organisational roles; they define "
                "functional areas to which roles are mapped. Registration of Documents is a single "
                "module with no online/offline classification in this master.",
                "High",
            ),
            (
                "FR-37",
                "The system shall allow Application Admin to map roles to Module Functions "
                "(Role–Module–Function mapping). Access to a module capability is granted only through "
                "this mapping (except Application Admin system privileges — see FR-51). Each mapping "
                "row shall reference an exact Role Master name (unique role name) and a valid Module "
                "Function code. Application Admin shall add, edit, and update mappings with audit trail.",
                "High",
            ),
            (
                "FR-39",
                "The system shall maintain a Module Function Master under each module (e.g. VIEW, ADD, "
                "EDIT, APPROVE, SIGN, PRINT, DOWNLOAD). Application Admin shall add, edit, and update "
                "module functions.",
                "High",
            ),
            (
                "FR-40",
                "The system shall maintain a Resource Master of APIs and URLs linked to Module Functions "
                "(resource type API or URL, HTTP method, path pattern). Application Admin shall add, "
                "edit, and update resources. Roles shall not store raw API/URL lists; access is resolved "
                "via Role → Module Function → Resource.",
                "High",
            ),
            (
                "FR-41",
                "At runtime the application shall enforce access as follows: resolve the user's session "
                "roles (for DSR Officers: from the session post fixed at login — FR-52 / FR-53); "
                "load Role–Module–Function claims; for each API/URL request look up the Resource Master "
                "to obtain the required Module Function; allow only if the user holds that function; "
                "otherwise deny (HTTP 403) and audit the attempt. Application Admin system principals "
                "are authorised for FN-UM-ADMIN outside Role–Module–Function mapping (FR-51).",
                "High",
            ),
            (
                "FR-50",
                "Role–Module–Function mapping entries shall be validated against exact Role Master names "
                "(referential integrity). The system shall reject save of a mapping whose Role name does "
                "not match an active Role Master entry character-for-character (e.g. bare FDA is invalid "
                "when only FDA (Admin), FDA (Enforcement), etc. exist). Same integrity pattern as "
                "Hierarchy ↔ Posts (FR-44).",
                "High",
            ),
            (
                "FR-51",
                "Application Admin is a system-level / deployment-seeded actor, not a Role Master seed "
                "role and not created via the normal User Master creation workflow. Privileges for "
                "maintaining masters (including FN-UM-ADMIN resources) are granted to this principal "
                "outside Role–Module–Function mapping. Application Admin shall not appear as a DSR "
                "(or other) row in the Role–Module–Function example catalogue.",
                "High",
            ),
            (
                "FR-42",
                "Application Admin shall create, edit, enable/disable, and update Module Master, "
                "Module Function Master, Resource Master, and Role–Module–Function mapping tables "
                "from the User Management admin console. All changes shall be audit-logged.",
                "High",
            ),
            (
                "FR-38",
                "A user's effective module access for a session shall be the union of Module Functions "
                "mapped to the roles held in that session: for DSR Officers, roles derived from the "
                "current session post (login-selected post or a vacant subordinate post chosen during "
                "login under FR-53; fixed for the session) via Post–Role mapping; for Other Department "
                "users, the single assigned role; for Citizens, assigned Citizen roles.",
                "High",
            ),
            ("FR-18", "The system shall restrict access to features, APIs, and URLs based on the user's assigned role(s) and Role–Module–Function / Resource mappings.", "High"),
            ("FR-19", "The system shall maintain DSR organizational divisions and DSR roles within the unified Role Master (Role Category = DSR) as defined in Section 6.5.1 and the DSR Officer Hierarchy Master in Section 6.5.7.", "High"),
            (
                "FR-43",
                "The system shall maintain a DSR Officer Hierarchy Master representing the reporting "
                "structure of Department officers (organisational chart). Each hierarchy node shall "
                "reference a Post from the Posts Master (not a Role) and an optional parent hierarchy "
                "node (also a Post). Application Admin shall add, edit, reorder, enable/disable, and "
                "update hierarchy nodes with audit trail. Reporting lines are position-based so that "
                "a specific post holder has one parent post in the tree.",
                "High",
            ),
            (
                "FR-44",
                "The DSR Officer Hierarchy Master shall seed and support the structure: Additional Chief "
                "Secretary / Principal Secretary / Secretary → Inspector General of Registration & "
                "Commissioner of Stamps → Divisions (Admin Law & Computers; Vigilance; Computers; "
                "Enforcement; Intelligence & Audit; DIGR CVC) with subordinate posts as listed in "
                "Section 6.5.7. Every hierarchy node Post must exist in Posts Master; every mapped Role "
                "in Post–Role mapping must exist in Role Master with a unique role name.",
                "High",
            ),
            (
                "FR-24",
                "The system shall maintain a Sanctioned Posts Master that references Posts Master entries "
                "and records sanctioned strength (approved headcount) for each Post at each office. "
                "Sanctioned strength applies only to DSR Posts Master entries.",
                "High",
            ),
            (
                "FR-25",
                "The system shall allow Application Admin to configure and update sanctioned strength "
                "per Post per office in the Sanctioned Posts Master, with audit trail of changes.",
                "High",
            ),
            (
                "FR-26",
                "DSR department users shall be assigned only to posts defined in the Posts Master and "
                "sanctioned in the Sanctioned Posts Master. A user may hold multiple post occupancies. "
                "Assignment to an unlisted post or to a non-vacant (over-capacity) sanctioned post shall be blocked.",
                "High",
            ),
            (
                "FR-27",
                "The system shall display vacant vs occupied sanctioned posts per Post per office "
                "and prevent over-capacity assignment. When assigning posts to a DSR "
                "Officer, only vacant sanctioned posts shall be selectable.",
                "High",
            ),
            (
                "FR-32",
                "The system shall provide dedicated step-by-step workflows for role assignment during "
                "user creation for DSR Officers (Section 6.6.1) and Other Department users (Section 6.6.2).",
                "High",
            ),
        ],
    )

    replace_table_rows(
        doc.tables[8],
        ["ID", "Requirement", "Priority"],
        [
            ("FR-20", "The system shall allow administrators to create, edit, suspend, and deactivate user accounts by category.", "High"),
            ("FR-21", "The system shall allow administrators to search and filter users by category, role, office, division, and status.", "Medium"),
            ("FR-22", "The system shall log all administrative actions performed on user accounts.", "High"),
            ("FR-23", "The system shall notify a user via SMS/email when their account is created, suspended, or deactivated.", "Medium"),
        ],
    )

    replace_table_rows(
        doc.tables[9],
        ["Category", "Requirement"],
        [
            ("Security", "No password storage — authentication is OTP-based only; OTP sent only to registered mobile."),
            ("Security", "All data in transit shall be encrypted using TLS 1.2 or higher."),
            ("Security", "Biometric data for departmental users shall comply with Aadhaar Act, 2016 and UIDAI guidelines."),
            ("Security", "Security-question answers shall be stored hashed/encrypted; recovery attempts rate-limited and audited."),
            ("Performance", "OTP dispatch (SMS to mobile only) shall complete within 5 seconds of user request."),
            ("Performance", "Login and authentication requests shall complete within 2 seconds after OTP/biometric verification."),
            ("Availability", "The module shall maintain 99.9% uptime, excluding scheduled maintenance."),
            ("Usability", "Registration and OTP login workflows shall be completable on desktop and mobile browsers."),
            ("Auditability", "All create, update, delete, login, mobile-update, security-question recovery, and access-control actions shall be logged with timestamp and actor."),
            ("Compliance", "The module shall comply with Karnataka e-Governance, MeitY/CERT-In, and applicable data protection norms."),
        ],
    )

    replace_table_rows(
        doc.tables[10],
        ["Risk", "Impact", "Mitigation"],
        [
            ("OTP delivery failure (SMS)", "High", "Retry mechanism; allow mobile update after login; lost-mobile recovery via security question (FR-56)."),
            ("Biometric device unavailability", "Medium", "Define fallback procedure for DSR and Other Department users per security policy."),
            ("Security-question answer forgotten", "Medium", "Admin-assisted recovery with strong identity verification and audit."),
            ("Over-capacity post assignment", "High", "Enforce sanctioned strength validation; block assignment when role at office is full."),
            ("Post occupancy end date / relieving not enforced", "High", "Scheduled job at/after 11:59 PM of Relieving Date (FR-58) and optional End Date (FR-30); audit alert."),
            ("Unauthorised relieving", "High", "Scope list to offices under actor (FR-59); allow relieve only if session Post is immediate parent of target Post (FR-57)."),
            ("Other Department end date not enforced", "Medium", "Scheduled job to deactivate Other Department users on optional End Date; audit alert."),
            ("Scope creep beyond OTP-only auth", "Medium", "Maintain strict change-control; no password management in KAVERI 3.0."),
            ("Delayed stakeholder sign-off", "Low", "Set clear review deadlines and escalation path."),
        ],
    )

    replace_table_rows(
        doc.tables[11],
        ["Term", "Definition"],
        [
            ("BRD", "Business Requirements Document"),
            ("RBAC", "Role-Based Access Control"),
            ("OTP", "One-Time Password — primary authentication credential sent only to registered mobile (SMS); no static password"),
            ("Email (Username)", "User's email address used as the login Username for all three user categories (FR-04)"),
            ("Security question", "Question selected at registration with a secret answer; used for lost-mobile recovery before login (FR-55, FR-56)"),
            ("DSR", "Department of Stamps and Registration, Government of Karnataka"),
            ("UAT", "User Acceptance Testing"),
            ("IGR", "Inspector General of Registration"),
            ("DIGR", "Deputy Inspector General of Registration"),
            ("AIGR", "Assistant Inspector General of Registration"),
            ("Sanctioned post", "A Posts Master entry sanctioned at a specific office with approved strength; Post + Office + Strength"),
            ("Sanctioned strength", "The approved number of occupants for a given Post at a given office"),
            ("Posts Master", "Separate catalogue of DSR establishment posts (distinct from Role Master)"),
            ("Post–Role mapping", "Table linking a Post to one or more Roles; one post may map to multiple roles"),
            ("Office", "A concrete DSR office instance (e.g. IGR Head Office, DRO Mysuru, SRO Yeshwanthapura)"),
            ("Office Hierarchy Master", "Tree of DSR offices: IGR Head Office (root) → District Registrar Offices → Sub-Registrar Offices (FR-59)"),
            ("Office span", "The actor's session Office plus all descendant offices in the Office Hierarchy Master; used to scope Transfer Out / Relieving lists (FR-57)"),
            ("Role Master", "Single master of all roles for Citizens, DSR Officers, and Other Department users; differentiated by Role Category"),
            ("User Master", "Single master of all users across categories; differentiated by User Category"),
            ("Role Category", "Attribute on a role in the Role Master: Citizen, DSR, or Other Department"),
            ("User Category", "Attribute on a user in the User Master: Public (Citizen), DSR Officer, or Other Department"),
            ("Module Master", "Catalogue of application modules (e.g. Registration of Documents, Marriage Registration, Encumbrance Search, Certified Copy)"),
            ("Module Function Master", "Catalogue of privilege verbs under a module — e.g. VIEW, ADD, EDIT, APPROVE, SIGN, PRINT, DOWNLOAD"),
            ("Resource Master", "Catalogue of APIs and URLs linked to a Module Function (type, HTTP method, path pattern)"),
            ("Role–Module–Function mapping", "Association of a Role Master role (exact unique name) to one or more Module Functions; Application Admin system privileges are not represented here (FR-51)"),
            ("Module function", "Privilege verb within a module — e.g. VIEW, ADD, EDIT, APPROVE, SIGN, PRINT, DOWNLOAD"),
            ("Resource", "An API endpoint or UI URL/route protected by a Module Function"),
            ("Access enforcement", "Runtime check: request path/method → Resource → required Module Function → user's role claims"),
            ("Application Admin", "System-level / deployment-seeded actor who maintains masters; not a Role Master seed role; privileges outside Role–Module–Function mapping (FR-51)"),
            ("DSR Officer Hierarchy Master", "Master of reporting relationships among DSR Posts (Posts Master) as per the Department organisational chart"),
            ("Hierarchy node", "A node in the DSR Officer Hierarchy Master linking a Post to an optional parent Post node"),
            ("Unique role name", "Distinct Role Master name including division context where needed (e.g. FDA (Admin), FDA (Enforcement)) — no duplicate bare FDA/SDA/Typist entries"),
            ("Other Department role", "A role in the Role Master with Role Category = Other Department; exactly one such role is assigned per Other Department user"),
            ("Post occupancy", "Assignment of a DSR user to a sanctioned post at an office; roles follow from Post–Role mapping for the login-selected post (FR-52); no Primary/Secondary distinction"),
            ("Relieving / Transfer out", "Process by which a superior removes a DSR Officer from a post occupancy within the superior's office span, capturing Relieving Date and Relieving Order (FR-57–FR-59); mapping ends after 11:59 PM of that date (FR-58)"),
            ("Relieving Date", "Calendar date captured during relieving; user–post mapping is removed after 11:59 PM of this date"),
            ("Relieving Order", "Order number / reference (and optional uploaded order document) recorded when relieving an officer"),
            ("Login post selection", "After authentication, when a DSR Officer has multiple active post occupancies, the mandatory step to choose one Post (shown with Office details) for the session (FR-52)"),
            ("Session post", "The Post + Office context fixed for the DSR session at login — either the login-selected post or a vacant subordinate post chosen during login (FR-53); cannot be changed mid-session; session roles derive from this post via Post–Role mapping"),
            ("Vacant subordinate switch", "Login-time only drop-down choice allowing a DSR Officer to act under a vacant child/descendant post of the login-selected post at the same office; each option shows Role with Post (FR-53); not available after login"),
            ("Session Post–Role display", "After login, persistent read-only display of the logged-in Post with its mapped Role(s) and Office (FR-54)"),
            ("Parent department", "The government department to which an Other Department user belongs (e.g. Revenue, Treasury, Police)"),
            ("Account End Date", "Optional date on an Other Department user account; if set, the system deactivates the user on that date"),
        ],
    )

    replace_table_rows(
        doc.tables[12],
        ["Name", "Role", "Signature", "Date"],
        [
            ("Prashanth", "Product Owner", "", ""),
            ("Prabhakar Naik", "Domain Expert", "", ""),
            ("Kaveri IT Cell Lead", "IT Security / Engineering", "", ""),
        ],
    )

    rbac_fr_tbl_el = doc.tables[7]._tbl
    rbac_heading = next(p for p in doc.paragraphs if p.text.strip() == "6.5 Role-Based Access Control (RBAC)")
    rbac_parent = rbac_heading._parent
    rbac_fr_tbl_el.getparent().remove(rbac_fr_tbl_el)
    rbac_heading._element.addnext(rbac_fr_tbl_el)

    sub1 = insert_heading_after(
        rbac_fr_tbl_el, "6.5.1 Unified Role Master and User Master", "Heading 3", parent=rbac_parent
    )
    note1 = insert_paragraph_after(
        sub1,
        "The system shall maintain a single Role Master and a single User Master for all user types. "
        "There shall be no separate masters for Citizens, DSR Officers, or Other Department users. "
        "Differentiation is by Role Category on roles and User Category on users. When assigning a role, "
        "the system shall filter the Role Master to roles matching the user's User Category (FR-34).",
        "Normal",
    )
    cat_diff_tbl = insert_table_after(
        note1,
        ["Master", "Differentiator", "Values", "Purpose"],
        [
            (
                "User Master",
                "User Category",
                "Public (Citizen); DSR Officer; Other Department",
                "One user store; category drives auth, creation path, and eligible roles",
            ),
            (
                "Role Master",
                "Role Category",
                "Citizen; DSR; Other Department",
                "One role store; category filters which roles can be assigned to which users",
            ),
        ],
    )
    div_intro = insert_heading_after(
        cat_diff_tbl,
        "DSR roles in the Role Master (Role Category = DSR) are organised by division. Role names "
        "shall be unique (e.g. FDA (Admin), FDA (Enforcement) — not a single shared FDA). Summary:",
        "Normal",
        parent=rbac_parent,
    )
    div_tbl = insert_table_after(
        div_intro,
        ["Division", "Roles (Role Category = DSR) — unique names"],
        [
            ("Secretariat", "ACS / Principal Secretary / Secretary"),
            ("Top Management", "IGR"),
            (
                "Division 1 — Admin, Law & Computers",
                "DIGR (Admin, Law & Computers), AIGR (Admin), HQA (Admin), SRO (Admin), FDA (Admin), SDA (Admin), Typist (Admin), HQA (RTI), FDA (RTI), SDA (RTI), Statistical Inspector, Accountant Superintendent (Admin)",
            ),
            ("Division 2 — Vigilance", "DIGR (Vigilance), Law Officer"),
            (
                "Division 3 — Computers",
                "AIGR (Computers), System Integrator, PMU, Application Developer, HQA / Project Manager (Comp), SRO (Comp), FDA (Computers), SDA (Computers)",
            ),
            ("Division 4 — Enforcement", "DIGR (Enforcement), DRO, HQA (Enforcement), Sub-Registrar (SR), FDA (Enforcement), SDA (Enforcement), DEO"),
            (
                "Division 5 — Intelligence & Audit",
                "DIGR (Intelligence), AIGR (Audit), HQA (Audit), Superintendent (Audit), FDA (Audit), SDA (Audit), Typist (Audit)",
            ),
            ("Division 6 — CVC", "DIGR CVC, JD Town Planning"),
        ],
    )

    sub2 = insert_heading_after(
        div_tbl, "6.5.2 User Categories and Authentication", "Heading 3", parent=rbac_parent
    )
    note2 = insert_paragraph_after(
        sub2,
        "All users are stored in the single User Master, differentiated by User Category. "
        "All user categories authenticate without passwords. Login uses Email as Username; OTP is "
        "sent only to the registered mobile. Captcha is required before OTP dispatch. "
        "DSR Officers and Other Department users also verify Biometrics. Mobile may be updated after "
        "login (FR-13); lost mobile before login uses security-question recovery (FR-56). "
        "DSR Officers with multiple active post occupancies must select one post (with office details) "
        "after authentication before entering the application (FR-52). During that same login flow "
        "they may choose a vacant subordinate post under the selected post via a Role-with-Post "
        "drop-down (FR-53). After login completes, the home/header shows the logged-in Post with "
        "mapped Role(s) (FR-54); post cannot be switched until the next login.",
        "Normal",
    )
    cat_tbl = insert_table_after(
        note2,
        ["User Category", "Description", "Authentication"],
        [
            ("Public users (Citizens)", "Citizens accessing Kaveri portal services — User Category = Public (Citizen)", "Email (Username) + Captcha + OTP to mobile"),
            ("Department users (DSR Officers)", "Officers and staff of DSR — User Category = DSR Officer", "Email (Username) + Captcha + OTP to mobile + Biometrics; then FR-52/53"),
            (
                "Other Department users",
                "Officers/staff from other government departments — User Category = Other Department",
                "Email (Username) + Captcha + OTP to mobile + Biometrics",
            ),
        ],
    )

    auth_flow_intro = insert_heading_after(
        cat_tbl,
        "Login and mobile recovery (all categories — FR-05–FR-07, FR-10–FR-13, FR-55, FR-56):",
        "Normal",
        parent=rbac_parent,
    )
    auth_flow_note = insert_paragraph_after(
        auth_flow_intro,
        "Normal login uses Email as Username. OTP is never emailed for authentication — only SMS to "
        "registered mobile. If the user cannot receive OTP because the mobile is lost or changed, "
        "the pre-login recovery path uses the security question set at registration.",
        "Normal",
    )
    auth_flow_tbl = insert_table_after(
        auth_flow_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "Enter Email (Username) and Captcha", "User", "Email = Username (FR-04)"),
            ("2", "Validate Captcha; look up user by Email", "System", "FR-11"),
            ("3a", "Normal path — dispatch OTP to registered mobile only", "System", "FR-10; not to email"),
            ("3b", "Lost mobile — enter recovery; present security question", "User / System", "FR-56"),
            ("3c", "If answer correct — allow update of mobile number before login", "User / System", "Then OTP to new mobile"),
            ("4", "Enter OTP (+ Biometrics for DSR / Other Department)", "User", "FR-05 / FR-06 / FR-07"),
            ("5", "On success — continue (DSR: FR-52/53; then home)", "System", "Session established"),
            ("6", "After login — user may update mobile from profile if required", "User", "FR-13"),
        ],
    )

    login_post_intro = insert_heading_after(
        auth_flow_tbl,
        "Login post selection (DSR Officers — FR-52):",
        "Normal",
        parent=rbac_parent,
    )
    login_post_note = insert_paragraph_after(
        login_post_intro,
        "Applies only to DSR Officers. After Email (Username) + Captcha + OTP (to mobile) + Biometrics "
        "succeed, the system "
        "loads the officer's active sanctioned-post occupancies. If more than one is active, the user "
        "must choose which post to work under for this session. Each choice shows Post Name and Office "
        "details so dual-charge / multi-office officers can pick the correct context.",
        "Normal",
    )
    login_post_tbl = insert_table_after(
        login_post_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "Authenticate (Email + Captcha + OTP to mobile + Biometrics)", "DSR Officer / Auth", "FR-06"),
            ("2", "Load active post occupancies (Post + Office)", "System", "Vacant/ended occupancies excluded"),
            ("3a", "If exactly one active post — auto-select; continue to home", "System", "No selection UI"),
            ("3b", "If two or more active posts — show post selection list", "System / UI", "Mandatory; cannot skip (FR-52)"),
            ("4", "Display each option as Post Name + Office Name (+ Office Code)", "UI", "Office details mandatory on each row"),
            ("5", "User selects exactly one post and confirms", "DSR Officer", "Selection stored on session"),
            ("6", "Resolve candidate session roles from login-selected post via Post–Role mapping", "UM / Auth", "May still choose vacant subordinate (FR-53)"),
            ("7", "During login — Role + Post drop-down for vacant subordinates under login post", "DSR Officer / UI", "FR-53; before home"),
            ("8", "Lock session post; resolve Module Function claims; enter home; show Post + Role (FR-54)", "UM / Auth / UI", "No further post switch this session"),
        ],
    )

    switch_intro = insert_heading_after(
        login_post_tbl,
        "Vacant subordinate post choice during login only (DSR Officers — FR-53) and post display (FR-54):",
        "Normal",
        parent=rbac_parent,
    )
    switch_note = insert_paragraph_after(
        switch_intro,
        "Post switching / acting-under-vacant-subordinate shall happen only during login. After "
        "authentication and after the login-selected post is determined (FR-52), and before the home "
        "page, the system shall offer vacant posts under that post in the Hierarchy Master at the same "
        "Office. Options shall be shown in a drop-down where each entry displays the Role together with "
        "the Post (and Office). Cascading: only vacant children are offered; under each vacant child, "
        "vacant grandchildren are also offered. Occupied posts are not offered and stop further descent. "
        "The officer may stay on the original login-selected post. Once the user confirms and enters "
        "the application, the session post is fixed — the user cannot switch post after login. After "
        "login, the system shall show the logged-in Post with its mapped Role(s) (FR-54).",
        "Normal",
    )
    switch_tbl = insert_table_after(
        switch_note,
        ["Step", "Action", "Actor / System", "Notes"],
        [
            ("1", "Complete authentication; determine login-selected Post + Office", "DSR Officer / System", "FR-06, FR-52"),
            ("2", "Before home — read Hierarchy Master children of login-selected Post", "System", "FR-43; login-time only"),
            ("3", "For each child at same Office — if vacant, add as option", "System", "Vacancy = Occupied < Sanctioned strength (FR-24)"),
            ("4", "If child is vacant, repeat for its children (cascade)", "System", "Stop cascade under occupied posts"),
            ("5", "Show drop-down: each option = Role + Post (+ Office); include login post", "UI", "FR-53; e.g. \"Sub-Registrar (SR) — Sub-Registrar / Office A\""),
            ("6", "User selects one option from the drop-down and confirms", "DSR Officer", "May keep login post or choose vacant subordinate"),
            ("7", "Lock session post; derive roles via Post–Role mapping; enter home", "UM / Auth", "FR-38, FR-47"),
            ("8", "Display logged-in Post with mapped Role(s) on home/header", "UI", "FR-54; read-only context"),
            ("9", "After login — block any further post switch for this session", "System / UI", "Mandatory (FR-53); re-login to change"),
            ("10", "Audit log choice (login post, chosen post, role(s), office, actor, timestamp)", "System", "Mandatory"),
        ],
    )
    switch_ex_intro = insert_heading_after(
        switch_tbl,
        "Illustrative example (Office A — Sub-Registrar login; FR-53 drop-down during login):",
        "Normal",
        parent=rbac_parent,
    )
    switch_ex_tbl = insert_table_after(
        switch_ex_intro,
        ["Hierarchy (Office A)", "Vacancy", "Drop-down entry (Role + Post)"],
        [
            ("Sub-Registrar (login-selected)", "Occupied by self", "Sub-Registrar (SR) — Sub-Registrar / Office A"),
            ("→ FDA under Sub-Registrar", "Vacant", "FDA (Enforcement) — FDA (Enforcement) / Office A"),
            ("→ → SDA under FDA", "Vacant", "SDA (Enforcement) — SDA (Enforcement) / Office A"),
            ("→ FDA under Sub-Registrar", "Occupied", "Not listed; cascade under FDA blocked"),
            ("After home entered", "—", "Header shows chosen Post + mapped Role(s) (FR-54); no switch"),
        ],
    )

    display_intro = insert_heading_after(
        switch_ex_tbl,
        "Logged-in Post and Role display after login (FR-54):",
        "Normal",
        parent=rbac_parent,
    )
    display_note = insert_paragraph_after(
        display_intro,
        "On the home page and/or application header, the system shall show the session Post Name, "
        "Office, and the Role(s) mapped to that Post (Post–Role mapping). Example: "
        "\"Post: Sub-Registrar | Office: SRO Yeshwanthapura | Role: Sub-Registrar (SR)\". "
        "This display does not allow changing the post after login.",
        "Normal",
    )

    sub3 = insert_heading_after(
        display_note, "6.5.3 Posts Master, Post–Role Mapping, and Sanctioned Posts", "Heading 3", parent=rbac_parent
    )
    note3 = insert_paragraph_after(
        sub3,
        "Posts Master is separate from Role Master. Posts shall be division-specific where the "
        "organisation chart distinguishes them (FR-49), so that Post–Role mapping does not over-provision "
        "access across Admin, Vigilance, and Enforcement. Sanctioned Posts Master references Posts Master "
        "(Post + Office + strength). Post–Role mapping links each Post to its corresponding unique Role(s). "
        "DSR Officers are assigned vacant sanctioned posts; at login the officer selects one active "
        "post when multiple are assigned (FR-52). Session roles are those mapped to the current session "
        "post (login-selected or vacant subordinate chosen during login — FR-53; fixed thereafter). "
        "Over-capacity assignment is blocked (FR-45–FR-49).",
        "Normal",
    )
    posts_intro = insert_paragraph_after(note3, "Posts Master — seed rows (division-specific; Application Admin maintains):", "Normal")
    posts_tbl = insert_table_after(
        posts_intro,
        ["Post Code", "Post Name", "Division / Branch", "Status"],
        [
            ("POST-ACS-SEC", "Additional Chief Secretary / Principal Secretary / Secretary", "Secretariat", "Active"),
            ("POST-IGR", "Inspector General of Registration & Commissioner of Stamps", "Top Management", "Active"),
            ("POST-DIGR-ADMIN", "DIGR (Admin, Law & Computers)", "Admin, Law & Computers", "Active"),
            ("POST-AIGR-ADMIN", "AIGR (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-HQA-ADMIN", "HQA (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-SR-ADMIN", "Sub Registrar (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-ACCT-SUP", "Accountant Superintendent (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-FDA-ADMIN", "FDA (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-SDA-ADMIN", "SDA (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-TYPIST-ADMIN", "Typist (Admin)", "Admin, Law & Computers — Administration", "Active"),
            ("POST-HQA-RTI", "HQA (RTI)", "Admin, Law & Computers — RTI & Statistics", "Active"),
            ("POST-FDA-RTI", "FDA (RTI)", "Admin, Law & Computers — RTI & Statistics", "Active"),
            ("POST-SDA-RTI", "SDA (RTI)", "Admin, Law & Computers — RTI & Statistics", "Active"),
            ("POST-SI", "Statistical Inspector", "Admin, Law & Computers — RTI & Statistics", "Active"),
            ("POST-DIGR-VIG", "DIGR (Vigilance)", "Vigilance", "Active"),
            ("POST-LAW-OFF", "Law Officer", "Vigilance", "Active"),
            ("POST-AIGR-COMP", "AIGR (Computers)", "Computers", "Active"),
            ("POST-SI-INT", "System Integrator", "Computers — Development", "Active"),
            ("POST-APP-DEV", "Application Developer", "Computers — Development", "Active"),
            ("POST-PMU", "PMU", "Computers — Development", "Active"),
            ("POST-HQA-COMP", "HQA / Project Manager (Comp)", "Computers — Operations", "Active"),
            ("POST-SR-COMP", "Sub Registrar (Comp)", "Computers — Operations", "Active"),
            ("POST-FDA-COMP", "FDA (Computers)", "Computers — Operations", "Active"),
            ("POST-SDA-COMP", "SDA (Computers)", "Computers — Operations", "Active"),
            ("POST-DIGR-ENF", "DIGR (Enforcement)", "Enforcement", "Active"),
            ("POST-DRO", "District Registrar", "Enforcement", "Active"),
            ("POST-HQA-ENF", "HQA (Enforcement)", "Enforcement", "Active"),
            ("POST-SR", "Sub-Registrar", "Enforcement / Field", "Active"),
            ("POST-FDA-ENF", "FDA (Enforcement)", "Enforcement", "Active"),
            ("POST-SDA-ENF", "SDA (Enforcement)", "Enforcement", "Active"),
            ("POST-DEO", "Data Entry Operator", "Field / SRO", "Active"),
            ("POST-DIGR-INT", "DIGR (Intelligence)", "Intelligence & Audit", "Active"),
            ("POST-AIGR-AUDIT", "AIGR (Audit)", "Intelligence & Audit", "Active"),
            ("POST-HQA-AUDIT", "HQA (Audit)", "Intelligence & Audit", "Active"),
            ("POST-SUP-AUDIT", "Superintendent (Audit)", "Intelligence & Audit", "Active"),
            ("POST-FDA-AUDIT", "FDA (Audit)", "Intelligence & Audit", "Active"),
            ("POST-SDA-AUDIT", "SDA (Audit)", "Intelligence & Audit", "Active"),
            ("POST-TYPIST-AUDIT", "Typist (Audit)", "Intelligence & Audit", "Active"),
            ("POST-DIGR-CVC", "DIGR CVC", "DIGR CVC", "Active"),
            ("POST-JD-TP", "JD Town Planning", "DIGR CVC", "Active"),
        ],
    )
    map_pr_intro = insert_heading_after(
        posts_tbl,
        "Post–Role mapping — example rows (division-aligned; avoid cross-division over-provisioning):",
        "Normal",
        parent=rbac_parent,
    )
    post_role_tbl = insert_table_after(
        map_pr_intro,
        ["Post Code", "Post Name", "Mapped Role (unique name)", "Notes"],
        [
            ("POST-ACS-SEC", "ACS / Principal Secretary / Secretary", "ACS / Principal Secretary / Secretary", "1:1 — hierarchy root"),
            ("POST-IGR", "IGR & Commissioner of Stamps", "IGR", "1:1"),
            ("POST-DIGR-ADMIN", "DIGR (Admin, Law & Computers)", "DIGR (Admin, Law & Computers)", "Not mapped to Vigilance/Enforcement DIGR roles"),
            ("POST-DIGR-VIG", "DIGR (Vigilance)", "DIGR (Vigilance)", "1:1 — division-specific"),
            ("POST-DIGR-ENF", "DIGR (Enforcement)", "DIGR (Enforcement)", "1:1 — division-specific"),
            ("POST-DIGR-INT", "DIGR (Intelligence)", "DIGR (Intelligence)", "1:1"),
            ("POST-AIGR-ADMIN", "AIGR (Admin)", "AIGR (Admin)", "1:1"),
            ("POST-AIGR-COMP", "AIGR (Computers)", "AIGR (Computers)", "1:1"),
            ("POST-AIGR-AUDIT", "AIGR (Audit)", "AIGR (Audit)", "1:1"),
            ("POST-FDA-ADMIN", "FDA (Admin)", "FDA (Admin)", "Unique role — not shared FDA"),
            ("POST-FDA-ENF", "FDA (Enforcement)", "FDA (Enforcement)", "Unique role"),
            ("POST-FDA-AUDIT", "FDA (Audit)", "FDA (Audit)", "Unique role"),
            ("POST-SDA-ADMIN", "SDA (Admin)", "SDA (Admin)", "Unique role"),
            ("POST-SDA-ENF", "SDA (Enforcement)", "SDA (Enforcement)", "Unique role"),
            ("POST-SR", "Sub-Registrar", "Sub-Registrar (SR)", "Field signing post"),
            ("POST-DEO", "Data Entry Operator", "DEO", "1:1"),
            ("POST-DRO", "District Registrar", "DRO", "1:1"),
        ],
    )
    sanc_intro = insert_heading_after(
        post_role_tbl,
        "Sanctioned Posts Master — references Posts Master (example rows):",
        "Normal",
        parent=rbac_parent,
    )
    sanc_tbl = insert_table_after(
        sanc_intro,
        ["Office", "Post Code", "Post Name", "Sanctioned Strength", "Occupied", "Vacant"],
        [
            ("SRO Yeshwanthapura", "POST-SR", "Sub-Registrar", "1", "1", "0"),
            ("SRO Yeshwanthapura", "POST-FDA-ENF", "FDA (Enforcement)", "2", "1", "1"),
            ("SRO Yeshwanthapura", "POST-SDA-ENF", "SDA (Enforcement)", "1", "0", "1"),
            ("SRO Yeshwanthapura", "POST-DEO", "Data Entry Operator", "2", "1", "1"),
            ("DRO Mysuru", "POST-DRO", "District Registrar", "1", "1", "0"),
            ("DRO Mysuru", "POST-HQA-ENF", "HQA (Enforcement)", "1", "0", "1"),
            ("IGRO Bangalore", "POST-DIGR-ADMIN", "DIGR (Admin, Law & Computers)", "1", "1", "0"),
            ("IGRO Bangalore", "POST-DIGR-ENF", "DIGR (Enforcement)", "1", "0", "1"),
        ],
    )
    admin_posts_intro = insert_heading_after(
        sanc_tbl,
        "Application Admin maintenance:",
        "Normal",
        parent=rbac_parent,
    )
    admin_posts_tbl = insert_table_after(
        admin_posts_intro,
        ["Master / Table", "Application Admin may", "FR"],
        [
            ("Posts Master", "Add, edit, enable/disable posts", "FR-46"),
            ("Post–Role mapping", "Map one Post to one or more Roles; add/edit/remove mappings", "FR-47"),
            ("Sanctioned Posts Master", "Set strength per Post per Office; view occupied/vacant", "FR-24, FR-25, FR-48"),
        ],
    )

    sub4 = insert_heading_after(
        admin_posts_tbl, "6.5.4 Role Master — Seed Roles by Role Category", "Heading 3", parent=rbac_parent
    )
    note4 = insert_paragraph_after(
        sub4,
        "All roles below reside in the single Role Master, differentiated by Role Category. Seed roles "
        "are listed for Citizen, Other Department, and DSR (Department). Application Admin may add further "
        "roles under the same Role Categories. Other Department users are assigned exactly one role with "
        "Role Category = Other Department. DSR Officers receive roles only through Post–Role mapping "
        "after sanctioned post assignment. Sanctioned posts reference Posts Master and apply only to "
        "DSR establishment posts.",
        "Normal",
    )
    citizen_intro = insert_paragraph_after(
        note4,
        "Roles with Role Category = Citizen:",
        "Normal",
    )
    citizen_tbl = insert_table_after(
        citizen_intro,
        ["Role", "Description"],
        [
            ("Citizen", "Default role assigned on instant self-registration; access to citizen portal services"),
            ("Marriage Applicant", "Apply for Hindu Marriage / Special Marriage registration and related citizen actions"),
            ("Document Registration Applicant", "Initiate and track document registration applications"),
            ("Certified Copy / EC Applicant", "Apply for Encumbrance Certificate (EC) and certified copy of registered documents"),
            ("Firm / Society Applicant", "Apply for firm / society related registration services where offered on the portal"),
            ("Stamp Duty / Challan Payer", "Pay stamp duty / registration fees and view payment history for citizen transactions"),
        ],
    )
    od_intro = insert_heading_after(
        citizen_tbl,
        "Roles with Role Category = Other Department:",
        "Normal",
        parent=rbac_parent,
    )
    other_tbl = insert_table_after(
        od_intro,
        ["Role", "Typical Parent Department", "Description"],
        [
            ("Revenue Verification Officer", "Revenue", "Verify land / revenue particulars linked to registration applications"),
            ("Bhoomi Cross-check User", "Revenue", "Cross-check Bhoomi / RTC data against registration records"),
            ("Treasury Payment Verifier", "Treasury / Finance", "Verify treasury / payment status for registration fees"),
            ("Khajane Reconciliation User", "Treasury / Finance", "Reconcile Khajane-II receipts with Kaveri transactions"),
            ("Police Enquiry Officer", "Police", "View / respond to police enquiry requests on registered documents"),
            ("FIR Verification User", "Police", "Verify FIR / crime particulars where required for registration workflow"),
            ("ULB Document Verifier", "Urban Local Body (ULB)", "Verify municipal / ULB documents submitted with applications"),
            ("Read-only MIS Viewer", "Any (cross-department)", "Read-only access to assigned MIS / dashboards; no transactional actions"),
            ("Document Upload User", "Any (cross-department)", "Upload supporting documents for assigned inter-department workflows"),
            ("Inter-Department Enquiry User", "Any (cross-department)", "Raise / respond to inter-department enquiries on applications"),
        ],
    )
    dsr_intro = insert_heading_after(
        other_tbl,
        "Roles with Role Category = DSR (Department) — unique role names (no duplicate FDA/SDA/Typist):",
        "Normal",
        parent=rbac_parent,
    )
    dsr_roles_tbl = insert_table_after(
        dsr_intro,
        ["Division", "Role (unique name)", "Description"],
        [
            ("Secretariat", "ACS / Principal Secretary / Secretary", "Additional Chief Secretary / Principal Secretary / Secretary — hierarchy root"),
            ("Top Management", "IGR", "Inspector General of Registration & Commissioner of Stamps"),
            ("Division 1 — Admin, Law & Computers", "DIGR (Admin, Law & Computers)", "Deputy IGR for Admin, Law & Computers"),
            ("Division 1 — Admin, Law & Computers", "AIGR (Admin)", "Assistant IGR (Admin)"),
            ("Division 1 — Admin, Law & Computers", "HQA (Admin)", "Head Quarter Assistant (Admin)"),
            ("Division 1 — Admin, Law & Computers", "SRO (Admin)", "Sub Registrar (Admin)"),
            ("Division 1 — Admin, Law & Computers", "Accountant Superintendent (Admin)", "Accountant Superintendent (Admin)"),
            ("Division 1 — Admin, Law & Computers", "FDA (Admin)", "First Division Assistant — Admin"),
            ("Division 1 — Admin, Law & Computers", "SDA (Admin)", "Second Division Assistant — Admin"),
            ("Division 1 — Admin, Law & Computers", "Typist (Admin)", "Typist — Admin"),
            ("Division 1 — Admin, Law & Computers", "HQA (RTI)", "Head Quarter Assistant (RTI)"),
            ("Division 1 — Admin, Law & Computers", "FDA (RTI)", "First Division Assistant — RTI & Statistics"),
            ("Division 1 — Admin, Law & Computers", "SDA (RTI)", "Second Division Assistant — RTI & Statistics"),
            ("Division 1 — Admin, Law & Computers", "Statistical Inspector", "Statistical Inspector"),
            ("Division 2 — Vigilance", "DIGR (Vigilance)", "Deputy IGR (Vigilance)"),
            ("Division 2 — Vigilance", "Law Officer", "Departmental Law Officer"),
            ("Division 3 — Computers", "AIGR (Computers)", "Assistant IGR (Computers)"),
            ("Division 3 — Computers", "System Integrator", "System Integrator / SI support"),
            ("Division 3 — Computers", "PMU", "Project Management Unit"),
            ("Division 3 — Computers", "Application Developer", "Application development support"),
            ("Division 3 — Computers", "HQA / Project Manager (Comp)", "Head Quarter Assistant / Project Manager (Computers)"),
            ("Division 3 — Computers", "SRO (Comp)", "Sub Registrar (Computers)"),
            ("Division 3 — Computers", "FDA (Computers)", "First Division Assistant — Computers"),
            ("Division 3 — Computers", "SDA (Computers)", "Second Division Assistant — Computers"),
            ("Division 4 — Enforcement", "DIGR (Enforcement)", "Deputy IGR (Enforcement)"),
            ("Division 4 — Enforcement", "DRO", "District Registrar / DRO"),
            ("Division 4 — Enforcement", "HQA (Enforcement)", "Head Quarter Assistant (Enforcement)"),
            ("Division 4 — Enforcement", "Sub-Registrar (SR)", "Sub-Registrar (office head / signing)"),
            ("Division 4 — Enforcement", "FDA (Enforcement)", "First Division Assistant — Enforcement"),
            ("Division 4 — Enforcement", "SDA (Enforcement)", "Second Division Assistant — Enforcement"),
            ("Division 5 — Intelligence & Audit", "DIGR (Intelligence)", "Deputy IGR (Intelligence)"),
            ("Division 5 — Intelligence & Audit", "AIGR (Audit)", "Assistant IGR (Audit)"),
            ("Division 5 — Intelligence & Audit", "HQA (Audit)", "Head Quarter Assistant (Audit)"),
            ("Division 5 — Intelligence & Audit", "Superintendent (Audit)", "Superintendent (Audit)"),
            ("Division 5 — Intelligence & Audit", "FDA (Audit)", "First Division Assistant — Audit"),
            ("Division 5 — Intelligence & Audit", "SDA (Audit)", "Second Division Assistant — Audit"),
            ("Division 5 — Intelligence & Audit", "Typist (Audit)", "Typist — Audit"),
            ("Division 6 — CVC", "DIGR CVC", "Deputy IGR (CVC)"),
            ("Division 6 — CVC", "JD Town Planning", "Joint Director, Town Planning"),
            ("Field / SRO (common)", "DEO", "Data Entry Operator — SRO operational role"),
        ],
    )

    sub5 = insert_heading_after(
        dsr_roles_tbl, "6.5.5 DSR Officer Post Assignment", "Heading 3", parent=rbac_parent
    )
    primary_note = insert_paragraph_after(
        sub5,
        "DSR Officers are assigned one or more vacant sanctioned posts (Posts Master via Sanctioned "
        "Posts Master). There is no Primary or Secondary role/post assignment. At login, if multiple "
        "active posts exist, the officer selects one post shown with office details (FR-52); session "
        "roles = roles mapped to the session post fixed at login via Post–Role mapping. During login "
        "the officer may choose a vacant subordinate under that login post using the Hierarchy Master "
        "(FR-53); after login the post cannot be switched. Division-specific posts "
        "prevent over-provisioning across Admin/Vigilance/Enforcement (FR-47, FR-49). Other Department "
        "users and Citizens use Role Master directly (see Sections 6.5.1, 6.5.4 and 6.6.2).",
        "Normal",
    )
    primary_tbl = insert_table_after(
        primary_note,
        ["User Category", "What is assigned", "Roles in session", "End Date", "Notes"],
        [
            (
                "DSR Officer",
                "One or more vacant sanctioned posts (Post + Office)",
                "Roles from Post–Role mapping for session post fixed at login (FR-52 / FR-53)",
                "Optional per post occupancy",
                "Multi-post → choose at login; vacant subordinate option at login only; no mid-session switch",
            ),
            (
                "Other Department",
                "Exactly one role from Role Master",
                "That single role",
                "Optional account End Date",
                "No Posts Master / sanctioned posts",
            ),
            (
                "Citizen",
                "Citizen role(s) from Role Master (e.g. on self-registration)",
                "Assigned Citizen role(s)",
                "N/A",
                "No sanctioned posts",
            ),
        ],
    )

    sub6 = insert_heading_after(
        primary_tbl, "6.5.6 Module, Function, Resource Masters and Access Enforcement", "Heading 3", parent=rbac_parent
    )
    note6 = insert_paragraph_after(
        sub6,
        "DSR organisational roles (SR, FDA (Enforcement), DEO, etc.) are not named after application "
        "services. Access is modelled as: User → Role(s) → Module Function(s) → Resource(s) (API/URL). "
        "Registration of Documents is a single module (no online/offline classification). "
        "Role names in Role–Module–Function mapping must match Role Master exactly (FR-50). "
        "Application Admin is a system-level actor outside Role Master / Role–Module–Function mapping "
        "(FR-51) and maintains all masters in this section (FR-42).",
        "Normal",
    )
    hier_intro = insert_paragraph_after(
        note6,
        "Privilege hierarchy (masters):",
        "Normal",
    )
    hier_tbl = insert_table_after(
        hier_intro,
        ["Level", "Master", "Maintained by", "Purpose"],
        [
            ("1", "Module Master", "Application Admin", "Business modules (Registration of Documents, Marriage Registration, …)"),
            ("2", "Module Function Master", "Application Admin", "Functions under each module (VIEW, ADD, APPROVE, SIGN, …)"),
            ("3", "Resource Master", "Application Admin", "APIs and URLs linked to each Module Function"),
            ("4", "Role–Module–Function mapping", "Application Admin", "Which roles may perform which Module Functions"),
        ],
    )

    mod_intro = insert_heading_after(
        hier_tbl, "Module Master — seed rows:", "Normal", parent=rbac_parent
    )
    mod_tbl = insert_table_after(
        mod_intro,
        ["Module Code", "Module Name", "Description", "Status"],
        [
            ("MOD-DOC-REG", "Registration of Documents", "Registration of documents", "Active"),
            ("MOD-MARRIAGE", "Marriage Registration", "Marriage registration (Hindu Marriage and Special Marriage)", "Active"),
            ("MOD-EC", "Encumbrance Search", "Search and issue of Encumbrance Certificate (EC)", "Active"),
            ("MOD-CC", "Certified Copy", "Application and issue of certified copies of registered documents", "Active"),
            ("MOD-STAMP", "Stamp Duty / Payments", "Stamp duty assessment, challan, and payment reconciliation", "Active"),
            ("MOD-FIRM", "Firm / Society Registration", "Firm and society related registration services", "Active"),
            ("MOD-UM", "User Management", "Identity, roles, posts, modules, functions, resources, and mappings", "Active"),
            ("MOD-MIS", "MIS / Dashboards", "Operational and management reporting", "Active"),
        ],
    )

    fn_intro = insert_heading_after(
        mod_tbl, "Module Function Master — example rows:", "Normal", parent=rbac_parent
    )
    fn_tbl = insert_table_after(
        fn_intro,
        ["Function Code", "Module", "Function Name", "Description"],
        [
            ("FN-DOC-VIEW", "Registration of Documents", "VIEW", "View document registration applications / records"),
            ("FN-DOC-ADD", "Registration of Documents", "ADD", "Create / initiate document registration"),
            ("FN-DOC-EDIT", "Registration of Documents", "EDIT", "Edit draft / in-progress applications"),
            ("FN-DOC-APPROVE", "Registration of Documents", "APPROVE", "Approve application for registration"),
            ("FN-DOC-SIGN", "Registration of Documents", "SIGN", "Digitally sign registered document"),
            ("FN-DOC-PRINT", "Registration of Documents", "PRINT", "Print registration extracts / slips"),
            ("FN-MAR-VIEW", "Marriage Registration", "VIEW", "View marriage applications"),
            ("FN-MAR-ADD", "Marriage Registration", "ADD", "Create marriage application"),
            ("FN-MAR-APPROVE", "Marriage Registration", "APPROVE", "Approve / register marriage"),
            ("FN-EC-APPLY", "Encumbrance Search", "APPLY", "Apply for EC"),
            ("FN-EC-ISSUE", "Encumbrance Search", "ISSUE", "Issue / download EC"),
            ("FN-CC-APPLY", "Certified Copy", "APPLY", "Apply for certified copy"),
            ("FN-CC-ISSUE", "Certified Copy", "ISSUE", "Issue / download certified copy"),
            ("FN-UM-ADMIN", "User Management", "ADMIN", "Maintain users, roles, modules, functions, resources, mappings"),
        ],
    )

    res_intro = insert_heading_after(
        fn_tbl, "Resource Master (API / URL) — example rows:", "Normal", parent=rbac_parent
    )
    res_tbl = insert_table_after(
        res_intro,
        ["Resource Code", "Type", "Method", "Path / URL Pattern", "Module Function", "Status"],
        [
            ("RES-DOC-GET", "API", "GET", "/api/v1/documents/{id}", "FN-DOC-VIEW", "Active"),
            ("RES-DOC-LIST", "API", "GET", "/api/v1/documents", "FN-DOC-VIEW", "Active"),
            ("RES-DOC-CREATE", "API", "POST", "/api/v1/documents", "FN-DOC-ADD", "Active"),
            ("RES-DOC-UPDATE", "API", "PUT", "/api/v1/documents/{id}", "FN-DOC-EDIT", "Active"),
            ("RES-DOC-APPROVE", "API", "POST", "/api/v1/documents/{id}/approve", "FN-DOC-APPROVE", "Active"),
            ("RES-DOC-SIGN", "API", "POST", "/api/v1/documents/{id}/sign", "FN-DOC-SIGN", "Active"),
            ("RES-UI-DOC-VIEW", "URL", "GET", "/app/documents/view", "FN-DOC-VIEW", "Active"),
            ("RES-UI-DOC-CREATE", "URL", "GET", "/app/documents/create", "FN-DOC-ADD", "Active"),
            ("RES-MAR-CREATE", "API", "POST", "/api/v1/marriages", "FN-MAR-ADD", "Active"),
            ("RES-MAR-APPROVE", "API", "POST", "/api/v1/marriages/{id}/approve", "FN-MAR-APPROVE", "Active"),
            ("RES-EC-APPLY", "API", "POST", "/api/v1/encumbrance/applications", "FN-EC-APPLY", "Active"),
            ("RES-CC-ISSUE", "API", "POST", "/api/v1/certified-copies/{id}/issue", "FN-CC-ISSUE", "Active"),
            ("RES-UM-ROLES", "API", "PUT", "/api/v1/admin/roles/{id}/module-functions", "FN-UM-ADMIN", "Active"),
        ],
    )

    map_intro = insert_heading_after(
        res_tbl,
        "Role–Module–Function mapping — example rows (exact Role Master names only; FR-50):",
        "Normal",
        parent=rbac_parent,
    )
    map_note = insert_paragraph_after(
        map_intro,
        "Every Role column value must match an active Role Master name character-for-character. "
        "Document-registration functions for First Division Assistants are seeded for FDA (Enforcement) "
        "only (Enforcement / SRO document-registration path). FDA (Admin), FDA (RTI), FDA (Computers), "
        "and FDA (Audit) are not granted these functions unless Domain Expert adds explicit separate rows. "
        "Sub-Registrar (SR) and DEO match §6.5.4 unchanged. Application Admin is not listed here (FR-51).",
        "Normal",
    )
    map_tbl = insert_table_after(
        map_note,
        ["Role", "Role Category", "Module Function", "Allowed"],
        [
            ("Citizen", "Citizen", "FN-DOC-ADD", "Yes"),
            ("Citizen", "Citizen", "FN-DOC-VIEW", "Yes"),
            ("Citizen", "Citizen", "FN-EC-APPLY", "Yes"),
            ("Citizen", "Citizen", "FN-CC-APPLY", "Yes"),
            ("Document Registration Applicant", "Citizen", "FN-DOC-ADD", "Yes"),
            ("Document Registration Applicant", "Citizen", "FN-DOC-VIEW", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-VIEW", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-APPROVE", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-SIGN", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-DOC-PRINT", "Yes"),
            ("Sub-Registrar (SR)", "DSR", "FN-MAR-APPROVE", "Yes"),
            ("DEO", "DSR", "FN-DOC-VIEW", "Yes"),
            ("DEO", "DSR", "FN-DOC-ADD", "Yes"),
            ("DEO", "DSR", "FN-DOC-EDIT", "Yes"),
            ("DEO", "DSR", "FN-MAR-ADD", "Yes"),
            ("FDA (Enforcement)", "DSR", "FN-DOC-VIEW", "Yes"),
            ("FDA (Enforcement)", "DSR", "FN-DOC-ADD", "Yes"),
            ("FDA (Enforcement)", "DSR", "FN-DOC-EDIT", "Yes"),
            ("Revenue Verification Officer", "Other Department", "FN-DOC-VIEW", "Yes"),
        ],
    )

    app_admin_note = insert_heading_after(
        map_tbl,
        "Application Admin (system actor — Option B): Application Admin is not a Role Master seed role "
        "and is not assigned via sanctioned posts or the normal user-creation flow. It is seeded at "
        "deployment as a system principal with authority to maintain User Management masters "
        "(Module, Module Function, Resource, Role–Module–Function, Posts, Hierarchy, Sanctioned Posts, "
        "etc.). FN-UM-ADMIN and related Resource Master rows (e.g. RES-UM-ROLES) describe the capability "
        "surface for that principal; they are not granted through a Role–Module–Function mapping row "
        "labelled Application Admin | DSR (FR-51).",
        "Normal",
        parent=rbac_parent,
    )

    enf_intro = insert_heading_after(
        app_admin_note, "How the application enforces access (runtime):", "Normal", parent=rbac_parent
    )
    enf_tbl = insert_table_after(
        enf_intro,
        ["Step", "Action", "Actor / Component", "Notes"],
        [
            ("1", "User authenticates (Email + Captcha + OTP to mobile; + Biometrics if DSR/Other Dept)", "User / Auth service", "Credentials verified"),
            ("1a", "If DSR Officer with multiple active posts — select Post + Office for session", "User / UI", "Mandatory; display Post Name with Office details (FR-52); single post auto-selected"),
            ("1b", "During login — Role + Post drop-down for vacant subordinates (then lock session post)", "User / UI", "FR-53; no switch after home"),
            ("1c", "Show logged-in Post with mapped Role(s) on home/header", "UI", "FR-54; read-only"),
            ("2", "Load session role(s)", "Auth / UM service", "DSR: roles from session post fixed at login via Post–Role map; Other Department: one role"),
            ("3", "Resolve Role–Module–Function mappings into session claims (function codes)", "UM / Auth service", "Functions for session roles (FR-38)"),
            ("4", "User calls an API or opens a URL", "Client / Browser", "Method + path"),
            ("5", "Look up Resource Master for matching Type + Method + Path pattern", "API Gateway / Middleware", "If no resource found → deny or treat as public only if explicitly marked"),
            ("6", "Obtain required Module Function from the matched Resource", "API Gateway / Middleware", "e.g. POST /api/v1/documents/{id}/approve → FN-DOC-APPROVE"),
            ("7", "Allow if session claims include that Module Function; else HTTP 403", "API Gateway / Middleware", "FR-41; Application Admin system principal may hold FN-UM-ADMIN outside mapping (FR-51); failed attempts audit-logged"),
            ("8", "UI menus / buttons shown only for Module Functions present in session claims", "Front-end", "Same mapping; hide unauthorised screens"),
        ],
    )

    admin_intro = insert_heading_after(
        enf_tbl, "Application Admin maintenance (CRUD):", "Normal", parent=rbac_parent
    )
    admin_crud_tbl = insert_table_after(
        admin_intro,
        ["Master / Table", "Application Admin may", "Audit"],
        [
            ("Module Master", "Add, edit, enable/disable modules", "Mandatory (FR-42)"),
            ("Module Function Master", "Add, edit, enable/disable functions under a module", "Mandatory (FR-42)"),
            ("Resource Master", "Add, edit, enable/disable API and URL resources; link to Module Function", "Mandatory (FR-42)"),
            ("Role–Module–Function mapping", "Grant or revoke Module Functions for Role Master roles only (exact names; FR-50); not used for Application Admin (FR-51)", "Mandatory (FR-42)"),
            ("DSR Officer Hierarchy Master", "Add, edit, reorder, enable/disable hierarchy nodes (§6.5.7)", "Mandatory (FR-43)"),
            ("Office Hierarchy Master", "Add, edit, enable/disable offices; set parent (IGR → DRO → SRO) (§6.5.8)", "Mandatory (FR-59)"),
        ],
    )

    sub7 = insert_heading_after(
        admin_crud_tbl,
        "6.5.7 DSR Officer Hierarchy Master",
        "Heading 3",
        parent=rbac_parent,
    )
    note7 = insert_paragraph_after(
        sub7,
        "The system shall maintain the reporting hierarchy for Department officers (DSR) as a "
        "Hierarchy Master aligned to the Department organisational chart. Hierarchy nodes reference "
        "Posts from the Posts Master (not Roles) — reporting is position-based (FR-43). Each child "
        "Post has exactly one parent Post in the tree. The same parent–child links drive vacant "
        "subordinate post choice during login only (FR-53; no mid-session switch) and Transfer out / "
        "relieving by hierarchy superior (FR-57, FR-58). Application Admin shall add, edit, reorder, "
        "enable/disable, and update hierarchy nodes. The hierarchy does not replace Sanctioned Posts "
        "or Role–Module–Function access control.",
        "Normal",
    )
    hier_seed_intro = insert_paragraph_after(
        note7,
        "Seed hierarchy (Post → Parent Post) — uses unique Posts Master codes. Enforcement SRO chain "
        "is Sub-Registrar → FDA (Enforcement) → SDA (Enforcement) under DRO so FR-53 / FR-57 examples apply:",
        "Normal",
    )
    hier_seed_tbl = insert_table_after(
        hier_seed_intro,
        ["Level", "Division / Branch", "Post Code", "Post Name", "Reports To (Parent Post Code)"],
        [
            ("0", "Secretariat", "POST-ACS-SEC", "ACS / Principal Secretary / Secretary", "— (root)"),
            ("1", "Top Management", "POST-IGR", "IGR & Commissioner of Stamps", "POST-ACS-SEC"),
            ("2", "Admin, Law & Computers", "POST-DIGR-ADMIN", "DIGR (Admin, Law & Computers)", "POST-IGR"),
            ("3", "Admin — Administration", "POST-AIGR-ADMIN", "AIGR (Admin)", "POST-DIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-HQA-ADMIN", "HQA (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-SR-ADMIN", "Sub Registrar (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-ACCT-SUP", "Accountant Superintendent (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-FDA-ADMIN", "FDA (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-SDA-ADMIN", "SDA (Admin)", "POST-AIGR-ADMIN"),
            ("4", "Admin — Administration", "POST-TYPIST-ADMIN", "Typist (Admin)", "POST-AIGR-ADMIN"),
            ("3", "Admin — RTI & Statistics", "POST-HQA-RTI", "HQA (RTI)", "POST-DIGR-ADMIN"),
            ("4", "Admin — RTI & Statistics", "POST-FDA-RTI", "FDA (RTI)", "POST-HQA-RTI"),
            ("4", "Admin — RTI & Statistics", "POST-SDA-RTI", "SDA (RTI)", "POST-HQA-RTI"),
            ("4", "Admin — RTI & Statistics", "POST-SI", "Statistical Inspector", "POST-HQA-RTI"),
            ("2", "Vigilance", "POST-DIGR-VIG", "DIGR (Vigilance)", "POST-IGR"),
            ("3", "Vigilance", "POST-LAW-OFF", "Law Officer", "POST-DIGR-VIG"),
            ("2", "Computers", "POST-AIGR-COMP", "AIGR (Computers)", "POST-IGR"),
            ("3", "Computers — Development", "POST-SI-INT", "System Integrator", "POST-AIGR-COMP"),
            ("3", "Computers — Development", "POST-APP-DEV", "Application Developer", "POST-AIGR-COMP"),
            ("3", "Computers — Development", "POST-PMU", "PMU", "POST-AIGR-COMP"),
            ("3", "Computers — Operations", "POST-HQA-COMP", "HQA / Project Manager (Comp)", "POST-AIGR-COMP"),
            ("4", "Computers — Operations", "POST-SR-COMP", "Sub Registrar (Comp)", "POST-HQA-COMP"),
            ("4", "Computers — Operations", "POST-FDA-COMP", "FDA (Computers)", "POST-HQA-COMP"),
            ("4", "Computers — Operations", "POST-SDA-COMP", "SDA (Computers)", "POST-HQA-COMP"),
            ("2", "Enforcement", "POST-DIGR-ENF", "DIGR (Enforcement)", "POST-IGR"),
            ("3", "Enforcement", "POST-DRO", "District Registrar", "POST-DIGR-ENF"),
            ("3", "Enforcement", "POST-HQA-ENF", "HQA (Enforcement)", "POST-DIGR-ENF"),
            ("4", "Enforcement — under DRO (field)", "POST-SR", "Sub-Registrar", "POST-DRO"),
            ("5", "Enforcement — SRO", "POST-FDA-ENF", "FDA (Enforcement)", "POST-SR"),
            ("6", "Enforcement — SRO", "POST-SDA-ENF", "SDA (Enforcement)", "POST-FDA-ENF"),
            ("5", "Enforcement — SRO", "POST-DEO", "DEO", "POST-SR"),
            ("2", "Intelligence & Audit", "POST-DIGR-INT", "DIGR (Intelligence)", "POST-IGR"),
            ("3", "Intelligence & Audit", "POST-AIGR-AUDIT", "AIGR (Audit)", "POST-DIGR-INT"),
            ("4", "Intelligence & Audit", "POST-HQA-AUDIT", "HQA (Audit)", "POST-AIGR-AUDIT"),
            ("4", "Intelligence & Audit", "POST-SUP-AUDIT", "Superintendent (Audit)", "POST-AIGR-AUDIT"),
            ("4", "Intelligence & Audit", "POST-FDA-AUDIT", "FDA (Audit)", "POST-AIGR-AUDIT"),
            ("4", "Intelligence & Audit", "POST-SDA-AUDIT", "SDA (Audit)", "POST-AIGR-AUDIT"),
            ("4", "Intelligence & Audit", "POST-TYPIST-AUDIT", "Typist (Audit)", "POST-AIGR-AUDIT"),
            ("2", "DIGR CVC", "POST-DIGR-CVC", "DIGR CVC", "POST-IGR"),
            ("3", "DIGR CVC", "POST-JD-TP", "JD Town Planning", "POST-DIGR-CVC"),
        ],
    )
    hier_attr_intro = insert_heading_after(
        hier_seed_tbl,
        "Hierarchy Master attributes (Application Admin maintains):",
        "Normal",
        parent=rbac_parent,
    )
    hier_attr_tbl = insert_table_after(
        hier_attr_intro,
        ["Attribute", "Description", "Mandatory"],
        [
            ("Hierarchy Node ID", "System-generated unique identifier", "Yes (system)"),
            ("Post", "Post from Posts Master (unique Post Code)", "Yes"),
            ("Parent Node", "Immediate reporting parent node (null for root); parent must also be a Post", "No for root"),
            ("Division / Branch", "Division label (e.g. Enforcement, Computers)", "Yes"),
            ("Display Order", "Sort order among siblings under the same parent", "Yes"),
            ("Level", "Depth in the tree (0 = root)", "Yes (system/derived)"),
            ("Is Active", "Enable / disable node without deleting history", "Yes"),
            ("Effective From / To", "Optional validity of the hierarchy link", "No"),
        ],
    )
    hier_admin_intro = insert_heading_after(
        hier_attr_tbl,
        "Maintain hierarchy — Application Admin workflow:",
        "Normal",
        parent=rbac_parent,
    )
    hier_wf_tbl = insert_table_after(
        hier_admin_intro,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → DSR Officer Hierarchy", "Application Admin", "FR-43"),
            ("2", "View tree / list of hierarchy nodes by division", "System", "Filter by Division, Active status"),
            ("3", "Add node — select Post from Posts Master, parent Post node, division, display order", "Application Admin", "Post must exist; FR-43"),
            ("4", "Edit node — change parent, order, division, or active flag", "Application Admin", "Cannot create circular parent links"),
            ("5", "Disable node (soft) if post/structure superseded", "Application Admin", "History retained; audit logged"),
            ("6", "Save — hierarchy available for organisational views and admin span rules", "System", "FR-44 seed structure; each Post has one parent"),
        ],
    )

    sub8 = insert_heading_after(
        hier_wf_tbl,
        "6.5.8 Office Hierarchy Master",
        "Heading 3",
        parent=rbac_parent,
    )
    office_note = insert_paragraph_after(
        sub8,
        "The system shall maintain the hierarchy of DSR offices separately from the Officer (Post) "
        "Hierarchy. IGR Office is the Head Office. Under the Head Office are District Registrar "
        "Offices. Under each District Registrar Office are Sub-Registrar Offices (FR-59). This "
        "hierarchy scopes Transfer Out / Relieving so a superior sees only offices (and posts) under "
        "them (FR-57).",
        "Normal",
    )
    office_seed_intro = insert_paragraph_after(
        office_note,
        "Seed Office Hierarchy (illustrative):",
        "Normal",
    )
    office_seed_tbl = insert_table_after(
        office_seed_intro,
        ["Level", "Office Type", "Office Code (example)", "Office Name (example)", "Parent Office"],
        [
            ("0", "Head Office", "OFF-IGR", "IGR Office (Head Office)", "— (root)"),
            ("1", "District Registrar Office", "OFF-DRO-MYS", "DRO Mysuru", "OFF-IGR"),
            ("1", "District Registrar Office", "OFF-DRO-BLR", "DRO Bengaluru", "OFF-IGR"),
            ("2", "Sub-Registrar Office", "OFF-SRO-YESH", "SRO Yeshwanthapura", "OFF-DRO-BLR"),
            ("2", "Sub-Registrar Office", "OFF-SRO-JAY", "SRO Jayanagar", "OFF-DRO-BLR"),
            ("2", "Sub-Registrar Office", "OFF-SRO-MYS-E", "SRO Mysuru East", "OFF-DRO-MYS"),
        ],
    )
    office_span_intro = insert_heading_after(
        office_seed_tbl,
        "Office span for Transfer Out / Relieving (examples):",
        "Normal",
        parent=rbac_parent,
    )
    insert_table_after(
        office_span_intro,
        ["Actor session Office", "Offices shown for relieving", "Notes"],
        [
            ("IGR Head Office (OFF-IGR)", "All DRO offices and all SRO offices under them", "Full state span"),
            ("DRO Bengaluru (OFF-DRO-BLR)", "OFF-DRO-BLR + SRO Yeshwanthapura, SRO Jayanagar, …", "Only that district's SROs"),
            ("SRO Yeshwanthapura (OFF-SRO-YESH)", "OFF-SRO-YESH only", "Posts under actor at that SRO only"),
        ],
    )

    admin_heading = next(
        p for p in doc.paragraphs if p.text.strip() == "6.7 Administrative User Management"
    )
    wf_heading_el = OxmlElement("w:p")
    admin_heading._element.addprevious(wf_heading_el)
    wf_heading = Paragraph(wf_heading_el, admin_heading._parent)
    wf_heading.style = "Heading 2"
    wf_heading.add_run("6.6 User Creation and Role Assignment Workflow")

    wf_intro = insert_paragraph_after(
        wf_heading,
        "The system shall provide dedicated step-by-step workflows to assign access during user creation. "
        "For DSR Officers, assign one or more vacant sanctioned posts (roles via Post–Role mapping; "
        "no Primary/Secondary; at login the officer selects one post when multiple are active — FR-52). "
        "Transfer out / relieving is scoped to offices under the actor (§6.6.3, FR-57–FR-59). "
        "For Other Department users, select exactly one role.",
        "Normal",
    )

    dsr_sub = insert_heading_after(wf_intro, "6.6.1 DSR Officer User Creation with Post Assignment", "Heading 3")
    dsr_tbl = insert_table_after(
        dsr_sub,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → Add DSR Department User", "Admin", "Authorised admin role only (FR-02)"),
            ("2", "Enter user particulars (name, Email as Username, mobile, KGID, photo, ID proof)", "Admin", "Email = Username (FR-04)"),
            ("2a", "Select security question and capture answer", "Admin / Officer", "Mandatory (FR-55)"),
            ("3", "Assign one or more vacant sanctioned posts (Post + Office from Posts / Sanctioned Posts Masters)", "Admin", "At least one required; FR-17, FR-30, FR-45, FR-48"),
            ("4", "Optionally set End Date per post occupancy", "Admin", "If set, occupancy freed on date; FR-30"),
            ("5", "System shows roles available via Post–Role mapping for each assigned post", "System", "FR-47; login will use selected post only (FR-52)"),
            ("6", "Upload approval letter where applicable", "Admin", "Should"),
            ("7", "Capture biometrics", "Admin / Officer", "Mandatory for DSR users (FR-06)"),
            ("8", "Review post occupancies and mapped roles; confirm", "Admin", "Multi-post users choose post at each login"),
            ("9", "Save — account active; occupied count updated per assigned post", "System", "Blocked if no post or any post not vacant"),
        ],
    )

    other_sub = insert_heading_after(
        dsr_tbl, "6.6.2 Other Department User Creation with Role Assignment", "Heading 3", parent=rbac_parent
    )
    other_note = insert_paragraph_after(
        other_sub,
        "Other Department users are stored in the same User Master (User Category = Other Department). "
        "Exactly one role is selected from the Role Master filtered by Role Category = Other Department. "
        "There is no Primary/Secondary model for Other Department users. An optional End Date may be "
        "entered; if entered, the system shall deactivate the user on that date.",
        "Normal",
    )
    other_tbl = insert_table_after(
        other_note,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open User Management → Add Other Department User", "Admin", "Authorised admin role only (FR-03)"),
            ("2", "Enter user particulars (name, Email as Username, mobile, photo, ID proof)", "Admin", "Email = Username (FR-04); User Category = Other Department"),
            ("2a", "Select security question and capture answer", "Admin / User", "Mandatory (FR-55)"),
            ("3", "Enter parent department and designation", "Admin", "e.g. Revenue, Treasury, Police"),
            ("4", "Assign exactly one role from Role Master (Role Category = Other Department)", "Admin", "Mandatory; FR-29, FR-34"),
            ("5", "Optionally enter Account End Date", "Admin", "Not mandatory; FR-33"),
            ("6", "Upload authorisation letter / NOC from parent department", "Admin", "Should"),
            ("7", "Capture biometrics", "Admin / User", "Mandatory (FR-07)"),
            ("8", "Review role and End Date (if any) and confirm", "Admin", ""),
            ("9", "Save — account active with module access for the assigned role", "System", "Blocked if no role selected"),
            ("10", "If End Date was entered and is reached — deactivate user; block login", "System", "FR-33; audit logged"),
        ],
    )

    relieve_sub = insert_heading_after(
        other_tbl,
        "6.6.3 Transfer Out / Relieving Process (DSR Officers)",
        "Heading 3",
        parent=rbac_parent,
    )
    relieve_note = insert_paragraph_after(
        relieve_sub,
        "A DSR Officer may be relieved from currently assigned post occupancy(ies). When Transfer Out / "
        "Relieving is opened, the system shall list only offices (and post occupancies) under the "
        "actor's session Office per the Office Hierarchy Master — IGR Head Office → District Registrar "
        "Offices → Sub-Registrar Offices (FR-59). Within that office span, relieving is allowed only "
        "where the actor's session Post is the immediate parent of the target Post in the Officer "
        "Hierarchy (§6.5.7, FR-57). Capture Relieving Date and Relieving Order. Mapping is removed "
        "after 11:59 PM of the Relieving Date (FR-58).",
        "Normal",
    )
    relieve_ex_intro = insert_paragraph_after(
        relieve_note,
        "Examples (Office Hierarchy + Officer Hierarchy):",
        "Normal",
    )
    relieve_ex_tbl = insert_table_after(
        relieve_ex_intro,
        ["Actor (session)", "Offices shown", "May relieve (posts under actor)", "Example"],
        [
            (
                "District Registrar @ DRO Bengaluru",
                "DRO Bengaluru + its SRO offices only",
                "Sub-Registrar at those SROs (POST-SR reports to POST-DRO)",
                "DRO Bengaluru relieves Sub-Registrar of SRO Yeshwanthapura — not SRO Mysuru East",
            ),
            (
                "Sub-Registrar @ SRO Yeshwanthapura",
                "SRO Yeshwanthapura only",
                "FDA / DEO under Sub-Registrar at that SRO",
                "SR Yeshwanthapura relieves FDA at Yeshwanthapura — not another SRO",
            ),
            (
                "IGR / DIGR @ IGR Head Office",
                "All DRO and SRO offices under Head Office",
                "Posts that report to actor's session Post across those offices",
                "Head Office span sees statewide offices under IGR",
            ),
            (
                "FDA (Enforcement) @ SRO Yeshwanthapura",
                "SRO Yeshwanthapura only",
                "SDA under FDA at that SRO",
                "FDA relieves SDA at same office",
            ),
        ],
    )
    relieve_wf_intro = insert_heading_after(
        relieve_ex_tbl,
        "Relieving workflow:",
        "Normal",
        parent=rbac_parent,
    )
    insert_table_after(
        relieve_wf_intro,
        ["Step", "Action", "Actor", "Notes"],
        [
            ("1", "Open Transfer Out / Relieving", "Superior", "FR-57"),
            ("2", "System lists only offices in actor's office span (Office Hierarchy)", "System", "FR-59; own office + descendants"),
            ("3", "Within those offices, list post occupancies where actor Post is immediate parent", "System", "Officer Hierarchy FR-43"),
            ("4", "Select officer / post occupancy to relieve", "Superior", "Cannot see outside span"),
            ("5", "Enter Relieving Date and Relieving Order (upload if applicable)", "Superior", "Mandatory (FR-57)"),
            ("6", "Confirm relieving", "Superior", "Audit-logged"),
            ("7", "Until end of Relieving Date — occupancy remains active", "System", "Officer may still use post until day closes"),
            ("8", "After 11:59 PM of Relieving Date — de-allocate user–post mapping", "System", "FR-58; occupied count −1; post vacant"),
            ("9", "If no remaining post occupancies — block / limit login until new assignment", "System", "Per policy; audit"),
        ],
    )

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v3.9"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
