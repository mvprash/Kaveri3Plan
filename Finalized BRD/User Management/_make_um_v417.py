# -*- coding: utf-8 -*-
"""Build BRD_User_Management_v4.17.docx from v4.16.

Brief §1.1 Purpose, §1.2 Background, and §1.3 Scope (detail remains in Section 6).
"""
from __future__ import annotations

import copy
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(__file__).resolve().parent
SRC = BASE / "BRD_User_Management_v4.16.docx"
DST = BASE / "BRD_User_Management_v4.17.docx"

PURPOSE = (
    "This BRD defines business requirements for the KAVERI 3.0 User Management "
    "module — identity, passwordless authentication, sanctioned post occupancy, "
    "RBAC, and officer lifecycle for Citizens, DSR Officers, and Other Department "
    "users. It is the agreed basis for design, development, testing, and sign-off."
)

BACKGROUND = (
    "KAVERI 3.0 requires a single platform service to manage users, roles, posts, "
    "and module access for the department and its citizens. It replaces fragmented "
    "Kaveri 2.0 user administration with one User Master and one Role Master, "
    "OTP-only login (no passwords), post-based DSR access, and Application "
    "Admin–maintained privilege mapping. Detailed authentication, occupancy, "
    "transfer, absence, and RBAC rules are specified in Section 6."
)

SCOPE_ITEMS = [
    "User registration and profile management for Citizens, DSR Officers, and "
    "Other Department users (single User Master)",
    "Passwordless authentication, session policy, Citizen lost-mobile reset, DSR "
    "post selection, and additional charge (Sections 6.2–6.5)",
    "Unified Role Master; Posts and Sanctioned Posts masters; office and officer "
    "hierarchies; RBAC via Module, Function, and Resource mapping (Section 6.5)",
    "DSR lifecycle: post assignment, Transfer Out/In, occupancy refresh, Temporary "
    "Absence and Temporary Charge (Section 6.6)",
    "Administrative user management, audit logging, and reporting (Sections 6.7–8)",
]


def set_cell_text(cell: _Cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def add_table_row(table: Table, values: list[str]) -> None:
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    last_tr.addnext(new_tr)
    row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell_text(row.cells[i], val)


def find_revision_table(doc: Document) -> Table:
    for tbl in doc.tables:
        if (
            len(tbl.columns) >= 4
            and tbl.rows
            and tbl.rows[0].cells[0].text.strip() == "Version"
            and tbl.rows[0].cells[3].text.strip() in ("Description", "Summary of change")
        ):
            return tbl
    raise RuntimeError("revision history table not found")


def find_field_table(doc: Document) -> Table:
    for tbl in doc.tables:
        if (
            tbl.rows
            and tbl.rows[0].cells[0].text.strip() == "Field"
            and len(tbl.columns) == 2
        ):
            return tbl
    raise RuntimeError("field/value table not found")


def set_field(table: Table, field: str, value: str) -> None:
    for row in table.rows:
        if row.cells[0].text.strip() == field:
            set_cell_text(row.cells[1], value)
            return
    raise KeyError(field)


def find_section_paragraphs(doc: Document) -> dict[str, list[Paragraph]]:
    """Map section heading → content paragraphs until next Heading 1/2."""
    sections: dict[str, list[Paragraph]] = {}
    current: str | None = None
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if style.startswith("Heading") and text.startswith("1."):
            current = text
            sections[current] = []
            continue
        if current and text.startswith("3. Business"):
            break
        if current and current.startswith("1.") and text:
            sections[current].append(para)
    return sections


def build() -> Document:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))
    sections = find_section_paragraphs(doc)

    replace_paragraph_text(sections["1.1 Purpose"][0], PURPOSE)
    replace_paragraph_text(sections["1.2 Background"][0], BACKGROUND)

    scope_paras = sections["1.3 Scope"]
    # First paragraph is "In Scope:" — keep; replace following with consolidated list
    content = [p for p in scope_paras if p.text.strip() != "In Scope:"]
    for i, item in enumerate(SCOPE_ITEMS):
        if i < len(content):
            replace_paragraph_text(content[i], item)
        else:
            break
    for extra in content[len(SCOPE_ITEMS) :]:
        delete_paragraph(extra)

    fields = find_field_table(doc)
    set_field(fields, "Version", "4.17")
    set_field(fields, "Last updated", "2026-09-02")

    rev = find_revision_table(doc)
    add_table_row(
        rev,
        [
            "4.17",
            "02-Sep-2026",
            "Nandha Kumar",
            "§1.1 Purpose, §1.2 Background, and §1.3 Scope shortened; "
            "detailed requirements unchanged in Section 6",
        ],
    )

    core = doc.core_properties
    core.title = "BRD — User Management Module (KAVERI 3.0) v4.17"
    core.author = "Nandha Kumar"
    core.subject = "BRD-K3-UM-001"

    return doc


def main() -> None:
    doc = build()
    target = DST
    try:
        doc.save(target)
    except PermissionError:
        target = DST.with_name(DST.stem + "_unlocked" + DST.suffix)
        doc.save(target)
        print("ORIGINAL LOCKED (open in Word) — saved instead as:")
    claude_dir = BASE.parent.parent / "Claude"
    if claude_dir.is_dir():
        try:
            shutil.copy2(target, claude_dir / target.name)
            print(f"Mirrored: {claude_dir / target.name}")
        except Exception as exc:
            print(f"Claude mirror skipped: {exc}")
    print(f"{target} ({target.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
