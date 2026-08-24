"""Create BRD_Hindu_Marriage_v0.9.docx from v0.8 — add Special Marriage
(Intended Marriage) and Other Forms notice + registration To-Be processes.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Hindu_Marriage_v0.8.docx"
DST = BASE / "BRD_Hindu_Marriage_v0.9.docx"
DIAG = BASE / "Process Diagrams"

IMG = {
    "ima_notice_online": DIAG
    / "Special Marriage (Intended Marriage) Notice generation-Online.png",
    "ima_notice_offline": DIAG
    / "Special Marriage (Intended Marriage) Notice generation- offline.png",
    "ima_reg": DIAG / "Special Marriage (Intended Marriage) Marriage Registration.png",
    "of_notice_online": DIAG
    / "Special Marriage Other Forms Notice generation-Online.png",
    "of_notice_offline": DIAG
    / "Special Marriage Other Forms Notice generation- offline.png",
    "of_reg": DIAG / "Special Marriage Other Forms Marriage Registration.png",
}


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    # clear content
    for child in list(new_p):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
            # keep structure simple — remove all runs/bookmarks later
            pass
    # Remove all children except pPr
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> Table:
    """Insert a table immediately after paragraph; return Table."""
    # Create table at end, then move XML after paragraph
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    # Bold header
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    tbl = table._tbl
    # detach from end
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def insert_picture_after(paragraph: Paragraph, image_path: Path, width_inches: float = 6.2) -> Paragraph:
    """Insert a paragraph containing a picture after the given paragraph."""
    pic_para = insert_paragraph_after(paragraph, "", style="Normal")
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return pic_para


def _style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not _style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(
        f"Paragraph not found: exact={exact!r} contains={contains!r} heading_only={heading_only}"
    )


def find_para_index(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> int:
    for i, p in enumerate(doc.paragraphs):
        if heading_only and not _style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return i
        if contains is not None and contains in t:
            return i
    raise KeyError(
        f"Paragraph not found: exact={exact!r} contains={contains!r} heading_only={heading_only}"
    )


def para_before(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    idx = find_para_index(doc, exact=exact, contains=contains, heading_only=heading_only)
    if idx == 0:
        raise KeyError("No paragraph before target")
    return doc.paragraphs[idx - 1]


def replace_in_all(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_para_text(p, p.text.replace(old, new))
                        n += 1
    return n


def add_version_row(doc: Document, version: str, date: str, author: str, summary: str, approver: str) -> None:
    table = doc.tables[1]
    row = table.add_row()
    vals = [version, date, author, summary, approver]
    for i, v in enumerate(vals):
        row.cells[i].text = v


def set_doc_control_field(doc: Document, field: str, value: str) -> None:
    table = doc.tables[0]
    for row in table.rows:
        if row.cells[0].text.strip() == field:
            row.cells[1].text = value
            return
    raise KeyError(field)


def append_table_row(table: Table, values: list[str]) -> None:
    row = table.add_row()
    for i, v in enumerate(values):
        row.cells[i].text = v

def find_table_by_header(doc: Document, first_cell: str):
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == first_cell:
            return table
    raise KeyError(f"Table with header {first_cell!r} not found")



def build_section_block(
    doc: Document,
    after: Paragraph,
    heading: str,
    image_path: Path,
    figure_caption: str,
    intro: str,
    flow_intro: str,
    steps: list[list[str]],
    key_chars: str,
) -> Paragraph:
    """Insert a full To-Be process section after `after`; return last paragraph."""
    h = insert_paragraph_after(after, heading, style="Heading 3")
    pic = insert_picture_after(h, image_path)
    cap = insert_paragraph_after(pic, figure_caption, style="Normal")
    intro_p = insert_paragraph_after(cap, intro, style="Normal")
    flow_p = insert_paragraph_after(intro_p, flow_intro, style="Normal")
    insert_table_after(flow_p, steps, doc)
    # Find the table we just inserted — next sibling of flow_p
    # Key characteristics after the table: insert after flow_p's next (the tbl),
    # by creating para after flow then moving after tbl is awkward; insert after flow,
    # then move key para after table.
    key_p = insert_paragraph_after(flow_p, key_chars, style="Normal")
    # Move key_p after the table that follows flow_p
    # Structure now: flow_p -> key_p -> table  (wrong order) OR flow -> table -> key
    # insert_table_after puts table right after flow_p, which pushes key? 
    # Actually: insert_table_after(flow) does flow.addnext(tbl), so if key was after flow,
    # order depends on timing. We inserted table AFTER creating key? No — table then key.
    # Current order: flow -> table -> key. Good if we inserted table before key.
    # Wait: we did insert_table_after(flow) THEN insert_paragraph_after(flow, key).
    # insert_paragraph_after(flow) inserts immediately after flow, BEFORE table.
    # So order is: flow -> key -> table. Fix by moving key after table.
    tbl = flow_p._p.getnext()
    # If next is key_p, then table is after that
    if tbl is key_p._p:
        tbl = key_p._p.getnext()
    if tbl is not None and tbl.tag == qn("w:tbl"):
        key_p._p.getparent().remove(key_p._p)
        tbl.addnext(key_p._p)
    blank = insert_paragraph_after(key_p, "", style="Normal")
    return blank


def main() -> None:
    for key, path in IMG.items():
        if not path.exists():
            raise FileNotFoundError(path)

    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # --- Document control ---
    set_doc_control_field(doc, "Version", "0.9")
    set_doc_control_field(
        doc,
        "Legal basis (primary)",
        "The Hindu Marriage Act, 1955; The Special Marriage Act, 1954 (Central Act 43 of 1954)",
    )
    set_doc_control_field(
        doc,
        "State rules (primary)",
        "Registration of Hindu Marriage (Karnataka) Rules, 1966; Special Marriage (Karnataka) Rules, 1961",
    )
    set_doc_control_field(
        doc,
        "Related inputs",
        "Acts_Rules/Marriage/Hindu Marriage Act, 1955.pdf; Acts_Rules/Marriage/The Special Marriage Act, 1954.pdf; "
        "Acts_Rules/Marriage/REGISTRATIONOFHINDUMARRIAGE_KARNATAKARULES_1966.docx; "
        "Acts_Rules/Marriage/SpecialMarriage(Karnataka)Rules1961.pdf; "
        "Acts_Rules/Marriage/SpecialMarriageFees.docx; Acts_Rules/Marriage/hindu marriage forms.pdf; "
        "Acts_Rules/Marriage/RD48MNMU2023-Notification-marriage.pdf",
    )
    set_doc_control_field(doc, "Last updated", "2026-08-24")
    add_version_row(
        doc,
        "0.9",
        "2026-08-24",
        "Nandha Kumar",
        "Added Special Marriage (Intended Marriage) notice Online/Offline + marriage registration, "
        "and Special Marriage Other Forms notice Online/Offline + marriage registration To-Be "
        "processes with approved diagrams; legal refs to SMA 1954 & Karnataka Rules 1961",
        "Prashanth",
    )

    # Title
    set_para_text(
        find_para(doc, exact="Marriage Registration Module — Hindu Marriage (Kaveri 3.0)"),
        "Marriage Registration Module — Hindu Marriage & Special Marriage (Kaveri 3.0)",
    )

    # Related process flows link
    for row in doc.tables[2].rows:
        if "PROC-K3-MRG-HMA-TOBE-001" in row.cells[0].text:
            row.cells[2].text = (
                "Process Diagrams/hindu marriage Online.png, Process Diagrams/hindu marriage Offline.png, "
                "Process Diagrams/Special Marriage (Intended Marriage) Notice generation-Online.png, "
                "Process Diagrams/Special Marriage (Intended Marriage) Notice generation- offline.png, "
                "Process Diagrams/Special Marriage (Intended Marriage) Marriage Registration.png, "
                "Process Diagrams/Special Marriage Other Forms Notice generation-Online.png, "
                "Process Diagrams/Special Marriage Other Forms Notice generation- offline.png, "
                "Process Diagrams/Special Marriage Other Forms Marriage Registration.png"
            )

    # --- Contents TOC ---
    toc_map = [
        (
            "7.5 Application status model (channel-aware)",
            "7.11 Application status model (channel-aware)",
        ),
    ]
    for old, new in toc_map:
        replace_in_all(doc, old, new)

    # Insert TOC entries after 7.4 line
    toc_74 = find_para(doc, exact="7.4 To-Be process — Hindu Marriage Offline")
    toc_entries = [
        "7.5 To-Be process — Special Marriage (Intended Marriage) Notice Generation Online",
        "7.6 To-Be process — Special Marriage (Intended Marriage) Notice Generation Offline",
        "7.7 To-Be process — Special Marriage (Intended Marriage) Marriage Registration",
        "7.8 To-Be process — Special Marriage Other Forms Notice Generation Online",
        "7.9 To-Be process — Special Marriage Other Forms Notice Generation Offline",
        "7.10 To-Be process — Special Marriage Other Forms Marriage Registration",
    ]
    anchor = toc_74
    for entry in toc_entries:
        anchor = insert_paragraph_after(anchor, entry, style="Normal")

    # Also update Contents for legal sections
    toc_33 = find_para(doc, exact="3.3 Statutory forms mapping")
    legal_toc = [
        "3.4 Primary legislation — Special Marriage Act, 1954",
        "3.5 Special Marriage (Karnataka) Rules, 1961",
        "3.6 Special Marriage statutory forms mapping",
    ]
    anchor = toc_33
    for entry in legal_toc:
        anchor = insert_paragraph_after(anchor, entry, style="Normal")

    # --- Executive summary — light touch ---
    for p in doc.paragraphs:
        if p.text.startswith("This document presents a comprehensive assessment of the existing Hindu Marriage"):
            set_para_text(
                p,
                "This document presents a comprehensive assessment of the existing Hindu Marriage "
                "Registration process and Special Marriage (Intended Marriage / Other Forms) notice "
                "and registration processes, covering both online and offline service channels. It "
                "identifies key operational, procedural, and user experience challenges currently "
                "impacting applicants and Sub-Registrar offices.",
            )
            break

    # --- Scope 2.1 ---
    scope_heading = find_para(doc, exact="2.1 In scope (Hindu Marriage — Phase [1])")
    set_para_text(scope_heading, "2.1 In scope (Hindu Marriage & Special Marriage — Phase [1])")
    # Update TOC scope line
    replace_in_all(
        doc,
        "2.1 In scope (Hindu Marriage — Phase [1])",
        "2.1 In scope (Hindu Marriage & Special Marriage — Phase [1])",
    )

    sma_scope = [
        "Special Marriage — Intended Marriage (Chapter II, SMA 1954): notice of intended marriage "
        "(Second Schedule) Online & Offline, 30-day publication / objection period, solemnization "
        "and certificate (Fourth Schedule) within statutory validity (diagram gate ≥30 and ≤90 days; "
        "aligns to Sec. 7 & Sec. 14).",
        "Special Marriage — Other Forms (Chapter III, SMA 1954 Sec. 15–16): notice/publication for "
        "registration of marriages celebrated in other forms, objection handling, and entry of "
        "certificate in Marriage Certificate Book (Fifth Schedule).",
        "Special Marriage process diagrams (§7.5–7.10): Notice Generation Online/Offline and "
        "Marriage Registration for both Intended Marriage and Other Forms.",
        "Special Marriage fees as per Special Marriage (Karnataka) Rules, 1961 / SpecialMarriageFees.docx.",
    ]
    anchor = para_before(doc, exact="2.2 Out of scope (unless PO promotes)", heading_only=True)
    for bullet in sma_scope:
        anchor = insert_paragraph_after(anchor, bullet, style="List Bullet")

    # Update Hindu process diagram scope bullet if present
    for p in doc.paragraphs:
        if p.text.startswith("Two processing channels per approved process diagrams (7.3–7.4)"):
            set_para_text(
                p,
                "Hindu Marriage processing channels per approved process diagrams (§7.3–7.4): "
                "Hindu Marriage Online and Hindu Marriage Offline.",
            )
            break

    # --- Assumptions: Marriage Officers ---
    for row in doc.tables[3].rows:
        if row.cells[0].text.strip() == "A-04":
            row.cells[1].text = (
                "Sub-Registrars act as Marriage Officers under SMA 1954 / Special Marriage "
                "(Karnataka) Rules, 1961 for Special Marriage services in Kaveri 3.0"
            )
            row.cells[2].text = "Domain Expert"
            break

    # --- Legal sections 3.4–3.6 ---
    insert_after = para_before(doc, exact="4. Stakeholders and actors", heading_only=True)

    h34 = insert_paragraph_after(
        insert_after,
        "3.4 Primary legislation — Special Marriage Act, 1954 (selected sections)",
        style="Heading 3",
    )
    intro34 = insert_paragraph_after(
        h34,
        "Source: Acts_Rules/Marriage/The Special Marriage Act, 1954.pdf. Applicable to solemnization "
        "of special marriages (Chapter II) and registration of marriages celebrated in other forms "
        "(Chapter III).",
        style="Normal",
    )
    sma_act_rows = [
        ["Section", "Topic", "System implication"],
        [
            "Sec. 4",
            "Conditions relating to solemnization",
            "Eligibility / prerequisite checks before notice",
        ],
        [
            "Sec. 5",
            "Notice of intended marriage (Second Schedule)",
            "Generate & capture notice; jurisdiction: ≥30 days residence of at least one party",
        ],
        [
            "Sec. 6",
            "Marriage Notice Book and publication",
            "Enter notice; publish on portal and/or notice board; transmit if party resides elsewhere",
        ],
        [
            "Sec. 7",
            "Objection to marriage (30 days)",
            "30-day countdown from publication; accept/record objections",
        ],
        [
            "Sec. 8–9",
            "Procedure / enquiry on objection",
            "SR enquiry, summon parties; uphold or reject objection",
        ],
        [
            "Sec. 11",
            "Declaration by parties and witnesses",
            "Declarations before solemnization; three witnesses",
        ],
        [
            "Sec. 12–13",
            "Solemnization and certificate of marriage",
            "Office solemnization; Fourth Schedule certificate",
        ],
        [
            "Sec. 14",
            "New notice if not solemnized within three months",
            "Block registration outside validity; require fresh notice",
        ],
        [
            "Sec. 15",
            "Registration of marriages celebrated in other forms",
            "Other Forms eligibility (ceremony performed; ages; residence ≥30 days; etc.)",
        ],
        [
            "Sec. 16",
            "Procedure for registration (Other Forms)",
            "Public notice, 30-day objections, Fifth Schedule certificate + 3 witnesses",
        ],
    ]
    insert_table_after(intro34, sma_act_rows, doc)

    insert_after = para_before(doc, exact="4. Stakeholders and actors", heading_only=True)
    h35 = insert_paragraph_after(
        insert_after,
        "3.5 Special Marriage (Karnataka) Rules, 1961",
        style="Heading 3",
    )
    p35 = insert_paragraph_after(
        h35,
        "Source: Acts_Rules/Marriage/SpecialMarriage(Karnataka)Rules1961.pdf and "
        "Acts_Rules/Marriage/SpecialMarriageFees.docx. Rules prescribe Marriage Officer procedures, "
        "notice publication, forms, and fee schedule for Special Marriage services in Karnataka. "
        "System shall apply the notified fee schedule for first payment (notice) and second payment "
        "(registration / solemnization) as shown in process diagrams.",
        style="Normal",
    )

    insert_after = para_before(doc, exact="4. Stakeholders and actors", heading_only=True)
    h36 = insert_paragraph_after(
        insert_after,
        "3.6 Special Marriage statutory forms mapping",
        style="Heading 3",
    )
    intro36 = insert_paragraph_after(
        h36,
        "Schedules under the Special Marriage Act, 1954 (as applicable in Kaveri 3.0):",
        style="Normal",
    )
    sma_form_rows = [
        ["Form / Schedule", "Act ref", "Purpose", "Generated by"],
        [
            "Notice of Intended Marriage (Second Schedule)",
            "Sec. 5 / Sec. 6",
            "Statutory notice for Intended Marriage",
            "System + SR generation; portal / notice-board publication",
        ],
        [
            "Declarations (Third Schedule)",
            "Sec. 11",
            "Declarations by parties and witnesses before solemnization",
            "System / office capture",
        ],
        [
            "Certificate of Marriage (Fourth Schedule)",
            "Sec. 13",
            "Certificate after solemnization under Chapter II",
            "System on SR DSC",
        ],
        [
            "Certificate of Marriage (Fifth Schedule)",
            "Sec. 16",
            "Certificate for marriages registered under Chapter III (Other Forms)",
            "System on SR DSC",
        ],
    ]
    insert_table_after(intro36, sma_form_rows, doc)

    # --- Glossary additions (table 6) ---
    glossary = find_table_by_header(doc, "Term")
    glossary_rows = [
        [
            "Special Marriage (Intended Marriage) Notice",
            "Chapter II notice of intended marriage under SMA 1954 Sec. 5–7; Online or Offline notice generation before solemnization",
            "Process Diagrams/Special Marriage (Intended Marriage) Notice generation-*.png",
        ],
        [
            "Special Marriage (Intended Marriage) Registration",
            "Post-notice solemnization & certificate issuance when timeline ≥30 and ≤90 days and no valid objection",
            "Process Diagrams/Special Marriage (Intended Marriage) Marriage Registration.png",
        ],
        [
            "Special Marriage Other Forms",
            "Chapter III registration of marriages celebrated in other forms (SMA Sec. 15–16), with notice publication and objection period",
            "Process Diagrams/Special Marriage Other Forms *.png",
        ],
        [
            "Marriage Notice Book",
            "Statutory book in which notices under Sec. 5 / Sec. 16 are entered and open to inspection (Sec. 6)",
            "SMA 1954",
        ],
        [
            "30-day countdown",
            "Statutory objection / waiting period from notice publication (Sec. 7 / Sec. 16)",
            "SMA 1954; process diagrams",
        ],
        [
            "First Payment / Second Payment",
            "Fee collected at notice stage (first) and at registration/solemnization stage (second) per Karnataka Special Marriage fee schedule",
            "SpecialMarriageFees.docx; process diagrams",
        ],
    ]
    for row in glossary_rows:
        append_table_row(glossary, row)

    # --- Channel model table T8 ---
    channel = find_table_by_header(doc, "Service Type")
    channel_rows = [
        [
            "Special Marriage (Intended Marriage) Notice — Online",
            "e-KYC bride & bridegroom, document upload, eSign, First Payment, appointment, portal notice display",
            "SR verification, notice generation, DEO select",
            "Online notice",
        ],
        [
            "Special Marriage (Intended Marriage) Notice — Offline",
            "Capture details, document upload, First Payment, appointment",
            "SR verification, notice generation, DEO photo/print/sign/scan/upload, paste on notice board",
            "Offline notice",
        ],
        [
            "Special Marriage (Intended Marriage) Marriage Registration",
            "Select approved notice, Second Payment, e-KYC witnesses, schedule visit, download certificate",
            "Objection enquiry (if any), SR verification, solemnization, DEO certificate & signatures, SR DSC",
            "Registration",
        ],
        [
            "Special Marriage Other Forms Notice — Online",
            "Same Online notice pattern as Intended Marriage notice generation (current approved diagram)",
            "SR verification, notice generation, DEO select; portal publication + 30-day countdown",
            "Online notice",
        ],
        [
            "Special Marriage Other Forms Notice — Offline",
            "Same Offline notice pattern as Intended Marriage notice generation (current approved diagram)",
            "SR verification, notice generation, DEO physical publication + 30-day countdown",
            "Offline notice",
        ],
        [
            "Special Marriage Other Forms Marriage Registration",
            "Select approved notice (timeline ≥30 & ≤90 days), Second Payment, e-KYC witnesses, schedule visit",
            "Objection enquiry (if any), SR verification, registration/solemnization steps per diagram, DEO certificate, SR DSC",
            "Registration",
        ],
    ]
    for row in channel_rows:
        append_table_row(channel, row)

    # Channel model narrative
    for p in doc.paragraphs:
        if p.text.startswith("The Hindu Marriage Registration service supports two delivery channels"):
            set_para_text(
                p,
                "The Marriage Registration module supports Hindu Marriage (Online / Offline) and "
                "Special Marriage services — Intended Marriage notice (Online / Offline) followed by "
                "marriage registration, and Other Forms notice (Online / Offline) followed by marriage "
                "registration. Notice channels are initiated through the citizen portal; registration "
                "continues from an approved / published notice after the statutory waiting period.",
            )
            break

    # Common intake channel selection
    for p in doc.paragraphs:
        if "Select channel: Hindu Marriage Online / Hindu Marriage Offline" in p.text:
            set_para_text(
                p,
                "Select channel / service path: Hindu Marriage Online / Hindu Marriage Offline; "
                "or Special Marriage (Intended Marriage) Notice Online / Offline; "
                "or Special Marriage Other Forms Notice Online / Offline. "
                "Channel is chosen before prerequisites (per updated diagrams). "
                "Marriage Registration for Special Marriage is initiated later by selecting an "
                "approved notice (see §7.7 / §7.10).",
            )
            break

    # Step tables for Special Marriage To-Be sections
    ima_online_steps = [
        ["#", "Step", "Lane", "Notes"],
        [
            "8",
            "If Aadhaar available → e-KYC on Bride & Bridegroom details",
            "System / Citizen",
            "Per Online notice diagram",
        ],
        [
            "9",
            "Review summary and proceed document uploading",
            "System / Citizen",
            "Summary of captured particulars",
        ],
        [
            "10",
            "Upload Identity Proof, Photo, Age Proof, Address Proof (Bridegroom & Bride)",
            "Citizen",
            "Mandatory supporting documents for notice",
        ],
        ["11", "Proceed with eSign", "System / Citizen", "Citizen eSign on notice application"],
        [
            "12",
            "SR Verification (decision)",
            "Sub Registrar",
            "Approve or Reject",
        ],
        [
            "12a",
            "Reject → return to e-KYC / review / upload",
            "Sub Registrar → Citizen",
            "Citizen corrects and resubmits",
        ],
        [
            "13",
            "First Payment",
            "System / Citizen",
            "Notice fee per Karnataka Special Marriage fee schedule",
        ],
        [
            "14",
            "Schedule appointment with SR",
            "System / Citizen",
            "After first payment",
        ],
        [
            "15",
            "SR Generates Notice",
            "Sub Registrar",
            "Statutory notice (Second Schedule / Sec. 5–6)",
        ],
        ["16", "Selects DEO", "Sub Registrar", "Assign publication / office follow-up"],
        [
            "17",
            "Marriage notice displayed in portal",
            "System",
            "Sec. 6 publication (digital)",
        ],
        [
            "18",
            "30-day countdown starts",
            "System",
            "Objection period per Sec. 7",
        ],
    ]

    ima_offline_steps = [
        ["#", "Step", "Lane", "Notes"],
        [
            "8",
            "If Aadhaar unavailable → capture Bride & Bridegroom details",
            "System / Citizen",
            "Manual capture path on Offline diagram",
        ],
        [
            "9",
            "Review summary and proceed document uploading",
            "System / Citizen",
            "",
        ],
        [
            "10",
            "Upload Identity / Age / Address proofs (Bridegroom & Bride)",
            "Citizen",
            "Mandatory documents",
        ],
        ["11", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
        [
            "11a",
            "Reject → return to capture Bride / Bridegroom details",
            "Sub Registrar → Citizen",
            "Resubmit",
        ],
        ["12", "First Payment", "System / Citizen", "Notice fee"],
        ["13", "Schedule appointment with SR", "System / Citizen", ""],
        [
            "14",
            "SR Generates Notice",
            "Sub Registrar",
            "Statutory notice generation",
        ],
        ["15", "Selects DEO", "Sub Registrar", "Assign FDA/SDA/DEO"],
        [
            "16",
            "Capture individual photos of Bride & Bridegroom",
            "FDA / SDA / DEO",
            "Office visit activity",
        ],
        [
            "17",
            "Download, Print, Sign, Scan and Upload notice",
            "FDA / SDA / DEO",
            "Physical notice processed into system",
        ],
        [
            "18",
            "Paste form on respective Notice Board",
            "FDA / SDA / DEO",
            "Sec. 6(2) conspicuous place publication",
        ],
        [
            "19",
            "30-day countdown starts",
            "System",
            "Objection period per Sec. 7",
        ],
    ]

    ima_reg_steps = [
        ["#", "Step", "Lane", "Notes"],
        ["1", "Citizen Login portal", "Citizen", "Continue from published notice"],
        ["2", "Select Notice", "Citizen", "Choose approved / published notice"],
        [
            "3",
            "Validate timeline ≥ 30 days and ≤ 90 days",
            "System",
            "If NO → No Action allowed (Sec. 7 / Sec. 14)",
        ],
        [
            "4",
            "If any Objection?",
            "System",
            "Branch to SR enquiry when objection exists",
        ],
        [
            "4a",
            "Conduct enquiry by summoning all parties",
            "Sub Registrar",
            "Sec. 8–9",
        ],
        [
            "4b",
            "Valid objection → update reason; Notice removal from portal (Objected)",
            "Sub Registrar / System",
            "Process stops for solemnization",
        ],
        [
            "4c",
            "Objection invalid → continue",
            "Sub Registrar",
            "Proceed to Second Payment",
        ],
        [
            "5",
            "Second Payment",
            "System / Citizen",
            "Registration / solemnization fee",
        ],
        ["6", "e-KYC on Witness details", "System / Citizen", "Three witnesses (Sec. 11)"],
        ["7", "Schedule Visit", "System / Citizen", "Office visit for solemnization"],
        ["8", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
        [
            "8a",
            "Reject → return to Schedule Visit",
            "Sub Registrar → System",
            "Reschedule / correct",
        ],
        [
            "9",
            "Marriage solemnization",
            "Sub Registrar",
            "Sec. 12; declarations Sec. 11",
        ],
        ["10", "Assigns to DEO", "Sub Registrar", "Certificate production"],
        ["11", "Joint Photo capturing", "FDA / SDA / DEO", ""],
        [
            "12",
            "Generate Marriage Certificate",
            "FDA / SDA / DEO",
            "Fourth Schedule",
        ],
        [
            "13",
            "Capture signs of Bride, Bridegroom, Witness",
            "FDA / SDA / DEO",
            "",
        ],
        ["14", "Upload signed copy", "FDA / SDA / DEO", ""],
        ["15", "Digital Signature (DSC)", "Sub Registrar", ""],
        [
            "16",
            "Marriage Certificate Issued",
            "Citizen / System",
            "Download from portal",
        ],
    ]

    # Other Forms notice — same step pattern as Intended (diagrams currently identical hashes)
    of_online_steps = [row[:] for row in ima_online_steps]
    of_online_steps[1][3] = (
        "Per Other Forms Online notice diagram (mirrors Intended Marriage Online notice)"
    )
    of_offline_steps = [row[:] for row in ima_offline_steps]
    of_offline_steps[1][3] = (
        "Per Other Forms Offline notice diagram (mirrors Intended Marriage Offline notice)"
    )

    of_reg_steps = [row[:] for row in ima_reg_steps]
    # Adjust certificate schedule note for Other Forms
    for row in of_reg_steps:
        if "Fourth Schedule" in row[3]:
            row[3] = "Fifth Schedule (Sec. 16)"
        if row[1] == "Marriage solemnization":
            row[1] = "Marriage registration / solemnization (per Other Forms diagram)"
            row[3] = "Chapter III Sec. 15–16; diagram includes SR registration step"

    sections = [
        {
            "heading": "7.5 To-Be process — Special Marriage (Intended Marriage) Notice Generation Online",
            "img": IMG["ima_notice_online"],
            "caption": "Figure: Special Marriage (Intended Marriage) Notice Generation — Online",
            "intro": (
                "Statutory basis: SMA 1954 Sec. 5–7 (Notice of intended marriage; Marriage Notice Book "
                "and publication; Objection). Citizen applies Online for notice of intended marriage; "
                "SR generates notice; system publishes on portal and starts 30-day countdown."
            ),
            "flow": "Flow (continuing from §7.2 common intake — Online notice-specific steps):",
            "steps": ima_online_steps,
            "key": (
                "Key characteristics: e-KYC on Bride & Bridegroom when Aadhaar available; document "
                "upload + eSign; SR verification before First Payment & appointment; SR generates "
                "notice and selects DEO; portal display of marriage notice; 30-day countdown (Sec. 7)."
            ),
        },
        {
            "heading": "7.6 To-Be process — Special Marriage (Intended Marriage) Notice Generation Offline",
            "img": IMG["ima_notice_offline"],
            "caption": "Figure: Special Marriage (Intended Marriage) Notice Generation — Offline",
            "intro": (
                "Statutory basis: SMA 1954 Sec. 5–7. Offline notice path uses portal intake but "
                "completes publication via office: DEO captures photos, prints/signs/scans/uploads "
                "notice, and pastes on the notice board before the 30-day countdown starts."
            ),
            "flow": "Flow (continuing from §7.2 common intake — Offline notice-specific steps):",
            "steps": ima_offline_steps,
            "key": (
                "Key characteristics: capture without mandatory e-KYC path when Aadhaar unavailable; "
                "SR verification → First Payment → appointment; SR generates notice & selects DEO; "
                "physical notice-board paste (Sec. 6(2)); 30-day countdown."
            ),
        },
        {
            "heading": "7.7 To-Be process — Special Marriage (Intended Marriage) Marriage Registration",
            "img": IMG["ima_reg"],
            "caption": "Figure: Special Marriage (Intended Marriage) Marriage Registration",
            "intro": (
                "Statutory basis: SMA 1954 Sec. 7–14 (objection, enquiry, declarations, solemnization, "
                "certificate; new notice if not solemnized within three months). Citizen selects a "
                "published notice only when timeline is ≥30 and ≤90 days; system blocks otherwise."
            ),
            "flow": "Flow (after notice publication / 30-day period — registration steps):",
            "steps": ima_reg_steps,
            "key": (
                "Key characteristics: notice selection with statutory timeline gate; objection enquiry "
                "branch (valid objection removes notice); Second Payment; witness e-KYC; schedule "
                "visit; SR solemnization; DEO joint photo, certificate, signatures; SR DSC; certificate "
                "issuance (Fourth Schedule)."
            ),
        },
        {
            "heading": "7.8 To-Be process — Special Marriage Other Forms Notice Generation Online",
            "img": IMG["of_notice_online"],
            "caption": "Figure: Special Marriage Other Forms Notice Generation — Online",
            "intro": (
                "Statutory basis: SMA 1954 Sec. 15–16 (registration of marriages celebrated in other "
                "forms — public notice and 30-day objections). Online notice generation for Other Forms "
                "follows the approved Online notice diagram (currently aligned to the Intended Marriage "
                "Online notice swimlanes)."
            ),
            "flow": "Flow (continuing from §7.2 common intake — Other Forms Online notice steps):",
            "steps": of_online_steps,
            "key": (
                "Key characteristics: Online e-KYC / upload / eSign → SR verification → First Payment "
                "& appointment → SR notice generation → portal publication → 30-day countdown "
                "(Sec. 16 public notice period)."
            ),
        },
        {
            "heading": "7.9 To-Be process — Special Marriage Other Forms Notice Generation Offline",
            "img": IMG["of_notice_offline"],
            "caption": "Figure: Special Marriage Other Forms Notice Generation — Offline",
            "intro": (
                "Statutory basis: SMA 1954 Sec. 15–16. Offline Other Forms notice generation follows "
                "the approved Offline notice diagram (currently aligned to the Intended Marriage "
                "Offline notice swimlanes), including physical notice-board publication."
            ),
            "flow": "Flow (continuing from §7.2 common intake — Other Forms Offline notice steps):",
            "steps": of_offline_steps,
            "key": (
                "Key characteristics: Offline capture/upload → SR verification → First Payment & "
                "appointment → SR notice → DEO print/sign/scan/upload and notice-board paste → "
                "30-day countdown."
            ),
        },
        {
            "heading": "7.10 To-Be process — Special Marriage Other Forms Marriage Registration",
            "img": IMG["of_reg"],
            "caption": "Figure: Special Marriage Other Forms Marriage Registration",
            "intro": (
                "Statutory basis: SMA 1954 Sec. 15–16 (conditions for registration; procedure including "
                "public notice, objections, and Fifth Schedule certificate signed by parties and three "
                "witnesses). Registration proceeds from an approved notice within the ≥30 and ≤90 day "
                "window shown on the process diagram."
            ),
            "flow": "Flow (after Other Forms notice publication — registration steps):",
            "steps": of_reg_steps,
            "key": (
                "Key characteristics: select notice with timeline gate; objection enquiry branch; "
                "Second Payment; witness e-KYC; schedule visit; SR verification and registration "
                "steps per diagram; DEO certificate production & signatures; SR DSC; certificate "
                "issuance (Fifth Schedule)."
            ),
        },
    ]

    for p in doc.paragraphs:
        if p.style and str(p.style.name).startswith("Heading") and "Application status model" in p.text:
            set_para_text(p, "7.11 Application status model (channel-aware)")
            break

    cursor = para_before(doc, exact="7.11 Application status model (channel-aware)", heading_only=True)

    for sec in sections:
        cursor = build_section_block(
            doc,
            after=cursor,
            heading=sec["heading"],
            image_path=sec["img"],
            figure_caption=sec["caption"],
            intro=sec["intro"],
            flow_intro=sec["flow"],
            steps=sec["steps"],
            key_chars=sec["key"],
        )

    # Ensure status heading number
    for p in doc.paragraphs:
        if p.style and str(p.style.name).startswith("Heading") and "Application status model" in p.text:
            set_para_text(p, "7.11 Application status model (channel-aware)")
            break

    # --- Appendix A ---
    appendix_items = [
        "The Special Marriage Act, 1954 — Acts_Rules/Marriage/The Special Marriage Act, 1954.pdf",
        "Special Marriage (Karnataka) Rules, 1961 — Acts_Rules/Marriage/SpecialMarriage(Karnataka)Rules1961.pdf",
        "Special Marriage fees — Acts_Rules/Marriage/SpecialMarriageFees.docx",
        "Approved process diagram — Special Marriage (Intended Marriage) Notice Online — Process Diagrams/Special Marriage (Intended Marriage) Notice generation-Online.png",
        "Approved process diagram — Special Marriage (Intended Marriage) Notice Offline — Process Diagrams/Special Marriage (Intended Marriage) Notice generation- offline.png",
        "Approved process diagram — Special Marriage (Intended Marriage) Marriage Registration — Process Diagrams/Special Marriage (Intended Marriage) Marriage Registration.png",
        "Approved process diagram — Special Marriage Other Forms Notice Online — Process Diagrams/Special Marriage Other Forms Notice generation-Online.png",
        "Approved process diagram — Special Marriage Other Forms Notice Offline — Process Diagrams/Special Marriage Other Forms Notice generation- offline.png",
        "Approved process diagram — Special Marriage Other Forms Marriage Registration — Process Diagrams/Special Marriage Other Forms Marriage Registration.png",
    ]
    # Find last appendix bullet and append after
    last_app = None
    for p in doc.paragraphs:
        if p.text.startswith("Approved process diagram — Hindu Marriage Offline"):
            last_app = p
        if p.text.startswith("Kaveri 3.0 Marriage prototype"):
            last_app = p
    if last_app is None:
        last_app = find_para(doc, contains="Appendix A")
    anchor = last_app
    for item in appendix_items:
        anchor = insert_paragraph_after(anchor, item, style="List Bullet")

    # Update Hindu appendix paths to Acts_Rules prefix where useful — optional light touch
    for p in doc.paragraphs:
        if p.text.startswith("The Hindu Marriage Act, 1955 — Marriage/"):
            set_para_text(
                p,
                "The Hindu Marriage Act, 1955 — Acts_Rules/Marriage/Hindu Marriage Act, 1955.pdf",
            )
        elif p.text.startswith("Registration of Hindu Marriage (Karnataka) Rules, 1966 — Marriage/"):
            set_para_text(
                p,
                "Registration of Hindu Marriage (Karnataka) Rules, 1966 — "
                "Acts_Rules/Marriage/REGISTRATIONOFHINDUMARRIAGE_KARNATAKARULES_1966.docx",
            )
        elif p.text.startswith("Statutory forms — Marriage/"):
            set_para_text(
                p,
                "Statutory forms — Acts_Rules/Marriage/hindu marriage forms.pdf, Acts_Rules/Marriage/Form1.pdf",
            )
        elif p.text.startswith("Marriage fee / process notification — Marriage/"):
            set_para_text(
                p,
                "Marriage fee / process notification — Acts_Rules/Marriage/RD48MNMU2023-Notification-marriage.pdf (validate)",
            )

    doc.save(str(DST))
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
