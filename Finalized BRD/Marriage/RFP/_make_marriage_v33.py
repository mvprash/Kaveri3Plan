# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.22.docx from v1.21.

Add §7.5.1 Rectified As-Is pain points under §7.5 What is new in Kaveri 3.0,
listing each §6.1 pain point with how Kaveri 3.0 closes it.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.21.docx"
DST = BASE / "BRD_Marriage_v1.22.docx"

RECTIFIED_INTRO = (
    "The following As-Is pain points from §6.1 (Kaveri 2.0 workshops, ServiceDesk "
    "tickets and department discussions) are closed in Kaveri 3.0. Cross-references "
    "point to the To-Be process, functional requirements, fallbacks or NFRs that "
    "implement the fix."
)

RECTIFIED_ROWS = [
    [
        "Sr.No",
        "Pain Point (As-Is)",
        "How rectified in Kaveri 3.0",
        "BRD ref",
    ],
    [
        "1",
        "Not mobile-friendly",
        "Mobile-responsive bilingual UI; guided step flow usable on phone browsers",
        "§10 UI; §15.4 NFR-MRG-VAPT-002",
    ],
    [
        "2",
        "Incorrect address on certificate",
        "Structured party address capture without erroneous sub-district on Form II-A",
        "§8.1.3–8.1.5; FR-HMA-008, FR-HMA-010/011; §8.1.16",
    ],
    [
        "3",
        "Complex area/village selection",
        "MDM-driven jurisdiction / area selection with simplified routing",
        "§8.1.2; FR-HMA-005, FR-HMA-008; §11 Integrations",
    ],
    [
        "4",
        "Restrictive document upload rules",
        "Relaxed size/dimension rules with clear validation messages for supporting documents",
        "§8.1.9; FR-HMA-065, FR-HMA-018/019",
    ],
    [
        "5",
        "Missing Save action for Witness 3",
        "Witness capture supports Save for all three witnesses before submission",
        "§8.1.6; FR-HMA-012–014; BR-HMA-001",
    ],
    [
        "6",
        "Edit does not return to correct step",
        "Edit from in-progress application resumes at the correct channel-aware status step",
        "§7.1.2.4; §8.1.15; FR-HMA-073/074; BR-HMA-014, BR-HMA-017",
    ],
    [
        "7",
        "Documents not visible during scrutiny",
        "SRO scrutiny view shows all uploaded supporting documents with the application",
        "§8.1.11; FR-HMA-026",
    ],
    [
        "8",
        "No edit flexibility before payment",
        "Citizen may navigate back and edit captured data until payment / eSign gate",
        "§7.1.2; §8.1.13; FR-HMA-052; BR-HMA-010",
    ],
    [
        "9",
        "No processing priority order",
        "FIFO / priority handling for officer queues with cycle-time MIS",
        "§8.1.11; FR-HMA-077; §8.6 FR-HMA-042",
    ],
    [
        "10",
        "Password-protected uploads not handled",
        "Clear reject/retry for password-protected documents (FB-MRG-003)",
        "§17 FB-MRG-003; §8.1.9 FR-HMA-065",
    ],
    [
        "11",
        "Limited notification channels",
        "Lifecycle notifications via both SMS and e-Mail",
        "§8.5; FR-HMA-036–038; FR-SMA-054; FB-MRG-004",
    ],
    [
        "12",
        "No operational dashboard visibility",
        "MIS / dashboard for daily registrations, pendency and officer workload",
        "§8.6; FR-HMA-041–045; FR-SMA-055–058",
    ],
    [
        "13",
        "Wrong DEO allocation and no reassignment",
        "SR allocates to DEO; officers can reassign without Service Desk",
        "§7.1.2.3; §8.1.14; FR-HMA-069, FR-HMA-088; §16 RS-MRG-003",
    ],
    [
        "14",
        "Notice workflow state errors",
        "Channel-aware notice / registration status model; post-period actions gated correctly",
        "§7.2.2.4, §7.3.2.3; §8.2.8; FR-SMA-019–032, FR-SMA-024–026",
    ],
    [
        "15",
        "Blank challan / verify-challan loop",
        "Integrated payment with receipt and resilient payment fallback",
        "§8.1.10; §17 FB-MRG-001; NFR-MRG-PAY-001; FR-HMA-025",
    ],
    [
        "16",
        "Stuck on summary page — cannot go to next step",
        "Summary review then proceed / eSign with explicit next-step navigation",
        "§7.1.2; §8.1.13; FR-HMA-051, FR-HMA-052; FR-SMA-012",
    ],
    [
        "17",
        "DOB and party details not reflecting",
        "Party particulars (including DOB) persist and display on summary, forms and certificate",
        "§8.1.4–8.1.5; §8.2.3; FR-HMA-058, FR-HMA-089; FR-SMA-009/062/063",
    ],
    [
        "18",
        "Place of marriage missing in summary/report",
        "Place of marriage captured once and shown on summary, forms and MIS",
        "§8.1.3; §8.2.2; FR-HMA-017; FR-HMA-051",
    ],
    [
        "19",
        "Cross-reference to registered marriage missing",
        "Search / cross-reference to registered marriage records (post-registration)",
        "§8.4; FR-HMA-034; §12.1 Core entities",
    ],
    [
        "20",
        "Signatures missing on digitally signed certificate",
        "Bride, bridegroom and witness signatures rendered on DSC-signed certificate",
        "§8.1.16; §8.3.3; FR-HMA-054/080; FR-SMA-040/048",
    ],
    [
        "21",
        "Notice number and notice date not displaying",
        "Notice number and notice date shown after generation / publication",
        "§8.2.5–8.2.7; FR-SMA-014/016/021; FR-SMA-055",
    ],
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
    last: bool = False,
) -> Paragraph:
    matches: list[Paragraph] = []
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            matches.append(p)
        elif contains is not None and contains in t:
            matches.append(p)
    if not matches:
        raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")
    return matches[-1] if last else matches[0]


def section_exists(doc: Document, exact: str) -> bool:
    try:
        find_para(doc, exact=exact, heading_only=True)
        return True
    except KeyError:
        return False


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_elements_before(anchor: Paragraph, elements: list) -> None:
    """Insert XML elements immediately before anchor paragraph, in order.

    Each addprevious targets the same anchor, so later inserts land closer to
    the anchor; iterate forward so the last element is nearest to §8.
    """
    for el in elements:
        anchor._p.addprevious(el)


def make_paragraph(style_para: Paragraph, text: str, style: str) -> Paragraph:
    new_p = deepcopy(style_para._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    # Temporarily attach after style_para so Paragraph wrapper works, then detach
    style_para._p.addnext(new_p)
    para = Paragraph(new_p, style_para._parent)
    para.style = style
    set_para_text(para, text)
    new_p.getparent().remove(new_p)
    return para


def make_table(doc: Document, rows: list[list[str]]) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    return Table(tbl, doc)


def find_whats_new_table(doc: Document) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:3] == ["#", "Capability", "What is new in Kaveri 3.0"]:
            return table
    raise KeyError("§7.5 What's new capability table not found")


def add_rectified_pain_points(doc: Document) -> None:
    if section_exists(doc, "7.5.1 Rectified As-Is pain points"):
        return

    section8 = find_para(doc, exact="8. Functional requirements", heading_only=True, last=True)
    # Build heading + intro + table, insert before §8 (after capability table)
    h = make_paragraph(section8, "7.5.1 Rectified As-Is pain points", "Heading 4")
    intro = make_paragraph(section8, RECTIFIED_INTRO, "Normal")
    tbl = make_table(doc, RECTIFIED_ROWS)
    insert_elements_before(section8, [h._p, intro._p, tbl._tbl])


def update_contents(doc: Document) -> None:
    for i, p in enumerate(doc.paragraphs[:140]):
        if p.text.strip() == "7.5 What is new in Kaveri 3.0" and style_name(p) == "Normal":
            nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            if "7.5.1" in nxt:
                return
            insert_paragraph_after(p, "7.5.1 Rectified As-Is pain points", style="Normal")
            return
    raise KeyError("Contents entry for 7.5 not found")


def update_intro(doc: Document) -> None:
    """Ensure §7.5 intro mentions both capabilities and rectified pain points."""
    intro = find_para(
        doc,
        contains="This section summarises material enhancements in Kaveri 3.0",
        last=True,
    )
    new_text = (
        "This section summarises material enhancements in Kaveri 3.0 compared with the "
        "legacy Kaveri 2.0 Marriage Registration module (§6). Capability highlights "
        "are listed below; §7.5.1 maps each As-Is pain point from §6.1 to the Kaveri 3.0 "
        "closure. Cross-references point to To-Be process (§7.1–§7.4) and functional "
        "requirements (§8)."
    )
    set_para_text(intro, new_text)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.22")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-03")
    version_table = doc.tables[1]
    last_change = version_table.rows[-1].cells[3].text.strip()
    if "Rectified As-Is pain points" not in last_change:
        add_version_row(
            version_table,
            [
                "1.22",
                "2026-09-03",
                "Nandha Kumar",
                "Add §7.5.1 Rectified As-Is pain points under What is new in Kaveri 3.0",
                "Prashanth",
            ],
        )

    find_whats_new_table(doc)  # preserve user-edited capability table
    update_intro(doc)
    add_rectified_pain_points(doc)
    update_contents(doc)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    find_para(doc2, exact="7.5.1 Rectified As-Is pain points", heading_only=True)
    for table in doc2.tables:
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:2] == ["Sr.No", "Pain Point (As-Is)"]:
            assert len(table.rows) == 22, len(table.rows)
            break
    else:
        raise AssertionError("Rectified pain points table not found")
    print("Verification OK")


if __name__ == "__main__":
    main()
