# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.14.docx from v1.13.

Add 'Addressed in (BRD ref)' column to the existing §6.1 As-Is pain points table,
mapping each pain point to the To-Be sections that close it.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.13.docx"
DST = BASE / "BRD_Marriage_v1.14.docx"

ADDRESSED_IN = {
    "1": "§10 UI; §15.4 NFR-MRG-VAPT-002 (mobile-responsive interfaces)",
    "2": "§8.1.3–8.1.5; FR-HMA-008, FR-HMA-010/011; §8.1.16 FR-HMA-030/080/081",
    "3": "§8.1.2; FR-HMA-005, FR-HMA-008; §11 Integrations (MDM / address master)",
    "4": "§8.1.9; FR-HMA-065, FR-HMA-018/019",
    "5": "§8.1.6; FR-HMA-012–014; BR-HMA-001",
    "6": "§7.1.2.4; §8.1.15; FR-HMA-073/074; BR-HMA-014, BR-HMA-017",
    "7": "§8.1.11; FR-HMA-026",
    "8": "§7.1.2; §8.1.13; FR-HMA-052; BR-HMA-010",
    "9": "§8.1.11; FR-HMA-077; §8.6 FR-HMA-042 (cycle-time MIS)",
    "10": "§17 FB-MRG-003; §8.1.9 FR-HMA-065",
    "11": "§8.5; FR-HMA-036–038; FR-SMA-054; FB-MRG-004",
    "12": "§8.6; FR-HMA-041–045; FR-SMA-055–058",
    "13": "§7.1.2.3; §8.1.14; FR-HMA-069, FR-HMA-088; §16 RS-MRG-003",
    "14": "§7.2.2.4, §7.3.2.3; §8.2.8; FR-SMA-019–032, FR-SMA-024–026",
    "15": "§8.1.10; §17 FB-MRG-001; NFR-MRG-PAY-001; FR-HMA-025, FR-SMA-052",
    "16": "§7.1.2; §8.1.13; FR-HMA-051, FR-HMA-052; FR-SMA-012; FR-HMA-038",
    "17": "§8.1.4–8.1.5; §8.2.3; FR-HMA-058, FR-HMA-089; FR-SMA-009/062/063/066",
    "18": "§8.1.3; §8.2.2; FR-HMA-017; FR-HMA-051",
    "19": "§8.4; FR-HMA-034; §12.1 Core entities",
    "20": "§8.1.16; §8.3.3; FR-HMA-054/080; FR-SMA-040/048; BR-HMA-001, BR-SMA-011",
    "21": "§8.2.5–8.2.7; FR-SMA-014/016/021; FR-SMA-055",
}

INTRO = (
    "Pain points evidenced from Kaveri 2.0 workshops, ServiceDesk tickets and "
    "department discussions. The Addressed in column maps each item to the To-Be "
    "process (§7), functional requirements (§8), fallbacks (§17) or risks (§16) "
    "that close it."
)


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_pain_points_table(doc: Document) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:4] == ["Sr.No", "Pain Point", "Description", "Source"]:
            return table
    raise KeyError("As-Is pain points table not found")


def add_column_to_table(table: Table, header: str, values_by_row: dict[str, str]) -> None:
    """Append a column; values_by_row keyed by Sr.No in column 0."""
    for ri, row in enumerate(table.rows):
        new_tc = deepcopy(row.cells[-1]._tc)
        row._tr.append(new_tc)
        cell = row.cells[-1]
        if ri == 0:
            set_cell_text(cell, header)
        else:
            sr_no = row.cells[0].text.strip()
            set_cell_text(cell, values_by_row.get(sr_no, ""))


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.14")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    add_version_row(
        doc.tables[1],
        [
            "1.14",
            "2026-09-02",
            "Nandha Kumar",
            "Add 'Addressed in (BRD ref)' column to §6.1 As-Is pain points table — "
            "map each pain point to §7 / §8 / §16 / §17 closures",
            "Prashanth",
        ],
    )

    heading = find_para(doc, exact="6.1 As-Is pain points", heading_only=True)
    nxt = heading._element.getnext()
    if nxt is None or INTRO not in "".join(
        t.text or "" for t in nxt.iter(qn("w:t"))
    ):
        insert_paragraph_after(heading, INTRO, style="Normal")

    pain_table = find_pain_points_table(doc)
    hdr = [c.text.strip() for c in pain_table.rows[0].cells]
    if len(hdr) == 4:
        add_column_to_table(pain_table, "Addressed in (BRD ref)", ADDRESSED_IN)
    elif len(hdr) == 5 and hdr[4] == "Addressed in (BRD ref)":
        for ri in range(1, len(pain_table.rows)):
            sr_no = pain_table.rows[ri].cells[0].text.strip()
            set_cell_text(pain_table.rows[ri].cells[4], ADDRESSED_IN.get(sr_no, ""))
    else:
        raise ValueError(f"Unexpected pain points table headers: {hdr}")

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    pain = find_pain_points_table(doc2)
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    print("Headers:", [c.text.strip() for c in pain.rows[0].cells])
    print("Pain tables in doc:", sum(
        1 for t in doc2.tables
        if t.rows and t.rows[0].cells[0].text.strip() == "Sr.No"
    ))
    print("Row 13 ref:", pain.rows[13].cells[4].text[:60])


if __name__ == "__main__":
    main()
