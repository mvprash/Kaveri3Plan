"""Create BRD_Hindu_Marriage_v1.2.docx from v1.1 — reorganize §7 (To-Be) by service
with Channel models + Process Diagram (Common intake / Online / Offline / Status).
Only Section 7 body content is rewritten; TOC and document control are updated for v1.2.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\Prashanth\Official\Kaveri 3.0\Kaveri3Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Hindu_Marriage_v1.1.docx"
DST = BASE / "BRD_Hindu_Marriage_v1.2.docx"
DIAG = BASE / "Process Diagrams"

IMG = {
    "hma_online": DIAG / "hindu marriage Online.png",
    "hma_offline": DIAG / "hindu marriage Offline.png",
    "ima_notice_online": DIAG
    / "Special Marriage (Intended Marriage) Notice generation-Online.png",
    "ima_notice_offline": DIAG
    / "Special Marriage (Intended Marriage) Notice generation- offline.png",
    "ima_reg": DIAG / "Special Marriage (Intended Marriage) Marriage Registration.png",
    "of_notice_online": DIAG
    / "Special Marriage Other Forms Notice generation-Online.png",
    "of_notice_offline": DIAG
    / "Special Marriage Other Forms Notice generation- offline.png",
    "of_reg": DIAG / "Special Marriage Other Forms Marriage Registration.png",
}

CHANNEL_HEADER = ["Service Type", "Online Activities", "Office Activities", "Mode"]
FLOW_HEADER = ["#", "Step", "Lane", "Notes"]
STATUS_HEADER = ["Status", "Description", "Channel", "Actor", "Next states"]

TOC_OLD = [
    "7.1 Channel model",
    "7.2 Common intake steps (both channels)",
    "7.3 To-Be process — Hindu Marriage Online",
    "7.4 To-Be process — Hindu Marriage Offline",
    "7.5 To-Be process — Special Marriage (Intended Marriage) Notice Generation Online",
    "7.6 To-Be process — Special Marriage (Intended Marriage) Notice Generation Offline",
    "7.7 To-Be process — Special Marriage (Intended Marriage) Marriage Registration",
    "7.8 To-Be process — Special Marriage Other Forms Notice Generation Online",
    "7.9 To-Be process — Special Marriage Other Forms Notice Generation Offline",
    "7.10 To-Be process — Special Marriage Other Forms Marriage Registration",
    "7.11 Application status model (channel-aware)",
]

TOC_NEW = [
    "7.1 Hindu Marriage",
    "7.1.1 Channel models",
    "7.1.2 Process Diagram",
    "7.1.2.1 Common intake steps",
    "7.1.2.2 Online",
    "7.1.2.3 Offline (In Person)",
    "7.1.2.4 Application Status Model",
    "7.2 Special Marriage (Intended Marriage) Notice Generation",
    "7.2.1 Channel models",
    "7.2.2 Process Diagram",
    "7.2.2.1 Common intake steps",
    "7.2.2.2 Online",
    "7.2.2.3 Offline (In Person)",
    "7.2.2.4 Application Status Model",
    "7.3 Special Marriage (Intended Marriage) Marriage Registration",
    "7.3.1 Channel models",
    "7.3.2 Process Diagram",
    "7.3.2.1 Common intake steps",
    "7.3.2.2 Offline (In Person)",
    "7.3.2.3 Application Status Model",
    "7.4 Special Marriage (Other Forms) Notice Generation",
    "7.4.1 Channel models",
    "7.4.2 Process Diagram",
    "7.4.2.1 Common intake steps",
    "7.4.2.2 Online",
    "7.4.2.3 Offline (In Person)",
    "7.4.2.4 Application Status Model",
    "7.5 Special Marriage (Other Forms) Marriage Registration",
    "7.5.1 Channel models",
    "7.5.2 Process Diagram",
    "7.5.2.1 Common intake steps",
    "7.5.2.2 Offline (In Person)",
    "7.5.2.3 Application Status Model",
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def _style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not _style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def insert_picture_after(
    paragraph: Paragraph, image_path: Path, width_inches: float = 6.2
) -> Paragraph:
    pic_para = insert_paragraph_after(paragraph, "", style="Normal")
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return pic_para


def clear_between(start: Paragraph, end: Paragraph) -> None:
    """Remove all body siblings strictly between start and end paragraphs."""
    el = start._p.getnext()
    while el is not None and el is not end._p:
        nxt = el.getnext()
        el.getparent().remove(el)
        el = nxt


def set_doc_control_field(doc: Document, field: str, value: str) -> None:
    for row in doc.tables[0].rows:
        if row.cells[0].text.strip() == field:
            row.cells[1].text = value
            return
    raise KeyError(field)


def append_table_row(table: Table, values: list[str]) -> None:
    row = table.add_row()
    for i, v in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def add_heading(cursor: Paragraph, text: str, level: int) -> Paragraph:
    return insert_paragraph_after(cursor, text, style=f"Heading {level}")


def add_normal(cursor: Paragraph, text: str = "") -> Paragraph:
    return insert_paragraph_after(cursor, text, style="Normal")


def add_list_items(cursor: Paragraph, items: list[str]) -> Paragraph:
    for item in items:
        cursor = insert_paragraph_after(cursor, item, style="List Number")
    return cursor


def add_table(cursor: Paragraph, rows: list[list[str]], doc: Document) -> Paragraph:
    insert_table_after(cursor, rows, doc)
    tbl = cursor._p.getnext()
    trailing = insert_paragraph_after(cursor, "", style="Normal")
    if tbl is not None and tbl.tag == qn("w:tbl"):
        trailing._p.getparent().remove(trailing._p)
        tbl.addnext(trailing._p)
    return trailing


def add_figure(
    cursor: Paragraph, image_key: str, caption: str
) -> Paragraph:
    cursor = insert_picture_after(cursor, IMG[image_key])
    cursor = add_normal(cursor, caption)
    return cursor


# --------------------------------------------------------------------------------------
# Content blocks
# --------------------------------------------------------------------------------------

HMA_CHANNEL = [
    CHANNEL_HEADER,
    [
        "Hindu Marriage Online",
        "Registration, e-KYC (Bride), eSign, Payment (post-approval), Tracking",
        "Verification, DSC",
        "Online",
    ],
    [
        "Hindu Marriage Offline",
        "Registration, Payment, Appointment Booking, Printout",
        "SR Stage 1, SR→DEO allocation, Form upload, SR Stage 2, DSC",
        "Offline",
    ],
]

HMA_INTAKE = [
    "START — citizen initiates the service.",
    "LogOn to Portal — authenticated citizen session.",
    "Start a new Application.",
    "Select Marriage Registration service.",
    "Select channel: Hindu Marriage Online / Hindu Marriage Offline. Channel is chosen before prerequisites.",
    "Read and continue with Prerequisite for marriage and complete declaration — single combined acknowledgement screen (eligibility, documents, channel implications and statutory declarations).",
    "Enter / capture Marriage details, Bride details, Bridegroom details, Witness details — persisted to the application record. Online channel: e-KYC on Bride details (see 7.1.2.2). This step is the re-entry point for SR rejection loops that return to citizen data entry in both diagrams.",
]

HMA_ONLINE_FLOW = [
    FLOW_HEADER,
    [
        "7",
        "Capture marriage / e-KYC/Face Authentication on Bride details / Bridegroom / 3 witness details",
        "System / Citizen",
        "Per Online diagram: e-KYC / Face Authentication applied on Bride details during capture (Bridegroom e-KYC / Face Authentication)",
    ],
    [
        "8",
        "Select Sub-Registrar office and review summary of updated information",
        "System / Citizen",
        "Jurisdiction routing",
    ],
    [
        "9",
        "Submit Form I (Memorandum) and Form IA (Application)",
        "System / Citizen",
        "Form I & Form IA generated for the selected office",
    ],
    ["10", "Proceed with eSign", "System / Citizen", "Citizen eSign on Form I and Form IA"],
    ["11", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    [
        "11a",
        "Reject → return to Enter Marriage / Bride / Bridegroom / Witness details",
        "Sub Registrar → System",
        "Citizen corrects and resubmits; refusal reason recorded",
    ],
    [
        "12",
        "Proceed for Online Payment",
        "System / Citizen",
        "Payment occurs only after SR approval",
    ],
    ["13", "SR Digitally signs", "Sub Registrar", "DSC applied"],
    [
        "14",
        "Marriage certificate Generated",
        "Sub Registrar / System",
        "Form II-A available for download",
    ],
]

HMA_OFFLINE_FLOW = [
    FLOW_HEADER,
    [
        "8",
        "SR Verification — Stage 1 (decision) on captured application data",
        "Sub Registrar",
        "Approve or Reject",
    ],
    [
        "8a",
        "Reject → return to Enter Marriage / Bride / Bridegroom / Witness details",
        "Sub Registrar → System",
        "Citizen corrects and resubmits",
    ],
    [
        "9",
        "Makes Payment and schedule appointment",
        "System / Citizen",
        "Payment and slot booking in one step, after Stage 1 approval",
    ],
    [
        "10",
        "Printout taken on Form-1",
        "System / Citizen",
        "Citizen prints the statutory forms",
    ],
    [
        "11",
        "SR allocates to DEO",
        "Sub Registrar",
        "Explicit allocation step per updated Offline diagram before DEO upload",
    ],
    [
        "12",
        "Parties and witnesses sign physically; citizen attends the SRO on the appointment date",
        "Citizen (offline activity)",
        "Not a system step; precondition for DEO upload",
    ],
    [
        "13",
        "Check the form on signature and uploads on portal",
        "Data Entry Operator",
        "DEO verifies signatures are present/complete, scans and uploads",
    ],
    [
        "14",
        "SR Verification — Stage 2 (decision) on the uploaded signed forms",
        "Sub Registrar",
        "Approve or Reject",
    ],
    [
        "14a",
        "Reject → return to DEO check / upload step",
        "Sub Registrar → DEO",
        "Re-check or re-upload; does not go back to citizen data entry",
    ],
    ["15", "SR Digitally Signs", "Sub Registrar", "DSC applied"],
    [
        "16",
        "Marriage certificate Issued",
        "Sub Registrar / System",
        "Form II-A issued",
    ],
]

HMA_STATUS = [
    STATUS_HEADER,
    ["Draft", "Saved not submitted", "Both", "Citizen", "Channel selected"],
    [
        "Channel selected",
        "Online or Offline chosen (before prerequisites)",
        "Both",
        "Citizen",
        "Prerequisite & declaration completed",
    ],
    [
        "Prerequisite & declaration completed",
        "Combined prerequisite + declaration acknowledged",
        "Both",
        "Citizen",
        "Details captured",
    ],
    [
        "Details captured",
        "Marriage / bride (Online: e-KYC / Face Authentication) / bridegroom / witness details saved",
        "Both",
        "Citizen",
        "Office selected (Online) / Pending SR verification Stage 1 (Offline)",
    ],
    [
        "Office selected & summary reviewed",
        "SRO office chosen, summary confirmed",
        "Online",
        "Citizen",
        "Form I & Form IA submitted",
    ],
    [
        "Form I & Form IA submitted",
        "Form I (Memorandum) and Form IA submitted",
        "Online",
        "Citizen",
        "eSign pending",
    ],
    ["eSign pending", "Awaiting citizen eSign", "Online", "Citizen", "Pending SR verification"],
    [
        "Pending SR verification",
        "Awaiting SR scrutiny",
        "Online / Offline (Stage 1)",
        "SR",
        "Approved for payment / Rejected — data",
    ],
    [
        "Rejected — data correction",
        "Sent back to citizen data entry",
        "Both",
        "SR",
        "Details captured",
    ],
    ["Approved for payment", "SR approved; fee payable", "Both", "SR", "Payment completed"],
    [
        "Payment completed",
        "Fee paid, receipt issued",
        "Both",
        "System",
        "Pending SR digital signature (Online) / Appointment scheduled (Offline)",
    ],
    ["Appointment scheduled", "SRO visit slot booked", "Offline", "Citizen", "Forms printed"],
    ["Forms printed", "Form printout taken", "Offline", "Citizen", "Allocated to DEO"],
    [
        "Allocated to DEO",
        "SR has allocated the application to a DEO",
        "Offline",
        "SR",
        "Awaiting signed-form upload",
    ],
    [
        "Awaiting signed-form upload",
        "Physically signed forms pending at SRO",
        "Offline",
        "Citizen / DEO",
        "Signed forms uploaded",
    ],
    [
        "Signed forms uploaded",
        "DEO checked signatures and uploaded",
        "Offline",
        "DEO",
        "Pending SR verification — Stage 2",
    ],
    [
        "Pending SR verification — Stage 2",
        "Awaiting SR scrutiny of signed forms",
        "Offline",
        "SR",
        "Pending SR digital signature / Rejected — upload",
    ],
    [
        "Rejected — upload",
        "Sent back to DEO for re-check / re-upload",
        "Offline",
        "SR",
        "Signed forms uploaded",
    ],
    [
        "Pending SR digital signature",
        "Awaiting DSC",
        "Both",
        "SR",
        "Registered",
    ],
    [
        "Registered",
        "Serial / page / volume assigned, Form II endorsed",
        "Both",
        "SR",
        "Certificate issued",
    ],
    [
        "Certificate issued",
        "Form II-A Generated / downloadable",
        "Both",
        "System",
        "Closed",
    ],
    ["Closed", "No further action", "Both", "System", "—"],
]

IMA_NOTICE_CHANNEL = [
    CHANNEL_HEADER,
    [
        "Special Marriage (Intended Marriage) Notice — Online",
        "e-KYC bride & bridegroom, document upload, eSign, First Payment, appointment, portal notice display",
        "SR verification, notice generation, DEO select",
        "Online notice",
    ],
    [
        "Special Marriage (Intended Marriage) Notice — Offline",
        "Capture details, document upload, First Payment, appointment",
        "SR verification, notice generation, DEO photo/print/sign/scan/upload, paste on notice board",
        "Offline notice",
    ],
]

NOTICE_INTAKE_IMA = [
    "START — citizen initiates the service.",
    "LogOn to Portal — authenticated citizen session.",
    "Start a new Application.",
    "Select Marriage Registration service.",
    "Select channel / service path: Special Marriage (Intended Marriage) Notice Online / Offline. Channel is chosen before prerequisites. Marriage Registration is initiated later by selecting an approved notice (see 7.3).",
    "Read and continue with Prerequisite for marriage and complete declaration — single combined acknowledgement screen (eligibility, documents, channel implications and statutory declarations).",
    "Enter / capture Marriage details, Bride details, Bridegroom details — persisted to the notice application record. This step is the re-entry point for SR rejection loops that return to citizen data entry.",
]

IMA_NOTICE_ONLINE = [
    FLOW_HEADER,
    [
        "8",
        "If Aadhaar available → e-KYC on Bride & Bridegroom details",
        "System / Citizen",
        "Per Online notice diagram",
    ],
    [
        "9",
        "Review summary and proceed document uploading",
        "System / Citizen",
        "Summary of captured particulars",
    ],
    [
        "10",
        "Upload Identity Proof, Photo, Age Proof, Address Proof (Bridegroom & Bride)",
        "Citizen",
        "Mandatory supporting documents for notice",
    ],
    ["11", "Proceed with eSign", "System / Citizen", "Citizen eSign on notice application"],
    ["12", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    [
        "12a",
        "Reject → return to e-KYC / review / upload",
        "Sub Registrar → Citizen",
        "Citizen corrects and resubmits",
    ],
    [
        "13",
        "First Payment",
        "System / Citizen",
        "Notice fee per Karnataka Special Marriage fee schedule",
    ],
    ["14", "Schedule appointment with SR", "System / Citizen", "After first payment"],
    [
        "15",
        "SR Generates Notice",
        "Sub Registrar",
        "Statutory notice (Second Schedule / Sec. 5–6)",
    ],
    ["16", "Selects DEO", "Sub Registrar", "Assign publication / office follow-up"],
    ["17", "Marriage notice displayed in portal", "System", "Sec. 6 publication (digital)"],
    ["18", "30-day countdown starts", "System", "Objection period per Sec. 7"],
]

IMA_NOTICE_OFFLINE = [
    FLOW_HEADER,
    [
        "8",
        "If Aadhaar unavailable → capture Bride & Bridegroom details",
        "System / Citizen",
        "Manual capture path on Offline diagram",
    ],
    ["9", "Review summary and proceed document uploading", "System / Citizen", ""],
    [
        "10",
        "Upload Identity / Age / Address proofs (Bridegroom & Bride)",
        "Citizen",
        "Mandatory documents",
    ],
    ["11", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    [
        "11a",
        "Reject → return to capture Bride / Bridegroom details",
        "Sub Registrar → Citizen",
        "Resubmit",
    ],
    ["12", "First Payment", "System / Citizen", "Notice fee"],
    ["13", "Schedule appointment with SR", "System / Citizen", ""],
    ["14", "SR Generates Notice", "Sub Registrar", "Statutory notice generation"],
    ["15", "Selects DEO", "Sub Registrar", "Assign FDA/SDA/DEO"],
    [
        "16",
        "Capture individual photos of Bride & Bridegroom",
        "FDA / SDA / DEO",
        "Office visit activity",
    ],
    [
        "17",
        "Download, Print, Sign, Scan and Upload notice",
        "FDA / SDA / DEO",
        "Physical notice processed into system",
    ],
    [
        "18",
        "Paste form on respective Notice Board",
        "FDA / SDA / DEO",
        "Sec. 6(2) conspicuous place publication",
    ],
    ["19", "30-day countdown starts", "System", "Objection period per Sec. 7"],
]

SM_NOTICE_STATUS = [
    STATUS_HEADER,
    ["Draft", "Saved not submitted", "Both", "Citizen", "Channel selected"],
    [
        "Channel selected",
        "Online or Offline notice channel chosen (before prerequisites)",
        "Both",
        "Citizen",
        "Prerequisite & declaration completed",
    ],
    [
        "Prerequisite & declaration completed",
        "Combined prerequisite + declaration acknowledged",
        "Both",
        "Citizen",
        "Details captured",
    ],
    [
        "Details captured",
        "Bride / bridegroom particulars saved (Online: e-KYC where Aadhaar available)",
        "Both",
        "Citizen",
        "Notice application submitted",
    ],
    [
        "Notice application submitted",
        "Special Marriage notice application submitted (Online: eSigned)",
        "SM Online / Offline",
        "Citizen",
        "Pending SR verification — notice",
    ],
    [
        "Pending SR verification — notice",
        "Awaiting Marriage Officer scrutiny of the notice application",
        "SM Online / Offline",
        "SR",
        "Notice approved / Rejected — notice data",
    ],
    [
        "Rejected — notice data",
        "Returned to citizen for correction of party details or documents",
        "SM Online / Offline",
        "SR",
        "Notice application submitted",
    ],
    [
        "Notice approved",
        "SR approved; first payment payable",
        "SM Online / Offline",
        "SR",
        "First payment completed",
    ],
    [
        "First payment completed",
        "Notice fee paid, receipt issued",
        "SM Online / Offline",
        "System",
        "Notice generated",
    ],
    [
        "Notice generated",
        "Statutory notice generated and entered in the Marriage Notice Book",
        "SM Online / Offline",
        "SR",
        "Notice published",
    ],
    [
        "Notice published",
        "Published on portal (Online) or pasted on notice board after DEO upload (Offline)",
        "SM Online / Offline",
        "System / DEO",
        "Objection period running",
    ],
    [
        "Objection period running",
        "30-day statutory countdown in progress",
        "SM Online / Offline",
        "System",
        "Notice valid for registration / Objection filed",
    ],
    [
        "Objection filed",
        "Objection recorded in the Marriage Notice Book",
        "SM Online / Offline",
        "Public / SR",
        "Under objection enquiry",
    ],
    [
        "Under objection enquiry",
        "SR enquiry by summoning parties (decide within 30 days)",
        "SM Online / Offline",
        "SR",
        "Objected — closed / Notice valid for registration",
    ],
    [
        "Objected — closed",
        "Objection upheld; notice removed from portal and tagged Objected",
        "SM Online / Offline",
        "SR",
        "Closed",
    ],
    [
        "Notice valid for registration",
        "Notice is ≥ 30 and ≤ 90 days old with no valid objection",
        "SM Online / Offline",
        "Citizen",
        "Proceed to Marriage Registration (see 7.3 / 7.5)",
    ],
    [
        "Notice expired",
        "Validity window lapsed; fresh notice required",
        "SM Online / Offline",
        "System",
        "Closed",
    ],
    ["Closed", "No further action", "Both", "System", "—"],
]

IMA_REG_CHANNEL = [
    CHANNEL_HEADER,
    [
        "Special Marriage (Intended Marriage) Marriage Registration",
        "Select approved notice, Second Payment, e-KYC witnesses, schedule visit, download certificate",
        "Objection enquiry (if any), SR verification, solemnization, DEO certificate & signatures, SR DSC",
        "Registration (In Person)",
    ],
]

REG_INTAKE = [
    "Citizen Login portal — continue from a published / approved notice.",
    "Select Notice — choose the approved notice for which registration / solemnization is sought.",
    "Validate timeline ≥ 30 days and ≤ 90 days — if outside the window, no action is allowed (Sec. 7 / Sec. 14).",
    "If any Objection? — branch to SR enquiry when an objection exists (Sec. 8–9).",
    "Second Payment — registration / solemnization fee.",
    "e-KYC on Witness details — three witnesses (Sec. 11).",
    "Schedule Visit — office visit for solemnization / registration.",
]

IMA_REG_OFFLINE = [
    FLOW_HEADER,
    ["1", "Citizen Login portal", "Citizen", "Continue from published notice"],
    ["2", "Select Notice", "Citizen", "Choose approved / published notice"],
    [
        "3",
        "Validate timeline ≥ 30 days and ≤ 90 days",
        "System",
        "If NO → No Action allowed (Sec. 7 / Sec. 14)",
    ],
    ["4", "If any Objection?", "System", "Branch to SR enquiry when objection exists"],
    ["4a", "Conduct enquiry by summoning all parties", "Sub Registrar", "Sec. 8–9"],
    [
        "4b",
        "Valid objection → update reason; Notice removal from portal (Objected)",
        "Sub Registrar / System",
        "Process stops for solemnization",
    ],
    ["4c", "Objection invalid → continue", "Sub Registrar", "Proceed to Second Payment"],
    ["5", "Second Payment", "System / Citizen", "Registration / solemnization fee"],
    ["6", "e-KYC on Witness details", "System / Citizen", "Three witnesses (Sec. 11)"],
    ["7", "Schedule Visit", "System / Citizen", "Office visit for solemnization"],
    ["8", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    [
        "8a",
        "Reject → return to Schedule Visit",
        "Sub Registrar → System",
        "Reschedule / correct",
    ],
    ["9", "Marriage solemnization", "Sub Registrar", "Sec. 12; declarations Sec. 11"],
    ["10", "Assigns to DEO", "Sub Registrar", "Certificate production"],
    ["11", "Joint Photo capturing", "FDA / SDA / DEO", ""],
    ["12", "Generate Marriage Certificate", "FDA / SDA / DEO", "Fourth Schedule"],
    ["13", "Capture signs of Bride, Bridegroom, Witness", "FDA / SDA / DEO", ""],
    ["14", "Upload signed copy", "FDA / SDA / DEO", ""],
    ["15", "Digital Signature (DSC)", "Sub Registrar", ""],
    ["16", "Marriage Certificate Issued", "Citizen / System", "Download from portal"],
]

SM_REG_STATUS = [
    STATUS_HEADER,
    [
        "Notice valid for registration",
        "Entry point — notice ≥ 30 and ≤ 90 days with no valid objection",
        "SM Registration",
        "Citizen",
        "Second payment completed",
    ],
    [
        "Second payment completed",
        "Registration / solemnization fee paid",
        "SM Registration",
        "System",
        "Visit scheduled",
    ],
    [
        "Visit scheduled",
        "Solemnization / registration visit booked; witness e-KYC done",
        "SM Registration",
        "Citizen",
        "Pending SR verification — visit",
    ],
    [
        "Pending SR verification — visit",
        "SR verification before solemnization",
        "SM Registration",
        "SR",
        "Solemnized / Rejected — visit",
    ],
    [
        "Rejected — visit",
        "Returned to visit scheduling with reason",
        "SM Registration",
        "SR",
        "Visit scheduled",
    ],
    [
        "Solemnized",
        "Marriage solemnized (Chapter II) or conditions satisfied for registration (Chapter III)",
        "SM Registration",
        "SR",
        "Allocated to DEO — certificate",
    ],
    [
        "Allocated to DEO — certificate",
        "Joint photo, certificate generation and signature capture",
        "SM Registration",
        "DEO",
        "Signed certificate uploaded",
    ],
    [
        "Signed certificate uploaded",
        "Signed certificate uploaded for SR digital signature",
        "SM Registration",
        "DEO",
        "Pending SR digital signature",
    ],
    [
        "Pending SR digital signature",
        "Awaiting DSC",
        "SM Registration",
        "SR",
        "Certificate issued",
    ],
    [
        "Certificate issued",
        "Marriage certificate downloadable from portal",
        "SM Registration",
        "System",
        "Closed",
    ],
    ["Closed", "No further action", "SM Registration", "System", "—"],
]

OF_NOTICE_CHANNEL = [
    CHANNEL_HEADER,
    [
        "Special Marriage Other Forms Notice — Online",
        "Same Online notice pattern as Intended Marriage notice generation (current approved diagram)",
        "SR verification, notice generation, DEO select; portal publication + 30-day countdown",
        "Online notice",
    ],
    [
        "Special Marriage Other Forms Notice — Offline",
        "Same Offline notice pattern as Intended Marriage notice generation (current approved diagram)",
        "SR verification, notice generation, DEO physical publication + 30-day countdown",
        "Offline notice",
    ],
]

NOTICE_INTAKE_OF = [
    "START — citizen initiates the service.",
    "LogOn to Portal — authenticated citizen session.",
    "Start a new Application.",
    "Select Marriage Registration service.",
    "Select channel / service path: Special Marriage Other Forms Notice Online / Offline. Channel is chosen before prerequisites. Marriage Registration is initiated later by selecting an approved notice (see 7.5).",
    "Read and continue with Prerequisite for marriage and complete declaration — single combined acknowledgement screen (eligibility, documents, channel implications and statutory declarations for Other Forms / Sec. 15).",
    "Enter / capture Marriage details, Bride details, Bridegroom details — persisted to the notice application record. This step is the re-entry point for SR rejection loops that return to citizen data entry.",
]

OF_NOTICE_ONLINE = [
    FLOW_HEADER,
    [
        "8",
        "If Aadhaar available → e-KYC on Bride & Bridegroom details",
        "System / Citizen",
        "Per Other Forms Online notice diagram (mirrors Intended Marriage Online notice)",
    ],
    [
        "9",
        "Review summary and proceed document uploading",
        "System / Citizen",
        "Summary of captured particulars",
    ],
    [
        "10",
        "Upload Identity Proof, Photo, Age Proof, Address Proof (Bridegroom & Bride)",
        "Citizen",
        "Mandatory supporting documents for notice",
    ],
    ["11", "Proceed with eSign", "System / Citizen", "Citizen eSign on notice application"],
    ["12", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    [
        "12a",
        "Reject → return to e-KYC / review / upload",
        "Sub Registrar → Citizen",
        "Citizen corrects and resubmits",
    ],
    [
        "13",
        "First Payment",
        "System / Citizen",
        "Notice fee per Karnataka Special Marriage fee schedule",
    ],
    ["14", "Schedule appointment with SR", "System / Citizen", "After first payment"],
    [
        "15",
        "SR Generates Notice",
        "Sub Registrar",
        "Statutory notice (Second Schedule / Sec. 5–6)",
    ],
    ["16", "Selects DEO", "Sub Registrar", "Assign publication / office follow-up"],
    ["17", "Marriage notice displayed in portal", "System", "Sec. 6 publication (digital)"],
    ["18", "30-day countdown starts", "System", "Objection period per Sec. 7"],
]

OF_NOTICE_OFFLINE = [
    FLOW_HEADER,
    [
        "8",
        "If Aadhaar unavailable → capture Bride & Bridegroom details",
        "System / Citizen",
        "Per Other Forms Offline notice diagram (mirrors Intended Marriage Offline notice)",
    ],
    ["9", "Review summary and proceed document uploading", "System / Citizen", ""],
    [
        "10",
        "Upload Identity / Age / Address proofs (Bridegroom & Bride)",
        "Citizen",
        "Mandatory documents",
    ],
    ["11", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    [
        "11a",
        "Reject → return to capture Bride / Bridegroom details",
        "Sub Registrar → Citizen",
        "Resubmit",
    ],
    ["12", "First Payment", "System / Citizen", "Notice fee"],
    ["13", "Schedule appointment with SR", "System / Citizen", ""],
    ["14", "SR Generates Notice", "Sub Registrar", "Statutory notice generation"],
    ["15", "Selects DEO", "Sub Registrar", "Assign FDA/SDA/DEO"],
    [
        "16",
        "Capture individual photos of Bride & Bridegroom",
        "FDA / SDA / DEO",
        "Office visit activity",
    ],
    [
        "17",
        "Download, Print, Sign, Scan and Upload notice",
        "FDA / SDA / DEO",
        "Physical notice processed into system",
    ],
    [
        "18",
        "Paste form on respective Notice Board",
        "FDA / SDA / DEO",
        "Sec. 6(2) conspicuous place publication",
    ],
    ["19", "30-day countdown starts", "System", "Objection period per Sec. 7"],
]

OF_REG_CHANNEL = [
    CHANNEL_HEADER,
    [
        "Special Marriage Other Forms Marriage Registration",
        "Select approved notice (timeline ≥30 & ≤90 days), Second Payment, e-KYC witnesses, schedule visit",
        "Objection enquiry (if any), SR verification, registration/solemnization steps per diagram, DEO certificate, SR DSC",
        "Registration (In Person)",
    ],
]

OF_REG_OFFLINE = [
    FLOW_HEADER,
    ["1", "Citizen Login portal", "Citizen", "Continue from published notice"],
    ["2", "Select Notice", "Citizen", "Choose approved / published notice"],
    [
        "3",
        "Validate timeline ≥ 30 days and ≤ 90 days",
        "System",
        "If NO → No Action allowed (Sec. 7 / Sec. 14)",
    ],
    ["4", "If any Objection?", "System", "Branch to SR enquiry when objection exists"],
    ["4a", "Conduct enquiry by summoning all parties", "Sub Registrar", "Sec. 8–9"],
    [
        "4b",
        "Valid objection → update reason; Notice removal from portal (Objected)",
        "Sub Registrar / System",
        "Process stops for solemnization",
    ],
    ["4c", "Objection invalid → continue", "Sub Registrar", "Proceed to Second Payment"],
    ["5", "Second Payment", "System / Citizen", "Registration / solemnization fee"],
    ["6", "e-KYC on Witness details", "System / Citizen", "Three witnesses (Sec. 11)"],
    ["7", "Schedule Visit", "System / Citizen", "Office visit for solemnization"],
    ["8", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    [
        "8a",
        "Reject → return to Schedule Visit",
        "Sub Registrar → System",
        "Reschedule / correct",
    ],
    [
        "9",
        "Marriage registration / solemnization (per Other Forms diagram)",
        "Sub Registrar",
        "Chapter III Sec. 15–16; diagram includes SR registration step",
    ],
    ["10", "Assigns to DEO", "Sub Registrar", "Certificate production"],
    ["11", "Joint Photo capturing", "FDA / SDA / DEO", ""],
    ["12", "Generate Marriage Certificate", "FDA / SDA / DEO", "Fifth Schedule (Sec. 16)"],
    ["13", "Capture signs of Bride, Bridegroom, Witness", "FDA / SDA / DEO", ""],
    ["14", "Upload signed copy", "FDA / SDA / DEO", ""],
    ["15", "Digital Signature (DSC)", "Sub Registrar", ""],
    ["16", "Marriage Certificate Issued", "Citizen / System", "Download from portal"],
]


def update_toc(doc: Document) -> None:
    """Replace old §7 TOC entries with the new hierarchy."""
    # Find first old TOC entry
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == TOC_OLD[0]:
            start_idx = i
            break
    if start_idx is None:
        raise KeyError("TOC start for §7 not found")

    # Verify contiguous old entries
    for offset, expected in enumerate(TOC_OLD):
        actual = doc.paragraphs[start_idx + offset].text.strip()
        if actual != expected:
            raise KeyError(
                f"TOC mismatch at offset {offset}: expected {expected!r}, got {actual!r}"
            )

    # Replace text of first min(len) paragraphs; insert extras or clear leftovers
    n_old = len(TOC_OLD)
    n_new = len(TOC_NEW)
    # Update existing slots
    for i in range(min(n_old, n_new)):
        set_para_text(doc.paragraphs[start_idx + i], TOC_NEW[i])

    if n_new > n_old:
        # Insert remaining after last updated old slot
        anchor = doc.paragraphs[start_idx + n_old - 1]
        for entry in TOC_NEW[n_old:]:
            anchor = insert_paragraph_after(anchor, entry, style="Normal")
    elif n_new < n_old:
        # Clear leftover old entries (should not happen with our TOC)
        for i in range(n_new, n_old):
            set_para_text(doc.paragraphs[start_idx + i], "")


def main() -> None:
    for key, path in IMG.items():
        if not path.exists():
            raise FileNotFoundError(f"{key}: {path}")

    if not SRC.exists():
        raise FileNotFoundError(SRC)

    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # --- Document control ---
    set_doc_control_field(doc, "Version", "1.2")
    set_doc_control_field(doc, "Last updated", "2026-08-25")
    append_table_row(
        doc.tables[1],
        [
            "1.2",
            "2026-08-25",
            "Nandha Kumar",
            "Reorganized §7 (To-Be) by service — Hindu Marriage and Special Marriage "
            "(Intended / Other Forms) Notice Generation and Marriage Registration — each with "
            "Channel models, Process Diagram (Common intake / Online / Offline as applicable) "
            "and Application Status Model",
            "Prashanth",
        ],
    )

    # --- TOC ---
    update_toc(doc)

    # --- Rebuild Section 7 ---
    h7 = find_para(doc, exact="7. Future state (To-Be)", heading_only=True)
    h8 = find_para(doc, exact="8. Functional requirements", heading_only=True)
    clear_between(h7, h8)

    cursor = h7
    cursor = add_normal(
        cursor,
        "The Marriage Registration To-Be model is organized by service. Each service documents "
        "its Channel models, Process Diagram (Common intake steps, Online and/or Offline "
        "(In Person) flows) and Application Status Model. The module covers Hindu Marriage "
        "(Online / Offline) and Special Marriage — Intended Marriage notice (Online / Offline) "
        "followed by marriage registration, and Other Forms notice (Online / Offline) followed "
        "by marriage registration. Notice channels are initiated through the citizen portal; "
        "registration continues from an approved / published notice after the statutory waiting period.",
    )

    # ========== 7.1 Hindu Marriage ==========
    cursor = add_heading(cursor, "7.1 Hindu Marriage", 3)

    cursor = add_heading(cursor, "7.1.1 Channel models", 4)
    cursor = add_normal(
        cursor,
        "Hindu Marriage supports Online and Offline (In Person) channels. Channel is selected "
        "before prerequisites and drives subsequent screens, payment timing and office tasks.",
    )
    cursor = add_table(cursor, HMA_CHANNEL, doc)

    cursor = add_heading(cursor, "7.1.2 Process Diagram", 4)

    cursor = add_heading(cursor, "7.1.2.1 Common intake steps", 5)
    cursor = add_normal(
        cursor,
        "Identical in both Hindu Marriage diagrams (Citizens and System lanes):",
    )
    cursor = add_list_items(cursor, HMA_INTAKE)

    cursor = add_heading(cursor, "7.1.2.2 Online", 5)
    cursor = add_figure(cursor, "hma_online", "Figure: Hindu Marriage Online process")
    cursor = add_normal(
        cursor,
        "Flow (continuing from 7.1.2.1 step 7 — Online-specific steps):",
    )
    cursor = add_table(cursor, HMA_ONLINE_FLOW, doc)
    cursor = add_normal(
        cursor,
        "Key characteristics: channel before combined prerequisite+declaration; e-KYC / Face "
        "Authentication on Bride details during capture; office selection + summary; Form I & "
        "Form IA + eSign; no printout, no appointment, no DEO; single SR verification; payment "
        "only after SR approval; fully digital signature chain.",
    )

    cursor = add_heading(cursor, "7.1.2.3 Offline (In Person)", 5)
    cursor = add_figure(cursor, "hma_offline", "Figure: Hindu Marriage Offline process")
    cursor = add_normal(
        cursor,
        "Flow (continuing from 7.1.2.1 step 7 — Offline-specific steps):",
    )
    cursor = add_table(cursor, HMA_OFFLINE_FLOW, doc)
    cursor = add_normal(
        cursor,
        "Key characteristics: channel before combined prerequisite+declaration; SR Verification "
        "Stage 1 on captured data before payment; payment + appointment bundled; printout of "
        "Form I, Form IA & II-A; SR allocates to DEO; DEO signature check and upload; SR "
        "Verification Stage 2 on uploaded signed forms (reject returns to DEO); SR DSC then certificate.",
    )

    cursor = add_heading(cursor, "7.1.2.4 Application Status Model", 5)
    cursor = add_normal(
        cursor,
        "Channel-aware application statuses for Hindu Marriage Online and Offline:",
    )
    cursor = add_table(cursor, HMA_STATUS, doc)

    # ========== 7.2 SM Intended Notice ==========
    cursor = add_heading(
        cursor, "7.2 Special Marriage (Intended Marriage) Notice Generation", 3
    )

    cursor = add_heading(cursor, "7.2.1 Channel models", 4)
    cursor = add_normal(
        cursor,
        "Intended Marriage notice generation supports Online and Offline (In Person) channels. "
        "The selected notice channel drives publication mode (portal vs notice board) and office tasks.",
    )
    cursor = add_table(cursor, IMA_NOTICE_CHANNEL, doc)

    cursor = add_heading(cursor, "7.2.2 Process Diagram", 4)

    cursor = add_heading(cursor, "7.2.2.1 Common intake steps", 5)
    cursor = add_normal(
        cursor,
        "Identical in both Intended Marriage notice diagrams (Citizens and System lanes):",
    )
    cursor = add_list_items(cursor, NOTICE_INTAKE_IMA)

    cursor = add_heading(cursor, "7.2.2.2 Online", 5)
    cursor = add_figure(
        cursor,
        "ima_notice_online",
        "Figure: Special Marriage (Intended Marriage) Notice Generation — Online",
    )
    cursor = add_normal(
        cursor,
        "Flow (continuing from 7.2.2.1 common intake — Online notice-specific steps):",
    )
    cursor = add_table(cursor, IMA_NOTICE_ONLINE, doc)
    cursor = add_normal(
        cursor,
        "Key characteristics: e-KYC on Bride & Bridegroom when Aadhaar available; document "
        "upload + eSign; SR verification before First Payment & appointment; SR generates notice "
        "and selects DEO; portal display of marriage notice; 30-day countdown (Sec. 7).",
    )

    cursor = add_heading(cursor, "7.2.2.3 Offline (In Person)", 5)
    cursor = add_figure(
        cursor,
        "ima_notice_offline",
        "Figure: Special Marriage (Intended Marriage) Notice Generation — Offline",
    )
    cursor = add_normal(
        cursor,
        "Flow (continuing from 7.2.2.1 common intake — Offline notice-specific steps):",
    )
    cursor = add_table(cursor, IMA_NOTICE_OFFLINE, doc)
    cursor = add_normal(
        cursor,
        "Key characteristics: capture without mandatory e-KYC path when Aadhaar unavailable; "
        "SR verification → First Payment → appointment; SR generates notice & selects DEO; "
        "physical notice-board paste (Sec. 6(2)); 30-day countdown.",
    )

    cursor = add_heading(cursor, "7.2.2.4 Application Status Model", 5)
    cursor = add_normal(
        cursor,
        "Application statuses for Special Marriage (Intended Marriage) Notice Generation "
        "(Online / Offline):",
    )
    cursor = add_table(cursor, SM_NOTICE_STATUS, doc)

    # ========== 7.3 SM Intended Registration ==========
    cursor = add_heading(
        cursor, "7.3 Special Marriage (Intended Marriage) Marriage Registration", 3
    )

    cursor = add_heading(cursor, "7.3.1 Channel models", 4)
    cursor = add_normal(
        cursor,
        "Marriage Registration for Intended Marriage continues from an approved / published "
        "notice after the statutory waiting period. Completion is Offline (In Person) at the "
        "Sub-Registrar office (solemnization, DEO certificate production and SR DSC). There is "
        "no separate fully-online registration channel.",
    )
    cursor = add_table(cursor, IMA_REG_CHANNEL, doc)

    cursor = add_heading(cursor, "7.3.2 Process Diagram", 4)

    cursor = add_heading(cursor, "7.3.2.1 Common intake steps", 5)
    cursor = add_normal(
        cursor,
        "Portal steps before the in-person office visit (after notice publication / 30-day period):",
    )
    cursor = add_list_items(cursor, REG_INTAKE)

    cursor = add_heading(cursor, "7.3.2.2 Offline (In Person)", 5)
    cursor = add_figure(
        cursor,
        "ima_reg",
        "Figure: Special Marriage (Intended Marriage) Marriage Registration",
    )
    cursor = add_normal(
        cursor,
        "Flow (after notice publication / 30-day period — registration / solemnization steps):",
    )
    cursor = add_table(cursor, IMA_REG_OFFLINE, doc)
    cursor = add_normal(
        cursor,
        "Key characteristics: notice selection with statutory timeline gate; objection enquiry "
        "branch (valid objection removes notice); Second Payment; witness e-KYC; schedule visit; "
        "SR solemnization; DEO joint photo, certificate, signatures; SR DSC; certificate issuance "
        "(Fourth Schedule).",
    )

    cursor = add_heading(cursor, "7.3.2.3 Application Status Model", 5)
    cursor = add_normal(
        cursor,
        "Application statuses for Special Marriage (Intended Marriage) Marriage Registration:",
    )
    cursor = add_table(cursor, SM_REG_STATUS, doc)

    # ========== 7.4 SM Other Forms Notice ==========
    cursor = add_heading(
        cursor, "7.4 Special Marriage (Other Forms) Notice Generation", 3
    )

    cursor = add_heading(cursor, "7.4.1 Channel models", 4)
    cursor = add_normal(
        cursor,
        "Other Forms notice generation supports Online and Offline (In Person) channels, "
        "mirroring the Intended Marriage notice pattern with Other Forms eligibility "
        "(Chapter III / Sec. 15).",
    )
    cursor = add_table(cursor, OF_NOTICE_CHANNEL, doc)

    cursor = add_heading(cursor, "7.4.2 Process Diagram", 4)

    cursor = add_heading(cursor, "7.4.2.1 Common intake steps", 5)
    cursor = add_normal(
        cursor,
        "Identical in both Other Forms notice diagrams (Citizens and System lanes):",
    )
    cursor = add_list_items(cursor, NOTICE_INTAKE_OF)

    cursor = add_heading(cursor, "7.4.2.2 Online", 5)
    cursor = add_figure(
        cursor,
        "of_notice_online",
        "Figure: Special Marriage Other Forms Notice Generation — Online",
    )
    cursor = add_normal(
        cursor,
        "Flow (continuing from 7.4.2.1 common intake — Other Forms Online notice steps):",
    )
    cursor = add_table(cursor, OF_NOTICE_ONLINE, doc)
    cursor = add_normal(
        cursor,
        "Key characteristics: Online e-KYC / upload / eSign → SR verification → First Payment "
        "& appointment → SR notice generation → portal publication → 30-day countdown "
        "(Sec. 16 public notice period).",
    )

    cursor = add_heading(cursor, "7.4.2.3 Offline (In Person)", 5)
    cursor = add_figure(
        cursor,
        "of_notice_offline",
        "Figure: Special Marriage Other Forms Notice Generation — Offline",
    )
    cursor = add_normal(
        cursor,
        "Flow (continuing from 7.4.2.1 common intake — Other Forms Offline notice steps):",
    )
    cursor = add_table(cursor, OF_NOTICE_OFFLINE, doc)
    cursor = add_normal(
        cursor,
        "Key characteristics: Offline capture/upload → SR verification → First Payment & "
        "appointment → SR notice → DEO print/sign/scan/upload and notice-board paste → "
        "30-day countdown.",
    )

    cursor = add_heading(cursor, "7.4.2.4 Application Status Model", 5)
    cursor = add_normal(
        cursor,
        "Application statuses for Special Marriage (Other Forms) Notice Generation "
        "(Online / Offline):",
    )
    # Reuse notice status table with Other Forms registration cross-ref already saying 7.3 / 7.5
    cursor = add_table(cursor, SM_NOTICE_STATUS, doc)

    # ========== 7.5 SM Other Forms Registration ==========
    cursor = add_heading(
        cursor, "7.5 Special Marriage (Other Forms) Marriage Registration", 3
    )

    cursor = add_heading(cursor, "7.5.1 Channel models", 4)
    cursor = add_normal(
        cursor,
        "Marriage Registration for Other Forms continues from an approved / published notice "
        "after the statutory waiting period. Completion is Offline (In Person) at the "
        "Sub-Registrar office. There is no separate fully-online registration channel.",
    )
    cursor = add_table(cursor, OF_REG_CHANNEL, doc)

    cursor = add_heading(cursor, "7.5.2 Process Diagram", 4)

    cursor = add_heading(cursor, "7.5.2.1 Common intake steps", 5)
    cursor = add_normal(
        cursor,
        "Portal steps before the in-person office visit (after Other Forms notice publication):",
    )
    cursor = add_list_items(cursor, REG_INTAKE)

    cursor = add_heading(cursor, "7.5.2.2 Offline (In Person)", 5)
    cursor = add_figure(
        cursor,
        "of_reg",
        "Figure: Special Marriage Other Forms Marriage Registration",
    )
    cursor = add_normal(
        cursor,
        "Flow (after Other Forms notice publication — registration steps):",
    )
    cursor = add_table(cursor, OF_REG_OFFLINE, doc)
    cursor = add_normal(
        cursor,
        "Key characteristics: select notice with timeline gate; objection enquiry branch; "
        "Second Payment; witness e-KYC; schedule visit; SR verification and registration steps "
        "per diagram; DEO certificate production & signatures; SR DSC; certificate issuance "
        "(Fifth Schedule).",
    )

    cursor = add_heading(cursor, "7.5.2.3 Application Status Model", 5)
    cursor = add_normal(
        cursor,
        "Application statuses for Special Marriage (Other Forms) Marriage Registration:",
    )
    cursor = add_table(cursor, SM_REG_STATUS, doc)

    doc.save(str(DST))
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
