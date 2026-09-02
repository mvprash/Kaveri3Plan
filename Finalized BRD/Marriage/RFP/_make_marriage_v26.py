# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.16.docx from v1.15.

- §3.5: add Form IA statutory declaration text and Sec. 17 / Sec. 18 penalty highlights.
- FR-HMA-091: display Sec. 17 & Sec. 18 in red at Form I / Form IA submission (both channels).
- FR-SMA-067: after second payment, display First Schedule prohibited-relationship details
  before visit / appointment scheduling during Special Marriage solemnization.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.15.docx"
DST = BASE / "BRD_Marriage_v1.16.docx"

HIGHLIGHT_RED = RGBColor(0xC0, 0x00, 0x00)

FR_HMA_091 = (
    "At Form I / Form IA submission (Hindu Marriage Online: immediately before citizen "
    "eSign; Hindu Marriage Offline: immediately before SR Verification Stage 1 submission), "
    "system shall display statutory highlight panels for HMA Sec. 17 (Punishment of bigamy) "
    "and Sec. 18 (Punishment for contravention of certain other conditions for a Hindu "
    "marriage) in red (#C00000 or department-approved accessible equivalent), using the "
    "wording in §3.5; citizen shall explicitly acknowledge before submission proceeds"
)

FR_SMA_067 = (
    "During Special Marriage solemnization (Intended Marriage path), after successful "
    "second payment (FR-SMA-033 / FR-SMA-050) and before visit / appointment scheduling "
    "(FR-SMA-035), system shall display the degrees of prohibited relationship under the "
    "First Schedule to the Special Marriage Act, 1954, together with any party "
    "relationship particulars already captured; citizen shall explicitly acknowledge before "
    "the appointment booking step is enabled"
)

FR_SMA_033_ACCEPT = (
    "Receipt generated; prohibited-relationship acknowledgement gate (FR-SMA-067) before "
    "scheduling; payment gate on scheduling"
)

FR_SMA_035_ACCEPT = (
    "Slot, date, time and office recorded; scheduling blocked until FR-SMA-067 acknowledged"
)


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def add_fr_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_red_paragraph_after(paragraph: Paragraph, text: str, *, bold: bool = True) -> Paragraph:
    new_para = insert_paragraph_after(paragraph, "", style="Normal")
    run = new_para.add_run(text)
    run.font.color.rgb = HIGHLIGHT_RED
    run.bold = bold
    return new_para


def find_table_by_header(doc: Document, first_cell: str, second_cell: str | None = None) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr0 = table.rows[0].cells[0].text.strip()
        hdr1 = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
        if hdr0 == first_cell and (second_cell is None or hdr1 == second_cell):
            return table
    raise KeyError(f"Table not found: {first_cell!r} / {second_cell!r}")


def find_table_containing(doc: Document, req_id: str) -> Table:
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == req_id:
                return table
    raise KeyError(f"Table not found containing {req_id!r}")


def find_fr_row(table: Table, req_id: str):
    for row in table.rows:
        if row.cells[0].text.strip() == req_id:
            return row
    raise KeyError(f"FR row not found: {req_id}")


def paragraph_after_table(table: Table) -> Paragraph:
    tbl = table._tbl
    nxt = tbl.getnext()
    while nxt is not None and nxt.tag != qn("w:p"):
        nxt = nxt.getnext()
    if nxt is None:
        raise KeyError("No paragraph found after table")
    return Paragraph(nxt, table._parent)


def replace_in_paragraphs(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    set_cell_text(cell, cell.text.replace(old, new))


def add_hma_declaration_section(anchor: Paragraph) -> None:
    """Insert Form IA declaration and Sec. 17 / 18 highlights before Sample forms."""
    if anchor.text.strip() == "Sample forms":
        cursor = insert_paragraph_after(anchor, "", style="Normal")
        anchor._p.getparent().remove(anchor._p)
        anchor = cursor
    else:
        cursor = anchor

    cursor = insert_paragraph_after(
        cursor, "Form IA — Statutory declaration (Rule 4(2))", style="Heading 4"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Source: Acts_Rules/Marriage/hindu marriage forms.pdf — FORM-IA (See rule 4(2)). "
        "Kaveri 3.0 shall present the approved gazette wording verbatim on the combined "
        "prerequisite + declaration screen and on generated Form IA outputs.",
        style="Normal",
    )
    cursor = insert_paragraph_after(
        cursor,
        "Opening narrative (system-populated): Marriage in accordance with the Hindu "
        "Marriage Act, 1955 (Act No. 25 of 1955) has been solemnized between the "
        "undersigned parties on [date of marriage], and they request that particulars "
        "be entered in the Hindu Marriages Register.",
        style="Normal",
    )
    cursor = insert_paragraph_after(cursor, "We hereby declare,", style="Normal")
    cursor = insert_paragraph_after(
        cursor,
        "(I) That a valid marriage was solemnized between us and that the marriage is "
        "capable of being registered under section 8 of the Hindu Marriage Act, 1955 "
        "(Central Act 25 of 1955).",
        style="Normal",
    )
    cursor = insert_paragraph_after(
        cursor,
        "(II) That the conditions laid down in section 5 of the said Act have been satisfied.",
        style="Normal",
    )
    cursor = insert_paragraph_after(
        cursor,
        "(III) That the particulars given in the application are true to the best of our "
        "knowledge and belief.",
        style="Normal",
    )
    cursor = insert_paragraph_after(
        cursor,
        "Signatures: (1) Husband (2) Wife — physical signature (Offline) or eSign (Online) "
        "per channel rules.",
        style="Normal",
    )

    cursor = insert_paragraph_after(
        cursor,
        "Statutory penalty highlights — Sec. 17 & Sec. 18 (displayed in red at submission)",
        style="Heading 4",
    )
    cursor = insert_paragraph_after(
        cursor,
        "The following extracts from the Hindu Marriage Act, 1955 shall be shown to the "
        "parties in red highlight panels when Form I / Form IA is submitted (Online and "
        "Offline) — see FR-HMA-091. Wording below is the statutory reference text for "
        "legal review; UI may abbreviate for readability without altering legal meaning.",
        style="Normal",
    )
    cursor = insert_red_paragraph_after(
        cursor,
        "Sec. 17 — Punishment of bigamy: Any marriage between two Hindus solemnized after "
        "the commencement of this Act is void if at the date of such marriage either party "
        "had a husband or wife living; and the provisions of sections 494 and 495 of the "
        "Indian Penal Code (45 of 1860), shall apply accordingly.",
    )
    cursor = insert_red_paragraph_after(
        cursor,
        "Sec. 18 — Punishment for contravention of certain other conditions for a Hindu "
        "marriage: Contravention of section 5(iii) (minimum age — bridegroom 21 years, "
        "bride 18 years): imprisonment which may extend to two years, or fine which may "
        "extend to one lakh rupees, or both. Contravention of section 5(iv) (degrees of "
        "prohibited relationship) or section 5(v) (sapinda relationship): simple "
        "imprisonment which may extend to one month, or fine which may extend to one "
        "thousand rupees, or both.",
    )
    insert_paragraph_after(cursor, "", style="Normal")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.16")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    add_version_row(
        doc.tables[1],
        [
            "1.16",
            "2026-09-02",
            "Nandha Kumar",
            "§3.5 Form IA declaration + Sec. 17/18 red highlights; FR-HMA-091; "
            "FR-SMA-067 prohibited-relationship gate after second payment",
            "Prashanth",
        ],
    )

    hma_table = find_table_by_header(doc, "Form", "Rule ref")
    sample_anchor = find_para(doc, exact="Sample forms", heading_only=True)
    add_hma_declaration_section(sample_anchor)

    replace_in_paragraphs(
        doc,
        "8.1.8 Form IA — application and declarations (FR-HMA-083–085)",
        "8.1.8 Form IA — application and declarations (FR-HMA-083–085, FR-HMA-091)",
    )
    replace_in_paragraphs(
        doc,
        "8.3.3 Solemnization, declarations and certificate (FR-SMA-033–041)",
        "8.3.3 Solemnization, declarations and certificate (FR-SMA-033–041, FR-SMA-067)",
    )

    hma_fr_table = find_table_containing(doc, "FR-HMA-084")
    add_fr_row(
        hma_fr_table,
        [
            "FR-HMA-091",
            FR_HMA_091,
            "Must",
            "Red highlight panels for Sec. 17 and Sec. 18 per §3.5; acknowledgement "
            "timestamped; applies to Hindu Marriage Online (pre-eSign) and Offline "
            "(pre Stage 1 submission); submission blocked until acknowledged",
        ],
    )

    sma_fr_table = find_table_containing(doc, "FR-SMA-033")
    set_cell_text(find_fr_row(sma_fr_table, "FR-SMA-033").cells[3], FR_SMA_033_ACCEPT)
    set_cell_text(find_fr_row(sma_fr_table, "FR-SMA-035").cells[3], FR_SMA_035_ACCEPT)
    add_fr_row(
        sma_fr_table,
        [
            "FR-SMA-067",
            FR_SMA_067,
            "Must",
            "First Schedule text displayed; party relationship context shown where "
            "captured; acknowledgement timestamped; appointment booking disabled until "
            "acknowledged; applies to Intended Marriage solemnization (Online and Offline)",
        ],
    )

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    find_para(doc2, contains="Form IA — Statutory declaration")
    find_para(doc2, contains="Sec. 17 — Punishment of bigamy")
    find_table_containing(doc2, "FR-HMA-091")
    find_table_containing(doc2, "FR-SMA-067")
    print("Verification OK")


if __name__ == "__main__":
    main()
