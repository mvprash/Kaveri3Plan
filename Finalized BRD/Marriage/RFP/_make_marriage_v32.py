# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.21.docx from v1.20.

Add §7.5 What is new in Kaveri 3.0 — summary of enhancements over legacy Kaveri 2.0,
placed after §7.4 and before §8.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.20.docx"
DST = BASE / "BRD_Marriage_v1.21.docx"

INTRO = (
    "This section summarises material enhancements in Kaveri 3.0 compared with the "
    "legacy Kaveri 2.0 Marriage Registration module (§6). Each item is implemented "
    "in the To-Be process (§7.1–§7.4) and corresponding functional requirements (§8). "
    "As-Is pain points from §6.1 are explicitly closed by these capabilities."
)

WHATS_NEW_ROWS = [
    ["#", "Capability", "What is new in Kaveri 3.0", "Legacy gap / pain point closed", "BRD ref"],
    [
        "1",
        "Dual-channel services",
        "Hindu Marriage and Special Marriage (Notice and Registration) support "
        "Online (portal) and Offline (In Person) channels with shared intake and "
        "channel-specific branches",
        "Single-mode legacy flows; limited citizen choice",
        "§7.1–§7.3; §8.1.12–8.1.14; §8.2; §8.3",
    ],
    [
        "2",
        "Identity verification",
        "Mandatory e-KYC / Face Authentication on Online paths; Offline Aadhaar "
        "YES/NO branch with manual capture fallback; OTP SMS states verification reason",
        "Inconsistent identity capture; manual re-entry errors",
        "§7.1.2.2–7.1.2.3; §7.2.2.2–7.2.2.3; FR-HMA-051–058; FR-SMA-009–012",
    ],
    [
        "3",
        "Prerequisite and declarations",
        "Combined read-and-continue prerequisite + statutory declaration screen "
        "before data capture; Sec. 17 & Sec. 18 penalty text displayed in red at "
        "Form I / Form IA submission",
        "Scattered eligibility checks; weak statutory awareness",
        "§7.1.2.1 step 6; §3.5; FR-HMA-046–050; FR-HMA-091",
    ],
    [
        "4",
        "Statutory forms and certificate",
        "Auto-generated Form I, Form IA, Form II endorsement and Form II-A "
        "certificate; SR DSC; register serial/page/volume; QR integrity on certificate",
        "Missing signatures on digital certificate (PP-20); incomplete register entry",
        "§7.1.2.2–7.1.2.3; §8.1.16; FR-HMA-078–082",
    ],
    [
        "5",
        "Special Marriage unified workflows",
        "Shared Notice Generation and Registration for Intended Marriage and Other "
        "Forms; timeline gates (≥30 / ≤90 days); objection enquiry branch; Second, "
        "Third and Fourth Schedule automation",
        "Notice workflow state errors (PP-14); missing notice number/date (PP-21)",
        "§7.2–§7.3; §8.2–§8.3; FR-SMA-014–048",
    ],
    [
        "6",
        "Sakala integration",
        "GSC number on payment; bidirectional lifecycle sync with "
        "https://sakala.kar.nic.in/; citizen tracking and appeal alignment",
        "No statutory service guarantee tracking",
        "§3.7; §7 (payment milestones); §8.7; FR-HMA-092–093; FR-SMA-068–069",
    ],
    [
        "7",
        "Null and Void endorsement",
        "SRO back-office workflow to record certified court order and flag registered "
        "Hindu marriage Null and Void (HMA 1955 Sec. 11–12)",
        "Manual register annotation only; no structured nullity workflow",
        "§7.4; §8.1.17; FR-HMA-094–099; BR-HMA-021",
    ],
    [
        "8",
        "Scrutiny and office operations",
        "Documents visible during SR scrutiny; edit-before-payment; channel-aware "
        "status models; DEO reassignment without Service Desk; processing priority "
        "and pendency MIS",
        "PP-6, PP-7, PP-8, PP-9, PP-12, PP-13",
        "§7.1.2.4; §8.1.11; §8.1.14–8.1.15; §8.6; FR-HMA-069, FR-HMA-088",
    ],
    [
        "9",
        "Data capture and address quality",
        "Structured party/witness capture; MDM-driven address; place of marriage in "
        "summary and reports; cross-reference search for registered marriages",
        "PP-2, PP-3, PP-5, PP-17, PP-18, PP-19",
        "§8.1.2–8.1.6; §8.2.2–8.2.3; §8.4; FR-HMA-008, FR-HMA-034",
    ],
    [
        "10",
        "Payments and notifications",
        "Integrated fee payment with receipt (Form VI); resilient payment fallback; "
        "SMS and e-Mail notifications across application lifecycle",
        "Blank challan / verify-challan loop (PP-15); limited notification channels (PP-11)",
        "§8.1.10; §8.5; §17 FB-MRG-001, FB-MRG-004; FR-HMA-020–025",
    ],
    [
        "11",
        "Integrations ecosystem",
        "Payment gateway, Aadhaar/e-KYC, DigiLocker, Kutumba portal, Civil Registration "
        "System, Labor Department and Sakala — per §11 integration matrix",
        "Siloed legacy module with partial integrations",
        "§2.1; §11 Integrations",
    ],
    [
        "12",
        "Citizen experience",
        "Mobile-responsive bilingual UI (English + Kannada); guided step flow with "
        "Save/resume; summary review before submit and eSign",
        "Not mobile-friendly (PP-1); stuck on summary page (PP-16)",
        "§2.1; §10 UI; §15.4 NFR-MRG-VAPT-002",
    ],
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
    last: bool = False,
) -> Paragraph:
    matches: list[Paragraph] = []
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            matches.append(p)
        elif contains is not None and contains in t:
            matches.append(p)
    if not matches:
        raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")
    return matches[-1] if last else matches[0]


def section_exists(doc: Document, exact: str) -> bool:
    try:
        find_para(doc, exact=exact, heading_only=True)
        return True
    except KeyError:
        return False


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)


def add_whats_new_section(doc: Document) -> None:
    if section_exists(doc, "7.5 What is new in Kaveri 3.0"):
        return

    section8 = find_para(doc, exact="8. Functional requirements", heading_only=True, last=True)

    heading = deepcopy(section8._p)
    for child in list(heading):
        if child.tag != qn("w:pPr"):
            heading.remove(child)
    section8._p.addprevious(heading)
    h_para = Paragraph(heading, section8._parent)
    h_para.style = "Heading 3"
    set_para_text(h_para, "7.5 What is new in Kaveri 3.0")

    cursor = insert_paragraph_after(h_para, INTRO, style="Normal")
    insert_table_after(cursor, WHATS_NEW_ROWS, doc)


def update_contents(doc: Document) -> None:
    for i, p in enumerate(doc.paragraphs[:140]):
        if p.text.strip() == "7.4.1 Process steps" and style_name(p) == "Normal":
            nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            if "7.5 What is new" in nxt:
                return
            insert_paragraph_after(p, "7.5 What is new in Kaveri 3.0", style="Normal")
            return
    raise KeyError("Contents entry for 7.4.1 not found")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.21")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    version_table = doc.tables[1]
    last_change = version_table.rows[-1].cells[3].text.strip()
    if "What is new in Kaveri 3.0" not in last_change:
        add_version_row(
            version_table,
            [
                "1.21",
                "2026-09-02",
                "Nandha Kumar",
                "Add §7.5 What is new in Kaveri 3.0 summary table after §7.4",
                "Prashanth",
            ],
        )

    add_whats_new_section(doc)
    update_contents(doc)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    find_para(doc2, exact="7.5 What is new in Kaveri 3.0", heading_only=True)
    for table in doc2.tables:
        for row in table.rows:
            if row.cells[0].text.strip() == "1" and "Dual-channel" in row.cells[1].text:
                assert len(table.rows) == 13
                break
        else:
            continue
        break
    else:
        raise AssertionError("What's new table not found")
    print("Verification OK")


if __name__ == "__main__":
    main()
