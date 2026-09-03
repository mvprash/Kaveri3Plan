# -*- coding: utf-8 -*-
"""Complete Sakala actor row in §4 of BRD_Marriage_BRD_v1.22.docx."""
from __future__ import annotations

import shutil
import sys
import time
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
DST = BASE / "BRD_Marriage_BRD_v1.22.docx"
TMP = BASE / "BRD_Marriage_BRD_v1.22._sakala_tmp.docx"

SAKALA = [
    "Sakala (Karnataka Guarantee of Services)",
    "Statutory time-bound service delivery platform (Sakala Mission / NIC)",
    "Issue GSC on acceptance/payment; track statutory timelines; sync lifecycle "
    "(in-process / delivered / rejected); citizen tracking and appeals",
    "Both (GSC and status sync for Online and Offline Hindu / Special Marriage services)",
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    row = table.rows[-1]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def find_actors_table(doc: Document) -> Table:
    for t in doc.tables:
        if not t.rows:
            continue
        hdr = [c.text.strip() for c in t.rows[0].cells]
        if hdr[:4] == ["Actor", "Description", "Primary goals", "Channel involvement"]:
            return t
    raise KeyError("§4 actors table not found")


def main() -> None:
    doc = Document(str(DST))
    actors = find_actors_table(doc)

    sakala_row = None
    for row in actors.rows[1:]:
        if row.cells[0].text.strip().startswith("Sakala"):
            sakala_row = row
            break
    if sakala_row is None:
        actors._tbl.append(deepcopy(actors.rows[-1]._tr))
        sakala_row = actors.rows[-1]

    for ci, val in enumerate(SAKALA):
        set_cell_text(sakala_row.cells[ci], val)

    last = doc.tables[1].rows[-1].cells[3].text.strip()
    if "Sakala actor" not in last:
        add_version_row(
            doc.tables[1],
            [
                "1.22",
                "2026-09-03",
                "Nandha Kumar",
                "§4 Stakeholders: complete Sakala actor Description, Primary goals and Channel involvement",
                "Prashanth",
            ],
        )

    doc.save(str(TMP))
    print(f"Wrote temp {TMP}")

    replaced = False
    last_err: Exception | None = None
    for attempt in range(8):
        try:
            shutil.move(str(TMP), str(DST))
            replaced = True
            break
        except PermissionError as e:
            last_err = e
            time.sleep(1.5)
    if not replaced:
        print(
            f"Could not overwrite {DST.name} (file locked). "
            f"Close Word and rename {TMP.name} → {DST.name}. "
            f"Error: {last_err}"
        )
        return

    doc2 = Document(str(DST))
    actors2 = find_actors_table(doc2)
    row = next(r for r in actors2.rows if r.cells[0].text.strip().startswith("Sakala"))
    assert all(c.text.strip() for c in row.cells[:4]), "Sakala columns still empty"
    print(f"Updated {DST}")
    for c in row.cells:
        print(f"  - {c.text}")


if __name__ == "__main__":
    main()
