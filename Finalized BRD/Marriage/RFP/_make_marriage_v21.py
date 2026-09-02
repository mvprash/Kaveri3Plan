# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.11.docx from v1.10.

- OTP during e-KYC / Face Authentication shall state the reason it is requested.
- §7.1.2.3 Hindu Marriage Offline: Aadhaar YES/NO branch with e-KYC or manual entry.
- §7.2.2.2 / 7.2.2.3: "Whether Marriage already taken place or not?" wording.
- FR-SMA-003: Sec. 15 conditions at notice filing date.
- Form II included in Hindu Marriage Online and Offline registration.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.10.docx"
DST = BASE / "BRD_Marriage_v1.11.docx"

MARRIAGE_TAKEN_OLD = "Whether Marriage Taken place or Not?"
MARRIAGE_TAKEN_NEW = "Whether Marriage already taken place or not?"
MARRIAGE_TAKEN_BRANCH_OLD = "Whether Marriage Taken place branch"
MARRIAGE_TAKEN_BRANCH_NEW = "Whether Marriage already taken place or not? branch"

FR_SMA_003_NEW = (
    "System shall enforce Sec. 15 conditions for Other Forms: ceremony already "
    "performed and parties living together as husband and wife since; neither party "
    "has more than one spouse living; both parties ≥ 21 years at notice filing date; "
    "not within prohibited degrees; residence in the district ≥ 30 days immediately "
    "preceding the application"
)

HMA_OFFLINE_STEP_7 = [
    "7",
    "Whether Aadhaar information available? If Yes → e-KYC / Face Authentication on "
    "Bride & Bridegroom details; else Enter Bride & Bridegroom details (manual)",
    "System / Citizen",
    "Per Offline diagram (Aadhaar YES/NO branch); witness details captured with "
    "party particulars",
]

FORM_II_ROW = [
    "Form II",
    "Rule 4(4)",
    "Endorsement on reverse of memorandum and duplicate: date received; serial no.; "
    "page; volume of Register under HMA 1955; Registrar signature",
    "System on SR digital signature (both channels); blank template in Offline printout "
    "pack before endorsement",
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def add_fr_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def find_table_containing(doc: Document, exact: str) -> Table:
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == exact:
                return table
    raise KeyError(f"Table not found containing {exact!r}")


def find_fr_row(table: Table, req_id: str):
    for row in table.rows:
        if row.cells[0].text.strip() == req_id:
            return row
    raise KeyError(f"FR row not found: {req_id}")


def replace_text(text: str) -> str:
    text = text.replace(MARRIAGE_TAKEN_OLD, MARRIAGE_TAKEN_NEW)
    text = text.replace(MARRIAGE_TAKEN_BRANCH_OLD, MARRIAGE_TAKEN_BRANCH_NEW)
    return text


def replace_in_document(doc: Document) -> None:
    for p in doc.paragraphs:
        new = replace_text(p.text)
        if new != p.text:
            set_para_text(p, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                new = replace_text(cell.text)
                if new != cell.text:
                    set_cell_text(cell, new)


def insert_row_after(table: Table, row_index: int, values: list[str]) -> None:
    new_tr = deepcopy(table.rows[row_index]._tr)
    table._tbl.insert(row_index + 1, new_tr)
    set_row(table, row_index + 1, values)


def insert_form_ii_mapping_row(table: Table) -> None:
    """Insert Form II between Form IA and Form II-A; preserve Form I → IA → II → II-A order."""
    if len(table.rows) < 2:
        raise ValueError("Forms mapping table has no data rows")
    template_tr = deepcopy(table.rows[1]._tr)
    rows_by_form: dict[str, list[str]] = {}
    for row in table.rows[1:]:
        key = row.cells[0].text.strip()
        rows_by_form[key] = [c.text.strip() for c in row.cells]
    rows_by_form["Form II"] = FORM_II_ROW
    ordered = ["Form I", "Form IA", "Form II", "Form II-A"]
    missing = [f for f in ordered if f not in rows_by_form]
    if missing:
        raise KeyError(f"Missing forms in mapping table: {missing}")
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for form in ordered:
        table._tbl.append(deepcopy(template_tr))
        set_row(table, len(table.rows) - 1, rows_by_form[form])


def replace_in_all_cells(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    set_cell_text(cell, cell.text.replace(old, new))


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Cover + version history
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.11")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-01")
    add_version_row(
        doc.tables[1],
        [
            "1.11",
            "2026-09-01",
            "Nandha Kumar",
            "OTP reason in e-KYC/Face Authentication SMS; Hindu Marriage Offline Aadhaar "
            "e-KYC branch (§7.1.2.3); notice wording 'already taken place'; FR-SMA-003 "
            "notice-filing-date age gate; Form II in Hindu Online/Offline registration",
            "Prashanth",
        ],
    )

    # Global wording: Whether Marriage already taken place or not?
    replace_in_document(doc)

    # §2 scope — statutory artefacts
    set_para_text(
        find_para(
            doc,
            contains="Statutory artefacts: Form I (Memorandum), Form IA (Application), Form II-A (Certificate)",
        ),
        "Statutory artefacts: Form I (Memorandum), Form IA (Application), Form II "
        "(Endorsement — Rule 4(4)), Form II-A (Certificate)",
    )

    # §3.3 / forms mapping — insert Form II
    insert_form_ii_mapping_row(doc.tables[4])

    # §7.1.2.1 common intake — offline e-KYC reference
    set_para_text(
        find_para(
            doc,
            contains="Enter / capture Marriage details, Bride details, Bridegroom details, Witness details",
        ),
        "Enter / capture Marriage details, Bride details, Bridegroom details, Witness details "
        "— persisted to the application record. Online channel: e-KYC / Face Authentication "
        "on Bride details (see 7.1.2.2). Offline channel: whether Aadhaar information is "
        "available — if yes, e-KYC / Face Authentication on Bride & Bridegroom details; if no, "
        "manual data entry (see 7.1.2.3). This step is the re-entry point for SR rejection "
        "loops that return to citizen data entry in both diagrams.",
    )

    # §7.1.2.2 Online key characteristics — Form II
    set_para_text(
        find_para(
            doc,
            contains="Key characteristics: channel before combined prerequisite+declaration; e-KYC / Face Authentication",
        ),
        "Key characteristics: channel before combined prerequisite+declaration; e-KYC / "
        "Face Authentication on Bride details during capture; office selection + summary; "
        "Form I & Form IA + eSign; SR digital signature applies Form II endorsement and "
        "issues Form II-A certificate; no printout, no appointment, no DEO; single SR "
        "verification; payment only after SR approval; fully digital signature chain.",
    )

    # §7.1.2.3 Offline — e-KYC branch + Form II printout
    set_para_text(
        find_para(
            doc,
            contains="Key characteristics: channel before combined prerequisite+declaration; SR Verification Stage 1",
        ),
        "Key characteristics: channel before combined prerequisite+declaration; whether "
        "Aadhaar available — e-KYC / Face Authentication on Bride & Bridegroom or manual "
        "data entry (Aadhaar YES/NO branch); SR Verification Stage 1 on captured data "
        "before payment; payment + appointment bundled; printout of Form I, Form IA, "
        "Form II (blank endorsement) & Form II-A; SR allocates to DEO; DEO signature check "
        "and upload; SR Verification Stage 2 on uploaded signed forms (reject returns to "
        "DEO); SR DSC applies Form II endorsement then certificate.",
    )

    # Hindu Offline flow table — insert step 7 (Aadhaar / e-KYC branch)
    insert_row_after(doc.tables[11], 0, HMA_OFFLINE_STEP_7)

    # Hindu Offline printout step — Form II in printout
    for row in doc.tables[11].rows:
        if row.cells[0].text.strip() == "10" and "Printout" in row.cells[1].text:
            set_cell_text(
                row.cells[1],
                "Printout taken on Form I, Form IA, Form II (blank endorsement) and Form II-A",
            )
            set_cell_text(
                row.cells[3],
                "Citizen prints the statutory forms per Karnataka Hindu Marriage forms",
            )
            break

    # Hindu Online certificate step — Form II endorsement
    for row in doc.tables[10].rows:
        if row.cells[0].text.strip() == "14":
            set_cell_text(
                row.cells[1],
                "Form II endorsement applied; marriage certificate (Form II-A) generated",
            )
            set_cell_text(
                row.cells[3],
                "Form II endorsement (Rule 4(4)) then Form II-A available for download",
            )
            break

    # Hindu status model — offline e-KYC in Details captured
    set_cell_text(
        doc.tables[12].rows[4].cells[1],
        "Marriage / bride / bridegroom / witness details saved (Online: e-KYC / Face "
        "Authentication on Bride details; Offline: e-KYC / Face Authentication on Bride "
        "& Bridegroom when Aadhaar available, else manual)",
    )

    # §8.1 section headers — Form II
    set_para_text(
        find_para(doc, exact="8.1.14 Offline channel — printout, DEO upload"),
        "8.1.14 Offline channel — printout (Form I, Form IA, Form II), DEO upload",
    )
    set_para_text(
        find_para(doc, exact="8.1.16 Digital signature and certificate issuance"),
        "8.1.16 Digital signature, Form II endorsement and certificate issuance (Form II-A)",
    )

    # FR-SMA-003
    fr003 = find_fr_row(find_table_containing(doc, "FR-SMA-003"), "FR-SMA-003")
    set_cell_text(fr003.cells[1], FR_SMA_003_NEW)

    # BR-SMA-008 — align with notice filing date
    br008 = find_fr_row(find_table_containing(doc, "BR-SMA-008"), "BR-SMA-008")
    set_cell_text(
        br008.cells[1],
        "Other Forms: both parties must have completed 21 years at the notice filing date",
    )

    # FR-HMA-061 — Form II in offline printout pack
    fr061 = find_fr_row(find_table_containing(doc, "FR-HMA-061"), "FR-HMA-061")
    set_cell_text(
        fr061.cells[1],
        "System shall generate a printout of Form I, Form IA and Form II (blank endorsement "
        "template) with exact statutory wordings",
    )
    set_cell_text(
        fr061.cells[3],
        "Legal sign-off on templates; Form II per Rule 4(4) (`hindu marriage forms.pdf` p.4); "
        "Kannada rendering correct",
    )

    # FR-HMA-080 — Form II endorsement before certificate
    fr080 = find_fr_row(find_table_containing(doc, "FR-HMA-080"), "FR-HMA-080")
    set_cell_text(
        fr080.cells[1],
        "On SR digital signature: assign serial no., page, volume; generate Form II "
        "endorsement per Rule 4(4); update register; then issue Form II-A (certificate)",
    )
    set_cell_text(
        fr080.cells[3],
        "Form II populated with date received, serial/page/volume and Registrar signature; "
        "then Form II-A issued — both channels",
    )

    # New FR-HMA-089 — OTP reason during e-KYC / Face Authentication
    hma_online = find_table_containing(doc, "FR-HMA-058")
    add_fr_row(
        hma_online,
        [
            "FR-HMA-089",
            "During Aadhaar e-KYC / Face Authentication for bride and bridegroom, the OTP "
            "SMS (or equivalent one-time code message) sent to the party shall state the "
            "reason why the OTP is requested (e.g. identity verification for Hindu Marriage "
            "registration application)",
            "Must",
            "OTP / SMS template reviewed by department; reason text visible before code entry; "
            "applies to Hindu Marriage Online and Offline e-KYC paths",
        ],
    )

    # New FR-HMA-090 — Offline Aadhaar / e-KYC branch
    hma_offline = find_table_containing(doc, "FR-HMA-059")
    add_fr_row(
        hma_offline,
        [
            "FR-HMA-090",
            "Offline channel: system shall ask whether Aadhaar information is available for "
            "bride and bridegroom; if yes, perform e-KYC / Face Authentication; if no, "
            "allow manual capture of bride and bridegroom particulars (per Offline diagram "
            "7.1.2.3)",
            "Must",
            "Aadhaar YES/NO decision node before SR Verification Stage 1; same validation "
            "rules as manual path; failure fallback per RS-MRG-002",
        ],
    )

    # New FR-SMA-066 — OTP reason for Special Marriage e-KYC
    sma_ekyc = find_table_containing(doc, "FR-SMA-009")
    add_fr_row(
        sma_ekyc,
        [
            "FR-SMA-066",
            "During Aadhaar e-KYC / Face Authentication for bride and bridegroom in Special "
            "Marriage notice generation, the OTP SMS (or equivalent one-time code message) "
            "sent to the party shall state the reason why the OTP is requested (e.g. "
            "identity verification for Special Marriage notice application)",
            "Must",
            "OTP / SMS template reviewed by department; reason text visible before code entry; "
            "applies to Intended Marriage and Other Forms notice channels (Online and Offline)",
        ],
    )

    # Remaining Form II references in stakeholders / UI / integrations
    for old, new in (
        (
            "physical signature on printed Form I / IA / II-A",
            "physical signature on printed Form I / IA / II / II-A",
        ),
        (
            "Check signatures on printed Form I / IA / II-A and upload to portal",
            "Check signatures on printed Form I / IA / II / II-A and upload to portal",
        ),
        ("Printout — Form I, IA, II-A", "Printout — Form I, IA, II, II-A"),
        ("Form I, IA, II-A", "Form I, IA, II, II-A"),
        (
            "DEO-uploaded signed Form I, Form IA & II-A",
            "DEO-uploaded signed Form I, Form IA, Form II & II-A",
        ),
    ):
        replace_in_all_cells(doc, old, new)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    # Verification
    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    checks = [
        ("FR-SMA-003", FR_SMA_003_NEW[:60]),
        ("FR-HMA-089", "reason why the OTP"),
        ("FR-HMA-090", "Aadhaar information is available"),
        ("FR-SMA-066", "reason why the OTP"),
    ]
    for rid, needle in checks:
        row = find_fr_row(find_table_containing(doc2, rid), rid)
        ok = needle in row.cells[1].text
        print(f"  {rid}: {'OK' if ok else 'MISSING'}")
    form_ii_rows = [
        r.cells[0].text.strip()
        for r in doc2.tables[4].rows
        if r.cells[0].text.strip() == "Form II"
    ]
    print(f"  Form II mapping rows: {len(form_ii_rows)}")
    offline_step7 = [
        r.cells[1].text[:50]
        for r in doc2.tables[11].rows
        if r.cells[0].text.strip() == "7"
    ]
    print(f"  Offline step 7: {offline_step7[0] if offline_step7 else 'MISSING'}...")
    count_new = sum(
        1 for p in doc2.paragraphs if MARRIAGE_TAKEN_NEW in p.text
    )
    count_old = sum(
        1 for p in doc2.paragraphs if MARRIAGE_TAKEN_OLD in p.text
    )
    print(f"  'already taken place' paragraphs: {count_new}; old wording left: {count_old}")


if __name__ == "__main__":
    main()
