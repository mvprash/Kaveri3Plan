# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.6.docx from v1.5.

Align Special Marriage §7 process narratives, channel/status/step tables,
embedded figures, and conflicting FR-SMA / BR-SMA rows with the updated
Special Marriage process diagrams under Process Diagrams/Special Marriage/.
"""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.5.docx"
DST = BASE / "BRD_Marriage_v1.6.docx"
DIAG = BASE / "Process Diagrams" / "Special Marriage"

IMG_ONLINE = DIAG / (
    "Special Marriage (Intended Marriage Notice) 2-SpecialMarriageNoticeOnline.drawio (1).png"
)
IMG_OFFLINE = DIAG / (
    "Special Marriage (Intended Marriage Notice) 2-SpecialMarriageNoticeOffline.drawio (1).png"
)
IMG_REG_INT = DIAG / (
    "Special Marriage (Intended Marriage Notice) 2-SpecialMarriageRegistration_IntendedMarriage.drawio.png"
)
IMG_REG_OTH = DIAG / (
    "Special Marriage (Intended Marriage Notice) 2-SpecialMarriageRegistration_Others.drawio.png"
)

NEW_PROC_LINKS = (
    "Process Diagrams/hindu marriage Online.png, "
    "Process Diagrams/hindu marriage Offline.png, "
    "Process Diagrams/Special Marriage/Special Marriage (Intended Marriage Notice) "
    "2-SpecialMarriageNoticeOnline.drawio (1).png, "
    "Process Diagrams/Special Marriage/Special Marriage (Intended Marriage Notice) "
    "2-SpecialMarriageNoticeOffline.drawio (1).png, "
    "Process Diagrams/Special Marriage/Special Marriage (Intended Marriage Notice) "
    "2-SpecialMarriageRegistration_IntendedMarriage.drawio.png, "
    "Process Diagrams/Special Marriage/Special Marriage (Intended Marriage Notice) "
    "2-SpecialMarriageRegistration_Others.drawio.png"
)


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        set_cell_text(row.cells[ci], val)


def clear_extra_rows(table: Table, keep_rows: int) -> None:
    """Remove trailing rows so table has exactly keep_rows (incl. header)."""
    tbl = table._tbl
    while len(table.rows) > keep_rows:
        tbl.remove(table.rows[-1]._tr)


def ensure_rows(table: Table, needed: int) -> None:
    """Ensure table has at least `needed` rows by cloning the last data row."""
    while len(table.rows) < needed:
        last = table.rows[-1]._tr
        table._tbl.append(deepcopy(last))


def replace_step_table(table: Table, rows: list[list[str]]) -> None:
    """Replace data rows; rows includes header as rows[0]."""
    ensure_rows(table, len(rows))
    for ri, vals in enumerate(rows):
        set_row(table, ri, vals)
    clear_extra_rows(table, len(rows))


def add_version_row(table: Table, values: list[str]) -> None:
    last = table.rows[-1]._tr
    new_tr = deepcopy(last)
    table._tbl.append(new_tr)
    set_row(table, len(table.rows) - 1, values)


def replace_image_blob(doc: Document, r_id: str, png_path: Path) -> None:
    part = doc.part.related_parts[r_id]
    part._blob = png_path.read_bytes()


# --- Step table content -------------------------------------------------------

ONLINE_NOTICE_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    [
        "8",
        "If Aadhaar available → e-KYC on Bride & Bridegroom details; else Enter Bride & Bridegroom details",
        "System / Citizen",
        "Per Online notice diagram (Aadhaar YES/NO branch)",
    ],
    [
        "9",
        "Review summary and proceed document uploading",
        "System / Citizen",
        "Summary of captured particulars",
    ],
    [
        "10",
        "Upload Identity Proof, Photo, Age Proof, Address Proof (Bridegroom & Bride)",
        "Citizen",
        "Mandatory supporting documents including individual photographs",
    ],
    [
        "11",
        "SR Verification (decision)",
        "Sub Registrar",
        "Approve or Reject",
    ],
    [
        "11a",
        "Reject → return to Prerequisite & declaration",
        "Sub Registrar → Citizen",
        "Citizen corrects and resubmits",
    ],
    [
        "12",
        "First Payment",
        "System / Citizen",
        "Notice fee per Karnataka Special Marriage fee schedule",
    ],
    [
        "13",
        "Notice Generated",
        "System",
        "Statutory notice generated after first payment (no Online DEO / appointment)",
    ],
    [
        "14",
        "Marriage notice displayed in portal",
        "System",
        "Sec. 6 publication (digital)",
    ],
    [
        "15",
        "Proceed with e-sign",
        "Citizen",
        "Citizen eSign on the generated notice after portal display",
    ],
    [
        "16",
        "30-day countdown starts",
        "System",
        "Objection period per Sec. 7 / Sec. 16",
    ],
]

OFFLINE_NOTICE_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    [
        "8",
        "If Aadhaar available → e-KYC on Bride & Bridegroom details; else Enter Bride & Bridegroom details",
        "System / Citizen",
        "Per Offline notice diagram (Aadhaar YES/NO branch)",
    ],
    [
        "9",
        "Review summary and proceed document uploading",
        "System / Citizen",
        "Summary of captured particulars",
    ],
    [
        "10",
        "Upload Identity / Age / Address proofs (Bridegroom & Bride)",
        "Citizen",
        "Mandatory documents",
    ],
    [
        "11",
        "SR Verification (decision)",
        "Sub Registrar",
        "Approve or Reject",
    ],
    [
        "11a",
        "Reject → return to Prerequisite & declaration",
        "Sub Registrar → Citizen",
        "Citizen corrects and resubmits",
    ],
    ["12", "First Payment", "System / Citizen", "Notice fee"],
    ["13", "Schedule appointment with SR", "System / Citizen", "After first payment"],
    ["14", "SR Generates Notice", "Sub Registrar", "Statutory notice generation"],
    ["15", "Selects DEO", "Sub Registrar", "Assign FDA/SDA/DEO"],
    [
        "16",
        "Capture individual photos of Bride & Bridegroom",
        "FDA / SDA / DEO",
        "Office visit activity",
    ],
    [
        "17",
        "Download, Print, Sign, Scan and Upload notice",
        "FDA / SDA / DEO",
        "Physical notice processed into system",
    ],
    [
        "18",
        "Paste form on respective Notice Board",
        "FDA / SDA / DEO",
        "Sec. 6(2) conspicuous place publication",
    ],
    ["19", "30-day countdown starts", "System", "Objection period per Sec. 7 / Sec. 16"],
]

REG_INTENDED_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    ["1", "Citizen Login portal", "Citizen", "Continue from published notice"],
    ["2", "Select Notice", "Citizen", "Choose approved / published notice"],
    [
        "3",
        "Validate timeline ≥ 30 days and ≤ 90 days",
        "System",
        "If NO → No Action allowed (Sec. 7 / Sec. 14)",
    ],
    ["4", "If any Objection?", "System", "Branch to SR enquiry when objection exists"],
    ["4a", "Conduct enquiry by summoning all parties", "Sub Registrar", "Sec. 8–9"],
    [
        "4b",
        "Valid objection → update reason; Notice removal from portal (Objected)",
        "Sub Registrar / System",
        "Process stops for solemnization",
    ],
    ["4c", "Objection invalid → continue", "Sub Registrar", "Proceed to Second Payment"],
    ["5", "Second Payment", "System / Citizen", "Registration / solemnization fee"],
    ["6", "Schedule Visit", "System / Citizen", "Office visit for solemnization"],
    ["7", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    ["7a", "Reject → return to Schedule Visit", "Sub Registrar → System", "Reschedule / correct"],
    ["8", "Assigns to DEO", "Sub Registrar", "Certificate / declaration production"],
    ["9", "Joint Photo capturing", "FDA / SDA / DEO", ""],
    [
        "10",
        "If Aadhaar available for Witness → e-KYC Witness Details; else Enter Witness Details",
        "FDA / SDA / DEO",
        "Three witnesses (Sec. 11) — after joint photo",
    ],
    ["11", "Generates Declaration", "FDA / SDA / DEO", "Third Schedule declarations"],
    ["12", "Sign & Upload Declaration", "FDA / SDA / DEO", "Signed declaration uploaded"],
    ["13", "Marriage solemnization", "Sub Registrar", "Sec. 12; declarations Sec. 11"],
    ["14", "Generate Marriage Certificate", "FDA / SDA / DEO", "Fourth Schedule"],
    ["15", "Capture signs of Bride, Bridegroom, Witness", "FDA / SDA / DEO", ""],
    ["16", "Scan signed copy", "FDA / SDA / DEO", ""],
    ["17", "Digital Signature (DSC)", "Sub Registrar", ""],
    ["18", "Marriage Certificate Issued", "Citizen / System", "Download from portal"],
]

REG_OTHERS_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    ["1", "Citizen Login portal", "Citizen", "Continue from published notice"],
    ["2", "Select Notice", "Citizen", "Choose approved / published notice"],
    [
        "3",
        "Validate timeline ≥ 30 days",
        "System",
        "If NO → No Action allowed (Other Forms diagram — no ≤ 90 gate)",
    ],
    ["4", "If any Objection?", "System", "Branch to SR enquiry when objection exists"],
    ["4a", "Conduct enquiry by summoning all parties", "Sub Registrar", "Sec. 8–9"],
    [
        "4b",
        "Valid objection → update reason; Notice removal from portal (Objected)",
        "Sub Registrar / System",
        "Process stops for registration",
    ],
    ["4c", "Objection invalid → continue", "Sub Registrar", "Proceed to Second Payment"],
    ["5", "Second Payment", "System / Citizen", "Registration fee"],
    ["6", "Schedule Visit", "System / Citizen", "Office visit for registration"],
    ["7", "SR Verification (decision)", "Sub Registrar", "Approve or Reject"],
    ["7a", "Reject → return to Schedule Visit", "Sub Registrar → System", "Reschedule / correct"],
    ["8", "Assigns to DEO", "Sub Registrar", "Certificate / declaration production"],
    ["9", "Joint Photo capturing", "FDA / SDA / DEO", ""],
    [
        "10",
        "If Aadhaar available for Witness → e-KYC Witness Details; else Enter Witness Details",
        "FDA / SDA / DEO",
        "Three witnesses — after joint photo",
    ],
    ["11", "Generates Declaration", "FDA / SDA / DEO", "Declarations for registration"],
    ["12", "Sign & Upload Declaration", "FDA / SDA / DEO", "Signed declaration uploaded"],
    [
        "13",
        "Marriage registration / solemnization (per Other Forms diagram)",
        "Sub Registrar",
        "Chapter III Sec. 15–16",
    ],
    ["14", "Generate Marriage Certificate", "FDA / SDA / DEO", "Fifth Schedule (Sec. 16)"],
    ["15", "Capture signs of Bride, Bridegroom, Witness", "FDA / SDA / DEO", ""],
    ["16", "Scan signed copy", "FDA / SDA / DEO", ""],
    ["17", "Digital Signature (DSC)", "Sub Registrar", ""],
    ["18", "Marriage Certificate Issued", "Citizen / System", "Download from portal"],
]


def main() -> None:
    for p in (IMG_ONLINE, IMG_OFFLINE, IMG_REG_INT, IMG_REG_OTH):
        if not p.exists():
            raise FileNotFoundError(p)

    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # --- Version control ------------------------------------------------------
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.6")
    set_cell_text(doc.tables[0].rows[12].cells[1], "2026-08-25")
    add_version_row(
        doc.tables[1],
        [
            "1.6",
            "2026-08-25",
            "Nandha Kumar",
            "Aligned §7.2–7.5 Special Marriage process steps, channel/status tables, "
            "embedded figures and FR-SMA-013/029/034/037/039 / BR-SMA-003 with updated "
            "Special Marriage Notice Online/Offline and Registration (Intended / Other Forms) diagrams",
            "Prashanth",
        ],
    )

    # --- Replace embedded figures (rId8 Online, rId9 Offline, rId10 Intended reg, rId11 Others reg)
    replace_image_blob(doc, "rId8", IMG_ONLINE)
    replace_image_blob(doc, "rId9", IMG_OFFLINE)
    replace_image_blob(doc, "rId10", IMG_REG_INT)
    replace_image_blob(doc, "rId11", IMG_REG_OTH)

    # --- Related docs & glossary paths ----------------------------------------
    set_cell_text(doc.tables[2].rows[2].cells[2], NEW_PROC_LINKS)
    set_cell_text(
        doc.tables[7].rows[12].cells[2],
        "Process Diagrams/Special Marriage/…NoticeOnline…png / …NoticeOffline…png",
    )
    set_cell_text(
        doc.tables[7].rows[13].cells[2],
        "Process Diagrams/Special Marriage/…Registration_IntendedMarriage…png",
    )
    set_cell_text(
        doc.tables[7].rows[14].cells[2],
        "Process Diagrams/Special Marriage/…Registration_Others…png",
    )

    # Appendix / approved diagram references
    set_para_text(
        doc.paragraphs[488],
        "Approved process diagram — Special Marriage (Intended Marriage/Other Forms) Notice Online — "
        "Process Diagrams/Special Marriage/Special Marriage (Intended Marriage Notice) "
        "2-SpecialMarriageNoticeOnline.drawio (1).png",
    )
    set_para_text(
        doc.paragraphs[489],
        "Approved process diagram — Special Marriage (Intended Marriage/Other Forms) Notice Offline — "
        "Process Diagrams/Special Marriage/Special Marriage (Intended Marriage Notice) "
        "2-SpecialMarriageNoticeOffline.drawio (1).png",
    )
    set_para_text(
        doc.paragraphs[490],
        "Approved process diagram — Special Marriage (Intended Marriage) Marriage Registration — "
        "Process Diagrams/Special Marriage/Special Marriage (Intended Marriage Notice) "
        "2-SpecialMarriageRegistration_IntendedMarriage.drawio.png",
    )
    set_para_text(
        doc.paragraphs[491],
        "Approved process diagram — Special Marriage (Intended Marriage/Other Forms) Notice Online "
        "(shared with Intended) — Process Diagrams/Special Marriage/Special Marriage "
        "(Intended Marriage Notice) 2-SpecialMarriageNoticeOnline.drawio (1).png",
    )
    set_para_text(
        doc.paragraphs[492],
        "Approved process diagram — Special Marriage (Intended Marriage/Other Forms) Notice Offline "
        "(shared with Intended) — Process Diagrams/Special Marriage/Special Marriage "
        "(Intended Marriage Notice) 2-SpecialMarriageNoticeOffline.drawio (1).png",
    )
    set_para_text(
        doc.paragraphs[493],
        "Approved process diagram — Special Marriage Other Forms Marriage Registration — "
        "Process Diagrams/Special Marriage/Special Marriage (Intended Marriage Notice) "
        "2-SpecialMarriageRegistration_Others.drawio.png",
    )

    # --- §7.2 Intended Notice narratives --------------------------------------
    set_para_text(
        doc.paragraphs[236],
        "Enter / capture Bride details, Bridegroom details — persisted to the notice application "
        "record. Marriage details (date / place / ceremony of an already solemnized marriage) are "
        "not captured for Intended Marriage notice. SR rejection returns to Prerequisite & "
        "declaration for correction and resubmission.",
    )
    set_para_text(
        doc.paragraphs[239],
        "Figure: Special Marriage (Intended Marriage/Other Forms) Notice Generation — Online",
    )
    set_para_text(
        doc.paragraphs[242],
        "Key characteristics: e-KYC on Bride & Bridegroom when Aadhaar available; document upload "
        "including individual photos; SR verification before First Payment; System generates notice "
        "after payment and displays it on the portal; citizen e-sign on the generated notice; "
        "no Online appointment or DEO; 30-day countdown (Sec. 7).",
    )
    set_para_text(
        doc.paragraphs[245],
        "Figure: Special Marriage (Intended Marriage/Other Forms) Notice Generation — Offline",
    )
    set_para_text(
        doc.paragraphs[248],
        "Key characteristics: Aadhaar YES (e-KYC) / NO (manual) capture; SR verification → First "
        "Payment → appointment; SR generates notice & selects DEO; individual photo capture; "
        "physical notice-board paste (Sec. 6(2)); 30-day countdown. SR rejection returns to "
        "Prerequisite & declaration.",
    )

    # --- §7.3 Intended Registration narratives --------------------------------
    set_para_text(
        doc.paragraphs[264],
        "Schedule Visit — office visit for solemnization / registration.",
    )
    set_para_text(
        doc.paragraphs[265],
        "Office visit continues per 7.3.2.2 — SR verification, DEO joint photo, witness e-KYC/"
        "manual capture, generate & sign/upload declaration, solemnization, certificate and SR DSC.",
    )
    set_para_text(
        doc.paragraphs[271],
        "Key characteristics: notice selection with ≥30 and ≤90 day gate; objection enquiry branch "
        "(valid objection removes notice); Second Payment; schedule visit; SR verification; DEO "
        "joint photo then witness capture (e-KYC or manual), declaration generate/sign/upload, "
        "solemnization, certificate & signatures; SR DSC; certificate issuance (Fourth Schedule).",
    )

    # --- §7.4 Other Forms Notice narratives -----------------------------------
    set_para_text(
        doc.paragraphs[288],
        "Enter / capture Marriage details, Bride details, Bridegroom details — persisted to the "
        "notice application record. Marriage details (date, place and form of the already celebrated "
        "ceremony) are mandatory for Other Forms notice only. SR rejection returns to Prerequisite "
        "& declaration for correction and resubmission.",
    )
    set_para_text(
        doc.paragraphs[291],
        "Figure: Special Marriage (Intended Marriage/Other Forms) Notice Generation — Online",
    )
    set_para_text(
        doc.paragraphs[294],
        "Key characteristics: Online e-KYC / upload (incl. photos) → SR verification → First Payment "
        "→ System notice generation → portal display → citizen e-sign → 30-day countdown "
        "(Sec. 16). No Online appointment or DEO.",
    )
    set_para_text(
        doc.paragraphs[297],
        "Figure: Special Marriage (Intended Marriage/Other Forms) Notice Generation — Offline",
    )
    set_para_text(
        doc.paragraphs[300],
        "Key characteristics: Offline Aadhaar YES/NO capture/upload → SR verification → First "
        "Payment & appointment → SR notice → DEO photos/print/sign/scan/upload and notice-board "
        "paste → 30-day countdown. SR rejection returns to Prerequisite & declaration.",
    )

    # --- §7.5 Other Forms Registration narratives -----------------------------
    set_para_text(
        doc.paragraphs[313],
        "Validate timeline ≥ 30 days — if outside the window, no action is allowed "
        "(Other Forms registration diagram; no ≤ 90 day upper gate).",
    )
    set_para_text(
        doc.paragraphs[316],
        "Schedule Visit — office visit for registration.",
    )
    set_para_text(
        doc.paragraphs[317],
        "Office visit continues per 7.5.2.2 — SR verification, DEO joint photo, witness e-KYC/"
        "manual capture, generate & sign/upload declaration, registration steps, certificate and SR DSC.",
    )
    set_para_text(
        doc.paragraphs[323],
        "Key characteristics: select notice with ≥30 day gate (no ≤90); objection enquiry branch; "
        "Second Payment; schedule visit; SR verification; DEO joint photo then witness capture, "
        "declaration generate/sign/upload, registration/certificate & signatures; SR DSC; "
        "certificate issuance (Fifth Schedule).",
    )

    # --- Channel model tables -------------------------------------------------
    set_row(
        doc.tables[13],
        1,
        [
            "Special Marriage (Intended Marriage) Notice — Online",
            "e-KYC bride & bridegroom, document upload (incl. photos), First Payment, portal notice, eSign",
            "SR verification; System notice generation (no DEO / appointment)",
            "Online notice",
        ],
    )
    set_row(
        doc.tables[13],
        2,
        [
            "Special Marriage (Intended Marriage) Notice — Offline",
            "Capture details (e-KYC or manual), document upload, First Payment, appointment",
            "SR verification, notice generation, DEO photo/print/sign/scan/upload, paste on notice board",
            "Offline notice",
        ],
    )
    set_row(
        doc.tables[17],
        1,
        [
            "Special Marriage (Intended Marriage) Marriage Registration",
            "Select approved notice (≥30 & ≤90 days), Second Payment, schedule visit, download certificate",
            "Objection enquiry (if any), SR verification, DEO joint photo / witness / declaration / certificate, solemnization, SR DSC",
            "Registration (In Person)",
        ],
    )
    set_row(
        doc.tables[20],
        1,
        [
            "Special Marriage Other Forms Notice — Online",
            "Same Online notice pattern as Intended Marriage (e-KYC/upload/First Payment/portal notice/eSign)",
            "SR verification; System notice generation; portal publication + 30-day countdown",
            "Online notice",
        ],
    )
    set_row(
        doc.tables[20],
        2,
        [
            "Special Marriage Other Forms Notice — Offline",
            "Same Offline notice pattern as Intended Marriage notice generation",
            "SR verification, notice generation, DEO physical publication + 30-day countdown",
            "Offline notice",
        ],
    )
    set_row(
        doc.tables[24],
        1,
        [
            "Special Marriage Other Forms Marriage Registration",
            "Select approved notice (timeline ≥30 days), Second Payment, schedule visit, download certificate",
            "Objection enquiry (if any), SR verification, DEO joint photo / witness / declaration / certificate, SR DSC",
            "Registration (In Person)",
        ],
    )

    # --- Step tables ----------------------------------------------------------
    replace_step_table(doc.tables[14], ONLINE_NOTICE_STEPS)  # Intended Online
    replace_step_table(doc.tables[15], OFFLINE_NOTICE_STEPS)  # Intended Offline
    replace_step_table(doc.tables[21], ONLINE_NOTICE_STEPS)  # Other Forms Online
    replace_step_table(doc.tables[22], OFFLINE_NOTICE_STEPS)  # Other Forms Offline
    replace_step_table(doc.tables[18], REG_INTENDED_STEPS)
    replace_step_table(doc.tables[25], REG_OTHERS_STEPS)

    # --- Status models: visit scheduled / Other Forms entry -------------------
    # T19 Intended Registration
    set_cell_text(
        doc.tables[19].rows[3].cells[1],
        "Solemnization / registration visit booked",
    )
    # T26 Other Forms Registration
    set_cell_text(
        doc.tables[26].rows[1].cells[1],
        "Entry point — notice ≥ 30 days from publication with no valid objection",
    )
    set_cell_text(
        doc.tables[26].rows[3].cells[1],
        "Registration visit booked",
    )

    # --- FR / BR / entity / RTM -----------------------------------------------
    # FR-SMA-013
    set_cell_text(
        doc.tables[45].rows[7].cells[1],
        "Online channel: after System notice generation and portal display, both parties shall "
        "eSign the generated notice",
    )
    set_cell_text(
        doc.tables[45].rows[7].cells[3],
        "eSign artefacts stored immutably with timestamp; eSign occurs after notice publication",
    )

    # FR-SMA-029
    set_cell_text(
        doc.tables[50].rows[1].cells[1],
        "Citizen shall initiate marriage registration by selecting a published notice; system shall "
        "validate timeline — Intended Marriage: ≥ 30 and ≤ 90 days from publication; Other Forms: "
        "≥ 30 days from publication",
    )
    set_cell_text(
        doc.tables[50].rows[1].cells[3],
        "Per §7.3 / §7.5 timeline decision (Intended ≤90; Other Forms ≥30 only)",
    )

    # FR-SMA-034 (witness table + solemnization table)
    set_cell_text(
        doc.tables[51].rows[1].cells[1],
        "At the Sub-Registrar office after joint photo capture, system shall capture exactly three "
        "witnesses via e-KYC when Aadhaar is available, otherwise by manual entry",
    )
    set_cell_text(
        doc.tables[51].rows[1].cells[3],
        "Three witnesses mandatory before declaration / solemnization (not before visit scheduling)",
    )
    set_cell_text(
        doc.tables[53].rows[2].cells[1],
        "At the Sub-Registrar office after joint photo capture, system shall capture exactly three "
        "witnesses via e-KYC when Aadhaar is available, otherwise by manual entry",
    )
    set_cell_text(
        doc.tables[53].rows[2].cells[3],
        "Three witnesses mandatory before declaration / solemnization (not before visit scheduling)",
    )

    # FR-SMA-037
    set_cell_text(
        doc.tables[53].rows[5].cells[1],
        "DEO shall generate the declarations of the parties and the three witnesses in the Third "
        "Schedule form; parties/witnesses shall sign and the signed declaration shall be uploaded "
        "before solemnization / registration",
    )
    set_cell_text(
        doc.tables[53].rows[5].cells[3],
        "Sec. 11 — generate → sign & upload declaration before solemnization (per registration diagram)",
    )

    # FR-SMA-039
    set_cell_text(
        doc.tables[53].rows[6].cells[1],
        "DEO shall capture the joint photograph; after witness capture and signed declaration upload, "
        "generate the Certificate of Marriage in the Fourth Schedule, capture signatures of bride, "
        "bridegroom and three witnesses, scan and upload the signed copy",
    )

    # FR-SMA-054
    set_cell_text(
        doc.tables[58].rows[1].cells[1],
        "System shall send SMS / email on notice submission, SR approval or rejection, notice "
        "publication, objection filed, objection decision, Offline appointment / registration visit "
        "scheduling, and certificate issue",
    )

    # BR-SMA-003
    set_cell_text(
        doc.tables[59].rows[23].cells[1],
        "Registration may proceed only when the notice meets the path-specific timeline — Intended "
        "Marriage: ≥ 30 and ≤ 90 days; Other Forms: ≥ 30 days from publication",
    )
    set_cell_text(
        doc.tables[59].rows[23].cells[2],
        "Sec. 14 + process diagrams",
    )

    # Data entities T60
    set_cell_text(
        doc.tables[60].rows[26].cells[1],
        "Identity, photo, age and address proofs; Online eSign after notice generation / portal display",
    )
    set_cell_text(
        doc.tables[60].rows[29].cells[1],
        "Select published notice — Intended: ≥ 30 and ≤ 90 days; Other Forms: ≥ 30 days",
    )
    set_cell_text(
        doc.tables[60].rows[30].cells[0],
        "Visit scheduling and office witness capture",
    )
    set_cell_text(
        doc.tables[60].rows[30].cells[1],
        "Solemnization visit booking after second payment; three-witness e-KYC/manual after joint photo at office",
    )
    set_cell_text(
        doc.tables[60].rows[31].cells[1],
        "Third Schedule declaration generate/sign/upload at office, then solemnization / registration",
    )

    # RTM T63 FR-SMA-029
    set_cell_text(
        doc.tables[63].rows[5].cells[2],
        "Validate notice age — Intended ≥30 & ≤90 days; Other Forms ≥30 days — before registration",
    )

    doc.save(str(DST))
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
