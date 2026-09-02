# -*- coding: utf-8 -*-
"""Update BRD_Marriage_v1.18.docx in place.

Add 'Implemented in (§7)' column to §3 section, rules and notification tables,
mapping each statutory reference to the To-Be process sub-section where it is
implemented.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.18.docx"
DST = BASE / "BRD_Marriage_v1.18.docx"

COLUMN_HEADER = "Implemented in (§7)"

INTRO = (
    "The Implemented in (§7) column maps each Act section, Rule or notification "
    "to the Future state (To-Be) process sub-section where Kaveri 3.0 implements "
    "the requirement. Items with no citizen-facing To-Be step are marked —."
)

HMA_SECTIONS: dict[str, str] = {
    "Sec. 2": "§7.1.2.1 (service selection / Hindu marriage eligibility)",
    "Sec. 3": "§7.1.2.1 (sapinda and prohibited-relationship capture)",
    "Sec. 5": "§7.1.2.1 (prerequisite, declarations and Form IA)",
    "Sec. 7": "§7.1.2.1 (marriage details / ceremony evidence)",
    "Sec. 8": "§7.1.2.2, §7.1.2.3, §7.1.2.4 (registration process and status)",
    "Sec. 11–12": "— (optional SRO flags; not in To-Be workflow)",
    "Sec. 17": "§7.1.2.1 (no living spouse declaration messaging)",
    "Sec. 18": "— (citizen / SRO awareness only)",
}

SMA_SECTIONS: dict[str, str] = {
    "Sec. 4": "§7.2.2.1, §7.3.2.1 (solemnization eligibility checks)",
    "Sec. 5": "§7.2.2.1–§7.2.2.4 (notice generation and capture)",
    "Sec. 6": "§7.2.2.2, §7.2.2.3 (Marriage Notice Book and publication)",
    "Sec. 7": "§7.2.2.4, §7.3.2.3 (30-day objection window)",
    "Sec. 8–9": "§7.2.2.4, §7.3.2.3 (objection enquiry workflow)",
    "Sec. 11": "§7.3.2.1–§7.3.2.2 (Third Schedule declarations)",
    "Sec. 12–13": "§7.3.2.2 (solemnization and Fourth Schedule certificate)",
    "Sec. 14": "§7.3.2.3 (notice validity / fresh-notice gate)",
    "Sec. 15": "§7.2.2.1, §7.3.2.1 (Other Forms eligibility)",
    "Sec. 16": "§7.2.2.x, §7.3.2.x (Other Forms notice and registration)",
    "Sec. 47–48": "§7.3.2.2, §7.3.2.3 (register entry; extracts — partial back-office)",
}

HMA_RULES: dict[str, str] = {
    "Rule 2": "§7.1.2.1 (glossary-driven eligibility validation)",
    "Rule 3(1) + S.O. 4896": "§7.1.2.2 step 8, §7.1.2.3 (SRO office / jurisdiction routing)",
    "Rule 4(1)": "§7.1.2.2, §7.1.2.3 (Online / Offline filing and office selection)",
    "Rule 4(2)": "§7.1.2.1–§7.1.2.3 (Form IA linked to memorandum)",
    "Rule 4(3)": "§7.1.2.1, §7.1.2.3 (three witnesses; eSign / physical signature)",
    "Rule 4(4)": "§7.1.2.2, §7.1.2.3 (Form II endorsement / register entry)",
    "Rule 4(5)": "§7.1.2.2, §7.1.2.3 (Form II-A certificate issuance)",
    "Rule 5": "— (monthly Form III batch — §8.6 back-office)",
    "Rule 6": "— (Forms IV & V indexes — §8.6 MIS / search)",
    "Rule 6A": "§7.1.2.2, §7.1.2.3 (defect remedy, forward, written refusal)",
    "Rule 8 + Schedule": "§7.1.2.2, §7.1.2.3 (payment and certified-extract fees)",
    "Rule 9": "§7.1.2.2, §7.1.2.3 (Form VI payment receipt)",
    "Rule 10": "— (records retention — NFR / archival policy)",
}

SMA_RULES: dict[str, str] = {
    "Rule 3": "§7.2.2.1 (Marriage Officer office master / display)",
    "Rule 4": "§7.2.2.1–§7.2.2.4 (Second Schedule notice generation)",
    "Rule 6": "§7.2.2.4, §7.3.2.3 (objection enquiry and decision)",
    "Rule 7": "§7.2.2.x, §7.3.2.x (Other Forms Form III notice path)",
    "Rule 8": "§7.3.2.2 (optional out-of-office solemnization fee)",
    "Rule 9": "§7.3.2.2 (Marriage Certificate Book / numbering)",
    "Rule 10": "§7.2.2.x, §7.3.2.x (first and second payment per fee schedule)",
    "Rule 11": "— (quarterly Form IV export — §8.6 back-office)",
}

NOTIFICATIONS: dict[str, str] = {
    "S.O. 4896 / HD 6 CIM 61": "§7.1.2.2, §7.1.2.3 (Sub-Registrar as Registrar routing)",
    "Registration of Hindu Marriages (Karnataka) (Amendment) Rules, 1999": (
        "§7.1.2.1–§7.1.2.3 (Forms IA & II-A; SR scrutiny and refusal)"
    ),
    "G.S.R. 314 / G.S.R. 394 / HD 5 PIM 69": "§7.1.2.2, §7.1.2.3 (Rule 4(1) jurisdiction / memorandum)",
    "RD/48/MNMU/2023 — Registration of Hindu Marriage (Karnataka) (Amendment) Rules, 2024": (
        "§7.1.1, §7.1.2.2 (Online channel; electronic register and certificate)"
    ),
    "Special Marriage (Karnataka) Rules, 1961": "§7.2, §7.3 (Special Marriage notice and registration)",
    "Karnataka Guarantee of Services to Citizens Rules, 2012": (
        "§3.7; §8.7 (GSC acknowledgement and lifecycle — no separate §7 process)"
    ),
}


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
    last: bool = False,
) -> Paragraph:
    matches: list[Paragraph] = []
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            matches.append(p)
        elif contains is not None and contains in t:
            matches.append(p)
    if not matches:
        raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")
    return matches[-1] if last else matches[0]


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_table_by_header(doc: Document, first_cell: str, second_cell: str | None = None) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr0 = table.rows[0].cells[0].text.strip()
        hdr1 = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
        if hdr0 == first_cell and (second_cell is None or hdr1 == second_cell):
            return table
    raise KeyError(f"Table not found: {first_cell!r} / {second_cell!r}")


def find_section_table(doc: Document, third_header: str) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:2] == ["Section", "Topic"] and len(hdr) >= 3 and hdr[2] == third_header:
            return table
    raise KeyError(f"Section table with third header {third_header!r} not found")


def add_column_to_table(
    table: Table,
    header: str,
    values_by_key: dict[str, str],
    key_col: int = 0,
) -> None:
    hdr = [c.text.strip() for c in table.rows[0].cells]
    if len(hdr) >= 4 and hdr[-1] == header:
        for ri in range(1, len(table.rows)):
            key = table.rows[ri].cells[key_col].text.strip()
            set_cell_text(table.rows[ri].cells[-1], values_by_key.get(key, ""))
        return

    for ri, row in enumerate(table.rows):
        new_tc = deepcopy(row.cells[-1]._tc)
        row._tr.append(new_tc)
        cell = row.cells[-1]
        if ri == 0:
            set_cell_text(cell, header)
        else:
            key = row.cells[key_col].text.strip()
            set_cell_text(cell, values_by_key.get(key, ""))


def ensure_section3_intro(doc: Document) -> None:
    heading = find_para(doc, exact="3. Legal and regulatory reference", heading_only=True, last=True)
    nxt = heading._element.getnext()
    if nxt is not None and INTRO in "".join(t.text or "" for t in nxt.iter(qn("w:t"))):
        return
    insert_paragraph_after(heading, INTRO, style="Normal")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if SRC.resolve() != DST.resolve():
        shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.18")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    version_table = doc.tables[1]
    last_change = version_table.rows[-1].cells[3].text.strip()
    if "Implemented in (§7)" not in last_change:
        add_version_row(
            version_table,
            [
                "1.18",
                "2026-09-02",
                "Nandha Kumar",
                "Add 'Implemented in (§7)' column to §3 section, rules and notification tables",
                "Prashanth",
            ],
        )

    ensure_section3_intro(doc)

    hma_sections = find_section_table(doc, "relevance")
    add_column_to_table(hma_sections, COLUMN_HEADER, HMA_SECTIONS)

    sma_sections = find_section_table(doc, "BRD relevance")
    add_column_to_table(sma_sections, COLUMN_HEADER, SMA_SECTIONS)

    hma_rules = find_table_by_header(doc, "Rule", "Requirement")
    if "Special Marriage" in hma_rules.rows[1].cells[1].text:
        raise ValueError("First Rule table is not HMA rules")
    add_column_to_table(hma_rules, COLUMN_HEADER, HMA_RULES)

    sma_rules_tables = [
        t
        for t in doc.tables
        if t.rows
        and t.rows[0].cells[0].text.strip() == "Rule"
        and t.rows[0].cells[1].text.strip() == "Requirement"
        and len(t.rows) > 1
        and "Marriage Officer" in t.rows[1].cells[1].text
    ]
    if len(sma_rules_tables) != 1:
        raise ValueError(f"Expected one SMA rules table, found {len(sma_rules_tables)}")
    add_column_to_table(sma_rules_tables[0], COLUMN_HEADER, SMA_RULES)

    notifications = find_table_by_header(doc, "Instrument", "Date / No.")
    add_column_to_table(notifications, COLUMN_HEADER, NOTIFICATIONS)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())

    hma = find_section_table(doc2, "relevance")
    hdr = [c.text.strip() for c in hma.rows[0].cells]
    assert hdr[-1] == COLUMN_HEADER, hdr
    assert "§7.1.2.1" in hma.rows[1].cells[-1].text

    notif = find_table_by_header(doc2, "Instrument", "Date / No.")
    assert notif.rows[0].cells[-1].text.strip() == COLUMN_HEADER
    assert "§7.1.2.2" in notif.rows[1].cells[-1].text

    sma_rules = [
        t
        for t in doc2.tables
        if t.rows
        and t.rows[0].cells[0].text.strip() == "Rule"
        and t.rows[0].cells[1].text.strip() == "Requirement"
        and len(t.rows) > 1
        and "Marriage Officer" in t.rows[1].cells[1].text
    ][0]

    missing: list[str] = []
    for table, mapping in [
        (hma, HMA_SECTIONS),
        (find_section_table(doc2, "BRD relevance"), SMA_SECTIONS),
        (find_table_by_header(doc2, "Rule", "Requirement"), HMA_RULES),
        (sma_rules, SMA_RULES),
        (notif, NOTIFICATIONS),
    ]:
        for ri in range(1, len(table.rows)):
            key = table.rows[ri].cells[0].text.strip()
            if key not in mapping:
                missing.append(key)
    if missing:
        raise AssertionError(f"Unmapped keys: {missing}")

    print("Verification OK")


if __name__ == "__main__":
    main()
