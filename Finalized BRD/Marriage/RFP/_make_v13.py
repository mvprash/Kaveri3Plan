"""Create BRD_Hindu_Marriage_v1.3.docx from v1.2 — reorganize §8 (Functional requirements)
by service, aligned with §7: Hindu Marriage; Special Marriage Notice Generation;
Special Marriage Marriage Registration; then Post-registration / Notifications / Reports.
Existing FR tables are redistributed under the new hierarchy; TOC and document control updated.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\Prashanth\Official\Kaveri 3.0\Kaveri3Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Hindu_Marriage_v1.2.docx"
DST = BASE / "BRD_Hindu_Marriage_v1.3.docx"

# Old TOC §8 entries (exact text from v1.2)
TOC_OLD = [
    "8. Functional requirements",
    "8.1 Eligibility and module entry (FR-HMA-001–004)",
    "8.2 Jurisdiction and office routing (FR-HMA-005–006)",
    "8.3 Data capture — marriage details (FR-HMA-007–009)",
    "8.4 Data capture — bridegroom (FR-HMA-010)",
    "8.5 Data capture — bride (FR-HMA-011)",
    "8.6 Data capture — witnesses (FR-HMA-012–014)",
    "8.7 Form I — memorandum generation and completeness (FR-HMA-015–017)",
    "8.7A Form IA — application and declarations (FR-HMA-083–085)",
    "8.8 Supporting documents and memorandum (FR-HMA-018–019)",
    "8.9 Fees and payments (FR-HMA-020–025)",
    "8.10 SRO scrutiny and registration (FR-HMA-026–030)",
    "8.11 Post-registration services (FR-HMA-031–035)",
    "8.12 Notifications (FR-HMA-036–038)",
    "8.13 Reports and MIS (FR-HMA-039–045)",
    "8.14 Channel selection and prerequisite acknowledgement (FR-HMA-046–050)",
    "8.15 Online channel — office, Form I & Form IA, eSign (FR-HMA-051–058, 086–087)",
    "8.16 Offline channel — printout, DEO upload (FR-HMA-059–069)",
    "8.17 SR verification (FR-HMA-070–077)",
    "8.18 Digital signature and certificate issuance (FR-HMA-078–082)",
    "8.19 Special Marriage — service selection and eligibility (FR-SMA-001–006)",
    "8.20 Notice of intended marriage — data capture and jurisdiction (FR-SMA-007–013)",
    "8.21 Notice generation, Marriage Notice Book and publication (FR-SMA-014–021)",
    "8.22 Objection management and enquiry (FR-SMA-022–028)",
    "8.23 Notice validity, timeline gate and fresh notice (FR-SMA-029–032)",
    "8.24 Solemnization, declarations and certificate (FR-SMA-033–041)",
    "8.25 Special Marriage Other Forms — application and registration (FR-SMA-042–048)",
    "8.26 Special Marriage fees and payments (FR-SMA-049–053)",
    "8.27 Special Marriage notifications, reports and audit (FR-SMA-054–060)",
]

TOC_NEW = [
    "8. Functional requirements",
    "8.1 Hindu Marriage",
    "8.1.1 Eligibility and module entry (FR-HMA-001–004)",
    "8.1.2 Jurisdiction and office routing (FR-HMA-005–006)",
    "8.1.3 Data capture — marriage details (FR-HMA-007–009)",
    "8.1.4 Data capture — bridegroom (FR-HMA-010)",
    "8.1.5 Data capture — bride (FR-HMA-011)",
    "8.1.6 Data capture — witnesses (FR-HMA-012–014)",
    "8.1.7 Form I — memorandum generation and completeness (FR-HMA-015–017)",
    "8.1.8 Form IA — application and declarations (FR-HMA-083–085)",
    "8.1.9 Supporting documents and memorandum (FR-HMA-018–019)",
    "8.1.10 Fees and payments (FR-HMA-020–025)",
    "8.1.11 SRO scrutiny and registration (FR-HMA-026–030)",
    "8.1.12 Channel selection and prerequisite acknowledgement (FR-HMA-046–050)",
    "8.1.13 Online channel — office, Form I & Form IA, eSign (FR-HMA-051–058, 086–087)",
    "8.1.14 Offline channel — printout, DEO upload (FR-HMA-059–069)",
    "8.1.15 SR verification (FR-HMA-070–077)",
    "8.1.16 Digital signature and certificate issuance (FR-HMA-078–082)",
    "8.2 Special Marriage (Intended Marriage/Other Forms) Notice Generation",
    "8.2.1 Eligibility and module entry (FR-SMA-001–006)",
    "8.2.2 Jurisdiction and office routing",
    "8.2.3 Data capture — marriage details (FR-SMA-007–013)",
    "8.2.4 Data capture — bridegroom",
    "8.2.5 Data capture — bride",
    "8.2.6 Data capture — witnesses",
    "8.2.7 Form I — memorandum generation and completeness",
    "8.2.8 Form IA — application and declarations",
    "8.2.9 Supporting documents and memorandum",
    "8.2.10 Fees and payments (FR-SMA-049–053)",
    "8.2.11 SRO scrutiny and registration",
    "8.2.12 Channel selection and prerequisite acknowledgement",
    "8.2.13 Online channel — office, Form I & Form IA, eSign",
    "8.2.14 Offline channel — printout, DEO upload",
    "8.2.15 SR verification",
    "8.2.16 Notice generation, Marriage Notice Book and publication (FR-SMA-014–021)",
    "8.2.17 Notice validity, timeline gate and fresh notice (FR-SMA-029–032)",
    "8.2.18 Digital signature and certificate issuance",
    "8.3 Special Marriage (Intended Marriage/Other Forms) Marriage Registration",
    "8.3.1 Eligibility and module entry",
    "8.3.2 Data capture — witnesses",
    "8.3.3 Objection management and enquiry (FR-SMA-022–028)",
    "8.3.4 Fees and payments",
    "8.3.5 SRO scrutiny and registration",
    "8.3.6 Solemnization, declarations and certificate (FR-SMA-033–041)",
    "8.3.7 Special Marriage Other Forms — application and registration (FR-SMA-042–048)",
    "8.3.8 Offline channel — printout, DEO upload",
    "8.3.9 SR verification",
    "8.3.10 Digital signature and certificate issuance",
    "8.4 Post-registration services (FR-HMA-031–035)",
    "8.5 Notifications (FR-HMA-036–038, FR-SMA-054)",
    "8.6 Reports and MIS (FR-HMA-039–045, FR-SMA-055–060)",
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def _style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not _style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clear_between(start: Paragraph, end: Paragraph) -> None:
    el = start._p.getnext()
    while el is not None and el is not end._p:
        nxt = el.getnext()
        el.getparent().remove(el)
        el = nxt


def set_doc_control_field(doc: Document, field: str, value: str) -> None:
    for row in doc.tables[0].rows:
        if row.cells[0].text.strip() == field:
            row.cells[1].text = value
            return
    raise KeyError(field)


def append_table_row(table: Table, values: list[str]) -> None:
    row = table.add_row()
    for i, v in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def add_heading(cursor: Paragraph, text: str, level: int) -> Paragraph:
    return insert_paragraph_after(cursor, text, style=f"Heading {level}")


def add_normal(cursor: Paragraph, text: str = "") -> Paragraph:
    return insert_paragraph_after(cursor, text, style="Normal")


def normalize_heading_key(text: str) -> str:
    """Map old Heading 3 text to a stable key."""
    t = text.strip()
    # Prefer explicit number prefixes
    for key in (
        "8.7A",
        "8.10",
        "8.11",
        "8.12",
        "8.13",
        "8.14",
        "8.15",
        "8.16",
        "8.17",
        "8.18",
        "8.19",
        "8.20",
        "8.21",
        "8.22",
        "8.23",
        "8.24",
        "8.25",
        "8.26",
        "8.27",
        "8.1",
        "8.2",
        "8.3",
        "8.4",
        "8.5",
        "8.6",
        "8.7",
        "8.8",
        "8.9",
    ):
        if t.startswith(key + " ") or t == key:
            return key
    raise KeyError(f"Unknown §8 heading: {t!r}")


def extract_s8_blocks(doc: Document) -> dict[str, list]:
    """Return deepcopied body elements for each old §8.x Heading-3 block (excluding the heading)."""
    body = doc.element.body
    children = list(body.iterchildren())

    s8_idx = s9_idx = None
    for i, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        p = Paragraph(child, doc)
        if _style_name(p) == "Heading 2":
            t = p.text.strip()
            if t.startswith("8.") and "Functional" in t:
                s8_idx = i
            elif s8_idx is not None and t.startswith("9."):
                s9_idx = i
                break
    if s8_idx is None or s9_idx is None:
        raise KeyError("Could not locate §8 / §9 Heading 2 boundaries")

    h3_indices: list[tuple[int, str]] = []
    for i in range(s8_idx + 1, s9_idx):
        child = children[i]
        if child.tag != qn("w:p"):
            continue
        p = Paragraph(child, doc)
        if _style_name(p) == "Heading 3":
            h3_indices.append((i, normalize_heading_key(p.text)))

    blocks: dict[str, list] = {}
    for n, (start_i, key) in enumerate(h3_indices):
        end_i = h3_indices[n + 1][0] if n + 1 < len(h3_indices) else s9_idx
        # Content after the Heading 3 itself
        elems = [deepcopy(children[j]) for j in range(start_i + 1, end_i)]
        blocks[key] = elems
        print(f"Extracted {key}: {len(elems)} elements")
    return blocks, s8_idx, s9_idx


def insert_block(cursor: Paragraph, elements: list) -> Paragraph:
    """Append copies of elements after cursor. Return a paragraph at the end of the block."""
    if not elements:
        return cursor
    anchor = cursor._p
    last_p_el = None
    last_node = None
    for el in elements:
        node = deepcopy(el)
        anchor.addnext(node)
        anchor = node
        last_node = node
        if node.tag == qn("w:p"):
            last_p_el = node
    parent = cursor._parent
    if last_node is not None and last_node.tag == qn("w:tbl"):
        # Trailing empty para after table for next insertions
        trailing = insert_paragraph_after(Paragraph(last_p_el, parent) if last_p_el else cursor, "", style="Normal")
        # insert_paragraph_after adds after last_p_el/cursor, not after table — fix position
        if last_p_el is None:
            # no paragraph in block; trailing was after cursor — move after table
            trailing._p.getparent().remove(trailing._p)
            last_node.addnext(trailing._p)
            return trailing
        # Move trailing to immediately after table
        trailing._p.getparent().remove(trailing._p)
        last_node.addnext(trailing._p)
        return trailing
    if last_p_el is not None:
        return Paragraph(last_p_el, parent)
    return cursor


def insert_block_skip_leading_blanks(cursor: Paragraph, elements: list) -> Paragraph:
    """Insert block content, dropping leading empty paragraphs."""
    trimmed = list(elements)
    while trimmed and trimmed[0].tag == qn("w:p"):
        # empty?
        texts = "".join(t.text or "" for t in trimmed[0].iter(qn("w:t")))
        if texts.strip():
            break
        trimmed.pop(0)
    # Also drop trailing empty paras (keep tables)
    while trimmed and trimmed[-1].tag == qn("w:p"):
        texts = "".join(t.text or "" for t in trimmed[-1].iter(qn("w:t")))
        if texts.strip():
            break
        trimmed.pop()
    return insert_block(cursor, trimmed)


def update_toc(doc: Document) -> None:
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == TOC_OLD[0]:
            start_idx = i
            break
    if start_idx is None:
        raise KeyError("TOC start for §8 not found")

    for offset, expected in enumerate(TOC_OLD):
        actual = doc.paragraphs[start_idx + offset].text.strip()
        if actual != expected:
            raise KeyError(
                f"TOC mismatch at offset {offset}: expected {expected!r}, got {actual!r}"
            )

    n_old = len(TOC_OLD)
    n_new = len(TOC_NEW)
    for i in range(min(n_old, n_new)):
        set_para_text(doc.paragraphs[start_idx + i], TOC_NEW[i])

    if n_new > n_old:
        anchor = doc.paragraphs[start_idx + n_old - 1]
        for entry in TOC_NEW[n_old:]:
            anchor = insert_paragraph_after(anchor, entry, style="Normal")
    elif n_new < n_old:
        for i in range(n_new, n_old):
            set_para_text(doc.paragraphs[start_idx + i], "")


def add_sub(
    cursor: Paragraph,
    heading: str,
    *,
    blocks: dict[str, list] | None = None,
    source_key: str | None = None,
    note: str | None = None,
) -> Paragraph:
    cursor = add_heading(cursor, heading, 4)
    if note:
        cursor = add_normal(cursor, note)
    if source_key:
        if blocks is None:
            raise KeyError(f"blocks required for source_key={source_key}")
        cursor = insert_block_skip_leading_blanks(cursor, blocks[source_key])
    return cursor


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)

    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # --- Document control ---
    set_doc_control_field(doc, "Version", "1.3")
    set_doc_control_field(doc, "Last updated", "2026-08-25")
    append_table_row(
        doc.tables[1],
        [
            "1.3",
            "2026-08-25",
            "Nandha Kumar",
            "Reorganized §8 (Functional requirements) by service — Hindu Marriage; "
            "Special Marriage (Intended Marriage/Other Forms) Notice Generation; "
            "Special Marriage (Intended Marriage/Other Forms) Marriage Registration; "
            "Post-registration services; Notifications; Reports and MIS — aligned with §7",
            "Prashanth",
        ],
    )

    # --- Extract blocks BEFORE TOC mutation changes paragraph indices mid-doc? TOC is before body §8.
    # Extract first while structure intact.
    blocks, _, _ = extract_s8_blocks(doc)

    # --- TOC ---
    update_toc(doc)

    # --- Rebuild Section 8 body ---
    # Re-find headings after TOC inserts (body H2 texts unchanged)
    h8 = find_para(doc, exact="8. Functional requirements", heading_only=True)
    h9 = find_para(doc, exact="9. Business rules", heading_only=True)
    clear_between(h8, h9)

    cursor = h8
    cursor = add_normal(
        cursor,
        "Functional requirements are organized by service, aligned with §7 (To-Be). "
        "Each service groups eligibility, data capture, channel, scrutiny and issuance "
        "requirements. Hindu Marriage uses FR-HMA-*; Special Marriage uses FR-SMA-*. "
        "Cross-cutting post-registration, notification and MIS requirements follow the "
        "service-specific blocks.",
    )

    # ========== 8.1 Hindu Marriage ==========
    cursor = add_heading(cursor, "8.1 Hindu Marriage", 3)
    cursor = add_sub(cursor, "8.1.1 Eligibility and module entry", blocks=blocks, source_key="8.1")
    cursor = add_sub(cursor, "8.1.2 Jurisdiction and office routing", blocks=blocks, source_key="8.2")
    cursor = add_sub(cursor, "8.1.3 Data capture — marriage details", blocks=blocks, source_key="8.3")
    cursor = add_sub(cursor, "8.1.4 Data capture — bridegroom", blocks=blocks, source_key="8.4")
    cursor = add_sub(cursor, "8.1.5 Data capture — bride", blocks=blocks, source_key="8.5")
    cursor = add_sub(cursor, "8.1.6 Data capture — witnesses", blocks=blocks, source_key="8.6")
    cursor = add_sub(cursor, "8.1.7 Form I — memorandum generation and completeness", blocks=blocks, source_key="8.7")
    cursor = add_sub(cursor, "8.1.8 Form IA — application and declarations", blocks=blocks, source_key="8.7A")
    cursor = add_sub(cursor, "8.1.9 Supporting documents and memorandum", blocks=blocks, source_key="8.8")
    cursor = add_sub(cursor, "8.1.10 Fees and payments", blocks=blocks, source_key="8.9")
    cursor = add_sub(cursor, "8.1.11 SRO scrutiny and registration", blocks=blocks, source_key="8.10")
    cursor = add_sub(cursor,
        "8.1.12 Channel selection and prerequisite acknowledgement",
        blocks=blocks,
        source_key="8.14",
    )
    cursor = add_sub(cursor,
        "8.1.13 Online channel — office, Form I & Form IA, eSign",
        blocks=blocks,
        source_key="8.15",
    )
    cursor = add_sub(cursor, "8.1.14 Offline channel — printout, DEO upload", blocks=blocks, source_key="8.16")
    cursor = add_sub(cursor, "8.1.15 SR verification", blocks=blocks, source_key="8.17")
    cursor = add_sub(cursor, "8.1.16 Digital signature and certificate issuance", blocks=blocks, source_key="8.18")

    # ========== 8.2 Special Marriage Notice Generation ==========
    cursor = add_heading(
        cursor,
        "8.2 Special Marriage (Intended Marriage/Other Forms) Notice Generation",
        3,
    )
    cursor = add_sub(cursor,
        "8.2.1 Eligibility and module entry",
        blocks=blocks,
        source_key="8.19",
    )
    cursor = add_sub(
        cursor,
        "8.2.2 Jurisdiction and office routing",
        note=(
            "Notice shall be routed to the Marriage Officer of the district in which at least "
            "one party has resided (FR-SMA-008). Full capture and routing requirements are in "
            "§8.2.3 (FR-SMA-007–013)."
        ),
    )
    cursor = add_sub(cursor,
        "8.2.3 Data capture — marriage details",
        blocks=blocks,
        source_key="8.20",
    )
    cursor = add_sub(
        cursor,
        "8.2.4 Data capture — bridegroom",
        note=(
            "Bridegroom particulars (name, condition, age, dwelling place, and supporting proofs) "
            "are captured as part of the Second Schedule notice data in FR-SMA-007 and "
            "FR-SMA-009–011 (§8.2.3)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.5 Data capture — bride",
        note=(
            "Bride particulars and e-KYC / Face Authentication (where Aadhaar is available) are "
            "captured under FR-SMA-007 and FR-SMA-009–011 (§8.2.3)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.6 Data capture — witnesses",
        note=(
            "Witnesses are not required at notice generation. Three witnesses are captured at "
            "marriage registration (see §8.3.2 and FR-SMA-034)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.7 Form I — memorandum generation and completeness",
        note=(
            "Not applicable to Special Marriage. Notice of intended marriage / other-forms "
            "public notice uses the Second Schedule (or prescribed Other Forms notice) — see "
            "§8.2.16 (FR-SMA-014–021) and §8.3.7 (FR-SMA-044)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.8 Form IA — application and declarations",
        note=(
            "Not applicable as HMA Form IA. Combined prerequisite and statutory declarations "
            "for Special Marriage notice are enforced under FR-SMA-004 (§8.2.1)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.9 Supporting documents and memorandum",
        note=(
            "Identity, photograph, age and address proof upload for both parties is required "
            "under FR-SMA-011 (§8.2.3)."
        ),
    )
    cursor = add_sub(cursor,
        "8.2.10 Fees and payments",
        blocks=blocks,
        source_key="8.26",
    )
    cursor = add_sub(
        cursor,
        "8.2.11 SRO scrutiny and registration",
        note=(
            "SR scrutinises the notice application before first payment and notice generation "
            "(see §7.2 / §7.4). Related requirements: FR-SMA-013 (eSign before SR verification), "
            "FR-SMA-014 (notice generation) and FR-SMA-049 (first payment after SR approval)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.12 Channel selection and prerequisite acknowledgement",
        note=(
            "Service path and Online / Offline notice channel are selected before prerequisites "
            "(FR-SMA-001, FR-SMA-004, FR-SMA-005 in §8.2.1). Channel drives publication mode and "
            "office tasks."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.13 Online channel — office, Form I & Form IA, eSign",
        note=(
            "Online notice channel: e-KYC on bride and bridegroom where Aadhaar is available "
            "(FR-SMA-009), eSign of the notice application (FR-SMA-013), and portal publication "
            "of the generated notice (FR-SMA-017). Analogous to Hindu Marriage online Form I / "
            "Form IA + eSign; Special Marriage uses the statutory notice schedules."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.14 Offline channel — printout, DEO upload",
        note=(
            "Offline notice channel: after SR verification and first payment, SR assigns "
            "FDA / SDA / DEO for photograph capture and notice-board publication tasks "
            "(FR-SMA-018). Printout / physical steps follow §7.2.2.3 / §7.4.2.3."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.2.15 SR verification",
        note=(
            "SR verifies the notice application before first payment and notice generation "
            "(§7.2 / §7.4). Rejection returns the citizen to data capture. Approval unlocks "
            "FR-SMA-049 (first payment) and FR-SMA-014 (notice generation)."
        ),
    )
    cursor = add_sub(cursor,
        "8.2.16 Notice generation, Marriage Notice Book and publication",
        blocks=blocks,
        source_key="8.21",
    )
    cursor = add_sub(cursor,
        "8.2.17 Notice validity, timeline gate and fresh notice",
        blocks=blocks,
        source_key="8.23",
    )
    cursor = add_sub(
        cursor,
        "8.2.18 Digital signature and certificate issuance",
        note=(
            "At notice stage, parties eSign the notice application online (FR-SMA-013). "
            "Marriage certificate digital signature and issuance occur at registration "
            "(§8.3.10; FR-SMA-039–041 / FR-SMA-048)."
        ),
    )

    # ========== 8.3 Special Marriage Marriage Registration ==========
    cursor = add_heading(
        cursor,
        "8.3 Special Marriage (Intended Marriage/Other Forms) Marriage Registration",
        3,
    )
    cursor = add_sub(
        cursor,
        "8.3.1 Eligibility and module entry",
        note=(
            "Citizen initiates marriage registration from a published / valid notice "
            "(FR-SMA-029–032 in §8.2.17). Sec. 4 (Intended) / Sec. 15 (Other Forms) conditions "
            "remain enforced from service selection (FR-SMA-002–003, FR-SMA-006)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.3.2 Data capture — witnesses",
        note=(
            "System shall perform e-KYC on three witnesses and capture Third Schedule "
            "declarations with the parties (FR-SMA-034, FR-SMA-037 in §8.3.6)."
        ),
    )
    cursor = add_sub(cursor,
        "8.3.3 Objection management and enquiry",
        blocks=blocks,
        source_key="8.22",
    )
    cursor = add_sub(
        cursor,
        "8.3.4 Fees and payments",
        note=(
            "Second payment (registration / solemnization fee) is collected after the objection "
            "check and before the solemnization visit (FR-SMA-033, FR-SMA-050). Additional fee "
            "applies for solemnization outside the Marriage Officer’s office (FR-SMA-051). "
            "Fee master configurability: FR-SMA-049–053 (§8.2.10)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.3.5 SRO scrutiny and registration",
        note=(
            "SR verifies the registration application before solemnization; rejection returns "
            "the application for correction (FR-SMA-036). On satisfaction, certificate entry "
            "and DSC follow §8.3.6 / §8.3.7."
        ),
    )
    cursor = add_sub(cursor,
        "8.3.6 Solemnization, declarations and certificate",
        blocks=blocks,
        source_key="8.24",
    )
    cursor = add_sub(cursor,
        "8.3.7 Special Marriage Other Forms — application and registration",
        blocks=blocks,
        source_key="8.25",
    )
    cursor = add_sub(
        cursor,
        "8.3.8 Offline channel — printout, DEO upload",
        note=(
            "Marriage registration for Intended Marriage and Other Forms is Offline (In Person) "
            "per §7.3 / §7.5. DEO captures the joint photograph and generates the Certificate of "
            "Marriage for SR digital signature (FR-SMA-039; Other Forms FR-SMA-046–048)."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.3.9 SR verification",
        note=(
            "SR verification before solemnization / final entry is required (FR-SMA-036; "
            "Other Forms objection hearing FR-SMA-045). Channel-aware verification mirrors the "
            "Offline registration diagrams in §7.3.2.2 / §7.5.2.2."
        ),
    )
    cursor = add_sub(
        cursor,
        "8.3.10 Digital signature and certificate issuance",
        note=(
            "SR applies digital signature to the Certificate of Marriage; system enters the "
            "certificate in the Marriage Certificate Book and issues it for citizen download "
            "(FR-SMA-039–041). Other Forms final entry and digital issue: FR-SMA-046–048."
        ),
    )

    # ========== 8.4–8.6 Cross-cutting ==========
    cursor = add_heading(cursor, "8.4 Post-registration services", 3)
    cursor = insert_block_skip_leading_blanks(cursor, blocks["8.11"])

    cursor = add_heading(cursor, "8.5 Notifications", 3)
    cursor = insert_block_skip_leading_blanks(cursor, blocks["8.12"])
    cursor = add_normal(
        cursor,
        "Special Marriage notifications (in addition to Hindu Marriage FR-HMA-036–038):",
    )
    # 8.27 contains notif + reports + audit — place full table here with a note, and again under reports?
    # Better: place full 8.27 under Reports with intro covering notifications pointer,
    # and under Notifications only call out FR-SMA-054.
    cursor = add_normal(
        cursor,
        "FR-SMA-054: System shall send SMS / email on notice submission, SR approval or rejection, "
        "notice publication, objection events, solemnization schedule and certificate issuance "
        "(full Special Marriage notification / reports / audit set: FR-SMA-054–060 — see §8.6).",
    )

    cursor = add_heading(cursor, "8.6 Reports and MIS", 3)
    cursor = insert_block_skip_leading_blanks(cursor, blocks["8.13"])
    cursor = add_normal(
        cursor,
        "Special Marriage notifications, reports and audit (FR-SMA-054–060):",
    )
    cursor = insert_block_skip_leading_blanks(cursor, blocks["8.27"])

    doc.save(str(DST))
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
