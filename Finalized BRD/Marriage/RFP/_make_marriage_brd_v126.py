# -*- coding: utf-8 -*-
"""Create BRD_Marriage_BRD_v1.26.docx — brief Scope without Acts/Rules detail."""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_BRD_v1.25.docx"
DST = BASE / "BRD_Marriage_BRD_v1.26.docx"

# Brief In-scope bullets — no Acts, Rules, Schedules, or statutory section cites.
BRIEF_BULLETS = [
    "Hindu Marriage registration — Online and Offline channels.",
    "Special Marriage — Intended Marriage: notice generation (Online & Offline), "
    "objection period, solemnization and registration.",
    "Special Marriage — Other Forms: notice generation (Online & Offline), "
    "objection handling and registration.",
    "Citizen portal and SRO desk workflows: apply, pay fee, scrutiny, "
    "approve/reject, register, and issue certificate.",
    "Null and Void endorsement of a registered marriage on court order.",
    "Integrations: payment, Aadhaar/e-KYC, DigiLocker, SMS, e-Mail, Kutumba, "
    "Civil Registration System, Labor Department, and Sakala.",
    "Bilingual UI (English + Kannada); audit trail, role-based access, and MIS/reporting.",
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_heading(doc: Document, contains: str) -> Paragraph:
    for p in doc.paragraphs:
        st = str(p.style.name) if p.style else ""
        if st.startswith("Heading") and contains in p.text:
            return p
    raise KeyError(contains)


def replace_in_scope_bullets(doc: Document) -> None:
    heading = find_heading(doc, "In scope (Hindu Marriage & Special Marriage)")
    # Collect following list bullets until next heading
    to_delete: list[Paragraph] = []
    started = False
    for p in doc.paragraphs:
        if p._p is heading._p:
            started = True
            continue
        if not started:
            continue
        st = str(p.style.name) if p.style else ""
        if st.startswith("Heading"):
            break
        if p.text.strip() or "List" in st:
            to_delete.append(p)

    for p in to_delete:
        delete_paragraph(p)

    # Re-find heading after deletes
    heading = find_heading(doc, "In scope (Hindu Marriage & Special Marriage)")
    anchor = heading
    for text in BRIEF_BULLETS:
        anchor = insert_paragraph_after(anchor, text, style="List Bullet")


def update_doc_control(doc: Document) -> None:
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.26")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-03")
    vt = doc.tables[1]
    vt._tbl.append(deepcopy(vt.rows[-1]._tr))
    vals = [
        "1.26",
        "2026-09-03",
        "Nandha Kumar",
        "Scope (§2.i): shorten In-scope bullets — remove Acts, Rules, Schedules "
        "and statutory section detail (kept in §3)",
        "Prashanth",
    ]
    row = vt.rows[-1]
    for ci, val in enumerate(vals):
        set_cell_text(row.cells[ci], val)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))
    replace_in_scope_bullets(doc)
    update_doc_control(doc)
    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    bullets = []
    started = False
    for p in doc2.paragraphs:
        st = str(p.style.name) if p.style else ""
        if st.startswith("Heading") and "In scope" in p.text:
            started = True
            continue
        if started and st.startswith("Heading"):
            break
        if started and p.text.strip():
            bullets.append(p.text.strip())
    assert len(bullets) == len(BRIEF_BULLETS), bullets
    for b in bullets:
        assert "HMA" not in b and "SMA" not in b and "Rule" not in b
        assert "Schedule" not in b and "Section 8" not in b
    print("Verification OK —", len(bullets), "brief bullets")
    for b in bullets:
        print(" -", b)


if __name__ == "__main__":
    main()
