# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.9.docx from v1.8.

Add §16 Risk and Mitigation, §17 System Fallbacks & Error Handling, and
§18 Training and Change Management.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\Prashanth\Official\Kaveri 3.0\Kaveri3Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.8.docx"
DST = BASE / "BRD_Marriage_v1.9.docx"

TOC_ENTRIES = [
    "16. Risk and Mitigation Strategy (RS-MRG-001–003)",
    "17. System Fallbacks & Error Handling (FB-MRG-001–004)",
    "18. Training and Change Management",
    "18.1 Target audience",
    "18.2 Training delivery",
    "18.3 Citizen change management",
    "18.4 Post-Go-Live support",
]

RISK_ROWS = [
    ["Risk ID", "Risk", "Mitigation", "Related requirements"],
    [
        "RS-MRG-001",
        "Users select the wrong marriage legal path because of confusing legal "
        "terminology (for example, confusing Intended Marriage Notice with "
        "Registration of Other Forms, or Hindu Marriage with Special Marriage)",
        "Implement a Needs-Based Wizard on the landing page rather than a legal-act "
        "selector. The wizard asks simple questions (status of the marriage — already "
        "solemnized vs intended; religion / customary form of the parties) and routes "
        "the citizen to the correct service path automatically",
        "FR-HMA-001; FR-SMA-001 / FR-SMA-004 / FR-SMA-005; channel choice FR-HMA-047",
    ],
    [
        "RS-MRG-002",
        "Prolonged downtime of third-party dependencies (e-KYC / Aadhaar, eSign, or "
        "payment gateway) blocks applications mid-flow",
        "Build asynchronous / resumable workflows. If e-KYC / Aadhaar fails, provide "
        "a seamless fallback to manual data entry paired with mandatory document "
        "uploads for Sub-Registrar manual scrutiny. Payment and eSign failures follow §17",
        "FR-SMA-009 / FR-SMA-010; FR-HMA-058; NFR-MRG-AVA-001; FB-MRG-001 / FB-MRG-002",
    ],
    [
        "RS-MRG-003",
        "Misallocation of applications to Data Entry Operators (DEOs) causes blocked "
        "Offline workflows (As-Is pain point 13 — officers cannot reassign without Service Desk)",
        "The Sub-Registrar shall have explicit system controls to reassign, recall or "
        "override DEO allocations without IT Helpdesk intervention (FR-HMA-088)",
        "FR-HMA-069; FR-HMA-088; NFR-MRG-AUD-001",
    ],
]

FALLBACK_ROWS = [
    ["Req ID", "Requirement", "Priority", "Acceptance criteria"],
    [
        "FB-MRG-001",
        "Payment gateway timeouts: if a payment transaction drops or times out, the "
        "system shall automatically poll the Treasury / payment gateway for a final "
        "status. The user shall be locked from making a duplicate payment until a "
        "definitive Failed or Success status is returned",
        "Must",
        "Same control as NFR-MRG-PAY-001; UI pay action disabled until terminal status; "
        "no double-debit; aligns with FR-HMA-025 / FR-SMA-052",
    ],
    [
        "FB-MRG-002",
        "eSign and DSC failures: if a citizen eSign or an SR DSC attempt fails mid-process, "
        "the system shall save the current state as eSign pending or Pending SR digital "
        "signature. The user shall retry signing without re-entering Form I, Form IA, "
        "notice particulars or other already-captured data",
        "Must",
        "Status remains eSign pending / Pending SR digital signature; retry resumes the "
        "signing step only; aligns with FR-HMA-056 / FR-HMA-078–079 and SMA eSign / DSC FRs",
    ],
    [
        "FB-MRG-003",
        "Document upload exceptions: the system shall detect and block password-protected "
        "PDFs and unsupported file formats at upload time, and immediately display a clear "
        "localized (English / Kannada) error message instructing the user how to correct the file",
        "Must",
        "Password-protected and invalid-type files rejected before persist; bilingual "
        "message; closes As-Is pain point 10; aligns with FR-HMA-065",
    ],
    [
        "FB-MRG-004",
        "Notification gateway delays: if the SMS or email gateway is unreachable, the "
        "system shall queue notifications locally and retry periodically. Critical workflow "
        "steps (notice generation, certificate issuance) shall not be halted because "
        "notification delivery failed",
        "Must",
        "Workflow continues; notifications drain from a retry queue; aligns with "
        "NFR-MRG-AVA-001, FR-HMA-036 and FR-SMA-054",
    ],
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def clear_extra_rows(table: Table, keep_rows: int) -> None:
    tbl = table._tbl
    while len(table.rows) > keep_rows:
        tbl.remove(table.rows[-1]._tr)


def ensure_rows(table: Table, needed: int) -> None:
    while len(table.rows) < needed:
        table._tbl.append(deepcopy(table.rows[-1]._tr))


def fill_table(table: Table, rows: list[list[str]]) -> None:
    ensure_rows(table, len(rows))
    for ri, vals in enumerate(rows):
        set_row(table, ri, vals)
    clear_extra_rows(table, len(rows))


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    p_pr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not p_pr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_para = insert_paragraph_after(paragraph, text, style)
    new_para._p.getparent().remove(new_para._p)
    paragraph._p.addprevious(new_para._p)
    return new_para


def find_table_containing(doc: Document, exact: str) -> Table:
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == exact:
                return table
    raise KeyError(f"Table not found containing {exact!r}")


def find_rtm(doc: Document) -> Table:
    for table in doc.tables:
        if len(table.rows[0].cells) >= 4 and table.rows[0].cells[1].text.strip() == "Act/Rule/Form":
            return table
    raise KeyError("RTM table not found")


def insert_table_after(paragraph: Paragraph, source: Table, rows: list[list[str]]) -> Paragraph:
    tbl = deepcopy(source._tbl)
    paragraph._p.addnext(tbl)
    table = Table(tbl, paragraph._parent)
    fill_table(table, rows)
    trailing = insert_paragraph_after(paragraph, "", style="Normal")
    trailing._p.getparent().remove(trailing._p)
    tbl.addnext(trailing._p)
    return Paragraph(trailing._p, paragraph._parent)


def add_fr_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.9")
    set_cell_text(doc.tables[0].rows[12].cells[1], "2026-08-26")
    add_version_row(
        doc.tables[1],
        [
            "1.9",
            "2026-08-26",
            "Nandha Kumar",
            "Added §16 Risk and Mitigation Strategy, §17 System Fallbacks & Error Handling, "
            "and §18 Training and Change Management; added FR-HMA-088 (SR reassign / recall / "
            "override of DEO allocation without Helpdesk)",
            "Prashanth",
        ],
    )

    # TOC after 15.4
    toc15 = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("15.4 Security Audit") and not style_name(p).startswith("Heading"):
            toc15 = p
            break
    if toc15 is None:
        raise KeyError("TOC 15.4 not found")
    cursor = toc15
    for entry in TOC_ENTRIES:
        cursor = insert_paragraph_after(cursor, entry, style="Normal")

    template = find_table_containing(doc, "FR-SMA-001")

    # FR-HMA-088 — DEO reassignment (closes pain point 13; supports RS-MRG-003)
    deo = find_table_containing(doc, "FR-HMA-069")
    add_fr_row(
        deo,
        [
            "FR-HMA-088",
            "Sub-Registrar shall be able to reassign, recall or override a DEO allocation "
            "from the SRO workbench without IT Helpdesk intervention",
            "Must",
            "Reassign / recall audited (actor, from-DEO, to-DEO, timestamp); unblocks "
            "As-Is pain point 13; RS-MRG-003",
        ],
    )

    uat = find_para(doc, contains="UAT scope: Test scenarios derived from FR-HMA-*")
    set_para_text(
        uat,
        "UAT scope: Test scenarios derived from FR-HMA-* and FR-SMA-* (see 13 RTM), "
        "NFR-MRG-* (see 15), fallbacks FB-MRG-* (see 17), BR-HMA-* / BR-SMA-*, statutory "
        "forms Form I / IA / II / II-A, and the Special Marriage Second, Third, Fourth and "
        "Fifth Schedules. Performance, security, availability and VAPT evidence in §15, "
        "and fallback / DEO-reassignment scenarios in §16–17, are Go-Live gates.",
    )

    appendix = find_para(doc, exact="Appendix A — References", heading_only=True)
    cursor = insert_paragraph_before(
        appendix, "16. Risk and Mitigation Strategy", style="Heading 2"
    )
    cursor = insert_paragraph_after(
        cursor,
        "This section outlines potential operational, technical and user-adoption risks "
        "associated with the new Marriage Registration module and the corresponding "
        "mitigation plans. Mitigations are mandatory design constraints, not optional guidance.",
        style="Normal",
    )
    cursor = insert_table_after(cursor, template, RISK_ROWS)

    cursor = insert_paragraph_after(
        cursor, "17. System Fallbacks & Error Handling", style="Heading 2"
    )
    cursor = insert_paragraph_after(
        cursor,
        "To ensure a seamless user experience, the system shall handle exceptions and "
        "integration failures without losing citizen data. §15.3 states the availability "
        "NFR; this section specifies the operational fallbacks for payment, signing, "
        "uploads and notifications.",
        style="Normal",
    )
    cursor = insert_table_after(cursor, template, FALLBACK_ROWS)

    cursor = insert_paragraph_after(
        cursor, "18. Training and Change Management", style="Heading 2"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Transitioning from legacy Kaveri 2.0 to Kaveri 3.0 requires structured change "
        "management so that department staff and citizens adopt the new Marriage "
        "Registration module.",
        style="Normal",
    )

    cursor = insert_paragraph_after(cursor, "18.1 Target audience", style="Heading 3")
    cursor = insert_paragraph_after(
        cursor,
        "Training covers Sub-Registrars (Marriage Officers), First Division Assistants "
        "(FDAs) / Second Division Assistants (SDAs) / Data Entry Operators (DEOs), and "
        "IT Helpdesk personnel.",
        style="Normal",
    )

    cursor = insert_paragraph_after(cursor, "18.2 Training delivery", style="Heading 3")
    cursor = insert_paragraph_after(
        cursor,
        "Conduct phased, role-based workshops for SRO staff focusing on digital DSC "
        "processes, DEO allocation / reassignment queues, and handling the 30-day statutory "
        "countdown for Special Marriages.",
        style="List Bullet",
    )
    cursor = insert_paragraph_after(
        cursor,
        "Distribute Standard Operating Procedures (SOPs), quick-reference cards and video "
        "tutorials for daily office tasks (Online vs Offline, notice publication, objection "
        "enquiry, certificate issue).",
        style="List Bullet",
    )

    cursor = insert_paragraph_after(
        cursor, "18.3 Citizen change management", style="Heading 3"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Deploy tooltips, dynamic help text and the guided questionnaire (Needs-Based "
        "Wizard — RS-MRG-001) to educate citizens on statutory rules — including the "
        "mandatory 30-day notice period for Intended Marriages — directly in the UI "
        "(English and Kannada).",
        style="Normal",
    )

    cursor = insert_paragraph_after(cursor, "18.4 Post-Go-Live support", style="Heading 3")
    cursor = insert_paragraph_after(
        cursor,
        "Establish a dedicated hyper-care IT support channel for the first 90 days after "
        "launch to resolve operational bottlenecks rapidly, particularly digital-signature "
        "(eSign / DSC) configuration and fee reconciliation.",
        style="Normal",
    )

    rtm = find_rtm(doc)
    for vals in (
        [
            "FR-HMA-088",
            "Process / Offline",
            "SR reassign, recall or override DEO allocation without Helpdesk",
            "8.1.14 / 16",
            "SRO workbench",
            "TC-HMA-___",
            "Draft",
        ],
        [
            "FB-MRG-001",
            "Payment / Treasury",
            "Poll gateway on timeout; lock UI against duplicate payment",
            "17",
            "Payment",
            "TC-NFR-___",
            "Draft",
        ],
        [
            "RS-MRG-001",
            "UX / service selection",
            "Needs-Based Wizard routes citizen to the correct marriage path",
            "16 / 18.3",
            "Landing / wizard",
            "TC-HMA-___",
            "Draft",
        ],
    ):
        add_fr_row(rtm, vals)

    doc.save(str(DST))
    print(f"Wrote {DST}")
    doc2 = Document(str(DST))
    print("--- new headings ---")
    for p in doc2.paragraphs:
        t = p.text.strip()
        if t.startswith(("16.", "17.", "18.", "Appendix")):
            print(f"  [{style_name(p)}] {t}")


if __name__ == "__main__":
    main()
