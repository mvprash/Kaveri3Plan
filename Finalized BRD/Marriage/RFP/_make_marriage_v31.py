# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.20.docx from v1.19.

Add §8.1.17 Null and Void Endorsement (Court Order) functional requirements
(FR-HMA-094–099), business rule BR-HMA-021, UI screen, data entity and RTM rows.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.19.docx"
DST = BASE / "BRD_Marriage_v1.20.docx"

FR_ROWS = [
    ["Req ID", "Requirement", "Priority", "Acceptance criteria"],
    [
        "FR-HMA-094",
        "System shall provide an SRO-only back-office function to initiate Null and Void "
        "endorsement on a registered Hindu marriage upon presentation of a certified court "
        "order (HMA 1955 Sec. 11–12); function shall not be available via citizen self-service",
        "Must",
        "SRO role only; §7.4 step 1",
    ],
    [
        "FR-HMA-095",
        "System shall allow Sub-Registrar to search registered Hindu marriage records by "
        "registration number, party name or Form II-A certificate number and verify court "
        "order particulars (court name, case number, order date, parties) against the "
        "register entry before endorsement",
        "Must",
        "Hard stop on mismatch; §7.4 steps 2–3",
    ],
    [
        "FR-HMA-096",
        "System shall require mandatory upload/scan of certified court order copy before "
        "Null and Void endorsement is saved; artefact retained per records policy (Rule 10)",
        "Must",
        "Endorsement blocked without court order artefact; §7.4 step 4",
    ],
    [
        "FR-HMA-097",
        "System shall allow Sub-Registrar to record nullity type as declared in court order "
        "— Void (Sec. 11) or Voidable / annulled (Sec. 12) — and flag the marriage "
        "registration as Null and Void",
        "Must",
        "Nullity type captured; §7.4 steps 5–6",
    ],
    [
        "FR-HMA-098",
        "On confirmation, system shall update register entry and certificate status: Form II-A "
        "marked invalid; Hindu Marriages Register shows NULL AND VOID endorsement; prior "
        "Registered / Certificate Issued status superseded",
        "Must",
        "Register and certificate reflect endorsement; §7.4 step 7",
    ],
    [
        "FR-HMA-099",
        "System shall prevent issuance of certified extracts for null/void-endorsed records "
        "unless the extract carries the null/void notation; audit trail shall record officer, "
        "timestamp, court order reference and prior/new status (Rule 8)",
        "Must",
        "Extracts show notation; audit log retained; §7.4 steps 8–9",
    ],
]

BR_HMA_021 = [
    "BR-HMA-021",
    "Null and Void endorsement requires a certified court order; Kaveri does not "
    "adjudicate nullity — SRO records the order and flags the registration only",
    "HMA 1955 Sec. 11–12; §7.4",
    "Hard stop without court order artefact; FR-HMA-094–099",
]

UI_ROW = [
    "Null and Void endorsement (SRO back-office)",
    "Search registered marriage, capture court order, apply Null and Void flag",
    "Offline",
    "HMA 1955 Sec. 11–12",
    "§7.4; FR-HMA-094–099",
]

RTM_SUMMARY_ROWS = [
    [
        "FR-HMA-094",
        FR_ROWS[1][1],
        "Must",
        FR_ROWS[1][3],
    ],
    [
        "FR-HMA-095",
        FR_ROWS[2][1],
        "Must",
        FR_ROWS[2][3],
    ],
    [
        "FR-HMA-096",
        FR_ROWS[3][1],
        "Must",
        FR_ROWS[3][3],
    ],
    [
        "FR-HMA-097",
        FR_ROWS[4][1],
        "Must",
        FR_ROWS[4][3],
    ],
    [
        "FR-HMA-098",
        FR_ROWS[5][1],
        "Must",
        FR_ROWS[5][3],
    ],
    [
        "FR-HMA-099",
        FR_ROWS[6][1],
        "Must",
        FR_ROWS[6][3],
    ],
]

RTM_FULL_ROWS = [
    [
        "FR-HMA-094",
        "HMA 1955 Sec. 11–12",
        "SRO-only Null and Void endorsement entry",
        "8.1.17",
        "Null and Void endorsement (SRO back-office)",
        "TC-HMA-___",
        "Draft",
    ],
    [
        "FR-HMA-095",
        "HMA 1955 Sec. 11–12",
        "Search register and verify court order particulars",
        "8.1.17",
        "Null and Void endorsement (SRO back-office)",
        "TC-HMA-___",
        "Draft",
    ],
    [
        "FR-HMA-096",
        "HMA 1955 Sec. 11–12; Rule 10",
        "Mandatory court order artefact upload",
        "8.1.17",
        "Null and Void endorsement (SRO back-office)",
        "TC-HMA-___",
        "Draft",
    ],
    [
        "FR-HMA-097",
        "HMA 1955 Sec. 11–12",
        "Record nullity type and flag Null and Void",
        "8.1.17",
        "Null and Void endorsement (SRO back-office)",
        "TC-HMA-___",
        "Draft",
    ],
    [
        "FR-HMA-098",
        "HMA 1955 Sec. 11–12; Rule 4(4)",
        "Update register entry and invalidate Form II-A",
        "8.1.17",
        "Null and Void endorsement (SRO back-office)",
        "TC-HMA-___",
        "Draft",
    ],
    [
        "FR-HMA-099",
        "Rule 8",
        "Extract notation and null/void audit trail",
        "8.1.17",
        "Null and Void endorsement (SRO back-office)",
        "TC-HMA-___",
        "Draft",
    ],
]

ENTITIES_OLD = (
    "Certificate (Form II-A), SakalaTransaction (GSC number, service/sub-service "
    "codes, processing status, transmission log), Special Marriage entities:"
)
ENTITIES_NEW = (
    "Certificate (Form II-A), NullityEndorsement (court order artefact, nullity type, "
    "endorsement date, prior status, endorsing officer), SakalaTransaction (GSC number, "
    "service/sub-service codes, processing status, transmission log), Special Marriage "
    "entities:"
)


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def add_table_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
    last: bool = False,
) -> Paragraph:
    matches: list[Paragraph] = []
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            matches.append(p)
        elif contains is not None and contains in t:
            matches.append(p)
    if not matches:
        raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")
    return matches[-1] if last else matches[0]


def section_exists(doc: Document, exact: str) -> bool:
    try:
        find_para(doc, exact=exact, heading_only=True)
        return True
    except KeyError:
        return False


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)


def find_table_containing(doc: Document, req_id: str) -> Table:
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == req_id:
                return table
    raise KeyError(f"Table not found containing {req_id!r}")


def find_table_by_header(doc: Document, first_cell: str, second_cell: str | None = None) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr0 = table.rows[0].cells[0].text.strip()
        hdr1 = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
        if hdr0 == first_cell and (second_cell is None or hdr1 == second_cell):
            return table
    raise KeyError(f"Table not found: {first_cell!r} / {second_cell!r}")


def replace_in_paragraphs(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    set_cell_text(cell, cell.text.replace(old, new))


def add_functional_requirements(doc: Document) -> None:
    if section_exists(doc, "8.1.17 Null and Void Endorsement (Court Order)"):
        return

    section82 = find_para(
        doc,
        exact="8.2 Special Marriage (Intended Marriage/Other Forms) Notice Generation",
        heading_only=True,
        last=True,
    )

    heading = deepcopy(section82._p)
    for child in list(heading):
        if child.tag != qn("w:pPr"):
            heading.remove(child)
    section82._p.addprevious(heading)
    h_para = Paragraph(heading, section82._parent)
    h_para.style = "Heading 4"
    set_para_text(h_para, "8.1.17 Null and Void Endorsement (Court Order)")

    cursor = insert_paragraph_after(h_para, "(Ref: §7.4)", style="Normal")
    insert_table_after(cursor, FR_ROWS, doc)


def update_contents(doc: Document) -> None:
    for i, p in enumerate(doc.paragraphs[:130]):
        if (
            p.text.strip() == "8.1.16 Digital signature and certificate issuance (FR-HMA-078–082)"
            and style_name(p) == "Normal"
        ):
            nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            if "8.1.17" in nxt:
                return
            insert_paragraph_after(
                p,
                "8.1.17 Null and Void Endorsement (Court Order) (FR-HMA-094–099)",
                style="Normal",
            )
            return
    raise KeyError("Contents entry for 8.1.16 not found")


def add_business_rule(doc: Document) -> None:
    br_table = find_table_containing(doc, "BR-HMA-020")
    if any(r.cells[0].text.strip() == "BR-HMA-021" for r in br_table.rows):
        return
    add_table_row(br_table, BR_HMA_021)


def add_ui_screen(doc: Document) -> None:
    ui_table = find_table_by_header(doc, "Screen / step", "Purpose")
    if any(
        "Null and Void endorsement" in r.cells[0].text for r in ui_table.rows
    ):
        return
    add_table_row(ui_table, UI_ROW)


def add_rtm_rows(doc: Document) -> None:
    rtm_summary = find_table_containing(doc, "FR-HMA-093")
    if any(r.cells[0].text.strip() == "FR-HMA-094" for r in rtm_summary.rows):
        return
    for row_data in RTM_SUMMARY_ROWS:
        add_table_row(rtm_summary, row_data)

    rtm_full = find_table_by_header(doc, "Req ID", "Act/Rule/Form")
    for row_data in RTM_FULL_ROWS:
        add_table_row(rtm_full, row_data)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.20")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    version_table = doc.tables[1]
    last_change = version_table.rows[-1].cells[3].text.strip()
    if "FR-HMA-094" not in last_change:
        add_version_row(
            version_table,
            [
                "1.20",
                "2026-09-02",
                "Nandha Kumar",
                "Add §8.1.17 Null/Void endorsement FRs (FR-HMA-094–099), BR-HMA-021, RTM",
                "Prashanth",
            ],
        )

    add_functional_requirements(doc)
    update_contents(doc)
    add_business_rule(doc)
    add_ui_screen(doc)
    add_rtm_rows(doc)
    replace_in_paragraphs(doc, ENTITIES_OLD, ENTITIES_NEW)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    find_para(doc2, exact="8.1.17 Null and Void Endorsement (Court Order)", heading_only=True)
    find_table_containing(doc2, "FR-HMA-094")
    find_table_containing(doc2, "BR-HMA-021")
    print("Verification OK")


if __name__ == "__main__":
    main()
