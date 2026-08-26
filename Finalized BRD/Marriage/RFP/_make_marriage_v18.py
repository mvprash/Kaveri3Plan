# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.8.docx from v1.7.

Add §15 Non-functional requirements (performance/scalability, security/privacy,
availability/error handling, VAPT) for the Marriage Registration module.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\Prashanth\Official\Kaveri 3.0\Kaveri3Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.7.docx"
DST = BASE / "BRD_Marriage_v1.8.docx"

TOC_ENTRIES = [
    "15. Non-functional requirements",
    "15.1 Performance and Scalability (NFR-MRG-PERF-001–002, NFR-MRG-SCALE-001)",
    "15.2 Security and Data Privacy (NFR-MRG-SEC-001, NFR-MRG-PRIV-001, NFR-MRG-AUD-001)",
    "15.3 System Availability and Error Handling (NFR-MRG-AVA-001, NFR-MRG-PAY-001)",
    "15.4 Security Audit and Compliance — VAPT Policy (NFR-MRG-VAPT-001–004)",
]

PERF_ROWS = [
    ["Req ID", "Requirement", "Priority", "Acceptance criteria"],
    [
        "NFR-MRG-PERF-001",
        "The system shall process standard user-interface interactions within 2 seconds "
        "and complete external API calls (including the Treasury payment gateway) within 5 seconds",
        "Must",
        "Measured at p95 under the concurrency load in NFR-MRG-PERF-002; UI ≤ 2 s; external API ≤ 5 s",
    ],
    [
        "NFR-MRG-PERF-002",
        "The portal shall support a minimum of 5,000 concurrent citizen sessions and "
        "2,500 concurrent internal-user (SRO / DEO) sessions without degradation in response time or error rate",
        "Must",
        "Load-test evidence: no SLA breach vs NFR-MRG-PERF-001 at the stated concurrency",
    ],
    [
        "NFR-MRG-SCALE-001",
        "Infrastructure shall automatically scale to handle a 300% surge above baseline "
        "traffic during culturally significant or auspicious marriage dates",
        "Must",
        "Autoscaling runbook and surge test: 3× baseline load served without manual scale-up",
    ],
]

SEC_ROWS = [
    ["Req ID", "Requirement", "Priority", "Acceptance criteria"],
    [
        "NFR-MRG-SEC-001",
        "All data in transit shall be secured using TLS 1.3, and all data at rest "
        "(including PII, documents and certificates) shall be encrypted using AES-256",
        "Must",
        "TLS 1.3 only on public and internal endpoints; AES-256 at rest confirmed in SDC / hosting design",
    ],
    [
        "NFR-MRG-PRIV-001",
        "Sensitive Personally Identifiable Information (PII), including Aadhaar numbers "
        "and biometric references, shall be strictly masked in system logs and in "
        "unauthorized database views",
        "Must",
        "Logs and support/DB views show masked Aadhaar / biometric refs; no raw values for unauthorized roles",
    ],
    [
        "NFR-MRG-AUD-001",
        "All critical state changes (application approvals, rejections, certificate "
        "issuances and DEO allocations) shall generate an immutable, timestamped audit "
        "log tied to the actor’s ID",
        "Must",
        "Append-only audit event with actor ID, timestamp, previous/next state and reason where applicable; "
        "aligns with FR-HMA-075 and FR-SMA-059",
    ],
]

AVA_ROWS = [
    ["Req ID", "Requirement", "Priority", "Acceptance criteria"],
    [
        "NFR-MRG-AVA-001",
        "If external integrations (eSign, DigiLocker or SMS / email gateways) experience "
        "downtime, the system shall display a clear, user-friendly error message and "
        "safely pause the workflow rather than failing the application completely",
        "Must",
        "Integration outage: citizen/officer sees bilingual reason; application state remains resumable; no data loss",
    ],
    [
        "NFR-MRG-PAY-001",
        "On a payment-gateway timeout, the system shall automatically poll the gateway "
        "for the final transaction status and lock the UI so the user cannot initiate a duplicate payment",
        "Must",
        "Timeout → poll until terminal status or defined retry window; pay action disabled until status resolved; "
        "no double-debit; aligns with FR-HMA-025 / FR-SMA-052",
    ],
]

VAPT_ROWS = [
    ["Req ID", "Requirement", "Priority", "Acceptance criteria"],
    [
        "NFR-MRG-VAPT-001",
        "As a strict infrastructure policy, the application shall undergo a comprehensive "
        "Vulnerability Assessment and Penetration Testing (VAPT) prior to production deployment",
        "Must",
        "VAPT report accepted by Security before Go-Live; no production release without a passed audit",
    ],
    [
        "NFR-MRG-VAPT-002",
        "VAPT scope shall be comprehensive: web application, APIs, integrated payment "
        "gateways and mobile-responsive interfaces (if any). Testing shall cover OWASP Top 10 "
        "vulnerabilities (https://owasp.org/), business-logic flaws and unauthorized data access",
        "Must",
        "Scope statement lists web, API, payment and responsive UI; OWASP Top 10, logic flaws and data-access tests evidenced",
    ],
    [
        "NFR-MRG-VAPT-003",
        "Any critical or high-severity vulnerability identified during the audit shall be "
        "remediated and verified through a re-audit before Go-Live. Tools-only automated "
        "scanning is insufficient; manual validation of automated findings is required",
        "Must",
        "Zero open Critical/High at Go-Live; retest report; manual validation notes against automated findings",
    ],
    [
        "NFR-MRG-VAPT-004",
        "After deployment, the application shall undergo mandatory VAPT at least once "
        "annually, or whenever significant architectural changes, feature additions or "
        "infrastructure migrations occur",
        "Must",
        "Annual VAPT scheduled; change-triggered VAPT in release checklist for architecture / feature / infra change",
    ],
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def clear_extra_rows(table: Table, keep_rows: int) -> None:
    tbl = table._tbl
    while len(table.rows) > keep_rows:
        tbl.remove(table.rows[-1]._tr)


def ensure_rows(table: Table, needed: int) -> None:
    while len(table.rows) < needed:
        last = table.rows[-1]._tr
        table._tbl.append(deepcopy(last))


def fill_table(table: Table, rows: list[list[str]]) -> None:
    ensure_rows(table, len(rows))
    for ri, vals in enumerate(rows):
        set_row(table, ri, vals)
    clear_extra_rows(table, len(rows))


def add_version_row(table: Table, values: list[str]) -> None:
    last = table.rows[-1]._tr
    table._tbl.append(deepcopy(last))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    p_pr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not p_pr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_table_by_cell(doc: Document, row: int, col: int, exact: str) -> Table:
    for table in doc.tables:
        if len(table.rows) <= row:
            continue
        cells = table.rows[row].cells
        if col >= len(cells):
            continue
        if cells[col].text.strip() == exact:
            return table
    raise KeyError(f"Table not found for [{row},{col}]={exact!r}")


def insert_table_after(paragraph: Paragraph, source: Table, rows: list[list[str]]) -> Paragraph:
    """Clone `source` after `paragraph`, fill `rows`, return a trailing empty Normal para."""
    tbl = deepcopy(source._tbl)
    paragraph._p.addnext(tbl)
    table = Table(tbl, paragraph._parent)
    fill_table(table, rows)
    trailing = insert_paragraph_after(paragraph, "", style="Normal")
    trailing._p.getparent().remove(trailing._p)
    tbl.addnext(trailing._p)
    return Paragraph(trailing._p, paragraph._parent)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.8")
    set_cell_text(doc.tables[0].rows[12].cells[1], "2026-08-26")
    add_version_row(
        doc.tables[1],
        [
            "1.8",
            "2026-08-26",
            "Nandha Kumar",
            "Added §15 Non-functional requirements — performance and scalability, "
            "security and data privacy, availability and error handling, and VAPT policy "
            "for the Marriage Registration module",
            "Prashanth",
        ],
    )

    # Executive summary — include NFR in the deliverable list.
    try:
        es = find_para(
            doc,
            contains="defines the functional , outlines the data requirements",
        )
        set_para_text(
            es,
            "Based on the analysis of the current state, the document proposes an enhanced "
            "future-state workflow aimed at improving accessibility, efficiency, transparency "
            "and service delivery. In addition, it defines the functional and non-functional "
            "requirements, outlines the data requirements required for effective implementation "
            "and management, and provides a reference to a detailed traceability matrix to "
            "ensure alignment between business objectives, requirements and solution deliverables.",
        )
    except KeyError:
        pass

    # TOC
    toc14 = None
    for p in doc.paragraphs:
        if p.text.strip() == "14. Acceptance and sign-off" and not style_name(p).startswith(
            "Heading"
        ):
            toc14 = p
            break
    if toc14 is None:
        raise KeyError("TOC entry for 14 not found")
    cursor = toc14
    for entry in TOC_ENTRIES:
        cursor = insert_paragraph_after(cursor, entry, style="Normal")

    template = find_table_by_cell(doc, 1, 0, "FR-SMA-001")

    uat = find_para(doc, contains="UAT scope: Test scenarios derived from FR-HMA-*")
    set_para_text(
        uat,
        "UAT scope: Test scenarios derived from FR-HMA-* and FR-SMA-* (see 13 RTM), "
        "NFR-MRG-* (see 15), BR-HMA-* / BR-SMA-*, statutory forms Form I / IA / II / II-A, "
        "and the Special Marriage Second, Third, Fourth and Fifth Schedules. Performance, "
        "security, availability and VAPT evidence in §15 are Go-Live gates, not optional extras.",
    )

    cursor = insert_paragraph_after(uat, "15. Non-functional requirements", style="Heading 2")
    cursor = insert_paragraph_after(
        cursor,
        "To ensure the Kaveri 3.0 Marriage Registration module operates reliably, securely "
        "and efficiently, the following non-functional parameters shall be integrated into "
        "the system architecture. These requirements apply to Hindu Marriage and Special "
        "Marriage (Intended Marriage / Other Forms) on both Online and Offline channels. "
        "Owners to validate: Solution Architect, DevOps / SDC, Security, DBA, Ops and Product Owner.",
        style="Normal",
    )

    cursor = insert_paragraph_after(
        cursor, "15.1 Performance and Scalability", style="Heading 3"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Response time, concurrency and elasticity targets below are mandatory load-test "
        "gates before production. Surge capacity is required for culturally significant / "
        "auspicious marriage dates.",
        style="Normal",
    )
    cursor = insert_table_after(cursor, template, PERF_ROWS)

    cursor = insert_paragraph_after(
        cursor, "15.2 Security and Data Privacy", style="Heading 3"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Encryption, PII protection and auditability are baseline controls. Aadhaar / e-KYC "
        "usage remains subject to department and UIDAI approval (see also 2.4 Constraints).",
        style="Normal",
    )
    cursor = insert_table_after(cursor, template, SEC_ROWS)

    cursor = insert_paragraph_after(
        cursor, "15.3 System Availability and Error Handling", style="Heading 3"
    )
    cursor = insert_paragraph_after(
        cursor,
        "The module shall fail safely when an external dependency is unavailable, and shall "
        "not allow duplicate fee collection when a payment response is delayed.",
        style="Normal",
    )
    cursor = insert_table_after(cursor, template, AVA_ROWS)

    cursor = insert_paragraph_after(
        cursor,
        "15.4 Security Audit and Compliance (VAPT Policy)",
        style="Heading 3",
    )
    cursor = insert_paragraph_after(
        cursor,
        "Vulnerability Assessment and Penetration Testing is a mandatory infrastructure "
        "policy for this module. Automated scanning alone does not satisfy the gate.",
        style="Normal",
    )
    cursor = insert_table_after(cursor, template, VAPT_ROWS)

    # RTM sample rows
    rtm = find_table_by_cell(doc, 0, 1, "Act/Rule/Form")
    rtm_extra = [
        [
            "NFR-MRG-PERF-001",
            "Architecture / SLA",
            "UI ≤ 2 s; external API (Treasury) ≤ 5 s",
            "15.1",
            "Citizen / officer UI",
            "TC-NFR-___",
            "Draft",
        ],
        [
            "NFR-MRG-PERF-002",
            "Architecture / SLA",
            "5,000 citizen + 2,500 SRO/DEO concurrent sessions without degradation",
            "15.1",
            "Load test",
            "TC-NFR-___",
            "Draft",
        ],
        [
            "NFR-MRG-VAPT-001",
            "Security policy",
            "Comprehensive VAPT before production deployment",
            "15.4",
            "—",
            "TC-NFR-___",
            "Draft",
        ],
    ]
    for vals in rtm_extra:
        last = rtm.rows[-1]._tr
        rtm._tbl.append(deepcopy(last))
        set_row(rtm, len(rtm.rows) - 1, vals)

    # Appendix — OWASP
    last_ref = None
    for p in doc.paragraphs:
        if style_name(p) in ("List Bullet", "List Paragraph") and p.text.strip().startswith(
            "Approved process diagram — Special Marriage Other Forms"
        ):
            last_ref = p
    if last_ref is None:
        last_ref = find_para(doc, contains="Approved process diagram — Special Marriage Other Forms")
    insert_paragraph_after(
        last_ref,
        "OWASP Top 10 — https://owasp.org/ (VAPT scope, NFR-MRG-VAPT-002)",
        style="List Bullet",
    )

    doc.save(str(DST))
    print(f"Wrote {DST}")
    doc2 = Document(str(DST))
    print("--- §15 headings ---")
    for p in doc2.paragraphs:
        if p.text.strip().startswith("15"):
            print(f"  [{style_name(p)}] {p.text.strip()}")


if __name__ == "__main__":
    main()
