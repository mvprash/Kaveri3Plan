# -*- coding: utf-8 -*-
"""Create BRD_Marriage_BRD_v1.24.docx from v1.23.

User updated Form bullets/structure under statutory forms. This version:
- Restores numbered headings for §§2–7 so 3.x / 7.x references resolve
- Updates In-scope bullets with 3.x / 7.x cross-refs
- Restores §3 Implemented-in intro; normalizes §3 implementation column headers
- Updates Contents for new Form sub-headings; restores version history
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_BRD_v1.23.docx"
HISTORY_SRC = BASE / "BRD_Marriage_BRD_v1.22.docx"
DST = BASE / "BRD_Marriage_BRD_v1.24.docx"

IMPL_INTRO = (
    "The Refer section 7 for Implementation column maps each Act section, Rule or "
    "notification to the Future state (To-Be) process sub-section (§7) where Kaveri 3.0 "
    "implements the requirement. Items with no citizen-facing To-Be step are marked —. "
    "Companion FRS/NFR detail is in FRS_and_NFRS_Marriage_v1.22.docx."
)

IMPL_HEADER = "Refer section 7 for Implementation"

# Exact heading text (as in v1.23) → numbered heading text
HEADING_MAP = {
    "Scope": "2. Scope",
    "In scope (Hindu Marriage & Special Marriage)": (
        "2.1 In scope (Hindu Marriage & Special Marriage)"
    ),
    "Legal and regulatory reference": "3. Legal and regulatory reference",
    "Applicable Acts": "3.1 Applicable Acts",
    "Relevant sections followed by the Department for Marriage Registration": (
        "3.2 Relevant sections followed by the Department for Marriage Registration"
    ),
    "Hindu Marriage Act, 1955 (selected sections for registration)": (
        "3.2.1 Hindu Marriage Act, 1955 (selected sections for registration)"
    ),
    "Special Marriage Act, 1954 (selected sections)": (
        "3.2.2 Special Marriage Act, 1954 (selected sections)"
    ),
    "Relevant rules followed by the Department for Marriage Registration": (
        "3.3 Relevant rules followed by the Department for Marriage Registration"
    ),
    "Registration of Hindu Marriage (Karnataka) Rules, 1966": (
        "3.3.1 Registration of Hindu Marriage (Karnataka) Rules, 1966"
    ),
    "Special Marriage (Karnataka) Rules, 1961": (
        "3.3.2 Special Marriage (Karnataka) Rules, 1961"
    ),
    "Relevant notifications issued by the Department for Marriage Registration": (
        "3.4 Relevant notifications issued by the Department for Marriage Registration"
    ),
    "Hindu Marriage statutory forms mapping": "3.5 Hindu Marriage statutory forms mapping",
    "Form 1": "3.5.1 Form I",
    "Form IA": "3.5.2 Form IA",
    "Form II": "3.5.3 Form II",
    "Form II - A": "3.5.4 Form II-A",
    "Statutory penalty highlights — Sec. 17 & Sec. 18 (displayed in red at submission)": (
        "3.5.5 Statutory penalty highlights — Sec. 17 & Sec. 18 "
        "(displayed in red at submission)"
    ),
    "Special Marriage statutory forms mapping": (
        "3.6 Special Marriage statutory forms mapping"
    ),
    "Form - Notice of Intended Marriage": "3.6.1 Form — Notice of Intended Marriage",
    "Form - Declaration": "3.6.2 Form — Declaration",
    "Form – Certificate – Special Marriage": (
        "3.6.3 Form — Certificate — Special Marriage"
    ),
    "Form – Certificate – Special Marriage (Other Forms)": (
        "3.6.4 Form — Certificate — Special Marriage (Other Forms)"
    ),
    "Sakala — Karnataka Guarantee of Services": (
        "3.7 Sakala — Karnataka Guarantee of Services"
    ),
    "Stakeholders and actors": "4. Stakeholders and actors",
    "Definitions and glossary": "5. Definitions and glossary",
    "Current state": "6. Current state",
    "As-Is pain points": "6.1 As-Is pain points",
    "Future state (To-Be)": "7. Future state (To-Be)",
    "Hindu Marriage": "7.1 Hindu Marriage",
    "Special Marriage Notice Generation (Intended Marriage / Other Forms)": (
        "7.2 Special Marriage Notice Generation (Intended Marriage / Other Forms)"
    ),
    "Special Marriage Marriage Registration (Intended Marriage / Other Forms)": (
        "7.3 Special Marriage Marriage Registration (Intended Marriage / Other Forms)"
    ),
    "Marriage — Null and Void Endorsement (Court Order)": (
        "7.4 Marriage — Null and Void Endorsement (Court Order)"
    ),
    "What is new in Kaveri 3.0": "7.5 What is new in Kaveri 3.0",
    "Rectified As-Is pain points": "7.5.1 Rectified As-Is pain points",
}

# Context-sensitive renames for repeated titles (Channel models, Online, …)
# Applied in document order via a simple state machine.
CONTEXT_HEADINGS = [
    # after 7.1
    ("7.1", "Channel models", "7.1.1 Channel models"),
    ("7.1", "Process Diagram", "7.1.2 Process Diagram"),
    ("7.1", "Common intake steps", "7.1.2.1 Common intake steps"),
    ("7.1", "Online", "7.1.2.2 Online"),
    ("7.1", "Offline (In Person)", "7.1.2.3 Offline (In Person)"),
    ("7.1", "Application Status Model", "7.1.2.4 Application Status Model"),
    # after 7.2
    ("7.2", "Channel models", "7.2.1 Channel models"),
    ("7.2", "Process Diagram", "7.2.2 Process Diagram"),
    ("7.2", "Common intake steps", "7.2.2.1 Common intake steps"),
    ("7.2", "Online", "7.2.2.2 Online"),
    ("7.2", "Offline (In Person)", "7.2.2.3 Offline (In Person)"),
    ("7.2", "Application Status Model", "7.2.2.4 Application Status Model"),
    # after 7.3
    ("7.3", "Channel models", "7.3.1 Channel models"),
    ("7.3", "Process Diagram", "7.3.2 Process Diagram"),
    ("7.3", "Common intake steps", "7.3.2.1 Common intake steps"),
    ("7.3", "Offline (In Person)", "7.3.2.2 Offline (In Person)"),
    ("7.3", "Application Status Model", "7.3.2.3 Application Status Model"),
    # after 7.4
    ("7.4", "Process steps", "7.4.1 Process steps"),
]

BULLET_UPDATES = [
    (
        "Statutory artefacts: Form I (Memorandum), Form IA (Application), Form II "
        "(Endorsement — Rule 4(4)), Form II-A (Certificate)",
        "Statutory artefacts: Form I (Memorandum), Form IA (Application), Form II "
        "(Endorsement — Rule 4(4)), Form II-A (Certificate) — see §3.5.",
    ),
    (
        "References/Endorsements to flag the marriage as NULL and VOID",
        "References/Endorsements to flag the marriage as NULL and VOID — see §7.4.",
    ),
    (
        "Integrations: [payment, Aadhaar/eKYC, DigiLocker, SMS, e-Mail, Kutumba portal, "
        "Civil Registration System, Labor Department, Sakala (Karnataka Guarantee of "
        "Services)].",
        "Integrations: [payment, Aadhaar/eKYC, DigiLocker, SMS, e-Mail, Kutumba portal, "
        "Civil Registration System, Labor Department, Sakala (Karnataka Guarantee of "
        "Services)] — Sakala regulatory reference §3.7.",
    ),
    (
        "Special Marriage — Intended Marriage (Chapter II, SMA 1954): notice of intended "
        "marriage (Second Schedule) Online & Offline, 30-day publication / objection "
        "period, solemnization and certificate (Fourth Schedule) within statutory "
        "validity (diagram gate ≥30 and ≤90 days; aligns to Sec. 7 & Sec. 14).",
        "Special Marriage — Intended Marriage (Chapter II, SMA 1954): notice of intended "
        "marriage (Second Schedule) Online & Offline, 30-day publication / objection "
        "period, solemnization and certificate (Fourth Schedule) within statutory "
        "validity (diagram gate ≥30 and ≤90 days; aligns to Sec. 7 & Sec. 14) — forms "
        "§3.6; process §7.2–§7.3.",
    ),
    (
        "Special Marriage — Other Forms (Chapter III, SMA 1954 Sec. 15–16): "
        "notice/publication for registration of marriages celebrated in other forms, "
        "objection handling, and entry of certificate in Marriage Certificate Book "
        "(Fifth Schedule).",
        "Special Marriage — Other Forms (Chapter III, SMA 1954 Sec. 15–16): "
        "notice/publication for registration of marriages celebrated in other forms, "
        "objection handling, and entry of certificate in Marriage Certificate Book "
        "(Fifth Schedule) — forms §3.6; process §7.2–§7.3.",
    ),
]

SAKALA_SOURCE = (
    "Source: §3.1 (Karnataka Act 1 of 2012); §3.4 (Karnataka Guarantee of Services to "
    "Citizens Rules, 2012); Sakala Mission portal — https://sakala.kar.nic.in/. Marriage "
    "registration services offered through the Department"
)

CONTENTS_FORM_ENTRIES = [
    ("3.5 Hindu Marriage statutory forms mapping", [
        "3.5.1 Form I",
        "3.5.2 Form IA",
        "3.5.3 Form II",
        "3.5.4 Form II-A",
        "3.5.5 Statutory penalty highlights — Sec. 17 & Sec. 18",
    ]),
    ("3.6 Special Marriage statutory forms mapping", [
        "3.6.1 Form — Notice of Intended Marriage",
        "3.6.2 Form — Declaration",
        "3.6.3 Form — Certificate — Special Marriage",
        "3.6.4 Form — Certificate — Special Marriage (Other Forms)",
    ]),
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def restore_headings(doc: Document) -> int:
    """Restore numbered headings. Returns count renamed."""
    count = 0
    section7_ctx = ""  # '7.1' | '7.2' | '7.3' | '7.4'
    for p in doc.paragraphs:
        if not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if not t:
            continue

        # Unique map first
        if t in HEADING_MAP:
            set_para_text(p, HEADING_MAP[t])
            count += 1
            nt = HEADING_MAP[t]
            if nt.startswith("7.1 "):
                section7_ctx = "7.1"
            elif nt.startswith("7.2 "):
                section7_ctx = "7.2"
            elif nt.startswith("7.3 "):
                section7_ctx = "7.3"
            elif nt.startswith("7.4 "):
                section7_ctx = "7.4"
            continue

        # Already numbered — track context
        if t.startswith("7.1 "):
            section7_ctx = "7.1"
        elif t.startswith("7.2 "):
            section7_ctx = "7.2"
        elif t.startswith("7.3 "):
            section7_ctx = "7.3"
        elif t.startswith("7.4 "):
            section7_ctx = "7.4"

        # Context-sensitive repeated titles
        for ctx, exact, numbered in CONTEXT_HEADINGS:
            if section7_ctx == ctx and t == exact:
                set_para_text(p, numbered)
                count += 1
                break
    return count


def update_bullets(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        for old, new in BULLET_UPDATES:
            if t == old:
                set_para_text(p, new)
                n += 1
                break
    return n


def restore_impl_intro(doc: Document) -> None:
    legal = None
    for p in doc.paragraphs:
        if style_name(p).startswith("Heading") and "Legal and regulatory reference" in p.text:
            legal = p
            break
    if legal is None:
        raise KeyError("§3 Legal heading not found")
    # Check if intro already present
    nxt = legal._p.getnext()
    if nxt is not None and nxt.tag.endswith("}p"):
        texts = "".join(t.text or "" for t in nxt.iter(qn("w:t")))
        if "maps each Act section" in texts or "Refer section 7" in texts:
            # update text
            for p in doc.paragraphs:
                if "maps each Act section" in p.text or (
                    "Implemented in" in p.text and "maps each" in p.text
                ):
                    set_para_text(p, IMPL_INTRO)
                    return
            return
    insert_paragraph_after(legal, IMPL_INTRO, style="Normal")


def normalize_impl_headers(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        if not table.rows:
            continue
        for ci, cell in enumerate(table.rows[0].cells):
            h = cell.text.strip().replace("\n", " ")
            if h in (
                "Implemented in (7)",
                "Implemented in (§7)",
                "Refer section 7 for Implementation",
            ) or ("Implement" in h and "7" in h) or h.startswith("Refer section 7"):
                if h != IMPL_HEADER:
                    set_cell_text(cell, IMPL_HEADER)
                    n += 1
                else:
                    n += 1  # already ok
    return n


def update_sakala_source(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith("Source:") and "Sakala Mission" in p.text:
            # Keep trailing sentence after portal URL if present in original longer form
            # Replace with corrected cross-refs; preserve rest of paragraph if longer
            rest = ""
            marker = "Marriage registration services offered through the Department"
            if marker in p.text:
                idx = p.text.find(marker)
                rest = p.text[idx + len(marker) :]
            set_para_text(p, SAKALA_SOURCE + rest)
            return


def update_contents(doc: Document) -> None:
    """Insert Form sub-entries under 3.5 / 3.6 in Contents; drop orphan 2.2–2.4 if present."""
    # Remove orphan TOC lines 2.2–2.4 (no body sections)
    to_remove = []
    in_contents = False
    for p in doc.paragraphs:
        t = p.text.strip()
        st = style_name(p)
        if st.startswith("Heading") and t == "Contents":
            in_contents = True
            continue
        if in_contents and st.startswith("Heading"):
            break
        if not in_contents:
            continue
        if t.startswith("2.2 ") or t.startswith("2.3 ") or t.startswith("2.4 "):
            to_remove.append(p._p)
        # Refresh 3.5 / 3.6 parent labels if unnumbered leftovers
        if t == "3.5 Hindu Marriage statutory forms mapping" or t.startswith("3.5 "):
            pass
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # Insert form sub-entries after 3.5 / 3.6 TOC lines if missing
    for parent_label, children in CONTENTS_FORM_ENTRIES:
        parent_para = None
        in_contents = False
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            st = style_name(p)
            if st.startswith("Heading") and t == "Contents":
                in_contents = True
                continue
            if in_contents and st.startswith("Heading"):
                break
            if not in_contents:
                continue
            if t.startswith(parent_label.split(" ", 1)[0] + " ") and parent_label.split(" ", 1)[1][:20] in t:
                parent_para = p
                # check if children already follow
                nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
                if any(nxt.startswith(c.split(" ")[0]) for c in children):
                    parent_para = None  # already present
                break
        if parent_para is None:
            continue
        anchor = parent_para
        for child in children:
            anchor = insert_paragraph_after(anchor, child, style="Normal")


def restore_version_history(doc: Document, history_doc: Document) -> None:
    src_vt = history_doc.tables[1]
    dst_vt = doc.tables[1]
    # Clear existing data rows (keep header)
    while len(dst_vt.rows) > 1:
        dst_vt._tbl.remove(dst_vt.rows[-1]._tr)
    # Copy all history rows from v1.22
    for row in src_vt.rows[1:]:
        dst_vt._tbl.append(deepcopy(row._tr))
    # Append 1.23 / 1.24 notes
    for values in [
        [
            "1.23",
            "2026-09-03",
            "Nandha Kumar",
            "Restructure §3.5/§3.6 Form headings; complete Sakala actor in §4; "
            "align §3 implementation column headers",
            "Prashanth",
        ],
        [
            "1.24",
            "2026-09-03",
            "Nandha Kumar",
            "Restore numbered §3.x / §7.x headings; update In-scope bullets with "
            "§3.5–§3.7 / §7.2–§7.4 refs; refresh Contents Form entries and version history",
            "Prashanth",
        ],
    ]:
        dst_vt._tbl.append(deepcopy(dst_vt.rows[-1]._tr))
        row = dst_vt.rows[-1]
        for ci, val in enumerate(values):
            set_cell_text(row.cells[ci], val)


def update_doc_control(doc: Document) -> None:
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.24")
    set_cell_text(
        doc.tables[0].rows[3].cells[1],
        "Draft / In review — BRD part (Executive summary through §7.5)",
    )
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-03")


def update_75_intros(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.startswith("This section summarises material enhancements"):
            set_para_text(
                p,
                "This section summarises material enhancements in Kaveri 3.0 compared with "
                "the legacy Kaveri 2.0 Marriage Registration module (§6). Capability "
                "highlights are listed below; §7.5.1 maps each As-Is pain point from §6.1 "
                "to the Kaveri 3.0 closure. Cross-references point to To-Be process in this "
                "BRD (§7.1–§7.4) and to functional / non-functional requirements in "
                "FRS_and_NFRS_Marriage_v1.22.docx.",
            )
        if p.text.startswith("The following As-Is pain points from"):
            set_para_text(
                p,
                "The following As-Is pain points from §6.1 (Kaveri 2.0 workshops, "
                "ServiceDesk tickets and department discussions) are closed in Kaveri 3.0. "
                "Cross-references point first to the To-Be process in this BRD (§7.1–§7.5), "
                "then to functional requirements, fallbacks or NFRs in "
                "FRS_and_NFRS_Marriage_v1.22.docx that implement the fix.",
            )


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))
    history = Document(str(HISTORY_SRC)) if HISTORY_SRC.exists() else doc

    n_head = restore_headings(doc)
    n_bullets = update_bullets(doc)
    restore_impl_intro(doc)
    n_hdr = normalize_impl_headers(doc)
    update_sakala_source(doc)
    update_contents(doc)
    restore_version_history(doc, history)
    update_doc_control(doc)
    update_75_intros(doc)

    doc.save(str(DST))
    print(f"Wrote {DST}")
    print(f"  headings restored: {n_head}; bullets updated: {n_bullets}; impl headers: {n_hdr}")

    # Verify
    doc2 = Document(str(DST))
    heads = [
        p.text.strip()
        for p in doc2.paragraphs
        if style_name(p).startswith("Heading") and p.text.strip()
    ]
    assert any(h.startswith("3.5.1") for h in heads), heads
    assert any(h.startswith("7.1.2.2") for h in heads), "7.1.2.2 missing"
    assert any(h.startswith("7.4 ") for h in heads), "7.4 missing"
    assert any("see §3.5" in p.text for p in doc2.paragraphs), "bullet §3.5 ref missing"
    assert any("see §7.4" in p.text for p in doc2.paragraphs), "bullet §7.4 ref missing"
    assert doc2.tables[0].rows[2].cells[1].text.strip() == "1.24"
    assert len(doc2.tables[1].rows) > 10, "version history not restored"
    print("Verification OK")


if __name__ == "__main__":
    main()
