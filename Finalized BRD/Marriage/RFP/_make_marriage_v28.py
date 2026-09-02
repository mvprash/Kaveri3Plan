# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.18.docx from v1.17.

Add Sakala Act and Rules to §3; fix §3.7 placement (after §3.6, before §4).
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.17.docx"
DST = BASE / "BRD_Marriage_v1.18.docx"

SAKALA_ACT_ROW = [
    "The Karnataka Guarantee of Services to Citizens Act, 2011",
    "Karnataka Act 1 of 2012",
    "Time-bound guarantee of citizen-related services; GSC acknowledgement; online status "
    "monitoring; appeals and compensatory cost (Secs. 3–13); amended by Karnataka Act 36 of "
    "2012 and Karnataka Act 31 of 2014",
]

SAKALA_SECTION_ROWS = [
    ["Section", "Topic", "BRD relevance"],
    [
        "Sec. 2",
        "Definitions — citizen related service, designated officer, competent officer, "
        "appellate authority, stipulated time, public authority",
        "Glossary; Sakala service / office mapping; officer roles in pendency and appeals",
    ],
    [
        "Sec. 3",
        "Right to obtain service within stipulated time",
        "Statutory basis for notified marriage-service timelines tracked from payment / "
        "acknowledgement",
    ],
    [
        "Sec. 4",
        "Notification of services, designated officers and stipulated time limits (Schedule)",
        "Marriage services listed in Sakala Schedule; service codes master in Kaveri",
    ],
    [
        "Sec. 5",
        "Providing services within stipulated time; acknowledgement; rejection with written "
        "reasons and appeal information",
        "GSC on acceptance; rejection reason sync to Sakala; FR-HMA-093 / FR-SMA-069",
    ],
    [
        "Sec. 6",
        "Monitoring application status online; application number (GSC)",
        "Bidirectional status sync with https://sakala.kar.nic.in/; FR-HMA-092–093",
    ],
    [
        "Sec. 7",
        "E-governance of services through mutual understanding",
        "Legal basis for Kaveri ↔ Sakala electronic integration",
    ],
    [
        "Sec. 8–9",
        "Compensatory cost to citizen; liability of defaulting public servant",
        "MIS on overdue marriage applications; appeal trigger when timeline breached",
    ],
    [
        "Sec. 13",
        "Appeal by aggrieved citizen (Appeal-I to competent officer; Appeal-II to appellate "
        "authority)",
        "Rejection / delay appeal workflow; Sakala portal Appeal-I / Appeal-II",
    ],
    [
        "Schedule",
        "Notified services with designated officer, competent officer, appellate authority "
        "and stipulated time",
        "Hindu Marriage registration; Special Marriage notice and registration services — "
        "confirm codes with Sakala Mission",
    ],
]

SAKALA_RULE_ROWS = [
    ["Rule", "Requirement", "System feature"],
    [
        "Rule 4",
        "Manner of receiving application and issuing acknowledgement to applicants",
        "GSC / acknowledgement slip on payment; receipt format; FR-HMA-092 / FR-SMA-068",
    ],
    [
        "Rule 5",
        "Public holidays shall not be counted in stipulated time",
        "Working-day calculator for Sakala timeline display and overdue flags",
    ],
    [
        "Rule 6",
        "Manner of seeking payment of compensatory cost",
        "Appeal outcome linkage when service delayed beyond stipulated time",
    ],
    [
        "Rule 11",
        "Procedure for decision on First or Second Appeal",
        "Appeal status sync; officer workflow alignment with Sakala appeal stages",
    ],
    [
        "Rule 16",
        "Maintenance of records of all disposed cases under the Act",
        "Audit trail; Sakala transmission log retention",
    ],
    [
        "Rule 18",
        "Monitoring of implementation",
        "Department / Sakala MIS; pendency reconciliation reports",
    ],
]

SAKALA_NOTIFICATION_ROW = [
    "Karnataka Guarantee of Services to Citizens Rules, 2012",
    "Gazette 29-Feb-2012 (published 14-Feb-2012 draft)",
    "Subordinate rules under Sec. 19 of the Sakala Act — acknowledgement, appeals, "
    "compensatory cost and monitoring procedures",
    "GSC issuance, appeal workflow, timeline calculation — §3.3.3; FR-HMA-092–093",
    "Karnataka Guarantee of Services to Citizens Rules, 2012 (DPAR / Sakala Mission)",
]

SAKALA_37_SERVICES = [
    "Hindu Marriage registration (online filing) — GSC issued on successful fee payment; "
    "statutory service timeline commences from payment date (typically 1 working day for "
    "certificate delivery after payment, per department practice).",
    "Special Marriage — Intended Marriage notice and registration.",
    "Special Marriage — Other Forms notice and registration.",
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para_index(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
    last: bool = False,
) -> int:
    matches: list[int] = []
    for i, p in enumerate(doc.paragraphs):
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            matches.append(i)
        elif contains is not None and contains in t:
            matches.append(i)
    if not matches:
        raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")
    return matches[-1] if last else matches[0]


def find_para(doc: Document, **kwargs) -> Paragraph:
    return doc.paragraphs[find_para_index(doc, **kwargs)]


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def find_table_by_header(doc: Document, first_cell: str, second_cell: str | None = None) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr0 = table.rows[0].cells[0].text.strip()
        hdr1 = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
        if hdr0 == first_cell and (second_cell is None or hdr1 == second_cell):
            return table
    raise KeyError(f"Table not found: {first_cell!r} / {second_cell!r}")


def remove_misplaced_sakala_37(doc: Document) -> None:
    """Remove §3.7 block wrongly inserted after §4 in v1.17."""
    to_delete: list[Paragraph] = []
    capture = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "3.7 Sakala — Karnataka Guarantee of Services" and style_name(p) == "Heading 3":
            prev = p._p.getprevious()
            while prev is not None and prev.tag != qn("w:p"):
                prev = prev.getprevious()
            if prev is not None:
                prev_para = Paragraph(prev, p._parent)
                if prev_para.text.strip().startswith("4. Stakeholders"):
                    capture = True
        if capture:
            to_delete.append(p)
            if t.startswith("5. Definitions"):
                to_delete.pop()  # keep §5 heading
                break
    for p in to_delete:
        delete_paragraph(p)


def add_sakala_act_to_applicable_acts(doc: Document) -> None:
    acts = find_table_by_header(doc, "Act", "Central Act No.")
    acts._tbl.append(deepcopy(acts.rows[-1]._tr))
    set_row(acts, len(acts.rows) - 1, SAKALA_ACT_ROW)


def add_sakala_sections(doc: Document) -> None:
    anchor = find_para(
        doc,
        exact="Source: Acts_Rules/Marriage/The Special Marriage Act, 1954.pdf",
        last=True,
    )
    cursor = insert_paragraph_after(
        anchor,
        "3.2.3 Karnataka Guarantee of Services to Citizens Act, 2011 (Sakala)",
        style="Heading 4",
    )
    cursor = insert_paragraph_after(
        cursor,
        "Source: Karnataka Act 1 of 2012 (Karnataka Guarantee of Services to Citizens Act, "
        "2011); amendments — Karnataka Act 36 of 2012, Karnataka Act 31 of 2014 "
        "(Karnataka Sakala Services (Amendment) Act, 2014 — electronic service delivery). "
        "Portal: https://sakala.kar.nic.in/",
        style="Normal",
    )
    insert_table_after(cursor, SAKALA_SECTION_ROWS, doc)


def add_sakala_rules(doc: Document) -> None:
    anchor = find_para(
        doc,
        contains="Source: Acts_Rules/Marriage/SpecialMarriage(Karnataka)Rules1961.pdf",
        last=True,
    )
    cursor = insert_paragraph_after(
        anchor,
        "3.3.3 Karnataka Guarantee of Services to Citizens Rules, 2012",
        style="Heading 4",
    )
    cursor = insert_paragraph_after(
        cursor,
        "Source: Karnataka Guarantee of Services to Citizens Rules, 2012 (made under "
        "Sec. 19 of the Sakala Act; published in Karnataka Gazette, effective from date "
        "of publication). Administered by DPAR (e-Governance) / Sakala Mission.",
        style="Normal",
    )
    insert_table_after(cursor, SAKALA_RULE_ROWS, doc)


def add_sakala_37_section(doc: Document) -> None:
    """Insert §3.7 after §3.6 content, immediately before §4 Stakeholders."""
    stakeholders = find_para(doc, exact="4. Stakeholders and actors", last=True)
    prev = stakeholders._p.getprevious()
    anchor = stakeholders
    while prev is not None and prev.tag == qn("w:p"):
        para = Paragraph(prev, stakeholders._parent)
        if para.text.strip():
            anchor = para
            break
        prev = prev.getprevious()

    cursor = insert_paragraph_after(
        anchor, "3.7 Sakala — Karnataka Guarantee of Services", style="Heading 3"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Source: §3.1 (Karnataka Act 1 of 2012); §3.2.3 (selected sections); §3.3.3 "
        "(Rules, 2012); Sakala Mission portal — https://sakala.kar.nic.in/. Marriage "
        "registration services offered through the Department of Stamps and Registration "
        "are notified Sakala services. Every eligible application shall receive a "
        "Guarantee of Services to Citizen (GSC) number so citizens can track status, "
        "receive time-bound service and file appeals (Appeal-I / Appeal-II) when service "
        "is delayed or rejected.",
        style="Normal",
    )
    cursor = insert_paragraph_after(
        cursor,
        "Notified marriage-related Sakala services (department baseline — confirm codes "
        "with Sakala Mission during integration design):",
        style="Normal",
    )
    for item in SAKALA_37_SERVICES:
        cursor = insert_paragraph_after(cursor, item, style="List Bullet")
    insert_paragraph_after(
        cursor,
        "Kaveri 3.0 shall integrate with the Sakala platform per NIC / Sakala Mission "
        "approved web methods for GSC registration, status updates (acceptance, "
        "in-process, delivered, rejected), party-detail upload and orphan-settle retry "
        "— see §8.7 and §11.",
        style="Normal",
    )


def update_contents(doc: Document) -> None:
    # Only update contents block (first ~120 paragraphs)
    for p in doc.paragraphs[:120]:
        if p.text.strip() == "3.2.2 Special Marriage Act, 1954 (selected sections)":
            insert_paragraph_after(
                p,
                "3.2.3 Karnataka Guarantee of Services to Citizens Act, 2011 (Sakala)",
                style="Normal",
            )
            break
    for p in doc.paragraphs[:120]:
        if p.text.strip() == "3.3.2 Special Marriage (Karnataka) Rules, 1961":
            insert_paragraph_after(
                p,
                "3.3.3 Karnataka Guarantee of Services to Citizens Rules, 2012",
                style="Normal",
            )
            break


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.18")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    add_version_row(
        doc.tables[1],
        [
            "1.18",
            "2026-09-02",
            "Nandha Kumar",
            "§3 Sakala Act (Karnataka Act 1 of 2012) + Rules 2012 in §3.1/3.2.3/3.3.3; "
            "fix §3.7 placement before §4",
            "Prashanth",
        ],
    )

    remove_misplaced_sakala_37(doc)
    add_sakala_act_to_applicable_acts(doc)
    add_sakala_sections(doc)
    add_sakala_rules(doc)
    add_sakala_37_section(doc)
    update_contents(doc)

    notifications = find_table_by_header(doc, "Instrument", "Date / No.")
    notifications._tbl.append(deepcopy(notifications.rows[-1]._tr))
    set_row(notifications, len(notifications.rows) - 1, SAKALA_NOTIFICATION_ROW)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    acts = find_table_by_header(doc2, "Act", "Central Act No.")
    assert any("Sakala" in r.cells[0].text or "Guarantee" in r.cells[0].text for r in acts.rows)
    find_para(doc2, contains="3.2.3 Karnataka Guarantee of Services", last=True)
    find_para(doc2, contains="3.3.3 Karnataka Guarantee of Services", last=True)
    # §3.7 must precede §4 in body
    idx_37 = find_para_index(
        doc2, exact="3.7 Sakala — Karnataka Guarantee of Services", last=True
    )
    idx_4 = find_para_index(doc2, exact="4. Stakeholders and actors", last=True)
    assert idx_37 < idx_4, f"§3.7 at {idx_37} should be before §4 at {idx_4}"
    find_table_by_header(doc2, "Section", "Topic")
    print("Verification OK")


if __name__ == "__main__":
    main()
