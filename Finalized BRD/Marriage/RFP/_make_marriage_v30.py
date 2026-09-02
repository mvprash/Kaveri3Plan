# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.19.docx from v1.18.

Add §7.4 Hindu Marriage — Null and Void Endorsement (Court Order) workflow
(HMA 1955 Sec. 11–12): SRO receives court order and flags marriage as Null and Void.
Steps only — no process diagram. Update §3.2.1 Implemented in (§7) mapping.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.18.docx"
DST = BASE / "BRD_Marriage_v1.19.docx"

NULL_VOID_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    [
        "1",
        "Sub-Registrar receives certified copy of court order declaring the marriage "
        "void or annulled (nullity decree)",
        "Sub Registrar",
        "HMA 1955 Sec. 11–12; physical or scanned copy presented at SRO office",
    ],
    [
        "2",
        "Sub-Registrar searches and opens the registered Hindu marriage record in Kaveri",
        "Sub Registrar",
        "By registration number, party name or Form II-A certificate number",
    ],
    [
        "3",
        "Sub-Registrar verifies court order particulars against the register entry "
        "(court name, case number, order date, parties)",
        "Sub Registrar",
        "Hard stop if parties or registration number do not match",
    ],
    [
        "4",
        "Sub-Registrar uploads / scans the court order copy",
        "Sub Registrar",
        "Mandatory artefact; retained per records policy (Rule 10)",
    ],
    [
        "5",
        "Sub-Registrar records nullity type as declared in the court order — Void "
        "(Sec. 11) or Voidable / annulled (Sec. 12)",
        "Sub Registrar",
        "Selection driven by court order wording",
    ],
    [
        "6",
        "Sub-Registrar flags the marriage registration as Null and Void",
        "Sub Registrar",
        "Endorsement on register entry; supersedes prior Registered / Certificate Issued status",
    ],
    [
        "7",
        "System updates register entry and certificate status",
        "System",
        "Form II-A marked invalid; Hindu Marriages Register shows NULL AND VOID endorsement",
    ],
    [
        "8",
        "System prevents issuance of certified extracts without null/void notation",
        "System",
        "Any subsequent extract shall reflect the endorsement (Rule 8)",
    ],
    [
        "9",
        "Audit trail recorded",
        "System",
        "Officer, timestamp, court order reference, prior and new status",
    ],
]

SEC_11_12_IMPL = "§7.4 (court-order null/void endorsement)"


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
    last: bool = False,
) -> Paragraph:
    matches: list[Paragraph] = []
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            matches.append(p)
        elif contains is not None and contains in t:
            matches.append(p)
    if not matches:
        raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")
    return matches[-1] if last else matches[0]


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)


def find_section_table(doc: Document, third_header: str) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:2] == ["Section", "Topic"] and len(hdr) >= 3 and hdr[2] == third_header:
            return table
    raise KeyError(f"Section table with third header {third_header!r} not found")


def update_sec_11_12_mapping(doc: Document) -> None:
    hma = find_section_table(doc, "relevance")
    for ri in range(1, len(hma.rows)):
        if hma.rows[ri].cells[0].text.strip() == "Sec. 11–12":
            set_cell_text(
                hma.rows[ri].cells[2],
                "Court nullity endorsement — SRO flags registered marriage Null and Void "
                "on receipt of court order (Sec. 11 void; Sec. 12 voidable / annulled)",
            )
            if len(hma.rows[ri].cells) >= 4:
                set_cell_text(hma.rows[ri].cells[3], SEC_11_12_IMPL)
            break
    else:
        raise KeyError("Sec. 11–12 row not found in HMA sections table")


def section_exists(doc: Document, exact: str) -> bool:
    try:
        find_para(doc, exact=exact, heading_only=True)
        return True
    except KeyError:
        return False


def add_null_void_workflow(doc: Document) -> None:
    """Insert §7.4 immediately before §8 (after 7.3.2.3 status table)."""
    if section_exists(doc, "7.4 Hindu Marriage — Null and Void Endorsement (Court Order)"):
        return

    section8 = find_para(doc, exact="8. Functional requirements", heading_only=True, last=True)

    heading = deepcopy(section8._p)
    for child in list(heading):
        if child.tag != qn("w:pPr"):
            heading.remove(child)
    section8._p.addprevious(heading)
    h_para = Paragraph(heading, section8._parent)
    h_para.style = "Heading 3"
    set_para_text(h_para, "7.4 Hindu Marriage — Null and Void Endorsement (Court Order)")

    cursor = h_para
    cursor = insert_paragraph_after(
        cursor,
        "Back-office workflow triggered when a competent court passes an order declaring a "
        "registered Hindu marriage void (HMA 1955 Sec. 11) or annulling a voidable marriage "
        "(Sec. 12). Kaveri does not adjudicate nullity; the Sub-Registrar records the court "
        "order and flags the existing registration as Null and Void. This workflow is "
        "initiated only on presentation of a certified court order copy — not from citizen "
        "self-service or automated eligibility checks.",
        style="Normal",
    )
    cursor = insert_paragraph_after(cursor, "7.4.1 Process steps", style="Heading 4")
    cursor = insert_paragraph_after(
        cursor,
        "Offline (In Person) at the Sub-Registrar office — steps only (no process diagram):",
        style="Normal",
    )
    insert_table_after(cursor, NULL_VOID_STEPS, doc)


def update_contents(doc: Document) -> None:
    for i, p in enumerate(doc.paragraphs[:140]):
        if (
            p.text.strip() == "7.3.2.3 Application Status Model"
            and style_name(p) == "Normal"
        ):
            nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            if "7.4 Hindu Marriage" in nxt:
                return
            p1 = insert_paragraph_after(
                p,
                "7.4 Hindu Marriage — Null and Void Endorsement (Court Order)",
                style="Normal",
            )
            insert_paragraph_after(p1, "7.4.1 Process steps", style="Normal")
            return
    raise KeyError("Contents entry for 7.3.2.3 not found")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.19")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    version_table = doc.tables[1]
    last_change = version_table.rows[-1].cells[3].text.strip()
    if "Null and Void Endorsement" not in last_change:
        add_version_row(
            version_table,
            [
                "1.19",
                "2026-09-02",
                "Nandha Kumar",
                "Add §7.4 Hindu Marriage null/void court-order endorsement workflow (Sec. 11–12)",
                "Prashanth",
            ],
        )

    add_null_void_workflow(doc)
    update_sec_11_12_mapping(doc)
    update_contents(doc)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    find_para(
        doc2,
        exact="7.4 Hindu Marriage — Null and Void Endorsement (Court Order)",
        heading_only=True,
    )
    hma = find_section_table(doc2, "relevance")
    sec_row = None
    for ri in range(1, len(hma.rows)):
        if hma.rows[ri].cells[0].text.strip() == "Sec. 11–12":
            sec_row = hma.rows[ri]
            break
    assert sec_row is not None
    assert SEC_11_12_IMPL in sec_row.cells[-1].text
    print("Verification OK")


if __name__ == "__main__":
    main()
