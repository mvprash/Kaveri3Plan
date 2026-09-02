# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.17.docx from v1.16.

Add Sakala (Karnataka Guarantee of Services) integration requirements:
- §3.7 regulatory reference
- §8.7 functional requirements (FR-HMA-092–093, FR-SMA-068–069)
- §11 Integrations table row
- Glossary, data entity, UI screen, fallback FB-MRG-005, appendix reference
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.16.docx"
DST = BASE / "BRD_Marriage_v1.17.docx"

SCOPE_INTEGRATIONS_OLD = (
    "Integrations: [payment, Aadhaar/eKYC, DigiLocker, SMS, e-Mail, Kutumba portal, "
    "Civil Registration System, Labor Department]."
)
SCOPE_INTEGRATIONS_NEW = (
    "Integrations: [payment, Aadhaar/eKYC, DigiLocker, SMS, e-Mail, Kutumba portal, "
    "Civil Registration System, Labor Department, Sakala (Karnataka Guarantee of Services)]."
)

IFACE_OLD = (
    "Interface requirements: [API list TBD by Architect — must now include eSign, "
    "DSC signing, appointment slots and scan upload]"
)
IFACE_NEW = (
    "Interface requirements: [API list TBD by Architect — must now include eSign, "
    "DSC signing, appointment slots, scan upload and Sakala GSC lifecycle APIs "
    "(acceptance, status, delivery/rejection, appeal sync) per Sakala Mission / NIC "
    "interface specification]"
)

ENTITIES_OLD = (
    "Application/Memorandum, Party (Bride/Bridegroom), Witness, MarriageEvent, "
    "Document, Payment, ScrutinyDecision, RegisterEntry (serial/page/volume), "
    "Certificate (Form II-A),  Special Marriage entities:"
)
ENTITIES_NEW = (
    "Application/Memorandum, Party (Bride/Bridegroom), Witness, MarriageEvent, "
    "Document, Payment, ScrutinyDecision, RegisterEntry (serial/page/volume), "
    "Certificate (Form II-A), SakalaTransaction (GSC number, service/sub-service "
    "codes, processing status, transmission log), Special Marriage entities:"
)

FR_HMA_092 = (
    "Upon successful Hindu Marriage fee payment (FR-HMA-022 / FR-HMA-024), system "
    "shall obtain or register a Sakala Guarantee of Services to Citizen (GSC) number "
    "for the notified Hindu Marriage registration service and display it on the "
    "payment receipt (Form VI equivalent), citizen application dashboard and "
    "certificate download screen; GSC shall be issued only after payment reconciliation "
    "and shall link to the Kaveri application number"
)

FR_HMA_093 = (
    "System shall transmit Hindu Marriage Sakala lifecycle events to the Sakala "
    "platform (https://sakala.kar.nic.in/) — application acceptance after payment, "
    "in-process milestones, service delivery on Form II-A issuance, or rejection with "
    "written reason — using department-approved Sakala service codes, sub-service codes "
    "and SRO office-code mapping; party particulars shall conform to the Sakala schema; "
    "failed or orphaned transmissions shall be logged and retried without blocking the "
    "citizen workflow (see FB-MRG-005)"
)

FR_SMA_068 = (
    "For Special Marriage, system shall register or update Sakala GSC on successful "
    "first payment (notice fee — FR-SMA-049) and on successful second payment "
    "(registration / solemnization fee — FR-SMA-033 / FR-SMA-050), using the "
    "notified Sakala service mapping for Intended Marriage vs Other Forms; GSC shall "
    "appear on receipts and the citizen status view"
)

FR_SMA_069 = (
    "System shall sync Special Marriage Sakala status for notice publication, "
    "objection enquiry outcomes, solemnization / certificate delivery (Fourth or Fifth "
    "Schedule) or rejection to Sakala with party particulars, enclosure flags and "
    "office routing per Sakala Mission interface; SR pendency and Sakala officer "
    "views shall reflect the same application identity (no duplicate or orphan GSC)"
)

FB_MRG_005 = (
    "Sakala integration failures: if the Sakala platform is unreachable or returns a "
    "non-terminal error during GSC registration or status upload, the system shall "
    "complete the Kaveri workflow, persist the application with a pending Sakala sync "
    "flag, queue the payload for automatic retry (orphan-settle pattern), and surface "
    "the GSC to the citizen once synchronization succeeds; officers shall see a "
    "bilingual Sakala sync status indicator on the application"
)

SAKALA_INTEGRATION_ROW = [
    "Sakala (Karnataka Guarantee of Services)",
    "Bidirectional",
    "GSC acknowledgement on payment; service lifecycle sync (acceptance, in-process, "
    "delivery/rejection); party particulars upload; appeal status exchange; citizen "
    "tracking at https://sakala.kar.nic.in/",
    "Both",
    "Integration / Sakala Mission",
    "TBD",
]

UI_SAKALA_ROW = [
    "GSC acknowledgement and tracking",
    "Display Sakala / GSC number, statutory service timeline and link to "
    "https://sakala.kar.nic.in/ status tracking after payment",
    "Both",
    "Karnataka Guarantee of Services Act, 2011",
    "§8.7; FR-HMA-092, FR-SMA-068",
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def add_fr_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para_index(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
    last: bool = False,
) -> int:
    matches: list[int] = []
    for i, p in enumerate(doc.paragraphs):
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            matches.append(i)
        elif contains is not None and contains in t:
            matches.append(i)
    if not matches:
        raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")
    return matches[-1] if last else matches[0]


def find_para(doc: Document, **kwargs) -> Paragraph:
    return doc.paragraphs[find_para_index(doc, **kwargs)]


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def find_table_by_header(doc: Document, first_cell: str, second_cell: str | None = None) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr0 = table.rows[0].cells[0].text.strip()
        hdr1 = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
        if hdr0 == first_cell and (second_cell is None or hdr1 == second_cell):
            return table
    raise KeyError(f"Table not found: {first_cell!r} / {second_cell!r}")


def find_table_containing(doc: Document, req_id: str) -> Table:
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == req_id:
                return table
    raise KeyError(f"Table not found containing {req_id!r}")


def replace_in_paragraphs(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    set_cell_text(cell, cell.text.replace(old, new))


def add_sakala_regulatory_section(before_stakeholders: Paragraph) -> None:
    cursor = insert_paragraph_after(
        before_stakeholders, "3.7 Sakala — Karnataka Guarantee of Services", style="Heading 3"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Source: Karnataka Guarantee of Services to Citizens Act, 2011 (and Amendment Act, "
        "2014); Sakala Mission portal — https://sakala.kar.nic.in/. Marriage registration "
        "services offered through the Department of Stamps and Registration are notified "
        "Sakala services. Every eligible application shall receive a Guarantee of Services "
        "to Citizen (GSC) number so citizens can track status, receive time-bound service "
        "and file appeals (Appeal-I / Appeal-II) when service is delayed or rejected.",
        style="Normal",
    )
    cursor = insert_paragraph_after(
        cursor,
        "Notified marriage-related Sakala services (department baseline — confirm codes "
        "with Sakala Mission during integration design):",
        style="Normal",
    )
    sakala_services = [
        "Hindu Marriage registration (online filing) — GSC issued on successful fee "
        "payment; statutory service timeline commences from payment date (typically 1 "
        "working day for certificate delivery after payment, per department practice).",
        "Special Marriage — Intended Marriage notice and registration.",
        "Special Marriage — Other Forms notice and registration.",
    ]
    for item in sakala_services:
        p = insert_paragraph_after(cursor, item, style="List Bullet")
        cursor = p
    insert_paragraph_after(
        cursor,
        "Kaveri 3.0 shall integrate with the Sakala platform per NIC / Sakala Mission "
        "approved web methods for GSC registration, status updates (acceptance, "
        "in-process, delivered, rejected), party-detail upload and orphan-settle retry "
        "— see §8.7 and §11.",
        style="Normal",
    )


def add_sakala_functional_section(after_86: Paragraph, doc: Document) -> None:
    cursor = insert_paragraph_after(
        after_86, "8.7 Sakala integration (Karnataka Guarantee of Services)", style="Heading 3"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Marriage registration modules shall comply with the Karnataka Guarantee of "
        "Services to Citizens Act, 2011 by issuing a GSC number and synchronizing "
        "service lifecycle data with the Sakala platform (https://sakala.kar.nic.in/). "
        "Integration shall use department-maintained Sakala service / sub-service code "
        "masters and SRO-to-Sakala office-code mapping. Citizens track applications on "
        "the Sakala portal using the GSC; officers reconcile Kaveri pendency with Sakala "
        "officer views.",
        style="Normal",
    )
    fr_header = ["Req ID", "Requirement", "Priority", "Acceptance criteria"]
    fr_rows = [
        fr_header,
        ["FR-HMA-092", FR_HMA_092, "Must", "GSC on receipt and dashboard; linked to application"],
        [
            "FR-HMA-093",
            FR_HMA_093,
            "Must",
            "Lifecycle events transmitted; retry queue for failures; audit log retained",
        ],
        [
            "FR-SMA-068",
            FR_SMA_068,
            "Must",
            "GSC on first and second payments per path; displayed on receipt and status",
        ],
        [
            "FR-SMA-069",
            FR_SMA_069,
            "Must",
            "Notice/registration milestones synced; no duplicate GSC per application",
        ],
    ]
    insert_table_after(cursor, fr_rows, doc)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.17")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    add_version_row(
        doc.tables[1],
        [
            "1.17",
            "2026-09-02",
            "Nandha Kumar",
            "Sakala integration: §3.7 regulatory ref; §8.7 FR-HMA-092–093 / FR-SMA-068–069; "
            "§11 integration row; glossary, entity, UI, FB-MRG-005",
            "Prashanth",
        ],
    )

    replace_in_paragraphs(doc, SCOPE_INTEGRATIONS_OLD, SCOPE_INTEGRATIONS_NEW)
    replace_in_paragraphs(doc, IFACE_OLD, IFACE_NEW)
    replace_in_paragraphs(doc, ENTITIES_OLD, ENTITIES_NEW)

    # Contents TOC entries
    toc_86 = find_para(doc, exact="8.6 Reports and MIS (FR-HMA-039–045, FR-SMA-055–060)")
    insert_paragraph_after(
        toc_86,
        "8.7 Sakala integration (FR-HMA-092–093, FR-SMA-068–069)",
        style="Normal",
    )
    toc_36 = find_para(doc, exact="3.6 Special Marriage statutory forms mapping")
    insert_paragraph_after(
        toc_36,
        "3.7 Sakala — Karnataka Guarantee of Services",
        style="Normal",
    )

    # §3.7 body (before §4 Stakeholders — last occurrence)
    stakeholders = find_para(doc, exact="4. Stakeholders and actors", last=True)
    add_sakala_regulatory_section(stakeholders)

    # §8.7 body (after §8.6)
    sec_86 = find_para(doc, exact="8.6 Reports and MIS", last=True)
    add_sakala_functional_section(sec_86, doc)

    # §11 Integrations table
    integrations = find_table_by_header(doc, "Integration")
    integrations._tbl.append(deepcopy(integrations.rows[-1]._tr))
    set_row(integrations, len(integrations.rows) - 1, SAKALA_INTEGRATION_ROW)

    # UI screen
    ui_table = find_table_by_header(doc, "Screen / step", "Purpose")
    ui_table._tbl.append(deepcopy(ui_table.rows[-1]._tr))
    set_row(ui_table, len(ui_table.rows) - 1, UI_SAKALA_ROW)

    # Glossary
    glossary = find_table_by_header(doc, "Term")
    for term, definition in [
        (
            "GSC (Guarantee of Services to Citizen)",
            "Unique Sakala acknowledgement number issued for notified services; enables "
            "citizens to track application status and file appeals at "
            "https://sakala.kar.nic.in/",
        ),
        (
            "Sakala",
            "Karnataka Guarantee of Services to Citizens programme (Act, 2011); mandates "
            "time-bound delivery and accountability for notified government services",
        ),
    ]:
        glossary._tbl.append(deepcopy(glossary.rows[-1]._tr))
        set_row(glossary, len(glossary.rows) - 1, [term, definition])

    # Fallback FB-MRG-005
    fb_table = find_table_containing(doc, "FB-MRG-004")
    add_fr_row(
        fb_table,
        [
            "FB-MRG-005",
            FB_MRG_005,
            "Must",
            "Citizen workflow not blocked; GSC shown when sync completes; officer sees sync status",
        ],
    )

    # RTM summary tables (post-registration / notifications area)
    rtm_post = find_table_containing(doc, "FR-HMA-033")
    add_fr_row(
        rtm_post,
        [
            "FR-HMA-092",
            FR_HMA_092,
            "Must",
            "GSC on receipt and dashboard after Hindu Marriage payment",
        ],
    )
    add_fr_row(
        rtm_post,
        [
            "FR-HMA-093",
            FR_HMA_093,
            "Must",
            "Sakala lifecycle sync with retry queue",
        ],
    )

    rtm_notif = find_table_containing(doc, "FR-SMA-056")
    add_fr_row(
        rtm_notif,
        [
            "FR-SMA-068",
            FR_SMA_068,
            "Must",
            "GSC on Special Marriage first and second payments",
        ],
    )
    add_fr_row(
        rtm_notif,
        [
            "FR-SMA-069",
            FR_SMA_069,
            "Must",
            "Special Marriage Sakala milestone sync",
        ],
    )

    # Full RTM traceability (template table)
    rtm_full = find_table_by_header(doc, "Req ID", "Act/Rule/Form")
    for row_data in [
        [
            "FR-HMA-092",
            "Karnataka Guarantee of Services Act, 2011",
            "GSC on Hindu Marriage payment",
            "8.7",
            "GSC acknowledgement and tracking",
            "TC-HMA-___",
            "Draft",
        ],
        [
            "FR-HMA-093",
            "Karnataka Guarantee of Services Act, 2011",
            "Sakala lifecycle sync for Hindu Marriage",
            "8.7",
            "GSC acknowledgement and tracking",
            "TC-HMA-___",
            "Draft",
        ],
        [
            "FR-SMA-068",
            "Karnataka Guarantee of Services Act, 2011",
            "GSC on Special Marriage payments",
            "8.7",
            "GSC acknowledgement and tracking",
            "TC-SMA-___",
            "Draft",
        ],
        [
            "FR-SMA-069",
            "Karnataka Guarantee of Services Act, 2011",
            "Sakala lifecycle sync for Special Marriage",
            "8.7",
            "GSC acknowledgement and tracking",
            "TC-SMA-___",
            "Draft",
        ],
    ]:
        rtm_full._tbl.append(deepcopy(rtm_full.rows[-1]._tr))
        set_row(rtm_full, len(rtm_full.rows) - 1, row_data)

    # Appendix reference
    appendix_anchor = find_para(
        doc, contains="OWASP Top 10 — https://owasp.org/", last=True
    )
    insert_paragraph_after(
        appendix_anchor,
        "Sakala (Karnataka Guarantee of Services) — https://sakala.kar.nic.in/; "
        "Karnataka Guarantee of Services to Citizens Act, 2011",
        style="Normal",
    )

    # Update NFR-MRG-AVA-001 to mention Sakala
    ava_table = find_table_containing(doc, "NFR-MRG-AVA-001")
    for row in ava_table.rows:
        if row.cells[0].text.strip() == "NFR-MRG-AVA-001":
            old = row.cells[1].text.strip()
            if "Sakala" not in old:
                set_cell_text(
                    row.cells[1],
                    old.replace(
                        "DigiLocker",
                        "DigiLocker, Sakala",
                    ),
                )
            break

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    find_para(doc2, contains="3.7 Sakala — Karnataka Guarantee of Services")
    find_para(doc2, contains="8.7 Sakala integration")
    find_table_containing(doc2, "FR-HMA-092")
    find_table_containing(doc2, "FR-SMA-069")
    find_table_containing(doc2, "FB-MRG-005")
    int_table = find_table_by_header(doc2, "Integration")
    assert any(
        "Sakala" in row.cells[0].text for row in int_table.rows
    ), "Sakala integration row missing"
    print("Verification OK")


if __name__ == "__main__":
    main()
