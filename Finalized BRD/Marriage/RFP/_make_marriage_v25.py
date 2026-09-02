# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.15.docx from v1.14.

Add sample statutory forms from Acts_Rules/Marriage/ to §3.5 (Hindu Marriage)
and §3.6 (Special Marriage).
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.14.docx"
DST = BASE / "BRD_Marriage_v1.15.docx"
ACTS = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Acts_Rules\Marriage")
MEDIA = BASE / "_sample_forms_media"

HMA_SAMPLES = [
    (
        MEDIA / "hindu_marriage_forms_p1.png",
        "Form I — Memorandum of marriage (page 1). "
        "Source: Acts_Rules/Marriage/Form1.pdf; hindu marriage forms.pdf",
    ),
    (
        MEDIA / "hindu_marriage_forms_p2.png",
        "Form I — Memorandum of marriage (page 2 — witness particulars). "
        "Source: Acts_Rules/Marriage/hindu marriage forms.pdf",
    ),
    (
        MEDIA / "hindu_marriage_forms_p3.png",
        "Form IA — Application for filing of marriage memorandum (Rule 4(2)). "
        "Source: Acts_Rules/Marriage/hindu marriage forms.pdf",
    ),
    (
        MEDIA / "hindu_marriage_forms_p4.png",
        "Form II — Endorsement on reverse of memorandum and duplicate (Rule 4(4)). "
        "Source: Acts_Rules/Marriage/hindu marriage forms.pdf",
    ),
    (
        MEDIA / "hindu_marriage_forms_p5.png",
        "Form II-A — Certificate of registration of marriage (Rule 4(5)). "
        "Source: Acts_Rules/Marriage/hindu marriage forms.pdf",
    ),
    (
        MEDIA / "hindu_marriage_forms_p6.png",
        "Form III — Certificate affixed to monthly duplicate memoranda (Rule 5(1)). "
        "Source: Acts_Rules/Marriage/hindu marriage forms.pdf",
    ),
]

SMA_SAMPLES = [
    (
        ACTS / "Notice_of_Intended_Marriage_Second_Schedule.png",
        "Second Schedule — Notice of intended marriage (Sec. 5 / Sec. 6). "
        "Source: Acts_Rules/Marriage/Notice_of_Intended_Marriage_Second_Schedule.png",
    ),
    (
        ACTS / "Declaration_Third_Schedule.png",
        "Third Schedule — Declarations by parties and witnesses (Sec. 11). "
        "Source: Acts_Rules/Marriage/Declaration_Third_Schedule.png",
    ),
    (
        ACTS / "Certificate_of_Marriage_Fourth_Schedule.png",
        "Fourth Schedule — Certificate of marriage after solemnization (Sec. 13). "
        "Source: Acts_Rules/Marriage/Certificate_of_Marriage_Fourth_Schedule.png",
    ),
    (
        ACTS / "SpecialOtherFormsCertificate.jpeg",
        "Fifth Schedule — Certificate of marriage celebrated in other forms (Sec. 16). "
        "Source: Acts_Rules/Marriage/SpecialOtherFormsCertificate.jpeg",
    ),
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_picture_after(
    paragraph: Paragraph, image_path: Path, width_inches: float = 6.2
) -> Paragraph:
    pic_para = insert_paragraph_after(paragraph, "", style="Normal")
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return pic_para


def find_table_by_header(doc: Document, first_cell: str, second_cell: str | None = None) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr0 = table.rows[0].cells[0].text.strip()
        hdr1 = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
        if hdr0 == first_cell and (second_cell is None or hdr1 == second_cell):
            return table
    raise KeyError(f"Table not found: {first_cell!r} / {second_cell!r}")


def paragraph_after_table(table: Table) -> Paragraph:
    tbl = table._tbl
    nxt = tbl.getnext()
    while nxt is not None and nxt.tag != qn("w:p"):
        nxt = nxt.getnext()
    if nxt is None:
        raise KeyError("No paragraph found after table")
    return Paragraph(nxt, table._parent)


def add_sample_forms(anchor: Paragraph, samples: list[tuple[Path, str]]) -> None:
    cursor = insert_paragraph_after(anchor, "Sample forms", style="Heading 4")
    cursor = insert_paragraph_after(
        cursor,
        "Statutory form layouts reproduced from Acts_Rules/Marriage/ for reference. "
        "Kaveri 3.0 generated outputs shall match the statutory wording and field "
        "structure without alteration unless legally approved.",
        style="Normal",
    )
    for image_path, caption in samples:
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        cursor = insert_picture_after(cursor, image_path)
        cursor = insert_paragraph_after(cursor, caption, style="Caption")
        cursor = insert_paragraph_after(cursor, "", style="Normal")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.15")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-02")
    add_version_row(
        doc.tables[1],
        [
            "1.15",
            "2026-09-02",
            "Nandha Kumar",
            "Add sample statutory forms from Acts_Rules/Marriage/ to §3.5 and §3.6",
            "Prashanth",
        ],
    )

    hma_table = find_table_by_header(doc, "Form", "Rule ref")
    sma_table = find_table_by_header(doc, "Form / Schedule", "Act ref")

    add_sample_forms(paragraph_after_table(hma_table), HMA_SAMPLES)
    add_sample_forms(paragraph_after_table(sma_table), SMA_SAMPLES)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    sample_count = sum(1 for p in doc2.paragraphs if p.text.startswith("Form I —") or p.text.startswith("Second Schedule"))
    pic_count = sum(1 for rel in doc2.part.rels.values() if "image" in rel.reltype)
    print("Sample captions:", sample_count)
    print("Embedded images (total doc):", pic_count)


if __name__ == "__main__":
    main()
