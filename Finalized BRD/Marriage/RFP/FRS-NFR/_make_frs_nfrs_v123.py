# -*- coding: utf-8 -*-
"""Create FRS_and_NFRS_Marriage_v1.23.docx from v1.22 with Scanning Module (§8.8)."""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP\FRS-NFR")
SRC = BASE / "FRS_and_NFRS_Marriage_v1.22.docx"
DST = BASE / "FRS_and_NFRS_Marriage_v1.23.docx"

IFACE_OLD = (
    "Interface requirements: [API list TBD by Architect — must now include eSign, "
    "DSC signing, appointment slots, scan upload and Sakala GSC lifecycle APIs "
    "(acceptance, status, delivery/rejection, appeal sync) per Sakala Mission / NIC "
    "interface specification]"
)
IFACE_NEW = (
    "Interface requirements: [API list TBD by Architect — must now include eSign, "
    "DSC signing, appointment slots, Scanning Module ingest / AV / index APIs, "
    "scan upload and Sakala GSC lifecycle APIs (acceptance, status, "
    "delivery/rejection, appeal sync) per Sakala Mission / NIC interface "
    "specification]"
)

ENTITIES_OLD = (
    "Application/Memorandum, Party (Bride/Bridegroom), Witness, MarriageEvent, "
    "Document, Payment, ScrutinyDecision, RegisterEntry (serial/page/volume), "
    "Certificate (Form II-A), NullityEndorsement (court order artefact, nullity "
    "type, endorsement date, prior status, endorsing officer), SakalaTransaction "
    "(GSC number, service/sub-service codes, processing status, transmission log), "
    "Special Marriage entities:"
)
ENTITIES_NEW = (
    "Application/Memorandum, Party (Bride/Bridegroom), Witness, MarriageEvent, "
    "Document, Payment, ScrutinyDecision, RegisterEntry (serial/page/volume), "
    "Certificate (Form II-A), NullityEndorsement (court order artefact, nullity "
    "type, endorsement date, prior status, endorsing officer), SakalaTransaction "
    "(GSC number, service/sub-service codes, processing status, transmission log), "
    "ScanArtefact (application link, document type, page count, checksum, AV "
    "result, ingest user/workstation, retention class), Special Marriage entities:"
)

FR_SCAN_001 = (
    "System shall provide a shared Scanning Module that accepts multi-page scanned "
    "artefacts from authorized FDA / SDA / DEO / SR roles via office workstation "
    "upload or connected scanner, for Marriage Offline (and shared platform) use"
)
FR_SCAN_002 = (
    "All artefacts ingested through the Scanning Module shall pass antivirus / "
    "malware scanning before persistence; infected or unscannable files shall be "
    "rejected with a clear bilingual (English / Kannada) message and shall not be "
    "linked to the application"
)
FR_SCAN_003 = (
    "Successfully ingested scans shall be indexed and linked to the Marriage "
    "application number, statutory document type (e.g. signed Form I / Form IA / "
    "Form II / Form II-A, notice, certificate, certified court order), office, "
    "channel and actor; stored in the Document / scan store with checksum, "
    "timestamp and retention class per records policy"
)
FR_SCAN_004 = (
    "Hindu Marriage and Special Marriage Offline (and Null and Void) steps that "
    "require scan / upload of signed forms, notices, certificates or court orders "
    "shall invoke the Scanning Module APIs rather than ad-hoc unmanaged file upload"
)
FR_SCAN_005 = (
    "Scanning Module shall support preview, page re-order / delete, minimum "
    "resolution and blank-page checks, and re-scan before final submit; password-"
    "protected PDFs and unsupported formats remain blocked (FB-MRG-003)"
)
FR_SCAN_006 = (
    "System shall retain an audit trail for each scan ingest — actor, role, "
    "office, workstation / device identifier, file hash, AV result, document type "
    "and linked application — queryable by SR / audit roles"
)
FR_SCAN_007 = (
    "Marriage workflows shall not mark Offline scan-mandatory steps complete "
    "until required artefacts are successfully ingested and indexed via the "
    "Scanning Module; partial or failed ingest shall leave the step pending"
)

FB_MRG_006 = (
    "Scanning Module failures: if antivirus, ingest or object-store persistence "
    "fails, the system shall not mark the Offline step complete, shall retain any "
    "temporary upload only as needed for retry, display a bilingual error with "
    "retry guidance, and log the failure for operations; the citizen / officer "
    "Marriage workflow shall remain resumable after successful re-scan"
)

SCAN_INTEGRATION_ROW = [
    "Scanning Module (shared platform)",
    "Inbound / Internal",
    "Secure scan ingest, antivirus gate, indexing hooks and Document / scan store "
    "persistence for DEO/SRO uploaded Marriage Offline artefacts (forms, notices, "
    "certificates, court orders); shared spine for later Document Registration",
    "Offline (primary); Both where upload applies",
    "Platform / Scanning",
    "TBD",
]

UI_SCAN_ROW = [
    "Scanning Module — ingest and preview",
    "Capture / upload multi-page scans, run AV and quality checks, preview and "
    "link artefacts to the Marriage application before Offline step completion",
    "Offline",
    "Registration Rules 1965 (scanning / digital storage practice); Phase 1 "
    "Scanning spine",
    "§8.8; FR-SCAN-001–007",
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def add_fr_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para_index(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
    last: bool = False,
) -> int:
    matches: list[int] = []
    for i, p in enumerate(doc.paragraphs):
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            matches.append(i)
        elif contains is not None and contains in t:
            matches.append(i)
    if not matches:
        raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")
    return matches[-1] if last else matches[0]


def find_para(doc: Document, **kwargs) -> Paragraph:
    return doc.paragraphs[find_para_index(doc, **kwargs)]


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def find_table_by_header(
    doc: Document, first_cell: str, second_cell: str | None = None
) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr0 = table.rows[0].cells[0].text.strip()
        hdr1 = (
            table.rows[0].cells[1].text.strip()
            if len(table.rows[0].cells) > 1
            else ""
        )
        if hdr0 == first_cell and (second_cell is None or hdr1 == second_cell):
            return table
    raise KeyError(f"Table not found: {first_cell!r} / {second_cell!r}")


def find_table_containing(doc: Document, req_id: str) -> Table:
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == req_id:
                return table
    raise KeyError(f"Table not found containing {req_id!r}")


def replace_in_paragraphs(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    set_cell_text(cell, cell.text.replace(old, new))


def add_scanning_functional_section(after_87: Paragraph, doc: Document) -> None:
    cursor = insert_paragraph_after(
        after_87, "8.8 Scanning Module (shared platform)", style="Heading 3"
    )
    cursor = insert_paragraph_after(
        cursor,
        "Marriage Registration — especially Offline Hindu Marriage and Special "
        "Marriage channels, and Null and Void court-order upload — shall consume "
        "the shared Kaveri 3.0 Scanning Module for secure ingest, antivirus "
        "scanning, indexing and storage of physical document artefacts. The "
        "Scanning Module is a Phase 1 platform capability (scan ingest, AV, "
        "indexing hooks, DEO/SRO upload path) and is the mandated path for "
        "Marriage Offline scan/upload rather than unmanaged ad-hoc file storage. "
        "It is designed for reuse by Document Registration and related modules.",
        style="Normal",
    )
    fr_header = ["Req ID", "Requirement", "Priority", "Acceptance criteria"]
    fr_rows = [
        fr_header,
        [
            "FR-SCAN-001",
            FR_SCAN_001,
            "Must",
            "Authorized roles can ingest multi-page scans from office workstation / scanner",
        ],
        [
            "FR-SCAN-002",
            FR_SCAN_002,
            "Must",
            "Infected files rejected before persist; bilingual message shown; no app link",
        ],
        [
            "FR-SCAN-003",
            FR_SCAN_003,
            "Must",
            "Artefact indexed to application + doc type; checksum and retention recorded",
        ],
        [
            "FR-SCAN-004",
            FR_SCAN_004,
            "Must",
            "Offline Marriage scan steps call Scanning Module APIs only",
        ],
        [
            "FR-SCAN-005",
            FR_SCAN_005,
            "Must",
            "Preview and quality checks available; blocked formats still rejected",
        ],
        [
            "FR-SCAN-006",
            FR_SCAN_006,
            "Must",
            "Ingest audit trail queryable by SR / audit roles",
        ],
        [
            "FR-SCAN-007",
            FR_SCAN_007,
            "Must",
            "Offline step stays pending until required scans successfully indexed",
        ],
    ]
    insert_table_after(cursor, fr_rows, doc)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    set_cell_text(doc.tables[0].rows[2].cells[1], "1.23")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-03")
    add_version_row(
        doc.tables[1],
        [
            "1.23",
            "2026-09-03",
            "Nandha Kumar",
            "Add §8.8 Scanning Module (FR-SCAN-001–007); integration, UI, entity, "
            "FB-MRG-006, RTM; interface API note",
            "Prashanth",
        ],
    )

    replace_in_paragraphs(doc, IFACE_OLD, IFACE_NEW)
    replace_in_paragraphs(doc, ENTITIES_OLD, ENTITIES_NEW)

    # Contents TOC
    toc_87 = find_para(
        doc, exact="8.7 Sakala integration (FR-HMA-092–093, FR-SMA-068–069)"
    )
    insert_paragraph_after(
        toc_87,
        "8.8 Scanning Module (FR-SCAN-001–007)",
        style="Normal",
    )

    # Fix Contents FB range if still 001–004
    for p in doc.paragraphs:
        if p.text.strip() == "17. System Fallbacks & Error Handling (FB-MRG-001–004)":
            set_para_text(
                p, "17. System Fallbacks & Error Handling (FB-MRG-001–006)"
            )
            break

    # §8.8 after Sakala heading (table follows heading; insert after heading so
    # new section appears after Sakala table when body order is resolved — we
    # insert after the Sakala FR table by finding the paragraph after the table.
    # Practical approach: insert after Heading 2 "9. Business rules" would put
    # section in wrong place. Insert after Sakala heading then move? Better:
    # find body child after Sakala table.
    sec_87 = find_para(
        doc, exact="8.7 Sakala integration (Karnataka Guarantee of Services)", last=True
    )
    # Insert immediately after Sakala table: walk body after sec_87._p
    body = doc.element.body
    children = list(body)
    sakala_p_idx = children.index(sec_87._p)
    insert_after_el = sec_87._p
    # Skip following tbl if present (Sakala FR table)
    if sakala_p_idx + 1 < len(children) and children[sakala_p_idx + 1].tag.endswith(
        "}tbl"
    ):
        # Create a temporary paragraph after the table as anchor
        # python-docx: insert_paragraph_after needs a Paragraph; use next heading
        # We'll insert scanning section before "9. Business rules"
        pass

    before_9 = find_para(doc, exact="9. Business rules", last=True)
    # Insert section just before §9 by inserting after previous sibling.
    # Find element immediately before §9 heading
    prev = before_9._p.getprevious()
    # If prev is Sakala table, create anchor paragraph after it via insert on a
    # dummy: insert_paragraph_after on a Paragraph wrapping prev if it's a p,
    # else insert new p after tbl then build section.
    from docx.oxml import OxmlElement

    if prev is not None and prev.tag.endswith("}tbl"):
        # Insert empty Normal paragraph after table, then build section from it
        new_p_el = OxmlElement("w:p")
        prev.addnext(new_p_el)
        anchor = Paragraph(new_p_el, before_9._parent)
        # Heading + intro replace empty anchor content
        set_para_text(anchor, "")
        # Rebuild: use add_scanning starting from a heading inserted via anchor
        # Replace: set anchor as heading 8.8
        anchor.style = "Heading 3"
        set_para_text(anchor, "8.8 Scanning Module (shared platform)")
        cursor = insert_paragraph_after(
            anchor,
            "Marriage Registration — especially Offline Hindu Marriage and Special "
            "Marriage channels, and Null and Void court-order upload — shall consume "
            "the shared Kaveri 3.0 Scanning Module for secure ingest, antivirus "
            "scanning, indexing and storage of physical document artefacts. The "
            "Scanning Module is a Phase 1 platform capability (scan ingest, AV, "
            "indexing hooks, DEO/SRO upload path) and is the mandated path for "
            "Marriage Offline scan/upload rather than unmanaged ad-hoc file storage. "
            "It is designed for reuse by Document Registration and related modules.",
            style="Normal",
        )
        fr_header = ["Req ID", "Requirement", "Priority", "Acceptance criteria"]
        fr_rows = [
            fr_header,
            [
                "FR-SCAN-001",
                FR_SCAN_001,
                "Must",
                "Authorized roles can ingest multi-page scans from office workstation / scanner",
            ],
            [
                "FR-SCAN-002",
                FR_SCAN_002,
                "Must",
                "Infected files rejected before persist; bilingual message shown; no app link",
            ],
            [
                "FR-SCAN-003",
                FR_SCAN_003,
                "Must",
                "Artefact indexed to application + doc type; checksum and retention recorded",
            ],
            [
                "FR-SCAN-004",
                FR_SCAN_004,
                "Must",
                "Offline Marriage scan steps call Scanning Module APIs only",
            ],
            [
                "FR-SCAN-005",
                FR_SCAN_005,
                "Must",
                "Preview and quality checks available; blocked formats still rejected",
            ],
            [
                "FR-SCAN-006",
                FR_SCAN_006,
                "Must",
                "Ingest audit trail queryable by SR / audit roles",
            ],
            [
                "FR-SCAN-007",
                FR_SCAN_007,
                "Must",
                "Offline step stays pending until required scans successfully indexed",
            ],
        ]
        insert_table_after(cursor, fr_rows, doc)
    else:
        # Fallback: insert after Sakala heading (may place before table)
        add_scanning_functional_section(sec_87, doc)

    # §11 Integrations
    integrations = find_table_by_header(doc, "Integration")
    # Update existing Document / scan store row purpose if present
    for row in integrations.rows:
        if row.cells[0].text.strip() == "Document / scan store":
            set_cell_text(
                row.cells[2],
                "Persistence backend for Scanning Module ingest — DEO-uploaded "
                "signed Form I, Form IA, Form II & II-A, notices, certificates "
                "and court-order copies (§8.8)",
            )
            break
    integrations._tbl.append(deepcopy(integrations.rows[-1]._tr))
    set_row(integrations, len(integrations.rows) - 1, SCAN_INTEGRATION_ROW)

    # UI
    ui_table = find_table_by_header(doc, "Screen / step", "Purpose")
    ui_table._tbl.append(deepcopy(ui_table.rows[-1]._tr))
    set_row(ui_table, len(ui_table.rows) - 1, UI_SCAN_ROW)

    # Fallback FB-MRG-006
    fb_table = find_table_containing(doc, "FB-MRG-005")
    add_fr_row(
        fb_table,
        [
            "FB-MRG-006",
            FB_MRG_006,
            "Must",
            "Step stays pending; bilingual retry; workflow resumable after successful scan",
        ],
    )

    # RTM full traceability
    rtm_full = find_table_by_header(doc, "Req ID", "Act/Rule/Form")
    for row_data in [
        [
            "FR-SCAN-001",
            "Registration Rules 1965 / Phase 1 Scanning spine",
            "Shared Scanning Module ingest for authorized office roles",
            "8.8",
            "Scanning Module — ingest and preview",
            "TC-SCAN-___",
            "Draft",
        ],
        [
            "FR-SCAN-002",
            "Security / AV policy",
            "Antivirus gate before scan persistence",
            "8.8",
            "Scanning Module — ingest and preview",
            "TC-SCAN-___",
            "Draft",
        ],
        [
            "FR-SCAN-003",
            "Records / retention policy",
            "Index and store scan artefacts with checksum",
            "8.8",
            "Scanning Module — ingest and preview",
            "TC-SCAN-___",
            "Draft",
        ],
        [
            "FR-SCAN-004",
            "Marriage Offline BRD / HLD",
            "Marriage Offline scan steps use Scanning Module APIs",
            "8.8",
            "Scanning Module — ingest and preview",
            "TC-SCAN-___",
            "Draft",
        ],
        [
            "FR-SCAN-005",
            "Usability / quality gate",
            "Preview, quality checks and re-scan",
            "8.8",
            "Scanning Module — ingest and preview",
            "TC-SCAN-___",
            "Draft",
        ],
        [
            "FR-SCAN-006",
            "Audit policy",
            "Scan ingest audit trail",
            "8.8",
            "Scanning Module — ingest and preview",
            "TC-SCAN-___",
            "Draft",
        ],
        [
            "FR-SCAN-007",
            "Marriage Offline completeness",
            "Offline step complete only after successful scan index",
            "8.8",
            "Scanning Module — ingest and preview",
            "TC-SCAN-___",
            "Draft",
        ],
    ]:
        rtm_full._tbl.append(deepcopy(rtm_full.rows[-1]._tr))
        set_row(rtm_full, len(rtm_full.rows) - 1, row_data)

    # Appendix reference
    try:
        appendix_anchor = find_para(
            doc, contains="Sakala (Karnataka Guarantee of Services)", last=True
        )
    except KeyError:
        appendix_anchor = find_para(
            doc, contains="OWASP Top 10", last=True
        )
    insert_paragraph_after(
        appendix_anchor,
        "Scanning Module (Kaveri 3.0 Phase 1 shared platform) — scan ingest, "
        "antivirus, indexing hooks and DEO/SRO upload path; Programme Plan "
        "Scanning spine",
        style="Normal",
    )

    doc.save(str(DST))
    print(f"Wrote {DST}")

    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    find_para(doc2, contains="8.8 Scanning Module")
    find_table_containing(doc2, "FR-SCAN-001")
    find_table_containing(doc2, "FR-SCAN-007")
    find_table_containing(doc2, "FB-MRG-006")
    int_table = find_table_by_header(doc2, "Integration")
    assert any(
        "Scanning Module" in row.cells[0].text for row in int_table.rows
    ), "Scanning integration row missing"
    # Confirm section order: 8.7 before 8.8 before 9
    headings = [
        p.text.strip()
        for p in doc2.paragraphs
        if style_name(p).startswith("Heading")
        and (
            p.text.strip().startswith("8.7")
            or p.text.strip().startswith("8.8")
            or p.text.strip().startswith("9.")
        )
    ]
    print("Heading order:", headings)
    assert headings.index([h for h in headings if h.startswith("8.7")][0]) < headings.index(
        [h for h in headings if h.startswith("8.8")][0]
    )
    assert headings.index([h for h in headings if h.startswith("8.8")][0]) < headings.index(
        [h for h in headings if h.startswith("9.")][0]
    )
    print("Verification OK")


if __name__ == "__main__":
    main()
