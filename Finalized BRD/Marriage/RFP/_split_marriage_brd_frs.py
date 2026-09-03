# -*- coding: utf-8 -*-
"""Split BRD_Marriage_v1.22.docx into BRD (≤§7.5) and FRS/NFRS (§8–Appendix A)."""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.22.docx"
DST_BRD = BASE / "BRD_Marriage_BRD_v1.22.docx"
DST_FRS = BASE / "FRS_and_NFRS_Marriage_v1.22.docx"
BACKUP_FULL = BASE / "BRD_Marriage_v1.22_combined.docx"

BRD_TITLE = "Business Requirements Document (BRD)"
FRS_TITLE = "Functional Requirements Specification (FRS) and Non-Functional Requirements (NFRs)"

SECTION8 = "8. Functional requirements"
EXEC_SUMMARY = "Executive summary"
TOC_START_BRD_CUT = "8. Functional requirements"
TOC_END = "Appendix A – References"
# TOC may use en-dash or hyphen variants
TOC_END_ALT = "Appendix A"


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    row = table.rows[-1]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def body_para_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def body_para_style(el) -> str:
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        return ""
    s = pPr.find(qn("w:pStyle"))
    return (s.get(qn("w:val")) or "") if s is not None else ""


def find_body_index(body, *, text_exact: str | None = None, text_startswith: str | None = None, style_prefix: str | None = None, last: bool = False) -> int:
    matches: list[int] = []
    for i, child in enumerate(body):
        if not child.tag.endswith("}p"):
            continue
        text = body_para_text(child)
        style = body_para_style(child)
        if style_prefix and not style.startswith(style_prefix):
            continue
        if text_exact is not None and text == text_exact:
            matches.append(i)
        elif text_startswith is not None and text.startswith(text_startswith):
            matches.append(i)
    if not matches:
        raise KeyError(f"Not found: exact={text_exact!r} startswith={text_startswith!r}")
    return matches[-1] if last else matches[0]


def delete_body_range(body, start: int, end: int) -> int:
    """Delete body children in [start, end). Returns count deleted."""
    children = list(body)
    deleted = 0
    for child in children[start:end]:
        body.remove(child)
        deleted += 1
    return deleted


def trim_toc_brd(doc: Document) -> None:
    """Remove Contents entries from §8 through Appendix A (keep ≤7.5.1)."""
    in_contents = False
    removing = False
    to_remove: list = []
    for p in doc.paragraphs:
        t = p.text.strip()
        style = p.style.name if p.style else ""
        if style.startswith("Heading") and t == "Contents":
            in_contents = True
            continue
        if in_contents and style.startswith("Heading") and t == EXEC_SUMMARY:
            break
        if not in_contents:
            continue
        if t.startswith(TOC_START_BRD_CUT) or t == "8. Functional requirements":
            removing = True
        if removing:
            to_remove.append(p._p)
            if t.startswith(TOC_END_ALT):
                break
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def trim_toc_frs(doc: Document) -> None:
    """Remove Contents entries for §1–§7.5.1 (keep §8–Appendix A)."""
    in_contents = False
    to_remove: list = []
    for p in doc.paragraphs:
        t = p.text.strip()
        style = p.style.name if p.style else ""
        if style.startswith("Heading") and t == "Contents":
            in_contents = True
            continue
        if in_contents and style.startswith("Heading") and t == EXEC_SUMMARY:
            break
        if not in_contents:
            continue
        # Keep from §8 onward
        if t.startswith("8.") or t.startswith("9.") or t.startswith("10.") or t.startswith("11.") or t.startswith("12.") or t.startswith("13.") or t.startswith("14.") or t.startswith("15.") or t.startswith("16.") or t.startswith("17.") or t.startswith("18.") or t.startswith(TOC_END_ALT):
            continue
        # Remove 1–7.x TOC lines
        if t[:1].isdigit() or t.startswith("1.") or t.startswith("2.") or t.startswith("3.") or t.startswith("4.") or t.startswith("5.") or t.startswith("6.") or t.startswith("7."):
            to_remove.append(p._p)
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def update_related_docs_brd(doc: Document) -> None:
    """Point related-docs table at companion FRS/NFRS file."""
    if len(doc.tables) < 3:
        return
    rel = doc.tables[2]
    # Ensure a row references the FRS/NFRS companion
    found = False
    for row in rel.rows[1:]:
        if "FRS" in row.cells[0].text.upper() or "FRS" in row.cells[1].text.upper():
            set_cell_text(row.cells[0], "FRS-NFRS-K3-MRG-001")
            set_cell_text(row.cells[1], "FRS and NFRs (Marriage) — §§8–Appendix A")
            set_cell_text(row.cells[2], "FRS_and_NFRS_Marriage_v1.22.docx")
            found = True
            break
    if not found and len(rel.rows) >= 2:
        rel._tbl.append(deepcopy(rel.rows[-1]._tr))
        set_cell_text(rel.rows[-1].cells[0], "FRS-NFRS-K3-MRG-001")
        set_cell_text(rel.rows[-1].cells[1], "FRS and NFRs (Marriage) — §§8–Appendix A")
        set_cell_text(rel.rows[-1].cells[2], "FRS_and_NFRS_Marriage_v1.22.docx")


def update_related_docs_frs(doc: Document) -> None:
    if len(doc.tables) < 3:
        return
    rel = doc.tables[2]
    # First data row: this document
    if len(rel.rows) >= 2:
        set_cell_text(rel.rows[1].cells[0], "FRS-NFRS-K3-MRG-001")
        set_cell_text(rel.rows[1].cells[1], "This document (FRS and NFRs)")
        set_cell_text(rel.rows[1].cells[2], "FRS_and_NFRS_Marriage_v1.22.docx")
    # Add / update BRD companion
    found = False
    for row in rel.rows[2:]:
        if "BRD" in row.cells[0].text.upper() or "BRD" in row.cells[1].text.upper():
            set_cell_text(row.cells[0], "BRD-K3-MRG-001")
            set_cell_text(row.cells[1], "BRD (Marriage) — Executive summary through §7.5")
            set_cell_text(row.cells[2], "BRD_Marriage_BRD_v1.22.docx")
            found = True
            break
    if not found:
        rel._tbl.append(deepcopy(rel.rows[-1]._tr))
        set_cell_text(rel.rows[-1].cells[0], "BRD-K3-MRG-001")
        set_cell_text(rel.rows[-1].cells[1], "BRD (Marriage) — Executive summary through §7.5")
        set_cell_text(rel.rows[-1].cells[2], "BRD_Marriage_BRD_v1.22.docx")


def make_brd(src: Path, dst: Path) -> None:
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    # Keep title as BRD; clarify module subtitle unchanged
    set_para_text(
        next(p for p in doc.paragraphs if p.text.strip() == BRD_TITLE or "Business Requirements Document" in p.text),
        BRD_TITLE,
    )

    # Document control
    set_cell_text(doc.tables[0].rows[1].cells[1], "BRD-K3-MRG-001")
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.22")
    set_cell_text(
        doc.tables[0].rows[3].cells[1],
        "Draft / In review — BRD part (Executive summary through §7.5)",
    )
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-03")

    last_change = doc.tables[1].rows[-1].cells[3].text.strip()
    if "Split into BRD" not in last_change:
        add_version_row(
            doc.tables[1],
            [
                "1.22",
                "2026-09-03",
                "Nandha Kumar",
                "Split: this file is BRD only (Executive summary through §7.5); FRS/NFRs in companion file",
                "Prashanth",
            ],
        )

    update_related_docs_brd(doc)
    trim_toc_brd(doc)

    # Re-resolve body after TOC edits
    body = doc.element.body
    idx8 = find_body_index(
        body,
        text_exact=SECTION8,
        style_prefix="Heading",
        last=True,
    )
    # Delete from §8 through end (before sectPr)
    end = len(body)
    if body[end - 1].tag.endswith("}sectPr"):
        end -= 1
    n = delete_body_range(body, idx8, end)
    doc.save(str(dst))
    print(f"Wrote BRD: {dst} (removed {n} body elements from §8 onward)")


def make_frs(src: Path, dst: Path) -> None:
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    # Title
    for p in doc.paragraphs:
        if "Business Requirements Document" in p.text or p.text.strip() == BRD_TITLE:
            set_para_text(p, FRS_TITLE)
            break

    # Document control
    set_cell_text(doc.tables[0].rows[1].cells[1], "FRS-NFRS-K3-MRG-001")
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.22")
    set_cell_text(
        doc.tables[0].rows[3].cells[1],
        "Draft / In review — FRS and NFRs (§§8–Appendix A)",
    )
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-03")

    last_change = doc.tables[1].rows[-1].cells[3].text.strip()
    if "Split into FRS" not in last_change and "FRS/NFRs" not in last_change:
        add_version_row(
            doc.tables[1],
            [
                "1.22",
                "2026-09-03",
                "Nandha Kumar",
                "Split: this file is FRS and NFRs (§§8–Appendix A); BRD (≤§7.5) in companion file",
                "Prashanth",
            ],
        )

    update_related_docs_frs(doc)
    trim_toc_frs(doc)

    body = doc.element.body
    idx_exec = find_body_index(
        body,
        text_exact=EXEC_SUMMARY,
        style_prefix="Heading",
        last=True,
    )
    idx8 = find_body_index(
        body,
        text_exact=SECTION8,
        style_prefix="Heading",
        last=True,
    )
    n = delete_body_range(body, idx_exec, idx8)
    doc.save(str(dst))
    print(f"Wrote FRS/NFRS: {dst} (removed {n} body elements before §8)")


def verify() -> None:
    brd = Document(str(DST_BRD))
    frs = Document(str(DST_FRS))

    brd_headings = [p.text.strip() for p in brd.paragraphs if p.style and p.style.name.startswith("Heading") and p.text.strip()]
    frs_headings = [p.text.strip() for p in frs.paragraphs if p.style and p.style.name.startswith("Heading") and p.text.strip()]

    assert any(h.startswith("7.5") for h in brd_headings), "BRD missing §7.5"
    assert any("7.5.1" in h for h in brd_headings), "BRD missing §7.5.1"
    assert not any(h.startswith("8.") for h in brd_headings), "BRD still has §8"
    assert not any(h.startswith("Appendix") for h in brd_headings), "BRD still has Appendix"

    assert any(h.startswith("8.") for h in frs_headings), "FRS missing §8"
    assert any(h.startswith("15.") for h in frs_headings), "FRS missing §15 NFRs"
    assert any(h.startswith("Appendix") for h in frs_headings), "FRS missing Appendix A"
    assert not any(h.startswith("7.") for h in frs_headings), "FRS still has §7"
    assert not any(h == EXEC_SUMMARY for h in frs_headings), "FRS still has Executive summary"

    assert BRD_TITLE in brd.paragraphs[1].text or any(BRD_TITLE in p.text for p in brd.paragraphs[:5])
    assert any("FRS" in p.text and "NFR" in p.text for p in frs.paragraphs[:5]), "FRS title missing"

    print("Verification OK")
    print(f"  BRD headings (sample): {[h for h in brd_headings if h[:1].isdigit() or h.startswith('7')][-5:]}")
    print(f"  FRS headings (sample): {[h for h in frs_headings if h.startswith(('8', '9', '15', 'Appendix'))][:6]}")


def main() -> None:
    if not SRC.exists() and not BACKUP_FULL.exists():
        raise FileNotFoundError(SRC)

    # Always work from the combined original. On first run, archive SRC as combined.
    if not BACKUP_FULL.exists():
        if not SRC.exists():
            raise FileNotFoundError(SRC)
        shutil.copy2(SRC, BACKUP_FULL)
        print(f"Backed up full document to {BACKUP_FULL}")
    else:
        print(f"Using combined backup: {BACKUP_FULL}")

    source = BACKUP_FULL
    make_brd(source, DST_BRD)
    make_frs(source, DST_FRS)
    verify()


if __name__ == "__main__":
    main()
