# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.10.docx from v1.9.

Align §7.2.2.1–7.2.2.4 narrative, Online/Offline notice step tables, status
model and related FR-SMA / §8.2 text with the updated Special Marriage Notice
Online / Offline process diagrams (Whether Marriage Taken place branch;
Other Forms Enter Marriage Details before Prerequisite; Online e-sign before
portal display).
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.9.docx"
DST = BASE / "BRD_Marriage_v1.10.docx"

ONLINE_NOTICE_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    [
        "9",
        "If Aadhaar available → e-KYC / Face Authentication on Bride & Bridegroom details; "
        "else Enter Bride & Bridegroom details",
        "System / Citizen",
        "Continues from 7.2.2.1 step 8; per Online notice diagram (Aadhaar YES/NO branch)",
    ],
    [
        "10",
        "Review summary and proceed document uploading",
        "System / Citizen",
        "Summary of captured particulars",
    ],
    [
        "11",
        "Upload Identity Proof, Photo, Age Proof, Address Proof (Bridegroom & Bride)",
        "Citizen",
        "Mandatory supporting documents including individual photographs",
    ],
    [
        "12",
        "SR Verification (decision)",
        "Sub Registrar",
        "Approve or Reject",
    ],
    [
        "12a",
        "Reject → return to Prerequisite & declaration",
        "Sub Registrar → Citizen",
        "Citizen corrects and resubmits (diagram returns to Prerequisite)",
    ],
    [
        "13",
        "First Payment",
        "System / Citizen",
        "Notice fee per Karnataka Special Marriage fee schedule",
    ],
    [
        "14",
        "Notice Generated",
        "System",
        "Statutory notice generated after first payment (no Online DEO / appointment)",
    ],
    [
        "15",
        "Proceed with e-sign",
        "Citizen",
        "Citizen eSign on the generated notice before portal publication",
    ],
    [
        "16",
        "Marriage notice displayed in portal",
        "System",
        "Sec. 6 publication (digital); follows e-sign per Online notice diagram",
    ],
    [
        "17",
        "30-day countdown starts",
        "System",
        "Objection period per Sec. 7 / Sec. 16 from publication / portal display",
    ],
]

OFFLINE_NOTICE_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    [
        "9",
        "If Aadhaar available → e-KYC / Face Authentication on Bride & Bridegroom details; "
        "else Enter Bride & Bridegroom details",
        "System / Citizen",
        "Continues from 7.2.2.1 step 8; per Offline notice diagram (Aadhaar YES/NO branch)",
    ],
    [
        "10",
        "Review summary and proceed document uploading",
        "System / Citizen",
        "Summary of captured particulars",
    ],
    [
        "11",
        "Upload Identity / Age / Address proofs (Bridegroom & Bride)",
        "Citizen",
        "Mandatory documents (individual photos captured later by DEO at office)",
    ],
    [
        "12",
        "SR Verification (decision)",
        "Sub Registrar",
        "Approve or Reject",
    ],
    [
        "12a",
        "Reject → return to Prerequisite & declaration",
        "Sub Registrar → Citizen",
        "Citizen corrects and resubmits (diagram returns to Prerequisite)",
    ],
    ["13", "First Payment", "System / Citizen", "Notice fee"],
    ["14", "Schedule appointment with SR", "System / Citizen", "After first payment"],
    ["15", "SR Generates Notice", "Sub Registrar", "Statutory notice generation"],
    ["16", "Selects DEO", "Sub Registrar", "Assign FDA/SDA/DEO"],
    [
        "17",
        "Capture individual photos of Bride & Bridegroom",
        "FDA / SDA / DEO",
        "Office visit activity",
    ],
    [
        "18",
        "Download, Print, Sign, Scan and Upload notice",
        "FDA / SDA / DEO",
        "Physical notice processed into system",
    ],
    [
        "19",
        "Paste form on respective Notice Board",
        "FDA / SDA / DEO",
        "Sec. 6(2) conspicuous place publication; also drives portal display",
    ],
    ["20", "30-day countdown starts", "System", "Objection period per Sec. 7 / Sec. 16"],
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def clear_extra_rows(table: Table, keep_rows: int) -> None:
    tbl = table._tbl
    while len(table.rows) > keep_rows:
        tbl.remove(table.rows[-1]._tr)


def ensure_rows(table: Table, needed: int) -> None:
    while len(table.rows) < needed:
        table._tbl.append(deepcopy(table.rows[-1]._tr))


def fill_table(table: Table, rows: list[list[str]]) -> None:
    ensure_rows(table, len(rows))
    for ri, vals in enumerate(rows):
        set_row(table, ri, vals)
    clear_extra_rows(table, len(rows))


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    p_pr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not p_pr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_table_containing(doc: Document, exact: str) -> Table:
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == exact:
                return table
    raise KeyError(f"Table not found containing {exact!r}")


def find_fr_row(table: Table, req_id: str):
    for row in table.rows:
        if row.cells[0].text.strip() == req_id:
            return row
    raise KeyError(f"FR row not found: {req_id}")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Cover + version history
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.10")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-08-27")
    add_version_row(
        doc.tables[1],
        [
            "1.10",
            "2026-08-27",
            "Nandha Kumar",
            "Aligned §7.2.2.1–7.2.2.4 Special Marriage Notice Online/Offline steps, "
            "status model and FR-SMA-001 / 013 / 061 / §8.2 with updated process "
            "diagrams (Whether Marriage Taken place branch; Other Forms Enter "
            "Marriage Details before Prerequisite; Online e-sign before portal display)",
            "Prashanth",
        ],
    )

    # §7.2 opening
    set_para_text(
        find_para(
            doc,
            contains="Intended Marriage (Chapter II) and Other Forms (Chapter III) share a single notice-generation",
        ),
        "Intended Marriage (Chapter II) and Other Forms (Chapter III) share a single "
        "notice-generation workflow. After Marriage Registration the citizen answers "
        "Whether Marriage Taken place or Not?: No routes to Special Marriage "
        "(Intended Marriage Notice); Yes routes to Special Marriage (Other Forms Notice) "
        "and then Enter Marriage Details (date, place and form / rites of the already "
        "celebrated ceremony) before Prerequisite. Both paths then share Prerequisite, "
        "Aadhaar / party capture, Online or Offline publication, first payment, "
        "eSign / notice-board and the 30-day countdown.",
    )

    # §7.2.1 channel note
    set_para_text(
        find_para(
            doc,
            contains="Notice generation supports Online and Offline (In Person) channels for both Intended Marriage and Other Forms.",
        ),
        "Notice generation supports Online and Offline (In Person) channels for both "
        "Intended Marriage and Other Forms. The selected notice channel drives "
        "publication mode (portal vs notice board) and office tasks. Service path "
        "(Intended Marriage vs Other Forms) is decided by Whether Marriage Taken place "
        "or Not? immediately after Marriage Registration; Other Forms additionally "
        "captures marriage details before Prerequisite.",
    )

    # §7.2.2.1 common intake
    set_para_text(
        find_para(
            doc,
            contains="Identical in both notice diagrams (Citizens and System lanes). Intended Marriage and Other Forms share these steps",
        ),
        "Identical in both notice diagrams (Citizens and System lanes). Intended Marriage "
        "and Other Forms share these steps; the path decision and Other Forms "
        "Enter Marriage Details branch are in steps 5–6:",
    )
    set_para_text(
        find_para(
            doc,
            contains="Select channel / service path: Special Marriage Notice Online / Offline, and path Intended Marriage or Other Forms.",
        ),
        "Whether Marriage Taken place or Not? — No: Special Marriage (Intended Marriage "
        "Notice); Yes: Special Marriage (Other Forms Notice). Online vs Offline is the "
        "selected notice channel (separate diagrams 7.2.2.2 / 7.2.2.3). Marriage "
        "Registration (certificate) is initiated later by selecting an approved notice "
        "(see 7.3).",
    )
    path_para = find_para(
        doc,
        contains="Whether Marriage Taken place or Not? — No: Special Marriage (Intended Marriage Notice)",
    )
    insert_paragraph_after(
        path_para,
        "If Other Forms (Yes): Enter Marriage Details — date, place and form / rites of "
        "the already celebrated ceremony (not required for Intended Marriage). Both "
        "paths then continue to Prerequisite.",
        style="List Number",
    )
    set_para_text(
        find_para(
            doc,
            contains="Enter / capture Bride details and Bridegroom details — persisted to the notice application record. Conditional branch:",
        ),
        "If Aadhaar information available: e-KYC / Face Authentication on Bride and "
        "Bridegroom details; else Enter Bride and Bridegroom details — persisted to the "
        "notice application record. Marriage-details fields are not captured at this "
        "step (Other Forms already captured them in step 6).",
    )
    # After insert, resolve list items by walking from the 7.2.2.1 heading
    h221 = find_para(doc, exact="7.2.2.1 Common intake steps", heading_only=True)
    list_paras: list[Paragraph] = []
    saw = False
    for p in doc.paragraphs:
        if p._p is h221._p:
            saw = True
            continue
        if saw:
            if style_name(p).startswith("Heading"):
                break
            if style_name(p) == "List Number" and p.text.strip():
                list_paras.append(p)
    if len(list_paras) < 8:
        raise RuntimeError(f"Expected 8 common-intake list items, found {len(list_paras)}")
    set_para_text(
        list_paras[3],
        "Marriage Registration — citizen selects the marriage registration service.",
    )
    set_para_text(
        list_paras[6],
        "Read and continue with Prerequisite for marriage and complete declaration — "
        "single combined acknowledgement screen (eligibility, documents, channel "
        "implications and statutory declarations). SR rejection of the notice "
        "application returns here for correction and resubmission.",
    )
    set_para_text(
        find_para(
            doc,
            exact="Flow (continuing from 7.2.2.1 common intake — Online notice-specific steps; shared by Intended Marriage and Other Forms):",
        ),
        "Flow (continuing from 7.2.2.1 step 8 — Online notice-specific steps; shared by "
        "Intended Marriage and Other Forms):",
    )
    set_para_text(
        find_para(
            doc,
            exact="Flow (continuing from 7.2.2.1 common intake — Offline notice-specific steps; shared by Intended Marriage and Other Forms):",
        ),
        "Flow (continuing from 7.2.2.1 step 8 — Offline notice-specific steps; shared by "
        "Intended Marriage and Other Forms):",
    )

    # Online key characteristics + Offline key characteristics
    set_para_text(
        find_para(
            doc,
            contains="Key characteristics (both paths): e-KYC / Face Authentication on Bride & Bridegroom when Aadhaar available; document upload including individual photos",
        ),
        "Key characteristics (both paths): after the Whether Marriage Taken place "
        "branch (Other Forms: Enter Marriage Details before Prerequisite), e-KYC / Face "
        "Authentication on Bride & Bridegroom when Aadhaar available; document upload "
        "including individual photos; SR verification before First Payment; System "
        "generates notice after payment; citizen e-sign on the generated notice; then "
        "Marriage notice displayed in portal; no Online appointment or DEO; 30-day "
        "countdown from publication (Sec. 7 / Sec. 16).",
    )
    set_para_text(
        find_para(
            doc,
            contains="Key characteristics (both paths): Aadhaar YES (e-KYC / Face Authentication) / NO (manual) capture; SR verification → First Payment → appointment",
        ),
        "Key characteristics (both paths): after the Whether Marriage Taken place "
        "branch (Other Forms: Enter Marriage Details before Prerequisite), Aadhaar YES "
        "(e-KYC / Face Authentication) / NO (manual) capture; SR verification → First "
        "Payment → appointment; SR generates notice & selects DEO; individual photo "
        "capture; physical notice-board paste (Sec. 6(2)) with portal display; 30-day "
        "countdown. SR rejection returns to Prerequisite & declaration.",
    )

    # Step tables under 7.2.2.2 / 7.2.2.3
    fill_table(doc.tables[14], ONLINE_NOTICE_STEPS)
    fill_table(doc.tables[15], OFFLINE_NOTICE_STEPS)

    # Status model Table 16
    status = doc.tables[16]
    # Row 4 Details captured
    set_cell_text(
        status.rows[4].cells[1],
        "Bride / bridegroom particulars saved (Online/Offline: e-KYC / Face Authentication "
        "where Aadhaar available). Other Forms only: marriage details already saved at "
        "path selection (before Prerequisite)",
    )
    # Row 5 Notice application submitted — remove premature Online eSigned
    set_cell_text(
        status.rows[5].cells[1],
        "Special Marriage notice application submitted for SR scrutiny (documents and "
        "particulars complete; Online citizen eSign of the statutory notice occurs after "
        "Notice Generated — see Notice generated / Notice published)",
    )
    # Row 7 Rejected — next state
    set_cell_text(status.rows[7].cells[3], "Prerequisite & declaration completed")
    set_cell_text(
        status.rows[7].cells[1],
        "Returned to citizen for correction; flow returns to Prerequisite & declaration "
        "per notice diagrams",
    )
    # Row 10 Notice generated — next toward eSign then publish (Online)
    set_cell_text(
        status.rows[10].cells[1],
        "Statutory notice generated and entered in the Marriage Notice Book. Online: "
        "citizen eSign is performed on the generated notice before portal display",
    )
    set_cell_text(
        status.rows[10].cells[3],
        "Notice published (after Online e-sign where applicable)",
    )
    # Row 11 Notice published
    set_cell_text(
        status.rows[11].cells[1],
        "Published on portal after Online e-sign, or pasted on notice board after DEO "
        "upload (Offline); publication timestamp starts the countdown",
    )

    # FR-SMA-001
    fr001 = find_fr_row(find_table_containing(doc, "FR-SMA-001"), "FR-SMA-001")
    set_cell_text(
        fr001.cells[1],
        "System shall route Special Marriage notice via Whether Marriage Taken place or "
        "Not? after Marriage Registration: No → Special Marriage (Intended Marriage) "
        "Notice; Yes → Special Marriage Other Forms Notice (with Enter Marriage Details "
        "before Prerequisite). Each path shall support an Online or Offline notice channel",
    )
    set_cell_text(
        fr001.cells[3],
        "Decision node and path labels match 7.2.2.1 / Online & Offline notice diagrams; "
        "four channel×path combinations map to 7.2–7.3",
    )

    # FR-SMA-013
    fr013 = find_fr_row(find_table_containing(doc, "FR-SMA-013"), "FR-SMA-013")
    set_cell_text(
        fr013.cells[1],
        "Online channel: after System notice generation, both parties shall eSign the "
        "generated notice; the system shall then display the notice on the portal "
        "(publication)",
    )
    set_cell_text(
        fr013.cells[3],
        "eSign artefacts stored immutably with timestamp; eSign occurs after Notice "
        "Generated and before portal publication (7.2.2.2 diagram)",
    )

    # FR-SMA-061
    fr061 = find_fr_row(find_table_containing(doc, "FR-SMA-061"), "FR-SMA-061")
    set_cell_text(
        fr061.cells[1],
        "Special Marriage Other Forms notice only: immediately after Other Forms path "
        "selection and before Prerequisite, system shall capture marriage details — "
        "date of ceremony, place of ceremony (with sufficient particulars to locate), "
        "and form / rites of the already celebrated marriage. These fields shall not be "
        "required for Intended Marriage notice",
    )
    set_cell_text(
        fr061.cells[3],
        "Other Forms: marriage-details step shown before Prerequisite; blocked until "
        "complete; Intended Marriage: marriage-details section hidden / not mandatory "
        "(7.2.2.1 steps 5–6)",
    )

    # §8.2.2 narrative
    set_para_text(
        find_para(
            doc,
            exact="Marriage details (date of ceremony, place of ceremony, and form / rites of the already celebrated marriage) shall be captured only for Special Marriage Other Forms notice generation.",
        ),
        "Marriage details (date of ceremony, place of ceremony, and form / rites of the "
        "already celebrated marriage) shall be captured only for Special Marriage Other "
        "Forms notice generation — immediately after the Other Forms path is selected "
        "(Whether Marriage Taken place or Not? = Yes) and before Prerequisite "
        "(FR-SMA-061; 7.2.2.1).",
    )

    # §8.2.5 eSign timing
    set_para_text(
        find_para(
            doc,
            exact="This form is required to be e-Signed during the filing of notice itself",
        ),
        "Online: this form (the generated notice) is required to be e-Signed after "
        "System notice generation and before portal publication (FR-SMA-013; 7.2.2.2). "
        "Offline: physical sign / scan / upload is performed by FDA/SDA/DEO after "
        "SR generates the notice (7.2.2.3).",
    )

    doc.save(str(DST))
    print(f"Wrote {DST}")

    # Verification dump
    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    print("--- 7.2.2.1 steps ---")
    capture = False
    for p in doc2.paragraphs:
        t = p.text.strip()
        if t.startswith("7.2.2.1"):
            capture = True
            continue
        if t.startswith("7.2.2.2"):
            break
        if capture and t:
            print(" ", t[:140])
    print("--- Online table steps ---")
    for row in doc2.tables[14].rows:
        print(" ", row.cells[0].text.strip(), row.cells[1].text.strip()[:90])
    print("--- Offline table first/last ---")
    print(" ", doc2.tables[15].rows[1].cells[0].text.strip(), doc2.tables[15].rows[1].cells[1].text.strip()[:80])
    print(" ", doc2.tables[15].rows[-1].cells[0].text.strip(), doc2.tables[15].rows[-1].cells[1].text.strip()[:80])
    print("--- FR-SMA-013 ---")
    fr = find_fr_row(find_table_containing(doc2, "FR-SMA-013"), "FR-SMA-013")
    print(" ", fr.cells[1].text.strip()[:160])


if __name__ == "__main__":
    main()
