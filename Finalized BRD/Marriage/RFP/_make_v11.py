"""Create BRD_Hindu_Marriage_v1.1.docx from v1.0 — add Special Marriage functional
requirements (§8.19–8.27, FR-SMA series) and dependent rules / status / UI / RTM updates.
"""

from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Hindu_Marriage_v1.0.docx"
DST = BASE / "BRD_Hindu_Marriage_v1.1.docx"

REQ_HEADER = ["Req ID", "Requirement", "Priority", "Acceptance criteria"]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def _style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para_index(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> int:
    for i, p in enumerate(doc.paragraphs):
        if heading_only and not _style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return i
        if contains is not None and contains in t:
            return i
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def find_para(doc: Document, **kwargs) -> Paragraph:
    return doc.paragraphs[find_para_index(doc, **kwargs)]


def para_before(doc: Document, **kwargs) -> Paragraph:
    idx = find_para_index(doc, **kwargs)
    if idx == 0:
        raise KeyError("No paragraph before target")
    return doc.paragraphs[idx - 1]


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    pPr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not pPr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], doc: Document) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def find_table_by_header(doc: Document, first_cell: str) -> Table:
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == first_cell:
            return table
    raise KeyError(f"Table with header {first_cell!r} not found")


def append_table_row(table: Table, values: list[str]) -> None:
    row = table.add_row()
    for i, v in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def set_doc_control_field(doc: Document, field: str, value: str) -> None:
    for row in doc.tables[0].rows:
        if row.cells[0].text.strip() == field:
            row.cells[1].text = value
            return
    raise KeyError(field)


def add_fr_section(
    doc: Document,
    after: Paragraph,
    heading: str,
    ref_note: str,
    rows: list[list[str]],
) -> Paragraph:
    """Insert Heading 3 + ref note + requirement table; return trailing paragraph."""
    h = insert_paragraph_after(after, heading, style="Heading 3")
    note = insert_paragraph_after(h, ref_note, style="Normal")
    insert_table_after(note, [REQ_HEADER] + rows, doc)
    tbl = note._p.getnext()
    trailing = insert_paragraph_after(note, "", style="Normal")
    if tbl is not None and tbl.tag == qn("w:tbl"):
        trailing._p.getparent().remove(trailing._p)
        tbl.addnext(trailing._p)
    return trailing


# --------------------------------------------------------------------------------------
# Requirement content
# --------------------------------------------------------------------------------------

S819 = [
    [
        "FR-SMA-001",
        "System shall offer Special Marriage service paths at service selection: Special Marriage "
        "(Intended Marriage) Notice and Special Marriage Other Forms, each with an Online or Offline "
        "notice channel",
        "Must",
        "Four selectable paths mapped to §7.5–7.10",
    ],
    [
        "FR-SMA-002",
        "System shall enforce Sec. 4 conditions for Intended Marriage: neither party has a spouse "
        "living; capacity to give valid consent; bridegroom ≥ 21 and bride ≥ 18 years; parties not "
        "within degrees of prohibited relationship",
        "Must",
        "Hard stop with statutory reason displayed",
    ],
    [
        "FR-SMA-003",
        "System shall enforce Sec. 15 conditions for Other Forms: ceremony already performed and "
        "parties living together as husband and wife since; neither party has more than one spouse "
        "living; both parties ≥ 21 years at registration; not within prohibited degrees; residence in "
        "the district ≥ 30 days immediately preceding the application",
        "Must",
        "Hard stop; conditions captured as declarations",
    ],
    [
        "FR-SMA-004",
        "After path and channel selection, system shall display a combined prerequisite and "
        "declaration screen specific to the selected Special Marriage path before data capture",
        "Must",
        "Acknowledgement timestamped and audited",
    ],
    [
        "FR-SMA-005",
        "Selected notice channel (Online / Offline) shall drive all subsequent screens, publication "
        "mode and office tasks",
        "Must",
        "Channel stored on the application; single channel per notice",
    ],
    [
        "FR-SMA-006",
        "System shall validate degrees of prohibited relationship against the First Schedule, "
        "allowing the custom/usage exception for Other Forms marriages celebrated before "
        "commencement of the Act",
        "Should",
        "Exception requires recorded justification",
    ],
]

S820 = [
    [
        "FR-SMA-007",
        "System shall capture notice particulars for both parties as per the Second Schedule: name, "
        "condition (unmarried / widow / widower / divorcee), occupation, age, dwelling place, "
        "permanent dwelling place if different, and length of residence",
        "Must",
        "All Second Schedule fields present on generated notice",
    ],
    [
        "FR-SMA-008",
        "System shall route the notice to the Marriage Officer of the district in which at least one "
        "party has resided for ≥ 30 days immediately preceding the notice date, validated against "
        "the address proof",
        "Must",
        "Jurisdiction derived and shown before submission",
    ],
    [
        "FR-SMA-009",
        "Online channel: system shall perform e-KYC / Face Authentication on bride and bridegroom "
        "details where Aadhaar information is available",
        "Must",
        "e-KYC reference stored against each party",
    ],
    [
        "FR-SMA-010",
        "Where Aadhaar information is unavailable, system shall allow manual capture of bride and "
        "bridegroom details with mandatory documentary proof",
        "Must",
        "Manual path flagged for SR scrutiny",
    ],
    [
        "FR-SMA-011",
        "System shall accept upload of identity proof, photograph, age proof and address proof for "
        "bridegroom and bride",
        "Must",
        "Submission blocked until all mandatory documents uploaded",
    ],
    [
        "FR-SMA-012",
        "System shall present a review summary of captured particulars before document upload and "
        "submission",
        "Must",
        "Citizen can return and edit any section from summary",
    ],
    [
        "FR-SMA-013",
        "Online channel: both parties shall eSign the notice application before it moves to SR "
        "verification",
        "Must",
        "eSign artefacts stored immutably with timestamp",
    ],
]

S821 = [
    [
        "FR-SMA-014",
        "Marriage Officer (Sub-Registrar) shall generate the statutory Notice of Intended Marriage in "
        "the Second Schedule format from the verified application",
        "Must",
        "Notice carries unique notice number and generation date",
    ],
    [
        "FR-SMA-015",
        "System shall enter a true copy of every notice in the Marriage Notice Book with running "
        "serial number and date of entry",
        "Must",
        "Sec. 6(1) — entry created automatically on generation",
    ],
    [
        "FR-SMA-016",
        "Marriage Notice Book shall be available for inspection without fee during office hours",
        "Should",
        "Read-only search by notice number, name and date",
    ],
    [
        "FR-SMA-017",
        "Online channel: system shall publish the generated notice on the portal and record the "
        "publication date and time",
        "Must",
        "Publication timestamp drives the countdown",
    ],
    [
        "FR-SMA-018",
        "Offline channel: SR shall assign an FDA / SDA / DEO who captures individual photographs of "
        "bride and bridegroom, downloads, prints, signs, scans and uploads the notice, and records "
        "pasting on the designated notice board",
        "Must",
        "Notice-board pasting recorded with date, office and officer ID",
    ],
    [
        "FR-SMA-019",
        "System shall start a 30-day countdown from the publication date and display days remaining "
        "to citizen and office users",
        "Must",
        "Countdown visible on application and notice record",
    ],
    [
        "FR-SMA-020",
        "Where a party is not permanently residing within the district of the notice, system shall "
        "transmit a copy of the notice to the Marriage Officer of the district of permanent residence "
        "for publication",
        "Must",
        "Sec. 6(3) — transmission logged and acknowledged",
    ],
    [
        "FR-SMA-021",
        "System shall issue an acknowledgement to the citizen with notice number, publication date "
        "and registration validity window",
        "Should",
        "Downloadable acknowledgement; SMS / email sent",
    ],
]

S822 = [
    [
        "FR-SMA-022",
        "System shall allow any person to file an objection before expiry of 30 days from notice "
        "publication, on the ground that the marriage would contravene a condition in Sec. 4",
        "Must",
        "Objection window closes automatically after day 30",
    ],
    [
        "FR-SMA-023",
        "System shall record the nature of the objection in writing in the Marriage Notice Book, and "
        "capture that it was read over and explained to and signed by the objector",
        "Must",
        "Sec. 7(3) — signed record attached to the notice",
    ],
    [
        "FR-SMA-024",
        "System shall block solemnization and registration while an objection is pending decision",
        "Must",
        "Hard gate on the registration flow",
    ],
    [
        "FR-SMA-025",
        "System shall provide an SR enquiry workflow to summon parties, record evidence and decide "
        "the objection within 30 days of the date of objection",
        "Must",
        "Sec. 8(1) — breach of 30 days raises an escalation alert",
    ],
    [
        "FR-SMA-026",
        "Where the objection is upheld, system shall record the objection reason, tag the notice as "
        "Objected and remove it from portal display",
        "Must",
        "Notice no longer selectable for registration",
    ],
    [
        "FR-SMA-027",
        "Where the objection is found invalid, system shall resume the registration flow at the "
        "second payment step",
        "Must",
        "Decision and reason recorded against the notice",
    ],
    [
        "FR-SMA-028",
        "On refusal to solemnize, system shall notify the parties of the right of appeal to the "
        "district court within 30 days and record the appeal outcome for compliance",
        "Must",
        "Sec. 8(2) — appeal decision recorded and acted upon",
    ],
]

S823 = [
    [
        "FR-SMA-029",
        "Citizen shall initiate marriage registration by selecting a published notice; system shall "
        "validate that the elapsed period is ≥ 30 days and ≤ 90 days from publication",
        "Must",
        "Per §7.7 / §7.10 timeline decision",
    ],
    [
        "FR-SMA-030",
        "Where the notice is outside the validity window, system shall allow no further action and "
        "display the statutory reason",
        "Must",
        "Notice tagged Expired after the window",
    ],
    [
        "FR-SMA-031",
        "Where the marriage is not solemnized within three months of the notice date, system shall "
        "require a fresh notice before the marriage may be solemnized",
        "Must",
        "Sec. 14 — new notice restarts publication and countdown",
    ],
    [
        "FR-SMA-032",
        "System shall alert parties before expiry of the notice validity window",
        "Should",
        "Reminder configurable (e.g. day 75 and day 85)",
    ],
]

S824 = [
    [
        "FR-SMA-033",
        "System shall collect the second payment (registration / solemnization fee) after the "
        "objection check is cleared and before visit scheduling",
        "Must",
        "Receipt generated; payment gate on scheduling",
    ],
    [
        "FR-SMA-034",
        "System shall perform e-KYC on witness details for three witnesses",
        "Must",
        "Three witnesses mandatory before scheduling",
    ],
    [
        "FR-SMA-035",
        "System shall schedule the solemnization visit with the Sub-Registrar and notify all parties",
        "Must",
        "Slot, date, time and office recorded",
    ],
    [
        "FR-SMA-036",
        "SR shall verify the application before solemnization; rejection shall return the application "
        "to the visit scheduling step with a recorded reason",
        "Must",
        "Approve / Reject with reason; no return to notice stage",
    ],
    [
        "FR-SMA-037",
        "System shall capture the declarations of the parties and the three witnesses in the Third "
        "Schedule form, signed in the presence of the Marriage Officer",
        "Must",
        "Sec. 11 — declarations countersigned by Marriage Officer",
    ],
    [
        "FR-SMA-038",
        "System shall record the place and form of solemnization, including solemnization at a place "
        "other than the Marriage Officer's office within reasonable distance on payment of the "
        "additional prescribed fee",
        "Must",
        "Sec. 12 — place recorded on certificate",
    ],
    [
        "FR-SMA-039",
        "DEO shall capture the joint photograph, generate the Certificate of Marriage in the Fourth "
        "Schedule, capture signatures of bride, bridegroom and three witnesses, and upload the signed "
        "copy",
        "Must",
        "Signed copy required before SR digital signature",
    ],
    [
        "FR-SMA-040",
        "System shall enter the certificate in the Marriage Certificate Book and apply the SR digital "
        "signature (DSC) before issue",
        "Must",
        "Sec. 13 — signature verifiable on the PDF",
    ],
    [
        "FR-SMA-041",
        "System shall issue the digitally signed certificate for download by the citizen as "
        "conclusive evidence of the marriage",
        "Must",
        "Certificate downloadable from portal; SMS / email sent",
    ],
]

S825 = [
    [
        "FR-SMA-042",
        "System shall accept an application for registration of a marriage celebrated in other forms, "
        "signed by both parties to the marriage",
        "Must",
        "Sec. 16 — both signatures / eSign mandatory",
    ],
    [
        "FR-SMA-043",
        "System shall capture the ceremony particulars and the declaration that the parties have been "
        "living together as husband and wife since the ceremony",
        "Must",
        "Declaration recorded with date of ceremony",
    ],
    [
        "FR-SMA-044",
        "System shall give public notice of the application in the prescribed manner (portal and / or "
        "notice board per selected channel) and allow a period of 30 days for objections",
        "Must",
        "Publication and countdown recorded as in §8.21",
    ],
    [
        "FR-SMA-045",
        "SR shall hear any objection received within the 30-day period before deciding the "
        "application",
        "Must",
        "Objection workflow of §8.22 applies",
    ],
    [
        "FR-SMA-046",
        "On being satisfied that all Sec. 15 conditions are fulfilled, system shall enter the "
        "certificate of marriage in the Marriage Certificate Book in the Fifth Schedule format, "
        "signed by the parties and three witnesses",
        "Must",
        "Fifth Schedule used for Other Forms, not Fourth",
    ],
    [
        "FR-SMA-047",
        "Where registration is refused, system shall record a reasoned order and notify the applicant "
        "of the right of appeal to the district court within 30 days",
        "Must",
        "Sec. 17 — appeal outcome recorded and complied with",
    ],
    [
        "FR-SMA-048",
        "On final entry of the certificate, system shall record the effect of registration and issue "
        "the digitally signed certificate to the parties",
        "Must",
        "Sec. 18 — deemed married under the Act from date of registration",
    ],
]

S826 = [
    [
        "FR-SMA-049",
        "System shall collect the first payment (notice fee) after SR verification approval, as per "
        "the Special Marriage (Karnataka) Rules fee schedule",
        "Must",
        "No payment before SR approval",
    ],
    [
        "FR-SMA-050",
        "System shall collect the second payment (registration / solemnization fee) before the visit "
        "is scheduled",
        "Must",
        "Scheduling blocked until payment success",
    ],
    [
        "FR-SMA-051",
        "System shall apply the additional prescribed fee where solemnization is at a place other "
        "than the Marriage Officer's office",
        "Should",
        "Fee auto-added when out-of-office venue selected",
    ],
    [
        "FR-SMA-052",
        "System shall generate a receipt with challan / transaction reference and handle payment "
        "failure, retry and refund",
        "Must",
        "Failed payments do not advance the application",
    ],
    [
        "FR-SMA-053",
        "Special Marriage fee heads shall be configurable in the fee master without code change",
        "Must",
        "Effective-dated fee versions retained",
    ],
]

S827 = [
    [
        "FR-SMA-054",
        "System shall send SMS / email on notice submission, SR approval or rejection, notice "
        "publication, objection filed, objection decision, appointment, and certificate issue",
        "Should",
        "Bilingual EN / KN templates",
    ],
    [
        "FR-SMA-055",
        "System shall provide a notice register report per office showing published, expired, "
        "objected and converted notices",
        "Must",
        "Filterable by office and period",
    ],
    [
        "FR-SMA-056",
        "System shall provide an objection register report with enquiry status and decision ageing "
        "against the 30-day statutory limit",
        "Must",
        "Overdue enquiries highlighted",
    ],
    [
        "FR-SMA-057",
        "System shall provide Marriage Notice Book and Marriage Certificate Book extracts separately "
        "for Chapter II and Chapter III registrations",
        "Must",
        "Extracts printable and exportable",
    ],
    [
        "FR-SMA-058",
        "System shall provide MIS counts of notices, objections, solemnizations and registrations by "
        "service path, channel, office and period",
        "Must",
        "Drill-down to application level",
    ],
    [
        "FR-SMA-059",
        "Every Special Marriage state transition shall raise an audit event capturing actor, role, "
        "timestamp and reason where applicable",
        "Must",
        "Audit trail immutable and reportable",
    ],
    [
        "FR-SMA-060",
        "System shall support transmission of copies of entries in marriage records to the prescribed "
        "authority",
        "Should",
        "Sec. 48 — periodic transmission with acknowledgement",
    ],
]

SECTIONS = [
    (
        "8.19 Special Marriage — service selection and eligibility",
        "(Ref: §7.5–7.10 intake; SMA 1954 Sec. 4 and Sec. 15)",
        S819,
    ),
    (
        "8.20 Notice of intended marriage — data capture and jurisdiction",
        "(Ref: §7.5 steps 8–11 and §7.6 steps 8–10; SMA 1954 Sec. 5, Second Schedule)",
        S820,
    ),
    (
        "8.21 Notice generation, Marriage Notice Book and publication",
        "(Ref: §7.5 steps 15–18 and §7.6 steps 14–19; SMA 1954 Sec. 5–6)",
        S821,
    ),
    (
        "8.22 Objection management and enquiry",
        "(Ref: §7.7 / §7.10 objection branch; SMA 1954 Sec. 7–9)",
        S822,
    ),
    (
        "8.23 Notice validity, timeline gate and fresh notice",
        "(Ref: §7.7 / §7.10 timeline decision; SMA 1954 Sec. 14)",
        S823,
    ),
    (
        "8.24 Solemnization, declarations and certificate (Intended Marriage)",
        "(Ref: §7.7 steps 5–16; SMA 1954 Sec. 11–13, Third and Fourth Schedules)",
        S824,
    ),
    (
        "8.25 Special Marriage Other Forms — application and registration",
        "(Ref: §7.8–7.10; SMA 1954 Chapter III, Sec. 15–18, Fifth Schedule)",
        S825,
    ),
    (
        "8.26 Special Marriage fees and payments",
        "(Ref: §7.5–7.10 first and second payment; Special Marriage (Karnataka) Rules, 1961)",
        S826,
    ),
    (
        "8.27 Special Marriage notifications, reports and audit",
        "(Ref: §7.5–7.10; SMA 1954 Sec. 6, 47–48)",
        S827,
    ),
]

TOC_ENTRIES = [
    "8.19 Special Marriage — service selection and eligibility (FR-SMA-001–006)",
    "8.20 Notice of intended marriage — data capture and jurisdiction (FR-SMA-007–013)",
    "8.21 Notice generation, Marriage Notice Book and publication (FR-SMA-014–021)",
    "8.22 Objection management and enquiry (FR-SMA-022–028)",
    "8.23 Notice validity, timeline gate and fresh notice (FR-SMA-029–032)",
    "8.24 Solemnization, declarations and certificate (FR-SMA-033–041)",
    "8.25 Special Marriage Other Forms — application and registration (FR-SMA-042–048)",
    "8.26 Special Marriage fees and payments (FR-SMA-049–053)",
    "8.27 Special Marriage notifications, reports and audit (FR-SMA-054–060)",
]

BR_SMA = [
    [
        "BR-SMA-001",
        "Notice may be published only after SR verification approval and first payment",
        "Process diagrams §7.5 / §7.6",
        "Publication gate; FR-SMA-014 / FR-SMA-049",
    ],
    [
        "BR-SMA-002",
        "Marriage may not be solemnized before expiry of 30 days from notice publication",
        "Sec. 7(2)",
        "Hard gate; FR-SMA-019 / FR-SMA-029",
    ],
    [
        "BR-SMA-003",
        "Registration may proceed only when the notice is ≥ 30 and ≤ 90 days old",
        "Sec. 14 + process diagram",
        "Timeline validation; FR-SMA-029 / FR-SMA-030",
    ],
    [
        "BR-SMA-004",
        "A pending objection blocks solemnization and registration",
        "Sec. 8(1)",
        "Hard stop; FR-SMA-024",
    ],
    [
        "BR-SMA-005",
        "Objection enquiry must be decided within 30 days of the objection",
        "Sec. 8(1)",
        "Ageing alert; FR-SMA-025",
    ],
    [
        "BR-SMA-006",
        "An upheld objection removes the notice from publication and closes the application",
        "Sec. 8",
        "Status Objected; FR-SMA-026",
    ],
    [
        "BR-SMA-007",
        "Notice jurisdiction requires ≥ 30 days residence of at least one party in the district",
        "Sec. 5",
        "Jurisdiction validation; FR-SMA-008",
    ],
    [
        "BR-SMA-008",
        "Other Forms: both parties must have completed 21 years at the time of registration",
        "Sec. 15(d)",
        "Validation; FR-SMA-003",
    ],
    [
        "BR-SMA-009",
        "Other Forms: the registration application must be signed by both parties",
        "Sec. 16",
        "Hard stop; FR-SMA-042",
    ],
    [
        "BR-SMA-010",
        "Declarations of both parties and three witnesses are required before solemnization",
        "Sec. 11",
        "Hard stop; FR-SMA-034 / FR-SMA-037",
    ],
    [
        "BR-SMA-011",
        "Certificate must be signed by the parties and three witnesses before SR digital signature",
        "Sec. 13 / Sec. 16",
        "Hard stop; FR-SMA-039 / FR-SMA-046",
    ],
    [
        "BR-SMA-012",
        "Certificate is issued only after the SR digital signature",
        "Process diagrams §7.7 / §7.10",
        "DSC pre-condition; FR-SMA-040",
    ],
    [
        "BR-SMA-013",
        "Fresh notice is required where the marriage is not solemnized within three months",
        "Sec. 14",
        "Restart notice flow; FR-SMA-031",
    ],
    [
        "BR-SMA-014",
        "Neither party may have a spouse living at the time of notice or registration",
        "Sec. 4(a) / Sec. 15(b)",
        "Declaration and validation; FR-SMA-002 / FR-SMA-003",
    ],
]

STATUS_ROWS = [
    [
        "Notice application submitted",
        "Special Marriage notice application submitted (Online: eSigned)",
        "SM Online / Offline",
        "Citizen",
        "Pending SR verification — notice",
    ],
    [
        "Pending SR verification — notice",
        "Awaiting Marriage Officer scrutiny of the notice application",
        "SM Online / Offline",
        "SR",
        "Notice approved / Rejected — notice data",
    ],
    [
        "Rejected — notice data",
        "Returned to citizen for correction of party details or documents",
        "SM Online / Offline",
        "SR",
        "Notice application submitted",
    ],
    [
        "Notice approved",
        "SR approved; first payment payable",
        "SM Online / Offline",
        "SR",
        "First payment completed",
    ],
    [
        "First payment completed",
        "Notice fee paid, receipt issued",
        "SM Online / Offline",
        "System",
        "Notice generated",
    ],
    [
        "Notice generated",
        "Statutory notice generated and entered in the Marriage Notice Book",
        "SM Online / Offline",
        "SR",
        "Notice published",
    ],
    [
        "Notice published",
        "Published on portal (Online) or pasted on notice board after DEO upload (Offline)",
        "SM Online / Offline",
        "System / DEO",
        "Objection period running",
    ],
    [
        "Objection period running",
        "30-day statutory countdown in progress",
        "SM Online / Offline",
        "System",
        "Notice valid for registration / Objection filed",
    ],
    [
        "Objection filed",
        "Objection recorded in the Marriage Notice Book",
        "SM Online / Offline",
        "Public / SR",
        "Under objection enquiry",
    ],
    [
        "Under objection enquiry",
        "SR enquiry by summoning parties (decide within 30 days)",
        "SM Online / Offline",
        "SR",
        "Objected — closed / Notice valid for registration",
    ],
    [
        "Objected — closed",
        "Objection upheld; notice removed from portal and tagged Objected",
        "SM Online / Offline",
        "SR",
        "Closed",
    ],
    [
        "Notice valid for registration",
        "Notice is ≥ 30 and ≤ 90 days old with no valid objection",
        "SM Online / Offline",
        "Citizen",
        "Second payment completed",
    ],
    [
        "Notice expired",
        "Validity window lapsed; fresh notice required",
        "SM Online / Offline",
        "System",
        "Closed",
    ],
    [
        "Second payment completed",
        "Registration / solemnization fee paid",
        "SM Registration",
        "System",
        "Visit scheduled",
    ],
    [
        "Visit scheduled",
        "Solemnization / registration visit booked; witness e-KYC done",
        "SM Registration",
        "Citizen",
        "Pending SR verification — visit",
    ],
    [
        "Pending SR verification — visit",
        "SR verification before solemnization",
        "SM Registration",
        "SR",
        "Solemnized / Rejected — visit",
    ],
    [
        "Rejected — visit",
        "Returned to visit scheduling with reason",
        "SM Registration",
        "SR",
        "Visit scheduled",
    ],
    [
        "Solemnized",
        "Marriage solemnized (Chapter II) or conditions satisfied for registration (Chapter III)",
        "SM Registration",
        "SR",
        "Allocated to DEO — certificate",
    ],
    [
        "Allocated to DEO — certificate",
        "Joint photo, certificate generation and signature capture",
        "SM Registration",
        "DEO",
        "Signed certificate uploaded",
    ],
    [
        "Signed certificate uploaded",
        "Signed certificate uploaded for SR digital signature",
        "SM Registration",
        "DEO",
        "Pending SR digital signature",
    ],
]

UI_ROWS = [
    [
        "Special Marriage path selection",
        "Choose Intended Marriage Notice or Other Forms, and Online / Offline notice channel",
        "SM Online / Offline",
        "Sec. 4 / Sec. 15",
        "§8.19; FR-SMA-001 / FR-SMA-005",
    ],
    [
        "Notice particulars (bride / bridegroom)",
        "Second Schedule particulars including condition, occupation and length of residence",
        "SM Online / Offline",
        "Sec. 5, Second Schedule",
        "§8.20; FR-SMA-007 / FR-SMA-008",
    ],
    [
        "Notice document upload and eSign",
        "Identity, photo, age and address proofs; eSign on Online path",
        "SM Online / Offline",
        "Sec. 5",
        "§8.20; FR-SMA-011 / FR-SMA-013",
    ],
    [
        "Notice publication and countdown",
        "Portal notice display and 30-day countdown; notice-board record for Offline",
        "SM Online / Offline",
        "Sec. 6–7",
        "§8.21; FR-SMA-017 / FR-SMA-018 / FR-SMA-019",
    ],
    [
        "Objection filing and enquiry",
        "Public objection intake and SR enquiry / decision workspace",
        "SM Online / Offline",
        "Sec. 7–9",
        "§8.22; FR-SMA-022 / FR-SMA-025",
    ],
    [
        "Notice selection for registration",
        "Select published notice with ≥ 30 and ≤ 90 day validation",
        "SM Registration",
        "Sec. 14",
        "§8.23; FR-SMA-029 / FR-SMA-030",
    ],
    [
        "Witness e-KYC and visit scheduling",
        "Three-witness e-KYC and solemnization visit booking after second payment",
        "SM Registration",
        "Sec. 11",
        "§8.24; FR-SMA-034 / FR-SMA-035",
    ],
    [
        "Declarations and solemnization",
        "Third Schedule declarations, place and form of solemnization",
        "SM Registration",
        "Sec. 11–12",
        "§8.24; FR-SMA-037 / FR-SMA-038",
    ],
    [
        "Certificate generation and issue",
        "Fourth Schedule (Chapter II) or Fifth Schedule (Chapter III) certificate, signatures and DSC",
        "SM Registration",
        "Sec. 13 / Sec. 16",
        "§8.24 / §8.25; FR-SMA-039 / FR-SMA-046",
    ],
]

INTEGRATION_ROWS = [
    [
        "Notice publication module (portal)",
        "Outbound",
        "Publish Special Marriage notices and run the 30-day countdown",
        "SM Online / Offline",
        "",
        "TBD",
    ],
    [
        "Objection intake service",
        "Inbound",
        "Public objection submission against a published notice",
        "SM Online / Offline",
        "",
        "TBD",
    ],
    [
        "Inter-office notice transmission",
        "Outbound",
        "Transmit notice copy to Marriage Officer of permanent residence district (Sec. 6(3))",
        "SM Online / Offline",
        "",
        "TBD",
    ],
    [
        "Payment gateway / Treasury (Special Marriage)",
        "Outbound",
        "First payment (notice) and second payment (registration / solemnization)",
        "SM Online / Offline",
        "",
        "TBD",
    ],
]

RTM_ROWS = [
    [
        "FR-SMA-001",
        "Sec. 4, 15 / SMA",
        "Offer Intended Marriage and Other Forms service paths with Online / Offline notice channel",
        "§8.19",
        "Special Marriage path selection",
        "TC-SMA-___",
        "Draft",
    ],
    [
        "FR-SMA-008",
        "Sec. 5 / SMA",
        "Route notice to Marriage Officer of district with ≥ 30 days residence",
        "§8.20",
        "Notice particulars",
        "TC-SMA-___",
        "Draft",
    ],
    [
        "FR-SMA-019",
        "Sec. 6–7 / SMA",
        "Start and display 30-day objection countdown from publication",
        "§8.21",
        "Notice publication and countdown",
        "TC-SMA-___",
        "Draft",
    ],
    [
        "FR-SMA-029",
        "Sec. 14 / SMA",
        "Validate notice age ≥ 30 and ≤ 90 days before registration",
        "§8.23",
        "Notice selection for registration",
        "TC-SMA-___",
        "Draft",
    ],
    [
        "FR-SMA-046",
        "Sec. 15–16 / SMA",
        "Enter Fifth Schedule certificate signed by parties and three witnesses",
        "§8.25",
        "Certificate generation and issue",
        "TC-SMA-___",
        "Draft",
    ],
]


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # --- Document control & version history ---
    set_doc_control_field(doc, "Version", "1.1")
    set_doc_control_field(doc, "Last updated", "2026-08-24")
    append_table_row(
        doc.tables[1],
        [
            "1.1",
            "2026-08-24",
            "Nandha Kumar",
            "Added Special Marriage functional requirements §8.19–8.27 (FR-SMA-001–060), business "
            "rules BR-SMA-001–014, Special Marriage application statuses, UI screens, integrations, "
            "data entities and RTM rows",
            "Prashanth",
        ],
    )

    # --- Contents entries ---
    anchor = find_para(
        doc, exact="8.18 Digital signature and certificate issuance (FR-HMA-078–082)"
    )
    for entry in TOC_ENTRIES:
        anchor = insert_paragraph_after(anchor, entry, style="Normal")

    for p in doc.paragraphs[:120]:
        if p.text.strip() == "9. Business rules (BR-HMA-001–020)":
            set_para_text(p, "9. Business rules (BR-HMA-001–020, BR-SMA-001–014)")
            break

    # --- §8.19–8.27 functional requirement sections ---
    cursor = para_before(doc, exact="9. Business rules", heading_only=True)
    for heading, ref_note, rows in SECTIONS:
        cursor = add_fr_section(doc, cursor, heading, ref_note, rows)

    # --- §9 business rules ---
    rules = find_table_by_header(doc, "Rule ID")
    for row in BR_SMA:
        append_table_row(rules, row)

    # --- §7.11 status model ---
    status = find_table_by_header(doc, "Status")
    for row in STATUS_ROWS:
        append_table_row(status, row)

    # --- §10 UI ---
    ui = find_table_by_header(doc, "Screen / step")
    for row in UI_ROWS:
        append_table_row(ui, row)

    # --- §11 Integrations ---
    integrations = find_table_by_header(doc, "Integration")
    for row in INTEGRATION_ROWS:
        append_table_row(integrations, row)

    # --- §12.1 entities ---
    for p in doc.paragraphs:
        if p.text.strip().startswith("Application/Memorandum, Party (Bride/Bridegroom)"):
            set_para_text(
                p,
                p.text.rstrip()
                + " Special Marriage entities: MarriageNotice (Second Schedule particulars), "
                "MarriageNoticeBookEntry, NoticePublication (portal / notice board), Objection and "
                "ObjectionEnquiry, NoticeValidity, Declaration (Third Schedule), "
                "MarriageCertificateBookEntry (Fourth / Fifth Schedule).",
            )
            break

    # --- §13 RTM ---
    rtm = find_table_by_header(doc, "Req ID")
    for table in doc.tables:
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:2] == ["Req ID", "Act/Rule/Form"]:
            rtm = table
            break
    for row in RTM_ROWS:
        append_table_row(rtm, row)

    # --- §14 UAT scope ---
    for p in doc.paragraphs:
        if p.text.strip().startswith("UAT scope: Test scenarios derived from FR-HMA-*"):
            set_para_text(
                p,
                "UAT scope: Test scenarios derived from FR-HMA-* and FR-SMA-* (see 13 RTM), "
                "BR-HMA-* / BR-SMA-*, statutory forms Form I / IA / II / II-A, and the Special "
                "Marriage Second, Third, Fourth and Fifth Schedules.",
            )
            break

    doc.save(str(DST))
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
