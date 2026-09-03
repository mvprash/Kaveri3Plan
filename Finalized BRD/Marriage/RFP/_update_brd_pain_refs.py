# -*- coding: utf-8 -*-
"""Add BRD (§7 To-Be) references to pain-point tables in BRD_Marriage_BRD_v1.22.docx.

§6.1 and §7.5.1 currently point mainly at FRS/NFR sections (§8+). After the split,
those live in FRS_and_NFRS_Marriage_v1.22.docx — this patch prepends in-document
BRD process refs (§7.x) and labels companion FRS/NFR refs clearly.
"""
from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
DST = BASE / "BRD_Marriage_BRD_v1.22.docx"

COL_HEADER = "Addressed in (BRD / FRS·NFR)"

# sr → full addressed-in text (BRD first, then FRS/NFRS companion)
ADDRESSED = {
    "1": (
        "BRD §7.5 (UI/UX); §7.1.1 / §7.2.1 Online channels; "
        "FRS/NFRS: §10 UI; §15.4 NFR-MRG-VAPT-002 (mobile-responsive interfaces)"
    ),
    "2": (
        "BRD §7.1.2.1 common intake (party address); §7.1.2.2 / §7.1.2.3 certificate issuance; "
        "FRS/NFRS: §8.1.3–8.1.5; FR-HMA-008, FR-HMA-010/011; §8.1.16 FR-HMA-030/080/081"
    ),
    "3": (
        "BRD §7.1.2.1 common intake (jurisdiction / area selection); "
        "FRS/NFRS: §8.1.2; FR-HMA-005, FR-HMA-008; §11 Integrations (MDM / address master)"
    ),
    "4": (
        "BRD §7.1.2.1 common intake (supporting documents); "
        "FRS/NFRS: §8.1.9; FR-HMA-065, FR-HMA-018/019"
    ),
    "5": (
        "BRD §7.1.2.1 common intake (witnesses); "
        "FRS/NFRS: §8.1.6; FR-HMA-012–014; BR-HMA-001"
    ),
    "6": (
        "BRD §7.1.2.4 Application Status Model; "
        "FRS/NFRS: §8.1.15; FR-HMA-073/074; BR-HMA-014, BR-HMA-017"
    ),
    "7": (
        "BRD §7.1.2.2 Online / §7.1.2.3 Offline (SRO scrutiny); "
        "FRS/NFRS: §8.1.11; FR-HMA-026"
    ),
    "8": (
        "BRD §7.1.2 process (edit until payment / eSign gate); "
        "FRS/NFRS: §8.1.13; FR-HMA-052; BR-HMA-010"
    ),
    "9": (
        "BRD §7.1.2.4 Application Status Model (officer queue); "
        "FRS/NFRS: §8.1.11; FR-HMA-077; §8.6 FR-HMA-042 (cycle-time MIS)"
    ),
    "10": (
        "BRD §7.1.2.1 common intake (document upload validation); "
        "FRS/NFRS: §17 FB-MRG-003; §8.1.9 FR-HMA-065"
    ),
    "11": (
        "BRD §7.1.2 / §7.2.2 process; §7.5 (payments and notifications); "
        "FRS/NFRS: §8.5; FR-HMA-036–038; FR-SMA-054; FB-MRG-004"
    ),
    "12": (
        "BRD §7.5 What is new in Kaveri 3.0; "
        "FRS/NFRS: §8.6; FR-HMA-041–045; FR-SMA-055–058"
    ),
    "13": (
        "BRD §7.1.2.3 Offline (SR allocates / reassigns DEO); "
        "FRS/NFRS: §8.1.14; FR-HMA-069, FR-HMA-088; §16 RS-MRG-003"
    ),
    "14": (
        "BRD §7.2.2.4 / §7.3.2.3 Application Status Model; "
        "FRS/NFRS: §8.2.8; FR-SMA-019–032, FR-SMA-024–026"
    ),
    "15": (
        "BRD §7.1.2.2 / §7.2.2.2 Online payment; §7.5 (payments); "
        "FRS/NFRS: §8.1.10; §17 FB-MRG-001; NFR-MRG-PAY-001; FR-HMA-025, FR-SMA-052"
    ),
    "16": (
        "BRD §7.1.2 / §7.2.2.2 process (summary → next step); "
        "FRS/NFRS: §8.1.13; FR-HMA-051, FR-HMA-052; FR-SMA-012; FR-HMA-038"
    ),
    "17": (
        "BRD §7.1.2.1 / §7.2.2.1 party particulars; "
        "FRS/NFRS: §8.1.4–8.1.5; §8.2.3; FR-HMA-058, FR-HMA-089; FR-SMA-009/062/063/066"
    ),
    "18": (
        "BRD §7.1.2.1 / §7.2.2.1 marriage details (place of marriage); "
        "FRS/NFRS: §8.1.3; §8.2.2; FR-HMA-017; FR-HMA-051"
    ),
    "19": (
        "BRD §7.5 What is new in Kaveri 3.0 (post-registration search); "
        "FRS/NFRS: §8.4; FR-HMA-034; §12.1 Core entities"
    ),
    "20": (
        "BRD §7.1.2.2 Online DSC / certificate; §7.3.2.2 solemnization & certificate; "
        "FRS/NFRS: §8.1.16; §8.3.3; FR-HMA-054/080; FR-SMA-040/048; BR-HMA-001, BR-SMA-011"
    ),
    "21": (
        "BRD §7.2.2 notice generation / publication; §7.2.2.4 status model; "
        "FRS/NFRS: §8.2.5–8.2.7; FR-SMA-014/016/021; FR-SMA-055"
    ),
}

# Shorter companion column for §7.5.1 (same BRD anchors; condensed FRS/NFRS)
RECTIFIED = {
    "1": "BRD §7.5; §7.1.1 / §7.2.1; FRS/NFRS: §10 UI; §15.4 NFR-MRG-VAPT-002",
    "2": "BRD §7.1.2.1; §7.1.2.2 / §7.1.2.3; FRS/NFRS: §8.1.3–8.1.5; FR-HMA-008, FR-HMA-010/011; §8.1.16",
    "3": "BRD §7.1.2.1; FRS/NFRS: §8.1.2; FR-HMA-005, FR-HMA-008; §11 Integrations",
    "4": "BRD §7.1.2.1; FRS/NFRS: §8.1.9; FR-HMA-065, FR-HMA-018/019",
    "5": "BRD §7.1.2.1; FRS/NFRS: §8.1.6; FR-HMA-012–014; BR-HMA-001",
    "6": "BRD §7.1.2.4; FRS/NFRS: §8.1.15; FR-HMA-073/074; BR-HMA-014, BR-HMA-017",
    "7": "BRD §7.1.2.2 / §7.1.2.3; FRS/NFRS: §8.1.11; FR-HMA-026",
    "8": "BRD §7.1.2; FRS/NFRS: §8.1.13; FR-HMA-052; BR-HMA-010",
    "9": "BRD §7.1.2.4; FRS/NFRS: §8.1.11; FR-HMA-077; §8.6 FR-HMA-042",
    "10": "BRD §7.1.2.1; FRS/NFRS: §17 FB-MRG-003; §8.1.9 FR-HMA-065",
    "11": "BRD §7.1.2 / §7.2.2; §7.5; FRS/NFRS: §8.5; FR-HMA-036–038; FR-SMA-054; FB-MRG-004",
    "12": "BRD §7.5; FRS/NFRS: §8.6; FR-HMA-041–045; FR-SMA-055–058",
    "13": "BRD §7.1.2.3; FRS/NFRS: §8.1.14; FR-HMA-069, FR-HMA-088; §16 RS-MRG-003",
    "14": "BRD §7.2.2.4 / §7.3.2.3; FRS/NFRS: §8.2.8; FR-SMA-019–032, FR-SMA-024–026",
    "15": "BRD §7.1.2.2 / §7.2.2.2; §7.5; FRS/NFRS: §8.1.10; §17 FB-MRG-001; NFR-MRG-PAY-001; FR-HMA-025",
    "16": "BRD §7.1.2 / §7.2.2.2; FRS/NFRS: §8.1.13; FR-HMA-051, FR-HMA-052; FR-SMA-012",
    "17": "BRD §7.1.2.1 / §7.2.2.1; FRS/NFRS: §8.1.4–8.1.5; §8.2.3; FR-HMA-058, FR-HMA-089; FR-SMA-009/062/063",
    "18": "BRD §7.1.2.1 / §7.2.2.1; FRS/NFRS: §8.1.3; §8.2.2; FR-HMA-017; FR-HMA-051",
    "19": "BRD §7.5; FRS/NFRS: §8.4; FR-HMA-034; §12.1 Core entities",
    "20": "BRD §7.1.2.2; §7.3.2.2; FRS/NFRS: §8.1.16; §8.3.3; FR-HMA-054/080; FR-SMA-040/048",
    "21": "BRD §7.2.2; §7.2.2.4; FRS/NFRS: §8.2.5–8.2.7; FR-SMA-014/016/021; FR-SMA-055",
}

INTRO_61 = (
    "Pain points evidenced from Kaveri 2.0 workshops, ServiceDesk tickets and "
    "department discussions. The Addressed in column maps each item to the To-Be "
    "process in this BRD (§7) and to functional / non-functional requirements, "
    "fallbacks and risks in the companion FRS and NFRs document "
    "(FRS_and_NFRS_Marriage_v1.22.docx — §§8–18)."
)

INTRO_751 = (
    "The following As-Is pain points from §6.1 (Kaveri 2.0 workshops, ServiceDesk "
    "tickets and department discussions) are closed in Kaveri 3.0. Cross-references "
    "point first to the To-Be process in this BRD (§7.1–§7.5), then to functional "
    "requirements, fallbacks or NFRs in FRS_and_NFRS_Marriage_v1.22.docx that "
    "implement the fix."
)

INTRO_75 = (
    "This section summarises material enhancements in Kaveri 3.0 compared with the "
    "legacy Kaveri 2.0 Marriage Registration module (§6). Capability highlights "
    "are listed below; §7.5.1 maps each As-Is pain point from §6.1 to the Kaveri 3.0 "
    "closure. Cross-references point to To-Be process in this BRD (§7.1–§7.4) and to "
    "functional / non-functional requirements in FRS_and_NFRS_Marriage_v1.22.docx."
)


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    row = table.rows[-1]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def find_table_by_headers(doc: Document, headers_prefix: list[str]) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = [c.text.strip().replace("\n", " ") for c in table.rows[0].cells]
        if hdr[: len(headers_prefix)] == headers_prefix:
            return table
        # allow renamed last column
        if len(hdr) >= len(headers_prefix) - 1 and hdr[: len(headers_prefix) - 1] == headers_prefix[:-1]:
            if "Addressed" in hdr[-1] or "BRD ref" in hdr[-1] or "FRS" in hdr[-1]:
                return table
    raise KeyError(f"Table not found for headers {headers_prefix}")


def update_pain_table(table: Table, mapping: dict[str, str], header_col: int) -> int:
    set_cell_text(table.rows[0].cells[header_col], COL_HEADER)
    updated = 0
    for row in table.rows[1:]:
        sr = row.cells[0].text.strip()
        if sr in mapping:
            set_cell_text(row.cells[header_col], mapping[sr])
            updated += 1
    return updated


def update_intro(doc: Document, contains: str, new_text: str) -> None:
    for p in doc.paragraphs:
        if contains in p.text:
            set_para_text(p, new_text)
            return
    raise KeyError(f"Intro not found containing {contains!r}")


def main() -> None:
    if not DST.exists():
        raise FileNotFoundError(DST)
    doc = Document(str(DST))

    t61 = find_table_by_headers(
        doc, ["Sr.No", "Pain Point", "Description", "Source", "Addressed in (BRD ref)"]
    )
    n61 = update_pain_table(t61, ADDRESSED, 4)

    t751 = find_table_by_headers(
        doc, ["Sr.No", "Pain Point (As-Is)", "How rectified in Kaveri 3.0", "BRD ref"]
    )
    n751 = update_pain_table(t751, RECTIFIED, 3)

    update_intro(doc, "Pain points evidenced from Kaveri 2.0", INTRO_61)
    update_intro(doc, "This section summarises material enhancements in Kaveri 3.0", INTRO_75)
    update_intro(doc, "The following As-Is pain points from", INTRO_751)

    last = doc.tables[1].rows[-1].cells[3].text.strip()
    if "BRD §7 pain-point refs" not in last:
        add_version_row(
            doc.tables[1],
            [
                "1.22",
                "2026-09-03",
                "Nandha Kumar",
                "§6.1 / §7.5.1: add BRD §7 To-Be refs alongside FRS/NFR companion refs",
                "Prashanth",
            ],
        )

    doc.save(str(DST))
    print(f"Updated {DST}")
    print(f"  §6.1 rows: {n61}; §7.5.1 rows: {n751}")

    # verify
    doc2 = Document(str(DST))
    t61b = find_table_by_headers(
        doc2, ["Sr.No", "Pain Point", "Description", "Source", COL_HEADER]
    )
    missing = []
    for row in t61b.rows[1:]:
        ref = row.cells[4].text
        if "BRD §7" not in ref:
            missing.append(row.cells[0].text.strip())
        if "FRS/NFRS:" not in ref:
            missing.append(f"{row.cells[0].text.strip()}:no-frs")
    if missing:
        raise AssertionError(f"Missing BRD/FRS labels: {missing}")
    print("Verification OK")


if __name__ == "__main__":
    main()
