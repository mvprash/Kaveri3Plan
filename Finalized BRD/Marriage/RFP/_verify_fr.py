import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

path = r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP\BRD_Marriage_v1.10.docx"
doc = Document(path)

def fr(req):
    for t in doc.tables:
        for row in t.rows:
            if row.cells[0].text.strip() == req:
                return row
    return None

for req in ("FR-SMA-001", "FR-SMA-013", "FR-SMA-061"):
    row = fr(req)
    print(req)
    print("  REQ:", row.cells[1].text.strip()[:200])
    print("  AC:", row.cells[3].text.strip()[:200])
    print()

print("Status Notice generated:", [c.text.strip()[:100] for c in doc.tables[16].rows[10].cells])
print("Status Notice published:", [c.text.strip()[:100] for c in doc.tables[16].rows[11].cells])
print("Cover version:", doc.tables[0].rows[2].cells[1].text.strip())
print("Last updated:", doc.tables[0].rows[11].cells[1].text.strip())
print("File OK")
