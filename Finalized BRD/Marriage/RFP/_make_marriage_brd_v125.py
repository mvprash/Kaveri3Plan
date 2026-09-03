# -*- coding: utf-8 -*-
"""Create BRD_Marriage_BRD_v1.25.docx from v1.23.

Update cross-references from old decimal §7 / §3 numbering (7.1.2.1, 3.7, …)
to the document's current Word auto-number scheme (7.i.b.A, 3.vii, …)
wherever those references appear (implementation tables, Contents, pain points,
process notes, etc.). Does NOT alter In-scope bullet wording beyond renumbering
existing 7.x / 3.x refs.
"""
from __future__ import annotations

import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_BRD_v1.23.docx"
DST = BASE / "BRD_Marriage_BRD_v1.25.docx"

# Old decimal → new auto-number labels (longest keys first).
# Section 7 under Future state (To-Be):
#   H3 lowerRoman: i–v; H4 lowerLetter: a/b; H5 upperLetter: A/B/C/D
#   (7.3 has no Online → Offline=B, Status=C)
# Section 3 under Legal:
#   H3 lowerRoman: i–vii; H4/H5 lowerLetter: a/b/…
REF_MAP: list[tuple[str, str]] = [
    # --- §7.5 ---
    ("7.5.1", "7.v.a"),
    ("7.5", "7.v"),
    # --- §7.4 ---
    ("7.4.1", "7.iv.a"),
    ("7.4", "7.iv"),
    # --- §7.3 (no Online H5) ---
    ("7.3.2.3", "7.iii.b.C"),
    ("7.3.2.2", "7.iii.b.B"),
    ("7.3.2.1", "7.iii.b.A"),
    ("7.3.2.x", "7.iii.b.x"),
    ("7.3.2", "7.iii.b"),
    ("7.3.1", "7.iii.a"),
    ("7.3", "7.iii"),
    # --- §7.2 ---
    ("7.2.2.4", "7.ii.b.D"),
    ("7.2.2.3", "7.ii.b.C"),
    ("7.2.2.2", "7.ii.b.B"),
    ("7.2.2.1", "7.ii.b.A"),
    ("7.2.2.x", "7.ii.b.x"),
    ("7.2.2", "7.ii.b"),
    ("7.2.1", "7.ii.a"),
    ("7.2", "7.ii"),
    # --- §7.1 ---
    ("7.1.2.4", "7.i.b.D"),
    ("7.1.2.3", "7.i.b.C"),
    ("7.1.2.2", "7.i.b.B"),
    ("7.1.2.1", "7.i.b.A"),
    ("7.1.2", "7.i.b"),
    ("7.1.1", "7.i.a"),
    ("7.1", "7.i"),
    # --- §3 (legal) — only standalone section refs, not Act "Sec. 3" ---
    ("3.5.5", "3.v.e"),
    ("3.5.4", "3.v.d"),
    ("3.5.3", "3.v.c"),
    ("3.5.2", "3.v.b"),
    ("3.5.1", "3.v.a"),
    ("3.6.4", "3.vi.d"),
    ("3.6.3", "3.vi.c"),
    ("3.6.2", "3.vi.b"),
    ("3.6.1", "3.vi.a"),
    ("3.3.2", "3.iii.b"),
    ("3.3.1", "3.iii.a"),
    ("3.2.2", "3.ii.b"),
    ("3.2.1", "3.ii.a"),
    ("3.7", "3.vii"),
    ("3.6", "3.vi"),
    ("3.5", "3.v"),
    ("3.4", "3.iv"),
    ("3.3", "3.iii"),
    ("3.2", "3.ii"),
    ("3.1", "3.i"),
]

# Contents lines: old TOC text → new TOC text
TOC_REPLACEMENTS: list[tuple[str, str]] = [
    ("1. Executive summary", "1. Executive summary"),
    ("2. Scope", "2. Scope"),
    ("2.1 In scope (Hindu Marriage & Special Marriage — Phase [1])",
     "2.i In scope (Hindu Marriage & Special Marriage — Phase [1])"),
    ("2.2 Out of scope (unless PO promotes)", "2.ii Out of scope (unless PO promotes)"),
    ("2.3 Assumptions", "2.iii Assumptions"),
    ("2.4 Constraints", "2.iv Constraints"),
    ("3. Legal and regulatory reference", "3. Legal and regulatory reference"),
    ("3.1 Applicable Acts", "3.i Applicable Acts"),
    ("3.2 Relevant sections followed by the Department for Marriage Registration",
     "3.ii Relevant sections followed by the Department for Marriage Registration"),
    ("3.3 Relevant rules followed by the Department for Marriage Registration",
     "3.iii Relevant rules followed by the Department for Marriage Registration"),
    ("3.4 Relevant notifications issued by the Department for Marriage Registration",
     "3.iv Relevant notifications issued by the Department for Marriage Registration"),
    ("3.5 Hindu Marriage statutory forms mapping",
     "3.v Hindu Marriage statutory forms mapping"),
    ("3.6 Special Marriage statutory forms mapping",
     "3.vi Special Marriage statutory forms mapping"),
    ("3.7 Sakala — Karnataka Guarantee of Services",
     "3.vii Sakala — Karnataka Guarantee of Services"),
    ("4. Stakeholders and actors", "4. Stakeholders and actors"),
    ("5. Definitions and glossary", "5. Definitions and glossary"),
    ("6. Current state (As-Is)", "6. Current state (As-Is)"),
    ("6.1 As-Is pain points", "6.i As-Is pain points"),
    ("7. Future state (To-Be)", "7. Future state (To-Be)"),
    ("7.1 Hindu Marriage", "7.i Hindu Marriage"),
    ("7.1.1 Channel models", "7.i.a Channel models"),
    ("7.1.2 Process Diagram", "7.i.b Process Diagram"),
    ("7.1.2.1 Common intake steps", "7.i.b.A Common intake steps"),
    ("7.1.2.2 Online", "7.i.b.B Online"),
    ("7.1.2.3 Offline (In Person)", "7.i.b.C Offline (In Person)"),
    ("7.1.2.4 Application Status Model", "7.i.b.D Application Status Model"),
    ("7.2 Special Marriage Notice Generation (Intended Marriage / Other Forms)",
     "7.ii Special Marriage Notice Generation (Intended Marriage / Other Forms)"),
    ("7.2.1 Channel models", "7.ii.a Channel models"),
    ("7.2.2 Process Diagram", "7.ii.b Process Diagram"),
    ("7.2.2.1 Common intake steps", "7.ii.b.A Common intake steps"),
    ("7.2.2.2 Online", "7.ii.b.B Online"),
    ("7.2.2.3 Offline (In Person)", "7.ii.b.C Offline (In Person)"),
    ("7.2.2.4 Application Status Model", "7.ii.b.D Application Status Model"),
    ("7.3 Special Marriage Marriage Registration (Intended Marriage / Other Forms)",
     "7.iii Special Marriage Marriage Registration (Intended Marriage / Other Forms)"),
    ("7.3.1 Channel models", "7.iii.a Channel models"),
    ("7.3.2 Process Diagram", "7.iii.b Process Diagram"),
    ("7.3.2.1 Common intake steps", "7.iii.b.A Common intake steps"),
    ("7.3.2.2 Offline (In Person)", "7.iii.b.B Offline (In Person)"),
    ("7.3.2.3 Application Status Model", "7.iii.b.C Application Status Model"),
    ("7.4 Hindu Marriage — Null and Void Endorsement (Court Order)",
     "7.iv Hindu Marriage — Null and Void Endorsement (Court Order)"),
    ("7.4.1 Process steps", "7.iv.a Process steps"),
    ("7.5 What is new in Kaveri 3.0", "7.v What is new in Kaveri 3.0"),
    ("7.5.1 Rectified As-Is pain points", "7.v.a Rectified As-Is pain points"),
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def remap_refs(text: str) -> str:
    """Replace old decimal section refs with new auto-number labels.

    Avoids rewriting Act/Rule 'Sec. N' by only matching BRD-style refs that
    appear as bare 7.x / 3.x tokens (optionally prefixed by §).
    """
    if not text:
        return text
    out = text
    for old, new in REF_MAP:
        # Match optional §, then old token, not followed by another digit/dot-digit
        # so 7.1 does not eat into 7.1.2.1 (we already replace longest first).
        pattern = re.compile(
            r"(?<![A-Za-z0-9])(§?)" + re.escape(old) + r"(?![0-9])"
        )
        out = pattern.sub(lambda m, n=new: (m.group(1) or "") + n, out)
    return out


def remap_paragraph(paragraph: Paragraph) -> bool:
    old = paragraph.text
    # Skip Contents exact-line handling here (done separately) — still OK to remap
    new = remap_refs(old)
    if new != old:
        set_para_text(paragraph, new)
        return True
    return False


def remap_table(table: Table) -> int:
    n = 0
    for row in table.rows:
        for cell in row.cells:
            # Remap each paragraph in the cell to preserve structure
            for p in cell.paragraphs:
                old = p.text
                new = remap_refs(old)
                if new != old:
                    set_para_text(p, new)
                    n += 1
    return n


def update_contents(doc: Document) -> int:
    n = 0
    in_contents = False
    for p in doc.paragraphs:
        st = str(p.style.name) if p.style else ""
        t = p.text.strip()
        if st.startswith("Heading") and t == "Contents":
            in_contents = True
            continue
        if in_contents and st.startswith("Heading"):
            break
        if not in_contents:
            continue
        for old, new in TOC_REPLACEMENTS:
            if t == old:
                set_para_text(p, new)
                n += 1
                break
        else:
            # Fallback: remap any residual decimal refs in TOC lines
            if remap_paragraph(p):
                n += 1
    return n


def update_doc_control(doc: Document) -> None:
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.25")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-03")
    # version history — append row
    vt = doc.tables[1]
    vt._tbl.append(deepcopy(vt.rows[-1]._tr))
    row = vt.rows[-1]
    vals = [
        "1.25",
        "2026-09-03",
        "Nandha Kumar",
        "Update §3 / §7 cross-references to match lettered/roman auto-numbering "
        "(e.g. 7.1.2.1 → 7.i.b.A) wherever those refs appear",
        "Prashanth",
    ]
    for ci, val in enumerate(vals):
        set_cell_text(row.cells[ci], val)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    n_toc = update_contents(doc)

    n_para = 0
    in_contents = False
    for p in doc.paragraphs:
        st = str(p.style.name) if p.style else ""
        t = p.text.strip()
        if st.startswith("Heading") and t == "Contents":
            in_contents = True
            continue
        if in_contents and st.startswith("Heading"):
            in_contents = False
        if in_contents:
            continue  # already handled
        # Do not rewrite Heading texts (they have no decimal numbers in v1.23)
        if st.startswith("Heading"):
            continue
        if remap_paragraph(p):
            n_para += 1

    n_tbl = 0
    for table in doc.tables:
        n_tbl += remap_table(table)

    update_doc_control(doc)
    doc.save(str(DST))
    print(f"Wrote {DST}")
    print(f"  Contents lines: {n_toc}; body paras: {n_para}; table paras: {n_tbl}")

    # Verify key remaps in HMA / SMA implementation tables
    doc2 = Document(str(DST))
    hma = doc2.tables[4]
    sma = doc2.tables[5]
    hma_impl = [r.cells[-1].text for r in hma.rows[1:]]
    sma_impl = [r.cells[-1].text for r in sma.rows[1:]]
    assert any("7.i.b.A" in x for x in hma_impl), hma_impl[:3]
    assert any("7.i.b.B" in x and "7.i.b.C" in x for x in hma_impl), hma_impl
    assert any("7.ii.b.A" in x for x in sma_impl), sma_impl[:3]
    assert any("7.iii.b.A" in x for x in sma_impl), sma_impl
    # Ensure old decimal process refs are gone from impl columns
    for x in hma_impl + sma_impl:
        assert "7.1.2" not in x and "7.2.2" not in x and "7.3.2" not in x, x
    # In-scope bullets still present (not rewritten into 'see §3.5' style)
    bullets = [p.text for p in doc2.paragraphs if "Statutory artefacts" in p.text]
    assert bullets and "see §3.5" not in bullets[0], bullets[0]
    print("Verification OK")
    print("  HMA Sec.2 impl:", hma_impl[0][:80])
    print("  SMA Sec.4 impl:", sma_impl[0][:80])


if __name__ == "__main__":
    main()
