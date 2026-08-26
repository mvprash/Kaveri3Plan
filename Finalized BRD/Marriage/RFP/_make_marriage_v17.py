# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.7.docx from v1.6.

Consolidate duplicity / redundancy:
- Merge Special Marriage notice workflows (former §7.2 and §7.4) into one
  workflow with an Other Forms data-capture branch.
- Merge Special Marriage registration workflows (former §7.3 and §7.5) into
  one workflow with timeline / certificate-schedule branches.
- Merge §8.2.4 / §8.2.5 bridegroom + bride field catalogues into Party Particulars.
- Keep FR-SMA-034 only in the witnesses section (remove identical restatement).
- Drop placeholder §8 sub-sections that only cross-reference earlier FRs.
- Keep §7 as process authority and §8 as testable requirements (no process restatement).
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\Prashanth\Official\Kaveri 3.0\Kaveri3Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.6.docx"
DST = BASE / "BRD_Marriage_v1.7.docx"

TOC_OLD_FROM_72 = [
    "7.2 Special Marriage (Intended Marriage) Notice Generation",
    "7.2.1 Channel models",
    "7.2.2 Process Diagram",
    "7.2.2.1 Common intake steps",
    "7.2.2.2 Online",
    "7.2.2.3 Offline (In Person)",
    "7.2.2.4 Application Status Model",
    "7.3 Special Marriage (Intended Marriage) Marriage Registration",
    "7.3.1 Channel models",
    "7.3.2 Process Diagram",
    "7.3.2.1 Common intake steps",
    "7.3.2.2 Offline (In Person)",
    "7.3.2.3 Application Status Model",
    "7.4 Special Marriage (Other Forms) Notice Generation",
    "7.4.1 Channel models",
    "7.4.2 Process Diagram",
    "7.4.2.1 Common intake steps",
    "7.4.2.2 Online",
    "7.4.2.3 Offline (In Person)",
    "7.4.2.4 Application Status Model",
    "7.5 Special Marriage (Other Forms) Marriage Registration",
    "7.5.1 Channel models",
    "7.5.2 Process Diagram",
    "7.5.2.1 Common intake steps",
    "7.5.2.2 Offline (In Person)",
    "7.5.2.3 Application Status Model",
    "8. Functional requirements",
    "8.1 Hindu Marriage",
    "8.1.1 Eligibility and module entry (FR-HMA-001–004)",
    "8.1.2 Jurisdiction and office routing (FR-HMA-005–006)",
    "8.1.3 Data capture — marriage details (FR-HMA-007–009)",
    "8.1.4 Data capture — bridegroom (FR-HMA-010)",
    "8.1.5 Data capture — bride (FR-HMA-011)",
    "8.1.6 Data capture — witnesses (FR-HMA-012–014)",
    "8.1.7 Form I — memorandum generation and completeness (FR-HMA-015–017)",
    "8.1.8 Form IA — application and declarations (FR-HMA-083–085)",
    "8.1.9 Supporting documents and memorandum (FR-HMA-018–019)",
    "8.1.10 Fees and payments (FR-HMA-020–025)",
    "8.1.11 SRO scrutiny and registration (FR-HMA-026–030)",
    "8.1.12 Channel selection and prerequisite acknowledgement (FR-HMA-046–050)",
    "8.1.13 Online channel — office, Form I & Form IA, eSign (FR-HMA-051–058, 086–087)",
    "8.1.14 Offline channel — printout, DEO upload (FR-HMA-059–069)",
    "8.1.15 SR verification (FR-HMA-070–077)",
    "8.1.16 Digital signature and certificate issuance (FR-HMA-078–082)",
    "8.2 Special Marriage (Intended Marriage/Other Forms) Notice Generation",
    "8.2.1 Eligibility and module entry (FR-SMA-001–006)",
    "8.2.2 Jurisdiction and office routing",
    "8.2.3 Data capture — marriage details (FR-SMA-007–013, FR-SMA-061)",
    "8.2.4 Data capture — bridegroom (FR-SMA-062)",
    "8.2.5 Data capture — bride (FR-SMA-063)",
    "8.2.6 Data capture — witnesses",
    "8.2.7 Form I — memorandum generation and completeness",
    "8.2.8 Form IA — application and declarations",
    "8.2.9 Supporting documents and memorandum",
    "8.2.10 Fees and payments (FR-SMA-049–053)",
    "8.2.11 SRO scrutiny and registration",
    "8.2.12 Channel selection and prerequisite acknowledgement",
    "8.2.13 Online channel — office, Form I & Form IA, eSign",
    "8.2.14 Offline channel — printout, DEO upload",
    "8.2.15 SR verification",
    "8.2.16 Notice generation, Marriage Notice Book and publication (FR-SMA-014–021)",
    "8.2.17 Notice validity, timeline gate and fresh notice (FR-SMA-029–032)",
    "8.2.18 Digital signature and certificate issuance",
    "8.3 Special Marriage (Intended Marriage/Other Forms) Marriage Registration",
    "8.3.1 Eligibility and module entry",
    "8.3.2 Data capture — witnesses (FR-SMA-034, FR-SMA-064–065)",
    "8.3.3 Objection management and enquiry (FR-SMA-022–028)",
    "8.3.4 Fees and payments",
    "8.3.5 SRO scrutiny and registration",
    "8.3.6 Solemnization, declarations and certificate (FR-SMA-033–041)",
    "8.3.7 Special Marriage Other Forms — application and registration (FR-SMA-042–048)",
    "8.3.8 Offline channel — printout, DEO upload",
    "8.3.9 SR verification",
    "8.3.10 Digital signature and certificate issuance",
    "8.4 Post-registration services (FR-HMA-031–035)",
    "8.5 Notifications (FR-HMA-036–038, FR-SMA-054)",
    "8.6 Reports and MIS (FR-HMA-039–045, FR-SMA-055–060)",
]

TOC_NEW_FROM_72 = [
    "7.2 Special Marriage Notice Generation (Intended Marriage / Other Forms)",
    "7.2.1 Channel models",
    "7.2.2 Process Diagram",
    "7.2.2.1 Common intake steps",
    "7.2.2.2 Online",
    "7.2.2.3 Offline (In Person)",
    "7.2.2.4 Application Status Model",
    "7.3 Special Marriage Marriage Registration (Intended Marriage / Other Forms)",
    "7.3.1 Channel models",
    "7.3.2 Process Diagram",
    "7.3.2.1 Common intake steps",
    "7.3.2.2 Offline (In Person)",
    "7.3.2.3 Application Status Model",
    "8. Functional requirements",
    "8.1 Hindu Marriage",
    "8.1.1 Eligibility and module entry (FR-HMA-001–004)",
    "8.1.2 Jurisdiction and office routing (FR-HMA-005–006)",
    "8.1.3 Data capture — marriage details (FR-HMA-007–009)",
    "8.1.4 Data capture — bridegroom (FR-HMA-010)",
    "8.1.5 Data capture — bride (FR-HMA-011)",
    "8.1.6 Data capture — witnesses (FR-HMA-012–014)",
    "8.1.7 Form I — memorandum generation and completeness (FR-HMA-015–017)",
    "8.1.8 Form IA — application and declarations (FR-HMA-083–085)",
    "8.1.9 Supporting documents and memorandum (FR-HMA-018–019)",
    "8.1.10 Fees and payments (FR-HMA-020–025)",
    "8.1.11 SRO scrutiny and registration (FR-HMA-026–030)",
    "8.1.12 Channel selection and prerequisite acknowledgement (FR-HMA-046–050)",
    "8.1.13 Online channel — office, Form I & Form IA, eSign (FR-HMA-051–058, 086–087)",
    "8.1.14 Offline channel — printout, DEO upload (FR-HMA-059–069)",
    "8.1.15 SR verification (FR-HMA-070–077)",
    "8.1.16 Digital signature and certificate issuance (FR-HMA-078–082)",
    "8.2 Special Marriage (Intended Marriage/Other Forms) Notice Generation",
    "8.2.1 Eligibility and module entry (FR-SMA-001–006)",
    "8.2.2 Data capture — marriage details (FR-SMA-007–013, FR-SMA-061)",
    "8.2.3 Data capture — party particulars (FR-SMA-062–063)",
    "8.2.4 Declaration (The Third Schedule)",
    "8.2.5 Notice of Intended Marriage (The Second Schedule)",
    "8.2.6 Fees and payments (FR-SMA-049–053)",
    "8.2.7 Notice generation, Marriage Notice Book and publication (FR-SMA-014–021)",
    "8.2.8 Notice validity, timeline gate and fresh notice (FR-SMA-029–032)",
    "8.3 Special Marriage (Intended Marriage/Other Forms) Marriage Registration",
    "8.3.1 Data capture — witnesses (FR-SMA-034, FR-SMA-064–065)",
    "8.3.2 Objection management and enquiry (FR-SMA-022–028)",
    "8.3.3 Solemnization, declarations and certificate (FR-SMA-033–041)",
    "8.3.4 Special Marriage Other Forms — application and registration (FR-SMA-042–048)",
    "8.4 Post-registration services (FR-HMA-031–035)",
    "8.5 Notifications (FR-HMA-036–038, FR-SMA-054)",
    "8.6 Reports and MIS (FR-HMA-039–045, FR-SMA-055–060)",
]

PLACEHOLDER_H4 = [
    "8.2.2 Jurisdiction and office routing",
    "8.2.6 Data capture — witnesses",
    "8.2.9 Supporting documents and memorandum",
    "8.2.11 SRO scrutiny and registration",
    "8.2.12 Channel selection and prerequisite acknowledgement",
    "8.2.13 Online channel — office, Form I & Notice of Intended Marriage , eSign",
    "8.2.14 Offline channel — printout, DEO upload",
    "8.2.15 SR verification",
    "8.2.18 Digital signature and certificate issuance",
    "8.3.1 Eligibility and module entry",
    "8.3.4 Fees and payments",
    "8.3.5 SRO scrutiny and registration",
    "8.3.8 Offline channel — printout, DEO upload",
    "8.3.9 SR verification",
    "8.3.10 Digital signature and certificate issuance",
]

NEW_82_H4 = [
    "8.2.1 Eligibility and module entry",
    "8.2.2 Data capture — marriage details",
    "8.2.3 Data capture — party particulars",
    "8.2.4 Declaration (The Third Schedule)",
    "8.2.5 Notice of Intended Marriage (The Second Schedule)",
    "8.2.6 Fees and payments",
    "8.2.7 Notice generation, Marriage Notice Book and publication",
    "8.2.8 Notice validity, timeline gate and fresh notice",
]

NEW_83_H4 = [
    "8.3.1 Data capture — witnesses",
    "8.3.2 Objection management and enquiry",
    "8.3.3 Solemnization, declarations and certificate",
    "8.3.4 Special Marriage Other Forms — application and registration",
]

# Longest-first replacements applied after structural edits.
TEXT_REPLACEMENTS: list[tuple[str, str]] = [
    (
        "7.5 steps 15–18 and 7.6 steps 14–19",
        "7.2.2.2 / 7.2.2.3",
    ),
    ("7.2.2.1 / 7.4.2.1 common intake", "7.2.2.1 common intake"),
    ("7.3.2.2 / 7.5.2.2", "7.3.2.2"),
    ("(Ref: 7.7 steps 5–16;", "(Ref: 7.3.2.2;"),
    ("(Ref: 7.7 / 7.10 timeline decision;", "(Ref: 7.3 timeline decision;"),
    ("(Ref: 7.7 / 7.10 objection branch;", "(Ref: 7.3 objection branch;"),
    ("(Ref: 7.8–7.10;", "(Ref: 7.3;"),
    ("(Ref: 7.5–7.10 intake;", "(Ref: 7.2–7.3 intake;"),
    ("(Ref: 7.5–7.10 first and second payment;", "(Ref: 7.2 / 7.3 first and second payment;"),
    ("(Ref: 7.5–7.10; SMA 1954 Sec. 6, 47–48)", "(Ref: 7.2–7.3; SMA 1954 Sec. 6, 47–48)"),
    ("mapped to 7.5–7.10", "mapped to 7.2–7.3"),
    ("process diagrams 7.5 / 7.6", "process diagrams 7.2"),
    ("Process diagrams 7.7 / 7.10", "process diagrams 7.3"),
    ("Per 7.3 / 7.5 timeline", "Per 7.3 timeline"),
    ("per 7.3 / 7.5 timeline", "per 7.3 timeline"),
    ("Process diagrams 7.5 / 7.6", "Process diagrams 7.2"),
    ("process diagrams 7.5 / 7.6", "process diagrams 7.2"),
    ("(7.2 steps 6–7)", "(7.1.2.1 steps 6–7)"),
    ("see 7.3 / 7.5", "see 7.3"),
    ("per 7.3 / 7.5.", "per 7.3."),
    ("(7.5–7.10):", "(7.2–7.3):"),
    ("diagrams (7.5–7.10)", "diagrams (7.2–7.3)"),
    ("(7.3–7.4):", "(7.1.2.2 / 7.1.2.3):"),
    ("(Ref: 7.2 steps 5–6 — both diagrams;", "(Ref: 7.1.2.1 steps 5–6 — both diagrams;"),
    ("(Ref: 7.3 steps 7 + 8–10)", "(Ref: 7.1.2.2)"),
    ("(Ref: 7.4 steps 9–13 including SR allocates to DEO)", "(Ref: 7.1.2.3 including SR allocates to DEO)"),
    ("(Ref: 7.3 step 11; 7.4 steps 8 and 14)", "(Ref: 7.1.2.2 / 7.1.2.3)"),
    ("(Ref: 7.3 steps 13–14; 7.4 steps 15–16)", "(Ref: 7.1.2.2 / 7.1.2.3)"),
    ("Hard gate 7.2 step 6;", "Hard gate 7.1.2.1 step 6;"),
    ("status from 7.5 model", "status from 7.1.2.4 model"),
    ("7.2 steps 2–3", "7.1.2.1 steps 2–3"),
    ("7.2 step 4", "7.1.2.1 step 4"),
    ("7.2 step 6; FR-HMA-046", "7.1.2.1 step 6; FR-HMA-046"),
    ("7.2 step 5; FR-HMA-047", "7.1.2.1 step 5; FR-HMA-047"),
    ("FR-SMA-043 (8.3.7)", "FR-SMA-043 (8.3.4)"),
    ("as in 8.21", "as in 8.2.7"),
    ("workflow of 8.22 applies", "workflow of 8.3.2 applies"),
    ("8.24 / 8.25;", "8.3.3 / 8.3.4;"),
    ("8.19;", "8.2.1;"),
    ("8.20;", "8.2.2;"),
    ("8.21;", "8.2.7;"),
    ("8.22;", "8.3.2;"),
    ("8.23;", "8.2.8;"),
    ("8.24;", "8.3.3;"),
    ("8.25;", "8.3.4;"),
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        set_cell_text(row.cells[ci], val)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def heading_level(paragraph: Paragraph) -> int | None:
    sn = style_name(paragraph)
    if sn.startswith("Heading "):
        try:
            return int(sn.split()[-1])
        except ValueError:
            return None
    return None


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    p_pr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child is not p_pr:
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_element(el) -> None:
    el.getparent().remove(el)


def delete_para(paragraph: Paragraph) -> None:
    delete_element(paragraph._p)


def heading_block_nodes(start: Paragraph, *, stop_at_level: int) -> list:
    """Body nodes from `start` through the node before the next heading of
    level <= stop_at_level (or sectPr)."""
    nodes = [start._p]
    nxt = start._p.getnext()
    parent = start._parent
    while nxt is not None and nxt.tag != qn("w:sectPr"):
        if nxt.tag == qn("w:p"):
            p = Paragraph(nxt, parent)
            lvl = heading_level(p)
            if lvl is not None and lvl <= stop_at_level:
                break
        nodes.append(nxt)
        nxt = nxt.getnext()
    return nodes


def delete_heading_block(start: Paragraph, *, stop_at_level: int) -> None:
    for node in heading_block_nodes(start, stop_at_level=stop_at_level):
        delete_element(node)


def move_element_after(el, anchor_el) -> None:
    el.getparent().remove(el)
    anchor_el.addnext(el)


def para_has_blip(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.findall(".//" + qn("a:blip")))


def previous_blip_para(caption: Paragraph) -> Paragraph:
    el = caption._p.getprevious()
    while el is not None:
        if el.tag == qn("w:p"):
            p = Paragraph(el, caption._parent)
            if para_has_blip(p):
                return p
        el = el.getprevious()
    raise KeyError("No image paragraph before caption")


def add_version_row(table: Table, values: list[str]) -> None:
    last = table.rows[-1]._tr
    new_tr = deepcopy(last)
    table._tbl.append(new_tr)
    set_row(table, len(table.rows) - 1, values)


def replace_in_runs(paragraph: Paragraph, old: str, new: str) -> bool:
    full = paragraph.text
    if old not in full:
        return False
    # Prefer run-0 rewrite when the match sits in concatenated text.
    set_para_text(paragraph, full.replace(old, new))
    return True


def replace_everywhere(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if replace_in_runs(p, old, new):
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_runs(p, old, new):
                        n += 1
    return n


def update_toc(doc: Document) -> None:
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == TOC_OLD_FROM_72[0] and not style_name(p).startswith("Heading"):
            start = i
            break
    if start is None:
        raise KeyError("TOC start for 7.2 not found")
    for offset, expected in enumerate(TOC_OLD_FROM_72):
        actual = doc.paragraphs[start + offset].text.strip()
        if actual != expected:
            raise KeyError(
                f"TOC mismatch at offset {offset}: expected {expected!r}, got {actual!r}"
            )
    n_old = len(TOC_OLD_FROM_72)
    n_new = len(TOC_NEW_FROM_72)
    for i in range(min(n_old, n_new)):
        set_para_text(doc.paragraphs[start + i], TOC_NEW_FROM_72[i])
    if n_new < n_old:
        # Delete leftover TOC paragraphs (from the end so indices stay valid).
        extras = [doc.paragraphs[start + i] for i in range(n_new, n_old)]
        for p in extras:
            delete_para(p)


def h4_between(doc: Document, h3_text: str, stop_h_prefix: str) -> list[Paragraph]:
    start = find_para(doc, exact=h3_text, heading_only=True)
    found: list[Paragraph] = []
    el = start._p.getnext()
    while el is not None and el.tag != qn("w:sectPr"):
        if el.tag == qn("w:p"):
            p = Paragraph(el, start._parent)
            t = p.text.strip()
            lvl = heading_level(p)
            if lvl == 2 or (lvl == 3 and t.startswith(stop_h_prefix)):
                break
            if lvl == 4:
                found.append(p)
        el = el.getnext()
    return found


def remove_table_row(table: Table, ri: int) -> None:
    table._tbl.remove(table.rows[ri]._tr)


def find_table_by_cell(doc: Document, row: int, col: int, exact: str) -> Table:
    for table in doc.tables:
        if len(table.rows) <= row:
            continue
        cells = table.rows[row].cells
        if col >= len(cells):
            continue
        if cells[col].text.strip() == exact:
            return table
    raise KeyError(f"Table not found for [{row},{col}]={exact!r}")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # --- Document control -----------------------------------------------------
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.7")
    set_cell_text(doc.tables[0].rows[12].cells[1], "2026-08-26")
    add_version_row(
        doc.tables[1],
        [
            "1.7",
            "2026-08-26",
            "Nandha Kumar",
            "Removed duplicity: combined Special Marriage notice (former 7.2/7.4) and "
            "registration (former 7.3/7.5) into single workflows with conditional "
            "branches; merged party-particulars field catalogue; dropped duplicate "
            "FR-SMA-034 and placeholder FR cross-reference sub-sections; §7 remains "
            "process authority, §8 testable requirements only",
            "Prashanth",
        ],
    )

    # --- §7 intro -------------------------------------------------------------
    set_para_text(
        find_para(doc, contains="The Marriage Registration To-Be model is organized by service."),
        "The Marriage Registration To-Be model is organized by service. Each service "
        "documents its Channel models, Process Diagram (Common intake steps, Online "
        "and/or Offline (In Person) flows) and Application Status Model. The module "
        "covers Hindu Marriage (Online / Offline) and Special Marriage notice "
        "generation (Intended Marriage and Other Forms share one workflow) followed "
        "by marriage registration (shared workflow with path-specific timeline and "
        "certificate-schedule branches). Notice channels are initiated through the "
        "citizen portal; registration is completed In Person at the Sub-Registrar office.",
    )

    # --- Combine §7.2 notice --------------------------------------------------
    h72 = find_para(
        doc,
        exact="7.2 Special Marriage (Intended Marriage) Notice Generation",
        heading_only=True,
    )
    set_para_text(
        h72,
        "7.2 Special Marriage Notice Generation (Intended Marriage / Other Forms)",
    )
    insert_paragraph_after(
        h72,
        "Intended Marriage (Chapter II) and Other Forms (Chapter III) share a single "
        "notice-generation workflow. The only functional branch at intake is data "
        "capture: Other Forms additionally requires date, place and form / rites of "
        "the already celebrated ceremony. Online and Offline publication, first "
        "payment, eSign / notice-board and the 30-day countdown are otherwise identical.",
        style="Normal",
    )
    set_para_text(
        find_para(
            doc,
            contains="Intended Marriage notice generation supports Online and Offline",
        ),
        "Notice generation supports Online and Offline (In Person) channels for both "
        "Intended Marriage and Other Forms. The selected notice channel drives "
        "publication mode (portal vs notice board) and office tasks. Service path "
        "(Intended Marriage vs Other Forms) is chosen with the channel and drives "
        "the marriage-details branch only.",
    )
    # T13 channel models (Intended Notice) — apply to both paths
    set_row(
        doc.tables[13],
        1,
        [
            "Special Marriage Notice — Online (Intended Marriage / Other Forms)",
            "e-KYC / Face Authentication bride & bridegroom, document upload (incl. photos), First Payment, portal notice, eSign",
            "SR verification; System notice generation (no DEO / appointment)",
            "Online notice",
        ],
    )
    set_row(
        doc.tables[13],
        2,
        [
            "Special Marriage Notice — Offline (Intended Marriage / Other Forms)",
            "Capture details (e-KYC / Face Authentication or manual), document upload, First Payment, appointment",
            "SR verification, notice generation, DEO photo/print/sign/scan/upload, paste on notice board",
            "Offline notice",
        ],
    )
    set_para_text(
        find_para(
            doc,
            exact="Identical in both Intended Marriage notice diagrams (Citizens and System lanes):",
        ),
        "Identical in both notice diagrams (Citizens and System lanes). Intended Marriage and Other Forms share these steps; the Other Forms branch is in step 7:",
    )
    set_para_text(
        find_para(
            doc,
            contains="Select channel / service path: Special Marriage (Intended Marriage) Notice Online / Offline.",
        ),
        "Select channel / service path: Special Marriage Notice Online / Offline, and "
        "path Intended Marriage or Other Forms. Channel is chosen before prerequisites. "
        "Marriage Registration is initiated later by selecting an approved notice (see 7.3).",
    )
    set_para_text(
        find_para(
            doc,
            contains="Enter / capture Bride details, Bridegroom details — persisted to the notice application",
        ),
        "Enter / capture Bride details and Bridegroom details — persisted to the notice "
        "application record. Conditional branch: if Other Forms, also capture marriage "
        "details (date, place and form of the already celebrated ceremony); if Intended "
        "Marriage, do not capture or require those fields. SR rejection returns to "
        "Prerequisite & declaration for correction and resubmission.",
    )
    set_para_text(
        find_para(
            doc,
            exact="Flow (continuing from 7.2.2.1 common intake — Online notice-specific steps):",
        ),
        "Flow (continuing from 7.2.2.1 common intake — Online notice-specific steps; shared by Intended Marriage and Other Forms):",
    )
    set_para_text(
        find_para(
            doc,
            contains="Key characteristics: e-KYC / Face Authentication on Bride & Bridegroom when Aadhaar available",
        ),
        "Key characteristics (both paths): e-KYC / Face Authentication on Bride & "
        "Bridegroom when Aadhaar available; document upload including individual photos; "
        "SR verification before First Payment; System generates notice after payment and "
        "displays it on the portal; citizen e-sign on the generated notice; no Online "
        "appointment or DEO; 30-day countdown (Sec. 7 / Sec. 16). Other Forms only: "
        "marriage details captured at intake.",
    )
    set_para_text(
        find_para(
            doc,
            exact="Flow (continuing from 7.2.2.1 common intake — Offline notice-specific steps):",
        ),
        "Flow (continuing from 7.2.2.1 common intake — Offline notice-specific steps; shared by Intended Marriage and Other Forms):",
    )
    set_para_text(
        find_para(
            doc,
            contains="Key characteristics: Aadhaar YES (e-KYC / Face Authentication) / NO (manual) capture; SR verification → First",
        ),
        "Key characteristics (both paths): Aadhaar YES (e-KYC / Face Authentication) / NO "
        "(manual) capture; SR verification → First Payment → appointment; SR generates "
        "notice & selects DEO; individual photo capture; physical notice-board paste "
        "(Sec. 6(2)); 30-day countdown. SR rejection returns to Prerequisite & declaration. "
        "Other Forms only: marriage details captured at intake.",
    )
    set_para_text(
        find_para(
            doc,
            contains="Application statuses for Special Marriage (Intended Marriage) Notice Generation",
        ),
        "Application statuses for Special Marriage Notice Generation (Online / Offline; Intended Marriage and Other Forms):",
    )
    # T16 Details captured / Notice valid
    set_cell_text(
        doc.tables[16].rows[4].cells[1],
        "Bride / bridegroom particulars saved (Online: e-KYC / Face Authentication where "
        "Aadhaar available). Other Forms only: marriage details (date / place / ceremony) also saved",
    )
    set_cell_text(
        doc.tables[16].rows[16].cells[1],
        "Path-specific: Intended Marriage — notice ≥ 30 and ≤ 90 days with no valid "
        "objection; Other Forms — notice ≥ 30 days with no valid objection",
    )
    set_cell_text(
        doc.tables[16].rows[16].cells[3],
        "Proceed to Marriage Registration (see 7.3)",
    )

    # --- Combine §7.3 registration -------------------------------------------
    h73 = find_para(
        doc,
        exact="7.3 Special Marriage (Intended Marriage) Marriage Registration",
        heading_only=True,
    )
    set_para_text(
        h73,
        "7.3 Special Marriage Marriage Registration (Intended Marriage / Other Forms)",
    )
    insert_paragraph_after(
        h73,
        "Intended Marriage and Other Forms share a single registration workflow after "
        "notice publication. Path-specific branches: (1) timeline gate — Intended "
        "Marriage ≥ 30 and ≤ 90 days from publication; Other Forms ≥ 30 days (no ≤ 90 "
        "upper gate in the Other Forms registration diagram); (2) certificate schedule "
        "— Fourth Schedule after Chapter II solemnization vs Fifth Schedule after "
        "Chapter III registration.",
        style="Normal",
    )
    set_para_text(
        find_para(
            doc,
            contains="Marriage Registration for Intended Marriage continues from an approved",
        ),
        "Marriage Registration continues from an approved / published notice after the "
        "statutory waiting period. Completion is Offline (In Person) at the Sub-Registrar "
        "office (solemnization or Chapter III registration, DEO certificate production "
        "and SR DSC). There is no separate fully-online registration channel.",
    )
    set_row(
        doc.tables[17],
        1,
        [
            "Special Marriage Marriage Registration (Intended Marriage / Other Forms)",
            "Select approved notice (Intended: ≥30 & ≤90 days; Other Forms: ≥30 days), Second Payment, schedule visit, download certificate",
            "Objection enquiry (if any), SR verification, DEO joint photo / witness / declaration / certificate, solemnization or Chapter III registration, SR DSC",
            "Registration (In Person)",
        ],
    )
    set_para_text(
        find_para(
            doc,
            exact="Portal steps before the in-person office visit (after notice publication / 30-day period):",
        ),
        "Portal and office steps are the same for Intended Marriage and Other Forms "
        "except the timeline gate and the certificate schedule. The process table in "
        "7.3.2.2 is the authoritative step list; path-specific branches are in the Notes column.",
    )
    # Drop the numbered list that duplicated the registration step table.
    for exact in (
        "Citizen Login portal — continue from a published / approved notice.",
        "Select Notice — choose the approved notice for which registration / solemnization is sought.",
        "Validate timeline ≥ 30 days and ≤ 90 days — if outside the window, no action is allowed (Sec. 7 / Sec. 14).",
        "If any Objection? — branch to SR enquiry when an objection exists (Sec. 8–9).",
        "Second Payment — registration / solemnization fee.",
        "Schedule Visit — office visit for solemnization / registration.",
        "Office visit continues per 7.3.2.2 — SR verification, DEO joint photo, witness e-KYC / Face Authentication/manual capture, generate & sign/upload declaration, solemnization, certificate and SR DSC.",
    ):
        delete_para(find_para(doc, exact=exact))

    set_para_text(
        find_para(
            doc,
            exact="Flow (after notice publication / 30-day period — registration / solemnization steps):",
        ),
        "Flow (after notice publication — registration / solemnization steps; shared by Intended Marriage and Other Forms):",
    )
    # T18 combined step notes
    set_cell_text(
        doc.tables[18].rows[3].cells[1],
        "Validate timeline — Intended: ≥ 30 and ≤ 90 days (Sec. 7 / Sec. 14); Other Forms: ≥ 30 days (no ≤ 90 gate)",
    )
    set_cell_text(
        doc.tables[18].rows[3].cells[3],
        "If NO → No Action allowed (path-specific gate)",
    )
    set_cell_text(
        doc.tables[18].rows[17].cells[1],
        "Marriage solemnization (Intended, Sec. 12) / registration (Other Forms, Chapter III Sec. 15–16)",
    )
    set_cell_text(
        doc.tables[18].rows[17].cells[3],
        "Chapter II solemnization vs Chapter III registration",
    )
    set_cell_text(
        doc.tables[18].rows[18].cells[1],
        "Generate Marriage Certificate — Fourth Schedule (Intended) / Fifth Schedule (Other Forms)",
    )
    set_cell_text(
        doc.tables[18].rows[18].cells[3],
        "Fourth Schedule Sec. 13 / Fifth Schedule Sec. 16",
    )
    set_para_text(
        find_para(
            doc,
            contains="Key characteristics: notice selection with ≥30 and ≤90 day gate",
        ),
        "Key characteristics: notice selection with path-specific timeline gate "
        "(Intended ≥30 and ≤90; Other Forms ≥30); objection enquiry branch (valid "
        "objection removes notice); Second Payment; schedule visit; SR verification; "
        "DEO joint photo then witness capture (e-KYC / Face Authentication or manual), "
        "declaration generate/sign/upload, solemnization or Chapter III registration, "
        "certificate & signatures; SR DSC; certificate issuance (Fourth or Fifth Schedule).",
    )
    set_para_text(
        find_para(
            doc,
            exact="Application statuses for Special Marriage (Intended Marriage) Marriage Registration:",
        ),
        "Application statuses for Special Marriage Marriage Registration (Intended Marriage / Other Forms):",
    )
    set_cell_text(
        doc.tables[19].rows[1].cells[1],
        "Entry point — Intended: notice ≥ 30 and ≤ 90 days with no valid objection; Other Forms: notice ≥ 30 days with no valid objection",
    )
    set_cell_text(
        doc.tables[19].rows[3].cells[1],
        "Solemnization visit (Intended) or registration visit (Other Forms) booked",
    )
    set_cell_text(
        doc.tables[19].rows[6].cells[1],
        "Marriage solemnized (Chapter II) or conditions satisfied for registration (Chapter III)",
    )

    # Move Other Forms registration figure into §7.3.2.2 (before deleting 7.5).
    of_caption = find_para(
        doc,
        exact="Figure: Special Marriage Other Forms Marriage Registration",
    )
    of_image = previous_blip_para(of_caption)
    intended_caption = find_para(
        doc,
        exact="Figure: Special Marriage (Intended Marriage) Marriage Registration",
    )
    move_element_after(of_caption._p, intended_caption._p)
    move_element_after(of_image._p, intended_caption._p)

    # --- Delete former §7.4 and §7.5 (duplicate notice + registration) -------
    h74 = find_para(
        doc,
        exact="7.4 Special Marriage (Other Forms) Notice Generation",
        heading_only=True,
    )
    h8 = find_para(doc, exact="8. Functional requirements", heading_only=True)
    el = h74._p
    while el is not None and el is not h8._p:
        nxt = el.getnext()
        delete_element(el)
        el = nxt

    # --- §8 intro: process vs FR ---------------------------------------------
    set_para_text(
        find_para(
            doc,
            contains="Functional requirements are organized by service, aligned with 7 (To-Be).",
        ),
        "Functional requirements are organized by service, aligned with §7 (To-Be). "
        "Section 7 is the process authority — channel models, diagrams, step tables and "
        "status models. Section 8 states testable functional requirements (FR-HMA-* / "
        "FR-SMA-*) and does not restate those process steps. Where a requirement "
        "implements a To-Be step, the FR cites the §7 reference. Cross-cutting "
        "post-registration, notification and MIS requirements follow the service-specific blocks.",
    )

    # --- Merge party particulars (8.2.4 + 8.2.5) ------------------------------
    h_bg = find_para(doc, exact="8.2.4 Data capture — bridegroom", heading_only=True)
    set_para_text(h_bg, "8.2.4 Data capture — party particulars")
    insert_paragraph_after(
        h_bg,
        "Bridegroom and bride share the same Second Schedule party-particulars schema. "
        "The field catalogue below is common to both parties. Age and condition "
        "validations differ by party and path as noted. Witnesses are not captured at "
        "notice generation — three witnesses are captured at registration (FR-SMA-034).",
        style="Normal",
    )
    t46 = find_table_by_cell(doc, 1, 0, "FR-SMA-062")
    t47 = find_table_by_cell(doc, 1, 0, "FR-SMA-063")
    # Insert FR-SMA-063 immediately after FR-SMA-062.
    t46.rows[2]._tr.addprevious(deepcopy(t47.rows[1]._tr))
    # After inserting FR-SMA-063 at index 2: Condition=8, Age=10, Permanent Address=12.
    set_cell_text(
        t46.rows[8].cells[0],
        "Condition (unmarried / widow / widower / divorcee)",
    )
    set_cell_text(
        t46.rows[10].cells[2],
        "Cross-check DOB; imported through e-KYC / Face Authentication when available. "
        "Intended Marriage: bridegroom ≥ 21, bride ≥ 18 (Sec. 4(c)). Other Forms: both "
        "parties ≥ 21 (Sec. 15(d))",
    )
    set_cell_text(
        t46.rows[12].cells[2],
        "Required if different from dwelling place; else record same-as-current",
    )
    delete_heading_block(
        find_para(doc, exact="8.2.5 Data capture — bride", heading_only=True),
        stop_at_level=4,
    )

    # --- Drop duplicate FR-SMA-034 from solemnization table -------------------
    t53 = find_table_by_cell(doc, 1, 0, "FR-SMA-033")
    if t53.rows[2].cells[0].text.strip() != "FR-SMA-034":
        raise KeyError(f"Expected FR-SMA-034 at T53 R2, got {t53.rows[2].cells[0].text!r}")
    remove_table_row(t53, 2)
    set_para_text(
        find_para(doc, exact="(Ref: 7.7 steps 5–16; SMA 1954 Sec. 11–13, Third and Fourth Schedules)"),
        "(Ref: 7.3.2.2; SMA 1954 Sec. 11–13, Third and Fourth Schedules). Witness capture "
        "is specified once in Data capture — witnesses (FR-SMA-034, FR-SMA-064–065) and is "
        "not repeated here.",
    )

    # --- Drop placeholder FR-only cross-reference sub-sections ----------------
    # Delete from the end so remaining heading texts stay unique.
    for heading in reversed(PLACEHOLDER_H4):
        delete_heading_block(
            find_para(doc, exact=heading, heading_only=True),
            stop_at_level=4,
        )

    # Fold FR-SMA-051 (only lived in deleted 8.3.4) into the fees intro.
    fees_intro = find_para(
        doc,
        contains="(Ref: 7.5–7.10 first and second payment",
    )
    set_para_text(
        fees_intro,
        "(Ref: 7.2 / 7.3 first and second payment; Special Marriage (Karnataka) Rules, 1961). "
        "Second payment is collected at registration (FR-SMA-033 / FR-SMA-050). Additional "
        "fee applies for solemnization outside the Marriage Officer’s office (FR-SMA-051).",
    )

    # --- Renumber remaining §8.2 / §8.3 Heading 4s ----------------------------
    remaining_82 = h4_between(
        doc,
        "8.2 Special Marriage (Intended Marriage/Other Forms) Notice Generation",
        "8.3",
    )
    if len(remaining_82) != len(NEW_82_H4):
        raise KeyError(
            f"§8.2 Heading 4 count {len(remaining_82)} != {len(NEW_82_H4)}: "
            + "; ".join(p.text.strip() for p in remaining_82)
        )
    for p, title in zip(remaining_82, NEW_82_H4):
        set_para_text(p, title)

    remaining_83 = h4_between(
        doc,
        "8.3 Special Marriage (Intended Marriage/Other Forms) Marriage Registration",
        "8.4",
    )
    if len(remaining_83) != len(NEW_83_H4):
        raise KeyError(
            f"§8.3 Heading 4 count {len(remaining_83)} != {len(NEW_83_H4)}: "
            + "; ".join(p.text.strip() for p in remaining_83)
        )
    for p, title in zip(remaining_83, NEW_83_H4):
        set_para_text(p, title)

    # --- Scope / in-scope bullets (before generic replacements) ---------------
    set_para_text(
        find_para(
            doc,
            contains="Hindu Marriage processing channels per approved process diagrams (7.3–7.4)",
        ),
        "Hindu Marriage processing channels per approved process diagrams (7.1.2.2 / 7.1.2.3): "
        "Hindu Marriage Online and Hindu Marriage Offline.",
    )
    set_para_text(
        find_para(
            doc,
            contains="Special Marriage process diagrams (7.5–7.10)",
        ),
        "Special Marriage process diagrams (7.2–7.3): shared Notice Generation Online/Offline "
        "and shared Marriage Registration for Intended Marriage and Other Forms, with "
        "conditional branches for Other Forms marriage details, timeline gate and certificate schedule.",
    )

    # --- Cross-references (stale 7.4–7.10 / 8.19–8.25 numbering) --------------
    for old, new in TEXT_REPLACEMENTS:
        replace_everywhere(doc, old, new)

    # RTM section column still uses old 8.19-style IDs in some cells.
    rtm = find_table_by_cell(doc, 0, 1, "Act/Rule/Form")
    rtm_map = {
        "8.19": "8.2.1",
        "8.20": "8.2.2",
        "8.21": "8.2.7",
        "8.23": "8.2.8",
        "8.25": "8.3.4",
    }
    for row in rtm.rows[1:]:
        cell = row.cells[3]
        key = cell.text.strip()
        if key in rtm_map:
            set_cell_text(cell, rtm_map[key])

    # --- Empty leftover Heading 5s under Hindu status -------------------------
    for p in list(doc.paragraphs):
        if heading_level(p) == 5 and not p.text.strip() and not para_has_blip(p):
            delete_para(p)

    # --- Appendix: drop duplicate notice-diagram bullets ----------------------
    dup_captions = [
        "Approved process diagram — Special Marriage (Intended Marriage/Other Forms) Notice Online "
        "(shared with Intended) — Process Diagrams/Special Marriage/Special Marriage "
        "(Intended Marriage Notice) 2-SpecialMarriageNoticeOnline.drawio (1).png",
        "Approved process diagram — Special Marriage (Intended Marriage/Other Forms) Notice Offline "
        "(shared with Intended) — Process Diagrams/Special Marriage/Special Marriage "
        "(Intended Marriage Notice) 2-SpecialMarriageNoticeOffline.drawio (1).png",
    ]
    for cap in dup_captions:
        delete_para(find_para(doc, exact=cap))

    # --- TOC last (body headings already renamed) -----------------------------
    update_toc(doc)

    doc.save(str(DST))
    print(f"Wrote {DST}")

    # Verification dump of remaining headings
    doc2 = Document(str(DST))
    print("--- headings ---")
    for p in doc2.paragraphs:
        if style_name(p).startswith("Heading") and p.text.strip():
            print(f"  [{style_name(p)}] {p.text.strip()}")


if __name__ == "__main__":
    main()
