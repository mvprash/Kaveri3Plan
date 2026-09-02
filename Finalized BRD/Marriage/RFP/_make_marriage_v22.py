# -*- coding: utf-8 -*-
"""Create BRD_Marriage_v1.12.docx from v1.11.

Align §7.2.2.1–7.2.2.2 Special Marriage Notice Online steps, status model and
FR-SMA-009/010 with the approved v1.11 Online notice process diagram (mandatory
e-KYC / Face Authentication after Prerequisite; no Aadhaar YES/NO branch on
Online). Offline notice diagram and steps retain the Aadhaar YES/NO branch.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout.reconfigure(encoding="utf-8")

BASE = Path(r"E:\MVP\Kaveri 3.0\Source Code\Kaveri 3 Plan\Finalized BRD\Marriage\RFP")
SRC = BASE / "BRD_Marriage_v1.11.docx"
DST = BASE / "BRD_Marriage_v1.12.docx"

ONLINE_NOTICE_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    [
        "8",
        "e-KYC / Face Authentication on Bride & Bridegroom details",
        "System / Citizen",
        "Continues from 7.2.2.1 step 7; mandatory on Online notice diagram (no manual path)",
    ],
    [
        "9",
        "Review summary and proceed document uploading",
        "System / Citizen",
        "Summary of captured particulars",
    ],
    [
        "10",
        "Upload Identity Proof, Photo, Age Proof, Address Proof (Bridegroom & Bride)",
        "Citizen",
        "Mandatory supporting documents including individual photographs",
    ],
    [
        "11",
        "SR Verification (decision)",
        "Sub Registrar",
        "Approve or Reject",
    ],
    [
        "11a",
        "Reject → return to Prerequisite & declaration",
        "Sub Registrar → Citizen",
        "Citizen corrects and resubmits (diagram returns to Prerequisite)",
    ],
    [
        "12",
        "First Payment",
        "System / Citizen",
        "Notice fee per Karnataka Special Marriage fee schedule",
    ],
    [
        "13",
        "Notice Generated",
        "System",
        "Statutory notice generated after first payment (no Online DEO / appointment)",
    ],
    [
        "14",
        "Proceed with e-sign",
        "Citizen",
        "Citizen eSign on the generated notice before portal publication",
    ],
    [
        "15",
        "Marriage notice displayed in portal",
        "System",
        "Sec. 6 publication (digital); follows e-sign per Online notice diagram",
    ],
    [
        "16",
        "30-day countdown starts",
        "System",
        "Objection period per Sec. 7 / Sec. 16 from publication / portal display",
    ],
]

OFFLINE_NOTICE_STEPS = [
    ["#", "Step", "Lane", "Notes"],
    [
        "8",
        "If Aadhaar available → e-KYC / Face Authentication on Bride & Bridegroom details; "
        "else Enter Bride & Bridegroom details",
        "System / Citizen",
        "Continues from 7.2.2.1 step 7; per Offline notice diagram (Aadhaar YES/NO branch)",
    ],
    [
        "9",
        "Review summary and proceed document uploading",
        "System / Citizen",
        "Summary of captured particulars",
    ],
    [
        "10",
        "Upload Identity / Age / Address proofs (Bridegroom & Bride)",
        "Citizen",
        "Mandatory documents (individual photos captured later by DEO at office)",
    ],
    [
        "11",
        "SR Verification (decision)",
        "Sub Registrar",
        "Approve or Reject",
    ],
    [
        "11a",
        "Reject → return to Prerequisite & declaration",
        "Sub Registrar → Citizen",
        "Citizen corrects and resubmits (diagram returns to Prerequisite)",
    ],
    ["12", "First Payment", "System / Citizen", "Notice fee"],
    ["13", "Schedule appointment with SR", "System / Citizen", "After first payment"],
    ["14", "SR Generates Notice", "Sub Registrar", "Statutory notice generation"],
    ["15", "Selects DEO", "Sub Registrar", "Assign FDA/SDA/DEO"],
    [
        "16",
        "Capture individual photos of Bride & Bridegroom",
        "FDA / SDA / DEO",
        "Office visit activity",
    ],
    [
        "17",
        "Download, Print, Sign, Scan and Upload notice",
        "FDA / SDA / DEO",
        "Physical notice processed into system",
    ],
    [
        "18",
        "Paste form on respective Notice Board",
        "FDA / SDA / DEO",
        "Sec. 6(2) conspicuous place publication; also drives portal display",
    ],
    ["19", "30-day countdown starts", "System", "Objection period per Sec. 7 / Sec. 16"],
]


def set_para_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def set_cell_text(cell: _Cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    set_para_text(paras[0], text)
    for p in paras[1:]:
        set_para_text(p, "")


def set_row(table: Table, ri: int, values: list[str]) -> None:
    row = table.rows[ri]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            set_cell_text(row.cells[ci], val)


def clear_extra_rows(table: Table, keep_rows: int) -> None:
    tbl = table._tbl
    while len(table.rows) > keep_rows:
        tbl.remove(table.rows[-1]._tr)


def ensure_rows(table: Table, needed: int) -> None:
    while len(table.rows) < needed:
        table._tbl.append(deepcopy(table.rows[-1]._tr))


def fill_table(table: Table, rows: list[list[str]]) -> None:
    ensure_rows(table, len(rows))
    for ri, vals in enumerate(rows):
        set_row(table, ri, vals)
    clear_extra_rows(table, len(rows))


def add_version_row(table: Table, values: list[str]) -> None:
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    set_row(table, len(table.rows) - 1, values)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style else ""


def find_para(
    doc: Document,
    exact: str | None = None,
    contains: str | None = None,
    heading_only: bool = False,
) -> Paragraph:
    for p in doc.paragraphs:
        if heading_only and not style_name(p).startswith("Heading"):
            continue
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            return p
    raise KeyError(f"Paragraph not found: exact={exact!r} contains={contains!r}")


def find_table_containing(doc: Document, exact: str) -> Table:
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == exact:
                return table
    raise KeyError(f"Table not found containing {exact!r}")


def find_fr_row(table: Table, req_id: str):
    for row in table.rows:
        if row.cells[0].text.strip() == req_id:
            return row
    raise KeyError(f"FR row not found: {req_id}")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # Cover + version history
    set_cell_text(doc.tables[0].rows[2].cells[1], "1.12")
    set_cell_text(doc.tables[0].rows[11].cells[1], "2026-09-01")
    add_version_row(
        doc.tables[1],
        [
            "1.12",
            "2026-09-01",
            "Nandha Kumar",
            "Align Special Marriage Notice Online steps (§7.2.2.1–7.2.2.2), status model "
            "and FR-SMA-009/010 with v1.11 Online notice diagram — mandatory e-KYC after "
            "Prerequisite; Aadhaar YES/NO branch retained for Offline only",
            "Prashanth",
        ],
    )

    # §7.2 overview
    set_para_text(
        find_para(
            doc,
            contains="Intended Marriage (Chapter II) and Other Forms (Chapter III) share a single notice-generation",
        ),
        "Intended Marriage (Chapter II) and Other Forms (Chapter III) share a single "
        "notice-generation workflow. After Marriage Registration the citizen answers "
        "Whether Marriage already taken place or not?: No routes to Special Marriage "
        "(Intended Marriage Notice); Yes routes to Special Marriage (Other Forms Notice) "
        "and then Enter Marriage Details (date, place and form / rites of the already "
        "celebrated ceremony) before Prerequisite. Both paths then share Prerequisite; "
        "party capture (Online: mandatory e-KYC / Face Authentication; Offline: Aadhaar "
        "YES/NO branch or manual); Online or Offline publication; first payment; "
        "eSign / notice-board and the 30-day countdown.",
    )

    # §7.2.2.1 common intake intro
    set_para_text(
        find_para(
            doc,
            contains="Identical in both notice diagrams (Citizens and System lanes). Intended Marriage and Other Forms share these steps",
        ),
        "Identical in both notice diagrams (Citizens and System lanes). Intended Marriage "
        "and Other Forms share these steps through Prerequisite (step 7); the path "
        "decision and Other Forms Enter Marriage Details branch are in steps 5–6. Party "
        "capture continues per the Online (7.2.2.2) or Offline (7.2.2.3) notice diagram:",
    )

    # Remove common-intake step 8 (Aadhaar branch — Offline-only per diagrams)
    delete_paragraph(
        find_para(
            doc,
            contains="If Aadhaar information available: e-KYC / Face Authentication on Bride and Bridegroom details; else Enter Bride and Bridegroom details",
        )
    )

    # §7.2.2.2 Online flow intro + key characteristics
    set_para_text(
        find_para(
            doc,
            contains="Flow (continuing from 7.2.2.1 step 8 — Online notice-specific steps; shared by Intended Marriage and Other Forms):",
        ),
        "Flow (continuing from 7.2.2.1 step 7 — Online notice-specific steps; shared by "
        "Intended Marriage and Other Forms):",
    )
    set_para_text(
        find_para(
            doc,
            contains="Key characteristics (both paths): after the Whether Marriage already taken place or not? branch (Other Forms: Enter Marriage Details before Prerequisite), e-KYC / Face Authentication on Bride & Bridegroom when Aadhaar available",
        ),
        "Key characteristics (both paths): after the Whether Marriage already taken place "
        "or not? branch (Other Forms: Enter Marriage Details before Prerequisite), "
        "mandatory e-KYC / Face Authentication on Bride & Bridegroom; document upload "
        "including individual photos; SR verification before First Payment; System "
        "generates notice after payment; citizen e-sign on the generated notice; then "
        "Marriage notice displayed in portal; no Online appointment or DEO; 30-day "
        "countdown from publication (Sec. 7 / Sec. 16).",
    )

    # §7.2.2.3 Offline flow intro (renumber reference only)
    set_para_text(
        find_para(
            doc,
            contains="Flow (continuing from 7.2.2.1 step 8 — Offline notice-specific steps; shared by Intended Marriage and Other Forms):",
        ),
        "Flow (continuing from 7.2.2.1 step 7 — Offline notice-specific steps; shared by "
        "Intended Marriage and Other Forms):",
    )

    # Step tables
    fill_table(doc.tables[14], ONLINE_NOTICE_STEPS)
    fill_table(doc.tables[15], OFFLINE_NOTICE_STEPS)

    # Channel model — Online row already omits manual path; clarify Offline
    set_cell_text(
        doc.tables[13].rows[2].cells[1],
        "Capture details (e-KYC / Face Authentication when Aadhaar available, else manual), "
        "document upload, First Payment, appointment",
    )

    # Status model — Details captured
    set_cell_text(
        doc.tables[16].rows[4].cells[1],
        "Online: bride / bridegroom particulars saved via mandatory e-KYC / Face "
        "Authentication. Offline: e-KYC / Face Authentication where Aadhaar available, "
        "else manual entry. Other Forms only: marriage details already saved at path "
        "selection (before Prerequisite)",
    )

    # §8.2.3 party particulars intro
    set_para_text(
        find_para(
            doc,
            contains="Bridegroom and bride share the same Second Schedule party-particulars schema",
        ),
        "Bridegroom and bride share the same Second Schedule party-particulars schema. "
        "The field catalogue below is common to both parties. Age and condition "
        "validations differ by party and path as noted. Online notice channel: mandatory "
        "e-KYC / Face Authentication on bride and bridegroom (FR-SMA-009; 7.2.2.2). "
        "Offline notice channel: e-KYC when Aadhaar is available, otherwise manual capture "
        "(FR-SMA-009 / FR-SMA-010; 7.2.2.3). Witnesses are not captured at notice "
        "generation — three witnesses are captured at registration (FR-SMA-034).",
    )

    # FR-SMA-009 / FR-SMA-010 — scope Online vs Offline
    fr009 = find_fr_row(find_table_containing(doc, "FR-SMA-009"), "FR-SMA-009")
    set_cell_text(
        fr009.cells[1],
        "Online channel: system shall perform e-KYC / Face Authentication on bride and "
        "bridegroom details (mandatory; per Online notice diagram 7.2.2.2). Offline "
        "channel: e-KYC / Face Authentication on bride and bridegroom where Aadhaar "
        "information is available (per Offline notice diagram 7.2.2.3)",
    )
    set_cell_text(
        fr009.cells[3],
        "Online: no manual bride/bridegroom capture path at notice generation; Offline: "
        "e-KYC when Aadhaar available; failure fallback per RS-MRG-002",
    )

    fr010 = find_fr_row(find_table_containing(doc, "FR-SMA-010"), "FR-SMA-010")
    set_cell_text(
        fr010.cells[1],
        "Offline channel only: where Aadhaar information is unavailable, system shall "
        "allow manual capture of bride and bridegroom details with mandatory documentary "
        "proof (per Offline notice diagram 7.2.2.3). Not applicable to Online notice "
        "channel",
    )
    set_cell_text(
        fr010.cells[3],
        "Manual path flagged for SR scrutiny; Online notice channel has no manual "
        "bride/bridegroom capture alternative",
    )

    fr066 = find_fr_row(find_table_containing(doc, "FR-SMA-066"), "FR-SMA-066")
    set_cell_text(
        fr066.cells[3],
        "OTP / SMS template reviewed by department; reason text visible before code entry; "
        "applies where e-KYC / Face Authentication is performed in Special Marriage notice "
        "generation (Online mandatory; Offline when Aadhaar available)",
    )

    doc.save(str(DST))
    print(f"Wrote {DST}")

    # Verification
    doc2 = Document(str(DST))
    print("Version:", doc2.tables[0].rows[2].cells[1].text.strip())
    print("Online step 8:", doc2.tables[14].rows[1].cells[1].text[:60])
    print("Offline step 8:", doc2.tables[15].rows[1].cells[1].text[:60])
    leftover = sum(
        1
        for p in doc2.paragraphs
        if "If Aadhaar information available: e-KYC / Face Authentication on Bride and Bridegroom details; else Enter Bride"
        in p.text
    )
    print(f"Common intake Aadhaar list item left: {leftover}")
    fr009_v = find_fr_row(find_table_containing(doc2, "FR-SMA-009"), "FR-SMA-009")
    print("FR-SMA-009 Online mandatory:", "mandatory" in fr009_v.cells[1].text)


if __name__ == "__main__":
    main()
